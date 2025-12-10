from telegram import (
    Update, InlineKeyboardButton, InlineKeyboardMarkup,
    ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove, Bot, InputFile
)
from telegram.ext import (
    ApplicationBuilder, CommandHandler, CallbackQueryHandler,
    MessageHandler, ContextTypes, ConversationHandler, filters, Application, CallbackContext
)
from flask import Flask, request
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor
from telegram.helpers import escape_markdown
from telegram.error import BadRequest
import telegram
import stripe
from stripe.error import StripeError
import threading
import asyncio
import os
import uuid
import requests
import re
import csv
import io
import time
import random  # Adicione esta linha no topo do arquivo com os outros imports
from decimal import Decimal, ROUND_HALF_UP
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.lib.units import mm
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import pickle
import logging
from typing import Dict, Any
import sys



    
   
    
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
STRIPE_PUBLIC_KEY = os.getenv("STRIPE_PUBLIC_KEY")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET")
ADMIN_CHAT_ID = os.getenv("ADMIN_CHAT_ID")


# ==================== VERIFICAÇÃO DE SEGURANÇA ====================

# 🔐 VERIFICAR SE AS CHAVES NÃO ESTÃO HARCODED
def verificar_seguranca():
    """Verificar se as chaves não estão hardcoded no código"""
    print("\n" + "="*60)
    print("🔍 VERIFICAÇÃO DE SEGURANÇA")
    print("="*60)
    
    # Lista de chaves que NÃO devem aparecer no código
    # ✅ CORREÇÃO: Apenas pedaços das chaves para não se auto-detectar
    chaves_perigosas = [
        
    ]
    
    # ✅ SUBSTITUA por isso:
    pedacos_chaves_perigosas = [
        "8416340654",
        "sk_test_51",
        "pk_test_51", 
        "whsec_"
    ]
    
    # Ler o próprio arquivo para verificar
    try:
        with open(__file__, 'r', encoding='utf-8') as f:
            codigo = f.read()
            
        for chave in pedacos_chaves_perigosas:
            # Contar quantas vezes aparece (exceto na própria função)
            ocorrencias = codigo.count(chave)
            
            # Se aparecer mais de 2 vezes (provavelmente na função + chamada)
            if ocorrencias > 2:
                print(f"⚠️ Possível chave hardcoded encontrada: {chave}...")
                return False
                
    except:
        pass
    
    print("✅ Código seguro: sem chaves hardcoded")
    return True


    
# ==================== CONFIGURAR STRIPE ====================

if STRIPE_SECRET_KEY:
    import stripe
    stripe.api_key = STRIPE_SECRET_KEY
    print("✅ Stripe configurado")
else:
    print("⚠️ Stripe não configurado - funcionalidade de pagamentos limitada")

# ==================== SEU CÓDIGO CONTINUA AQUI ====================

# ... resto do seu código (handlers, funções, etc.)

print("\n" + "="*60)
print("🤖 BOT CONFIGURADO COM SUCESSO")
print("="*60)



application = Application.builder().token(TELEGRAM_TOKEN).pool_timeout(30).build()

stripe.api_key = STRIPE_SECRET_KEY
bot = Bot(token=TELEGRAM_TOKEN)

MEU_CHAT_ID = ""  # ⚠️ SUBSTITUA pelo SEU chat ID real! 
# Estados da conversa
# Estados da conversa - ATUALIZAR COM TODOS OS ESTADOS
NOME, EMAIL, PAIS, CONTACTO, TIPO, ESTILO, PROFISSAO, OBJETOS, SUPER_HEROI, ELEMENTOS_FAMILY, ADULTOS_FAMILY, CRIANCAS_FAMILY, ANIMAIS_FAMILY, TAMANHO, FOTO, NOME_ANIMAL, TIPO_ANIMAL, TIPO_PERSONALIZADO, NOME_PECA, NOME_CARTOON, FRASE_CARTOON, NOME_PERSONALIZADO, FRASE_PERSONALIZADO, NOME_FAMILY, FRASE_FAMILY, AGUARDANDO_ID_PEDIDO, AGUARDANDO_PROBLEMA, GIFT_NOME, GIFT_EMAIL, GIFT_PAIS, GIFT_CONTACTO, GIFT_FOTO, GIFT_FIM, GIFT_NOME_BOX, GIFT_FRASE_BOX  = range(35)

# ======================= SISTEMA DE ESTATÍSTICAS =======================
PEDIDOS_REGISTO = {}
TIMERS_ATIVOS = {}

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

USER_SESSIONS = {}  # Dicionário para gerir sessões de usuários


        
PRECOS = {
    "cartoon_individual": 00.0,
    "cartoon_family": 130.0,
    "cartoon_animal": 40.0,  # ✅ CORRIGIDO: era 0.0
    "cartoon_custom": 90.0
}

PRECOS_ESTILO = {
    "Full Body": 80.0,
    "Bust": 50.0,
    "Voxel": 70.0,
    "Office": 90.0,
    "Superheroes": 70.0
}

PRECOS_TAMANHO = {
    " 6\" | 15.2cm ": 0.0,
    " 7\" | 17.8cm ": 0.0,
    " 8\" | 20.3cm ": 10.0,
    " 9\" | 22.9cm ": 15.0
}


# Tamanhos específicos para CADA estilo
TAMANHOS_POR_ESTILO = {
    "Full Body": {
        "6": {"nome": " 6\" | 15.2cm ", "preco": 0.0},
        "8": {"nome": " 8\" | 20.3cm ", "preco": 40.0},
        "10": {"nome": " 10\" | 25.4cm ", "preco": 70.0},
        "12": {"nome": " 12\" | 30.5cm ", "preco": 100.0}
    },
      "Bust": {
        "3.5": {"nome": " 3.5\" | 8.9cm ", "preco": 0.0}
    },
    "Voxel": {
        "6": {"nome": " 6\" | 15.2cm ", "preco": 0.0},
        "8": {"nome": " 8\" | 20.3cm ", "preco": 40.0},
        "10": {"nome": " 10\" | 25.4cm ", "preco": 70.0},
        "12": {"nome": " 12\" | 30.5cm ", "preco": 100.0}
    },
    "Office": {
        "6": {"nome": " 6\" | 15.2cm ", "preco": 0.0},
        "8": {"nome": " 8\" | 20.3cm ", "preco": 40.0},
        "10": {"nome": " 10\" | 25.4cm ", "preco": 70.0},
        "12": {"nome": " 12\" | 30.5cm ", "preco": 100.0}
    },
    "Superheroes": {
        "6": {"nome": " 6\" | 15.2cm ", "preco": 0.0},
        "8": {"nome": " 8\" | 20.3cm ", "preco": 40.0},
        "10": {"nome": " 10\" | 25.4cm ", "preco": 70.0},
        "12": {"nome": " 12\" | 30.5cm ", "preco": 100.0}
    }
}

PRECOS_TAMANHO_REDUZIDO = {
    "4.5": 0.0,
    "3.5": 5.0,
    "PORTA-CHAVES": 15.0,
    
}

# --- Impostos e frete por país ---
TAXAS_PAISES = {
    "portugal": {"imposto": 0.23, "frete": 10.0},
    "espanha": {"imposto": 0.23, "frete": 10.0},
    "franca": {"imposto": 0.23, "frete": 10.0},
    "alemanha": {"imposto": 0.23, "frete": 10.0},
    "belgica": {"imposto": 0.23, "frete": 10.0},
    "reino_unido": {"imposto": 0.00, "frete": 30.0},
    "estados_unidos": {"imposto": 0.00, "frete": 50.0},
    "canada": {"imposto": 0.00, "frete": 50.0},
    "paises_baixos": {"imposto": 0.23, "frete": 10.0},
    "brasil": {"imposto": 0.00, "frete": 70.0},
    "irlanda": {"imposto": 0.23, "frete": 10.0},
    "italia": {"imposto": 0.23, "frete": 10.0},
    "luxemburgo": {"imposto": 0.23, "frete": 10.0},

}



# Tamanhos específicos para Porta-Chaves
TAMANHOS_GIFT = {
    "padrao": {"nome": "🔑 Porta-Chaves", "preco": 20.0}
}



    #"portugal": {"imposto": 0.23, "frete": 10.0},
    #"espanha": {"imposto": 0.21, "frete": 10.0},
    #"franca": {"imposto": 0.20, "frete": 10.0},
    #"alemanha": {"imposto": 0.19, "frete": 10.0},
    #"belgica": {"imposto": 0.21, "frete": 10.0},
    #"reino_unido": {"imposto": 0.00, "frete": 30.0},
    #"estados_unidos": {"imposto": 0.00, "frete": 50.0},
    #"paises_baixos": {"imposto": 0.21, "frete": 10.0},
    #"brasil": {"imposto": 0.00, "frete": 70.0},
    #"irlanda": {"imposto": 0.23, "frete": 10.0},
    #"italia": {"imposto": 0.22, "frete": 10.0},
    #"luxemburgo": {"imposto": 0.17, "frete": 10.0},



 #-------------------------------------------------------------------

   # "republica_checa": {"imposto": 0.21, "frete": 7.5},
   # "suica": {"imposto": 0.00, "frete": 8.5},
   # "finlandia": {"imposto": 0.24, "frete": 9.5},
   #  "grecia": {"imposto": 0.24, "frete": 8.0},
   # "austria": {"imposto": 0.20, "frete": 6.5},
   # "bulgaria": {"imposto": 0.20, "frete": 8.0},
   # "chipre": {"imposto": 0.19, "frete": 9.0},
   # "croacia": {"imposto": 0.25, "frete": 6.5},
   # "eslovaquia": {"imposto": 0.20, "frete": 7.5},
   # "eslovenia": {"imposto": 0.22, "frete": 7.0},
   # "estonia": {"imposto": 0.22, "frete": 8.5},
   # "hungria": {"imposto": 0.27, "frete": 7.5},
   # "letonia": {"imposto": 0.21, "frete": 8.0},
   # "lituania": {"imposto": 0.21, "frete": 8.5},
   # "malta": {"imposto": 0.18, "frete": 9.0},
   # "polonia": {"imposto": 0.23, "frete": 7.0},
   # "romenia": {"imposto": 0.19, "frete": 8.0},

   # "angola": {"imposto": 0.15, "frete": 15.0},
   # "mocambique": {"imposto": 0.14, "frete": 15.0},
   # "cabo_verde": {"imposto": 0.10, "frete": 10.0},
   # "guine_bissau": {"imposto": 0.08, "frete": 10.0},
   # "sao_tome": {"imposto": 0.07, "frete": 9.0},
   # "timor_leste": {"imposto": 0.05, "frete": 10.0}


# Dicionário de países e prefixos
PAISES_PREFIXOS = {
  "portugal": "+351",
  "espanha": "+34",
  "franca": "+33",
  "alemanha": "+49",
  "belgica": "+32",
  "reino_unido": "+44",
  "estados_unidos": "+1",
  "canada": "+1",
  "paises_baixos": "+31",
  "brasil": "+55",
  "irlanda": "+353",
  "italia": "+39",
  "luxemburgo": "+352"
}

PRECOS_FAMILY = {
    "adulto": 60.0,      # Preço por adulto
    "crianca": 50.0,     # Preço por criança
    "animal": 40.0       # Preço por animal
}

TAMANHOS_FAMILY = {
    "6": {"nome": " 6\" | 15.2cm ", "preco": 0.0},
    "7.5": {"nome": " 7.5\" | 19.1cm ", "preco": 60.0}
}
# ======================= PREÇOS PARA ANIMAL =======================

TAMANHOS_ANIMAL = {
    "2.5": {"nome": " 2.5\" | 6.4cm ", "preco": 0.0},
    "3.5": {"nome": " 3.5\" | 8.9cm ", "preco": 10.0}
}

PEDIDOS_RECUSADOS = {}



executor = ThreadPoolExecutor()
# ======================= FUNÇÕES AUXILIARES =================
TEMPORIZADORES_ATIVOS = {}

EMAIL_CHAT_MAP = {}
# =======================  FLASK APP =================


CONTADOR_UTILIZADORES = {
    "data": datetime.now().date(),
    "contador": 0,
    "utilizadores_unicos": set()  # Para evitar duplicados no mesmo dia
}



app = Flask(__name__)

def run_flask():
    app.run(host="0.0.0.0", port=5000)


# ========================================





class UserSession:
    """Classe para gerir sessão de cada usuário de forma isolada"""
    def __init__(self, user_id: int):
        self.user_id = user_id
        self.data: Dict[str, Any] = {}
        self.timers: Dict[str, asyncio.Task] = {}
        self.last_activity = time.time()
    
    def update_state(self, key: str, value: Any):
        """Atualizar estado do usuário"""
        self.data[key] = value
        self.last_activity = time.time()
    
    def get_state(self, key: str, default=None):
        """Obter estado do usuário"""
        return self.data.get(key, default)
    
    def clear_state(self):
        """Limpar estado do usuário"""
        self.data.clear()
        # Cancelar todos os temporizadores
        for timer_id, timer in self.timers.items():
            if timer and not timer.done():
                timer.cancel()
        self.timers.clear()

# Dicionário global para armazenar sessões
USER_SESSIONS: Dict[int, UserSession] = {}

def get_user_session(user_id: int) -> UserSession:
    """Obter ou criar sessão do usuário"""
    if user_id not in USER_SESSIONS:
        USER_SESSIONS[user_id] = UserSession(user_id)
        print(f"📱 Nova sessão criada para user_id: {user_id}")
    return USER_SESSIONS[user_id]




async def cleanup_inactive_sessions():
    """Limpar sessões inativas periodicamente"""
    while True:
        await asyncio.sleep(3600)  # Verificar a cada hora
        current_time = asyncio.get_event_loop().time()
        inactive_users = []
        
        for user_id, session in USER_SESSIONS.items():
            # Se inativo por mais de 24 horas
            if current_time - session.last_activity > 86400:
                inactive_users.append(user_id)
        
        for user_id in inactive_users:
            session = USER_SESSIONS.pop(user_id, None)
            if session:
                await session.clear_state()
                logger.info(f"🧹 Sessão limpa para user_id: {user_id} (inativo)")




async def iniciar_temporizador_seguro(user_id: int, pedido_id: str, minutos: int, callback_func):
    """Iniciar temporizador de forma segura para múltiplos clientes"""
    session = get_user_session(user_id)
    
    async def timer_task():
        try:
            await asyncio.sleep(minutos * 60)
            await callback_func(pedido_id, user_id)
        except asyncio.CancelledError:
            print(f"⏰ Temporizador cancelado para pedido {pedido_id}")
        except Exception as e:
            print(f"❌ Erro no temporizador {pedido_id}: {e}")
    
    # Cancelar temporizador anterior se existir
    if pedido_id in session.timers:
        old_timer = session.timers.pop(pedido_id)
        if old_timer and not old_timer.done():
            old_timer.cancel()
    
    # Criar novo temporizador
    timer = asyncio.create_task(timer_task())
    session.timers[pedido_id] = timer
    print(f"⏰ Temporizador {minutos}min iniciado para pedido {pedido_id}, user {user_id}")

async def cancelar_temporizador_seguro(user_id: int, pedido_id: str):
    """Cancelar temporizador de forma segura"""
    session = USER_SESSIONS.get(user_id)
    if session and pedido_id in session.timers:
        timer = session.timers.pop(pedido_id)
        if timer and not timer.done():
            timer.cancel()
            print(f"⏰ Temporizador cancelado para pedido {pedido_id}")



























# 🔥 FILTROS PERSONALIZADOS PARA O GIFT
def is_gift_nome(update, context):
    return context.user_data.get('conversation_state') == GIFT_NOME

def is_gift_email(update, context):
    return context.user_data.get('conversation_state') == GIFT_EMAIL

def is_gift_contacto(update, context):
    return context.user_data.get('conversation_state') == GIFT_CONTACTO

def is_gift_foto(update, context):
    return context.user_data.get('conversation_state') == GIFT_FOTO

ESTATISTICAS = {
    "total_pedidos": 0,
    "pedidos_pagos": 0,
    "pedidos_expirados": 0,
    "problemas_reportados": 0,
    "tentativas_recuperacao": 0,
    "em_recuperacao": 0,
    "ofertas_aceites": 0,
    "ofertas_recusadas": 0
}

def atualizar_estatistica(tipo):
    """Atualiza as estatísticas globais"""
    try:
        # 🔥 ATUALIZA A ESTATÍSTICA ESPECÍFICA
        if tipo in ESTATISTICAS:
            ESTATISTICAS[tipo] += 1
        else:
            # Se for um tipo novo, cria com valor 1
            ESTATISTICAS[tipo] = 1
        
        # 🔥 MOSTRA RELATÓRIO
        print("\n" + "=" * 80)
        print("📊 ESTATÍSTICAS ATUALIZADAS")
        print("=" * 80)
        print(f"📦 Total de Pedidos: {ESTATISTICAS['total_pedidos']}")
        print(f"✅ Pedidos Pagos: {ESTATISTICAS['pedidos_pagos']}")
        print(f"❌ Pedidos Expirados: {ESTATISTICAS['pedidos_expirados']}")
        print(f"🚨 Problemas Reportados: {ESTATISTICAS['problemas_reportados']}")
        print(f"🔄 Tentativas de Recuperação: {ESTATISTICAS['tentativas_recuperacao']}")
        print(f"⏳ Em Recuperação: {ESTATISTICAS['em_recuperacao']}")
        print(f"🎉 Ofertas Aceites: {ESTATISTICAS['ofertas_aceites']}")
        print(f"😔 Ofertas Recusadas: {ESTATISTICAS['ofertas_recusadas']}")
        
        if ESTATISTICAS['total_pedidos'] > 0:
            taxa_conversao = (ESTATISTICAS['pedidos_pagos'] / ESTATISTICAS['total_pedidos']) * 100
            taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / ESTATISTICAS['total_pedidos']) * 100
            print(f"📈 Taxa de Conversão: {taxa_conversao:.1f}%")
            print(f"🔄 Taxa de Recuperação: {taxa_recuperacao:.1f}%")
        print("=" * 80 + "\n")
        
    except Exception as e:
        print(f"❌ Erro em atualizar_estatistica: {e}")


# 🔥 CONTADOR FINAL - RESET DIÁRIO ÀS 00:00
CONTADOR_UTILIZADORES = {
    "contador": 0,
    "utilizadores_unicos": set(),
    "reset_feito_hoje": False,
    "estatisticas_enviadas": False
}

async def enviar_estatisticas_diarias(utilizadores_antes):
    """Envia as estatísticas diárias para o chat especificado"""
    try:
        # ✅ CARREGAR CANAL DE LOGS DO .env
        CANAL_LOGS = os.getenv("CANAL_LOGS")
        
        if not CANAL_LOGS:
            print("⚠️ AVISO: CANAL_LOGS não configurado")
            return False
        
        try:
            chat_id = int(CANAL_LOGS)
        except ValueError:
            print("⚠️ AVISO: CANAL_LOGS inválido")
            return False


        data_hoje = datetime.now().strftime('%d/%m/%Y')
        hora_atual = datetime.now().strftime('%H:%M')
        
        mensagem = (
            f"📊 *ESTATÍSTICAS DIÁRIAS - {data_hoje}*\n"
            f"┌─────────────────────────────┐\n"
            f"│ 👥 Utilizadores do dia: {utilizadores_antes}\n"
            f"│ 🕐 Período: 00:00 - 23:59\n"
            f"│ 📈 Utilizadores únicos: {len(CONTADOR_UTILIZADORES['utilizadores_unicos'])}\n"
            f"└─────────────────────────────┘\n"
            f"\n"
            f"🔄 *Próximo reset: Amanhã às 00:00*"
        )
        
        await application.bot.send_message(
            chat_id=chat_id,
            text=mensagem,
            parse_mode='Markdown'
        )
        
        print(f"✅ ESTATÍSTICAS ENVIADAS para o chat {chat_id}")
        print(f"   • Utilizadores do dia: {utilizadores_antes}")
        print(f"   • Data: {data_hoje}")
        
        return True
        
    except Exception as e:
        print(f"❌ ERRO ao enviar estatísticas: {e}")
        return False

def verificar_reset_0000():
    """Verifica se precisa de reset à meia-noite (00:00)"""
    global CONTADOR_UTILIZADORES
    
    agora = datetime.now()
    hora_atual = agora.hour
    minuto_atual = agora.minute
    
    print(f"🔍 VERIFICAÇÃO RESET 00:00 | Agora: {hora_atual:02d}:{minuto_atual:02d} | Contador: {CONTADOR_UTILIZADORES['contador']} | Reset feito hoje: {CONTADOR_UTILIZADORES['reset_feito_hoje']}")
    
    # Se já passou da meia-noite (00:00) E ainda não resetou hoje
    ja_passou_0000 = hora_atual == 0 and minuto_atual >= 0
    
    if ja_passou_0000 and not CONTADOR_UTILIZADORES["reset_feito_hoje"]:
        utilizadores_antes = CONTADOR_UTILIZADORES["contador"]
        
        print(f"\n" + "="*70)
        print(f"🌙 🌙 🌙  RESET AUTOMÁTICO À MEIA-NOITE (00:00) 🌙 🌙 🌙")
        print(f"="*70)
        print(f"📊 ESTATÍSTICAS DO RESET:")
        print(f"   • Utilizadores ANTES do reset: {utilizadores_antes}")
        print(f"   • Data/hora: {agora.strftime('%d/%m/%Y %H:%M:%S')}")
        print(f"="*70)
        
        # 🔥 ENVIAR ESTATÍSTICAS ANTES DO RESET
        if not CONTADOR_UTILIZADORES["estatisticas_enviadas"] and utilizadores_antes > 0:
            print(f"   • 📤 ENVIANDO ESTATÍSTICAS (ANTES DO RESET)...")
            try:
                import asyncio
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    asyncio.create_task(enviar_estatisticas_diarias(utilizadores_antes))
                else:
                    loop.run_until_complete(enviar_estatisticas_diarias(utilizadores_antes))
            except Exception as e:
                print(f"   • ⚠️ Não foi possível enviar estatísticas: {e}")
        else:
            print(f"   • ℹ️ Nenhum utilizador hoje ou estatísticas já enviadas")
        
        # Fazer o reset
        CONTADOR_UTILIZADORES = {
            "contador": 0,
            "utilizadores_unicos": set(),
            "reset_feito_hoje": True,
            "estatisticas_enviadas": True
        }
        
        print(f"🔄 CONTADOR RESETADO para 0")
        print(f"   • Utilizadores DEPOIS do reset: {CONTADOR_UTILIZADORES['contador']}")
        print(f"="*70 + "\n")
        return True
    
    # 🔥 Resetar flag durante o dia (após a meia-noite)
    if hora_atual >= 1:
        CONTADOR_UTILIZADORES["reset_feito_hoje"] = False
        CONTADOR_UTILIZADORES["estatisticas_enviadas"] = False
    
    return False

def atualizar_contador_utilizadores(user_id=None):
    """Atualiza o contador de utilizadores ativos - reset à meia-noite"""
    global CONTADOR_UTILIZADORES
    
    # 🔥 SEMPRE verificar reset antes de qualquer operação
    reset_ocorreu = verificar_reset_0000()
    
    if reset_ocorreu:
        print(f"🔄 Reset ocorreu à meia-noite, contador agora é: {CONTADOR_UTILIZADORES['contador']}")
    
    # 🔥 INCREMENTAR CONTADOR (se for um user real)
    if user_id:
        if user_id not in CONTADOR_UTILIZADORES["utilizadores_unicos"]:
            CONTADOR_UTILIZADORES["contador"] += 1
            CONTADOR_UTILIZADORES["utilizadores_unicos"].add(user_id)
            print(f"👤 NOVO UTILIZADOR ATIVO: {user_id} | Total desde 00:00: {CONTADOR_UTILIZADORES['contador']}")
        else:
            print(f"👤 UTILIZADOR JÁ CONTABILIZADO: {user_id}")
    
    return CONTADOR_UTILIZADORES["contador"]

def obter_utilizadores_ativos():
    """Retorna o número de utilizadores ativos desde a meia-noite"""
    global CONTADOR_UTILIZADORES
    
    try:
        # 🔥 VERIFICAR SE É DICIONÁRIO OU INTEIRO
        if isinstance(CONTADOR_UTILIZADORES, dict):
            verificar_reset_0000()
            return CONTADOR_UTILIZADORES.get("contador", 0)
        elif isinstance(CONTADOR_UTILIZADORES, int):
            # Se for inteiro, converter para dicionário
            print(f"⚠️ CONTADOR_UTILIZADORES é int, convertendo para dict...")
            CONTADOR_UTILIZADORES = {
                "contador": CONTADOR_UTILIZADORES,
                "utilizadores_unicos": set(),
                "reset_feito_hoje": False,
                "estatisticas_enviadas": False
            }
            verificar_reset_0000()
            return CONTADOR_UTILIZADORES["contador"]
        else:
            print(f"⚠️ CONTADOR_UTILIZADORES tipo inválido: {type(CONTADOR_UTILIZADORES)}")
            # Reinicializar
            CONTADOR_UTILIZADORES = {
                "contador": 0,
                "utilizadores_unicos": set(),
                "reset_feito_hoje": False,
                "estatisticas_enviadas": False
            }
            return 0
            
    except Exception as e:
        print(f"❌ ERRO em obter_utilizadores_ativos: {e}")
        # Reinicializar em caso de erro
        CONTADOR_UTILIZADORES = {
            "contador": 0,
            "utilizadores_unicos": set(),
            "reset_feito_hoje": False,
            "estatisticas_enviadas": False
        }
        return 0
    

def forcar_reset_0000():
    """Força o reset do contador manualmente"""
    global CONTADOR_UTILIZADORES
    
    CONTADOR_UTILIZADORES = {
        "contador": 0,
        "utilizadores_unicos": set(),
        "reset_feito_hoje": True,
        "estatisticas_enviadas": False
    }
    
    print(f"🔄 RESET MANUAL FORÇADO!")
    print(f"   • Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    print(f"   • Contador: 0")

async def enviar_estatisticas_manualmente():
    """Função para enviar estatísticas manualmente"""
    try:
        utilizadores_atuais = CONTADOR_UTILIZADORES["contador"]
        success = await enviar_estatisticas_diarias(utilizadores_atuais)
        if success:
            print("✅ Estatísticas enviadas manualmente com sucesso!")
            CONTADOR_UTILIZADORES["estatisticas_enviadas"] = True
        else:
            print("❌ Falha ao enviar estatísticas manualmente")
        return success
    except Exception as e:
        print(f"❌ ERRO ao enviar estatísticas manualmente: {e}")
        return False

def ver_estado_contador_0000():
    """Mostra o estado atual do contador para meia-noite"""
    global CONTADOR_UTILIZADORES
    
    agora = datetime.now()
    hora_atual = agora.hour
    minuto_atual = agora.minute
    
    print(f"\n📊 ESTADO ATUAL DO CONTADOR 00:00:")
    print(f"   • Agora: {hora_atual:02d}:{minuto_atual:02d}")
    print(f"   • Reset programado: 00:00 (todos os dias)")
    
    # 🔥 VERIFICAR SE É DICIONÁRIO OU INTEIRO
    if isinstance(CONTADOR_UTILIZADORES, dict):
        print(f"   • Utilizadores desde último reset: {CONTADOR_UTILIZADORES.get('contador', 0)}")
        print(f"   • Utilizadores únicos: {len(CONTADOR_UTILIZADORES.get('utilizadores_unicos', set()))}")
        print(f"   • Reset feito hoje: {CONTADOR_UTILIZADORES.get('reset_feito_hoje', False)}")
        print(f"   • Estatísticas enviadas: {CONTADOR_UTILIZADORES.get('estatisticas_enviadas', False)}")
    elif isinstance(CONTADOR_UTILIZADORES, int):
        print(f"   • Utilizadores desde último reset: {CONTADOR_UTILIZADORES}")
        print(f"   • ⚠️ CONTADOR está como INT (deveria ser dict)")
        print(f"   • Reset feito hoje: N/A")
        print(f"   • Estatísticas enviadas: N/A")
        
        # 🔥 CONVERTER DE VOLTA PARA DICIONÁRIO
        CONTADOR_UTILIZADORES = {
            "contador": CONTADOR_UTILIZADORES,
            "utilizadores_unicos": set(),
            "reset_feito_hoje": False,
            "estatisticas_enviadas": False
        }
        print(f"   • 🔄 Convertido de volta para dicionário")
    else:
        print(f"   • ❌ Tipo desconhecido: {type(CONTADOR_UTILIZADORES)}")
        print(f"   • Valor: {CONTADOR_UTILIZADORES}")
    
    # Verificar se já passou da meia-noite
    ja_passou_0000 = hora_atual == 0
    
    if ja_passou_0000:
        print(f"   • ✅ JÁ PASSOU DA MEIA-NOITE")
        if isinstance(CONTADOR_UTILIZADORES, dict) and CONTADOR_UTILIZADORES.get("reset_feito_hoje", False):
            print(f"   • ✅ RESET JÁ FEITO HOJE")
        else:
            print(f"   • ❌ RESET PENDENTE! (deveria ter resetado à meia-noite)")
    else:
        # Calcular horas e minutos restantes
        horas_restantes = 23 - hora_atual
        minutos_restantes = 60 - minuto_atual
        if minutos_restantes == 60:
            minutos_restantes = 0
            horas_restantes += 1
        
        print(f"   • ⏳ Faltam {horas_restantes}h {minutos_restantes}m para as 00:00")
    
    print()















def normalizar_nome_pais(pais):
    """Normaliza o nome do país para coincidir com as chaves do dicionário"""
    normalizacao = {
        "portugal": "portugal",
        "espanha": "espanha", 
        "frança": "franca",
        "alemanha": "alemanha",
        "bélgica": "belgica",
        "reino unido": "reino_unido",
        "estados unidos": "estados_unidos",
        "canada": "canada",
        "países baixos": "paises_baixos",
        "brasil": "brasil",
        "irlanda": "irlanda",
        "itália": "italia",
        "luxemburgo": "luxemburgo"
    }
    return normalizacao.get(pais.lower(), "portugal")  # Default para Portugal







def get_simbolo_moeda(currency):
    """Retorna o símbolo da moeda"""
    simbolos = {
        "eur": "€",
        "gbp": "£", 
        "usd": "$",
        "cad": "C$",
        "brl": "R$"
    }
    return simbolos.get(currency, "€")




def get_moeda_do_pais(pais):
    """Versão DEFINITIVA SIMPLES - 5 países especiais (COM CANADÁ)"""
    if not pais:
        return "eur"
    
    pais = str(pais).lower().strip()
    print(f"🔍 País recebido: '{pais}'")
    
    # 🔥 AGORA 5 PAÍSES ESPECIAIS, TODO O RESTO É EUR
    if any(x in pais for x in ["reino_unido", "reino unido", "united kingdom", "uk"]):
        return "gbp"
    elif any(x in pais for x in ["estados_unidos", "estados unidos", "united states", "usa", "us"]):
        return "usd"
    elif any(x in pais for x in ["brasil", "brazil"]):
        return "brl"
    elif any(x in pais for x in ["canada", "canadá", "can"]):  # 🔥 NOVO: CANADÁ
        return "cad"
    else:
        # 🔥 TODOS OS OUTROS PAÍSES SÃO EUR
        return "eur"

# 🔥 🔥 🔥 TESTE ATUALIZADO COM CANADÁ
print("🧪 TESTE MOEDAS (COM CAD):")
test_paises = ["estados_unidos", "reino_unido", "brasil", "canada", "portugal", "frança"]
for p in test_paises:
    moeda = get_moeda_do_pais(p)
    print(f"   {p} → {moeda.upper()}")













CACHE_CAMBIO = {
    "taxas": None,
    "ultima_atualizacao": 0,
    "validade": 3600  # 1 hora em segundos
}

def obter_taxas_cambio_em_tempo_real():
    """Obtém taxas de câmbio - VERSÃO SIMPLES COM CAD"""
    agora = time.time()
    
    # 🔥 USAR CACHE SE AINDA FOR VÁLIDO
    if (CACHE_CAMBIO["taxas"] is not None and 
        agora - CACHE_CAMBIO["ultima_atualizacao"] < CACHE_CAMBIO["validade"]):
        print("💾 Usando taxas de câmbio em cache")
        return CACHE_CAMBIO["taxas"]
    
    print("🔄 Atualizando taxas de câmbio...")
    
    try:
        response = requests.get("https://api.frankfurter.app/latest?from=EUR", timeout=10)
        response.raise_for_status()
        data = response.json()
        
        # 🔥 TAXAS COM VALORES PADRÃO SE A API NÃO TIVER
        taxas = {
            "eur": Decimal("1.00"),
            "gbp": Decimal(str(data["rates"].get("GBP", 0.85))).quantize(Decimal('0.0001')),
            "usd": Decimal(str(data["rates"].get("USD", 1.08))).quantize(Decimal('0.0001')),
            "brl": Decimal(str(data["rates"].get("BRL", 5.80))).quantize(Decimal('0.0001')),
            "cad": Decimal(str(data["rates"].get("CAD", 1.45))).quantize(Decimal('0.0001'))  # 🔥 CAD
        }
        
        print(f"✅ Taxas: GBP={taxas['gbp']}, USD={taxas['usd']}, BRL={taxas['brl']}, CAD={taxas['cad']}")
        
        # 🔥 ATUALIZAR CACHE
        CACHE_CAMBIO["taxas"] = taxas
        CACHE_CAMBIO["ultima_atualizacao"] = agora
        
        return taxas
        
    except Exception as e:
        print(f"❌ Erro API: {e}")
        # 🔥 FALLBACK COM CAD
        taxas_fallback = {
            "eur": Decimal("1.00"),
            "gbp": Decimal("0.85"),
            "usd": Decimal("1.08"), 
            "brl": Decimal("5.80"),
            "cad": Decimal("1.45")  # 🔥 CAD
        }
        print("🔄 Usando fallback com CAD")
        return taxas_fallback
        



def calcular_total_por_moeda(context, pais=None):
    """Pega o total REAL e converte com taxas automáticas - VERSÃO CORRIGIDA"""
    
    # 🔥 USAR SUAS FUNÇÕES EXISTENTES
    if pais is None:
        pais = context.user_data.get("pais", "portugal")
    
    print(f"💰 Convertendo moeda para: {pais}")
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: NÃO ALTERAR O PAÍS NO CONTEXT!
    # Em vez disso, calcular o total com o país REAL
    
    # 🔥 PASSO 1: CALCULAR COM O PAÍS REAL (não forçar Portugal)
    totais_reais = calcular_total(context)  # ← Já usa o país correto do context
    
    total_eur = totais_reais['total']
    subtotal_eur = totais_reais['subtotal']
    imposto_eur = totais_reais['imposto']
    frete_eur = totais_reais['frete']
    
    print(f"📊 Total REAL em EUR para {pais}: €{total_eur:.2f}")
    print(f"   • Subtotal: €{subtotal_eur:.2f}")
    print(f"   • Imposto: €{imposto_eur:.2f}")
    print(f"   • Frete: €{frete_eur:.2f}")
    
    # 🔥 PASSO 2: USAR SUA get_moeda_do_pais EXISTENTE
    currency = get_moeda_do_pais(pais)
    simbolo = get_simbolo_moeda(currency)
    
    print(f"🌍 Moeda do país: {currency.upper()} {simbolo}")
    
    # 🔥 SE FOR EUR, NÃO PRECISA CONVERTER
    if currency == "eur":
        print("✅ País da Zona Euro - sem conversão necessária")
        return {
            'subtotal': subtotal_eur,
            'imposto': imposto_eur,
            'frete': frete_eur,
            'total': total_eur,
            'taxa': totais_reais['taxa'],
            'moeda': "EUR",
            'simbolo_moeda': "€",
            'pais': pais.title()
        }
    
    # 🔥 PASSO 3: OBTER TAXA DE CÂMBIO ATUAL
    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
    taxa = TAXAS_CAMBIO.get(currency, Decimal("1.0"))
    
    print(f"🔁 Taxa atual {currency.upper()}/EUR: {taxa}")
    
    # 🔥 PASSO 4: CONVERTER VALORES USANDO Decimal PARA PRECISÃO
    subtotal_convertido = float(Decimal(str(subtotal_eur)) * taxa)
    imposto_convertido = float(Decimal(str(imposto_eur)) * taxa)
    frete_convertido = float(Decimal(str(frete_eur)) * taxa)
    total_convertido = float(Decimal(str(total_eur)) * taxa)
    
    print(f"🔄 Conversão automática para {currency.upper()}:")
    print(f"   • Subtotal: €{subtotal_eur:.2f} → {simbolo}{subtotal_convertido:.2f}")
    print(f"   • Imposto: €{imposto_eur:.2f} → {simbolo}{imposto_convertido:.2f}")
    print(f"   • Frete: €{frete_eur:.2f} → {simbolo}{frete_convertido:.2f}")
    print(f"   • TOTAL: €{total_eur:.2f} → {simbolo}{total_convertido:.2f}")
    
    return {
        'subtotal': subtotal_convertido,
        'imposto': imposto_convertido,
        'frete': frete_convertido,
        'total': total_convertido,
        'taxa': totais_reais['taxa'],
        'moeda': currency.upper(),
        'simbolo_moeda': simbolo,
        'pais': pais.title()
    }






























def calcular_preco_tamanho_45(pedido):
    """Calcula o preço com 20% de desconto DIRETO no total - 20% EXATOS"""
    
    total_original = pedido['total']
    
    # 🔥 APLICAR 20% DE DESCONTO DIRETAMENTE NO TOTAL
    total_com_desconto = total_original * 0.80  # Isso é exatamente 20% de desconto
    
    print(f"🔍 DEBUG cálculo 4.5cm - 20% EXATOS:")
    print(f"   - Total original: €{total_original:.2f}")
    print(f"   - Total com 20% desconto: €{total_com_desconto:.2f}")
    print(f"   - Economia: €{total_original - total_com_desconto:.2f} (20% exato)")
    
    return total_com_desconto




def calcular_oferta_portachaves(pedido):
    """Calcula o preço final do porta-chaves com 70% desconto DIRETO no total - 70% EXATOS"""
    
    print(f"🔍🔍🔍 INICIANDO calcular_oferta_portachaves")
    print(f"   - Pedido ID: {pedido.get('id')}")
    print(f"   - País: {pedido.get('pais')}")
    
    # 🔥 OBTER MOEDA CORRETA BASEADA NO PAÍS (FORÇAR CORREÇÃO)
    pais = pedido["pais"].lower()
    
    # 🔥 DEFINIR MOEDA CORRETA PELO PAÍS
    moeda_por_pais = {
        "portugal": ("€", "EUR"),
        "espanha": ("€", "EUR"), 
        "frança": ("€", "EUR"),
        "alemanha": ("€", "EUR"),
        "itália": ("€", "EUR"),
        "bélgica": ("€", "EUR"),
        "países baixos": ("€", "EUR"),
        "luxemburgo": ("€", "EUR"),
        "irlanda": ("€", "EUR"),
        "estados unidos": ("$", "USD"),
        "canada": ("C$", "CAD"),
        "brasil": ("R$", "BRL"),
        "reino unido": ("£", "GBP")
    }
    
    # Buscar moeda pelo país ou usar padrão EUR
    moeda, codigo_moeda = moeda_por_pais.get(pais, ("€", "EUR"))
    
    print(f"   🔥 MOEDA DEFINIDA POR PAÍS:")
    print(f"      - País: {pais}")
    print(f"      - Moeda forçada: '{moeda}'")
    print(f"      - Código forçado: '{codigo_moeda}'")
    
    # 🔥 NORMALIZAR NOME DO PAÍS E OBTER IMPOSTO
    pais_normalizado = normalizar_nome_pais(pais)
    taxas_pais = TAXAS_PAISES.get(pais_normalizado, TAXAS_PAISES["portugal"])
    taxa_imposto_decimal = Decimal(str(taxas_pais["imposto"]))
    
    print(f"   - Taxa imposto: {taxas_pais['imposto']*100}%")
    
    # 🔥 USAR O VALOR ORIGINAL REAL DO PEDIDO INICIAL
    valor_original_pedido = pedido.get("total_original_real", pedido.get("total_original", pedido["total"]))
    
    print(f"   - Valor original real: {moeda}{valor_original_pedido:.2f}")
    
    # 🔥 APLICAR 70% DE DESCONTO DIRETAMENTE NO TOTAL
    total_com_desconto = Decimal(str(valor_original_pedido)) * Decimal("0.30")
    
    print(f"🔍 DEBUG cálculo porta-chaves - 70% EXATOS:")
    print(f"   - Total original: {moeda}{float(valor_original_pedido):.2f}")
    print(f"   - Total com 70% desconto: {moeda}{float(total_com_desconto):.2f}")
    print(f"   - Economia: {moeda}{float(Decimal(str(valor_original_pedido)) - total_com_desconto):.2f} (70% exato)")
    
    # 🔥 DEFINIR FRETE BASE EM EUR - COM CANADÁ
    if pais == "portugal":
        frete_base_eur = Decimal("6.50")
    elif pais in ["espanha", "frança", "franca", "alemanha", "bélgica", "belgica", "países baixos", "paises baixos", "holanda", "irlanda", "itália", "italia", "luxemburgo"]:
        frete_base_eur = Decimal("10.00")
    elif pais in ["brasil", "estados unidos", "canada"]:  # 🔥 CANADÁ MESMO FRETE QUE US/BR
        frete_base_eur = Decimal("25.00")  # 🔥 FRETE BASE EM EUR: 25€
    elif pais == "reino unido":
        frete_base_eur = Decimal("15.00")
    else:
        frete_base_eur = Decimal("15.00")
    
    print(f"   - Frete base (EUR): €{float(frete_base_eur):.2f}")
    
    # 🔥 CONVERTER FRETE DE EUR PARA MOEDA DO CLIENTE - COM CAD
    try:
        response = requests.get("https://api.frankfurter.app/latest?from=EUR", timeout=10)
        response.raise_for_status()
        data = response.json()
        
        print(f"   🔍 API Response:")
        print(f"      Base: {data['base']}")
        print(f"      Rates: {data['rates']}")
        
        # 🔥 BUSCAR TAXA CORRETA BASEADA NO CÓDIGO DA MOEDA - COM CAD
        if codigo_moeda == "EUR":
            taxa_cliente = Decimal("1.00")
            print(f"      ✅ Cliente usa EUR - taxa: 1.00")
        elif codigo_moeda in data['rates']:
            taxa_cliente = Decimal(str(data['rates'][codigo_moeda]))
            print(f"      ✅ Taxa EUR→{codigo_moeda}: {taxa_cliente}")
        else:
            print(f"      ⚠️ {codigo_moeda} não encontrado, usando fallback")
            taxas_fallback = {
                "USD": Decimal("1.1648"),
                "GBP": Decimal("0.8846"), 
                "BRL": Decimal("6.1764"),
                "CAD": Decimal("1.4500"),  # 🔥 NOVO: CAD
                "EUR": Decimal("1.00")
            }
            taxa_cliente = taxas_fallback.get(codigo_moeda, Decimal("1.00"))
        
        # 🔥 CONVERTER FRETE CORRETAMENTE
        frete = (frete_base_eur * taxa_cliente).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        
        print(f"   🔍 Conversão frete CORRIGIDA:")
        print(f"      Frete base: €{frete_base_eur}")
        print(f"      Taxa EUR→{codigo_moeda}: {taxa_cliente}")
        print(f"      Cálculo: {frete_base_eur} × {taxa_cliente} = {frete}")
        print(f"   ✅ Frete convertido: {moeda}{float(frete):.2f}")
        
    except Exception as e:
        print(f"   ⚠️  Erro na API: {e}")
        taxas_fallback = {
            "USD": Decimal("1.1648"),
            "GBP": Decimal("0.8846"), 
            "BRL": Decimal("6.1764"),
            "CAD": Decimal("1.4500"),  # 🔥 NOVO: CAD
            "EUR": Decimal("1.00")
        }
        taxa_cliente = taxas_fallback.get(codigo_moeda, Decimal("1.00"))
        frete = (frete_base_eur * taxa_cliente).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        print(f"   ✅ Frete fallback: {moeda}{float(frete):.2f}")
    
    # 🔥 CÁLCULO CORRETO COM/SEM IMPOSTO
    if taxa_imposto_decimal > 0:
        subtotal = ((total_com_desconto - frete) / (Decimal("1") + taxa_imposto_decimal)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        imposto = (subtotal * taxa_imposto_decimal).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        print(f"   - Cálculo COM imposto: B = ({total_com_desconto} - {frete}) / (1 + {taxa_imposto_decimal})")
    else:
        subtotal = (total_com_desconto - frete).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        imposto = Decimal("0.00")
        print(f"   - Cálculo SEM imposto: B = {total_com_desconto} - {frete}")
    
    print(f"   - Subtotal (B): {moeda}{float(subtotal):.2f}")
    print(f"   - Imposto: {moeda}{float(imposto):.2f}")
    print(f"   - Frete: {moeda}{float(frete):.2f}")
    
    # 🔥 VERIFICAR CONSISTÊNCIA
    total_calculado = subtotal + imposto + frete
    print(f"   - Verificação: {subtotal} + {imposto} + {frete} = {total_calculado}")
    print(f"   - Total esperado: {total_com_desconto}")
    
    print(f"🔍🔍🔍 FINAL calcular_oferta_portachaves")
    
    return {
        "nome": "Porta-chaves Especial",
        "subtotal": float(subtotal),
        "frete": float(frete),
        "imposto": float(imposto),
        "total": float(total_com_desconto),
        "descricao": "Cartoon in keychain format - Perfect to carry with you!",
        "economia": float(Decimal(str(valor_original_pedido)) - total_com_desconto),
        "moeda": moeda,
        "codigo_moeda": codigo_moeda,
        "valor_original": float(valor_original_pedido)
    }

def converter_codigo_para_simbolo(codigo_moeda):
    """Converte código da moeda para símbolo - COM CAD"""
    conversao = {
        "USD": "$",
        "EUR": "€", 
        "GBP": "£",
        "BRL": "R$",
        "CAD": "C$"  # 🔥 NOVO: Dólar Canadiano
    }
    return conversao.get(codigo_moeda.upper(), "€")










def cancelar_temporizador_30min(chat_id):
    """Cancela o temporizador de 30min para um chat específico"""
    if chat_id in TEMPORIZADORES_ATIVOS:
        task = TEMPORIZADORES_ATIVOS[chat_id]
        if not task.done():
            task.cancel()
            print(f"✅ Temporizador 30min CANCELADO para chat {chat_id}")
        del TEMPORIZADORES_ATIVOS[chat_id]
    else:
        print(f"ℹ️  Nenhum temporizador ativo para chat {chat_id}")


def calcular_total(context):
    # 🔥 CORREÇÃO: Melhor normalização do tipo
    tipo_bruto = context.user_data.get("tipo_cartoon", "")
    tipo = tipo_bruto.lower().replace(" ", "_").replace("🐱", "").replace("🐶", "").replace("👨‍👩‍👧", "").strip()
    
    estilo = context.user_data.get("estilo_cartoon", "")
    preco_tamanho = context.user_data.get("preco_tamanho", 0.0)
    pais = context.user_data.get("pais", "").lower().replace(" ", "_")

    print(f"🔧 DEBUG calcular_total - INÍCIO:")
    print(f"  Tipo cartoon bruto: {tipo_bruto}")
    print(f"  Tipo normalizado: {tipo}")
    print(f"  Preço tamanho: {preco_tamanho}")
    print(f"  País: {pais}")

    subtotal = 0.0

    # 🔥 CALCULAR PREÇO PARA ANIMAL (COM MELHOR DETECÇÃO)
    if "animal" in tipo or "tier" in tipo or "animale" in tipo or "mascota" in tipo or "pet" in tipo:
        print("🔧 Calculando preço para ANIMAL")
        preco_base = PRECOS["cartoon_animal"]  # ✅ Agora 40.0
        subtotal = preco_base + preco_tamanho
        print(f"  Preço base Animal: €{preco_base:.2f}")
        print(f"  Preço tamanho: €{preco_tamanho:.2f}")
        print(f"  Subtotal Animal: €{subtotal:.2f}")
    
    # 🔥 CALCULAR PREÇO PARA FAMILY/GRUPO - TODOS OS IDIOMAS DOS BOTÕES
    elif ("family" in tipo or 
          "grupo" in tipo or     # Português, Espanhol
          "group" in tipo or     # Inglês
          "gruppe" in tipo or    # Alemão
          "groupe" in tipo or    # Francês
          "gruppo" in tipo):     # Italiano
        print("🔧 Calculando preço para FAMILY/GRUPO")
        try:
            adultos = int(context.user_data.get("adultos_family", 0))
            criancas = int(context.user_data.get("criancas_family", 0)) 
            animais = int(context.user_data.get("animais_family", 0))
            
            subtotal = (adultos * PRECOS_FAMILY["adulto"] + 
                       criancas * PRECOS_FAMILY["crianca"] + 
                       animais * PRECOS_FAMILY["animal"] + 
                       preco_tamanho)
            print(f"  Adultos: {adultos} x €{PRECOS_FAMILY['adulto']:.2f}")
            print(f"  Crianças: {criancas} x €{PRECOS_FAMILY['crianca']:.2f}")
            print(f"  Animais: {animais} x €{PRECOS_FAMILY['animal']:.2f}")
            print(f"  Preço tamanho: €{preco_tamanho:.2f}")
            print(f"  Subtotal: €{subtotal:.2f}")
        except Exception as e:
            print(f"❌ Erro no cálculo family: {e}")
            subtotal = PRECOS["cartoon_family"]
    
    # 🔥 CALCULAR PREÇO PARA INDIVIDUAL - TODOS OS IDIOMAS DOS BOTÕES
    elif ("individual" in tipo or      # Português, Inglês, Espanhol
          "individuale" in tipo or     # Italiano
          "individuel" in tipo or      # Francês
          "individueller" in tipo or   # Alemão
          "einzel" in tipo or          # Alemão (alternativo)
          "einzeln" in tipo):          # Alemão
        print("🔧 Calculando preço para INDIVIDUAL")
        preco_base = PRECOS["cartoon_individual"]
        preco_estilo = PRECOS_ESTILO.get(estilo, 0)
        subtotal = preco_base + preco_estilo + preco_tamanho
        print(f"  Preço base: €{preco_base:.2f}")
        print(f"  Preço estilo: €{preco_estilo:.2f}")
        print(f"  Preço tamanho: €{preco_tamanho:.2f}")
        print(f"  Subtotal: €{subtotal:.2f}")

    # 🔥 CALCULAR PREÇO PARA PERSONALIZADO - TODOS OS IDIOMAS DOS BOTÕES
    elif ("personalizado" in tipo or  # Português, Espanhol
          "custom" in tipo or         # Inglês
          "personalizzato" in tipo or # Italiano
          "personalisiert" in tipo or # Alemão
          "personnalisé" in tipo or   # Francês
          "personal." in tipo or      # 🔥 para "Personal. Karikatur"
          "personal_" in tipo):       # 🔥 para tipo normalizado
        print("🔧 🔥 🔥 Calculando preço para PERSONALIZADO")
        preco_base = PRECOS["cartoon_custom"]  # €90.00
        subtotal = preco_base + preco_tamanho
        print(f"  Preço base Personalizado: €{preco_base:.2f}")
        print(f"  Preço tamanho: €{preco_tamanho:.2f}")
        print(f"  Subtotal Personalizado: €{subtotal:.2f}")    
    
    # 🔥 CALCULAR PREÇO PARA OUTROS TIPOS
    else:
        print("🔧 Calculando preço para OUTRO TIPO (fallback)")
        preco_base = PRECOS.get(tipo, 0)
        subtotal = preco_base + preco_tamanho
        print(f"  Preço tipo '{tipo}': €{preco_base:.2f}")
        print(f"  Preço tamanho: €{preco_tamanho:.2f}")
        print(f"  Subtotal: €{subtotal:.2f}")
    
    # Calcular impostos e frete
    taxas = TAXAS_PAISES.get(pais, {"imposto": 0.1, "frete": 10})
    imposto = subtotal * taxas["imposto"]
    frete = taxas["frete"]
    total = subtotal + imposto + frete
    
    print(f"💰 RESULTADO FINAL:")
    print(f"  Subtotal: €{subtotal:.2f}")
    print(f"  Imposto ({taxas['imposto']*100}%): €{imposto:.2f}")
    print(f"  Frete: €{frete:.2f}")
    print(f"  TOTAL: €{total:.2f}")
    
    return {
        "subtotal": subtotal,
        "imposto": imposto, 
        "frete": frete,
        "taxa": taxas["imposto"],
        "total": total
    }








#Gift 

async def gift_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Iniciar fluxo do gift (porta-chaves) COM ESCOLHA DE IDIOMA"""
    chat_id = update.message.chat_id
    user_id = update.effective_user.id
    
    print(f"🎁 GIFT INICIADO por user {user_id}")
    print("🎁 GIFT COMMAND - INICIANDO")
    print(f"🔍 Context user_data ANTES: {context.user_data}")

    # 🔥 CANCELAR TEMPORIZADORES
    try:
        cancelar_temporizador_30min(chat_id)
    except:
        pass
    
    # 🔥 LIMPAR DADOS
    context.user_data.clear()
    
    print(f"✅ Dados limpos e temporizador cancelado para chat {chat_id}")
    
    # 🔥 🔥 🔥 **PRIMEIRO: PEDIR PARA ESCOLHER IDIOMA (IGUAL AO START)**
    texto_escolha_idioma = "🌍 *Please choose your language / Por favor escolha seu idioma:*"
    
    keyboard = [
        [
            InlineKeyboardButton("🇵🇹 Português", callback_data="gift_idioma_portugues"),
            InlineKeyboardButton("🇺🇸 English", callback_data="gift_idioma_ingles")
        ],
        [
            InlineKeyboardButton("🇪🇸 Español", callback_data="gift_idioma_espanhol"),
            InlineKeyboardButton("🇮🇹 Italiano", callback_data="gift_idioma_italiano")
        ],
        [
            InlineKeyboardButton("🇩🇪 Deutsch", callback_data="gift_idioma_alemao"),
            InlineKeyboardButton("🇫🇷 Français", callback_data="gift_idioma_frances")
        ]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        texto_escolha_idioma,
        reply_markup=reply_markup,
        parse_mode="Markdown"
    )
    
    print(f"✅ Tela de escolha de idioma para GIFT mostrada para chat {chat_id}")

async def gift_selecionar_idioma(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para seleção de idioma NO GIFT - CORRIGIDO"""
    query = update.callback_query
    await query.answer()
    
    # Extrair idioma do callback_data
    idioma = query.data.replace("gift_idioma_", "")
    
    # Salvar idioma no user_data
    context.user_data['idioma'] = idioma
    
    print(f"✅ Idioma GIFT selecionado: {idioma} por user {query.from_user.id}")
    
    # 🔥 DEFINIR EXPLICITAMENTE O TIPO DE OFERTA
    context.user_data['oferta_tipo'] = 'oferta_surpresa'
    context.user_data['_gift_oferta_surpresa'] = True
    
    # 🔥 DEFINIR TIPO COMO GIFT
    context.user_data["tipo_cartoon"] = "Porta-Chaves 🎁"
    context.user_data["estilo_cartoon"] = "Gift"
    context.user_data["tamanho_cartoon"] = "Porta-Chaves Padrão"
    context.user_data["tamanho_key"] = "padrao"
    context.user_data["preco_tamanho"] = 20.0
    
    print(f"🎯 OFERTA_TIPO DEFINIDO: {context.user_data['oferta_tipo']}")
    print("✅ Dados gift inicializados")
    
    # 🔥 CORREÇÃO: MENSAGENS COM FORMATAÇÃO CORRETA (asterisco fechado)
    mensagens_confirmacao = {
        'portugues': "✅ *Idioma definido para Português!*\n\n🎁 *Vamos criar um Porta-Chaves personalizado!*",
        'ingles': "✅ *Language set to English!*\n\n🎁 *Let's create a personalized Keychain!*",
        'espanhol': "✅ *¡Idioma establecido en Español!*\n\n🎁 *¡Vamos a crear un Llavero personalizado!*",
        'italiano': "✅ *Lingua impostata su Italiano!*\n\n🎁 *Creiamo un Portachiavi personalizzato!*",
        'alemao': "✅ *Sprache auf Deutsch eingestellt!*\n\n🎁 *Erstellen wir einen personalisierten Schlüsselanhänger!*",
        'frances': "✅ *Langue définie sur Français!*\n\n🎁 *Créons un Porte-clés personnalisé !*"
    }
    
    # Apagar mensagem de escolha de idioma
    await query.delete_message()
    
    # 🔥 ARMAZENAR ID DA MENSAGEM DE CONFIRMAÇÃO PARA APAGAR DEPOIS
    mensagem_confirmacao = await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=mensagens_confirmacao.get(idioma, "✅ Idioma selecionado!"),
        parse_mode="Markdown"
    )
    
    # Salvar o ID da mensagem de confirmação
    context.user_data['mensagem_confirmacao_id'] = mensagem_confirmacao.message_id
    print(f"📝 ID da mensagem de confirmação guardado: {mensagem_confirmacao.message_id}")
    
    # 🔥 CORREÇÃO: TEXTOS COM ASTERISCO DE FECHAMENTO
    textos_nome = {
        'portugues': "*Antes de começarmos, qual é o seu nome?*",
        'ingles': "*Before we start, what's your name?*",
        'espanhol': "*Antes de empezar, ¿cuál es su nombre?*",
        'italiano': "*Prima di iniziare, qual è il tuo nome?*",
        'alemao': "*Bevor wir beginnen, wie ist Ihr Name?*",
        'frances': "*Avant de commencer, quel est votre nom ?*"
    }
    
    # Primeira pergunta (nome) no idioma escolhido e ARMAZENAR O ID
    mensagem_pergunta_nome = await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_nome.get(idioma, textos_nome['portugues']),
        parse_mode="Markdown"
    )
    
    # Salvar o ID da mensagem da pergunta
    context.user_data['mensagem_gift_nome_id'] = mensagem_pergunta_nome.message_id
    print(f"📝 ID da mensagem de pergunta nome guardado: {mensagem_pergunta_nome.message_id}")
    
    context.user_data['conversation_state'] = GIFT_NOME
    print(f"✅ Estado atualizado para: GIFT_NOME | Idioma: {idioma}")

    

# --- Receber nome do gift ---
async def receber_gift_nome(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # 🔥 OBTER IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Recebendo nome GIFT em: {idioma}")
    
    context.user_data["nome"] = update.message.text
    
    # 🔥 APAGAR MENSAGENS
    try:
        await update.message.delete()
        print("✅ Mensagem do usuário (gift nome) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem do usuário: {e}")
    
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id,
            message_id=update.message.message_id - 1
        )
        print("✅ Mensagem da pergunta (gift nome) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar pergunta do nome: {e}")
    
    # 🔥 TEXTOS POR IDIOMA PARA PERGUNTA DE EMAIL
    textos_email = {
        'portugues': "📧 Perfeito! Agora, qual é o seu email?",
        'ingles': "📧 Perfect! Now, what's your email?",
        'espanhol': "📧 ¡Perfecto! Ahora, ¿cuál es su email?",
        'italiano': "📧 Perfetto! Ora, qual è la tua email?",
        'alemao': "📧 Perfekt! Nun, was ist Ihre E-Mail?",
        'frances': "📧 Parfait ! Maintenant, quelle est votre adresse e-mail ?"
    }
    
    # Próxima pergunta (email) no idioma correto
    await update.message.reply_text(textos_email.get(idioma, textos_email['portugues']))
    
    context.user_data['conversation_state'] = GIFT_EMAIL

# --- Receber email do gift ---
async def receber_gift_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # 🔥 OBTER IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Recebendo email GIFT em: {idioma}")
    
    context.user_data["email"] = update.message.text
    
    # 🔥 APAGAR MENSAGENS
    try:
        await update.message.delete()
        print("✅ Mensagem do usuário (gift email) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem do usuário: {e}")
    
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id,
            message_id=update.message.message_id - 1
        )
        print("✅ Mensagem da pergunta (gift email) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar pergunta do email: {e}")
    
    # 🔥 TEXTOS POR IDIOMA PARA PERGUNTA DE PAÍS
    textos_pais = {
        'portugues': "🌍 De qual país você é?",
        'ingles': "🌍 Which country are you from?",
        'espanhol': "🌍 ¿De qué país es usted?",
        'italiano': "🌍 Di quale paese sei?",
        'alemao': "🌍 Aus welchem Land kommen Sie?",
        'frances': "🌍 De quel pays êtes-vous ?"
    }
    
    texto = textos_pais.get(idioma, textos_pais['portugues'])
    
    # 🔥 BOTÕES DE PAÍSES (MANTÉM OS MESMOS, MAS CALLBACKS ESPECÍFICOS PARA GIFT)
    keyboard = [
        [InlineKeyboardButton("🇺🇸 United States", callback_data="gift_pais_estados_unidos")],
        [InlineKeyboardButton("🇨🇦 Canada", callback_data="gift_pais_canada"),
         InlineKeyboardButton("🇬🇧 United Kingdom", callback_data="gift_pais_reino_unido")],
        [InlineKeyboardButton("🇧🇷 Brazil", callback_data="gift_pais_brasil"),
         InlineKeyboardButton("🇩🇪 Germany", callback_data="gift_pais_alemanha")],
        [InlineKeyboardButton("🇳🇱 Netherlands", callback_data="gift_pais_holanda"),
         InlineKeyboardButton("🇫🇷 France", callback_data="gift_pais_franca")],
        [InlineKeyboardButton("🇪🇸 Spain", callback_data="gift_pais_espanha"),
         InlineKeyboardButton("🇧🇪 Belgium", callback_data="gift_pais_belgica")],
        [InlineKeyboardButton("🇮🇹 Italy", callback_data="gift_pais_italia"),
         InlineKeyboardButton("🇵🇹 Portugal", callback_data="gift_pais_portugal")],
        [InlineKeyboardButton("🇮🇪 Ireland", callback_data="gift_pais_irlanda"),
         InlineKeyboardButton("🇱🇺 Luxembourg", callback_data="gift_pais_luxemburgo")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(texto, reply_markup=reply_markup)
    
    context.user_data['conversation_state'] = GIFT_PAIS


# --- Handler para seleção de país do gift ---
async def selecionar_gift_pais(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    pais_data = query.data
    
    # 🔥 OBTER IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Selecionando país GIFT em: {idioma}")
    
    if pais_data == "gift_pais_outro":
        # 🔥 TEXTOS POR IDIOMA PARA PAÍS MANUAL
        textos_outro_pais = {
            'portugues': "🌍 Por favor, digite o nome do seu país:",
            'ingles': "🌍 Please, enter your country name:",
            'espanhol': "🌍 Por favor, escriba el nombre de su país:",
            'italiano': "🌍 Per favore, inserisci il nome del tuo paese:",
            'alemao': "🌍 Bitte geben Sie den Namen Ihres Landes ein:",
            'frances': "🌍 Veuillez entrer le nom de votre pays :"
        }
        
        # 🔥 APAGAR MENSAGEM DOS PAÍSES
        try:
            await query.delete_message()
            print("✅ Mensagem dos países gift apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar mensagem dos países: {e}")
        
        await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_outro_pais.get(idioma, textos_outro_pais['portugues'])
        )
        
        context.user_data['aguardando_gift_pais_manual'] = True
        return
    
    nome_pais = pais_data.replace("gift_pais_", "")
    pais_formatado = nome_pais.replace("_", " ").title()
    prefixo = PAISES_PREFIXOS.get(nome_pais, "+??")
    
    context.user_data["pais"] = pais_formatado
    context.user_data["prefixo"] = prefixo
    
    print(f"✅ País GIFT selecionado: {pais_formatado} | Prefixo: {prefixo} | Idioma: {idioma}")
    
    # 🔥 APAGAR MENSAGEM DOS PAÍSES
    try:
        await query.delete_message()
        print("✅ Mensagem dos países gift apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem dos países: {e}")
    
    # 🔥 TEXTOS POR IDIOMA PARA CONFIRMAÇÃO DE PAÍS
    textos_confirmacao = {
        'portugues': {
            'titulo': "🌍 País selecionado:",
            'prefixo': "📞 Prefixo:",
            'pergunta': "\nAgora envie o seu número de telemóvel:"
        },
        'ingles': {
            'titulo': "🌍 Selected country:",
            'prefixo': "📞 Prefix:",
            'pergunta': "\nNow send your phone number:"
        },
        'espanhol': {
            'titulo': "🌍 País seleccionado:",
            'prefixo': "📞 Prefijo:",
            'pergunta': "\nAhora envíe su número de teléfono:"
        },
        'italiano': {
            'titulo': "🌍 Paese selezionato:",
            'prefixo': "📞 Prefisso:",
            'pergunta': "\nOra invia il tuo numero di telefono:"
        },
        'alemao': {
            'titulo': "🌍 Ausgewähltes Land:",
            'prefixo': "📞 Vorwahl:",
            'pergunta': "\nSenden Sie nun Ihre Telefonnummer:"
        },
        'frances': {
            'titulo': "🌍 Pays sélectionné :",
            'prefixo': "📞 Indicatif :",
            'pergunta': "\nMaintenant, envoyez votre numéro de téléphone :"
        }
    }
    
    textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
    
    # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
    texto = f"{textos['titulo']} *{pais_formatado}*\n"
    texto += f"{textos['prefixo']} {prefixo}\n"
    texto += f"{textos['pergunta']}"
    
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=texto,
        parse_mode="Markdown"
    )
    
    context.user_data['conversation_state'] = GIFT_CONTACTO


async def gift_text_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # 🔥 VERIFICAR SE É COMANDO PRIMEIRO (PROTEÇÃO EXTRA)
    if update.message.text.startswith('/'):
        print(f"🎁 Comando detectado no gift handler: {update.message.text} - ignorando")
        # 🔥 PASSA PARA O PRÓXIMO HANDLER
        await handle_message(update, context)
        return
    
    estado = context.user_data.get('conversation_state')
    print(f"🎁 GIFT TEXT HANDLER - Estado: {estado}")
    
    # 🔥 SÓ PROCESSAR SE FOR ESTADO DO GIFT
    if estado in [GIFT_NOME, GIFT_EMAIL, GIFT_CONTACTO, GIFT_NOME_BOX, GIFT_FRASE_BOX]:
        if estado == GIFT_NOME:
            print("✅ Processando GIFT_NOME")
            await receber_gift_nome(update, context)
        elif estado == GIFT_EMAIL:
            print("✅ Processando GIFT_EMAIL")
            await receber_gift_email(update, context)
        elif estado == GIFT_CONTACTO:
            print("✅ Processando GIFT_CONTACTO")
            await receber_gift_contacto(update, context)
        elif estado == GIFT_NOME_BOX:  # 🔥 NOVO
            print("✅ Processando GIFT_NOME_BOX")
            await receber_gift_nome_box(update, context)
        elif estado == GIFT_FRASE_BOX:  # 🔥 NOVO
            print("✅ Processando GIFT_FRASE_BOX")
            await receber_gift_frase_box(update, context)
    else:
        print(f"❌ Não é estado do Gift - passando para handler genérico")
        # 🔥 PASSA PARA O HANDLER GENÉRICO
        await handle_message(update, context)



# 🔥 HANDLER PARA FOTO DO GIFT
async def gift_foto_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler ESPECÍFICO apenas para fotos do GIFT e estados de PROBLEMA"""
    user_data = context.user_data if context.user_data is not None else {}
    estado = user_data.get('conversation_state')
    idioma = user_data.get('idioma', 'portugues')
    
    print(f"🎁 GIFT FOTO HANDLER - Estado: {estado}")
    
    # 🔥 SÓ PROCESSAR SE FOR ESTADO DO GIFT_FOTO (31)
    if estado == GIFT_FOTO:
        print("✅ Processando GIFT_FOTO - Chamando receber_gift_foto")
        await receber_gift_foto(update, context)
        return
    
    # 🔥 CORREÇÃO: Estados de PROBLEMA mais específicos - ADICIONAR 'problema_outro'
    elif estado in ['problema_outro', 'problema_foto', 'problema_preco', 'problema_pagamento', 'problema_entrega']:
        print(f"📸 É foto para estado de problema: {estado}")
        await receber_reportar_problema(update, context)
        return
    
    elif estado == FOTO_PROBLEMA:
        print(f"📸 É foto de PROBLEMA - FOTO_PROBLEMA - chamando receber_problema")
        await receber_problema(update, context)
        return
    
    elif estado == AGUARDANDO_SCREENSHOT_CARTOON:
        print(f"📸 É screenshot para cartoon (estado: {estado}) - chamando receber_screenshot_cartoon")
        await receber_screenshot_cartoon(update, context)
        return
    
    # 🔥 Verificar se é estado FOTO (14) para cartoon normal
    elif estado == FOTO:
        print("📸 É foto do cartoon normal - chamando receber_foto")
        await receber_foto(update, context)
        return
    
    # 🔥 Se não for nenhum dos estados acima, enviar mensagem informativa
    else:
        print(f"❓ Estado desconhecido para foto: {estado}")
        
        textos_mensagem = {
            'portugues': "📸 *Foto recebida!*\n\nPara usar esta foto:\n👉 /start - Criar cartoon\n👉 /help - Reportar problema",
            'ingles': "📸 *Photo received!*\n\nTo use this photo:\n👉 /start - Create cartoon\n👉 /help - Report problem",
            'espanhol': "📸 *¡Foto recibida!*\n\nPara usar esta foto:\n👉 /start - Crear cartoon\n👉 /help - Informar problema",
            'italiano': "📸 *Foto ricevuta!*\n\nPer usare questa foto:\n👉 /start - Creare cartoon\n👉 /help - Segnalare problema",
            'alemao': "📸 *Foto erhalten!*\n\nUm dieses Foto zu verwenden:\n👉 /start - Cartoon erstellen\n👉 /help - Problem melden",
            'frances': "📸 *Photo reçue !*\n\nPour utiliser cette photo :\n👉 /start - Créer dessin animé\n👉 /help - Signaler problème"
        }
        
        await update.message.reply_text(
            textos_mensagem.get(idioma, textos_mensagem['portugues']),
            parse_mode="Markdown"
        )






# --- Handler para receber contacto do gift ---
async def receber_gift_contacto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber contacto no GIFT com tradução completa"""
    try:
        # 🔥 OBTER IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"🌐 Recebendo contacto GIFT em: {idioma}")
        
        if update.message.contact:
            numero_completo = update.message.contact.phone_number
            context.user_data["contacto"] = numero_completo
        else:
            numero = update.message.text.strip()
            prefixo = context.user_data.get("prefixo", "+??")
            numero_completo = numero if numero.startswith("+") else f"{prefixo} {numero}"
            context.user_data["contacto"] = numero_completo

        print(f"✅ Contacto gift: {context.user_data['contacto']} | Idioma: {idioma}")

        # 🔥 APAGAR MENSAGENS
        try:
            await update.message.delete()
            print("✅ Mensagem do contacto apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar mensagem: {e}")
        
        try:
            await context.bot.delete_message(
                chat_id=update.message.chat_id,
                message_id=update.message.message_id - 1
            )
            print("✅ Mensagem da pergunta do contacto apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar pergunta: {e}")

        # 🔥 TEXTOS POR IDIOMA PARA RESUMO DO GIFT
        textos_resumo = {
            'portugues': {
                'titulo': "🎁 *RESUMO DO SEU PORTA-CHAVES*\n\n",
                'nome': "👤 *Nome:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *País:*",
                'telefone': "📱 *Telefone:*",
                'produto': "🎁 *Produto:* Porta-Chaves Personalizado\n",
                'continue': "\n⬇️ *Continue preenchendo abaixo* ⬇️"
            },
            'ingles': {
                'titulo': "🎁 *YOUR KEYCHAIN SUMMARY*\n\n",
                'nome': "👤 *Name:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Country:*",
                'telefone': "📱 *Phone:*",
                'produto': "🎁 *Product:* Personalized Keychain\n",
                'continue': "\n⬇️ *Continue filling below* ⬇️"
            },
            'espanhol': {
                'titulo': "🎁 *RESUMEN DE SU LLAVERO*\n\n",
                'nome': "👤 *Nombre:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *País:*",
                'telefone': "📱 *Teléfono:*",
                'produto': "🎁 *Producto:* Llavero Personalizado\n",
                'continue': "\n⬇️ *Continúe rellenando abajo* ⬇️"
            },
            'italiano': {
                'titulo': "🎁 *RIEPILOGO DEL TUO PORTACHIAVI*\n\n",
                'nome': "👤 *Nome:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Paese:*",
                'telefone': "📱 *Telefono:*",
                'produto': "🎁 *Prodotto:* Portachiavi Personalizzato\n",
                'continue': "\n⬇️ *Continua a compilare qui sotto* ⬇️"
            },
            'alemao': {
                'titulo': "🎁 *ZUSAMMENFASSUNG IHRES SCHLÜSSELANHÄNGERS*\n\n",
                'nome': "👤 *Name:*",
                'email': "📧 *E-Mail:*",
                'pais': "🌍 *Land:*",
                'telefone': "📱 *Telefon:*",
                'produto': "🎁 *Produkt:* Personalisierter Schlüsselanhänger\n",
                'continue': "\n⬇️ *Fahren Sie unten fort* ⬇️"
            },
            'frances': {
                'titulo': "🎁 *RÉSUMÉ DE VOTRE PORTE-CLÉS*\n\n",
                'nome': "👤 *Nom:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Pays:*",
                'telefone': "📱 *Téléphone:*",
                'produto': "🎁 *Produit:* Porte-clés Personnalisé\n",
                'continue': "\n⬇️ *Continuez à remplir ci-dessous* ⬇️"
            }
        }
        
        textos = textos_resumo.get(idioma, textos_resumo['portugues'])
        
        # 🔥 CONSTRUIR RESUMO TRADUZIDO
        resumo = f"{textos['titulo']}"
        user_data = context.user_data

        if "nome" in user_data:
            resumo += f"{textos['nome']} {user_data['nome']}\n"
        if "email" in user_data:
            resumo += f"{textos['email']} {user_data['email']}\n"
        if "pais" in user_data:
            resumo += f"{textos['pais']} {user_data['pais']}\n"
        if "contacto" in user_data:
            resumo += f"{textos['telefone']} {user_data['contacto']}\n"
        resumo += f"{textos['produto']}"
        resumo += f"{textos['continue']}"
        
        # 🔥 ENVIAR RESUMO
        msg = await context.bot.send_message(
            chat_id=update.message.chat_id, 
            text=resumo, 
            parse_mode="Markdown"
        )
        context.user_data['resumo_msg_id'] = msg.message_id

        # 🔥 🔥 🔥 NOVO: PERGUNTAR NOME PARA A BOX (TRADUZIDO)
        textos_nome_box = {
            'portugues': "🎭 *Escreva o nome pessoal ou alcunha que irá representar o porta-chaves, para a personalização da sua box!*\n\n",
            'ingles': "🎭 *Write the personal name or nickname that will represent the keychain, for the personalization of your box!*\n\n",
            'espanhol': "🎭 *¡Escriba el nombre personal o apodo que representará el llavero, para la personalización de su caja!*\n\n",
            'italiano': "🎭 *Scrivi il nome personale o il soprannome che rappresenterà il portachiavi, per la personalizzazione della tua scatola!*\n\n",
            'alemao': "🎭 *Schreiben Sie den persönlichen Namen oder Spitznamen, der den Schlüsselanhänger für die Personalisierung Ihrer Box repräsentieren wird!*\n\n",
            'frances': "🎭 *Écrivez le nom personnel ou le surnom qui représentera le porte-clés, pour la personnalisation de votre boîte !*\n\n"
        }
        
        mensagem_nome_gift = await update.message.reply_text(
            textos_nome_box.get(idioma, textos_nome_box['portugues']),
            parse_mode="Markdown"
        )
        
        context.user_data['mensagem_nome_gift_id'] = mensagem_nome_gift.message_id
        context.user_data['conversation_state'] = GIFT_NOME_BOX
        
        print(f"✅ Estado atualizado para GIFT_NOME_BOX | Idioma: {idioma}")
        
    except Exception as e:
        print(f"❌ ERRO em receber_gift_contacto: {e}")
        import traceback
        traceback.print_exc()
        
        # 🔥 TEXTOS POR IDIOMA PARA MENSAGEM DE ERRO
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtelo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )



async def receber_gift_nome_box(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber nome/alcunha para personalização da box do Gift COM TRADUÇÃO"""
    try:
        # 🔥 OBTER IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"🌐 Recebendo nome box GIFT em: {idioma}")
        
        # 🔥 COMANDO DE PULAR POR IDIOMA (IGUAL À FAMILY)
        comandos_pular = {
            'portugues': '/skip',
            'ingles': '/skip', 
            'espanhol': '/skip',
            'italiano': '/skip',
            'alemao': '/skip',
            'frances': '/skip'
        }
        
        comando_pular = comandos_pular.get(idioma, '/skip')
        
        # Verificar se é comando /skip ou mensagem normal (IGUAL À FAMILY)
        if update.message.text and update.message.text.strip() == comando_pular:
            # 🔥 TEXTOS POR IDIOMA PARA "NÃO ADICIONOU NOME"
            textos_sem_nome = {
                'portugues': "Não adicionou nome/alcunha",
                'ingles': "No name/nickname added",
                'espanhol': "No añadió nombre/apodo",
                'italiano': "Nessun nome/soprannome aggiunto",
                'alemao': "Kein Name/Spitzname hinzugefügt",
                'frances': "Aucun nom/surnom ajouté"
            }
            
            nome_gift = textos_sem_nome.get(idioma, textos_sem_nome['portugues'])
            print(f"✅ Usuário usou {comando_pular} para nome/alcunha: {nome_gift}")
        else:
            nome_gift = update.message.text
            print(f"✅ Usuário adicionou nome/alcunha: {nome_gift}")
        
        context.user_data["nome_gift"] = nome_gift
        
        # 🔥 REMOVER MENSAGENS (IGUAL À FAMILY)
        try:
            await update.message.delete()
            print("✅ Mensagem nome gift box do usuário apagada")
        except Exception as e:
            print(f"❌ Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER MENSAGEM DA PERGUNTA
        mensagem_nome_gift_id = context.user_data.get('mensagem_nome_gift_id')
        if mensagem_nome_gift_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_gift_id
                )
                print(f"✅ Mensagem pergunta nome gift box apagada: {mensagem_nome_gift_id}")
            except Exception as e:
                print(f"❌ Não foi possível apagar pergunta nome gift box: {e}")
        
        # 🔥 APAGAR RESUMO ANTERIOR ANTES DE ENVIAR O NOVO
        resumo_antigo_id = context.user_data.get('resumo_msg_id')
        if resumo_antigo_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=resumo_antigo_id
                )
                print(f"✅ Resumo anterior gift apagado: {resumo_antigo_id}")
            except Exception as e:
                print(f"❌ Erro ao apagar resumo anterior: {e}")

        # 🔥 ENVIAR NOVO RESUMO ATUALIZADO
        new_message_id = await enviar_resumo_gift(context, update.message.chat_id)
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTOS POR IDIOMA PARA PERGUNTA DA FRASE (COM /skip SEM BACKTICKS IGUAL À FAMILY)
        textos_frase = {
            'portugues': {
                'pergunta': "💬 *Quer partilhar alguma frase ou algo que queira ficar registado, para juntar na box?*\n\n",
                'instrucao': "*Pode escrever a sua frase ou enviar* /skip *para pular:*"  # 🔥 SEM BACKTICKS
            },
            'ingles': {
                'pergunta': "💬 *Do you want to share a phrase or something you want to be recorded, to add to the box?*\n\n",
                'instrucao': "*You can write your phrase or send* /skip *to skip:*"  # 🔥 SEM BACKTICKS
            },
            'espanhol': {
                'pergunta': "💬 *¿Quiere compartir alguna frase o algo que quiera que quede registrado, para agregar a la caja?*\n\n",
                'instrucao': "*Puede escribir su frase o enviar* /skip *para saltar:*"  # 🔥 SEM BACKTICKS
            },
            'italiano': {
                'pergunta': "💬 *Vuoi condividere una frase o qualcosa che vuoi che venga registrato, per aggiungere alla scatola?*\n\n",
                'instrucao': "*Puoi scrivere la tua frase o inviare* /skip *per saltare:*"  # 🔥 SEM BACKTICKS
            },
            'alemao': {
                'pergunta': "💬 *Möchten Sie einen Spruch oder algo, das aufgezeichnet werden soll, mitteilen, um ihn der Box hinzuzufügen?*\n\n",
                'instrucao': "*Sie können Ihren Spruch schreiben oder* /skip *senden, um zu überspringen:*"  # 🔥 SEM BACKTICKS
            },
            'frances': {
                'pergunta': "💬 *Voulez-vous partager une phrase ou quelque chose que vous souhaitez voir enregistré, à ajouter à la boîte ?*\n\n",
                'instrucao': "*Vous pouvez écrire votre phrase ou envoyer* /skip *pour passer:*"  # 🔥 SEM BACKTICKS
            }
        }
        
        textos = textos_frase.get(idioma, textos_frase['portugues'])
        
        # 🔥 PRÓXIMA PERGUNTA: FRASE PARA A BOX (IGUAL À FAMILY)
        mensagem_frase_gift = await update.message.reply_text(
            f"{textos['pergunta']}"
            f"{textos['instrucao']}",
            parse_mode="Markdown"
        )
        
        context.user_data['mensagem_frase_gift_id'] = mensagem_frase_gift.message_id
        context.user_data['conversation_state'] = GIFT_FRASE_BOX
        
        print(f"✅ Estado atualizado para GIFT_FRASE_BOX | Idioma: {idioma}")
        
    except Exception as e:
        print(f"❌ ERRO em receber_gift_nome_box: {e}")
        import traceback
        traceback.print_exc()
        
        # 🔥 TEXTOS POR IDIOMA PARA MENSAGEM DE ERRO
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtelo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )




async def receber_gift_frase_box(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber frase para registar na box do Gift COM TRADUÇÃO"""
    try:
        # 🔥 OBTER IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"🌐 Recebendo frase box GIFT em: {idioma}")
        
        # 🔥 COMANDO DE PULAR POR IDIOMA (IGUAL À FAMILY)
        comandos_pular = {
            'portugues': '/skip',
            'ingles': '/skip',
            'espanhol': '/skip',
            'italiano': '/skip',
            'alemao': '/skip',
            'frances': '/skip'
        }
        
        comando_pular = comandos_pular.get(idioma, '/skip')
        
        # 🔥 TEXTOS POR IDIOMA PARA "NÃO ADICIONOU FRASE"
        textos_sem_frase = {
            'portugues': "Não adicionou frase",
            'ingles': "No phrase added",
            'espanhol': "No añadió frase",
            'italiano': "Nessuna frase aggiunta",
            'alemao': "Keine Phrase hinzugefügt",
            'frances': "Aucune phrase ajoutée"
        }
        
        # Verificar se é comando /skip (IGUAL À FAMILY)
        if update.message.text and update.message.text.strip() == comando_pular:
            frase_gift = textos_sem_frase.get(idioma, textos_sem_frase['portugues'])
            print(f"✅ Usuário usou {comando_pular} para frase: {frase_gift}")
        else:
            frase_gift = update.message.text
            print(f"✅ Usuário adicionou frase: {frase_gift}")
        
        context.user_data["frase_gift"] = frase_gift
        
        # 🔥 REMOVER MENSAGENS (IGUAL À FAMILY)
        try:
            await update.message.delete()
            print("✅ Mensagem frase gift do usuário apagada")
        except Exception as e:
            print(f"❌ Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER MENSAGEM DA PERGUNTA
        mensagem_frase_gift_id = context.user_data.get('mensagem_frase_gift_id')
        if mensagem_frase_gift_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_frase_gift_id
                )
                print(f"✅ Mensagem pergunta frase gift apagada: {mensagem_frase_gift_id}")
            except Exception as e:
                print(f"❌ Não foi possível apagar pergunta frase gift: {e}")
        
        # 🔥 APAGAR RESUMO ANTERIOR ANTES DE ENVIAR O NOVO
        resumo_antigo_id = context.user_data.get('resumo_msg_id')
        if resumo_antigo_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=resumo_antigo_id
                )
                print(f"✅ Resumo anterior gift apagado: {resumo_antigo_id}")
            except Exception as e:
                print(f"❌ Erro ao apagar resumo anterior: {e}")

        # 🔥 ENVIAR NOVO RESUMO ATUALIZADO
        new_message_id = await enviar_resumo_gift(context, update.message.chat_id)
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTOS POR IDIOMA PARA PEDIDO DE FOTO
        textos_foto = {
            'portugues': {
                'titulo': "📸 *Perfeito! Agora envie a foto que deseja transformar em Porta-Chaves.*\n\n",
                'dica': "💡 *Dica:* Envie uma foto com boa iluminação e foco no rosto/objeto.",
                'skip_info': "\n\n*Nota:* Não é possível pular esta etapa - é necessária uma foto!"
            },
            'ingles': {
                'titulo': "📸 *Perfect! Now send the photo you want to transform into a Keychain.*\n\n",
                'dica': "💡 *Tip:* Send a photo with good lighting and focus on the face/object.",
                'skip_info': "\n\n*Note:* Cannot skip this step - a photo is required!"
            },
            'espanhol': {
                'titulo': "📸 *¡Perfecto! Ahora envíe la foto que desea transformar em Llavero.*\n\n",
                'dica': "💡 *Consejo:* Envíe una foto con buena iluminación y enfoque en el rostro/objeto.",
                'skip_info': "\n\n*Nota:* ¡No se puede saltar este paso - se requiere una foto!"
            },
            'italiano': {
                'titulo': "📸 *Perfetto! Ora invia la foto che desideri trasformare in Portachiavi.*\n\n",
                'dica': "💡 *Suggerimento:* Invia una foto con una buona illuminación e messa a fuoco sul viso/oggetto.",
                'skip_info': "\n\n*Nota:* Non puoi saltare questo passaggio - è necessaria una foto!"
            },
            'alemao': {
                'titulo': "📸 *Perfekt! Jetzt senden Sie das Foto, das Sie in einen Schlüsselanhänger verwandeln möchten.*\n\n",
                'dica': "💡 *Tipp:* Senden Sie ein Foto mit guter Beleuchtung e foco sul viso/oggetto.",
                'skip_info': "\n\n*Hinweis:* Dieser Schritt kann nicht übersprungen werden - ein Foto ist erforderlich!"
            },
            'frances': {
                'titulo': "📸 *Parfait ! Maintenant, envoyez la photo que vous souhaitez transformer em Porte-clés.*\n\n",
                'dica': "💡 *Astuce:* Envoyez une photo avec un bon éclairage et une mise au point sur le visage/l'objet.",
                'skip_info': "\n\n*Remarque:* Impossible de sauter cette étape - une photo est requise !"
            }
        }
        
        textos = textos_foto.get(idioma, textos_foto['portugues'])
        
        # 🔥 AGORA PEDIR A FOTO
        await update.message.reply_text(
            f"{textos['titulo']}"
            f"{textos['dica']}"
            f"{textos['skip_info']}",
            parse_mode="Markdown"
        )
        
        context.user_data['conversation_state'] = GIFT_FOTO
        print(f"✅ Estado atualizado para GIFT_FOTO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"❌ ERRO em receber_gift_frase_box: {e}")
        import traceback
        traceback.print_exc()
        
        # 🔥 TEXTOS POR IDIOMA PARA MENSAGEM DE ERRO
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtelo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )





# --- Enviar resumo do gift ---
async def enviar_resumo_gift(context, chat_id):
    """Enviar resumo específico para gift - APENAS UM RESUMO POR VEZ COM TRADUÇÃO"""
    try:
        # 🔥 OBTER IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"🌐 Enviando resumo GIFT em: {idioma}")
        
        # 🔥 DICIONÁRIO PARA CONVERTER PAÍSES PARA INGLÊS
        PAISES_PARA_INGLES = {
            'estados_unidos': 'United States',
            'canada': 'Canada',
            'reino_unido': 'United Kingdom',
            'brasil': 'Brazil',
            'alemanha': 'Germany',
            'paises_baixos': 'Netherlands',
            'holanda': 'Netherlands',
            'franca': 'France',
            'espanha': 'Spain',
            'belgica': 'Belgium',
            'italia': 'Italy',
            'portugal': 'Portugal',
            'irlanda': 'Ireland',
            'luxemburgo': 'Luxembourg'
        }
        
        def converter_pais_para_ingles(pais_key):
            """Converte o nome/callback do país para inglês"""
            if isinstance(pais_key, str):
                # Remove "pais_" se existir
                if pais_key.startswith('pais_'):
                    pais_key = pais_key[5:]
                # Remove acentos e converte para minúsculas para comparação
                pais_clean = pais_key.lower()
                # Mapeamento adicional para nomes em português
                mapeamento = {
                    'bélgica': 'belgica',
                    'bélgica (português)': 'belgica',
                    'frança': 'franca',
                    'espanha': 'espanha',
                    'alemanha': 'alemanha',
                    'itália': 'italia',
                    'irlanda': 'irlanda',
                    'luxemburgo': 'luxemburgo',
                    'países baixos': 'paises_baixos',
                    'holanda': 'paises_baixos',
                    'reino unido': 'reino_unido',
                    'estados unidos': 'estados_unidos',
                    'eua': 'estados_unidos'
                }
                pais_key = mapeamento.get(pais_clean, pais_key)
            return PAISES_PARA_INGLES.get(pais_key, pais_key.title())
        
        # 🔥 TEXTOS POR IDIOMA PARA OS CAMPOS
        textos_campos = {
            'portugues': {
                'titulo': "🎁 *RESUMO DO SEU PORTA-CHAVES*\n\n",
                'nome': "👤 *Nome:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *País:*",
                'telefone': "📱 *Telefone:*",
                'nome_gift': "🎭 *Nome na Box:*",
                'frase_gift': "💬 *Frase na Box:*",
                'produto': "🎁 *Produto:* Porta-Chaves Personalizado\n",
                'continue': "\n⬇️ *Continue preenchendo abaixo* ⬇️",
                'sem_frase': "Não adicionou frase"
            },
            'ingles': {
                'titulo': "🎁 *YOUR KEYCHAIN SUMMARY*\n\n",
                'nome': "👤 *Name:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Country:*",
                'telefone': "📱 *Phone:*",
                'nome_gift': "🎭 *Name on Box:*",
                'frase_gift': "💬 *Phrase on Box:*",
                'produto': "🎁 *Product:* Personalized Keychain\n",
                'continue': "\n⬇️ *Continue filling below* ⬇️",
                'sem_frase': "No phrase added"
            },
            'espanhol': {
                'titulo': "🎁 *RESUMEN DE SU LLAVERO*\n\n",
                'nome': "👤 *Nombre:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *País:*",
                'telefone': "📱 *Teléfono:*",
                'nome_gift': "🎭 *Nombre en Caja:*",
                'frase_gift': "💬 *Frase en Caja:*",
                'produto': "🎁 *Producto:* Llavero Personalizado\n",
                'continue': "\n⬇️ *Continúe rellenando abajo* ⬇️",
                'sem_frase': "No añadió frase"
            },
            'italiano': {
                'titulo': "🎁 *RIEPILOGO DEL TUO PORTACHIAVI*\n\n",
                'nome': "👤 *Nome:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Paese:*",
                'telefone': "📱 *Telefono:*",
                'nome_gift': "🎭 *Nome su Scatola:*",
                'frase_gift': "💬 *Frase su Scatola:*",
                'produto': "🎁 *Prodotto:* Portachiavi Personalizzato\n",
                'continue': "\n⬇️ *Continua a compilare qui sotto* ⬇️",
                'sem_frase': "Nessuna frase aggiunta"
            },
            'alemao': {
                'titulo': "🎁 *ZUSAMMENFASSUNG IHRES SCHLÜSSELANHÄNGERS*\n\n",
                'nome': "👤 *Name:*",
                'email': "📧 *E-Mail:*",
                'pais': "🌍 *Land:*",
                'telefone': "📱 *Telefon:*",
                'nome_gift': "🎭 *Name auf Box:*",
                'frase_gift': "💬 *Phrase auf Box:*",
                'produto': "🎁 *Produkt:* Personalisierter Schlüsselanhänger\n",
                'continue': "\n⬇️ *Fahren Sie unten fort* ⬇️",
                'sem_frase': "Keine Phrase hinzugefügt"
            },
            'frances': {
                'titulo': "🎁 *RÉSUMÉ DE VOTRE PORTE-CLÉS*\n\n",
                'nome': "👤 *Nom:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Pays:*",
                'telefone': "📱 *Téléphone:*",
                'nome_gift': "🎭 *Nom sur Boîte:*",
                'frase_gift': "💬 *Phrase sur Boîte:*",
                'produto': "🎁 *Produit:* Porte-clés Personnalisé\n",
                'continue': "\n⬇️ *Continuez à remplir ci-dessous* ⬇️",
                'sem_frase': "Aucune phrase ajoutée"
            }
        }
        
        textos = textos_campos.get(idioma, textos_campos['portugues'])
        user_data = context.user_data
        
        # 🔥 PRIMEIRO APAGAR QUALQUER RESUMO EXISTENTE
        resumo_antigo_id = context.user_data.get('resumo_msg_id')
        if resumo_antigo_id:
            try:
                await context.bot.delete_message(
                    chat_id=chat_id,
                    message_id=resumo_antigo_id
                )
                print(f"🗑️ Resumo anterior apagado: {resumo_antigo_id} | Idioma: {idioma}")
            except Exception as e:
                print(f"⚠️ Não foi possível apagar resumo anterior: {e}")

        # 🔥 CONSTRUIR NOVO RESUMO
        resumo = f"{textos['titulo']}"
        
        if "nome" in user_data:
            resumo += f"{textos['nome']} {user_data['nome']}\n"
        if "email" in user_data:
            resumo += f"{textos['email']} {user_data['email']}\n"
        if "pais" in user_data:
            # 🔥 CONVERTER PAÍS PARA INGLÊS
            pais_original = user_data['pais']
            pais_ingles = converter_pais_para_ingles(pais_original)
            resumo += f"{textos['pais']} {pais_ingles}\n"  # Mostra em inglês
        if "contacto" in user_data:
            resumo += f"{textos['telefone']} {user_data['contacto']}\n"
        if "nome_gift" in user_data:
            resumo += f"{textos['nome_gift']} {user_data['nome_gift']}\n"
        if "frase_gift" in user_data and user_data['frase_gift'] != textos['sem_frase']:
            resumo += f"{textos['frase_gift']} \"{user_data['frase_gift']}\"\n"
        
        resumo += f"{textos['produto']}"
        resumo += f"{textos['continue']}"
        
        # 🔥 ENVIAR NOVO RESUMO
        msg = await context.bot.send_message(
            chat_id=chat_id, 
            text=resumo, 
            parse_mode="Markdown"
        )
        
        print(f"✅ Novo resumo enviado com ID: {msg.message_id} | Idioma: {idioma}")
        print(f"✅ País mostrado como: {pais_ingles if 'pais' in user_data else 'N/A'} (em inglês)")
        return msg.message_id
        
    except Exception as e:
        print(f"❌ Erro em enviar_resumo_gift: {e}")
        import traceback
        traceback.print_exc()
        return None



# --- Receber foto do gift ---
async def receber_gift_foto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber foto do gift com tradução completa e países em inglês"""
    print("📸 DEBUG: receber_gift_foto foi chamado!")
    
    # 🔥 OBTER IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Processando foto GIFT em: {idioma}")
    
    # 🔥 DICIONÁRIO PARA CONVERTER PAÍSES PARA INGLÊS
    PAISES_PARA_INGLES = {
        'estados_unidos': 'United States',
        'canada': 'Canada',
        'reino_unido': 'United Kingdom',
        'brasil': 'Brazil',
        'alemanha': 'Germany',
        'paises_baixos': 'Netherlands',
        'holanda': 'Netherlands',
        'franca': 'France',
        'espanha': 'Spain',
        'belgica': 'Belgium',
        'italia': 'Italy',
        'portugal': 'Portugal',
        'irlanda': 'Ireland',
        'luxemburgo': 'Luxembourg'
    }
    
    def converter_pais_para_ingles(pais_key):
        """Converte o nome/callback do país para inglês"""
        if isinstance(pais_key, str):
            # Remove "pais_" se existir
            if pais_key.startswith('pais_'):
                pais_key = pais_key[5:]
            # Remove acentos e converte para minúsculas para comparação
            pais_clean = pais_key.lower()
            # Mapeamento adicional para nomes em português
            mapeamento = {
                'bélgica': 'belgica',
                'bélgica (português)': 'belgica',
                'frança': 'franca',
                'espanha': 'espanha',
                'alemanha': 'alemanha',
                'itália': 'italia',
                'irlanda': 'irlanda',
                'luxemburgo': 'luxemburgo',
                'países baixos': 'paises_baixos',
                'holanda': 'paises_baixos',
                'reino unido': 'reino_unido',
                'estados unidos': 'estados_unidos',
                'eua': 'estados_unidos'
            }
            pais_key = mapeamento.get(pais_clean, pais_key)
        return PAISES_PARA_INGLES.get(pais_key, pais_key.title())
    
    # Verificar se já temos foto
    if "foto_id" in context.user_data and context.user_data.get('conversation_state') != GIFT_FOTO:
        # 🔥 TEXTOS POR IDIOMA PARA FOTO JÁ EXISTENTE
        textos_foto_existente = {
            'portugues': "⚠️ Já recebemos a sua foto. Use o botão 'Mudar Foto' se quiser alterar.",
            'ingles': "⚠️ We already received your photo. Use the 'Change Photo' button if you want to change it.",
            'espanhol': "⚠️ Ya recibimos su foto. Use el botón 'Cambiar Foto' si desea cambiarla.",
            'italiano': "⚠️ Abbiamo già ricevuto la tua foto. Usa il pulsante 'Cambia Foto' se vuoi cambiarla.",
            'alemao': "⚠️ Wir haben Ihr Foto bereits erhalten. Verwenden Sie die Schaltfläche 'Foto ändern', wenn Sie es ändern möchten.",
            'frances': "⚠️ Nous avons déjà reçu votre photo. Utilisez le bouton 'Changer la Photo' si vous souhaitez la modifier."
        }
        
        print("📸 DEBUG: Foto gift já existe")
        await update.message.reply_text(
            textos_foto_existente.get(idioma, textos_foto_existente['portugues'])
        )
        return

    print("📸 DEBUG: Processando nova foto gift...")
    
    photo = update.message.photo[-1]
    file_id = photo.file_id
    
    # GERAR NOME DA FOTO
    nome_usuario = context.user_data.get('nome', 'Cliente')
    timestamp = datetime.now().strftime("%H%M%S")
    nome_foto = f"gift_{timestamp}.jpg"
    
    context.user_data["foto_id"] = file_id
    context.user_data["nome_foto"] = nome_foto

    # 🔥 TEXTOS POR IDIOMA PARA CONFIRMAÇÃO DE FOTO
    textos_confirmacao_foto = {
        'portugues': "📸 Foto recebida com sucesso!",
        'ingles': "📸 Photo received successfully!",
        'espanhol': "📸 ¡Foto recibida con éxito!",
        'italiano': "📸 Foto ricevuta con successo!",
        'alemao': "📸 Foto erfolgreich empfangen!",
        'frances': "📸 Photo reçue avec succès !"
    }
    
    await update.message.reply_text(
        textos_confirmacao_foto.get(idioma, textos_confirmacao_foto['portugues'])
    )

    # 🔥 TEXTOS POR IDIOMA PARA RESUMO FINAL
    textos_resumo_final = {
        'portugues': {
            'titulo': "✅ *Resumo Final do Porta-Chaves:*\n\n",
            'nome': "👤 *Nome:*",
            'email': "📧 *Email:*",
            'pais': "🌍 *País:*",
            'telefone': "📱 *Telefone:*",
            'nome_gift': "🎭 *Nome na Box:*",
            'frase_gift': "💬 *Frase na Box:*",
            'produto': "🎁 *Produto:* Porta-Chaves Personalizado\n",
            'tamanho': "📏 *Tamanho:* 2.5\" | 6.4cm\n",
            'foto': "📸 *Foto:* recebida ✅",
            'final': "\n🚀 *Estamos prontos para criar o seu Porta-Chaves personalizado!*",
            'sem_frase': "Não adicionou frase"
        },
        'ingles': {
            'titulo': "✅ *Final Keychain Summary:*\n\n",
            'nome': "👤 *Name:*",
            'email': "📧 *Email:*",
            'pais': "🌍 *Country:*",
            'telefone': "📱 *Phone:*",
            'nome_gift': "🎭 *Name on Box:*",
            'frase_gift': "💬 *Phrase on Box:*",
            'produto': "🎁 *Product:* Personalized Keychain\n",
            'tamanho': "📏 *Size:* 2.5\" | 6.4cm\n",
            'foto': "📸 *Photo:* received ✅",
            'final': "\n🚀 *We are ready to create your personalized Keychain!*",
            'sem_frase': "No phrase added"
        },
        'espanhol': {
            'titulo': "✅ *Resumen Final del Llavero:*\n\n",
            'nome': "👤 *Nombre:*",
            'email': "📧 *Email:*",
            'pais': "🌍 *País:*",
            'telefone': "📱 *Teléfono:*",
            'nome_gift': "🎭 *Nombre en Caja:*",
            'frase_gift': "💬 *Frase en Caja:*",
            'produto': "🎁 *Producto:* Llavero Personalizado\n",
            'tamanho': "📏 *Tamaño:* 2.5\" | 6.4cm\n",
            'foto': "📸 *Foto:* recibida ✅",
            'final': "\n🚀 *¡Estamos listos para crear su Llavero personalizado!*",
            'sem_frase': "No añadió frase"
        },
        'italiano': {
            'titulo': "✅ *Riepilogo Finale del Portachiavi:*\n\n",
            'nome': "👤 *Nome:*",
            'email': "📧 *Email:*",
            'pais': "🌍 *Paese:*",
            'telefone': "📱 *Telefono:*",
            'nome_gift': "🎭 *Nome su Scatola:*",
            'frase_gift': "💬 *Frase su Scatola:*",
            'produto': "🎁 *Prodotto:* Portachiavi Personalizzato\n",
            'tamanho': "📏 *Dimensione:* 2.5\" | 6.4cm\n",
            'foto': "📸 *Foto:* ricevuta ✅",
            'final': "\n🚀 *Siamo pronti a creare il tuo Portachiavi personalizzato!*",
            'sem_frase': "Nessuna frase aggiunta"
        },
        'alemao': {
            'titulo': "✅ *Endzusammenfassung Schlüsselanhänger:*\n\n",
            'nome': "👤 *Name:*",
            'email': "📧 *E-Mail:*",
            'pais': "🌍 *Land:*",
            'telefone': "📱 *Telefon:*",
            'nome_gift': "🎭 *Name auf Box:*",
            'frase_gift': "💬 *Phrase auf Box:*",
            'produto': "🎁 *Produkt:* Personalisierter Schlüsselanhänger\n",
            'tamanho': "📏 *Größe:* 2.5\" | 6.4cm\n",
            'foto': "📸 *Foto:* empfangen ✅",
            'final': "\n🚀 *Wir sind bereit, Ihren personalisierten Schlüsselanhänger zu erstellen!*",
            'sem_frase': "Keine Phrase hinzugefügt"
        },
        'frances': {
            'titulo': "✅ *Résumé Final du Porte-clés:*\n\n",
            'nome': "👤 *Nom:*",
            'email': "📧 *Email:*",
            'pais': "🌍 *Pays:*",
            'telefone': "📱 *Téléphone:*",
            'nome_gift': "🎭 *Nom sur Boîte:*",
            'frase_gift': "💬 *Phrase sur Boîte:*",
            'produto': "🎁 *Produit:* Porte-clés Personnalisé\n",
            'tamanho': "📏 *Taille:* 2.5\" | 6.4cm\n",
            'foto': "📸 *Photo:* reçue ✅",
            'final': "\n🚀 *Nous sommes prêts à créer votre Porte-clés personnalisé !*",
            'sem_frase': "Aucune phrase ajoutée"
        }
    }
    
    textos = textos_resumo_final.get(idioma, textos_resumo_final['portugues'])
    
    # 🔥 CONSTRUIR RESUMO FINAL
    nome_foto_resumo = context.user_data.get('nome_foto', 'foto.jpg')
    nome_gift = context.user_data.get('nome_gift', '')
    frase_gift = context.user_data.get('frase_gift', textos['sem_frase'])
    
    resumo = f"{textos['titulo']}"
    resumo += f"{textos['nome']} {context.user_data.get('nome', '')}\n"
    resumo += f"{textos['email']} {context.user_data.get('email', '')}\n"
    
    # 🔥 CONVERTER PAÍS PARA INGLÊS
    if "pais" in context.user_data:
        pais_original = context.user_data['pais']
        pais_ingles = converter_pais_para_ingles(pais_original)
        resumo += f"{textos['pais']} {pais_ingles}\n"  # 🔥 MOSTRAR EM INGLÊS
    else:
        resumo += f"{textos['pais']} \n"
    
    resumo += f"{textos['telefone']} {context.user_data.get('contacto', '')}\n"
    
    # 🔥 ADICIONAR CAMPOS DE PERSONALIZAÇÃO SE EXISTIREM
    if nome_gift and nome_gift != textos['sem_frase'].replace("frase", "nome/alcunha"):
        resumo += f"{textos['nome_gift']} {nome_gift}\n"
    
    if frase_gift and frase_gift != textos['sem_frase']:
        resumo += f"{textos['frase_gift']} \"{frase_gift}\"\n"
    
    resumo += f"{textos['produto']}"
    resumo += f"{textos['tamanho']}"
    resumo += f"{textos['foto']} (*{nome_foto_resumo}*)\n"
    resumo += f"{textos['final']}"
    
    # 🔥 TEXTOS POR IDIOMA PARA OS BOTÕES
    textos_botoes = {
        'portugues': {
            'finalizar': "💳 Finalizar Compra",
            'mudar_foto': "📸 Enganei-me na foto (mudar)",
            'voltar': "↩️ Voltar ao Início"
        },
        'ingles': {
            'finalizar': "💳 Finalize Purchase",
            'mudar_foto': "📸 Wrong photo (change)",
            'voltar': "↩️ Return to Start"
        },
        'espanhol': {
            'finalizar': "💳 Finalizar Compra",
            'mudar_foto': "📸 Me equivoqué en la foto (cambiar)",
            'voltar': "↩️ Volver al Inicio"
        },
        'italiano': {
            'finalizar': "💳 Finalizza Acquisto",
            'mudar_foto': "📸 Foto sbagliata (cambia)",
            'voltar': "↩️ Torna all'Inizio"
        },
        'alemao': {
            'finalizar': "💳 Kauf abschließen",
            'mudar_foto': "📸 Falsches Foto (ändern)",
            'voltar': "↩️ Zum Anfang zurück"
        },
        'frances': {
            'finalizar': "💳 Finaliser l'Achat",
            'mudar_foto': "📸 Mauvaise photo (changer)",
            'voltar': "↩️ Retour au Début"
        }
    }
    
    botoes_traduzidos = textos_botoes.get(idioma, textos_botoes['portugues'])
    
    botoes = [
        [InlineKeyboardButton(botoes_traduzidos['finalizar'], callback_data="finalizar_gift")],
        [InlineKeyboardButton(botoes_traduzidos['mudar_foto'], callback_data="mudar_gift_foto")],
        [InlineKeyboardButton(botoes_traduzidos['voltar'], callback_data="voltar_inicio")]
    ]
    
    await update.message.reply_text(
        resumo, 
        parse_mode="Markdown", 
        reply_markup=InlineKeyboardMarkup(botoes)
    )
    
    context.user_data['conversation_state'] = GIFT_FIM
    print(f"✅ Estado atualizado para GIFT_FIM | Idioma: {idioma}")
    print(f"✅ País mostrado como: {pais_ingles if 'pais' in context.user_data else 'N/A'} (em inglês)")









# --- Mudar foto do gift ---
async def mudar_gift_foto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para mudar a foto do gift COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 OBTER IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Mudando foto GIFT em: {idioma}")
    
    # Remover a foto atual
    context.user_data.pop("foto_id", None)
    context.user_data.pop("nome_foto", None)
    
    # 🔥 TEXTOS POR IDIOMA PARA CONFIRMAÇÃO DE REMOÇÃO
    textos_remocao = {
        'portugues': "🔄 Foto anterior removida. Pronto para receber nova foto!",
        'ingles': "🔄 Previous photo removed. Ready to receive new photo!",
        'espanhol': "🔄 Foto anterior eliminada. ¡Listo para recibir nueva foto!",
        'italiano': "🔄 Foto precedente rimossa. Pronto a ricevere nuova foto!",
        'alemao': "🔄 Vorheriges Foto entfernt. Bereit für neues Foto!",
        'frances': "🔄 Photo précédente supprimée. Prêt à recevoir une nouvelle photo !"
    }
    
    print(f"✅ Foto GIFT anterior removida | Idioma: {idioma}")
    
    # Remover a mensagem com os botões antigos
    await safe_delete_message(query)
    
    # 🔥 TEXTOS POR IDIOMA PARA PEDIDO DE NOVA FOTO
    textos_nova_foto = {
        'portugues': "📸 *Por favor, envie a nova foto para o Porta-Chaves:*\n\n💡 *Dica:* Uma foto nítida e bem iluminada garante melhor resultado!",
        'ingles': "📸 *Please, send the new photo for the Keychain:*\n\n💡 *Tip:* A sharp and well-lit photo ensures better results!",
        'espanhol': "📸 *Por favor, envíe la nueva foto para el Llavero:*\n\n💡 *Consejo:* ¡Una foto nítida y bien iluminada asegura mejores resultados!",
        'italiano': "📸 *Per favore, invia la nuova foto per il Portachiavi:*\n\n💡 *Suggerimento:* Una foto nitida e ben illuminata garantisce risultati migliori!",
        'alemao': "📸 *Bitte senden Sie das neue Foto für den Schlüsselanhänger:*\n\n💡 *Tipp:* Ein scharfes und gut beleuchtetes Foto sorgt für bessere Ergebnisse!",
        'frances': "📸 *Veuillez envoyer la nouvelle photo pour le Porte-clés:*\n\n💡 *Astuce:* Une photo nette et bien éclairée assure de meilleurs résultats !"
    }
    
    # Enviar confirmação de remoção
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_remocao.get(idioma, textos_remocao['portugues']),
        parse_mode="Markdown"
    )
    
    # Pedir nova foto
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_nova_foto.get(idioma, textos_nova_foto['portugues']),
        parse_mode="Markdown"
    )
    
    context.user_data['conversation_state'] = GIFT_FOTO
    print(f"✅ Estado redefinido para GIFT_FOTO | Idioma: {idioma}")








    # --- Handler para iniciar gift pelo botão ---
async def iniciar_gift_botao(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Iniciar fluxo gift a partir do botão do menu COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    chat_id = query.message.chat_id
    user_id = update.effective_user.id
    
    print(f"🎁 GIFT INICIADO VIA BOTÃO por user {user_id}")
    
    # 🔥 OBTER IDIOMA DO USER_DATA OU DO CONTEXT
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Iniciando GIFT via botão em: {idioma}")
    
    # 🔥 LIMPAR DADOS E CANCELAR TEMPORIZADORES
    try:
        cancelar_temporizador_30min(chat_id)
    except:
        pass
    
    context.user_data.clear()
    
    # 🔥 GUARDAR IDIOMA NOVAMENTE APÓS LIMPAR
    context.user_data['idioma'] = idioma
    
    # 🔥 DEFINIR TIPO COMO GIFT
    context.user_data["tipo_cartoon"] = {
        'portugues': "Porta-Chaves 🎁",
        'ingles': "Keychain 🎁",
        'espanhol': "Llavero 🎁",
        'italiano': "Portachiavi 🎁",
        'alemao': "Schlüsselanhänger 🎁",
        'frances': "Porte-clés 🎁"
    }.get(idioma, "Porta-Chaves 🎁")
    
    context.user_data["estilo_cartoon"] = "Gift"
    context.user_data["tamanho_cartoon"] = {
        'portugues': "Porta-Chaves Padrão",
        'ingles': "Standard Keychain",
        'espanhol': "Llavero Estándar",
        'italiano': "Portachiavi Standard",
        'alemao': "Standard Schlüsselanhänger",
        'frances': "Porte-clés Standard"
    }.get(idioma, "Porta-Chaves Padrão")
    
    context.user_data["tamanho_key"] = "padrao"
    context.user_data["preco_tamanho"] = 0.0
    context.user_data['oferta_tipo'] = 'oferta_surpresa'
    context.user_data['_gift_oferta_surpresa'] = True
    
    print(f"✅ Dados GIFT inicializados | Idioma: {idioma}")
    print(f"🎯 OFERTA_TIPO DEFINIDO: {context.user_data['oferta_tipo']}")
    
    # 🔥 APAGAR MENSAGEM DO MENU
    try:
        await query.delete_message()
        print("✅ Mensagem do menu apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem do menu: {e}")
    
    # 🔥 TEXTOS POR IDIOMA PARA PRIMEIRA PERGUNTA
    textos_inicio = {
        'portugues': "🎁 *Vamos criar um Porta-Chaves personalizado!*\n\nAntes de começarmos, qual é o seu nome?",
        'ingles': "🎁 *Let's create a personalized Keychain!*\n\nBefore we start, what's your name?",
        'espanhol': "🎁 *¡Vamos a crear un Llavero personalizado!*\n\nAntes de empezar, ¿cuál es su nombre?",
        'italiano': "🎁 *Creiamo un Portachiavi personalizzato!*\n\nPrima di iniziare, qual è il tuo nome?",
        'alemao': "🎁 *Erstellen wir einen personalisierten Schlüsselanhänger!*\n\nBevor wir beginnen, wie ist Ihr Name?",
        'frances': "🎁 *Créons un Porte-clés personnalisé !*\n\nAvant de commencer, quel est votre nom ?"
    }
    
    # Primeira pergunta (nome)
    await context.bot.send_message(
        chat_id=chat_id,
        text=textos_inicio.get(idioma, textos_inicio['portugues']),
        parse_mode="Markdown"
    )
    
    context.user_data['conversation_state'] = GIFT_NOME
    print(f"✅ Estado definido para GIFT_NOME | Idioma: {idioma}")





    


async def finalizar_gift(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Finalizar gift com tradução completa e países em inglês"""
    query = update.callback_query
    await query.answer()
    
    print("🔄 FINALIZAR_GIFT INICIADO")
    print(f"🔍 DEBUG - Chat ID: {query.message.chat_id}")
    print(f"🔍 DEBUG - User Data keys: {list(context.user_data.keys())}")
    
    # 🔥 OBTER IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Finalizando GIFT em: {idioma}")
    
    # 🔥 DICIONÁRIO PARA CONVERTER PAÍSES PARA INGLÊS
    PAISES_PARA_INGLES = {
        'estados_unidos': 'United States',
        'canada': 'Canada',
        'reino_unido': 'United Kingdom',
        'brasil': 'Brazil',
        'alemanha': 'Germany',
        'paises_baixos': 'Netherlands',
        'holanda': 'Netherlands',
        'franca': 'France',
        'espanha': 'Spain',
        'belgica': 'Belgium',
        'italia': 'Italy',
        'portugal': 'Portugal',
        'irlanda': 'Ireland',
        'luxemburgo': 'Luxembourg'
    }
    
    def converter_pais_para_ingles(pais_key):
        """Converte o nome/callback do país para inglês"""
        if isinstance(pais_key, str):
            # Remove "pais_" se existir
            if pais_key.startswith('pais_'):
                pais_key = pais_key[5:]
            # Remove acentos e converte para minúsculas para comparação
            pais_clean = pais_key.lower()
            # Mapeamento adicional para nomes em português
            mapeamento = {
                'bélgica': 'belgica',
                'bélgica (português)': 'belgica',
                'frança': 'franca',
                'espanha': 'espanha',
                'alemanha': 'alemanha',
                'itália': 'italia',
                'irlanda': 'irlanda',
                'luxemburgo': 'luxemburgo',
                'países baixos': 'paises_baixos',
                'holanda': 'paises_baixos',
                'reino unido': 'reino_unido',
                'estados unidos': 'estados_unidos',
                'eua': 'estados_unidos'
            }
            pais_key = mapeamento.get(pais_clean, pais_key)
        return PAISES_PARA_INGLES.get(pais_key, pais_key.title())
    
    # 🔥 CORREÇÃO CRÍTICA: VERIFICAR E MANTER O TIPO DE OFERTA CORRETO
    # 1. PRIMEIRO: Verificar se é um gift com oferta_surpresa
    if context.user_data.get('_gift_oferta_surpresa'):
        oferta_tipo_final = 'oferta_surpresa'
        print(f"🎯 GIFT DETECTADO - OFERTA_TIPO FORÇADO: {oferta_tipo_final}")
    # 2. SEGUNDO: Usar o oferta_tipo existente no context
    elif context.user_data.get('oferta_tipo'):
        oferta_tipo_final = context.user_data['oferta_tipo']
        print(f"🎯 OFERTA_TIPO EXISTENTE: {oferta_tipo_final}")
    # 3. TERCEIRO: Default para gift é oferta_surpresa
    else:
        oferta_tipo_final = 'oferta_surpresa'
        print(f"🎯 OFERTA_TIPO PADRÃO PARA GIFT: {oferta_tipo_final}")
    
    # Verificar se já existe pedido anterior
    if "pedido_id" in context.user_data:
        old_pedido_id = context.user_data["pedido_id"]
        print(f"🚨 ATENÇÃO: JÁ EXISTE pedido_id NO user_data: {old_pedido_id}")
        
        if old_pedido_id in PEDIDOS_REGISTO:
            status_antigo = PEDIDOS_REGISTO[old_pedido_id]["status"]
            print(f"🔍 Pedido anterior #{old_pedido_id} ainda no registo - Status: {status_antigo}")
            
            if status_antigo == "pendente":
                del PEDIDOS_REGISTO[old_pedido_id]
                print(f"🗑️ Pedido anterior #{old_pedido_id} removido do registo")
        
        del context.user_data["pedido_id"]
        print(f"✅ Pedido_id anterior #{old_pedido_id} removido do user_data")
    
    # Remover a mensagem anterior com botões
    await safe_delete_message(query)
    
    # 🔥 CALCULAR TOTAIS PARA GIFT
    pais_original = context.user_data.get("pais", "portugal")
    pais_ingles = converter_pais_para_ingles(pais_original)
    print(f"🌍 País original: {pais_original} -> Inglês: {pais_ingles}")
    totais = calcular_total_por_moeda(context, pais_original)
    
    # 🔥 OBTER TODOS OS DADOS
    nome = context.user_data.get("nome", "")
    email = context.user_data.get("email", "")
    contacto = context.user_data.get("contacto", "")
    nome_foto = context.user_data.get("nome_foto", "foto.jpg")
    nome_gift = context.user_data.get("nome_gift", "")
    frase_gift = context.user_data.get("frase_gift", "Não adicionou frase")

    print(f"🔍 DEBUG FINALIZAR_GIFT - Dados a guardar:")
    print(f"   • Nome: {nome}")
    print(f"   • Email: {email}")
    print(f"   • Contacto: {contacto}")
    print(f"   • País original: {pais_original}")
    print(f"   • País em inglês: {pais_ingles}")
    print(f"   • Nome Gift: {nome_gift}")
    print(f"   • Frase Gift: {frase_gift}")
    print(f"   • Oferta Tipo Final: {oferta_tipo_final}")
    print(f"   • Idioma: {idioma}")
    
    foto_recebida = "✅" if "foto_id" in context.user_data else "❌"

    # GERAR ID ÚNICO DO PEDIDO
    pedido_id = str(uuid.uuid4())[:8].upper()
    data_pedido = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    
    if pedido_id in PEDIDOS_REGISTO:
        print(f"🔄 CONFLITO: ID {pedido_id} já existe, gerando novo...")
        pedido_id = str(uuid.uuid4())[:8].upper()
        print(f"🆔 NOVO ID gerado: {pedido_id}")
    
    # 🔥 TEXTOS POR IDIOMA PARA "NÃO ADICIONOU FRASE"
    textos_sem_frase = {
        'portugues': "Não adicionou frase",
        'ingles': "No phrase added",
        'espanhol': "No añadió frase",
        'italiano': "Nessuna frase aggiunta",
        'alemao': "Keine Phrase hinzugefügt",
        'frances': "Aucune phrase ajoutée"
    }
    
    # Ajustar frase para o idioma correto
    if frase_gift == "Não adicionou frase":
        frase_gift = textos_sem_frase.get(idioma, textos_sem_frase['portugues'])
    
    # 🔥 TEXTOS POR IDIOMA PARA NOMES DE PRODUTOS
    textos_produtos = {
        'portugues': {
            'tipo_cartoon': "Porta-Chaves 🎁",
            'tamanho_cartoon': ' 2.5" | 6.4cm ',
            'produto': "Porta-Chaves Personalizado",
            'tamanho': "2.5\" | 6.4cm"
        },
        'ingles': {
            'tipo_cartoon': "Keychain 🎁",
            'tamanho_cartoon': ' 2.5" | 6.4cm ',
            'produto': "Personalized Keychain",
            'tamanho': "2.5\" | 6.4cm"
        },
        'espanhol': {
            'tipo_cartoon': "Llavero 🎁",
            'tamanho_cartoon': ' 2.5" | 6.4cm ',
            'produto': "Llavero Personalizado",
            'tamanho': "2.5\" | 6.4cm"
        },
        'italiano': {
            'tipo_cartoon': "Portachiavi 🎁",
            'tamanho_cartoon': ' 2.5" | 6.4cm ',
            'produto': "Portachiavi Personalizzato",
            'tamanho': "2.5\" | 6.4cm"
        },
        'alemao': {
            'tipo_cartoon': "Schlüsselanhänger 🎁",
            'tamanho_cartoon': ' 2.5" | 6.4cm ',
            'produto': "Personalisierter Schlüsselanhänger",
            'tamanho': "2.5\" | 6.4cm"
        },
        'frances': {
            'tipo_cartoon': "Porte-clés 🎁",
            'tamanho_cartoon': ' 2.5" | 6.4cm ',
            'produto': "Porte-clés Personnalisé",
            'tamanho': "2.5\" | 6.4cm"
        }
    }
    
    textos_prod = textos_produtos.get(idioma, textos_produtos['portugues'])
    
    # 🔥 🔥 🔥 CORREÇÃO: GUARDAR NO PEDIDOS_REGISTO PARA O TEMPORIZADOR FUNCIONAR
    # Mas ainda NÃO contar nas estatísticas - só quando pagar
    
    # Criar o objeto do pedido GIFT
    pedido_data = {
        "id": pedido_id,
        "data": data_pedido,
        "nome": nome,
        "email": email,
        "pais": pais_ingles,  # 🔥 GUARDAR EM INGLÊS NO PEDIDO
        "pais_original": pais_original,  # Guardar original também
        "contacto": contacto,
        "tipo_cartoon": textos_prod['tipo_cartoon'],
        "tamanho_cartoon": textos_prod['tamanho_cartoon'],
        "nome_foto": nome_foto,
        "foto_id": context.user_data.get("foto_id"),
        "nome_gift": nome_gift,
        "frase_gift": frase_gift,
        "idioma": idioma,  # 🔥 GUARDAR IDIOMA NO PEDIDO
        
        # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: USAR O TIPO DE OFERTA CORRETO
        "oferta_tipo": oferta_tipo_final,  # "oferta_surpresa", "pagamento_direto", "original", "tamanho_4.5", "portachaves"
        
        # 🔥 INFORMAÇÕES DE MOEDA
        "subtotal": totais['subtotal'],
        "imposto": totais['imposto'],
        "frete": totais['frete'],
        "total": totais['total'],
        "valor_original_real": totais['total'],
        "moeda": totais['moeda'],
        "simbolo_moeda": totais['simbolo_moeda'],
        "chat_id": query.message.chat_id,
        "status": "pendente",
        "data_expiracao": datetime.now() + timedelta(minutes=10),
        "tentativas_recuperacao": 0,
        "produto_tipo": "portachaves"  # 🔥 ADICIONADO: Tipo do produto (para estatísticas separadas)
    }
    
    # 🔥 GUARDAR NO PEDIDOS_REGISTO (PARA TEMPORIZADOR FUNCIONAR)
    PEDIDOS_REGISTO[pedido_id] = pedido_data
    
    # 🔥 GUARDAR TAMBÉM NO USER_DATA (PARA FACILITAR ACESSO)
    context.user_data["pedido_data"] = pedido_data.copy()
    context.user_data["pedido_id"] = pedido_id
    # 🔥 MANTER O OFERTA_TIPO NO CONTEXT PARA O PAGAR_STRIPE
    context.user_data["oferta_tipo"] = oferta_tipo_final

    print(f"✅ PEDIDO GIFT GUARDADO NO PEDIDOS_REGISTO: #{pedido_id}")
    print(f"📊 TIPO DE OFERTA: {oferta_tipo_final}")
    print(f"📊 PRODUTO: portachaves")
    print(f"📊 IDIOMA: {idioma}")
    print(f"📊 NÃO CONTADO NAS ESTATÍSTICAS (aguardando pagamento)")
    print(f"💰 Moeda do pedido: {totais['moeda']} {totais['simbolo_moeda']}")

    # 🔥 CÁLCULO DO PREÇO ANTERIOR E DESCONTO
    preco_anterior = totais['total'] / 0.70
    desconto = preco_anterior - totais['total']
    percentual_desconto = 30

    # 🔥 TEXTOS POR IDIOMA PARA RESUMO FINAL
    textos_resumo = {
        'portugues': {
            'titulo': "🎁 <b>RESUMO FINAL DO PORTA-CHAVES</b>\n\n",
            'id_pedido': "<b>🆔 ID do Pedido:</b>",
            'data': "<b>📅 Data:</b>",
            'pais_envio': "<b>🌍 País de Envio:</b>",
            'moeda': "<b>💰 Moeda:</b>",
            'tempo_pagar': "<b>⏰ Tempo para pagar:</b> 10 minutos\n\n",
            'dados_pessoais': "<b>👤 DADOS PESSOAIS:</b>",
            'nome': "• 👤 <b>Nome:</b>",
            'email': "• 📧 <b>Email:</b>",
            'pais': "• 🌍 <b>País:</b>",
            'telefone': "• 📱 <b>Telefone:</b>",
            'detalhes': "<b>🎁 DETALHES DO PORTA-CHAVES:</b>",
            'produto': "• 🎁 <b>Produto:</b>",
            'tamanho': "• 📏 <b>Tamanho:</b>",
            'nome_box': "• 🎭 <b>Nome na Box:</b>",
            'frase_box': "• 💬 <b>Frase na Box:</b>",
            'foto': "• 📸 <b>Foto:</b>",
            'valores': "<b>💵 VALORES:</b>",
            'preco_anterior': "• <b>Preço anterior:</b>",
            'desconto': "• <b>Desconto:</b>",
            'total_pagar': "• 💰 <b>TOTAL A PAGAR:",
            'impostos_frete': "• 📝 <b>Impostos e Frete Incluídos</b>\n\n",
            'dados_pedido': "<b>📊 DADOS DO PEDIDO:</b>",
            'id': "• 🆔 <b>ID:</b>",
            'data2': "• 📅 <b>Data:</b>",
            'pais2': "• 🌍 <b>País:</b>",
            'moeda2': "• 💰 <b>Moeda:</b>",
            'total': "• 💵 <b>Total:</b>",
            'expira': "• ⏰ <b>Expira:</b>",
            'aviso_tempo': "<b>⚠️ Tem 10 minutos para efetuar o pagamento!</b>",
            'guarde_id': "<b>Guarde o ID do pedido para referência futura!</b>\n\n",
            'clique_pagar': "<b>Clique abaixo para pagar:</b> 👇",
            'botao_pagar': "💳 Pagar com Cartão"
        },
        'ingles': {
            'titulo': "🎁 <b>FINAL KEYCHAIN SUMMARY</b>\n\n",
            'id_pedido': "<b>🆔 Order ID:</b>",
            'data': "<b>📅 Date:</b>",
            'pais_envio': "<b>🌍 Shipping Country:</b>",
            'moeda': "<b>💰 Currency:</b>",
            'tempo_pagar': "<b>⏰ Time to pay:</b> 10 minutes\n\n",
            'dados_pessoais': "<b>👤 PERSONAL DATA:</b>",
            'nome': "• 👤 <b>Name:</b>",
            'email': "• 📧 <b>Email:</b>",
            'pais': "• 🌍 <b>Country:</b>",
            'telefone': "• 📱 <b>Phone:</b>",
            'detalhes': "<b>🎁 KEYCHAIN DETAILS:</b>",
            'produto': "• 🎁 <b>Product:</b>",
            'tamanho': "• 📏 <b>Size:</b>",
            'nome_box': "• 🎭 <b>Name on Box:</b>",
            'frase_box': "• 💬 <b>Phrase on Box:</b>",
            'foto': "• 📸 <b>Photo:</b>",
            'valores': "<b>💵 VALUES:</b>",
            'preco_anterior': "• <b>Previous price:</b>",
            'desconto': "• <b>Discount:</b>",
            'total_pagar': "• 💰 <b>TOTAL TO PAY:",
            'impostos_frete': "• 📝 <b>Taxes and Shipping Included</b>\n\n",
            'dados_pedido': "<b>📊 ORDER DATA:</b>",
            'id': "• 🆔 <b>ID:</b>",
            'data2': "• 📅 <b>Date:</b>",
            'pais2': "• 🌍 <b>Country:</b>",
            'moeda2': "• 💰 <b>Currency:</b>",
            'total': "• 💵 <b>Total:</b>",
            'expira': "• ⏰ <b>Expires:</b>",
            'aviso_tempo': "<b>⚠️ You have 10 minutes to make the payment!</b>",
            'guarde_id': "<b>Save the order ID for future reference!</b>\n\n",
            'clique_pagar': "<b>Click below to pay:</b> 👇",
            'botao_pagar': "💳 Pay with Card"
        },
        'espanhol': {
            'titulo': "🎁 <b>RESUMEN FINAL DEL LLAVERO</b>\n\n",
            'id_pedido': "<b>🆔 ID del Pedido:</b>",
            'data': "<b>📅 Fecha:</b>",
            'pais_envio': "<b>🌍 País de Envío:</b>",
            'moeda': "<b>💰 Moneda:</b>",
            'tempo_pagar': "<b>⏰ Tiempo para pagar:</b> 10 minutos\n\n",
            'dados_pessoais': "<b>👤 DATOS PERSONALES:</b>",
            'nome': "• 👤 <b>Nombre:</b>",
            'email': "• 📧 <b>Email:</b>",
            'pais': "• 🌍 <b>País:</b>",
            'telefone': "• 📱 <b>Teléfono:</b>",
            'detalhes': "<b>🎁 DETALLES DEL LLAVERO:</b>",
            'produto': "• 🎁 <b>Producto:</b>",
            'tamanho': "• 📏 <b>Tamaño:</b>",
            'nome_box': "• 🎭 <b>Nombre en Caja:</b>",
            'frase_box': "• 💬 <b>Frase en Caja:</b>",
            'foto': "• 📸 <b>Foto:</b>",
            'valores': "<b>💵 VALORES:</b>",
            'preco_anterior': "• <b>Precio anterior:</b>",
            'desconto': "• <b>Descuento:</b>",
            'total_pagar': "• 💰 <b>TOTAL A PAGAR:",
            'impostos_frete': "• 📝 <b>Impuestos y Envío Incluidos</b>\n\n",
            'dados_pedido': "<b>📊 DATOS DEL PEDIDO:</b>",
            'id': "• 🆔 <b>ID:</b>",
            'data2': "• 📅 <b>Fecha:</b>",
            'pais2': "• 🌍 <b>País:</b>",
            'moeda2': "• 💰 <b>Moneda:</b>",
            'total': "• 💵 <b>Total:</b>",
            'expira': "• ⏰ <b>Expira:</b>",
            'aviso_tempo': "<b>⚠️ ¡Tienes 10 minutos para realizar el pago!</b>",
            'guarde_id': "<b>¡Guarda el ID del pedido para referencia futura!</b>\n\n",
            'clique_pagar': "<b>Haz clic abajo para pagar:</b> 👇",
            'botao_pagar': "💳 Pagar con Tarjeta"
        },
        'italiano': {
            'titulo': "🎁 <b>RIEPILOGO FINALE DEL PORTACHIAVI</b>\n\n",
            'id_pedido': "<b>🆔 ID Ordine:</b>",
            'data': "<b>📅 Data:</b>",
            'pais_envio': "<b>🌍 Paese di Spedizione:</b>",
            'moeda': "<b>💰 Valuta:</b>",
            'tempo_pagar': "<b>⏰ Tempo per pagare:</b> 10 minuti\n\n",
            'dados_pessoais': "<b>👤 DATI PERSONALI:</b>",
            'nome': "• 👤 <b>Nome:</b>",
            'email': "• 📧 <b>Email:</b>",
            'pais': "• 🌍 <b>Paese:</b>",
            'telefone': "• 📱 <b>Telefono:</b>",
            'detalhes': "<b>🎁 DETTAGLI PORTACHIAVI:</b>",
            'produto': "• 🎁 <b>Prodotto:</b>",
            'tamanho': "• 📏 <b>Dimensione:</b>",
            'nome_box': "• 🎭 <b>Nome su Scatola:</b>",
            'frase_box': "• 💬 <b>Frase su Scatola:</b>",
            'foto': "• 📸 <b>Foto:</b>",
            'valores': "<b>💵 VALORI:</b>",
            'preco_anterior': "• <b>Prezzo precedente:</b>",
            'desconto': "• <b>Sconto:</b>",
            'total_pagar': "• 💰 <b>TOTALE DA PAGARE:",
            'impostos_frete': "• 📝 <b>Tasse e Spedizione Incluse</b>\n\n",
            'dados_pedido': "<b>📊 DATI ORDINE:</b>",
            'id': "• 🆔 <b>ID:</b>",
            'data2': "• 📅 <b>Data:</b>",
            'pais2': "• 🌍 <b>Paese:</b>",
            'moeda2': "• 💰 <b>Valuta:</b>",
            'total': "• 💵 <b>Totale:</b>",
            'expira': "• ⏰ <b>Scade:</b>",
            'aviso_tempo': "<b>⚠️ Hai 10 minuti per effettuare il pagamento!</b>",
            'guarde_id': "<b>Conserva l'ID dell'ordine per riferimento futuro!</b>\n\n",
            'clique_pagar': "<b>Clicca sotto per pagare:</b> 👇",
            'botao_pagar': "💳 Paga con Carta"
        },
        'alemao': {
            'titulo': "🎁 <b>ENDZUSAMMENFASSUNG SCHLÜSSELANHÄNGER</b>\n\n",
            'id_pedido': "<b>🆔 Bestell-ID:</b>",
            'data': "<b>📅 Datum:</b>",
            'pais_envio': "<b>🌍 Versandland:</b>",
            'moeda': "<b>💰 Währung:</b>",
            'tempo_pagar': "<b>⏰ Zeit zum Bezahlen:</b> 10 Minuten\n\n",
            'dados_pessoais': "<b>👤 PERSÖNLICHE DATEN:</b>",
            'nome': "• 👤 <b>Name:</b>",
            'email': "• 📧 <b>E-Mail:</b>",
            'pais': "• 🌍 <b>Land:</b>",
            'telefone': "• 📱 <b>Telefon:</b>",
            'detalhes': "<b>🎁 SCHLÜSSELANHÄNGER DETAILS:</b>",
            'produto': "• 🎁 <b>Produkt:</b>",
            'tamanho': "• 📏 <b>Größe:</b>",
            'nome_box': "• 🎭 <b>Name auf Box:</b>",
            'frase_box': "• 💬 <b>Phrase auf Box:</b>",
            'foto': "• 📸 <b>Foto:</b>",
            'valores': "<b>💵 WERTE:</b>",
            'preco_anterior': "• <b>Vorheriger Preis:</b>",
            'desconto': "• <b>Rabatt:</b>",
            'total_pagar': "• 💰 <b>GESAMT ZU ZAHLEN:",
            'impostos_frete': "• 📝 <b>Steuern und Versand inklusive</b>\n\n",
            'dados_pedido': "<b>📊 BESTELLDATEN:</b>",
            'id': "• 🆔 <b>ID:</b>",
            'data2': "• 📅 <b>Datum:</b>",
            'pais2': "• 🌍 <b>Land:</b>",
            'moeda2': "• 💰 <b>Währung:</b>",
            'total': "• 💵 <b>Gesamtsumme:</b>",
            'expira': "• ⏰ <b>Läuft ab:</b>",
            'aviso_tempo': "<b>⚠️ Sie haben 10 Minuten, um die Zahlung vorzunehmen!</b>",
            'guarde_id': "<b>Bewahren Sie die Bestell-ID für zukünftige Referenz auf!</b>\n\n",
            'clique_pagar': "<b>Klicken Sie unten, um zu bezahlen:</b> 👇",
            'botao_pagar': "💳 Mit Karte bezahlen"
        },
        'frances': {
            'titulo': "🎁 <b>RÉSUMÉ FINAL DU PORTE-CLÉS</b>\n\n",
            'id_pedido': "<b>🆔 ID de Commande:</b>",
            'data': "<b>📅 Date:</b>",
            'pais_envio': "<b>🌍 Pays d'Expédition:</b>",
            'moeda': "<b>💰 Devise:</b>",
            'tempo_pagar': "<b>⏰ Temps pour payer:</b> 10 minutes\n\n",
            'dados_pessoais': "<b>👤 DONNÉES PERSONNELLES:</b>",
            'nome': "• 👤 <b>Nom:</b>",
            'email': "• 📧 <b>Email:</b>",
            'pais': "• 🌍 <b>Pays:</b>",
            'telefone': "• 📱 <b>Téléphone:</b>",
            'detalhes': "<b>🎁 DÉTAILS DU PORTE-CLÉS:</b>",
            'produto': "• 🎁 <b>Produit:</b>",
            'tamanho': "• 📏 <b>Taille:</b>",
            'nome_box': "• 🎭 <b>Nom sur Boîte:</b>",
            'frase_box': "• 💬 <b>Phrase sur Boîte:</b>",
            'foto': "• 📸 <b>Photo:</b>",
            'valores': "<b>💵 VALEURS:</b>",
            'preco_anterior': "• <b>Prix précédent:</b>",
            'desconto': "• <b>Réduction:</b>",
            'total_pagar': "• 💰 <b>TOTAL À PAYER:",
            'impostos_frete': "• 📝 <b>Taxes et Livraison Incluses</b>\n\n",
            'dados_pedido': "<b>📊 DONNÉES DE COMMANDE:</b>",
            'id': "• 🆔 <b>ID:</b>",
            'data2': "• 📅 <b>Date:</b>",
            'pais2': "• 🌍 <b>Pays:</b>",
            'moeda2': "• 💰 <b>Devise:</b>",
            'total': "• 💵 <b>Total:</b>",
            'expira': "• ⏰ <b>Expire:</b>",
            'aviso_tempo': "<b>⚠️ Vous avez 10 minutes pour effectuer le paiement !</b>",
            'guarde_id': "<b>Conservez l'ID de commande pour référence future !</b>\n\n",
            'clique_pagar': "<b>Cliquez ci-dessous pour payer:</b> 👇",
            'botao_pagar': "💳 Payer par Carte"
        }
    }
    
    textos = textos_resumo.get(idioma, textos_resumo['portugues'])
    
    # 🔥 CONSTRUIR TEXTO HTML TRADUZIDO COM PAÍS EM INGLÊS
    texto = f"""{textos['titulo']}
{textos['id_pedido']} {pedido_id}
{textos['data']} {data_pedido}
{textos['pais_envio']} {pais_ingles}  
{textos['moeda']} {totais['moeda']} {totais['simbolo_moeda']}
{textos['tempo_pagar']}
{textos['dados_pessoais']}
{textos['nome']} {nome}
{textos['email']} {email}
{textos['pais']} {pais_ingles}  
{textos['telefone']} {contacto}

{textos['detalhes']}
{textos['produto']} {textos_prod['produto']}
{textos['tamanho']} {textos_prod['tamanho']}"""

    # 🔥 CAMPOS DE PERSONALIZAÇÃO DA BOX
    if nome_gift and nome_gift != textos_sem_frase.get(idioma, textos_sem_frase['portugues']).replace("frase", "nome/alcunha"):
        texto += f"\n{textos['nome_box']} {nome_gift}"
    if frase_gift and frase_gift != textos_sem_frase.get(idioma, textos_sem_frase['portugues']):
        texto += f"\n{textos['frase_box']} \"{frase_gift}\""
    
    texto += f"""
{textos['foto']} {foto_recebida} ({nome_foto})

{textos['valores']}
{textos['preco_anterior']} {totais['simbolo_moeda']}{preco_anterior:.2f}❌ 
{textos['desconto']} {totais['simbolo_moeda']}{desconto:.2f} ({percentual_desconto}% OFF)
{textos['total_pagar']} {totais['simbolo_moeda']}{totais['total']:.2f}</b>
{textos['impostos_frete']}
{textos['dados_pedido']}
{textos['id']} {pedido_id}
{textos['data2']} {data_pedido}
{textos['pais2']} {pais_ingles}  
{textos['moeda2']} {totais['moeda']}
{textos['total']} {totais['simbolo_moeda']}{totais['total']:.2f}
{textos['expira']} {(datetime.now() + timedelta(minutes=10)).strftime("%d/%m/%Y %H:%M")}

{textos['aviso_tempo']}
{textos['guarde_id']}
{textos['clique_pagar']}"""

    # BOTÕES TRADUZIDOS
    botoes = [
        [InlineKeyboardButton(textos['botao_pagar'], callback_data="pagar_stripe")]
    ]
    
    # ENVIAR MENSAGEM
    try:
        mensagem = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=texto, 
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup(botoes)
        )
        print(f"✅ Resumo de pagamento gift enviado | Idioma: {idioma}")
        print(f"✅ País mostrado como: {pais_ingles} (em inglês)")
        
    except Exception as e:
        print(f"❌ Erro ao enviar com HTML: {e}")
        # Fallback para Markdown se HTML falhar
        try:
            texto_simples = f"🎁 RESUMO FINAL DO PORTA-CHAVES\n\nID: {pedido_id}\nTotal: {totais['simbolo_moeda']}{totais['total']:.2f}\n\nClique para pagar:"
            mensagem = await context.bot.send_message(
                chat_id=query.message.chat_id,
                text=texto_simples,
                reply_markup=InlineKeyboardMarkup(botoes)
            )
        except Exception as e2:
            print(f"❌ Erro também no fallback: {e2}")
            return
    
    # 🔥 🔥 🔥 AGORA SIM: TEMPORIZADOR (PEDIDO JÁ ESTÁ NO REGISTRO)
    print(f"⏰ Iniciando temporizador de 10min para pedido GIFT #{pedido_id} | Idioma: {idioma}")
    await iniciar_temporizador(context, pedido_id, query.message.chat_id, mensagem.message_id)









#create my cartoon

# --- Menu inicial ---
async def menu_inicial(update: Update, context: ContextTypes.DEFAULT_TYPE):
    texto = "👋 Olá! Bem-vindo à *GodsPlan*, vamos criar o seu cartoon?"
    keyboard = [[InlineKeyboardButton("CREATE MY CARTOON", callback_data="mycartoon")]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    if update.callback_query:
        await update.callback_query.edit_message_text(
            texto, reply_markup=reply_markup, parse_mode="Markdown"
        )
    else:
        await update.message.reply_text(
            texto, reply_markup=reply_markup, parse_mode="Markdown"
        )



# --- /start ---
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat_id
    user_id = update.effective_user.id
    
    # 🔥 OBTER SESSÃO DO USUÁRIO
    session = get_user_session(user_id)
    
    # 🔥 USAR DADOS DA SESSÃO EM VEZ DE context.user_data
    session.clear_state()  # Limpar estado anterior
    
    # 🔥 ATUALIZAR CONTADOR (se esta função existir)
    try:
        utilizadores_hoje = atualizar_contador_utilizadores(user_id)
        print(f"🚀 BOT INICIADO por user {user_id} | Utilizadores hoje: {utilizadores_hoje}")
    except:
        print(f"🚀 BOT INICIADO por user {user_id}")
    
    print(f"🔧 /start chamado para chat {chat_id}")
    
    # 🔥 CANCELAR TEMPORIZADORES
    try:
        cancelar_temporizador_30min(chat_id)
    except:
        pass
    
    # 🔥 LIMPAR DADOS
    context.user_data.clear()
    
    print(f"✅ Dados limpos e temporizador cancelado para chat {chat_id}")
    
    # 🔥 🔥 🔥 **PRIMEIRO: PEDIR PARA ESCOLHER IDIOMA**
    texto_escolha_idioma = "🌍 *Please choose your language / Por favor escolha seu idioma:*"
    
    keyboard = [
        [
            InlineKeyboardButton("🇵🇹 Português", callback_data="idioma_portugues"),
            InlineKeyboardButton("🇺🇸 English", callback_data="idioma_ingles")
        ],
        [
            InlineKeyboardButton("🇪🇸 Español", callback_data="idioma_espanhol"),
            InlineKeyboardButton("🇮🇹 Italiano", callback_data="idioma_italiano")
        ],
        [
            InlineKeyboardButton("🇩🇪 Deutsch", callback_data="idioma_alemao"),
            InlineKeyboardButton("🇫🇷 Français", callback_data="idioma_frances")
        ]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        texto_escolha_idioma,
        reply_markup=reply_markup,
        parse_mode="Markdown"
    )
    
    print(f"✅ Tela de escolha de idioma mostrada para chat {chat_id}")



async def selecionar_idioma(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para seleção de idioma"""
    query = update.callback_query
    await query.answer()
    
    # Extrair idioma do callback_data
    idioma = query.data.replace("idioma_", "")
    
    # Salvar idioma no user_data
    context.user_data['idioma'] = idioma
    
    print(f"✅ Idioma selecionado: {idioma} por user {query.from_user.id}")
    
    # Mensagens de confirmação em cada idioma
    mensagens_confirmacao = {
        'portugues': "✅ *Idioma definido para Português!*",
        'ingles': "✅ *Language set to English!*",
        'espanhol': "✅ *¡Idioma establecido en Español!*",
        'italiano': "✅ *Lingua impostata su Italiano!*",
        'alemao': "✅ *Sprache auf Deutsch eingestellt!*",
        'frances': "✅ *Langue définie sur Français!*"
    }
    
    # Apagar mensagem de escolha de idioma
    await query.delete_message()
    
    # Mostrar confirmação
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=mensagens_confirmacao.get(idioma, "✅ Idioma selecionado!"),
        parse_mode="Markdown"
    )
    
    # 🔥 AGORA MOSTRAR O MENU INICIAL NO IDIOMA ESCOLHIDO
    texto_intro = ""
    botao_texto = ""
    texto_botao = ""  # Texto do botão em cada idioma
    
    if idioma == 'portugues':
        texto_intro = """
🎨 *BEM-VINDO À GODSPLAN*

*Transforme seus momentos em arte!*

🎭 `/start` - Individual | Grupo | Animal | Personalizado\n
🎁 `/gift` - LIMITED EDITION | 🛑 APROVEITA AGORA! (30% OFF)\n
ℹ️ `/help` - Ajuda e Suporte
"""
        botao_texto = "💫 *Pronto para criar o seu cartoon personalizado?*"
        texto_botao = "🎭 CRIAR MEU CARTOON"
    
    elif idioma == 'ingles':
        texto_intro = """
🎨 *WELCOME TO GODSPLAN*

*Transform your moments into art!*

🎭 `/start` - Individual | Group | Animal | Custom\n
🎁 `/gift` - LIMITED EDITION | 🛑 TAKE ADVANTAGE NOW! (30% OFF)\n
ℹ️ `/help` - Help & Support
"""
        botao_texto = "💫 *Ready to create your personalized cartoon?*"
        texto_botao = "🎭 CREATE MY CARTOON"
    
    elif idioma == 'espanhol':
        texto_intro = """
🎨 *BIENVENIDO A GODSPLAN*

*¡Transforma tus momentos en arte!*

🎭 `/start` - Individual | Grupo | Animal | Personalizado\n
🎁 `/gift` - EDICIÓN LIMITADA | 🛑 ¡APROVECHA AHORA! (30% OFF)\n
ℹ️ `/help` - Ayuda y Soporte
"""
        botao_texto = "💫 *¿Listo para crear tu caricatura personalizada?*"
        texto_botao = "🎭 CREAR MI CARICATURA"
    
    elif idioma == 'italiano':
        texto_intro = """
🎨 *BENVENUTO IN GODSPLAN*

*Trasforma i tuoi momenti in arte!*

🎭 `/start` - Individuale | Gruppo | Animale | Personalizzato\n
🎁 `/gift` - EDIZIONE LIMITATA | 🛑 APPROFITTANE ORA! (30% OFF)\n
ℹ️ `/help` - Aiuto e Supporto
"""
        botao_texto = "💫 *Pronto per creare il tuo cartoon personalizzato?*"
        texto_botao = "🎭 CREA IL MIO CARTOON"
    
    elif idioma == 'alemao':
        texto_intro = """
🎨 *WILLKOMMEN BEI GODSPLAN*

*Verwandle deine Momente in Kunst!*

🎭 `/start` - Einzeln | Gruppe | Tier | Personalisiert\n
🎁 `/gift` - LIMITIERTE EDITION | 🛑 JETZT VORTEIL NUTZEN! (30% OFF)\n
ℹ️ `/help` - Hilfe & Support
"""
        botao_texto = "💫 *Bereit, deine personalisierte Karikatur zu erstellen?*"
        texto_botao = "🎭 MEINE KARIKATUR ERSTELLEN"
    
    elif idioma == 'frances':
        texto_intro = """
🎨 *BIENVENUE CHEZ GODSPLAN*

*Transformez vos moments en art !*

🎭 `/start` - Individuel | Groupe | Animal | Personnalisé\n
🎁 `/gift` - ÉDITION LIMITÉE | 🛑 PROFITEZ MAINTENANT ! (30% OFF)\n
ℹ️ `/help` - Aide & Support
"""
        botao_texto = "💫 *Prêt à créer votre dessin animé personnalisé ?*"
        texto_botao = "🎭 CRÉER MON DESSIN ANIMÉ"
    
    # Enviar introdução
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=texto_intro,
        parse_mode="Markdown"
    )
    
    # 🔥 BOTÃO TRADUZIDO POR IDIOMA
    keyboard = [
        [InlineKeyboardButton(texto_botao, callback_data="iniciar_cartoon")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=botao_texto,
        reply_markup=reply_markup,
        parse_mode="Markdown"
    )
    
    print(f"✅ Menu inicial mostrado no idioma: {idioma}")
    print(f"✅ Botão traduzido: {texto_botao}")



async def menu_inicial(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Menu inicial com botão traduzido"""
    # Pegar idioma do user_data (padrão: português)
    idioma = context.user_data.get('idioma', 'portugues')
    
    # Textos em cada idioma
    textos_menu = {
        'portugues': "👋 Olá! Bem-vindo à *GodsPlan*, vamos criar o seu cartoon?",
        'ingles': "👋 Hello! Welcome to *GodsPlan*, shall we create your cartoon?",
        'espanhol': "👋 ¡Hola! Bienvenido a *GodsPlan*, ¿vamos a crear tu caricatura?",
        'italiano': "👋 Ciao! Benvenuto in *GodsPlan*, creiamo il tuo cartoon?",
        'alemao': "👋 Hallo! Willkommen bei *GodsPlan*, sollen wir deine Karikatur erstellen?",
        'frances': "👋 Bonjour! Bienvenue chez *GodsPlan*, allons-nous créer votre dessin animé?"
    }
    
    # Botões em cada idioma
    botoes_menu = {
        'portugues': "🎭 CRIAR MEU CARTOON",
        'ingles': "🎭 CREATE MY CARTOON",
        'espanhol': "🎭 CREAR MI CARICATURA",
        'italiano': "🎭 CREA IL MIO CARTOON",
        'alemao': "🎭 MEINE KARIKATUR ERSTELLEN",
        'frances': "🎭 CRÉER MON DESSIN ANIMÉ"
    }
    
    texto = textos_menu.get(idioma, textos_menu['portugues'])
    texto_botao = botoes_menu.get(idioma, botoes_menu['portugues'])
    
    keyboard = [[InlineKeyboardButton(texto_botao, callback_data="mycartoon")]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    if update.callback_query:
        await update.callback_query.edit_message_text(
            texto, reply_markup=reply_markup, parse_mode="Markdown"
        )
    else:
        await update.message.reply_text(
            texto, reply_markup=reply_markup, parse_mode="Markdown"
        )
    
    print(f"✅ Menu inicial mostrado em: {idioma} | Botão: {texto_botao}")



# --- Iniciar criação de cartoon ---
async def iniciar_cartoon(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # Textos da primeira pergunta em cada idioma
    textos_inicio = {
        'portugues': "🔥 Vamos criar o nosso Cartoon 3D!\n\nAntes de começarmos, qual é o seu nome?",
        'ingles': "🔥 Let's create our 3D Cartoon!\n\nBefore we start, what's your name?",
        'espanhol': "🔥 ¡Vamos a crear nuestro Cartoon 3D!\n\nAntes de empezar, ¿cuál es tu nombre?",
        'italiano': "🔥 Creiamo il nostro Cartoon 3D!\n\nPrima di iniziare, qual è il tuo nome?",
        'alemao': "🔥 Lass uns unsere 3D-Karikatur erstellen!\n\nBevor wir anfangen, wie ist dein Name?",
        'frances': "🔥 Créons notre Dessin Animé 3D !\n\nAvant de commencer, quel est votre nom ?"
    }
    
    # Textos dos botões para verificar se é mensagem com botão
    textos_botao = {
        'portugues': "💫 *Pronto para criar o seu cartoon personalizado?*",
        'ingles': "💫 *Ready to create your personalized cartoon?*",
        'espanhol': "💫 *¿Listo para crear tu caricatura personalizada?*",
        'italiano': "💫 *Pronto per creare il tuo cartoon personalizado?*",
        'alemao': "💫 *Bereit, deine personalisierte Karikatur zu erstellen?*",
        'frances': "💫 *Prêt à créer votre dessin animé personnalisé ?*"
    }
    
    # 🔥 ACEITAR AMBOS OS PADRÕES
    if query.data in ["mycartoon", "iniciar_cartoon"]:
        # 🔥 APAGAR APENAS SE FOR MENSAGEM COM BOTÃO
        try:
            message_text = query.message.text or ""
            
            # Verificar se é mensagem com botão em QUALQUER idioma
            texto_botao_atual = textos_botao.get(idioma, textos_botao['portugues'])
            texto_botao_portugues = textos_botao['portugues']
            texto_botao_ingles = textos_botao['ingles']
            
            # Verifica se a mensagem contém algum dos textos de botão
            if (texto_botao_atual in message_text or 
                texto_botao_portugues in message_text or 
                texto_botao_ingles in message_text or
                "Pronto para criar" in message_text or 
                "Ready to create" in message_text or
                "CREATE MY CARTOON" in message_text):
                
                await query.delete_message()
                print(f"✅ Mensagem com botão apagada | Idioma: {idioma}")
        except Exception as e:
            print(f"❌ Erro ao apagar mensagem: {e}")
        
        # 🔥 PRIMEIRA PERGUNTA NO IDIOMA CORRETO
        await query.message.reply_text(textos_inicio.get(idioma, textos_inicio['portugues']))
        
        context.user_data['conversation_state'] = NOME
        print(f"✅ Fluxo de cartoon iniciado - estado: NOME | Idioma: {idioma}")



# --- Perguntas sequenciais ---
async def receber_nome(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["nome"] = update.message.text
    
    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 APAGAR MENSAGEM DO USUÁRIO (resposta)
    try:
        await update.message.delete()
        print("✅ Mensagem do usuário (nome) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem do usuário: {e}")
    
    # 🔥 APAGAR MENSAGEM DA PERGUNTA DO NOME
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id,
            message_id=update.message.message_id - 1
        )
        print("✅ Mensagem da pergunta (nome) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar pergunta do nome: {e}")
    
    # 🔥 PRÓXIMA PERGUNTA (EMAIL) NO IDIOMA CORRETO
    textos_email = {
        'portugues': "📧 Perfeito! Agora, qual é o seu email?",
        'ingles': "📧 Perfect! Now, what's your email?",
        'espanhol': "📧 ¡Perfecto! Ahora, ¿cuál es tu email?",
        'italiano': "📧 Perfetto! Ora, qual è la tua email?",
        'alemao': "📧 Perfekt! Nun, wie lautet deine E-Mail?",
        'frances': "📧 Parfait ! Maintenant, quelle est votre adresse e-mail ?"
    }
    
    await update.message.reply_text(textos_email.get(idioma, textos_email['portugues']))
    
    context.user_data['conversation_state'] = EMAIL
    print(f"✅ Estado atualizado para: EMAIL | Idioma: {idioma}")






async def receber_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["email"] = update.message.text
    
    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 APAGAR MENSAGEM DO USUÁRIO (resposta)
    try:
        await update.message.delete()
        print("✅ Mensagem do usuário (email) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem do usuário: {e}")
    
    # 🔥 APAGAR MENSAGEM DA PERGUNTA DO EMAIL
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id,
            message_id=update.message.message_id - 1
        )
        print("✅ Mensagem da pergunta (email) apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar pergunta do email: {e}")
    
    # 🔥 PRÓXIMA PERGUNTA (PAÍS) NO IDIOMA CORRETO
    textos_pais = {
        'portugues': "📋 De qual país você é?",
        'ingles': "📋 Which country are you from?",
        'espanhol': "📋 ¿De qué país eres?",
        'italiano': "📋 Di quale paese sei?",
        'alemao': "📋 Aus welchem Land kommst du?",
        'frances': "📋 De quel pays êtes-vous ?"
    }
    
    texto = textos_pais.get(idioma, textos_pais['portugues'])
    
    # Países (os mesmos em todos os idiomas, mas com emojis)
    keyboard = [
     [InlineKeyboardButton("🇺🇸 United States", callback_data="pais_estados_unidos")],
     [InlineKeyboardButton("🇨🇦 Canada", callback_data="pais_canada"),
     InlineKeyboardButton("🇬🇧 United Kingdom", callback_data="pais_reino_unido")],
     [InlineKeyboardButton("🇧🇷 Brazil", callback_data="pais_brasil"),
     InlineKeyboardButton("🇩🇪 Germany", callback_data="pais_alemanha")],
     [InlineKeyboardButton("🇳🇱 Netherlands", callback_data="pais_paises_baixos"),
     InlineKeyboardButton("🇫🇷 France", callback_data="pais_franca")],
     [InlineKeyboardButton("🇪🇸 Spain", callback_data="pais_espanha"),
     InlineKeyboardButton("🇧🇪 Belgium", callback_data="pais_belgica")],
     [InlineKeyboardButton("🇮🇹 Italy", callback_data="pais_italia"),
     InlineKeyboardButton("🇵🇹 Portugal", callback_data="pais_portugal")],
     [InlineKeyboardButton("🇮🇪 Ireland", callback_data="pais_irlanda"),
     InlineKeyboardButton("🇱🇺 Luxembourg", callback_data="pais_luxemburgo")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(texto, reply_markup=reply_markup)
    
    context.user_data['conversation_state'] = PAIS
    print(f"✅ Estado atualizado para: PAIS | Idioma: {idioma}")





# --- Handler para seleção de país ---
async def selecionar_pais(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    pais_data = query.data
    
    if pais_data == "pais_outro":
        # 🔥 APAGAR MENSAGEM DOS PAÍSES (pergunta)
        try:
            await query.delete_message()
            print("✅ Mensagem dos países apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar mensagem dos países: {e}")
        
        # Texto "digite país" por idioma
        textos_digite_pais = {
            'portugues': "Por favor, digite o nome do seu país:",
            'ingles': "Please enter your country name:",
            'espanhol': "Por favor, escribe el nombre de tu país:",
            'italiano': "Per favore, inserisci il nome del tuo paese:",
            'alemao': "Bitte gib den Namen deines Landes ein:",
            'frances': "Veuillez entrer le nom de votre pays :"
        }
        
        await query.message.reply_text(textos_digite_pais.get(idioma, textos_digite_pais['portugues']))
        context.user_data['aguardando_pais_manual'] = True
        return
    
    nome_pais = pais_data.replace("pais_", "")
    pais_formatado = nome_pais.replace("_", " ").title()
    prefixo = PAISES_PREFIXOS.get(nome_pais, "+??")
    
    context.user_data["pais"] = pais_formatado
    context.user_data["prefixo"] = prefixo
    
    # 🔥 APAGAR MENSAGEM DOS PAÍSES (pergunta)
    try:
        await query.delete_message()
        print("✅ Mensagem dos países apagada")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem dos países: {e}")
    
    # 🔥 TEXTO "AGORA ENVIE TELEFONE" POR IDIOMA
    textos_telefone = {
        'portugues': f"🌍 País selecionado: *{pais_formatado}*\n📞 Prefixo: {prefixo}\n\nAgora envie o seu número de telemóvel:",
        'ingles': f"🌍 Selected country: *{pais_formatado}*\n📞 Prefix: {prefixo}\n\nNow send your phone number:",
        'espanhol': f"🌍 País seleccionado: *{pais_formatado}*\n📞 Prefijo: {prefixo}\n\nAhora envía tu número de teléfono:",
        'italiano': f"🌍 Paese selezionato: *{pais_formatado}*\n📞 Prefisso: {prefixo}\n\nOra invia il tuo numero di telefono:",
        'alemao': f"🌍 Ausgewähltes Land: *{pais_formatado}*\n📞 Vorwahl: {prefixo}\n\nJetzt sende deine Telefonnummer:",
        'frances': f"🌍 Pays sélectionné : *{pais_formatado}*\n📞 Indicatif : {prefixo}\n\nMaintenant envoyez votre numéro de téléphone :"
    }
    
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_telefone.get(idioma, textos_telefone['portugues']),
        parse_mode="Markdown"
    )
    
    context.user_data['conversation_state'] = CONTACTO
    print(f"✅ Estado: CONTACTO | Idioma: {idioma}")



# --- Handler para país manual ---
async def receber_pais_manual(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('aguardando_pais_manual'):
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        pais_manual = update.message.text
        context.user_data["pais"] = pais_manual.title()
        context.user_data["prefixo"] = "+??"
        context.user_data['aguardando_pais_manual'] = False
        
        # 🔥 APAGAR MENSAGEM DO USUÁRIO (resposta)
        try:
            await update.message.delete()
            print("✅ Mensagem do usuário (país) apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar mensagem do usuário: {e}")
        
        # 🔥 APAGAR MENSAGEM DA PERGUNTA DO PAÍS
        try:
            await context.bot.delete_message(
                chat_id=update.message.chat_id,
                message_id=update.message.message_id - 1
            )
            print("✅ Mensagem da pergunta (país) apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar pergunta do país: {e}")
        
        # 🔥 TEXTO "PAÍS PERSONALIZADO + TELEFONE" POR IDIOMA
        textos_personalizado = {
            'portugues': f"🌍 País: *{pais_manual.title()}*\n📞 Prefixo: +?? (país personalizado)\n\nAgora envie o seu número de telemóvel:",
            'ingles': f"🌍 Country: *{pais_manual.title()}*\n📞 Prefix: +?? (custom country)\n\nNow send your phone number:",
            'espanhol': f"🌍 País: *{pais_manual.title()}*\n📞 Prefijo: +?? (país personalizado)\n\nAhora envía tu número de teléfono:",
            'italiano': f"🌍 Paese: *{pais_manual.title()}*\n📞 Prefisso: +?? (paese personalizzato)\n\nOra invia il tuo numero di telefono:",
            'alemao': f"🌍 Land: *{pais_manual.title()}*\n📞 Vorwahl: +?? (benutzerdefiniertes Land)\n\nJetzt sende deine Telefonnummer:",
            'frances': f"🌍 Pays : *{pais_manual.title()}*\n📞 Indicatif : +?? (pays personnalisé)\n\nMaintenant envoyez votre numéro de téléphone :"
        }
        
        await update.message.reply_text(
            textos_personalizado.get(idioma, textos_personalizado['portugues']),
            parse_mode="Markdown"
        )
        
        context.user_data['conversation_state'] = CONTACTO
        print(f"✅ Estado: CONTACTO | Idioma: {idioma}")




# --- Handler para receber contacto ---
async def receber_contacto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber e processar contacto do usuário"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        if update.message.contact:
            numero_completo = update.message.contact.phone_number
            context.user_data["contacto"] = numero_completo
        else:
            numero = update.message.text.strip()
            prefixo = context.user_data.get("prefixo", "+??")
            numero_completo = numero if numero.startswith("+") else f"{prefixo} {numero}"
            context.user_data["contacto"] = numero_completo

        print(f"✅ Contacto registrado: {context.user_data['contacto']} | Idioma: {idioma}")

        # 🔥 APAGAR MENSAGEM DO USUÁRIO (resposta)
        try:
            await update.message.delete()
            print("✅ Mensagem do usuário (contacto) apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar mensagem do usuário: {e}")
        
        # 🔥 APAGAR MENSAGEM DA PERGUNTA DO CONTACTO
        try:
            await context.bot.delete_message(
                chat_id=update.message.chat_id,
                message_id=update.message.message_id - 1
            )
            print("✅ Mensagem da pergunta (contacto) apagada")
        except Exception as e:
            print(f"❌ Erro ao apagar pergunta do contacto: {e}")

        # 🔥 ENVIAR RESUMO COM TODOS OS DADOS REGISTRADOS
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        
        if new_message_id:
            context.user_data['resumo_msg_id'] = new_message_id
            print(f"DEBUG: Resumo atualizado com ID: {new_message_id}")

        # 🔥 TEXTO "ESCOLHA TIPO DE CARTOON" POR IDIOMA
        textos_tipo_cartoon = {
            'portugues': "🎨 *Agora escolha o tipo de Cartoon:*",
            'ingles': "🎨 *Now choose the type of Cartoon:*",
            'espanhol': "🎨 *Ahora elige el tipo de Caricatura:*",
            'italiano': "🎨 *Ora scegli il tipo di Cartoon:*",
            'alemao': "🎨 *Wähle jetzt die Art der Karikatur:*",
            'frances': "🎨 *Maintenant choisissez le type de Dessin Animé :*"
        }
        
        texto = textos_tipo_cartoon.get(idioma, textos_tipo_cartoon['portugues'])
        
        # 🔥 BOTÕES TRADUZIDOS POR IDIOMA
        botoes_por_idioma = {
            'portugues': {
                'individual': "😎 Individual",
                'grupo': "👨‍👩‍👧 Grupo", 
                'animal': "🐱 Animal",
                'personalizado': "🎨 Personalizado"
            },
            'ingles': {
                'individual': "😎 Individual",
                'grupo': "👨‍👩‍👧 Group", 
                'animal': "🐱 Animal",
                'personalizado': "🎨 Custom"
            },
            'espanhol': {
                'individual': "😎 Individual",
                'grupo': "👨‍👩‍👧 Grupo", 
                'animal': "🐱 Animal",
                'personalizado': "🎨 Personalizado"
            },
            'italiano': {
                'individual': "😎 Individuale",
                'grupo': "👨‍👩‍👧 Gruppo", 
                'animal': "🐱 Animale",
                'personalizado': "🎨 Personalizzato"
            },
            'alemao': {
                'individual': "😎 Einzeln",
                'grupo': "👨‍👩‍👧 Gruppe", 
                'animal': "🐱 Tier",
                'personalizado': "🎨 Personalisiert"
            },
            'frances': {
                'individual': "😎 Individuel",
                'grupo': "👨‍👩‍👧 Groupe", 
                'animal': "🐱 Animal",
                'personalizado': "🎨 Personnalisé"
            }
        }
        
        botoes = botoes_por_idioma.get(idioma, botoes_por_idioma['portugues'])
        
        cartoon_keyboard = [
           [InlineKeyboardButton(botoes['individual'], callback_data="cartoon_individual"),
           InlineKeyboardButton(botoes['grupo'], callback_data="cartoon_family")],
           [InlineKeyboardButton(botoes['animal'], callback_data="cartoon_animal"),
           InlineKeyboardButton(botoes['personalizado'], callback_data="cartoon_custom")]
        ]
        
        await update.message.reply_text(
            texto, 
            reply_markup=InlineKeyboardMarkup(cartoon_keyboard),
            parse_mode="Markdown"
        )
        
        context.user_data['conversation_state'] = TIPO
        print(f"✅ Estado atualizado para: TIPO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_contacto: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))






async def receber_profissao(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber a profissão do usuário para estilo Office - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        profissao = update.message.text
        context.user_data["profissao"] = profissao
        
        # 🔥 CORREÇÃO: Tentar remover mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem profissão do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 CORREÇÃO: Tentar remover mensagem da pergunta
        mensagem_profissao_id = context.user_data.get('mensagem_profissao_id')
        if mensagem_profissao_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_profissao_id
                )
                print(f"DEBUG: Mensagem pergunta profissão apagada: {mensagem_profissao_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta profissão: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "PEDIR OBJETOS" POR IDIOMA
        textos_objetos = {
            'portugues': """🎯 *Excelente! Agora preciso que me digas 3 objetos que gostarias de ter ao teu lado no cartoon:*

• *Relacionado ao teu trabalho* 💼\n  
• *Relacionado ao teu hobby* 🎨\n
• *O que mais amas* ❤️

*Por exemplo:*
`Portátil, Guitarra, Café`

*Escreve os 3 objetos separados por vírgula:*""",
            
            'ingles': """🎯 *Excellent! Now I need you to tell me 3 objects you would like to have by your side in the cartoon:*

• *Related to your work* 💼\n  
• *Related to your hobby* 🎨\n
• *What you love most* ❤️

*For example:*
`Laptop, Guitar, Coffee`

*Write the 3 objects separated by commas:*""",
            
            'espanhol': """🎯 *¡Excelente! Ahora necesito que me digas 3 objetos que te gustaría tener a tu lado en la caricatura:*

• *Relacionado con tu trabajo* 💼\n  
• *Relacionado con tu hobby* 🎨\n
• *Lo que más amas* ❤️

*Por ejemplo:*
`Portátil, Guitarra, Café`

*Escribe los 3 objetos separados por comas:*""",
            
            'italiano': """🎯 *Eccellente! Ora ho bisogno che tu mi dica 3 oggetti che vorresti avere al tuo fianco nel cartoon:*

• *Relativo al tuo lavoro* 💼\n  
• *Relativo al tuo hobby* 🎨\n
• *Ciò che ami di più* ❤️

*Per esempio:*
`Computer portatile, Chitarra, Caffè`

*Scrivi i 3 oggetti separati da virgola:*""",
            
            'alemao': """🎯 *Ausgezeichnet! Jetzt brauche ich, dass du mir 3 Objekte nennst, die du neben dir im Cartoon haben möchtest:*

• *Bezogen auf deine Arbeit* 💼\n
• *Bezogen auf dein Hobby* 🎨\n
• *Was du am meisten liebst* ❤️

*Zum Beispiel:*
`Laptop, Gitarre, Kaffee`

*Schreibe die 3 Objekte durch Komma getrennt:*""",
            
            'frances': """🎯 *Excellent ! Maintenant j'ai besoin que vous me disiez 3 objets que vous aimeriez avoir à vos côtés dans le dessin animé :*

• *Lié à votre travail* 💼\n  
• *Lié à votre passe-temps* 🎨\n
• *Ce que vous aimez le plus* ❤️

*Par exemple :*
`Ordinateur portable, Guitare, Café`

*Écrivez les 3 objets séparés par des virgules :*"""
        }
        
        mensagem_objetos = await update.message.reply_text(
            textos_objetos.get(idioma, textos_objetos['portugues']),
            parse_mode="Markdown"
        )
        
        # 🔥 GUARDAR ID da mensagem para depois apagar
        context.user_data['mensagem_objetos_id'] = mensagem_objetos.message_id
        context.user_data['conversation_state'] = OBJETOS
        print(f"✅ Estado: OBJETOS | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_profissao: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))




async def receber_objetos(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber os 3 objetos personalizados para estilo Office - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        objetos_texto = update.message.text
        context.user_data["objetos_office"] = objetos_texto
        
        # 🔥 CORREÇÃO: Tentar remover mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem objetos do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem objetos usuário: {e}")
        
        # 🔥 CORREÇÃO: Tentar remover mensagem da pergunta
        mensagem_objetos_id = context.user_data.get('mensagem_objetos_id')
        if mensagem_objetos_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_objetos_id
                )
                print(f"DEBUG: Mensagem pergunta objetos apagada: {mensagem_objetos_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta objetos: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "ESCOLHER TAMANHO" POR IDIOMA
        textos_tamanho = {
            'portugues': "📏 *Perfeito! Agora escolhe o tamanho do teu Cartoon Office:*",
            'ingles': "📏 *Perfect! Now choose the size of your Office Cartoon:*",
            'espanhol': "📏 *¡Perfecto! Ahora elige el tamaño de tu Caricatura Office:*",
            'italiano': "📏 *Perfetto! Ora scegli la dimensione del tuo Cartoon Office:*",
            'alemao': "📏 *Perfekt! Wähle jetzt die Größe deiner Office-Karikatur:*",
            'frances': "📏 *Parfait ! Maintenant choisissez la taille de votre Dessin Animé Office :*"
        }
        
        # 🔥 AGORA mostrar os tamanhos para Office em GRADE 2xN
        estilo_escolhido = "Office"
        tamanhos_disponiveis = TAMANHOS_POR_ESTILO.get(estilo_escolhido, {})
        
        if not tamanhos_disponiveis:
            # 🔥 MENSAGEM DE ERRO SE NÃO HOUVER TAMANHOS
            textos_sem_tamanhos = {
                'portugues': "❌ Nenhum tamanho disponível para Office.",
                'ingles': "❌ No sizes available for Office.",
                'espanhol': "❌ No hay tamaños disponibles para Office.",
                'italiano': "❌ Nessuna dimensione disponibile per Office.",
                'alemao': "❌ Keine Größen für Office verfügbar.",
                'frances': "❌ Aucune taille disponible pour Office."
            }
            
            await update.message.reply_text(textos_sem_tamanhos.get(idioma, textos_sem_tamanhos['portugues']))
            return

        # Criar botões dos tamanhos em GRADE 2xN
        teclado = []
        tamanhos_lista = list(tamanhos_disponiveis.items())
        
        # Processar em pares (2 botões por linha)
        for i in range(0, len(tamanhos_lista), 2):
            linha = []
            # Primeiro botão da linha
            tamanho_key1, info_tamanho1 = tamanhos_lista[i]
            botao_texto1 = f"{info_tamanho1['nome']}"
            linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
            
            # Segundo botão da linha (se existir)
            if i + 1 < len(tamanhos_lista):
                tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
                botao_texto2 = f"{info_tamanho2['nome']}"
                linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
            
            teclado.append(linha)

        mensagem_tamanhos = await update.message.reply_text(
            textos_tamanho.get(idioma, textos_tamanho['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        
        # 🔥 GUARDAR ID da mensagem dos tamanhos
        context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
        context.user_data['conversation_state'] = TAMANHO
        print(f"✅ Estado: TAMANHO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_objetos: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))







async def receber_super_heroi(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber o super-herói escolhido para estilo Superheroes - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        super_heroi = update.message.text
        context.user_data["super_heroi"] = super_heroi
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem super-herói do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta
        mensagem_superheroi_id = context.user_data.get('mensagem_superheroi_id')
        if mensagem_superheroi_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_superheroi_id
                )
                print(f"DEBUG: Mensagem pergunta super-herói apagada: {mensagem_superheroi_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta super-herói: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 AGORA mostrar os tamanhos para Superheroes em GRADE 2xN
        estilo_escolhido = "Superheroes"
        tamanhos_disponiveis = TAMANHOS_POR_ESTILO.get(estilo_escolhido, {})
        
        if not tamanhos_disponiveis:
            # 🔥 MENSAGEM DE ERRO TRADUZIDA
            textos_sem_tamanhos = {
                'portugues': "❌ Nenhum tamanho disponível para Superheroes.",
                'ingles': "❌ No sizes available for Superheroes.",
                'espanhol': "❌ No hay tamaños disponibles para Superhéroes.",
                'italiano': "❌ Nessuna dimensione disponibile per Supereroi.",
                'alemao': "❌ Keine Größen für Superhelden verfügbar.",
                'frances': "❌ Aucune taille disponible pour Super-héros."
            }
            
            await update.message.reply_text(textos_sem_tamanhos.get(idioma, textos_sem_tamanhos['portugues']))
            return

        # 🔥 TEXTO "ESCOLHER TAMANHO" POR IDIOMA
        textos_tamanho = {
            'portugues': "📏 *Perfeito! Agora escolhe o tamanho do teu Cartoon Superheroes:*",
            'ingles': "📏 *Perfect! Now choose the size of your Superheroes Cartoon:*",
            'espanhol': "📏 *¡Perfecto! Ahora elige el tamaño de tu Caricatura de Superhéroes:*",
            'italiano': "📏 *Perfetto! Ora scegli la dimensione del tuo Cartoon Supereroi:*",
            'alemao': "📏 *Perfekt! Wähle jetzt die Größe deiner Superhelden-Karikatur:*",
            'frances': "📏 *Parfait ! Maintenant choisissez la taille de votre Dessin Animé Super-héros :*"
        }

        # Criar botões dos tamanhos em GRADE 2xN
        teclado = []
        tamanhos_lista = list(tamanhos_disponiveis.items())
        
        # Processar em pares (2 botões por linha)
        for i in range(0, len(tamanhos_lista), 2):
            linha = []
            # Primeiro botão da linha
            tamanho_key1, info_tamanho1 = tamanhos_lista[i]
            botao_texto1 = f"{info_tamanho1['nome']}"
            linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
            
            # Segundo botão da linha (se existir)
            if i + 1 < len(tamanhos_lista):
                tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
                botao_texto2 = f"{info_tamanho2['nome']}"
                linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
            
            teclado.append(linha)

        mensagem_tamanhos = await update.message.reply_text(
            textos_tamanho.get(idioma, textos_tamanho['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        
        context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
        context.user_data['conversation_state'] = TAMANHO
        print(f"✅ Estado: TAMANHO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_super_heroi: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))












# --- Enviar resumo ---
async def enviar_resumo(context, chat_id, message_id=None):
    """Enviar ou atualizar o resumo do pedido - COM TRADUÇÃO E PAÍSES EM INGLÊS"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        # 🔥 DICIONÁRIO PARA CONVERTER PAÍSES PARA INGLÊS
        PAISES_PARA_INGLES = {
            # callback_data → Nome em inglês para o resumo
            'estados_unidos': 'United States',
            'canada': 'Canada',
            'reino_unido': 'United Kingdom',
            'brasil': 'Brazil',
            'alemanha': 'Germany',
            'paises_baixos': 'Netherlands',
            'holanda': 'Netherlands',
            'franca': 'France',
            'espanha': 'Spain',
            'belgica': 'Belgium',
            'italia': 'Italy',
            'portugal': 'Portugal',
            'irlanda': 'Ireland',
            'luxemburgo': 'Luxembourg'
        }
        
        def converter_pais_para_ingles(pais_key):
            """Converte o nome/callback do país para inglês"""
            if isinstance(pais_key, str):
                # Remove "pais_" se existir
                if pais_key.startswith('pais_'):
                    pais_key = pais_key[5:]
                # Remove acentos e converte para minúsculas para comparação
                pais_clean = pais_key.lower()
                # Mapeamento adicional para nomes em português
                mapeamento = {
                    'bélgica': 'belgica',
                    'bélgica (português)': 'belgica',
                    'frança': 'franca',
                    'espanha': 'espanha',
                    'alemanha': 'alemanha',
                    'itália': 'italia',
                    'irlanda': 'irlanda',
                    'luxemburgo': 'luxemburgo',
                    'países baixos': 'paises_baixos',
                    'holanda': 'paises_baixos',
                    'reino unido': 'reino_unido',
                    'estados unidos': 'estados_unidos',
                    'eua': 'estados_unidos'
                }
                pais_key = mapeamento.get(pais_clean, pais_key)
            return PAISES_PARA_INGLES.get(pais_key, pais_key.title())
        
        # 🔥 DICIONÁRIOS DE TRADUÇÃO
        textos_titulo = {
            'portugues': "📋 *RESUMO DO SEU PEDIDO*",
            'ingles': "📋 *YOUR ORDER SUMMARY*",
            'espanhol': "📋 *RESUMEN DE TU PEDIDO*",
            'italiano': "📋 *RIEPILOGO DEL TUO ORDINE*",
            'alemao': "📋 *DEINE BESTELLÜBERSICHT*",
            'frances': "📋 *RÉSUMÉ DE VOTRE COMMANDE*"
        }
        
        textos_continuar = {
            'portugues': "*Continue preenchendo abaixo* ⬇️",
            'ingles': "*Continue filling in below* ⬇️",
            'espanhol': "*Continúe completando a continuación* ⬇️",
            'italiano': "*Continua a compilare qui sotto* ⬇️",
            'alemao': "*Fahren Sie unten fort* ⬇️",
            'frances': "*Continuez à remplir ci-dessous* ⬇️"
        }
        
        textos_campos = {
            'portugues': {
                'nome': "👤 *Nome:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *País:*",
                'contacto': "📱 *Telefone:*",
                'tipo_cartoon': "🎨 *Tipo de Cartoon:*",
                'estilo_cartoon': "🖌 *Estilo:*",
                'nome_cartoon': "🎭 *Nome no Cartoon:*",
                'frase_cartoon': "💬 *Frase na Box:*",
                'tipo_personalizado': "📦 *Tipo de Peça:*",
                'nome_peca_personalizado': "📝 *Nome da Peça:*",
                'nome_personalizado': "🎭 *Nome do Cartoon:*",
                'frase_personalizado': "💬 *Frase do Elemento:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nome da Família:*",
                'frase_family': "💬 *Frase da Família:*",
                'nome_animal': "🐾 *Nome do Animal:*",
                'tipo_animal': "🐕 *Tipo de Animal:*",
                'profissao': "💼 *Profissão:*",
                'objetos_office': "🎯 *Objetos Personalizados:*",
                'super_heroi': "🦸‍♂️ *Super-Herói:*",
                'elementos_family': "👥 *Total de Elementos:*",
                'adultos_family': "👨‍👩 *Adultos:*",
                'criancas_family': "👧🧒 *Crianças:*",
                'animais_family': "🐱🐶 *Animais:*",
                'tamanho_cartoon': "📏 *Tamanho:*",
                'foto': "📸 *Foto:*"
            },
            'ingles': {
                'nome': "👤 *Name:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Country:*",
                'contacto': "📱 *Phone:*",
                'tipo_cartoon': "🎨 *Cartoon Type:*",
                'estilo_cartoon': "🖌 *Style:*",
                'nome_cartoon': "🎭 *Name on Cartoon:*",
                'frase_cartoon': "💬 *Box Phrase:*",
                'tipo_personalizado': "📦 *Piece Type:*",
                'nome_peca_personalizado': "📝 *Piece Name:*",
                'nome_personalizado': "🎭 *Cartoon Name:*",
                'frase_personalizado': "💬 *Element Phrase:*",
                'nome_family': "👨‍👩‍👧‍👦 *Family Name:*",
                'frase_family': "💬 *Family Phrase:*",
                'nome_animal': "🐾 *Animal Name:*",
                'tipo_animal': "🐕 *Animal Type:*",
                'profissao': "💼 *Profession:*",
                'objetos_office': "🎯 *Custom Objects:*",
                'super_heroi': "🦸‍♂️ *Superhero:*",
                'elementos_family': "👥 *Total Elements:*",
                'adultos_family': "👨‍👩 *Adults:*",
                'criancas_family': "👧🧒 *Children:*",
                'animais_family': "🐱🐶 *Animals:*",
                'tamanho_cartoon': "📏 *Size:*",
                'foto': "📸 *Photo:*"
            },
            'espanhol': {
                'nome': "👤 *Nombre:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *País:*",
                'contacto': "📱 *Teléfono:*",
                'tipo_cartoon': "🎨 *Tipo de Caricatura:*",
                'estilo_cartoon': "🖌 *Estilo:*",
                'nome_cartoon': "🎭 *Nombre en la Caricatura:*",
                'frase_cartoon': "💬 *Frase en la Caja:*",
                'tipo_personalizado': "📦 *Tipo de Pieza:*",
                'nome_peca_personalizado': "📝 *Nombre de la Pieza:*",
                'nome_personalizado': "🎭 *Nombre de la Caricatura:*",
                'frase_personalizado': "💬 *Frase del Elemento:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nombre de la Familia:*",
                'frase_family': "💬 *Frase de la Familia:*",
                'nome_animal': "🐾 *Nombre del Animal:*",
                'tipo_animal': "🐕 *Tipo de Animal:*",
                'profissao': "💼 *Profesión:*",
                'objetos_office': "🎯 *Objetos Personalizados:*",
                'super_heroi': "🦸‍♂️ *Superhéroe:*",
                'elementos_family': "👥 *Total de Elementos:*",
                'adultos_family': "👨‍👩 *Adultos:*",
                'criancas_family': "👧🧒 *Niños:*",
                'animais_family': "🐱🐶 *Animales:*",
                'tamanho_cartoon': "📏 *Tamaño:*",
                'foto': "📸 *Foto:*"
            },
            'italiano': {
                'nome': "👤 *Nome:*",
                'email': "📧 *Email:*",
                'pais': "🌍 *Paese:*",
                'contacto': "📱 *Telefono:*",
                'tipo_cartoon': "🎨 *Tipo di Cartoon:*",
                'estilo_cartoon': "🖌 *Stile:*",
                'nome_cartoon': "🎭 *Nome sul Cartoon:*",
                'frase_cartoon': "💬 *Frase sulla Scatola:*",
                'tipo_personalizado': "📦 *Tipo di Pezzo:*",
                'nome_peca_personalizado': "📝 *Nome del Pezzo:*",
                'nome_personalizado': "🎭 *Nome del Cartoon:*",
                'frase_personalizado': "💬 *Frase dell'Elemento:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nome della Famiglia:*",
                'frase_family': "💬 *Frase della Famiglia:*",
                'nome_animal': "🐾 *Nome dell'Animale:*",
                'tipo_animal': "🐕 *Tipo di Animale:*",
                'profissao': "💼 *Professione:*",
                'objetos_office': "🎯 *Oggetti Personalizzati:*",
                'super_heroi': "🦸‍♂️ *Supereroe:*",
                'elementos_family': "👥 *Totale Elementi:*",
                'adultos_family': "👨‍👩 *Adulti:*",
                'criancas_family': "👧🧒 *Bambini:*",
                'animais_family': "🐱🐶 *Animali:*",
                'tamanho_cartoon': "📏 *Dimensione:*",
                'foto': "📸 *Foto:*"
            },
            'alemao': {
                'nome': "👤 *Name:*",
                'email': "📧 *E-Mail:*",
                'pais': "🌍 *Land:*",
                'contacto': "📱 *Telefon:*",
                'tipo_cartoon': "🎨 *Karikaturtyp:*",
                'estilo_cartoon': "🖌 *Stil:*",
                'nome_cartoon': "🎭 *Name auf der Karikatur:*",
                'frase_cartoon': "💬 *Box-Satz:*",
                'tipo_personalizado': "📦 *Stücktyp:*",
                'nome_peca_personalizado': "📝 *Stückname:*",
                'nome_personalizado': "🎭 *Karikaturname:*",
                'frase_personalizado': "💬 *Elementsatz:*",
                'nome_family': "👨‍👩‍👧‍👦 *Familienname:*",
                'frase_family': "💬 *Familiensatz:*",
                'nome_animal': "🐾 *Tiername:*",
                'tipo_animal': "🐕 *Tierart:*",
                'profissao': "💼 *Beruf:*",
                'objetos_office': "🎯 *Benutzerdefinierte Objekte:*",
                'super_heroi': "🦸‍♂️ *Superheld:*",
                'elementos_family': "👥 *Gesamtelemente:*",
                'adultos_family': "👨‍👩 *Erwachsene:*",
                'criancas_family': "👧🧒 *Kinder:*",
                'animais_family': "🐱🐶 *Tiere:*",
                'tamanho_cartoon': "📏 *Größe:*",
                'foto': "📸 *Foto:*"
            },
            'frances': {
                'nome': "👤 *Nom:*",
                'email': "📧 *E-mail:*",
                'pais': "🌍 *Pays:*",
                'contacto': "📱 *Téléphone:*",
                'tipo_cartoon': "🎨 *Type de Dessin Animé:*",
                'estilo_cartoon': "🖌 *Style:*",
                'nome_cartoon': "🎭 *Nom sur le Dessin Animé:*",
                'frase_cartoon': "💬 *Phrase sur la Boîte:*",
                'tipo_personalizado': "📦 *Type de Pièce:*",
                'nome_peca_personalizado': "📝 *Nom de la Pièce:*",
                'nome_personalizado': "🎭 *Nom du Dessin Animé:*",
                'frase_personalizado': "💬 *Phrase de l'Élément:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nom de Famille:*",
                'frase_family': "💬 *Phrase de Famille:*",
                'nome_animal': "🐾 *Nom de l'Animal:*",
                'tipo_animal': "🐕 *Type d'Animal:*",
                'profissao': "💼 *Profession:*",
                'objetos_office': "🎯 *Objets Personnalisés:*",
                'super_heroi': "🦸‍♂️ *Super-héros:*",
                'elementos_family': "👥 *Total des Éléments:*",
                'adultos_family': "👨‍👩 *Adultes:*",
                'criancas_family': "👧🧒 *Enfants:*",
                'animais_family': "🐱🐶 *Animaux:*",
                'tamanho_cartoon': "📏 *Taille:*",
                'foto': "📸 *Photo:*"
            }
        }
        
        # Pegar textos no idioma correto
        campos = textos_campos.get(idioma, textos_campos['portugues'])
        titulo = textos_titulo.get(idioma, textos_titulo['portugues'])
        continuar = textos_continuar.get(idioma, textos_continuar['portugues'])
        
        # Construir resumo
        resumo = f"{titulo}\n\n"
        user_data = context.user_data

        # Campos existentes com tradução
        if "nome" in user_data:
            resumo += f"{campos['nome']} {user_data['nome']}\n"
        if "email" in user_data:
            resumo += f"{campos['email']} {user_data['email']}\n"
        if "pais" in user_data:
            # 🔥 CONVERTER PAÍS PARA INGLÊS
            pais_original = user_data['pais']
            pais_ingles = converter_pais_para_ingles(pais_original)
            resumo += f"{campos['pais']} {pais_ingles}\n"
        if "contacto" in user_data:
            resumo += f"{campos['contacto']} {user_data['contacto']}\n"
        if "tipo_cartoon" in user_data:
            resumo += f"{campos['tipo_cartoon']} {user_data['tipo_cartoon']}\n"
        if "estilo_cartoon" in user_data:
            resumo += f"{campos['estilo_cartoon']} {user_data['estilo_cartoon']}\n"

        # CAMPOS PARA PERSONALIZAÇÃO DA BOX
        if "nome_cartoon" in user_data:
            resumo += f"{campos['nome_cartoon']} {user_data['nome_cartoon']}\n"
        if "frase_cartoon" in user_data:
            if user_data["frase_cartoon"] != "Não adicionou frase":
                resumo += f"{campos['frase_cartoon']} {user_data['frase_cartoon']}\n"

        # CAMPOS PARA PERSONALIZADO
        if "tipo_personalizado" in user_data:
            resumo += f"{campos['tipo_personalizado']} {user_data['tipo_personalizado']}\n"
        if "nome_peca_personalizado" in user_data:
            resumo += f"{campos['nome_peca_personalizado']} {user_data['nome_peca_personalizado']}\n"
        if "nome_personalizado" in user_data:
            resumo += f"{campos['nome_personalizado']} {user_data['nome_personalizado']}\n"
        if "frase_personalizado" in user_data and user_data['frase_personalizado'] != "Não adicionou frase":
            resumo += f"{campos['frase_personalizado']} \"{user_data['frase_personalizado']}\"\n"

        # CAMPOS PARA FAMILY
        if "nome_family" in user_data:
            resumo += f"{campos['nome_family']} {user_data['nome_family']}\n"
        if "frase_family" in user_data and user_data['frase_family'] != "Não adicionou frase":
            resumo += f"{campos['frase_family']} \"{user_data['frase_family']}\"\n"

        # CAMPOS PARA ANIMAL
        if "nome_animal" in user_data:
            resumo += f"{campos['nome_animal']} {user_data['nome_animal']}\n"
        if "tipo_animal" in user_data:
            resumo += f"{campos['tipo_animal']} {user_data['tipo_animal']}\n"
        
        # OUTROS CAMPOS PERSONALIZADOS
        if "profissao" in user_data:
            resumo += f"{campos['profissao']} {user_data['profissao']}\n"
        if "objetos_office" in user_data:
            resumo += f"{campos['objetos_office']} {user_data['objetos_office']}\n"
        if "super_heroi" in user_data:
            resumo += f"{campos['super_heroi']} {user_data['super_heroi']}\n"
        if "elementos_family" in user_data:
            resumo += f"{campos['elementos_family']} {user_data['elementos_family']}\n"
        if "adultos_family" in user_data:
            resumo += f"{campos['adultos_family']} {user_data['adultos_family']}\n"
        if "criancas_family" in user_data:
            resumo += f"{campos['criancas_family']} {user_data['criancas_family']}\n"
        if "animais_family" in user_data:
            resumo += f"{campos['animais_family']} {user_data['animais_family']}\n"
        
        if "tamanho_cartoon" in user_data:
            resumo += f"{campos['tamanho_cartoon']} {user_data['tamanho_cartoon']}\n"
        if "foto_id" in user_data:
            resumo += f"{campos['foto']} ✅ Recebida\n"

        resumo += f"\n{continuar}"
        
        # 🔥 Se temos um message_id, tentar editar a mensagem existente
        if message_id:
            try:
                await context.bot.edit_message_text(
                    chat_id=chat_id,
                    message_id=message_id,
                    text=resumo,
                    parse_mode="Markdown"
                )
                print(f"✅ Resumo editado | Idioma: {idioma}")
                return message_id
            except Exception as e:
                # Se falhar ao editar, enviar nova
                print(f"Erro ao editar mensagem: {e}")
                msg = await context.bot.send_message(
                    chat_id=chat_id, 
                    text=resumo, 
                    parse_mode="Markdown"
                )
                return msg.message_id
        else:
            # Enviar nova mensagem
            msg = await context.bot.send_message(
                chat_id=chat_id, 
                text=resumo, 
                parse_mode="Markdown"
            )
            print(f"✅ Novo resumo enviado | Idioma: {idioma}")
            return msg.message_id
            
    except Exception as e:
        print(f"Erro crítico em enviar_resumo: {e}")
        # Tentativa de fallback
        try:
            # Texto de fallback traduzido
            textos_fallback = {
                'portugues': "📋 *Resumo do pedido em atualização...*",
                'ingles': "📋 *Order summary updating...*",
                'espanhol': "📋 *Resumen del pedido actualizando...*",
                'italiano': "📋 *Riepilogo ordine in aggiornamento...*",
                'alemao': "📋 *Bestellübersicht wird aktualisiert...*",
                'frances': "📋 *Résumé de la commande en cours de mise à jour...*"
            }
            
            idioma = context.user_data.get('idioma', 'portugues')
            texto_fallback = textos_fallback.get(idioma, textos_fallback['portugues'])
            
            msg = await context.bot.send_message(
                chat_id=chat_id, 
                text=texto_fallback, 
                parse_mode="Markdown"
            )
            return msg.message_id
        except:
            return None












# --- Tipo de cartoon --- CORRIGIDO
async def cartoon_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 ATUALIZAR CONTADOR DE UTILIZADORES ATIVOS
    user_id = update.effective_user.id
    utilizadores_hoje = atualizar_contador_utilizadores(user_id)
    
    print(f"🎨 USER {user_id} INICIOU CRIAÇÃO DE CARTOON | Tipo: {query.data} | Idioma: {idioma} | Utilizadores hoje: {utilizadores_hoje}")

    # 🔥 TIPOS DE CARTOON TRADUZIDOS
    tipos_por_idioma = {
        'portugues': {
            "cartoon_individual": "Cartoon Individual 😎",
            "cartoon_family": "Cartoon Grupo 👨‍👩‍👧", 
            "cartoon_animal": "Cartoon Animal 🐱",
            "cartoon_custom": "Cartoon Personalizado 🎨"
        },
        'ingles': {
            "cartoon_individual": "Individual Cartoon 😎",
            "cartoon_family": "Group Cartoon 👨‍👩‍👧", 
            "cartoon_animal": "Animal Cartoon 🐱",
            "cartoon_custom": "Custom Cartoon 🎨"
        },
        'espanhol': {
            "cartoon_individual": "Caricatura Individual 😎",
            "cartoon_family": "Caricatura de Grupo 👨‍👩‍👧", 
            "cartoon_animal": "Caricatura Animal 🐱",
            "cartoon_custom": "Caricatura Personalizada 🎨"
        },
        'italiano': {
            "cartoon_individual": "Cartoon Individuale 😎",
            "cartoon_family": "Cartoon Gruppo 👨‍👩‍👧", 
            "cartoon_animal": "Cartoon Animale 🐱",
            "cartoon_custom": "Cartoon Personalizzato 🎨"
        },
        'alemao': {
            "cartoon_individual": "Einzel-Karikatur 😎",
            "cartoon_family": "Gruppen-Karikatur 👨‍👩‍👧", 
            "cartoon_animal": "Tier-Karikatur 🐱",
            "cartoon_custom": "Personalisiert Karikatur 🎨"
        },
        'frances': {
            "cartoon_individual": "Dessin Animé Individuel 😎",
            "cartoon_family": "Dessin Animé de Groupe 👨‍👩‍👧", 
            "cartoon_animal": "Dessin Animé Animal 🐱",
            "cartoon_custom": "Dessin Animé Personnalisé 🎨"
        }
    }
    
    tipos = tipos_por_idioma.get(idioma, tipos_por_idioma['portugues'])
    context.user_data["tipo_cartoon"] = tipos.get(query.data, "")
    
    # Atualizar resumo
    current_resumo_msg_id = context.user_data.get('resumo_msg_id')
    new_message_id = await enviar_resumo(
        context, 
        query.message.chat_id, 
        message_id=current_resumo_msg_id
    )
    context.user_data['resumo_msg_id'] = new_message_id

    # Remover mensagem com botões
    await query.delete_message()

    # 🔥 FLUXO ESPECIAL PARA FAMILY (AGORA COM PERSONALIZAÇÃO)
    if query.data == "cartoon_family":
        # 🔥 TEXTOS "NOME DA FAMÍLIA/GRUPO" POR IDIOMA
        textos_family = {
            'portugues': "👨‍👩‍👧‍👦 *Escreve um apelido, alcunha ou nome para a tua família / amigos!*\n\n*Exemplo:* `Família Silva`, `Os Aventureiros`",
            'ingles': "👨‍👩‍👧‍👦 *Write a nickname or name for your family/friends!*\n\n*Example:* `The Smith Family`, `The Adventurers`",
            'espanhol': "👨‍👩‍👧‍👦 *¡Escribe un apodo o nombre para tu familia/amigos!*\n\n*Ejemplo:* `Familia García`, `Los Aventureros`",
            'italiano': "👨‍👩‍👧‍👦 *Scrivi un soprannome o nome per la tua famiglia/amici!*\n\n*Esempio:* `Famiglia Rossi`, `Gli Avventurieri`",
            'alemao': "👨‍👩‍👧‍👦 *Schreibe einen Spitznamen oder Namen für deine Familie/Freunde!*\n\n*Beispiel:* `Familie Müller`, `Die Abenteurer`",
            'frances': "👨‍👩‍👧‍👦 *Écrivez un surnom ou un nom pour votre famille/amis !*\n\n*Exemple :* `Famille Dupont`, `Les Aventuriers`"
        }
        
        mensagem_nome_family = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_family.get(idioma, textos_family['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_family_id'] = mensagem_nome_family.message_id
        context.user_data['conversation_state'] = NOME_FAMILY
        print(f"✅ Fluxo: FAMILY | Idioma: {idioma}")
    
    # 🔥 FLUXO ESPECIAL PARA ANIMAL
    elif query.data == "cartoon_animal":
        # 🔥 TEXTOS "NOME DO ANIMAL" POR IDIOMA
        textos_animal = {
            'portugues': "🐾 *Perfeito! Escolheste Cartoon Animal.*\n\n📝 *Qual é o nome do animal?*",
            'ingles': "🐾 *Perfect! You chose Animal Cartoon.*\n\n📝 *What is the animal's name?*",
            'espanhol': "🐾 *¡Perfecto! Elegiste Caricatura Animal.*\n\n📝 *¿Cuál es el nombre del animal?*",
            'italiano': "🐾 *Perfetto! Hai scelto Cartoon Animale.*\n\n📝 *Qual è il nome dell'animale?*",
            'alemao': "🐾 *Perfekt! Du hast Tier-Karikatur gewählt.*\n\n📝 *Wie heißt das Tier?*",
            'frances': "🐾 *Parfait ! Vous avez choisi Dessin Animé Animal.*\n\n📝 *Quel est le nom de l'animal ?*"
        }
        
        mensagem_nome_animal = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_animal.get(idioma, textos_animal['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_animal_id'] = mensagem_nome_animal.message_id
        context.user_data['conversation_state'] = NOME_ANIMAL
        print(f"✅ Fluxo: ANIMAL | Idioma: {idioma}")
    
    # 🔥 FLUXO ESPECIAL PARA PERSONALIZADO
    elif query.data == "cartoon_custom":
        # 🔥 TEXTOS "PERSONALIZADO" POR IDIOMA
        textos_personalizado = {
            'portugues': """🎨 *Perfeito! Escolheste Cartoon Personalizado.*

🔄 *A GODSPLAN eterniza as suas memórias em 3D*

Escolha o tipo de peça personalizada:""",
            'ingles': """🎨 *Perfect! You chose Custom Cartoon.*

🔄 *GODSPLAN immortalizes your memories in 3D*

Choose the type of custom piece:""",
            'espanhol': """🎨 *¡Perfecto! Elegiste Caricatura Personalizada.*

🔄 *GODSPLAN inmortaliza tus recuerdos en 3D*

Elige el tipo de pieza personalizada:""",
            'italiano': """🎨 *Perfetto! Hai scelto Cartoon Personalizzato.*

🔄 *GODSPLAN immortalizza i tuoi ricordi in 3D*

Scegli il tipo di pezzo personalizzato:""",
            'alemao': """🎨 *Perfekt! Du hast Personalisierte Karikatur gewählt.*

🔄 *GODSPLAN verewigt deine Erinnerungen in 3D*

Wählen Sie die Art des personalisierten Stücks:""",
            'frances': """🎨 *Parfait ! Vous avez choisi Dessin Animé Personnalisé.*

🔄 *GODSPLAN immortalise vos souvenirs en 3D*

Choisissez le type de pièce personnalisée :"""
        }
        
        # 🔥 BOTÕES TRADUZIDOS PARA PERSONALIZADO
        botoes_personalizado_por_idioma = {
            'portugues': {
                'carro': "🚗 Carro",
                'peluche': "🧸 Peluche",
                'acessorio': "💍 Acessório",
                'outro': "📦 Outro"
            },
            'ingles': {
                'carro': "🚗 Car",
                'peluche': "🧸 Plush Toy",
                'acessorio': "💍 Accessory",
                'outro': "📦 Other"
            },
            'espanhol': {
                'carro': "🚗 Coche",
                'peluche': "🧸 Peluche",
                'acessorio': "💍 Accesorio",
                'outro': "📦 Otro"
            },
            'italiano': {
                'carro': "🚗 Auto",
                'peluche': "🧸 Peluche",
                'acessorio': "💍 Accessorio",
                'outro': "📦 Altro"
            },
            'alemao': {
                'carro': "🚗 Auto",
                'peluche': "🧸 Plüschtier",
                'acessorio': "💍 Accessoire",
                'outro': "📦 Andere"
            },
            'frances': {
                'carro': "🚗 Voiture",
                'peluche': "🧸 Peluche",
                'acessorio': "💍 Accessoire",
                'outro': "📦 Autre"
            }
        }
        
        botoes = botoes_personalizado_por_idioma.get(idioma, botoes_personalizado_por_idioma['portugues'])
        
        teclado = [
           [InlineKeyboardButton(botoes['carro'], callback_data="personalizado_carro"),
           InlineKeyboardButton(botoes['peluche'], callback_data="personalizado_peluche")],
           [InlineKeyboardButton(botoes['acessorio'], callback_data="personalizado_acessorio"),
           InlineKeyboardButton(botoes['outro'], callback_data="personalizado_outro")]
        ]
        
        mensagem_personalizado = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_personalizado.get(idioma, textos_personalizado['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        context.user_data['mensagem_personalizado_id'] = mensagem_personalizado.message_id
        context.user_data['conversation_state'] = TIPO_PERSONALIZADO
        print(f"✅ Fluxo: PERSONALIZADO | Idioma: {idioma}")
    
    # 🔥 FLUXOS EXISTENTES PARA INDIVIDUAL
    elif query.data == "cartoon_individual":
        # 🔥 TEXTOS "ESCOLHER ESTILO" POR IDIOMA
        textos_estilo = {
            'portugues': "🎨 *Escolha o estilo do seu Cartoon Individual:*",
            'ingles': "🎨 *Choose the style of your Individual Cartoon:*",
            'espanhol': "🎨 *Elige el estilo de tu Caricatura Individual:*",
            'italiano': "🎨 *Scegli lo stile del tuo Cartoon Individuale:*",
            'alemao': "🎨 *Wählen Sie den Stil Ihrer Einzel-Karikatur:*",
            'frances': "🎨 *Choisissez le style de votre Dessin Animé Individuel :*"
        }
        
        # Botões mantêm os mesmos nomes em inglês (estilos técnicos)
        teclado = [
          [InlineKeyboardButton("Full Body", callback_data="estilo_fullbody"),
          InlineKeyboardButton("Bust", callback_data="estilo_bust")],
          [InlineKeyboardButton("Voxel", callback_data="estilo_voxel"),
          InlineKeyboardButton("Office", callback_data="estilo_office")],
          [InlineKeyboardButton("Superheroes", callback_data="estilo_superheroes")]
        ]
        
        await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_estilo.get(idioma, textos_estilo['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        context.user_data['conversation_state'] = ESTILO
        print(f"✅ Fluxo: INDIVIDUAL | Idioma: {idioma}")
    
    # 🔥 FLUXO PARA OUTROS TIPOS (fallback)
    else:
        # 🔥 TEXTO "ENVIE FOTO" POR IDIOMA
        textos_foto = {
            'portugues': "📸 Agora envie a foto que deseja transformar em cartoon.",
            'ingles': "📸 Now send the photo you want to transform into a cartoon.",
            'espanhol': "📸 Ahora envía la foto que deseas transformar en caricatura.",
            'italiano': "📸 Ora invia la foto che desideri trasformare in cartoon.",
            'alemao': "📸 Senden Sie jetzt das Foto, das Sie in eine Karikatur verwandeln möchten.",
            'frances': "📸 Maintenant envoyez la photo que vous souhaitez transformer en dessin animé."
        }
        
        await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_foto.get(idioma, textos_foto['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['conversation_state'] = FOTO
        print(f"✅ Fluxo: OUTROS | Estado: FOTO | Idioma: {idioma}")





async def receber_nome_family(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber o nome/apelido da família - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        nome_family = update.message.text
        context.user_data["nome_family"] = nome_family
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem nome family do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta
        mensagem_nome_family_id = context.user_data.get('mensagem_nome_family_id')
        if mensagem_nome_family_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_family_id
                )
                print(f"DEBUG: Mensagem pergunta nome family apagada: {mensagem_nome_family_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta nome family: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTOS "FRASE DA FAMÍLIA" POR IDIOMA
        textos_frase_family = {
            'portugues': """💬 *Escreve uma frase que complete para enquadrar na box da tua família / amigos!*

*Exemplo:* `"Unidos para sempre nos bons momentos"* 
*Ou:* `"A nossa aventura só está a começar"* 
*Ou:* `"O amor é o nosso superpoder"* 

*Pode escrever a sua frase ou enviar* /skip *para continuar:*""",
            
            'ingles': """💬 *Write a phrase to frame in your family/friends box!*

*Example:* `"United forever in good times"* 
*Or:* `"Our adventure is just beginning"* 
*Or:* `"Love is our superpower"* 

*You can write your phrase or send* /skip *to continue:*""",
            
            'espanhol': """💬 *¡Escribe una frase para enmarcar en la caja de tu familia/amigos!*

*Ejemplo:* `"Unidos para siempre en los buenos momentos"* 
*O:* `"Nuestra aventura solo está comenzando"* 
*O:* `"El amor es nuestro superpoder"* 

*Puedes escribir tu frase o enviar* /skip *para continuar:*""",
            
            'italiano': """💬 *Scrivi una frase da incorniciare nella scatola della tua famiglia/amici!*

*Esempio:* `"Uniti per sempre nei bei momenti"* 
*O:* `"La nostra avventura è appena iniziata"* 
*O:* `"L'amore è il nostro superpotere"* 

*Puoi scrivere la tua frase o inviare* /skip *per continuare:*""",
            
            'alemao': """💬 *Schreibe einen Satz, der in deine Familien-/Freundebox gerahmt werden soll!*

*Beispiel:* `"Für immer in guten Zeiten vereint"* 
*Oder:* `"Unser Abenteuer fängt gerade erst an"* 
*Oder:* `"Liebe ist unsere Superkraft"* 

*Sie können Ihren Satz schreiben oder* /skip *senden, um fortzufahren:*""",
            
            'frances': """💬 *Écrivez une phrase à encadrer dans la boîte de votre famille/amis !*

*Exemple :* `"Unis pour toujours dans les bons moments"* 
*Ou :* `"Notre aventure ne fait que commencer"* 
*Ou :* `"L'amour est notre super-pouvoir"* 

*Vous pouvez écrire votre phrase ou envoyer* /skip *pour continuer :*"""
        }
        
        mensagem_frase_family = await update.message.reply_text(
            textos_frase_family.get(idioma, textos_frase_family['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_frase_family_id'] = mensagem_frase_family.message_id
        context.user_data['conversation_state'] = FRASE_FAMILY
        print(f"✅ Estado: FRASE_FAMILY | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_nome_family: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))

# --- Handler para receber frase da Family ---
async def receber_frase_family(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber a frase da família para a box - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        # 🔥 COMANDO DE PULAR POR IDIOMA
        comandos_pular = {
            'portugues': '/skip',
            'ingles': '/skip',
            'espanhol': '/skip',
            'italiano': '/skip',
            'alemao': '/skip',
            'frances': '/skip'
        }
        
        comando_pular = comandos_pular.get(idioma, '/pular')
        
        # Verificar se é comando /pular ou mensagem normal
        if update.message.text and update.message.text.strip() == comando_pular:
            frase_family = "Não adicionou frase"
            print(f"✅ Usuário escolheu pular a frase da family | Idioma: {idioma}")
        else:
            frase_family = update.message.text
            print(f"✅ Usuário adicionou frase da family: {frase_family} | Idioma: {idioma}")
        
        context.user_data["frase_family"] = frase_family
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem frase family do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta
        mensagem_frase_family_id = context.user_data.get('mensagem_frase_family_id')
        if mensagem_frase_family_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_frase_family_id
                )
                print(f"DEBUG: Mensagem pergunta frase family apagada: {mensagem_frase_family_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta frase family: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTOS "PERGUNTAR ELEMENTOS DA FAMÍLIA" POR IDIOMA
        textos_elementos = {
            'portugues': "👨‍👩‍👧‍👦 *Perfeito! Agora, quantos elementos vão estar no seu cartoon 3D, olhando para a imagem que irá enviar?*",
            'ingles': "👨‍👩‍👧‍👦 *Perfect! Now, how many elements will be in your 3D cartoon, looking at the image you will send?*",
            'espanhol': "👨‍👩‍👧‍👦 *¡Perfecto! Ahora, ¿cuántos elementos habrá en tu caricatura 3D, mirando la imagen que enviarás?*",
            'italiano': "👨‍👩‍👧‍👦 *Perfetto! Ora, quanti elementi saranno nel tuo cartoon 3D, guardando l'immagine che invierai?*",
            'alemao': "👨‍👩‍👧‍👦 *Perfekt! Wie viele Elemente werden in Ihrer 3D-Karikatur sein, wenn Sie sich das Bild ansehen, das Sie senden werden?*",
            'frances': "👨‍👩‍👧‍👦 *Parfait ! Maintenant, combien d'éléments y aura-t-il dans votre dessin animé 3D, en regardant l'image que vous enverrez ?*"
        }
        
        mensagem_elementos = await update.message.reply_text(
            textos_elementos.get(idioma, textos_elementos['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_elementos_id'] = mensagem_elementos.message_id
        context.user_data['conversation_state'] = ELEMENTOS_FAMILY
        print(f"✅ Estado: ELEMENTOS_FAMILY | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_frase_family: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))







# --- Escolha do estilo --- COM TRADUÇÃO
async def estilo_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    estilos = {
        "estilo_fullbody": "Full Body",
        "estilo_bust": "Bust",
        "estilo_voxel": "Voxel", 
        "estilo_office": "Office",
        "estilo_superheroes": "Superheroes"
    }

    estilo_escolhido = estilos.get(query.data, "")
    context.user_data["estilo_cartoon"] = estilo_escolhido
    
    # 🔥 OBTER IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 LOG CRÍTICO PARA DEBUG
    tipo_cartoon = context.user_data.get("tipo_cartoon", "NÃO DEFINIDO")
    print(f"🔍 DEBUG estilo_handler:")
    print(f"   • Tipo cartoon: '{tipo_cartoon}'")
    print(f"   • Estilo escolhido: '{estilo_escolhido}'")
    print(f"   • Idioma: {idioma}")
    
    # 🔥 LISTA COMPLETA DE TODOS OS NOMES POSSÍVEIS PARA "CARTOON INDIVIDUAL"
    nomes_individual = [
        # Português
        "Cartoon Individual 😎",
        "cartoon_individual",  # também pode ser salvo assim
        
        # Inglês
        "Individual Cartoon 😎", 
        "individual_cartoon",
        
        # Espanhol
        "Caricatura Individual 😎",
        "caricatura_individual",
        
        # Italiano
        "Cartoon Individuale 😎",
        "cartoon_individuale",
        
        # Alemão (AGORA CORRIGIDO)
        "Individueller Cartoon 😎",
        "Einzel-Karikatur 😎",  # ← ESTE ESTÁ FALTANDO!
        "individueller_cartoon",
        
        # Francês
        "Dessin Animé Individuel 😎",
        "dessin_animé_individuel",
    ]
    
    # VERIFICAR SE É INDIVIDUAL
    is_individual = tipo_cartoon in nomes_individual
    print(f"   • É individual? {is_individual}")
    print(f"   • Tipo '{tipo_cartoon}' está na lista? {'SIM' if tipo_cartoon in nomes_individual else 'NÃO'}")
    if tipo_cartoon not in nomes_individual:
        print(f"   • Valores na lista: {nomes_individual}")
    
    # Atualizar resumo
    current_resumo_msg_id = context.user_data.get('resumo_msg_id')
    new_message_id = await enviar_resumo(
        context, 
        query.message.chat_id, 
        message_id=current_resumo_msg_id
    )
    context.user_data['resumo_msg_id'] = new_message_id

    # Remover mensagem com botões
    await query.delete_message()

    # 🔥 FLUXO ESPECIAL PARA INDIVIDUAL
    if is_individual and estilo_escolhido in ["Full Body", "Bust", "Voxel", "Office", "Superheroes"]:
        print(f"🎯 INDO PARA FLUXO INDIVIDUAL (nome/alcunha)")
        
        textos_nome = {
            'portugues': "🎭 *Escreva o nome pessoal ou alcunha que irá representar o cartoon 3D, para a personalização da sua box!*\n\n",
            'ingles': "🎭 *Write the personal name or nickname that will represent the 3D cartoon, for your box personalization!*\n\n",
            'espanhol': "🎭 *¡Escribe el nome personal o apodo que representará la caricatura 3D, para la personalización de tu caja!*\n\n",
            'italiano': "🎭 *Scrivi il nome personale o soprannome che rappresenterà il cartoon 3D, per la personalizzazione della tua scatola!*\n\n",
            'alemao': "🎭 *Schreiben Sie den persönlichen Namen oder Spitznamen, der den 3D-Cartoon für die Personalisierung Ihrer Box darstellen wird!*\n\n",
            'frances': "🎭 *Écrivez le nome personnel ou le surnom qui représentera le dessin animé 3D, pour la personnalisation de votre boîte !*\n\n"
        }
        
        mensagem_nome_cartoon = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_nome.get(idioma, textos_nome['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_cartoon_id'] = mensagem_nome_cartoon.message_id
        context.user_data['conversation_state'] = NOME_CARTOON
    
    else:
        print(f"🎯 INDO PARA FLUXO NORMAL (tamanhos direto)")
        
        # MOSTRAR TAMANHOS ESPECÍFICOS PARA O ESTILO ESCOLHIDO
        tamanhos_disponiveis = TAMANHOS_POR_ESTILO.get(estilo_escolhido, {})
        
        if not tamanhos_disponiveis:
            textos_sem_tamanho = {
                'portugues': "❌ Nenhum tamanho disponível para este estilo.",
                'ingles': "❌ No sizes available for this style.",
                'espanhol': "❌ No hay tamaños disponibles para este estilo.",
                'italiano': "❌ Nessuna dimensione disponibile per questo stile.",
                'alemao': "❌ Keine Größen für diesen Stil verfügbar.",
                'frances': "❌ Aucune taille disponible pour ce style."
            }
            
            await context.bot.send_message(
                chat_id=query.message.chat_id,
                text=textos_sem_tamanho.get(idioma, textos_sem_tamanho['portugues'])
            )
            return

        # Criar botões dinamicamente em GRADE 2xN
        teclado = []
        tamanhos_lista = list(tamanhos_disponiveis.items())
        
        # Processar em pares (2 botões por linha)
        for i in range(0, len(tamanhos_lista), 2):
            linha = []
            # Primeiro botão da linha
            tamanho_key1, info_tamanho1 = tamanhos_lista[i]
            botao_texto1 = f"{info_tamanho1['nome']}"
            linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
            
            # Segundo botão da linha (se existir)
            if i + 1 < len(tamanhos_lista):
                tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
                botao_texto2 = f"{info_tamanho2['nome']}"
                linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
            
            teclado.append(linha)

        # Texto para escolha de tamanho por idioma
        textos_tamanho = {
            'portugues': f"📏 *Escolha o tamanho para {estilo_escolhido}:*",
            'ingles': f"📏 *Choose the size for {estilo_escolhido}:*",
            'espanhol': f"📏 *Elige el tamaño para {estilo_escolhido}:*",
            'italiano': f"📏 *Scegli la dimensione per {estilo_escolhido}:*",
            'alemao': f"📏 *Wählen Sie die Größe für {estilo_escolhido}:*",
            'frances': f"📏 *Choisissez la taille pour {estilo_escolhido} :*"
        }

        mensagem_tamanhos = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_tamanho.get(idioma, textos_tamanho['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
        context.user_data['conversation_state'] = TAMANHO



async def receber_nome_cartoon(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber o nome/alcunha para personalização da box - para todos os estilos do Individual - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        nome_cartoon = update.message.text
        context.user_data["nome_cartoon"] = nome_cartoon
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem nome cartoon do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta
        mensagem_nome_cartoon_id = context.user_data.get('mensagem_nome_cartoon_id')
        if mensagem_nome_cartoon_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_cartoon_id
                )
                print(f"DEBUG: Mensagem pergunta nome cartoon apagada: {mensagem_nome_cartoon_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta nome cartoon: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 VERIFICAR QUAL ESTILO PARA SABER O PRÓXIMO PASSO
        estilo_escolhido = context.user_data.get("estilo_cartoon", "")
        
        # 🔥 FLUXOS DIFERENCIADOS POR ESTILO
        if estilo_escolhido == "Full Body":
            # FULL BODY: Nome → Frase → Tamanhos
            textos_frase_cartoon = {
                'portugues': "💬 *Quer partilhar alguma frase para juntar na box?*\n\n*Pode escrever a sua frase ou simplesmente enviar* /skip *para continuar:*",
                'ingles': "💬 *Want to share a phrase to include in the box?*\n\n*You can write your phrase or simply send* /skip *to continue:*",
                'espanhol': "💬 *¿Quieres compartir alguna frase para incluir en la caja?*\n\n*Puedes escribir tu frase o simplemente enviar* /skip *para continuar:*",
                'italiano': "💬 *Vuoi condividere una frase da includere nella scatola?*\n\n*Puoi scrivere la tua frase o semplicemente inviare* /skip *per continuare:*",
                'alemao': "💬 *Möchten Sie einen Satz teilen, der in der Box enthalten sein soll?*\n\n*Sie können Ihren Satz schreiben oder einfach* /skip *senden, um fortzufahren:*",
                'frances': "💬 *Voulez-vous partager une phrase à inclure dans la boîte ?*\n\n*Vous pouvez écrire votre phrase ou simplement envoyer* /skip *pour continuer :*"
            }
            
            mensagem_frase_cartoon = await update.message.reply_text(
                textos_frase_cartoon.get(idioma, textos_frase_cartoon['portugues']),
                parse_mode="Markdown"
            )
            context.user_data['mensagem_frase_cartoon_id'] = mensagem_frase_cartoon.message_id
            context.user_data['conversation_state'] = FRASE_CARTOON
            print(f"✅ Estilo: {estilo_escolhido} | Fluxo: FRASE_CARTOON | Idioma: {idioma}")
            
        elif estilo_escolhido == "Office":
            # OFFICE: Nome → Profissão
            textos_profissao = {
                'portugues': "💼 *Perfeito! Agora, qual é a sua profissão?*",
                'ingles': "💼 *Perfect! Now, what is your profession?*",
                'espanhol': "💼 *¡Perfecto! Ahora, ¿cuál es tu profesión?*",
                'italiano': "💼 *Perfetto! Ora, qual è la tua professione?*",
                'alemao': "💼 *Perfekt! Nun, was ist Ihr Beruf?*",
                'frances': "💼 *Parfait ! Maintenant, quelle est votre profession ?*"
            }
            
            mensagem_profissao = await update.message.reply_text(
                textos_profissao.get(idioma, textos_profissao['portugues']),
                parse_mode="Markdown"
            )
            context.user_data['mensagem_profissao_id'] = mensagem_profissao.message_id
            context.user_data['conversation_state'] = PROFISSAO
            print(f"✅ Estilo: {estilo_escolhido} | Fluxo: PROFISSAO | Idioma: {idioma}")
            
        elif estilo_escolhido == "Superheroes":
            # SUPERHEROES: Nome → Super-herói
            textos_superheroi = {
                'portugues': "🦸‍♂️ *Perfeito! Agora, escreva o super-herói que gostaria de colocar no seu cartoon 3D:*",
                'ingles': "🦸‍♂️ *Perfect! Now, write the superhero you would like to put in your 3D cartoon:*",
                'espanhol': "🦸‍♂️ *¡Perfecto! Ahora, escribe el superhéroe que te gustaría poner en tu caricatura 3D:*",
                'italiano': "🦸‍♂️ *Perfetto! Ora, scrivi il supereroe che vorresti mettere nel tuo cartoon 3D:*",
                'alemao': "🦸‍♂️ *Perfekt! Schreiben Sie jetzt den Superhelden, den Sie in Ihren 3D-Cartoon setzen möchten:*",
                'frances': "🦸‍♂️ *Parfait ! Maintenant, écrivez le super-héros que vous aimeriez mettre dans votre dessin animé 3D :*"
            }
            
            mensagem_superheroi = await update.message.reply_text(
                textos_superheroi.get(idioma, textos_superheroi['portugues']),
                parse_mode="Markdown"
            )
            context.user_data['mensagem_superheroi_id'] = mensagem_superheroi.message_id
            context.user_data['conversation_state'] = SUPER_HEROI
            print(f"✅ Estilo: {estilo_escolhido} | Fluxo: SUPER_HEROI | Idioma: {idioma}")
            
        else:
            # BUST e VOXEL: Nome → Tamanhos
            tamanhos_disponiveis = TAMANHOS_POR_ESTILO.get(estilo_escolhido, {})
            
            if not tamanhos_disponiveis:
                # 🔥 MENSAGEM DE ERRO TRADUZIDA
                textos_sem_tamanhos = {
                    'portugues': "❌ Nenhum tamanho disponível para este estilo.",
                    'ingles': "❌ No sizes available for this style.",
                    'espanhol': "❌ No hay tamaños disponibles para este estilo.",
                    'italiano': "❌ Nessuna dimensione disponibile per questo stile.",
                    'alemao': "❌ Keine Größen für diesen Stil verfügbar.",
                    'frances': "❌ Aucune taille disponible pour ce style."
                }
                
                await update.message.reply_text(textos_sem_tamanhos.get(idioma, textos_sem_tamanhos['portugues']))
                return

            # 🔥 TEXTO "ESCOLHER TAMANHO" POR IDIOMA
            textos_tamanho = {
                'portugues': f"📏 *Perfeito! Agora escolhe o tamanho do teu Cartoon {estilo_escolhido}:*",
                'ingles': f"📏 *Perfect! Now choose the size of your {estilo_escolhido} Cartoon:*",
                'espanhol': f"📏 *¡Perfecto! Ahora elige el tamaño de tu Caricatura {estilo_escolhido}:*",
                'italiano': f"📏 *Perfetto! Ora scegli la dimensione del tuo Cartoon {estilo_escolhido}:*",
                'alemao': f"📏 *Perfekt! Wähle jetzt die Größe deiner {estilo_escolhido}-Karikatur:*",
                'frances': f"📏 *Parfait ! Maintenant choisissez la taille de votre Dessin Animé {estilo_escolhido} :*"
            }

            # Criar botões dos tamanhos em GRADE 2xN
            teclado = []
            tamanhos_lista = list(tamanhos_disponiveis.items())
            
            # Processar em pares (2 botões por linha)
            for i in range(0, len(tamanhos_lista), 2):
                linha = []
                # Primeiro botão da linha
                tamanho_key1, info_tamanho1 = tamanhos_lista[i]
                botao_texto1 = f"{info_tamanho1['nome']}"
                linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
                
                # Segundo botão da linha (se existir)
                if i + 1 < len(tamanhos_lista):
                    tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
                    botao_texto2 = f"{info_tamanho2['nome']}"
                    linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
                
                teclado.append(linha)

            mensagem_tamanhos = await update.message.reply_text(
                textos_tamanho.get(idioma, textos_tamanho['portugues']),
                parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup(teclado)
            )
            
            context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
            context.user_data['conversation_state'] = TAMANHO
            print(f"✅ Estilo: {estilo_escolhido} | Fluxo: TAMANHO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_nome_cartoon: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))





async def receber_frase_cartoon(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber a frase para registar na box"""
    try:
        # 🔥 OBTER IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        # 🔥 NORMALIZAR O TEXTO RECEBIDO
        texto_recebido = update.message.text.strip() if update.message.text else ""
        
        # Verificar se é comando /skip ou qualquer texto que indique "pular"
        # Aceitamos: /skip, skip, pular, saltar, etc. (flexível)
        texto_lower = texto_recebido.lower()
        
        # 🔥 PALAVRAS-CHAVE QUE INDICAM "PULAR" EM DIFERENTES IDIOMAS
        palavras_pular = [
            # Português
            "/skip", "/skip", "/pular", "pular", "saltar", "passar",
            # Inglês
            "/skip", "jump", "pass",
            # Espanhol
            "/saltar", "pasar", "/skip", 
            # Italiano
            "salta", "passa", "/skip",
            # Alemão
            "überspringen", "ueberspringen", "springen", "/skip",
            # Francês
            "passer", "saute", "/skip"
        ]
        
        # Verificar se o texto contém alguma palavra de pular
        is_pular = any(palavra in texto_lower for palavra in palavras_pular)
        
        frase_cartoon = ""
        
        print(f"🔍 Texto recebido: '{texto_recebido}' | Idioma: {idioma} | É pular? {is_pular}")
        
        if is_pular:
            # 🔥 TRADUZIR FRASE DE FALLBACK
            frases_sem_frase = {
                'portugues': "Não adicionou frase",
                'ingles': "No phrase added",
                'espanhol': "No añadió frase",
                'italiano': "Non ha aggiunto frase",
                'alemao': "Keinen Satz hinzugefügt",
                'frances': "Aucune phrase ajoutée"
            }
            frase_cartoon = frases_sem_frase.get(idioma, frases_sem_frase['portugues'])
            print(f"✅ Usuário escolheu pular a frase | Idioma: {idioma}")
        else:
            frase_cartoon = texto_recebido
            print(f"✅ Usuário adicionou frase: {frase_cartoon[:50]}... | Idioma: {idioma}")
        
        context.user_data["frase_cartoon"] = frase_cartoon
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta da frase
        mensagem_frase_cartoon_id = context.user_data.get('mensagem_frase_cartoon_id')
        if mensagem_frase_cartoon_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_frase_cartoon_id
                )
                print(f"DEBUG: Mensagem pergunta frase cartoon apagada: {mensagem_frase_cartoon_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta frase cartoon: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 AGORA mostrar os tamanhos para Full Body
        estilo_escolhido = "Full Body"
        tamanhos_disponiveis = TAMANHOS_POR_ESTILO.get(estilo_escolhido, {})
        
        if not tamanhos_disponiveis:
            textos_sem_tamanho = {
                'portugues': "❌ Nenhum tamanho disponível para Full Body.",
                'ingles': "❌ No sizes available for Full Body.",
                'espanhol': "❌ No hay tamaños disponibles para Full Body.",
                'italiano': "❌ Nessuna dimensione disponibile per Full Body.",
                'alemao': "❌ Keine Größen für Full Body verfügbar.",
                'frances': "❌ Aucune taille disponible pour Full Body."
            }
            await update.message.reply_text(textos_sem_tamanho.get(idioma, textos_sem_tamanho['portugues']))
            return

        # Criar botões dos tamanhos em GRADE 2xN
        teclado = []
        tamanhos_lista = list(tamanhos_disponiveis.items())
        
        for i in range(0, len(tamanhos_lista), 2):
            linha = []
            # Primeiro botão da linha
            tamanho_key1, info_tamanho1 = tamanhos_lista[i]
            botao_texto1 = f"{info_tamanho1['nome']}"
            linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
            
            # Segundo botão da linha (se existir)
            if i + 1 < len(tamanhos_lista):
                tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
                botao_texto2 = f"{info_tamanho2['nome']}"
                linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
            
            teclado.append(linha)

        # 🔥 TEXTO PARA ESCOLHA DE TAMANHO POR IDIOMA
        textos_tamanho = {
            'portugues': "📏 *Perfeito! Agora escolhe o tamanho do teu Cartoon Full Body:*",
            'ingles': "📏 *Perfect! Now choose the size of your Full Body Cartoon:*",
            'espanhol': "📏 *¡Perfecto! Ahora elige el tamaño de tu Caricatura Full Body:*",
            'italiano': "📏 *Perfetto! Ora scegli la dimensione del tuo Cartoon Full Body:*",
            'alemao': "📏 *Perfekt! Wählen Sie jetzt die Größe Ihres Full Body-Cartoons:*",
            'frances': "📏 *Parfait ! Maintenant choisissez la taille de votre Dessin Animé Full Body :*"
        }

        mensagem_tamanhos = await update.message.reply_text(
            textos_tamanho.get(idioma, textos_tamanho['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        
        context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
        context.user_data['conversation_state'] = TAMANHO
        
        print(f"✅ Fluxo continuou para tamanhos após frase | Idioma: {idioma}")
        
    except Exception as e:
        print(f"❌ ERRO em receber_frase_cartoon: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un erro. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))





async def pular_frase_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler específico para o comando /skip - ATUALIZADO"""
    print(f"🔍 /skip detectado")
    
    # 🔥 OBTER IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    state = context.user_data.get('conversation_state')
    print(f"DEBUG: Estado atual: {state} | Idioma: {idioma}")
    
    # 🔥 ESTADOS QUE SUPORTAM /skip
    if state == FRASE_CARTOON:
        print(f"✅ /skip no estado FRASE_CARTOON - processando...")
        await receber_frase_cartoon(update, context)
        
    elif state == FRASE_PERSONALIZADO:
        print(f"✅ /skip no estado FRASE_PERSONALIZADO - processando...")
        await receber_frase_personalizado(update, context)
        
    elif state == FRASE_FAMILY:
        print(f"✅ /skip no estado FRASE_FAMILY - processando...")
        await receber_frase_family(update, context)
        
    elif state == GIFT_NOME_BOX:
        print(f"✅ /skip no estado GIFT_NOME_BOX - processando...")
        # 🔥 ADICIONADO: Suporte a /skip para nome do Gift
        await receber_gift_nome_box(update, context)
        
    elif state == GIFT_FRASE_BOX:
        print(f"✅ /skip no estado GIFT_FRASE_BOX - processando...")
        await receber_gift_frase_box(update, context)
        
    else:
        print(f"❌ /skip em estado inválido: {state}")
        
        # 🔥 MENSAGEM DE ERRO POR IDIOMA
        textos_erro = {
            'portugues': "❌ Comando `/skip` não disponível neste momento.",
            'ingles': "❌ Command `/skip` not available at this time.",
            'espanhol': "❌ Comando `/skip` no disponible en este momento.",
            'italiano': "❌ Comando `/skip` non disponibile in questo momento.",
            'alemao': "❌ Befehl `/skip` derzeit nicht verfügbar.",
            'frances': "❌ Commande `/skip` non disponible pour le moment."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues']),
            parse_mode="Markdown"
        )




async def tipo_personalizado_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    print(f"🎯 tipo_personalizado_handler CHAMADO - callback_data: {query.data} | Idioma: {idioma}")
    
    # 🔥 TIPOS PERSONALIZADOS TRADUZIDOS
    tipos_personalizado_por_idioma = {
        'portugues': {
            "personalizado_carro": "Carro 🚗",
            "personalizado_peluche": "Peluche 🧸", 
            "personalizado_acessorio": "Acessório 💍",
            "personalizado_outro": "Outro 📦"
        },
        'ingles': {
            "personalizado_carro": "Car 🚗",
            "personalizado_peluche": "Plush Toy 🧸", 
            "personalizado_acessorio": "Accessory 💍",
            "personalizado_outro": "Other 📦"
        },
        'espanhol': {
            "personalizado_carro": "Coche 🚗",
            "personalizado_peluche": "Peluche 🧸", 
            "personalizado_acessorio": "Accesorio 💍",
            "personalizado_outro": "Otro 📦"
        },
        'italiano': {
            "personalizado_carro": "Auto 🚗",
            "personalizado_peluche": "Peluche 🧸", 
            "personalizado_acessorio": "Accessorio 💍",
            "personalizado_outro": "Altro 📦"
        },
        'alemao': {
            "personalizado_carro": "Auto 🚗",
            "personalizado_peluche": "Plüschtier 🧸", 
            "personalizado_acessorio": "Accessoire 💍",
            "personalizado_outro": "Andere 📦"
        },
        'frances': {
            "personalizado_carro": "Voiture 🚗",
            "personalizado_peluche": "Peluche 🧸", 
            "personalizado_acessorio": "Accessoire 💍",
            "personalizado_outro": "Autre 📦"
        }
    }
    
    tipos_personalizado = tipos_personalizado_por_idioma.get(idioma, tipos_personalizado_por_idioma['portugues'])
    tipo_personalizado = tipos_personalizado.get(query.data, "")
    context.user_data["tipo_personalizado"] = tipo_personalizado
    
    print(f"✅ Tipo personalizado guardado: {tipo_personalizado}")
    
    # 🔥 ATUALIZAR RESUMO IMEDIATAMENTE
    current_resumo_msg_id = context.user_data.get('resumo_msg_id')
    new_message_id = await enviar_resumo(
        context, 
        query.message.chat_id, 
        message_id=current_resumo_msg_id
    )
    context.user_data['resumo_msg_id'] = new_message_id
    
    # Remover mensagem anterior
    mensagem_personalizado_id = context.user_data.get('mensagem_personalizado_id')
    if mensagem_personalizado_id:
        try:
            await context.bot.delete_message(
                chat_id=query.message.chat_id,
                message_id=mensagem_personalizado_id
            )
            print("✅ Mensagem personalizado anterior apagada")
        except:
            print("❌ Não foi possível apagar mensagem personalizado anterior")
    
    try:
        await query.delete_message()
        print("✅ Mensagem callback apagada")
    except:
        print("❌ Não foi possível apagar mensagem callback")

    # 🔥 SE FOR "OUTRO", PERGUNTAR O NOME DA PEÇA PRIMEIRO
    if query.data == "personalizado_outro":
        print("🎯 Fluxo: Personalizado Outro - perguntando nome da peça")
        
        # 🔥 TEXTO "NOME DA PEÇA" POR IDIOMA
        textos_nome_peca = {
            'portugues': "📝 *Qual é o nome da sua peça personalizada?*\n\n*Exemplos:* Casa, Bicicleta, Instrumento Musical, etc.",
            'ingles': "📝 *What is the name of your custom piece?*\n\n*Examples:* House, Bicycle, Musical Instrument, etc.",
            'espanhol': "📝 *¿Cuál es el nombre de tu pieza personalizada?*\n\n*Ejemplos:* Casa, Bicicleta, Instrumento Musical, etc.",
            'italiano': "📝 *Qual è il nome del tuo pezzo personalizzato?*\n\n*Esempi:* Casa, Bicicletta, Strumento Musicale, etc.",
            'alemao': "📝 *Wie heißt Ihr benutzerdefiniertes Stück?*\n\n*Beispiele:* Haus, Fahrrad, Musikinstrument, etc.",
            'frances': "📝 *Quel est le nom de votre pièce personnalisée ?*\n\n*Exemples :* Maison, Vélo, Instrument de Musique, etc."
        }
        
        mensagem_nome_peca = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_nome_peca.get(idioma, textos_nome_peca['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_peca_id'] = mensagem_nome_peca.message_id
        context.user_data['conversation_state'] = NOME_PECA
        print(f"✅ Estado definido para NOME_PECA: {NOME_PECA} | Idioma: {idioma}")
        
    else:
        # 🔥 PARA OS OUTROS TIPOS, PERGUNTAR O NOME DO CARTOON
        print(f"🎯 Fluxo: Personalizado {tipo_personalizado} - perguntando nome do cartoon")
        
        # 🔥 TEXTO "NOME DO CARTOON" POR IDIOMA
        textos_nome_personalizado = {
            'portugues': "🎭 *Escreve um nome ao seu cartoon para colocar na sua box!*\n\n*Exemplo:* `Carro Especial`, `Peluche Mary`",
            'ingles': "🎭 *Write a name for your cartoon to put on your box!*\n\n*Example:* `Special Car`, `Mary's Plush`",
            'espanhol': "🎭 *¡Escribe un nombre para tu caricatura para poner en tu caja!*\n\n*Ejemplo:* `Coche Especial`, `Peluche María`",
            'italiano': "🎭 *Scrivi un nome per il tuo cartoon da mettere sulla tua scatola!*\n\n*Esempio:* `Auto Speciale`, `Peluche Maria`",
            'alemao': "🎭 *Schreibe einen Namen für deine Karikatur, der auf deine Box kommt!*\n\n*Beispiel:* `Besonderes Auto`, `Maries Plüschtier`",
            'frances': "🎭 *Écrivez un nom pour votre dessin animé à mettre sur votre boîte !*\n\n*Exemple :* `Voiture Spéciale`, `Peluche Marie`"
        }
        
        mensagem_nome_personalizado = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=textos_nome_personalizado.get(idioma, textos_nome_personalizado['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_personalizado_id'] = mensagem_nome_personalizado.message_id
        context.user_data['conversation_state'] = NOME_PERSONALIZADO
        print(f"✅ Estado definido para NOME_PERSONALIZADO: {NOME_PERSONALIZADO} | Idioma: {idioma}")




        

# --- Handler para receber nome da peça quando é "Outro" ---
async def receber_nome_peca(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber nome da peça personalizada quando escolhe 'Outro' - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        nome_peca = update.message.text
        context.user_data["nome_peca_personalizado"] = nome_peca
        
        # 🔥 ATUALIZAR RESUMO IMEDIATAMENTE
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id
        
        print(f"✅ Nome da peça personalizada guardado: {nome_peca} | Idioma: {idioma}")
        
        # Remover mensagens
        try:
            await update.message.delete()
        except:
            pass
            
        mensagem_nome_peca_id = context.user_data.get('mensagem_nome_peca_id')
        if mensagem_nome_peca_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_peca_id
                )
            except:
                pass
        
        # 🔥 TEXTO "NOME DO CARTOON PARA OUTRO" POR IDIOMA
        textos_nome_cartoon_outro = {
            'portugues': "🎭 *Escreve um nome ao seu cartoon para colocar na sua box!*\n\n*Exemplo:* `Casa dos Sonhos`, `Bicicleta Aventureira`",
            'ingles': "🎭 *Write a name for your cartoon to put on your box!*\n\n*Example:* `Dream House`, `Adventure Bike`",
            'espanhol': "🎭 *¡Escribe un nombre para tu caricatura para poner en tu caja!*\n\n*Ejemplo:* `Casa de los Sueños`, `Bicicleta Aventurera`",
            'italiano': "🎭 *Scrivi un nome per il tuo cartoon da mettere sulla tua scatola!*\n\n*Esempio:* `Casa dei Sogni`, `Bicicletta Avventurosa`",
            'alemao': "🎭 *Schreibe einen Namen für deine Karikatur, der auf deine Box kommt!*\n\n*Beispiel:* `Traumhaus`, `Abenteuer-Fahrrad`",
            'frances': "🎭 *Écrivez un nom pour votre dessin animé à mettre sur votre boîte !*\n\n*Exemple :* `Maison de Rêve`, `Vélo Aventure`"
        }
        
        mensagem_nome_personalizado = await update.message.reply_text(
            textos_nome_cartoon_outro.get(idioma, textos_nome_cartoon_outro['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_personalizado_id'] = mensagem_nome_personalizado.message_id
        context.user_data['conversation_state'] = NOME_PERSONALIZADO
        print(f"✅ Estado: NOME_PERSONALIZADO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_nome_peca: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))

# --- Handler para receber nome do cartoon personalizado ---
async def receber_nome_personalizado(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber o nome do cartoon personalizado - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        nome_personalizado = update.message.text
        context.user_data["nome_personalizado"] = nome_personalizado
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem nome personalizado do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta
        mensagem_nome_personalizado_id = context.user_data.get('mensagem_nome_personalizado_id')
        if mensagem_nome_personalizado_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_personalizado_id
                )
                print(f"DEBUG: Mensagem pergunta nome personalizado apagada: {mensagem_nome_personalizado_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta nome personalizado: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTOS "FRASE QUE DEFINE O ELEMENTO" POR IDIOMA
        textos_frase_personalizado = {
            'portugues': """💬 *Escreve uma frase que define o seu elemento!*

*Exemplo:* `"Minha companheira de aventuras"* 
*Ou:* `"Presente cheio de memórias"* 
*Ou:* `"Símbolo da minha paixão"* 

*Pode escrever a sua frase ou enviar* /skip *para continuar:*""",
            
            'ingles': """💬 *Write a phrase that defines your element!*

*Example:* `"My adventure companion"* 
*Or:* `"Gift full of memories"* 
*Or:* `"Symbol of my passion"* 

*You can write your phrase or send* /skip *to continue:*""",
            
            'espanhol': """💬 *¡Escribe una frase que defina tu elemento!*

*Ejemplo:* `"Mi compañera de aventuras"* 
*O:* `"Regalo lleno de recuerdos"* 
*O:* `"Símbolo de mi pasión"* 

*Puedes escribir tu frase o enviar* /skip *para continuar:*""",
            
            'italiano': """💬 *Scrivi una frase che definisce il tuo elemento!*

*Esempio:* `"Il mio compagno di avventure"* 
*O:* `"Regalo pieno di ricordi"* 
*O:* `"Simbolo della mia passione"* 

*Puoi scrivere la tua frase o inviare* /skip *per continuare:*""",
            
            'alemao': """💬 *Schreibe einen Satz, der dein Element definiert!*

*Beispiel:* `"Mein Abenteuerbegleiter"* 
*Oder:* `"Geschenk voller Erinnerungen"* 
*Oder:* `"Symbol meiner Leidenschaft"* 

*Sie können Ihren Satz schreiben oder* /skip *senden, um fortzufahren:*""",
            
            'frances': """💬 *Écrivez une phrase qui définit votre élément !*

*Exemple :* `"Mon compagnon d'aventure"* 
*Ou :* `"Cadeau plein de souvenirs"* 
*Ou :* `"Symbole de ma passion"* 

*Vous pouvez écrire votre phrase ou envoyer* /skip *pour continuer :*"""
        }
        
        mensagem_frase_personalizado = await update.message.reply_text(
            textos_frase_personalizado.get(idioma, textos_frase_personalizado['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_frase_personalizado_id'] = mensagem_frase_personalizado.message_id
        context.user_data['conversation_state'] = FRASE_PERSONALIZADO
        print(f"✅ Estado: FRASE_PERSONALIZADO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_nome_personalizado: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))

# --- Handler para receber frase do personalizado ---
async def receber_frase_personalizado(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber a frase que define o elemento personalizado - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        # 🔥 COMANDO DE PULAR POR IDIOMA
        comandos_pular = {
            'portugues': '/skip',
            'ingles': '/skip',
            'espanhol': '/skip',
            'italiano': '/skip',
            'alemao': '/skip',
            'frances': '/skip'
        }
        
        comando_pular = comandos_pular.get(idioma, '/pular')
        
        # Verificar se é comando /pular ou mensagem normal
        if update.message.text and update.message.text.strip() == comando_pular:
            frase_personalizado = "Não adicionou frase"
            print(f"✅ Usuário escolheu pular a frase do personalizado | Idioma: {idioma}")
        else:
            frase_personalizado = update.message.text
            print(f"✅ Usuário adicionou frase do personalizado: {frase_personalizado} | Idioma: {idioma}")
        
        context.user_data["frase_personalizado"] = frase_personalizado
        
        # 🔥 REMOVER mensagem do usuário
        try:
            await update.message.delete()
            print("DEBUG: Mensagem frase personalizado do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
        
        # 🔥 REMOVER mensagem da pergunta
        mensagem_frase_personalizado_id = context.user_data.get('mensagem_frase_personalizado_id')
        if mensagem_frase_personalizado_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_frase_personalizado_id
                )
                print(f"DEBUG: Mensagem pergunta frase personalizado apagada: {mensagem_frase_personalizado_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta frase personalizado: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 AGORA MOSTRAR TAMANHOS PARA PERSONALIZADO
        await mostrar_tamanhos_personalizado(context, update.message.chat_id)
        
    except Exception as e:
        print(f"ERRO em receber_frase_personalizado: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))

















    

async def tamanho_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # Extrair o tamanho do callback_data
    tamanho_key = query.data.replace("tamanho_", "")
    estilo_escolhido = context.user_data.get("estilo_cartoon", "")
    tipo_cartoon = context.user_data.get("tipo_cartoon", "")
    
    print(f"🔧 DEBUG tamanho_handler | Idioma: {idioma}:")
    print(f"  Tamanho selecionado: {tamanho_key}")
    print(f"  Tipo cartoon: {tipo_cartoon}")
    print(f"  Estilo: {estilo_escolhido}")

    # 🔥 CORREÇÃO: Normalizar o tipo para comparação
    tipo_normalizado = tipo_cartoon.lower().replace(" ", "_").replace("🐱", "").replace("🐶", "").replace("🎨", "").strip()
    tipo_lower = tipo_cartoon.lower()  # Para comparação direta
    print(f"  Tipo normalizado: {tipo_normalizado}")

    # 🔥 🔥 🔥 CORREÇÃO: BUSCAR PREÇO COM SUPORTE MULTI-IDIOMA
    # DETECTAR PERSONALIZADO (todos os idiomas)
    if (any(p in tipo_lower for p in ['personal', 'personal.', 'personalisiert', 'personalizzato', 
                                     'personnalisé', 'personalizado', 'custom', 'bespoke', 'especial',
                                     'individuell', 'benutzerdefiniert', 'maßgeschneidert', 'speciale',
                                     'spécial', 'su misura', 'customizado', 'customisé'])):
        print("🔧 🔥 🔥 Buscando preços para PERSONALIZADO")
        # Tamanhos específicos para Personalizado
        TAMANHOS_PERSONALIZADO = {
           "2.5": {"nome": " 2.5\" | 6.4cm ", "preco": 0.0},
           "3.5": {"nome": " 3.5\" | 8.9cm ", "preco": 5.0},
           "4.5": {"nome": " 4.5\" | 11.4cm ", "preco": 10.0},
           "6": {"nome": " 6\" | 15.2cm ", "preco": 25.0},
           "7": {"nome": " 7\" | 17.8cm ", "preco": 35.0},
           "8": {"nome": " 8\" | 20.3cm ", "preco": 55.0},
           "9": {"nome": " 9\" | 22.9cm ", "preco": 70.0},
           "10": {"nome": " 10\" | 25.4cm ", "preco": 90.0},
           "11": {"nome": " 11\" | 27.9cm ", "preco": 110.0},
           "12": {"nome": " 12\" | 30.5cm ", "preco": 150.0}
        }
        tamanhos_disponiveis = TAMANHOS_PERSONALIZADO
        info_tamanho = tamanhos_disponiveis.get(tamanho_key, {})
        
    # DETECTAR ANIMAL (todos os idiomas)
    elif (any(p in tipo_lower for p in ['animal', 'tier', 'animale', 'mascota', 'pet', 'bicho',
                                       'haustier', 'animal de compagnie', 'animale domestico',
                                       'creature', 'animais', 'animales', 'animali', 'tiere', 'animaux'])):
        print("🔧 Buscando preços para ANIMAL")
        tamanhos_disponiveis = TAMANHOS_ANIMAL
        info_tamanho = tamanhos_disponiveis.get(tamanho_key, {})
        
    # DETECTAR FAMILY/GRUPO (todos os idiomas)
    elif (any(p in tipo_lower for p in ['family', 'família', 'familia', 'grupo', 'group', 'gruppe',
                                       'groupe', 'gruppo', 'fam', 'families', 'familiar', 'familiare',
                                       'familien', 'familial', 'famille'])):
        print("🔧 Buscando preços para FAMILY")
        tamanhos_disponiveis = TAMANHOS_FAMILY
        info_tamanho = tamanhos_disponiveis.get(tamanho_key, {})
        
    else:
        print("🔧 Buscando preços para OUTRO TIPO")
        tamanhos_disponiveis = TAMANHOS_POR_ESTILO.get(estilo_escolhido, {})
        info_tamanho = tamanhos_disponiveis.get(tamanho_key, {})

    print(f"  Tamanhos disponíveis: {list(tamanhos_disponiveis.keys())}")
    print(f"  Info tamanho encontrada: {info_tamanho}")
    
    if info_tamanho:
        context.user_data["tamanho_cartoon"] = info_tamanho["nome"]
        context.user_data["tamanho_key"] = tamanho_key
        context.user_data["preco_tamanho"] = info_tamanho["preco"]
        print(f"✅ Tamanho guardado: {info_tamanho['nome']}, Preço: €{info_tamanho['preco']:.2f}")
    else:
        # 🔥 CORREÇÃO: Se não encontrar, usar fallback dos tamanhos de animal
        print("⚠️  Tamanho não encontrado, usando fallback animal")
        fallback_info = TAMANHOS_ANIMAL.get(tamanho_key, {"nome": f"{tamanho_key}cm", "preco": 0.0})
        context.user_data["tamanho_cartoon"] = fallback_info["nome"]
        context.user_data["tamanho_key"] = tamanho_key
        context.user_data["preco_tamanho"] = fallback_info["preco"]
        print(f"⚠️  Tamanho fallback: {fallback_info['nome']}, Preço: €{fallback_info['preco']:.2f}")

    print(f"  Preço tamanho guardado: €{context.user_data.get('preco_tamanho', 0):.2f}")

    # 🔥 DEBUG: TESTAR CÁLCULO IMEDIATAMENTE
    print("🐾 DEBUG - Testando cálculo:")
    totais_teste = calcular_total(context)
    print(f"🐾 RESULTADO TESTE: €{totais_teste['total']:.2f}")

    # Resto do código permanece igual...
    mensagem_tamanhos_id = context.user_data.get('mensagem_tamanhos_id')
    if mensagem_tamanhos_id:
        try:
            await context.bot.delete_message(
                chat_id=query.message.chat_id,
                message_id=mensagem_tamanhos_id
            )
            print(f"DEBUG: Mensagem tamanhos apagada: {mensagem_tamanhos_id}")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem tamanhos: {e}")
    
    try:
        await query.delete_message()
        print("DEBUG: Mensagem callback apagada com sucesso")
    except Exception as e:
        print(f"DEBUG: Não foi possível apagar mensagem callback: {e}")
        try:
            # 🔥 TEXTO "TAMANHO SELECIONADO" POR IDIOMA
            textos_tamanho_selecionado = {
                'portugues': "✅ Tamanho selecionado!",
                'ingles': "✅ Size selected!",
                'espanhol': "✅ ¡Tamaño seleccionado!",
                'italiano': "✅ Dimensione selezionata!",
                'alemao': "✅ Größe ausgewählt!",
                'frances': "✅ Taille sélectionnée !"
            }
            
            await query.edit_message_text(
                text=textos_tamanho_selecionado.get(idioma, textos_tamanho_selecionado['portugues']),
                reply_markup=None
            )
            print("DEBUG: Mensagem callback editada para esconder botões")
        except Exception as e2:
            print(f"DEBUG: Também não foi possível editar mensagem: {e2}")

    # Atualizar resumo
    current_resumo_msg_id = context.user_data.get('resumo_msg_id')
    new_message_id = await enviar_resumo(
        context, 
        query.message.chat_id, 
        message_id=current_resumo_msg_id
    )
    context.user_data['resumo_msg_id'] = new_message_id

    # 🔥 TEXTO "ENVIE FOTO" POR IDIOMA
    textos_envie_foto = {
        'portugues': "📸 *Perfeito! Agora envie a foto que deseja transformar em cartoon.*",
        'ingles': "📸 *Perfect! Now send the photo you want to transform into a cartoon.*",
        'espanhol': "📸 *¡Perfecto! Ahora envía la foto que deseas transformar en caricatura.*",
        'italiano': "📸 *Perfetto! Ora invia la foto che desideri trasformare in cartoon.*",
        'alemao': "📸 *Perfekt! Senden Sie jetzt das Foto, das Sie in eine Karikatur verwandeln möchten.*",
        'frances': "📸 *Parfait ! Maintenant envoyez la photo que vous souhaitez transformer en dessin animé.*"
    }

    # Pedir a foto
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_envie_foto.get(idioma, textos_envie_foto['portugues']),
        parse_mode="Markdown"
    )
    context.user_data['conversation_state'] = FOTO
    print(f"✅ Estado atualizado para: FOTO | Idioma: {idioma}")




    

# --- Receber foto ---
async def receber_foto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    print("DEBUG: receber_foto foi chamado!")
    print("🎯🎯🎯 RECEBER_FOTO CHAMADO!")
    
    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # Verificar se já temos foto (só bloqueia se não estiver no estado FOTO)
    if "foto_id" in context.user_data and context.user_data.get('conversation_state') != FOTO:
        print("DEBUG: Foto já existe e não está no estado FOTO")
        
        # 🔥 MENSAGEM "FOTO JÁ EXISTE" POR IDIOMA
        textos_foto_existe = {
            'portugues': "⚠️ Já recebemos a sua foto. Use o botão 'Mudar Foto' se quiser alterar.",
            'ingles': "⚠️ We already have your photo. Use the 'Change Photo' button if you want to change it.",
            'espanhol': "⚠️ Ya recibimos tu foto. Usa el botón 'Cambiar Foto' si quieres cambiarla.",
            'italiano': "⚠️ Abbiamo già la tua foto. Usa il pulsante 'Cambia Foto' se vuoi cambiarla.",
            'alemao': "⚠️ Wir haben Ihr Foto bereits. Verwenden Sie die Schaltfläche 'Foto ändern', wenn Sie es ändern möchten.",
            'frances': "⚠️ Nous avons déjà votre photo. Utilisez le bouton 'Changer la photo' si vous souhaitez la modifier."
        }
        
        await update.message.reply_text(textos_foto_existe.get(idioma, textos_foto_existe['portugues']))
        return

    print("DEBUG: Processando nova foto...")
    
    photo = update.message.photo[-1]
    file_id = photo.file_id
    
    # GERAR NOME DA FOTO
    nome_usuario = context.user_data.get('nome', 'Cliente')
    timestamp = datetime.now().strftime("%H%M%S")
    nome_foto = f"foto_{timestamp}.jpg"
    
    print(f"DEBUG: Nome da foto: {nome_foto}")
    print(f"DEBUG: Nome do usuário: {nome_usuario}")
    print(f"DEBUG: Idioma: {idioma}")
    
    context.user_data["foto_id"] = file_id
    context.user_data["nome_foto"] = nome_foto
    print(f"DEBUG: Nome guardado no user_data: {context.user_data.get('nome_foto')}")

    # 🔥 TEXTO "FOTO RECEBIDA" POR IDIOMA
    textos_foto_recebida = {
        'portugues': "📸 Foto recebida com sucesso!",
        'ingles': "📸 Photo received successfully!",
        'espanhol': "📸 ¡Foto recibida con éxito!",
        'italiano': "📸 Foto ricevuta con successo!",
        'alemao': "📸 Foto erfolgreich empfangen!",
        'frances': "📸 Photo reçue avec succès !"
    }
    
    await update.message.reply_text(textos_foto_recebida.get(idioma, textos_foto_recebida['portugues']))

    nome_foto_resumo = context.user_data.get('nome_foto', 'foto.jpg')
    print(f"DEBUG: Nome a mostrar no resumo: {nome_foto_resumo}")

    # 🔥 DICIONÁRIO PARA CONVERTER PAÍSES PARA INGLÊS
    PAISES_PARA_INGLES = {
        # callback_data → Nome em inglês para o resumo
        'estados_unidos': 'United States',
        'canada': 'Canada',
        'reino_unido': 'United Kingdom',
        'brasil': 'Brazil',
        'alemanha': 'Germany',
        'paises_baixos': 'Netherlands',
        'holanda': 'Netherlands',
        'franca': 'France',
        'espanha': 'Spain',
        'belgica': 'Belgium',
        'italia': 'Italy',
        'portugal': 'Portugal',
        'irlanda': 'Ireland',
        'luxemburgo': 'Luxembourg'
    }
    
    def converter_pais_para_ingles(pais_key):
        """Converte o nome/callback do país para inglês"""
        if isinstance(pais_key, str):
            # Remove "pais_" se existir
            if pais_key.startswith('pais_'):
                pais_key = pais_key[5:]
            # Remove acentos e converte para minúsculas para comparação
            pais_clean = pais_key.lower()
            # Mapeamento adicional para nomes em português
            mapeamento = {
                'bélgica': 'belgica',
                'bélgica (português)': 'belgica',
                'frança': 'franca',
                'espanha': 'espanha',
                'alemanha': 'alemanha',
                'itália': 'italia',
                'irlanda': 'irlanda',
                'luxemburgo': 'luxemburgo',
                'países baixos': 'paises_baixos',
                'holanda': 'paises_baixos',
                'reino unido': 'reino_unido',
                'estados unidos': 'estados_unidos',
                'eua': 'estados_unidos'
            }
            pais_key = mapeamento.get(pais_clean, pais_key)
        return PAISES_PARA_INGLES.get(pais_key, pais_key.title())

    # 🔥 CORREÇÃO: RESUMO FINAL COMPLETO COM TRADUÇÃO
    # Texto do título do resumo por idioma
    titulos_resumo = {
        'portugues': "✅ *Resumo Final do Pedido:*",
        'ingles': "✅ *Final Order Summary:*",
        'espanhol': "✅ *Resumen Final del Pedido:*",
        'italiano': "✅ *Riepilogo Finale dell'Ordine:*",
        'alemao': "✅ *Endgültige Bestellübersicht:*",
        'frances': "✅ *Résumé Final de la Commande :*"
    }
    
    # Textos dos campos por idioma
    campos_traduzidos = {
        'portugues': {
            'nome': "👤 Nome:",
            'email': "📧 Email:",
            'pais': "🌍 País:",
            'telefone': "📱 Telefone:",
            'tipo_cartoon': "🎨 Tipo de Cartoon:",
            'estilo': "🖌 Estilo:",
            'nome_family': "👨‍👩‍👧‍👦 Nome da Família:",
            'frase_family': "💬 Frase da Família:",
            'nome_cartoon': "🎭 Nome no Cartoon:",
            'frase_box': "💬 Frase na Box:",
            'tipo_peca': "📦 Tipo de Peça:",
            'nome_peca': "📝 Nome da Peça:",
            'profissao': "💼 Profissão:",
            'objetos': "🎯 Objetos Personalizados:",
            'super_heroi': "🦸‍♂️ Super-Herói:",
            'elementos': "👥 Total de Elementos:",
            'adultos': "👨‍👩 Adultos:",
            'criancas': "👧🧒 Crianças:",
            'animais': "🐱🐶 Animais:",
            'nome_animal': "🐾 Nome do Animal:",
            'tipo_animal': "🐕 Tipo de Animal:",
            'tamanho': "📏 Tamanho:",
            'foto': "📸 Foto:",
            'pronto': "🚀 Estamos prontos para criar o seu cartoon!"
        },
        'ingles': {
            'nome': "👤 Name:",
            'email': "📧 Email:",
            'pais': "🌍 Country:",
            'telefone': "📱 Phone:",
            'tipo_cartoon': "🎨 Cartoon Type:",
            'estilo': "🖌 Style:",
            'nome_family': "👨‍👩‍👧‍👦 Family Name:",
            'frase_family': "💬 Family Phrase:",
            'nome_cartoon': "🎭 Name on Cartoon:",
            'frase_box': "💬 Box Phrase:",
            'tipo_peca': "📦 Piece Type:",
            'nome_peca': "📝 Piece Name:",
            'profissao': "💼 Profession:",
            'objetos': "🎯 Custom Objects:",
            'super_heroi': "🦸‍♂️ Superhero:",
            'elementos': "👥 Total Elements:",
            'adultos': "👨‍👩 Adults:",
            'criancas': "👧🧒 Children:",
            'animais': "🐱🐶 Animals:",
            'nome_animal': "🐾 Animal Name:",
            'tipo_animal': "🐕 Animal Type:",
            'tamanho': "📏 Size:",
            'foto': "📸 Photo:",
            'pronto': "🚀 We are ready to create your cartoon!"
        },
        'espanhol': {
            'nome': "👤 Nombre:",
            'email': "📧 Email:",
            'pais': "🌍 País:",
            'telefone': "📱 Teléfono:",
            'tipo_cartoon': "🎨 Tipo de Caricatura:",
            'estilo': "🖌 Estilo:",
            'nome_family': "👨‍👩‍👧‍👦 Nombre de la Familia:",
            'frase_family': "💬 Frase de la Familia:",
            'nome_cartoon': "🎭 Nombre en la Caricatura:",
            'frase_box': "💬 Frase en la Caja:",
            'tipo_peca': "📦 Tipo de Pieza:",
            'nome_peca': "📝 Nombre de la Pieza:",
            'profissao': "💼 Profesión:",
            'objetos': "🎯 Objetos Personalizados:",
            'super_heroi': "🦸‍♂️ Superhéroe:",
            'elementos': "👥 Total de Elementos:",
            'adultos': "👨‍👩 Adultos:",
            'criancas': "👧🧒 Niños:",
            'animais': "🐱🐶 Animales:",
            'nome_animal': "🐾 Nombre del Animal:",
            'tipo_animal': "🐕 Tipo de Animal:",
            'tamanho': "📏 Tamaño:",
            'foto': "📸 Foto:",
            'pronto': "🚀 ¡Estamos listos para crear tu caricatura!"
        },
        'italiano': {
            'nome': "👤 Nome:",
            'email': "📧 Email:",
            'pais': "🌍 Paese:",
            'telefone': "📱 Telefono:",
            'tipo_cartoon': "🎨 Tipo di Cartoon:",
            'estilo': "🖌 Stile:",
            'nome_family': "👨‍👩‍👧‍👦 Nome della Famiglia:",
            'frase_family': "💬 Frase della Famiglia:",
            'nome_cartoon': "🎭 Nome sul Cartoon:",
            'frase_box': "💬 Frase sulla Scatola:",
            'tipo_peca': "📦 Tipo di Pezzo:",
            'nome_peca': "📝 Nome del Pezzo:",
            'profissao': "💼 Professione:",
            'objetos': "🎯 Oggetti Personalizzati:",
            'super_heroi': "🦸‍♂️ Supereroe:",
            'elementos': "👥 Totale Elementi:",
            'adultos': "👨‍👩 Adulti:",
            'criancas': "👧🧒 Bambini:",
            'animais': "🐱🐶 Animali:",
            'nome_animal': "🐾 Nome dell'Animale:",
            'tipo_animal': "🐕 Tipo di Animale:",
            'tamanho': "📏 Dimensione:",
            'foto': "📸 Foto:",
            'pronto': "🚀 Siamo pronti a creare il tuo cartoon!"
        },
        'alemao': {
            'nome': "👤 Name:",
            'email': "📧 E-Mail:",
            'pais': "🌍 Land:",
            'telefone': "📱 Telefon:",
            'tipo_cartoon': "🎨 Karikaturtyp:",
            'estilo': "🖌 Stil:",
            'nome_family': "👨‍👩‍👧‍👦 Familienname:",
            'frase_family': "💬 Familiensatz:",
            'nome_cartoon': "🎭 Name auf der Karikatur:",
            'frase_box': "💬 Box-Satz:",
            'tipo_peca': "📦 Stücktyp:",
            'nome_peca': "📝 Stückname:",
            'profissao': "💼 Beruf:",
            'objetos': "🎯 Benutzerdefinierte Objekte:",
            'super_heroi': "🦸‍♂️ Superheld:",
            'elementos': "👥 Gesamtelemente:",
            'adultos': "👨‍👩 Erwachsene:",
            'criancas': "👧🧒 Kinder:",
            'animais': "🐱🐶 Tiere:",
            'nome_animal': "🐾 Tiername:",
            'tipo_animal': "🐕 Tierart:",
            'tamanho': "📏 Größe:",
            'foto': "📸 Foto:",
            'pronto': "🚀 Wir sind bereit, Ihre Karikatur zu erstellen!"
        },
        'frances': {
            'nome': "👤 Nom:",
            'email': "📧 E-mail:",
            'pais': "🌍 Pays:",
            'telefone': "📱 Téléphone:",
            'tipo_cartoon': "🎨 Type de Dessin Animé:",
            'estilo': "🖌 Style:",
            'nome_family': "👨‍👩‍👧‍👦 Nom de Famille:",
            'frase_family': "💬 Phrase de Famille:",
            'nome_cartoon': "🎭 Nom sur le Dessin Animé:",
            'frase_box': "💬 Phrase sur la Boîte:",
            'tipo_peca': "📦 Type de Pièce:",
            'nome_peca': "📝 Nom de la Pièce:",
            'profissao': "💼 Profession:",
            'objetos': "🎯 Objets Personnalisés:",
            'super_heroi': "🦸‍♂️ Super-héros:",
            'elementos': "👥 Total des Éléments:",
            'adultos': "👨‍👩 Adultes:",
            'criancas': "👧🧒 Enfants:",
            'animais': "🐱🐶 Animaux:",
            'nome_animal': "🐾 Nom de l'Animal:",
            'tipo_animal': "🐕 Type d'Animal:",
            'tamanho': "📏 Taille:",
            'foto': "📸 Photo:",
            'pronto': "🚀 Nous sommes prêts à créer votre dessin animé !"
        }
    }
    
    campos = campos_traduzidos.get(idioma, campos_traduzidos['portugues'])
    
    # Construir resumo
    resumo = f"{titulos_resumo.get(idioma, titulos_resumo['portugues'])}\n\n"
    resumo += f"{campos['nome']} {context.user_data.get('nome', '')}\n"
    resumo += f"{campos['email']} {context.user_data.get('email', '')}\n"
    
    # 🔥 CONVERTER PAÍS PARA INGLÊS
    if "pais" in context.user_data:
        pais_original = context.user_data['pais']
        pais_ingles = converter_pais_para_ingles(pais_original)
        resumo += f"{campos['pais']} {pais_ingles}\n"
    else:
        resumo += f"{campos['pais']} \n"
    
    resumo += f"{campos['telefone']} {context.user_data.get('contacto', '')}\n"
    resumo += f"{campos['tipo_cartoon']} {context.user_data.get('tipo_cartoon', '')}\n"
    
    # 🔥 ADICIONAR ESTILO SE EXISTIR (para Individual)
    if "estilo_cartoon" in context.user_data:
        resumo += f"{campos['estilo']} {context.user_data.get('estilo_cartoon', '')}\n"

    # 🔥 CAMPOS ESPECÍFICOS DA FAMILY
    if "nome_family" in context.user_data:
        resumo += f"{campos['nome_family']} {context.user_data.get('nome_family', '')}\n"
    if "frase_family" in context.user_data and context.user_data.get('frase_family') != "Não adicionou frase":
        resumo += f"{campos['frase_family']} \"{context.user_data.get('frase_family', '')}\"\n"
    
    # 🔥 CAMPOS DE PERSONALIZAÇÃO DA BOX
    if "nome_cartoon" in context.user_data:
        resumo += f"{campos['nome_cartoon']} {context.user_data.get('nome_cartoon', '')}\n"
    if "frase_cartoon" in context.user_data and context.user_data.get('frase_cartoon') != "Não adicionou frase":
        resumo += f"{campos['frase_box']} \"{context.user_data.get('frase_cartoon', '')}\"\n"
    
    # 🔥 CAMPOS DO PERSONALIZADO
    if "tipo_personalizado" in context.user_data:
        resumo += f"{campos['tipo_peca']} {context.user_data.get('tipo_personalizado', '')}\n"
    if "nome_peca_personalizado" in context.user_data:
        resumo += f"{campos['nome_peca']} {context.user_data.get('nome_peca_personalizado', '')}\n"
    
    if "nome_personalizado" in context.user_data:
        resumo += f"{campos['nome_cartoon']} {context.user_data.get('nome_personalizado', '')}\n"
    if "frase_personalizado" in context.user_data and context.user_data.get('frase_personalizado') != "Não adicionou frase":
        resumo += f"{campos['frase_box']} \"{context.user_data.get('frase_personalizado', '')}\"\n"
    
    # 🔥 CAMPOS PERSONALIZADOS
    if "profissao" in context.user_data:
        resumo += f"{campos['profissao']} {context.user_data.get('profissao', '')}\n"
    if "objetos_office" in context.user_data:
        resumo += f"{campos['objetos']} {context.user_data.get('objetos_office', '')}\n"
    if "super_heroi" in context.user_data:
        resumo += f"{campos['super_heroi']} {context.user_data.get('super_heroi', '')}\n"
    
    # 🔥 CAMPOS DA FAMILY
    if "elementos_family" in context.user_data:
        resumo += f"{campos['elementos']} {context.user_data.get('elementos_family', '')}\n"
    if "adultos_family" in context.user_data:
        resumo += f"{campos['adultos']} {context.user_data.get('adultos_family', '')}\n"
    if "criancas_family" in context.user_data:
        resumo += f"{campos['criancas']} {context.user_data.get('criancas_family', '')}\n"
    if "animais_family" in context.user_data:
        resumo += f"{campos['animais']} {context.user_data.get('animais_family', '')}\n"
    if "nome_animal" in context.user_data:
        resumo += f"{campos['nome_animal']} {context.user_data.get('nome_animal', '')}\n"
    if "tipo_animal" in context.user_data:
        resumo += f"{campos['tipo_animal']} {context.user_data.get('tipo_animal', '')}\n"
     
    # 🔥 ADICIONAR TAMANHO E FOTO
    resumo += f"{campos['tamanho']} {context.user_data.get('tamanho_cartoon', '')}\n"
    resumo += f"{campos['foto']} recebida ✅ (*{nome_foto_resumo}*)\n\n"
    resumo += f"{campos['pronto']}"

    # 🔥 BOTÕES TRADUZIDOS
    botoes_por_idioma = {
        'portugues': {
            'finalizar': "💳 Finalizar Compra",
            'mudar_foto': "📸 Enganei-me na foto (mudar)",
            'voltar_inicio': "↩️ Enganei-me nos dados (voltar ao início)"
        },
        'ingles': {
            'finalizar': "💳 Complete Purchase",
            'mudar_foto': "📸 Wrong photo (change)",
            'voltar_inicio': "↩️ Wrong data (back to start)"
        },
        'espanhol': {
            'finalizar': "💳 Finalizar Compra",
            'mudar_foto': "📸 Me equivoqué de foto (cambiar)",
            'voltar_inicio': "↩️ Me equivoqué en los datos (volver al inicio)"
        },
        'italiano': {
            'finalizar': "💳 Completa Acquisto",
            'mudar_foto': "📸 Foto sbagliata (cambia)",
            'voltar_inicio': "↩️ Dati sbagliati (torna all'inizio)"
        },
        'alemao': {
            'finalizar': "💳 Kauf abschließen",
            'mudar_foto': "📸 Falsches Foto (ändern)",
            'voltar_inicio': "↩️ Falsche Daten (zurück zum Start)"
        },
        'frances': {
            'finalizar': "💳 Finaliser l'Achat",
            'mudar_foto': "📸 Mauvaise photo (changer)",
            'voltar_inicio': "↩️ Mauvaises données (retour au début)"
        }
    }
    
    botoes_texto = botoes_por_idioma.get(idioma, botoes_por_idioma['portugues'])
    
    botoes = [
        [InlineKeyboardButton(botoes_texto['finalizar'], callback_data="finalizar_compra")],
        [InlineKeyboardButton(botoes_texto['mudar_foto'], callback_data="mudar_foto")],
        [InlineKeyboardButton(botoes_texto['voltar_inicio'], callback_data="voltar_inicio")]
    ]
    
    await update.message.reply_text(resumo, parse_mode="Markdown", reply_markup=InlineKeyboardMarkup(botoes))
    context.user_data['conversation_state'] = "FIM"
    print(f"✅ Estado: FIM | Idioma: {idioma}")




    

async def mudar_foto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para mudar a foto - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # DEBUG: Verificar nome atual antes de remover
    print(f"DEBUG: Nome da foto atual antes de mudar: {context.user_data.get('nome_foto')} | Idioma: {idioma}")
    
    # Remover a foto atual E o nome da foto
    context.user_data.pop("foto_id", None)
    context.user_data.pop("nome_foto", None)
    
    # Remover a mensagem com os botões antigos
    await safe_delete_message(query)
    
    # 🔥 TEXTO "ENVIE NOVA FOTO" POR IDIOMA
    textos_nova_foto = {
        'portugues': "📸 *Por favor, envie a nova foto:*",
        'ingles': "📸 *Please send the new photo:*",
        'espanhol': "📸 *Por favor, envía la nueva foto:*",
        'italiano': "📸 *Per favore, invia la nuova foto:*",
        'alemao': "📸 *Bitte senden Sie das neue Foto:*",
        'frances': "📸 *Veuillez envoyer la nouvelle photo :*"
    }
    
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_nova_foto.get(idioma, textos_nova_foto['portugues']),
        parse_mode="Markdown"
    )
    context.user_data['conversation_state'] = FOTO
    print(f"✅ Estado: FOTO (mudar foto) | Idioma: {idioma}")










async def receber_elementos_family(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber número total de elementos da family - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        elementos = update.message.text
        context.user_data["elementos_family"] = elementos
        
        # Remover mensagens
        try:
            await update.message.delete()
        except:
            pass
            
        mensagem_elementos_id = context.user_data.get('mensagem_elementos_id')
        if mensagem_elementos_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_elementos_id
                )
            except:
                pass
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "QUANTOS ADULTOS" POR IDIOMA
        textos_adultos = {
            'portugues': "👨‍👩 *Quantos adultos vão estar no cartoon?*",
            'ingles': "👨‍👩 *How many adults will be in the cartoon?*",
            'espanhol': "👨‍👩 *¿Cuántos adultos estarán en la caricatura?*",
            'italiano': "👨‍👩 *Quanti adulti saranno nel cartoon?*",
            'alemao': "👨‍👩 *Wie viele Erwachsene werden in der Karikatur sein?*",
            'frances': "👨‍👩 *Combien d'adultes seront dans le dessin animé ?*"
        }
        
        mensagem_adultos = await update.message.reply_text(
            textos_adultos.get(idioma, textos_adultos['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_adultos_id'] = mensagem_adultos.message_id
        context.user_data['conversation_state'] = ADULTOS_FAMILY
        print(f"✅ Estado: ADULTOS_FAMILY | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_elementos_family: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))

async def receber_adultos_family(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber número de adultos da family - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        adultos = update.message.text
        context.user_data["adultos_family"] = adultos
        
        # Remover mensagens
        try:
            await update.message.delete()
        except:
            pass
            
        mensagem_adultos_id = context.user_data.get('mensagem_adultos_id')
        if mensagem_adultos_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_adultos_id
                )
            except:
                pass
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "QUANTAS CRIANÇAS" POR IDIOMA
        textos_criancas = {
            'portugues': "👧🧒 *Quantas crianças vão estar no cartoon?*",
            'ingles': "👧🧒 *How many children will be in the cartoon?*",
            'espanhol': "👧🧒 *¿Cuántos niños estarán en la caricatura?*",
            'italiano': "👧🧒 *Quanti bambini saranno nel cartoon?*",
            'alemao': "👧🧒 *Wie viele Kinder werden in der Karikatur sein?*",
            'frances': "👧🧒 *Combien d'enfants seront dans le dessin animé ?*"
        }
        
        mensagem_criancas = await update.message.reply_text(
            textos_criancas.get(idioma, textos_criancas['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_criancas_id'] = mensagem_criancas.message_id
        context.user_data['conversation_state'] = CRIANCAS_FAMILY
        print(f"✅ Estado: CRIANCAS_FAMILY | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_adultos_family: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))



        

async def receber_criancas_family(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber número de crianças da family - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        criancas = update.message.text
        context.user_data["criancas_family"] = criancas
        
        # Remover mensagens
        try:
            await update.message.delete()
        except:
            pass
            
        mensagem_criancas_id = context.user_data.get('mensagem_criancas_id')
        if mensagem_criancas_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_criancas_id
                )
            except:
                pass
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "QUANTOS ANIMAIS" POR IDIOMA
        textos_animais = {
            'portugues': "🐱🐶 *Quantos animais vão estar no cartoon?*",
            'ingles': "🐱🐶 *How many animals will be in the cartoon?*",
            'espanhol': "🐱🐶 *¿Cuántos animales estarán en la caricatura?*",
            'italiano': "🐱🐶 *Quanti animali saranno nel cartoon?*",
            'alemao': "🐱🐶 *Wie viele Tiere werden in der Karikatur sein?*",
            'frances': "🐱🐶 *Combien d'animaux seront dans le dessin animé ?*"
        }
        
        mensagem_animais = await update.message.reply_text(
            textos_animais.get(idioma, textos_animais['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_animais_id'] = mensagem_animais.message_id
        context.user_data['conversation_state'] = ANIMAIS_FAMILY
        print(f"✅ Estado: ANIMAIS_FAMILY | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_criancas_family: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))

async def receber_animais_family(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber número de animais da family - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        animais = update.message.text
        context.user_data["animais_family"] = animais
        
        # Remover mensagens
        try:
            await update.message.delete()
        except:
            pass
            
        mensagem_animais_id = context.user_data.get('mensagem_animais_id')
        if mensagem_animais_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_animais_id
                )
            except:
                pass
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "ESCOLHER TAMANHO FAMILY" POR IDIOMA
        textos_tamanho_family = {
            'portugues': "📏 *Escolha o tamanho do seu Cartoon Family:*",
            'ingles': "📏 *Choose the size of your Family Cartoon:*",
            'espanhol': "📏 *Elige el tamaño de tu Caricatura Familiar:*",
            'italiano': "📏 *Scegli la dimensione del tuo Cartoon Famiglia:*",
            'alemao': "📏 *Wählen Sie die Größe Ihrer Familien-Karikatur:*",
            'frances': "📏 *Choisissez la taille de votre Dessin Animé Familial :*"
        }

        # Mostrar tamanhos para Family em GRADE 2xN
        teclado = []
        tamanhos_lista = list(TAMANHOS_FAMILY.items())
        
        # Processar em pares (2 botões por linha)
        for i in range(0, len(tamanhos_lista), 2):
            linha = []
            # Primeiro botão da linha
            tamanho_key1, info_tamanho1 = tamanhos_lista[i]
            botao_texto1 = f"{info_tamanho1['nome']}"
            linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
            
            # Segundo botão da linha (se existir)
            if i + 1 < len(tamanhos_lista):
                tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
                botao_texto2 = f"{info_tamanho2['nome']}"
                linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
            
            teclado.append(linha)

        mensagem_tamanhos = await update.message.reply_text(
            textos_tamanho_family.get(idioma, textos_tamanho_family['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        
        context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
        context.user_data['conversation_state'] = TAMANHO
        print(f"✅ Estado: TAMANHO | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_animais_family: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))





async def receber_nome_animal(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber nome do animal - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        nome_animal = update.message.text
        context.user_data["nome_animal"] = nome_animal
        
        # 🔥 CORREÇÃO: Tentar remover mensagens COM TRY/EXCEPT
        try:
            await update.message.delete()
            print("DEBUG: Mensagem nome animal do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
            
        mensagem_nome_animal_id = context.user_data.get('mensagem_nome_animal_id')
        if mensagem_nome_animal_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_animal_id
                )
                print(f"DEBUG: Mensagem pergunta nome animal apagada: {mensagem_nome_animal_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta nome animal: {e}")
        
        # Atualizar resumo
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id

        # 🔥 TEXTO "QUAL OPÇÃO SE ENQUADRA" POR IDIOMA
        textos_opcoes_animal = {
            'portugues': "🐾 *Qual das opções se enquadra melhor com o seu animal?*",
            'ingles': "🐾 *Which option best fits your animal?*",
            'espanhol': "🐾 *¿Cuál de las opciones se ajusta mejor a tu animal?*",
            'italiano': "🐾 *Quale opzione si adatta meglio al tuo animale?*",
            'alemao': "🐾 *Welche Option passt am besten zu Ihrem Tier?*",
            'frances': "🐾 *Quelle option correspond le mieux à votre animal ?*"
        }
        
        # 🔥 BOTÕES DE TIPO DE ANIMAL POR IDIOMA
        botoes_animal_por_idioma = {
            'portugues': {
                'cao': "🐶 Cão",
                'gato': "🐱 Gato",
                'reptil': "🦎 Réptil",
                'ave': "🐦 Ave",
                'roedor': "🐹 Roedor"
            },
            'ingles': {
                'cao': "🐶 Dog",
                'gato': "🐱 Cat",
                'reptil': "🦎 Reptile",
                'ave': "🐦 Bird",
                'roedor': "🐹 Rodent"
            },
            'espanhol': {
                'cao': "🐶 Perro",
                'gato': "🐱 Gato",
                'reptil': "🦎 Reptil",
                'ave': "🐦 Ave",
                'roedor': "🐹 Roedor"
            },
            'italiano': {
                'cao': "🐶 Cane",
                'gato': "🐱 Gatto",
                'reptil': "🦎 Rettile",
                'ave': "🐦 Uccello",
                'roedor': "🐹 Roditore"
            },
            'alemao': {
                'cao': "🐶 Hund",
                'gato': "🐱 Katze",
                'reptil': "🦎 Reptil",
                'ave': "🐦 Vogel",
                'roedor': "🐹 Nagetier"
            },
            'frances': {
                'cao': "🐶 Chien",
                'gato': "🐱 Chat",
                'reptil': "🦎 Reptile",
                'ave': "🐦 Oiseau",
                'roedor': "🐹 Rongeur"
            }
        }
        
        botoes = botoes_animal_por_idioma.get(idioma, botoes_animal_por_idioma['portugues'])

        teclado = [
           [InlineKeyboardButton(botoes['cao'], callback_data="tipo_cao"),
           InlineKeyboardButton(botoes['gato'], callback_data="tipo_gato")],
           [InlineKeyboardButton(botoes['reptil'], callback_data="tipo_reptil"),
           InlineKeyboardButton(botoes['ave'], callback_data="tipo_ave")],
           [InlineKeyboardButton(botoes['roedor'], callback_data="tipo_roedor")]
        ]

        mensagem_tipo_animal = await update.message.reply_text(
            textos_opcoes_animal.get(idioma, textos_opcoes_animal['portugues']),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
        context.user_data['mensagem_tipo_animal_id'] = mensagem_tipo_animal.message_id
        context.user_data['conversation_state'] = TIPO_ANIMAL
        print(f"✅ Estado: TIPO_ANIMAL | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_nome_animal: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))










async def tipo_animal_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber tipo de animal escolhido - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()

    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 TIPOS DE ANIMAIS TRADUZIDOS
    tipos_animais_por_idioma = {
        'portugues': {
            "tipo_cao": "Cão 🐶",
            "tipo_gato": "Gato 🐱", 
            "tipo_reptil": "Réptil 🦎",
            "tipo_ave": "Ave 🐦",
            "tipo_roedor": "Roedor 🐹"
        },
        'ingles': {
            "tipo_cao": "Dog 🐶",
            "tipo_gato": "Cat 🐱", 
            "tipo_reptil": "Reptile 🦎",
            "tipo_ave": "Bird 🐦",
            "tipo_roedor": "Rodent 🐹"
        },
        'espanhol': {
            "tipo_cao": "Perro 🐶",
            "tipo_gato": "Gato 🐱", 
            "tipo_reptil": "Reptil 🦎",
            "tipo_ave": "Ave 🐦",
            "tipo_roedor": "Roedor 🐹"
        },
        'italiano': {
            "tipo_cao": "Cane 🐶",
            "tipo_gato": "Gatto 🐱", 
            "tipo_reptil": "Rettile 🦎",
            "tipo_ave": "Uccello 🐦",
            "tipo_roedor": "Roditore 🐹"
        },
        'alemao': {
            "tipo_cao": "Hund 🐶",
            "tipo_gato": "Katze 🐱", 
            "tipo_reptil": "Reptil 🦎",
            "tipo_ave": "Vogel 🐦",
            "tipo_roedor": "Nagetier 🐹"
        },
        'frances': {
            "tipo_cao": "Chien 🐶",
            "tipo_gato": "Chat 🐱", 
            "tipo_reptil": "Reptile 🦎",
            "tipo_ave": "Oiseau 🐦",
            "tipo_roedor": "Rongeur 🐹"
        }
    }
    
    tipos_animais = tipos_animais_por_idioma.get(idioma, tipos_animais_por_idioma['portugues'])
    tipo_animal = tipos_animais.get(query.data, "")
    context.user_data["tipo_animal"] = tipo_animal
    
    print(f"✅ Tipo de animal selecionado: {tipo_animal} | Idioma: {idioma}")
    
    # 🔥 CORREÇÃO: Tentar remover mensagem da pergunta COM TRY/EXCEPT
    mensagem_tipo_animal_id = context.user_data.get('mensagem_tipo_animal_id')
    if mensagem_tipo_animal_id:
        try:
            await context.bot.delete_message(
                chat_id=query.message.chat_id,
                message_id=mensagem_tipo_animal_id
            )
            print(f"DEBUG: Mensagem tipo animal apagada: {mensagem_tipo_animal_id}")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem tipo animal: {e}")
    
    # 🔥 CORREÇÃO: Tentar remover a mensagem do callback COM TRY/EXCEPT
    try:
        await query.delete_message()
        print("DEBUG: Mensagem callback animal apagada com sucesso")
    except Exception as e:
        print(f"DEBUG: Não foi possível apagar mensagem callback animal: {e}")
        # Alternativa: editar a mensagem para esconder os botões
        try:
            # 🔥 TEXTO "ANIMAL SELECIONADO" POR IDIOMA
            textos_selecionado = {
                'portugues': f"✅ {tipo_animal} selecionado!",
                'ingles': f"✅ {tipo_animal} selected!",
                'espanhol': f"✅ ¡{tipo_animal} seleccionado!",
                'italiano': f"✅ {tipo_animal} selezionato!",
                'alemao': f"✅ {tipo_animal} ausgewählt!",
                'frances': f"✅ {tipo_animal} sélectionné !"
            }
            
            await query.edit_message_text(
                text=textos_selecionado.get(idioma, textos_selecionado['portugues']),
                reply_markup=None
            )
            print("DEBUG: Mensagem callback animal editada para esconder botões")
        except Exception as e2:
            print(f"DEBUG: Também não foi possível editar mensagem animal: {e2}")
            # Em último caso, não fazemos nada

    # Atualizar resumo
    current_resumo_msg_id = context.user_data.get('resumo_msg_id')
    new_message_id = await enviar_resumo(
        context, 
        query.message.chat_id, 
        message_id=current_resumo_msg_id
    )
    context.user_data['resumo_msg_id'] = new_message_id

    # 🔥 TEXTO "QUAL TAMANHO PARA ANIMAL" POR IDIOMA
    textos_tamanho_animal = {
        'portugues': "📏 *Qual tamanho gostaria para o seu Cartoon Animal?*",
        'ingles': "📏 *What size would you like for your Animal Cartoon?*",
        'espanhol': "📏 *¿Qué tamaño te gustaría para tu Caricatura Animal?*",
        'italiano': "📏 *Che dimensione vorresti per il tuo Cartoon Animale?*",
        'alemao': "📏 *Welche Größe möchten Sie für Ihre Tier-Karikatur?*",
        'frances': "📏 *Quelle taille souhaitez-vous pour votre Dessin Animé Animal ?*"
    }

    # 🔥 CORREÇÃO: Mostrar tamanhos para Animal (AGORA DENTRO DA FUNÇÃO)
    teclado = []
    tamanhos_lista = list(TAMANHOS_ANIMAL.items())

    # Processar em pares (2 botões por linha)
    for i in range(0, len(tamanhos_lista), 2):
        linha = []
        # Primeiro botão da linha
        tamanho_key1, info_tamanho1 = tamanhos_lista[i]
        preco_extra1 = f" (+€{info_tamanho1['preco']})" if info_tamanho1['preco'] > 0 else ""
        botao_texto1 = f"{info_tamanho1['nome']}"  # {preco_extra1}
        linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
        
        # Segundo botão da linha (se existir)
        if i + 1 < len(tamanhos_lista):
            tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
            preco_extra2 = f" (+€{info_tamanho2['preco']})" if info_tamanho2['preco'] > 0 else ""
            botao_texto2 = f"{info_tamanho2['nome']}"  # {preco_extra2}
            linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
        
        teclado.append(linha)

    mensagem_tamanhos = await context.bot.send_message(
        chat_id=query.message.chat_id,
        text=textos_tamanho_animal.get(idioma, textos_tamanho_animal['portugues']),
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(teclado)
    )

    context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
    context.user_data['conversation_state'] = TAMANHO




async def receber_nome_peca(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber nome da peça personalizada quando escolhe 'Outro' - COM TRADUÇÃO"""
    try:
        # 🔥 PEGAR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        nome_peca = update.message.text
        context.user_data["nome_peca_personalizado"] = nome_peca
        
        # 🔥 ATUALIZAR RESUMO IMEDIATAMENTE
        current_resumo_msg_id = context.user_data.get('resumo_msg_id')
        new_message_id = await enviar_resumo(
            context, 
            update.message.chat_id, 
            message_id=current_resumo_msg_id
        )
        context.user_data['resumo_msg_id'] = new_message_id
        
        print(f"✅ Nome da peça personalizada guardado: {nome_peca} | Idioma: {idioma}")
        
        # Remover mensagens
        try:
            await update.message.delete()
            print("DEBUG: Mensagem nome peça do usuário apagada")
        except Exception as e:
            print(f"DEBUG: Não foi possível apagar mensagem usuário: {e}")
            
        mensagem_nome_peca_id = context.user_data.get('mensagem_nome_peca_id')
        if mensagem_nome_peca_id:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat_id,
                    message_id=mensagem_nome_peca_id
                )
                print(f"DEBUG: Mensagem pergunta nome peça apagada: {mensagem_nome_peca_id}")
            except Exception as e:
                print(f"DEBUG: Não foi possível apagar pergunta nome peça: {e}")
        
        # 🔥 🔥 🔥 CORREÇÃO: AGORA PERGUNTA O NOME DO CARTOON (igual aos outros tipos)
        print("🎯 Fluxo: Personalizado Outro - perguntando nome do cartoon")
        
        # 🔥 TEXTO "NOME DO CARTOON PARA OUTRO" POR IDIOMA
        textos_nome_cartoon_outro = {
            'portugues': "🎭 *Escreve um nome ao seu cartoon para colocar na sua box!*\n\n*Exemplo:* `Sonho do Avô Fernando`, `Volta ao Mundo`",
            'ingles': "🎭 *Write a name for your cartoon to put on your box!*\n\n*Example:* `Grandpa Fernando's Dream`, `Around the World`",
            'espanhol': "🎭 *¡Escribe un nombre para tu caricatura para poner en tu caja!*\n\n*Ejemplo:* `Sueño del Abuelo Fernando`, `Vuelta al Mundo`",
            'italiano': "🎭 *Scrivi un nome per il tuo cartoon da mettere sulla tua scatola!*\n\n*Esempio:* `Sogno del Nonno Fernando`, `Giro del Mondo`",
            'alemao': "🎭 *Schreibe einen Namen für deine Karikatur, der auf deine Box kommt!*\n\n*Beispiel:* `Opa Fernandos Traum`, `Weltreise`",
            'frances': "🎭 *Écrivez un nom pour votre dessin animé à mettre sur votre boîte !*\n\n*Exemple :* `Rêve du Grand-père Fernando`, `Tour du Monde`"
        }
        
        mensagem_nome_personalizado = await update.message.reply_text(
            textos_nome_cartoon_outro.get(idioma, textos_nome_cartoon_outro['portugues']),
            parse_mode="Markdown"
        )
        context.user_data['mensagem_nome_personalizado_id'] = mensagem_nome_personalizado.message_id
        context.user_data['conversation_state'] = NOME_PERSONALIZADO
        print(f"✅ Estado definido para NOME_PERSONALIZADO: {NOME_PERSONALIZADO} | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_nome_peca: {e}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, inténtalo de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(textos_erro.get(idioma, textos_erro['portugues']))






async def mostrar_tamanhos_personalizado(context, chat_id):
    """Mostrar opções de tamanho para Cartoon Personalizado - COM TRADUÇÃO"""
    
    # 🔥 PEGAR IDIOMA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # TAMANHOS PARA PERSONALIZADO
    TAMANHOS_PERSONALIZADO = {
        "2.5": {"nome": " 2.5\" | 6.4cm ", "preco": 0.0},
        "3.5": {"nome": " 3.5\" | 8.9cm ", "preco": 5.0},
        "4.5": {"nome": " 4.5\" | 11.4cm ", "preco": 10.0},
        "6": {"nome": " 6\" | 15.2cm ", "preco": 25.0},
        "7": {"nome": " 7\" | 17.8cm ", "preco": 35.0},
        "8": {"nome": " 8\" | 20.3cm ", "preco": 55.0},
        "9": {"nome": " 9\" | 22.9cm ", "preco": 70.0},
        "10": {"nome": " 10\" | 25.4cm ", "preco": 90.0},
        "11": {"nome": " 11\" | 27.9cm ", "preco": 110.0},
        "12": {"nome": " 12\" | 30.5cm ", "preco": 150.0}
    }
    
    # 🔥 TEXTO "ESCOLHA O TAMANHO" POR IDIOMA
    textos_tamanho_personalizado = {
        'portugues': "📏 *Escolha o tamanho do seu Cartoon Personalizado:*",
        'ingles': "📏 *Choose the size of your Custom Cartoon:*",
        'espanhol': "📏 *Elige el tamaño de tu Caricatura Personalizada:*",
        'italiano': "📏 *Scegli la dimensione del tuo Cartoon Personalizzato:*",
        'alemao': "📏 *Wählen Sie die Größe Ihrer Personalisierten Karikatur:*",
        'frances': "📏 *Choisissez la taille de votre Dessin Animé Personnalisé :*"
    }
    
    # Criar botões dos tamanhos em GRADE 2xN
    teclado = []
    tamanhos_lista = list(TAMANHOS_PERSONALIZADO.items())
    
    # Processar em pares (2 botões por linha)
    for i in range(0, len(tamanhos_lista), 2):
        linha = []
        # Primeiro botão da linha
        tamanho_key1, info_tamanho1 = tamanhos_lista[i]
        botao_texto1 = f"{info_tamanho1['nome']}"  # {preco_total:.0f}€
        linha.append(InlineKeyboardButton(botao_texto1, callback_data=f"tamanho_{tamanho_key1}"))
        
        # Segundo botão da linha (se existir)
        if i + 1 < len(tamanhos_lista):
            tamanho_key2, info_tamanho2 = tamanhos_lista[i + 1]
            botao_texto2 = f"{info_tamanho2['nome']}"  # {preco_total:.0f}€
            linha.append(InlineKeyboardButton(botao_texto2, callback_data=f"tamanho_{tamanho_key2}"))
        
        teclado.append(linha)

    mensagem_tamanhos = await context.bot.send_message(
        chat_id=chat_id,
        text=textos_tamanho_personalizado.get(idioma, textos_tamanho_personalizado['portugues']),
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(teclado)
    )
    
    context.user_data['mensagem_tamanhos_id'] = mensagem_tamanhos.message_id
    context.user_data['conversation_state'] = TAMANHO
    print(f"✅ Tamanhos personalizados mostrados | Idioma: {idioma}")








# --- Handler para voltar ao início ---
async def voltar_inicio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA ANTES DE LIMPAR
    idioma = context.user_data.get('idioma', 'portugues')
    
    print(f"🔄 Voltando ao início | Idioma mantido: {idioma}")
    
    # 🔥 LIMPAR TODOS OS DADOS, MAS MANTER O IDIOMA
    context.user_data.clear()
    context.user_data['idioma'] = idioma
    
    # 🔥 APAGAR MENSAGEM ATUAL
    try:
        await query.delete_message()
        print("✅ Mensagem atual apagada para voltar ao início")
    except Exception as e:
        print(f"❌ Erro ao apagar mensagem: {e}")
    
    # 🔥 TEXTOS DO MENU INICIAL POR IDIOMA
    textos_menu = {
        'portugues': {
            'saudacao': "👋 Olá! Bem-vindo à *GodsPlan*, vamos criar o seu cartoon?",
            'botao': "🎭 CRIAR MEU CARTOON"
        },
        'ingles': {
            'saudacao': "👋 Hello! Welcome to *GodsPlan*, shall we create your cartoon?",
            'botao': "🎭 CREATE MY CARTOON"
        },
        'espanhol': {
            'saudacao': "👋 ¡Hola! Bienvenido a *GodsPlan*, ¿vamos a crear tu caricatura?",
            'botao': "🎭 CREAR MI CARICATURA"
        },
        'italiano': {
            'saudacao': "👋 Ciao! Benvenuto in *GodsPlan*, creiamo il tuo cartoon?",
            'botao': "🎭 CREA IL MIO CARTOON"
        },
        'alemao': {
            'saudacao': "👋 Hallo! Willkommen bei *GodsPlan*, sollen wir Ihre Karikatur erstellen?",
            'botao': "🎭 MEINE KARIKATUR ERSTELLEN"
        },
        'frances': {
            'saudacao': "👋 Bonjour ! Bienvenue chez *GodsPlan*, allons-nous créer votre dessin animé ?",
            'botao': "🎭 CRÉER MON DESSIN ANIMÉ"
        }
    }
    
    # Obter textos para o idioma atual
    textos = textos_menu.get(idioma, textos_menu['portugues'])
    
    # SEMPRE MOSTRAR MENU INICIAL
    keyboard = [[InlineKeyboardButton(textos['botao'], callback_data="mycartoon")]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.message.reply_text(
        textos['saudacao'], 
        reply_markup=reply_markup, 
        parse_mode="Markdown"
    )
    
    print(f"✅ Menu inicial mostrado em {idioma}")



    




async def pagar_original(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para pagamento do pedido original - ATUALIZADO COM MESMA ESTRUTURA E TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    
    print(f"🎯 PAGAR_ORIGINAL CHAMADO - VERSÃO ATUALIZADA | Idioma: {idioma}")
    
    # Extrair pedido_id do callback_data
    pedido_id = query.data.replace("pagar_original_", "")
    print(f"🔍 Procurando pedido: {pedido_id}")
    
    if pedido_id not in PEDIDOS_REGISTO:
        print(f"❌ Pedido não encontrado no registro: {pedido_id}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado. Por favor, inicie um novo pedido.",
            'ingles': "❌ Order not found. Please start a new order.",
            'espanhol': "❌ Pedido no encontrado. Por favor, inicie un nuevo pedido.",
            'italiano': "❌ Ordine non trovato. Per favore, inizi un nuovo ordine.",
            'alemao': "❌ Bestellung nicht gefunden. Bitte beginnen Sie eine neue Bestellung.",
            'frances': "❌ Commande introuvable. Veuillez démarrer une nouvelle commande."
        }
        
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    chat_id = query.message.chat_id
    
    # 🔥 CANCELAR QUALQUER TEMPORIZADOR ATIVO
    await cancelar_temporizadores_pedido(pedido_id)
    
    print(f"✅ Pedido encontrado: #{pedido_id}")
    print(f"🔍 Chat ID do cliente: {chat_id}")

    try:
        # 🔥 PASSO 1: DEFINIR MÉTODOS DE PAGAMENTO POR PAÍS - MESMA ESTRUTURA DO pagar_stripe
        def get_payment_methods(pais):
            """Retorna métodos de pagamento baseado no país"""
            
            def get_country_code(pais_nome):
                mapeamento_paises = {
                    "portugal": "PT",
                    "espanha": "ES", 
                    "franca": "FR",
                    "alemanha": "DE",
                    "belgica": "BE",
                    "reino unido": "GB",
                    "estados unidos": "US",
                    "paises baixos": "NL",
                    "brasil": "BR",
                    "irlanda": "IE",
                    "italia": "IT",
                    "luxemburgo": "LU",
                    "canada": "CA"
                }
                return mapeamento_paises.get(pais_nome.lower(), pais_nome.upper())
            
            country_code = get_country_code(pais)
            print(f"🔍 País recebido: '{pais}' → Código: '{country_code}'")
            
            # 🔥 MESMO payment_methods_by_country DO pagar_stripe
            payment_methods_by_country = {
                "PT": ["card", "paypal", "link", "klarna", "mb_way", "sepa_debit"],
                "ES": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "FR": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "DE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "BE": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "GB": ["card", "paypal", "link", "klarna"],
                "US": ["card", "paypal", "link"],
                "NL": ["card", "paypal", "link", "klarna", "ideal", "sepa_debit"],
                "BR": ["card", "link"],
                "IE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "IT": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "LU": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "CA": ["card", "paypal", "link"]
            }
            
            methods = payment_methods_by_country.get(country_code, ["card", "link"])
            print(f"💳 Métodos de pagamento para {pais} ({country_code}): {methods}")
            return methods

        # 🔥 OBTER MÉTODOS REAIS PARA ESTE PAÍS
        metodos_reais = get_payment_methods(pedido['pais'])
        
        # 🔥 CRIAR TEXTO DINÂMICO DOS MÉTODOS COM TRADUÇÃO
        def formatar_metodos(metodos, pais, idioma):
            """Formata os métodos de pagamento para exibição"""
            # 🔥 NOMES DOS MÉTODOS POR IDIOMA
            nomes_metodos_por_idioma = {
                'portugues': {
                    "card": "Cartão",
                    "paypal": "PayPal", 
                    "link": "Link (inclui Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'ingles': {
                    "card": "Card",
                    "paypal": "PayPal", 
                    "link": "Link (includes Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA Debit",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'espanhol': {
                    "card": "Tarjeta",
                    "paypal": "PayPal", 
                    "link": "Link (incluye Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'italiano': {
                    "card": "Carta",
                    "paypal": "PayPal", 
                    "link": "Link (include Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Addebito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'alemao': {
                    "card": "Karte",
                    "paypal": "PayPal", 
                    "link": "Link (inkl. Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA-Lastschrift",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'frances': {
                    "card": "Carte",
                    "paypal": "PayPal", 
                    "link": "Link (inclut Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Prélèvement SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                }
            }
            
            nomes_metodos = nomes_metodos_por_idioma.get(idioma, nomes_metodos_por_idioma['portugues'])
            textos = []
            
            for metodo in metodos:
                if metodo in nomes_metodos:
                    textos.append(nomes_metodos[metodo])
                else:
                    textos.append(metodo.capitalize())
            
            return ", ".join(textos)

        # 🔥 PASSO 2: VERIFICAR CONFIGURAÇÃO DE WALLETS - MESMA LÓGICA
        def verificar_config_wallets():
            """Verifica se as wallets estão configuradas corretamente"""
            try:
                apple_domains = stripe.ApplePayDomain.list()
                print("🍎 Domínios Apple Pay configurados:")
                for domain in apple_domains.data:
                    print(f"   - {domain.domain}")
                
                # Verificar domínio atual
                seu_dominio = "unceased-bibliothecal-donette.ngrok-free.dev"
                dominios_apple = [d.domain for d in apple_domains.data]
                if seu_dominio in dominios_apple:
                    print("✅ Domínio ngrok configurado no Apple Pay!")
                    return True
                else:
                    print("⚠️ Domínio ngrok NÃO configurado no Apple Pay")
                    return False
                    
            except Exception as e:
                print(f"❌ Erro ao verificar wallets: {e}")
                return False

        wallets_configuradas = verificar_config_wallets()

        # 🔥 TEXTOS TRADUZIDOS PARA O CHECKOUT (shipping_message e submit_message)
        textos_checkout_messages = {
            'portugues': {
                "shipping_message": "📦 Enviaremos o seu Cartoon personalizado para este endereço!",
                "submit_message": "✨ Obrigado! Vamos criar um Cartoon incrível para si!"
            },
            'ingles': {
                "shipping_message": "📦 We'll send your personalized Cartoon to this address!",
                "submit_message": "✨ Thank you! We'll create an amazing Cartoon for you!"
            },
            'espanhol': {
                "shipping_message": "📦 ¡Enviaremos tu Cartoon personalizado a esta dirección!",
                "submit_message": "✨ ¡Gracias! ¡Crearemos un Cartoon increíble para ti!"
            },
            'italiano': {
                "shipping_message": "📦 Spediremo il tuo Cartoon personalizzato a questo indirizzo!",
                "submit_message": "✨ Grazie! Creeremo un Cartoon incredibile per te!"
            },
            'alemao': {
                "shipping_message": "📦 Wir senden Ihren personalisierten Cartoon an diese Adresse!",
                "submit_message": "✨ Danke! Wir erstellen einen fantastischen Cartoon für Sie!"
            },
            'frances': {
                "shipping_message": "📦 Nous enverrons votre Cartoon personnalisé à cette adresse !",
                "submit_message": "✨ Merci ! Nous créerons un Cartoon incroyable pour vous !"
            }
        }
        
        textos_messages = textos_checkout_messages.get(idioma, textos_checkout_messages['portugues'])
        
        # 🔥 DESCRIÇÕES DO PRODUTO POR IDIOMA
        descricoes_produto = {
            'portugues': f"Pedido #{pedido_id} - Para {pedido['nome']}",
            'ingles': f"Order #{pedido_id} - For {pedido['nome']}",
            'espanhol': f"Pedido #{pedido_id} - Para {pedido['nome']}",
            'italiano': f"Ordine #{pedido_id} - Per {pedido['nome']}",
            'alemao': f"Bestellung #{pedido_id} - Für {pedido['nome']}",
            'frances': f"Commande #{pedido_id} - Pour {pedido['nome']}"
        }
        
        descricao_produto = descricoes_produto.get(idioma, descricoes_produto['portugues'])
        
        # 🔥 NOMES DO PRODUTO POR IDIOMA
        nomes_produto = {
            'portugues': f"Cartoon Personalizado - {pedido['tipo_cartoon']}",
            'ingles': f"Personalized Cartoon - {pedido['tipo_cartoon']}",
            'espanhol': f"Cartoon Personalizado - {pedido['tipo_cartoon']}",
            'italiano': f"Cartoon Personalizzato - {pedido['tipo_cartoon']}",
            'alemao': f"Personaliserter Cartoon - {pedido['tipo_cartoon']}",
            'frances': f"Dessin Animé Personnalisé - {pedido['tipo_cartoon']}"
        }
        
        nome_produto = nomes_produto.get(idioma, nomes_produto['portugues'])
        
        # 🔥 PASSO 3: CRIAR SESSÃO STRIPE - MESMA ESTRUTURA
        print("🔗 Criando Checkout Session para pagamento original...")
        
        session_config = {
            "payment_method_types": metodos_reais,
            "mode": "payment",
            "customer_email": pedido["email"],
            
            # 🔥 CONFIGURAÇÃO PARA WALLETS
            "payment_method_options": {
                "card": {
                    "request_three_d_secure": "automatic"
                }
            },
            
            "shipping_address_collection": {
                "allowed_countries": [
                    "PT", "ES", "FR", "DE", "BE", "GB", "US", "NL", "BR", "IE", "IT", "LU", "CA"
                ]
            },
            
            # 🔥 MENSAGENS TRADUZIDAS PARA O CHECKOUT
            "custom_text": {
                "shipping_address": {
                    "message": textos_messages["shipping_message"]
                },
                "submit": {
                    "message": textos_messages["submit_message"]
                }
            },
            
            "line_items": [{
                "price_data": {
                    "currency": pedido["moeda"].lower(),
                    "product_data": {
                        "name": nome_produto,  # 🔥 NOME TRADUZIDO
                        "description": descricao_produto,  # 🔥 DESCRIÇÃO TRADUZIDA
                    },
                    "unit_amount": int(pedido["total"] * 100),
                },
                "quantity": 1
            }],
            
            # 🔥 URLs CORRETAS
            "success_url": f"https://t.me/plan3d_bot?start=payment_success_{pedido_id}",
            "cancel_url": f"https://t.me/plan3d_bot?start=payment_cancelled_{pedido_id}",
            
            "metadata": {
                "pedido_id": pedido_id,
                "chat_id": str(chat_id),
                "pais": pedido['pais'],
                "moeda": pedido["moeda"],
                "total_pago": str(pedido["total"]),
                "nome_cliente": pedido['nome'],
                "tipo_cartoon": pedido['tipo_cartoon'],
                "tipo_sessao": "original",
                "wallets_habilitadas": str(wallets_configuradas),
                "idioma": idioma  # 🔥 ADICIONAR IDIOMA AO METADATA
            },
            
            "expires_at": int((datetime.now() + timedelta(minutes=30)).timestamp()),
        }

        # 🔥 CONFIGURAÇÃO ESPECÍFICA PARA WALLETS - MESMA LÓGICA
        paises_com_wallets = ["Reino Unido", "Estados Unidos", "Brasil", "Irlanda", 
                            "França", "Alemanha", "Itália", "Espanha", "Portugal", 
                            "Países Baixos", "Bélgica", "Luxemburgo", "Canadá"]
        
        if pedido['pais'] in paises_com_wallets and "link" in metodos_reais:
            print(f"📱 Configurando Apple Pay/Google Pay para {pedido['pais']}")
            session_config["payment_method_options"]["link"] = {"persistent_token": None}

        # 🔥 CRIAR A SESSÃO
        session = stripe.checkout.Session.create(**session_config)

        print(f"✅ CHECKOUT SESSION CRIADA: {session.id}")
        print(f"🔗 URL do Checkout: {session.url}")

        # 🔥 PASSO 4: ATUALIZAR PEDIDO
        pedido["session_id_original"] = session.id
        pedido["payment_intent_id"] = session.payment_intent
        pedido["wallets_configuradas"] = wallets_configuradas
        pedido["data_pagamento_original"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pedido["idioma"] = idioma  # 🔥 GUARDAR IDIOMA NO PEDIDO
        
        print(f"📊 Pedido atualizado para pagamento original")

        # 🔥 PASSO 5: MENSAGEM FINAL COM INSTRUÇÕES CLARAS - COM TRADUÇÃO
        texto_metodos = formatar_metodos(metodos_reais, pedido['pais'], idioma)
        
        # 🔥 TEXTOS DO CHECKOUT POR IDIOMA
        textos_checkout = {
            'portugues': {
                'titulo': "💳 *CHECKOUT DE PAGAMENTO* 💳",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envio:*",
                'moeda': "💰 *Moeda:*",
                'total_pagar': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido:",
                'checkout_pedido': "📋 *No checkout será pedido:*",
                'endereco': "1️⃣ **Endereço de entrega completo**",
                'metodo': "2️⃣ **Método de pagamento**",
                'metodos_disponiveis': "💳 *Métodos disponíveis:*",
                'seguro': "🔒 *Pagamento 100% seguro via Stripe*",
                'tempo': "⏰ *Tem 30 minutos para efetuar o pagamento*",
                'clique_abaixo': "Clique abaixo para pagar: 👇",
                'botao': "💳 PAGAR AGORA →"
            },
            'ingles': {
                'titulo': "💳 *PAYMENT CHECKOUT* 💳",
                'cliente': "👤 *Customer:*",
                'pais': "🌍 *Shipping Country:*",
                'moeda': "💰 *Currency:*",
                'total_pagar': "💳 **TOTAL TO PAY:",
                'pedido': "🆔 **Order:",
                'checkout_pedido': "📋 *In checkout you will be asked for:*",
                'endereco': "1️⃣ **Complete shipping address**",
                'metodo': "2️⃣ **Payment method**",
                'metodos_disponiveis': "💳 *Available methods:*",
                'seguro': "🔒 *100% secure payment via Stripe*",
                'tempo': "⏰ *You have 30 minutes to complete payment*",
                'clique_abaixo': "Click below to pay: 👇",
                'botao': "💳 PAY NOW →"
            },
            'espanhol': {
                'titulo': "💳 *CHECKOUT DE PAGO* 💳",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envío:*",
                'moeda': "💰 *Moneda:*",
                'total_pagar': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido:",
                'checkout_pedido': "📋 *En el checkout se pedirá:*",
                'endereco': "1️⃣ **Dirección de envío completa**",
                'metodo': "2️⃣ **Método de pago**",
                'metodos_disponiveis': "💳 *Métodos disponibles:*",
                'seguro': "🔒 *Pago 100% seguro vía Stripe*",
                'tempo': "⏰ *Tienes 30 minutos para efectuar el pago*",
                'clique_abaixo': "Haz clic abajo para pagar: 👇",
                'botao': "💳 PAGAR AHORA →"
            },
            'italiano': {
                'titulo': "💳 *CHECKOUT DI PAGAMENTO* 💳",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *Paese di Spedizione:*",
                'moeda': "💰 *Valuta:*",
                'total_pagar': "💳 **TOTALE DA PAGARE:",
                'pedido': "🆔 **Ordine:",
                'checkout_pedido': "📋 *Nel checkout verrà richiesto:*",
                'endereco': "1️⃣ **Indirizzo di spedizione completo**",
                'metodo': "2️⃣ **Metodo di pagamento**",
                'metodos_disponiveis': "💳 *Metodi disponibili:*",
                'seguro': "🔒 *Pagamento 100% sicuro tramite Stripe*",
                'tempo': "⏰ *Hai 30 minuti per effettuare il pagamento*",
                'clique_abaixo': "Clicca qui sotto per pagare: 👇",
                'botao': "💳 PAGA ORA →"
            },
            'alemao': {
                'titulo': "💳 *ZAHLUNGS-CHECKOUT* 💳",
                'cliente': "👤 *Kunde:*",
                'pais': "🌍 *Versandland:*",
                'moeda': "💰 *Währung:*",
                'total_pagar': "💳 **GESAMTBETRAG ZU ZAHLEN:",
                'pedido': "🆔 **Bestellung:",
                'checkout_pedido': "📋 *Im Checkout wird angefordert:*",
                'endereco': "1️⃣ **Vollständige Lieferadresse**",
                'metodo': "2️⃣ **Zahlungsmethode**",
                'metodos_disponiveis': "💳 *Verfügbare Methoden:*",
                'seguro': "🔒 *100% sichere Zahlung über Stripe*",
                'tempo': "⏰ *Sie haben 30 Minuten für die Zahlung*",
                'clique_abaixo': "Klicken Sie unten zum Bezahlen: 👇",
                'botao': "💳 JETZT BEZAHLEN →"
            },
            'frances': {
                'titulo': "💳 *CHECKOUT DE PAIEMENT* 💳",
                'cliente': "👤 *Client:*",
                'pais': "🌍 *Pays de Livraison:*",
                'moeda': "💰 *Devise:*",
                'total_pagar': "💳 **TOTAL À PAYER:",
                'pedido': "🆔 **Commande:",
                'checkout_pedido': "📋 *Dans le checkout, il sera demandé:*",
                'endereco': "1️⃣ **Adresse de livraison complète**",
                'metodo': "2️⃣ **Méthode de paiement**",
                'metodos_disponiveis': "💳 *Méthodes disponibles:*",
                'seguro': "🔒 *Paiement 100% sécurisé via Stripe*",
                'tempo': "⏰ *Vous avez 30 minutes pour effectuer le paiement*",
                'clique_abaixo': "Cliquez ci-dessous pour payer : 👇",
                'botao': "💳 PAYER MAINTENANT →"
            }
        }
        
        textos = textos_checkout.get(idioma, textos_checkout['portugues'])

        await query.edit_message_text(
            text=(
                f"{textos['titulo']}\n\n"
                f"{textos['cliente']} {pedido['nome']}\n"
                f"{textos['pais']} {pedido['pais']}\n"
                f"{textos['moeda']} {pedido['moeda'].upper()} {pedido['simbolo_moeda']}\n\n"
                f"{textos['total_pagar']} {pedido['simbolo_moeda']}{pedido['total']:.2f}**\n"
                f"{textos['pedido']} #{pedido_id}**\n\n"
                f"{textos['checkout_pedido']}\n"
                f"{textos['endereco']}\n"
                f"{textos['metodo']}\n\n"
                f"{textos['metodos_disponiveis']} {texto_metodos}\n"
                f"{textos['seguro']}\n\n"
                f"{textos['tempo']}\n\n"
                f"{textos['clique_abaixo']}"
            ),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos['botao'], url=session.url)]
            ])
        )
        
        print(f"✅ Usuário redirecionado para Checkout (Pagamento Original) | Idioma: {idioma}")

        # 🔥 INICIAR TEMPORIZADOR (será cancelado pelo webhook quando pagamento for feito)
        await iniciar_temporizador_pagamento_original(context, pedido_id, chat_id, query.message.message_id, idioma)
        
    except Exception as e:
        print(f"❌ ERRO STRIPE NO PAGAMENTO ORIGINAL: {str(e)}")
        print(f"🔍 Tipo do erro: {type(e)}")
        
        import traceback
        print(f"🔍 Traceback completo: {traceback.format_exc()}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro_pagamento = {
            'portugues': {
                'erro': "❌ Erro no processamento do pagamento.",
                'tentar': "🔄 Tentar Novamente",
                'suporte': "📞 Suporte"
            },
            'ingles': {
                'erro': "❌ Error processing payment.",
                'tentar': "🔄 Try Again",
                'suporte': "📞 Support"
            },
            'espanhol': {
                'erro': "❌ Error en el procesamiento del pago.",
                'tentar': "🔄 Intentar de Nuevo",
                'suporte': "📞 Soporte"
            },
            'italiano': {
                'erro': "❌ Errore nell'elaborazione del pagamento.",
                'tentar': "🔄 Riprova",
                'suporte': "📞 Supporto"
            },
            'alemao': {
                'erro': "❌ Fehler bei der Zahlungsverarbeitung.",
                'tentar': "🔄 Erneut versuchen",
                'suporte': "📞 Support"
            },
            'frances': {
                'erro': "❌ Erreur lors du traitement du paiement.",
                'tentar': "🔄 Réessayer",
                'suporte': "📞 Support"
            }
        }
        
        textos_erro = textos_erro_pagamento.get(idioma, textos_erro_pagamento['portugues'])
        
        await query.edit_message_text(
            f"{textos_erro['erro']}\n"
            "Por favor, tente novamente em alguns segundos.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos_erro['tentar'], callback_data=f"pagar_original_{pedido_id}")],
                [InlineKeyboardButton(textos_erro['suporte'], callback_data=f"todas_recusadas_{pedido_id}")]
            ])
        )





async def iniciar_temporizador_pagamento_original(context, pedido_id, chat_id, message_id, idioma='portugues'):
    """Temporizador de 30 minutos para pagamento original - COM TRADUÇÃO"""
    try:
        print(f"⏰⏰⏰ INICIAR_TEMPORIZADOR_PAGAMENTO_ORIGINAL para #{pedido_id} (30 minutos) | Idioma: {idioma}")
        
        # 🔥 INICIAR TASK DIRETAMENTE
        task = asyncio.create_task(temporizador_pagamento_original_task(context, pedido_id, chat_id, message_id, idioma))
        PEDIDOS_REGISTO[pedido_id]["timer_task_original"] = task
        print(f"✅✅✅ Task temporizador pagamento original criada para #{pedido_id}")
        
    except Exception as e:
        print(f"❌❌❌ Erro ao iniciar temporizador de pagamento original: {e}")

async def temporizador_pagamento_original_task(context, pedido_id, chat_id, message_id, idioma='portugues'):
    """Task do temporizador de pagamento original - COM TRADUÇÃO"""
    try:
        print(f"⏰ Task temporizador pagamento original iniciada para #{pedido_id} | Idioma: {idioma}")
        await asyncio.sleep(1800)  # 30 minutos
        
        print(f"🔍 Verificando se pagamento original #{pedido_id} ainda está ativo...")
        
        if (pedido_id in PEDIDOS_REGISTO and 
            PEDIDOS_REGISTO[pedido_id].get("status") not in ["pago", "processando"] and
            "timer_task_original" in PEDIDOS_REGISTO[pedido_id]):
            
            pedido = PEDIDOS_REGISTO[pedido_id]
            
            print("=" * 70)
            print(f"❌ PAGAMENTO ORIGINAL EXPIRADO: #{pedido_id}")
            print(f"👤 {pedido['nome']} | 🎨 {pedido['tipo_cartoon']} | 💰 {pedido['simbolo_moeda']}{pedido['total']:.2f} EXPIRADO")
            print("=" * 70)
            
            # Atualizar estatísticas
            atualizar_estatistica("pedidos_expirados")
            
            # Atualizar status do pedido
            pedido["status"] = "expirado"
            pedido["data_expiracao"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            
            # 🔥 TEXTOS DE EXPIRAÇÃO POR IDIOMA
            textos_expiracao = {
                'portugues': {
                    'titulo': "❌ *PAGAMENTO EXPIRADO* ❌",
                    'mensagem': "O tempo para efetuar o pagamento do pedido original expirou.",
                    'pedido': "🆔 *Pedido:*",
                    'valor': "💰 *Valor:*",
                    'cliente': "👤 *Cliente:*",
                    'deseja': "*Deseja tentar novamente ou ver ofertas especiais?*",
                    'tentar': "🔄 Tentar Novamente",
                    'ofertas': "📞 Suporte"
                },
                'ingles': {
                    'titulo': "❌ *PAYMENT EXPIRED* ❌",
                    'mensagem': "The time to complete payment for the original order has expired.",
                    'pedido': "🆔 *Order:*",
                    'valor': "💰 *Amount:*",
                    'cliente': "👤 *Customer:*",
                    'deseja': "*Would you like to try again or see special offers?*",
                    'tentar': "🔄 Try Again",
                    'ofertas': "📞 Support"
                },
                'espanhol': {
                    'titulo': "❌ *PAGO EXPIRADO* ❌",
                    'mensagem': "El tiempo para efectuar el pago del pedido original ha expirado.",
                    'pedido': "🆔 *Pedido:*",
                    'valor': "💰 *Valor:*",
                    'cliente': "👤 *Cliente:*",
                    'deseja': "*¿Desea intentar de nuevo o ver ofertas especiales?*",
                    'tentar': "🔄 Intentar de Nuevo",
                    'ofertas': "📞 Soporte"
                },
                'italiano': {
                    'titulo': "❌ *PAGAMENTO SCADUTO* ❌",
                    'mensagem': "Il tempo per completare il pagamento dell'ordine originale è scaduto.",
                    'pedido': "🆔 *Ordine:*",
                    'valor': "💰 *Importo:*",
                    'cliente': "👤 *Cliente:*",
                    'deseja': "*Vuoi riprovare o vedere offerte speciali?*",
                    'tentar': "🔄 Riprova",
                    'ofertas': "📞 Supporto"
                },
                'alemao': {
                    'titulo': "❌ *ZAHLUNG ABGELAUFEN* ❌",
                    'mensagem': "Die Zeit für die Zahlung der ursprünglichen Bestellung ist abgelaufen.",
                    'pedido': "🆔 *Bestellung:*",
                    'valor': "💰 *Betrag:*",
                    'cliente': "👤 *Kunde:*",
                    'deseja': "*Möchten Sie es erneut versuchen oder Sonderangebote sehen?*",
                    'tentar': "🔄 Erneut versuchen",
                    'ofertas': "📞 Support"
                },
                'frances': {
                    'titulo': "❌ *PAIEMENT EXPIRÉ* ❌",
                    'mensagem': "Le temps pour effectuer le paiement de la commande originale a expiré.",
                    'pedido': "🆔 *Commande:*",
                    'valor': "💰 *Montant:*",
                    'cliente': "👤 *Client:*",
                    'deseja': "*Souhaitez-vous réessayer ou voir des offres spéciales ?*",
                    'tentar': "🔄 Réessayer",
                    'ofertas': "📞 Support"
                }
            }
            
            textos = textos_expiracao.get(idioma, textos_expiracao['portugues'])
            
            # MENSAGEM FINAL
            await context.bot.edit_message_text(
                chat_id=chat_id,
                message_id=message_id,
                text=(
                    f"{textos['titulo']}\n\n"
                    f"{textos['mensagem']}\n\n"
                    f"{textos['pedido']} #{pedido_id}\n"
                    f"{textos['valor']} {pedido.get('simbolo_moeda', '')}{pedido['total']:.2f}\n"
                    f"{textos['cliente']} {pedido['nome']}\n\n"
                    f"{textos['deseja']}"
                ),
                parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton(textos['tentar'], callback_data=f"recuperar_pagar_{pedido_id}")],
                    [InlineKeyboardButton(textos['ofertas'], callback_data=f"todas_recusadas_{pedido_id}")]
                ])
            )
            
            print(f"✅ Mensagem de expiração enviada para #{pedido_id} | Idioma: {idioma}")
            
    except asyncio.CancelledError:
        print(f"✅✅✅ Temporizador pagamento original CANCELADO - Pedido #{pedido_id} PAGO")
    except Exception as e:
        print(f"❌❌❌ Erro na task do temporizador de pagamento original: {e}")









async def processar_pagamento_direto(context, pedido, chat_id, message_id):
    """Processa pagamento direto para qualquer oferta - COM TRADUÇÃO"""
    
    # 🔥 PEGAR IDIOMA DO PEDIDO OU USER_DATA
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"💳 Processando pagamento direto | Idioma: {idioma}")
    
    # ✅ CORREÇÃO: Verificar se estamos em modo de teste de forma SEGURA
    modo_teste = False
    
    if STRIPE_SECRET_KEY:
        if STRIPE_SECRET_KEY.startswith('sk_test_'):
            modo_teste = True
            print(f"🔧 Modo de TESTE detectado")
        elif STRIPE_SECRET_KEY.startswith('sk_live_'):
            modo_teste = False
            print(f"🚀 Modo de PRODUÇÃO detectado")
        else:
            # Chave inválida
            print(f"⚠️ Formato de chave Stripe inválido")
            modo_teste = True
    else:
        # Stripe não configurado
        print(f"⚠️ Stripe não configurado")
        modo_teste = True
    
    # Se for modo teste, pode fazer pagamento simulado
    if modo_teste:
        # PAGAMENTO SIMULADO
        # 🔥 TEXTOS DE SUCESSO POR IDIOMA
        textos_sucesso = {
            'portugues': {
                'titulo': "🎉 *PAGAMENTO PROCESSADO COM SUCESSO!*",
                'pedido': "🆔 *Pedido:*",
                'valor': "💵 *Valor:*",
                'produto': "📦 *Produto:*",
                'agradecimento': "✨ *Obrigado pela sua encomenda!*",
                'mensagem': "A nossa equipa já começou a trabalhar na sua obra de arte exclusiva!"
            },
            'ingles': {
                'titulo': "🎉 *PAYMENT PROCESSED SUCCESSFULLY!*",
                'pedido': "🆔 *Order:*",
                'valor': "💵 *Amount:*",
                'produto': "📦 *Product:*",
                'agradecimento': "✨ *Thank you for your order!*",
                'mensagem': "Our team has already started working on your exclusive artwork!"
            },
            'espanhol': {
                'titulo': "🎉 *¡PAGO PROCESADO CON ÉXITO!*",
                'pedido': "🆔 *Pedido:*",
                'valor': "💵 *Valor:*",
                'produto': "📦 *Producto:*",
                'agradecimento': "✨ *¡Gracias por su pedido!*",
                'mensagem': "¡Nuestro equipo ya comenzó a trabajar en su obra de arte exclusiva!"
            },
            'italiano': {
                'titulo': "🎉 *PAGAMENTO ELABORATO CON SUCCESSO!*",
                'pedido': "🆔 *Ordine:*",
                'valor': "💵 *Importo:*",
                'produto': "📦 *Prodotto:*",
                'agradecimento': "✨ *Grazie per il tuo ordine!*",
                'mensagem': "Il nostro team ha già iniziato a lavorare sulla tua opera d'arte esclusiva!"
            },
            'alemao': {
                'titulo': "🎉 *ZAHLUNG ERFOLGREICH VERARBEITET!*",
                'pedido': "🆔 *Bestellung:*",
                'valor': "💵 *Betrag:*",
                'produto': "📦 *Produkt:*",
                'agradecimento': "✨ *Vielen Dank für Ihre Bestellung!*",
                'mensagem': "Unser Team hat bereits mit der Arbeit an Ihrem exklusiven Kunstwerk begonnen!"
            },
            'frances': {
                'titulo': "🎉 *PAIEMENT TRAITÉ AVEC SUCCÈS !*",
                'pedido': "🆔 *Commande:*",
                'valor': "💵 *Montant:*",
                'produto': "📦 *Produit:*",
                'agradecimento': "✨ *Merci pour votre commande !*",
                'mensagem': "Notre équipe a déjà commencé à travailler sur votre œuvre d'art exclusive !"
            }
        }
        
        textos = textos_sucesso.get(idioma, textos_sucesso['portugues'])
        
        # 🔥 OBTER SÍMBOLO DE MOEDA CORRETO
        simbolo_moeda = pedido.get('simbolo_moeda', '€')
        
        await context.bot.edit_message_text(
            chat_id=chat_id,
            message_id=message_id,
            text=(
                f"{textos['titulo']}\n\n"
                f"{textos['pedido']} #{pedido['id']}\n"
                f"{textos['valor']} {simbolo_moeda}{pedido['total']:.2f}\n"
                f"{textos['produto']} {pedido['tipo_cartoon']}\n\n"
                f"{textos['agradecimento']}\n"
                f"{textos['mensagem']}"
            ),
            parse_mode="Markdown"
        )
        
        # ATUALIZAR STATUS PARA PAGO
        pedido["status"] = "pago"
        pedido["data_pagamento"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        # Atualizar estatística (certifique-se de que esta função existe)
        if 'atualizar_estatistica' in globals():
            atualizar_estatistica("pedidos_pagos")
        
        print(f"✅ Pagamento simulado processado | Idioma: {idioma}")
        
    else:
        # PAGAMENTO REAL COM STRIPE
        try:
            # 🔥 TEXTOS DO PAGAMENTO REAL POR IDIOMA
            textos_pagamento_real = {
                'portugues': {
                    'titulo': "🔒 *Pagamento Seguro*",
                    'mensagem': "Clique abaixo para finalizar o pagamento:",
                    'botao': "💳 Pagar Agora"
                },
                'ingles': {
                    'titulo': "🔒 *Secure Payment*",
                    'mensagem': "Click below to complete payment:",
                    'botao': "💳 Pay Now"
                },
                'espanhol': {
                    'titulo': "🔒 *Pago Seguro*",
                    'mensagem': "Haz clic abajo para completar el pago:",
                    'botao': "💳 Pagar Ahora"
                },
                'italiano': {
                    'titulo': "🔒 *Pagamento Sicuro*",
                    'mensagem': "Clicca qui sotto per completare il pagamento:",
                    'botao': "💳 Paga Ora"
                },
                'alemao': {
                    'titulo': "🔒 *Sichere Zahlung*",
                    'mensagem': "Klicken Sie abaixo, para completar el pago:",
                    'botao': "💳 Jetzt Bezahlen"
                },
                'frances': {
                    'titulo': "🔒 *Paiement Sécurisé*",
                    'mensagem': "Cliquez ci-dessous para finaliser le paiement :",
                    'botao': "💳 Payer Maintenant"
                }
            }
            
            textos = textos_pagamento_real.get(idioma, textos_pagamento_real['portugues'])
            
            # ✅ CERTIFIQUE-SE QUE O STRIPE ESTÁ IMPORTADO
            import stripe
            
            session = stripe.checkout.Session.create(
                payment_method_types=["card"],
                mode="payment",
                customer_email=pedido["email"],
                line_items=[{
                    "price_data": {
                        "currency": pedido.get("moeda", "eur").lower(),
                        "product_data": {"name": f"{pedido['tipo_cartoon']} - {pedido.get('estilo_cartoon', '')}"},
                        "unit_amount": int(pedido["total"] * 100)
                    },
                    "quantity": 1
                }],
                success_url="https://teusite.com/sucesso",
                cancel_url="https://teusite.com/cancelado"
            )
            
            botoes = [[InlineKeyboardButton(textos['botao'], url=session.url)]]
            
            await context.bot.edit_message_text(
                chat_id=chat_id,
                message_id=message_id,
                text=(
                    f"{textos['titulo']}\n\n"
                    f"{textos['mensagem']}"
                ),
                parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup(botoes)
            )
            
            print(f"✅ Checkout Stripe criado | Idioma: {idioma}")
            
        except Exception as e:
            print(f"❌ Erro no Stripe: {e}")
            
            # 🔥 TEXTOS DE ERRO POR IDIOMA
            textos_erro = {
                'portugues': "❌ *Erro no processamento*\n\nPor favor, tente novamente.",
                'ingles': "❌ *Processing Error*\n\nPlease try again.",
                'espanhol': "❌ *Error en el procesamiento*\n\nPor favor, intente de nuevo.",
                'italiano': "❌ *Errore nell'elaborazione*\n\nPer favore, riprova.",
                'alemao': "❌ *Verarbeitungsfehler*\n\nBitte versuchen Sie es erneut.",
                'frances': "❌ *Erreur de traitement*\n\nVeuillez réessayer."
            }
            
            await context.bot.edit_message_text(
                chat_id=chat_id,
                message_id=message_id,
                text=textos_erro.get(idioma, textos_erro['portugues']),
                parse_mode="Markdown"
            )


async def mostrar_oferta_tamanho_45(context, pedido, chat_id, message_id):
    """Mostra oferta do tamanho 4.5cm - COM 20% EXATOS E TRADUÇÃO"""
    
    # 🔥 OBTER IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    # 🔥 OBTER PAÍS DO PEDIDO
    pais = pedido.get('pais', '').lower()
    print(f"🌍 País no pedido (para Klarna): {pais}")
    
    # 🔥 LISTA DE PAÍSES QUE NÃO SUPORTAM KLARNA
    paises_sem_klarna = ["canada", "brasil", "estados unidos", "united states", "usa", "us"]
    
    # 🔥 VERIFICAR SE O PAÍS ESTÁ NA LISTA DE NÃO SUPORTE
    tem_klarna = True
    for pais_sem_klarna in paises_sem_klarna:
        if pais_sem_klarna in pais:
            tem_klarna = False
            print(f"🚫 Klarna NÃO disponível para: {pais}")
            break
    
    if tem_klarna:
        print(f"✅ Klarna disponível para: {pais}")
    
    # 🔥 OBTER MOEDA E SÍMBOLO CORRETOS DO PEDIDO
    moeda = pedido.get('moeda', 'EUR')
    simbolo_moeda = pedido.get('simbolo_moeda', '€')
    total_original = pedido.get('total_pago_original', pedido.get('total', 0))
    
    # Calcular preço do tamanho 4.5cm com 20% exatos
    total_45 = calcular_preco_tamanho_45(pedido)
    
    # Calcular economia (será exatamente 20%)
    economia = total_original - total_45
    
    # 🔥 AGORA SERÁ SEMPRE 20%
    percentual_desconto = 20

    valor_3x = total_45 / 3

    # 🔥 GUARDAR VALORES REAIS SEM SOBRESCREVER TAMANHO ORIGINAL
    pedido['valor_original_real'] = total_original  # 🔥 GUARDAR VALOR ORIGINAL
    pedido['valor_oferta_45_real'] = total_45  # 🔥 GUARDAR VALOR OFERTA 4.5cm

    print(f"💰 OFERTA 4.5cm GUARDADA | Idioma: {idioma}:")
    print(f"   • Valor Original: {simbolo_moeda}{total_original:.2f}")
    print(f"   • Valor Oferta 4.5cm: {simbolo_moeda}{total_45:.2f}")
    print(f"   • Tamanho Original: {pedido.get('tamanho_original', 'N/A')}")
    print(f"   • Economia: {simbolo_moeda}{economia:.2f} ({percentual_desconto}%)")
    print(f"   • Klarna disponível: {tem_klarna}")
    
    # 🔥 SÓ GUARDAR TAMANHO ORIGINAL SE AINDA NÃO EXISTIR (PROTEÇÃO)
    if 'tamanho_original' not in pedido:
        pedido['tamanho_original'] = pedido.get('tamanho_cartoon', '')
        print(f"💰 TAMANHO ORIGINAL GUARDADO: {pedido['tamanho_original']}")

    print(f"🔍 DEBUG mostrar_oferta_tamanho_45 - ANTES | Idioma: {idioma}:")
    print(f"   • tamanho_original: {pedido.get('tamanho_original', 'NÃO ENCONTRADO')}")
    print(f"   • tamanho_cartoon: {pedido.get('tamanho_cartoon', 'NÃO ENCONTRADO')}")

    # 🔥 TEXTOS DA OFERTA POR IDIOMA (COM VERSÕES COM/SEM KLARNA)
    textos_oferta = {
        'portugues': {
            'titulo': "🎉 *Temos uma opção ESPETACULAR para si!*",
            'subtitulo': "🌟 *CARTOON 3D - EDIÇÃO COLECIONADOR*",
            'qualidade': "• Mesma qualidade premium do original",
            'personalizado': "• Totalmente personalizado como pediu", 
            'tamanho': "• Tamanho perfeito (4.5\" | 11.5cm) para a secretária",
            'acabamento': "• Acabamento Premium",
            'valor_original': "• *Valor Original:*",
            'desconto': "• *DESCONTO*",
            'oferta_exclusiva': "🔥 *Oferta Exclusiva:*",
            'klarna': "💳 *Klarna:* 3x de",
            'sem_juros': "SEM JUROS",
            'klarna_indisponivel': "💳 *Opções de pagamento flexíveis disponíveis*",
            'pega_unica': "*Uma peça única a um preço irresistível!* ✨",
            'unica_oportunidade': "*A única oportunidade de adquirir está aqui!* 👇",
            'botao_sim': "✅ Sim, Quero Adquirir!",
            'botao_nao': "❌ Recusar Última Oportunidade."
        },
        'ingles': {
            'titulo': "🎉 *We have a SPECTACULAR option for you!*",
            'subtitulo': "🌟 *3D CARTOON - COLLECTOR'S EDITION*",
            'qualidade': "• Same premium quality as the original",
            'personalizado': "• Fully customized as you requested", 
            'tamanho': "• Perfect size (4.5\" | 11.5cm) for your desk",
            'acabamento': "• Premium Finish",
            'valor_original': "• *Original Price:*",
            'desconto': "• *DISCOUNT*",
            'oferta_exclusiva': "🔥 *Exclusive Offer:*",
            'klarna': "💳 *Klarna:* 3 installments of",
            'sem_juros': "NO INTEREST",
            'klarna_indisponivel': "💳 *Flexible payment options available*",
            'pega_unica': "*A unique piece at an irresistible price!* ✨",
            'unica_oportunidade': "*The only opportunity to get it is here!* 👇",
            'botao_sim': "✅ Yes, I Want to Get It!",
            'botao_nao': "❌ Reject Last Opportunity."
        },
        'espanhol': {
            'titulo': "🎉 *¡Tenemos uma opción ESPECTACULAR para ti!*",
            'subtitulo': "🌟 *CARICATURA 3D - EDICIÓN COLECCIONISTA*",
            'qualidade': "• Misma calidad premium que el original",
            'personalizado': "• Totalmente personalizado como lo pediste", 
            'tamanho': "• Tamaño perfecto (4.5\" | 11.5cm) para el escritorio",
            'acabamento': "• Acabado Premium",
            'valor_original': "• *Precio Original:*",
            'desconto': "• *DESCUENTO*",
            'oferta_exclusiva': "🔥 *Oferta Exclusiva:*",
            'klarna': "💳 *Klarna:* 3 cuotas de",
            'sem_juros': "SIN INTERESES",
            'klarna_indisponivel': "💳 *Opciones de pago flexibles disponibles*",
            'pega_unica': "*¡Una pieza única a un precio irresistible!* ✨",
            'unica_oportunidade': "*¡La única oportunidad de adquirirla está aquí!* 👇",
            'botao_sim': "✅ Sí, ¡Quiero Adquirirla!",
            'botao_nao': "❌ Rechazar Última Oportunidad."
        },
        'italiano': {
            'titulo': "🎉 *Abbiamo un'opzione SPETTACOLARE per te!*",
            'subtitulo': "🌟 *CARTOON 3D - EDIZIONE DA COLLEZIONE*",
            'qualidade': "• Stessa qualità premium dell'originale",
            'personalizado': "• Totalmente personalizzato come richiesto", 
            'tamanho': "• Taglia perfetta (4.5\" | 11.5cm) per la scrivania",
            'acabamento': "• Finitura Premium",
            'valor_original': "• *Prezzo Originale:*",
            'desconto': "• *SCONTO*",
            'oferta_exclusiva': "🔥 *Offerta Esclusiva:*",
            'klarna': "💳 *Klarna:* 3 rate da",
            'sem_juros': "SENZA INTERESSI",
            'klarna_indisponivel': "💳 *Opzioni di pagamento flessibili disponibili*",
            'pega_unica': "*Un pezzo unico a un prezzo irresistibile!* ✨",
            'unica_oportunidade': "*L'unica opportunità per acquistarlo è qui!* 👇",
            'botao_sim': "✅ Sì, Voglio Acquistarlo!",
            'botao_nao': "❌ Rifiuta Ultima Opportunità."
        },
        'alemao': {
            'titulo': "🎉 *Wir haben eine SPEKTAKULÄRE Option für Sie!*",
            'subtitulo': "🌟 *3D-KARIKATUR - SAMMLEREDITION*",
            'qualidade': "• Gleiche Premium-Qualität wie das Original",
            'personalizado': "• Vollständig nach Ihren Wünschen personalisiert", 
            'tamanho': "• Perfekte Größe (4.5\" | 11.5cm) für den Schreibtisch",
            'acabamento': "• Premium-Finish",
            'valor_original': "• *Originalpreis:*",
            'desconto': "• *RABATT*",
            'oferta_exclusiva': "🔥 *Exklusives Angebot:*",
            'klarna': "💳 *Klarna:* 3 Raten à",
            'sem_juros': "OHNE ZINSEN",
            'klarna_indisponivel': "💳 *Flexible Zahlungsoptionen verfügbar*",
            'pega_unica': "*Ein einzigartiges Stück zu einem unwiderstehlichen Preis!* ✨",
            'unica_oportunidade': "*Die einzige Gelegenheit zum Erwerb ist hier!* 👇",
            'botao_sim': "✅ Ja, Ich Möchte Es Erwerben!",
            'botao_nao': "❌ Letzte Gelegenheit Ablehnen."
        },
        'frances': {
            'titulo': "🎉 *Nous avons une option SPECTACULAIRE pour vous !*",
            'subtitulo': "🌟 *DESSIN ANIMÉ 3D - ÉDITION COLLECTOR*",
            'qualidade': "• Même qualité premium que l'original",
            'personalizado': "• Entièrement personnalisé comme vous l'avez demandé", 
            'tamanho': "• Taille parfaite (4.5\" | 11.5cm) pour le bureau",
            'acabamento': "• Finition Premium",
            'valor_original': "• *Prix Original:*",
            'desconto': "• *RÉDUCTION*",
            'oferta_exclusiva': "🔥 *Offre Exclusive:*",
            'klarna': "💳 *Klarna:* 3 versements de",
            'sem_juros': "SANS INTÉRÊTS",
            'klarna_indisponivel': "💳 *Options de paiement flexibles disponibles*",
            'pega_unica': "*Une pièce unique à un prix irrésistible !* ✨",
            'unica_oportunidade': "*La seule opportunité d'acquisition est ici !* 👇",
            'botao_sim': "✅ Oui, Je Veux l'Acquérir !",
            'botao_nao': "❌ Refuser Dernière Opportunité."
        }
    }
    
    textos = textos_oferta.get(idioma, textos_oferta['portugues'])

    # 🔥 CONSTRUIR TEXTO BASE (COMUM PARA TODOS)
    texto = (
        f"{textos['titulo']}\n\n"
        
        f"{textos['subtitulo']}\n"
        f"{textos['qualidade']}\n"
        f"{textos['personalizado']}\n" 
        f"{textos['tamanho']}\n"
        f"{textos['acabamento']}\n"
        
        f"{textos['valor_original']} {simbolo_moeda}{total_original:.2f}❌\n"
        f"{textos['desconto']} {simbolo_moeda}{economia:.2f} ({percentual_desconto}% OFF!)\n\n"

        f"{textos['oferta_exclusiva']} {simbolo_moeda}{total_45:.2f}✅\n"
    )
    
    # 🔥 ADICIONAR INFORMAÇÃO DO KLARNA APENAS SE DISPONÍVEL
    if tem_klarna:
        texto += f"{textos['klarna']} {simbolo_moeda}{valor_3x:.2f} {textos['sem_juros']}\n\n"
    else:
        # Mostrar mensagem alternativa para países sem Klarna
        texto += f"{textos['klarna_indisponivel']}\n\n"
    
    # 🔥 ADICIONAR TEXTO FINAL
    texto += (
        f"{textos['pega_unica']}\n\n"
        f"{textos['unica_oportunidade']}"
    )
    
    # 🔥 ATUALIZAR PEDIDO SEM SOBRESCREVER TAMANHO ORIGINAL
    pedido["tamanho_cartoon"] = "4.5\" | 11.5cm (Oferta Especial)"  # ✅ Atualiza apenas o tamanho atual
    pedido["total_original"] = total_original
    pedido["total"] = total_45
    pedido["tipo_oferta"] = "tamanho_45"
    
    botoes = [
        [InlineKeyboardButton(textos['botao_sim'], callback_data=f"pagar_tamanho45_{pedido['id']}")],
        [InlineKeyboardButton(textos['botao_nao'], callback_data=f"sair_poferta45_{pedido['id']}")]
    ]
    
    await context.bot.edit_message_text(
        chat_id=chat_id,
        message_id=message_id,
        text=texto,
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(botoes)
    )
    
    print(f"✅ Oferta tamanho 4.5cm mostrada | Idioma: {idioma} | Klarna: {tem_klarna}")



#mostrarofertatamanho45








async def mostrar_oferta_portachaves(context, pedido, chat_id, message_id):
    """Mostra oferta final do porta-chaves - COM ATUALIZAÇÃO COMPLETA DOS VALORES E TRADUÇÃO"""
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    # 🔥 PROTEGER TAMANHO ORIGINAL IMEDIATAMENTE
    tamanho_original_protegido = pedido.get('tamanho_original', pedido.get('tamanho_cartoon', ''))
    print(f"🔒 TAMANHO ORIGINAL PROTEGIDO (porta-chaves): {tamanho_original_protegido} | Idioma: {idioma}")
    
    # 🔥 🔥 🔥 OBTER TAMANHO CORRETO DO PORTA-CHAVES
    tamanho_portachaves = pedido.get('tamanho_portachaves', '2.5" | 6.4cm')
    print(f"🔍 TAMANHO PORTA-CHAVES DEFINIDO: {tamanho_portachaves}")
    print(f"🔍 VAI DIRETO PORTA-CHAVES: {pedido.get('vai_direto_portachaves', False)}")
    
    # 🔥 IDENTIFICAR MOEDA CORRETA BASEADA NO PAÍS - COM CANADÁ
    def determinar_moeda_pais(pais):
        pais_lower = pais.lower()
        if pais_lower == "estados unidos":
            return "$", "USD"
        elif pais_lower == "brasil":
            return "R$", "BRL"
        elif pais_lower == "reino unido":
            return "£", "GBP"
        elif pais_lower in ["canada", "canadá"]:  # 🔥 NOVO: CANADÁ
            return "C$", "CAD"
        else:
            return "€", "EUR"  # Default para Europa
    
    moeda, codigo_moeda = determinar_moeda_pais(pedido["pais"])
    
    print(f"🔍🔍🔍 INICIANDO OFERTA PORTA-CHAVES PARA: #{pedido['id']} | Idioma: {idioma}")
    print(f"   - País: {pedido['pais']}")
    print(f"   - Moeda identificada: {moeda}")
    print(f"   - Código moeda: {codigo_moeda}")
    print(f"   - Valor atual no pedido: {moeda}{pedido['total']:.2f}")
    print(f"   - Tamanho Original Protegido: {tamanho_original_protegido}")
    print(f"   - Tamanho Porta-chaves: {tamanho_portachaves}")
    
    # 🔥 CALCULAR OFERTA DO PORTA-CHAVES
    oferta = calcular_oferta_portachaves(pedido)
    
    print(f"🔍🔍🔍 OFERTA CALCULADA:")
    print(f"   - Total oferta: {moeda}{oferta['total']:.2f}")
    print(f"   - Economia: {moeda}{oferta['economia']:.2f}")
    
    # 🔥🔥🔥 ATUALIZAR TODOS OS VALORES DO PEDIDO COM OS VALORES DA OFERTA
    if "total_original_real" not in pedido:
        pedido["total_original_real"] = pedido["total"]

    # 🔥 GUARDAR VALOR REAL DA OFERTA PORTA-CHAVES
    pedido["valor_oferta_portachaves_real"] = oferta["total"]
    
    print(f"💰 VALORES REAIS GUARDADOS:")
    print(f"   • Valor Original Real: {moeda}{pedido['valor_original_real']:.2f}")
    print(f"   • Valor Oferta Porta-chaves Real: {moeda}{pedido['valor_oferta_portachaves_real']:.2f}")
    print(f"   • Economia Real: {moeda}{pedido['valor_original_real'] - pedido['valor_oferta_portachaves_real']:.2f}")
    
    # Guardar originais
    pedido["subtotal_original"] = pedido["subtotal"]
    pedido["frete_original"] = pedido["frete"]  
    pedido["imposto_original"] = pedido["imposto"]
    pedido["total_original"] = pedido["total"]
    
    # 🔥 ATUALIZAR COM VALORES DA OFERTA
    pedido["subtotal"] = oferta["subtotal"]
    pedido["frete"] = oferta["frete"]
    pedido["imposto"] = oferta["imposto"]
    pedido["total"] = oferta["total"]  # 🔥 TOTAL DA OFERTA
    
    # 🔥 ATUALIZAR INFORMAÇÕES DO PRODUTO E MOEDA (SEM SOBRESCREVER TAMANHO ORIGINAL)
    pedido["tipo_original"] = pedido["tipo_cartoon"]
    pedido["tipo_cartoon"] = "Porta-chaves"
    
    # 🔥🔥🔥 CORREÇÃO: USAR TAMANHO DINÂMICO DO PORTA-CHAVES
    pedido["tamanho_cartoon"] = tamanho_portachaves  # ✅ Usar tamanho dinâmico (1.5" ou 2.5")
    pedido["tamanho_original"] = tamanho_original_protegido  # 🔥 MANTÉM o tamanho original protegido
    
    pedido["tipo_oferta"] = "portachaves"
    pedido["nome_oferta"] = oferta["nome"]
    pedido["economia"] = oferta["economia"]
    pedido["valor_original"] = oferta["valor_original"]
    
    # 🔥🔥🔥 FORÇAR MOEDA CORRETA NO PEDIDO
    pedido["moeda"] = moeda
    pedido["codigo_moeda"] = codigo_moeda
    pedido["simbolo_moeda"] = moeda
    
    print(f"🔍🔍🔍 PEDIDO ATUALIZADO:")
    print(f"   - Total anterior: {moeda}{pedido['total_original']:.2f}")
    print(f"   - Total oferta: {moeda}{pedido['total']:.2f}")
    print(f"   - Economia: {moeda}{pedido['economia']:.2f}")
    print(f"   - Moeda final: {pedido['moeda']} {pedido['codigo_moeda']}")
    print(f"   - Tamanho Original: {pedido['tamanho_original']}")
    print(f"   - Tamanho Cartoon: {pedido['tamanho_cartoon']}")  # 🔥 AGORA SERÁ 1.5" OU 2.5"
    
    # 🔥 FORMATAR VALORES PARA EXIBIÇÃO
    total_formatado = f"{moeda}{oferta['total']:.2f}"
    economia_formatado = f"{moeda}{oferta['economia']:.2f}"
    valor_original_formatado = f"{moeda}{oferta['valor_original']:.2f}"
    
    # 🔥 TEXTOS DA OFERTA PORTA-CHAVES POR IDIOMA
    textos_oferta = {
        'portugues': {
            'titulo': "🎁 *OFERTA ESPECIAL: PORTA-CHAVES COM 70% DE DESCONTO!* 🎁",
            'subtitulo': "*🔑 PORTA-CHAVES PREMIUM*",
            'miniatura': "• Seu cartoon em miniatura de luxo",
            'memorias': "• Leve suas memórias para todo lado",
            'presente': "• Presente único e personalizado",
            'acabamento': "• Acabamento premium resistente",
            'tamanho': "• Tamanho:",
            'comparativo': "💰 *COMPARATIVO DE VALORES:*",
            'original': " *Original*",
            'com_desconto': "✅ *Com 70% OFF*",
            'mensagem': "*A maneira mais acessível de ter o seu Cartoon 3D sempre consigo!* 🌟",
            'pergunta': "*Vai aproveitar esta oferta exclusiva?* 👇",
            'botao_sim': "✅ SIM, QUERO 70% OFF!",
            'botao_nao': "❌ Recusar Oferta Especial"
        },
        'ingles': {
            'titulo': "🎁 *SPECIAL OFFER: KEYCHAIN WITH 70% DISCOUNT!* 🎁",
            'subtitulo': "*🔑 PREMIUM KEYCHAIN*",
            'miniatura': "• Your cartoon in luxury miniature",
            'memorias': "• Take your memories everywhere",
            'presente': "• Unique and personalized gift",
            'acabamento': "• Resistant premium finish",
            'tamanho': "• Size:",
            'comparativo': "💰 *PRICE COMPARISON:*",
            'original': " *Original*",
            'com_desconto': "✅ *With 70% OFF*",
            'mensagem': "*The most affordable way to have your 3D Cartoon with you always!* 🌟",
            'pergunta': "*Will you take advantage of this exclusive offer?* 👇",
            'botao_sim': "✅ YES, I WANT 70% OFF!",
            'botao_nao': "❌ Reject Special Offer"
        },
        'espanhol': {
            'titulo': "🎁 *¡OFERTA ESPECIAL: LLAVERO CON 70% DE DESCUENTO!* 🎁",
            'subtitulo': "*🔑 LLAVERO PREMIUM*",
            'miniatura': "• Tu caricatura en miniatura de lujo",
            'memorias': "• Lleva tus recuerdos a todas partes",
            'presente': "• Regalo único y personalizado",
            'acabamento': "• Acabado premium resistente",
            'tamanho': "• Tamaño:",
            'comparativo': "💰 *COMPARATIVO DE PRECIOS:*",
            'original': " *Original*",
            'com_desconto': "✅ *Con 70% DESCUENTO*",
            'mensagem': "*¡La forma más asequible de tener tu Caricatura 3D siempre contigo!* 🌟",
            'pergunta': "*¿Vas a aprovechar esta oferta exclusiva?* 👇",
            'botao_sim': "✅ ¡SÍ, QUIERO 70% DESCUENTO!",
            'botao_nao': "❌ Rechazar Oferta Especial"
        },
        'italiano': {
            'titulo': "🎁 *OFFERTA SPECIALE: PORTA-CHIAVI CON 70% DI SCONTO!* 🎁",
            'subtitulo': "*🔑 PORTA-CHIAVI PREMIUM*",
            'miniatura': "• Il tuo cartoon in miniatura di lusso",
            'memorias': "• Porta i tuoi ricordi ovunque",
            'presente': "• Regalo unico e personalizzato",
            'acabamento': "• Finitura premium resistente",
            'tamanho': "• Dimensione:",
            'comparativo': "💰 *CONFRONTO PREZZI:*",
            'original': " *Originale*",
            'com_desconto': "✅ *Con 70% DI SCONTO*",
            'mensagem': "*Il modo più accessibile per avere il tuo Cartoon 3D sempre con te!* 🌟",
            'pergunta': "*Approfitterai di questa offerta esclusiva?* 👇",
            'botao_sim': "✅ SÌ, VOGLIO IL 70% DI SCONTO!",
            'botao_nao': "❌ Rifiuta Offerta Speciale"
        },
        'alemao': {
            'titulo': "🎁 *SONDERANGEBOT: SCHLÜSSELANHÄNGER MIT 70% RABATT!* 🎁",
            'subtitulo': "*🔑 PREMIUM-SCHLÜSSELANHÄNGER*",
            'miniatura': "• Ihre Karikatur in Luxus-Miniatur",
            'memorias': "• Nehmen Sie Ihre Erinnerungen überall hin mit",
            'presente': "• Einzigartiges und personalisiertes Geschenk",
            'acabamento': "• Widerstandsfähige Premium-Ausführung",
            'tamanho': "• Größe:",
            'comparativo': "💰 *PREISVERGLEICH:*",
            'original': " *Original*",
            'com_desconto': "✅ *Mit 70% RABATT*",
            'mensagem': "*Der günstigste Weg, Ihren 3D-Karikatur immer bei sich zu haben!* 🌟",
            'pergunta': "*Werden Sie dieses exklusive Angebot nutzen?* 👇",
            'botao_sim': "✅ JA, ICH WILL 70% RABATT!",
            'botao_nao': "❌ Sonderangebot Ablehnen"
        },
        'frances': {
            'titulo': "🎁 *OFFRE SPÉCIALE : PORTE-CLÉS AVEC 70% DE RÉDUCTION !* 🎁",
            'subtitulo': "*🔑 PORTE-CLÉS PREMIUM*",
            'miniatura': "• Votre dessin animé en miniature de luxe",
            'memorias': "• Emportez vos souvenirs partout",
            'presente': "• Cadeau unique et personnalisé",
            'acabamento': "• Finition premium résistante",
            'tamanho': "• Taille:",
            'comparativo': "💰 *COMPARAISON DES PRIX:*",
            'original': " *Original*",
            'com_desconto': "✅ *Avec 70% DE RÉDUCTION*",
            'mensagem': "*La façon la plus abordable d'avoir toujours votre Dessin Animé 3D avec vous !* 🌟",
            'pergunta': "*Allez-vous profiter de cette offre exclusive ?* 👇",
            'botao_sim': "✅ OUI, JE VEUX 70% DE RÉDUCTION !",
            'botao_nao': "❌ Refuser Offre Spéciale"
        }
    }
    
    textos = textos_oferta.get(idioma, textos_oferta['portugues'])
    
    # 🔥 🔥 🔥 TEXTO DINÂMICO COM TAMANHO CORRETO
    texto = (
        f"{textos['titulo']}\n\n"
        
        f"{textos['subtitulo']}\n"
        f"{textos['miniatura']}\n"
        f"{textos['memorias']}\n"
        f"{textos['presente']}\n"
        f"{textos['acabamento']}\n"
        f"{textos['tamanho']} {tamanho_portachaves}\n\n"  # 🔥 TAMANHO DINÂMICO
        
        f"{textos['comparativo']}\n"
        f"💰{valor_original_formatado}{textos['original']}❌\n\n"
        f"🎯 *{total_formatado}* {textos['com_desconto']}\n\n"
        
        f"{textos['mensagem']}\n\n"
        f"{textos['pergunta']}"
    )
    
    botoes = [
        [InlineKeyboardButton(textos['botao_sim'], callback_data=f"pagar_portachaves_{pedido['id']}")],
        [InlineKeyboardButton(textos['botao_nao'], callback_data=f"recusar_oferta_{pedido['id']}")]
    ]
    
    try:
        await context.bot.edit_message_text(
            chat_id=chat_id,
            message_id=message_id,
            text=texto,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(botoes)
        )
        print(f"✅ Oferta de porta-chaves {tamanho_portachaves} com 70% OFF exibida com sucesso! | Idioma: {idioma}")
        
    except Exception as e:
        print(f"❌ Erro ao exibir oferta de porta-chaves: {e}")
        # Tentar enviar nova mensagem em caso de erro
        await context.bot.send_message(
            chat_id=chat_id,
            text=texto,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(botoes)
        )





async def pagar_tamanho45(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para pagar oferta do tamanho 4.5cm - ATUALIZADO COM MESMA ESTRUTURA E TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    print(f"🎯 PAGAR_TAMANHO45 CHAMADO - VERSÃO ATUALIZADA")
    
    # Extrair pedido_id do callback_data
    pedido_id = query.data.replace("pagar_tamanho45_", "")
    print(f"🔍 Procurando pedido: {pedido_id}")
    
    if pedido_id not in PEDIDOS_REGISTO:
        print(f"❌ Pedido não encontrado no registro: {pedido_id}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro_pedido = {
            'portugues': "❌ Pedido não encontrado. Por favor, inicie um novo pedido.",
            'ingles': "❌ Order not found. Please start a new order.",
            'espanhol': "❌ Pedido no encontrado. Por favor, inicie un nuevo pedido.",
            'italiano': "❌ Ordine non trovato. Per favore, inizi un nuovo ordine.",
            'alemao': "❌ Bestellung nicht gefunden. Bitte beginnen Sie eine neue Bestellung.",
            'frances': "❌ Commande introuvable. Veuillez démarrer une nouvelle commande."
        }
        
        # 🔥 PEGAR IDIOMA DO CONTEXTO
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro_pedido.get(idioma, textos_erro_pedido['portugues']))
        return
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    chat_id = query.message.chat_id
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    # 🔥 CANCELAR QUALQUER TEMPORIZADOR ATIVO
    await cancelar_temporizadores_pedido(pedido_id)
    
    # ATUALIZAR ESTATÍSTICAS
    atualizar_estatistica("ofertas_aceites")
    
    print(f"✅ Pedido encontrado: #{pedido_id}")
    print(f"🔍 Chat ID do cliente: {chat_id}")
    print(f"🌐 Idioma do pedido: {idioma}")

    try:
        # 🔥 PASSO 1: DEFINIR MÉTODOS DE PAGAMENTO POR PAÍS - MESMA ESTRUTURA DO pagar_stripe
        def get_payment_methods(pais):
            """Retorna métodos de pagamento baseado no país"""
            
            def get_country_code(pais_nome):
                mapeamento_paises = {
                    "portugal": "PT",
                    "espanha": "ES", 
                    "franca": "FR",
                    "alemanha": "DE",
                    "belgica": "BE",
                    "reino unido": "GB",
                    "estados unidos": "US",
                    "paises baixos": "NL",
                    "brasil": "BR",
                    "irlanda": "IE",
                    "italia": "IT",
                    "luxemburgo": "LU",
                    "canada": "CA"
                }
                return mapeamento_paises.get(pais_nome.lower(), pais_nome.upper())
            
            country_code = get_country_code(pais)
            print(f"🔍 País recebido: '{pais}' → Código: '{country_code}'")
            
            # 🔥 MESMO payment_methods_by_country DO pagar_stripe
            payment_methods_by_country = {
                "PT": ["card", "paypal", "link", "klarna", "mb_way", "sepa_debit"],
                "ES": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "FR": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "DE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "BE": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "GB": ["card", "paypal", "link", "klarna"],
                "US": ["card", "paypal", "link"],
                "NL": ["card", "paypal", "link", "klarna", "ideal", "sepa_debit"],
                "BR": ["card", "link"],
                "IE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "IT": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "LU": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "CA": ["card", "paypal", "link"]
            }
            
            methods = payment_methods_by_country.get(country_code, ["card", "link"])
            print(f"💳 Métodos de pagamento para {pais} ({country_code}): {methods}")
            return methods

        # 🔥 OBTER MÉTODOS REAIS PARA ESTE PAÍS
        metodos_reais = get_payment_methods(pedido['pais'])
        
        # 🔥 CRIAR TEXTO DINÂMICO DOS MÉTODOS COM TRADUÇÃO
        def formatar_metodos(metodos, pais, idioma):
            """Formata os métodos de pagamento para exibição"""
            # 🔥 NOMES DOS MÉTODOS POR IDIOMA
            nomes_metodos_por_idioma = {
                'portugues': {
                    "card": "Cartão",
                    "paypal": "PayPal", 
                    "link": "Link (inclui Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'ingles': {
                    "card": "Card",
                    "paypal": "PayPal", 
                    "link": "Link (includes Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA Debit",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'espanhol': {
                    "card": "Tarjeta",
                    "paypal": "PayPal", 
                    "link": "Link (incluye Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'italiano': {
                    "card": "Carta",
                    "paypal": "PayPal", 
                    "link": "Link (include Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Addebito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'alemao': {
                    "card": "Karte",
                    "paypal": "PayPal", 
                    "link": "Link (inkl. Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA-Lastschrift",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'frances': {
                    "card": "Carte",
                    "paypal": "PayPal", 
                    "link": "Link (inclut Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Prélèvement SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                }
            }
            
            nomes_metodos = nomes_metodos_por_idioma.get(idioma, nomes_metodos_por_idioma['portugues'])
            textos = []
            
            for metodo in metodos:
                if metodo in nomes_metodos:
                    textos.append(nomes_metodos[metodo])
                else:
                    textos.append(metodo.capitalize())
            
            return ", ".join(textos)

        # 🔥 PASSO 2: VERIFICAR CONFIGURAÇÃO DE WALLETS - MESMA LÓGICA
        def verificar_config_wallets():
            """Verifica se as wallets estão configuradas corretamente"""
            try:
                apple_domains = stripe.ApplePayDomain.list()
                print("🍎 Domínios Apple Pay configurados:")
                for domain in apple_domains.data:
                    print(f"   - {domain.domain}")
                
                # Verificar domínio atual
                seu_dominio = "unceased-bibliothecal-donette.ngrok-free.dev"
                dominios_apple = [d.domain for d in apple_domains.data]
                if seu_dominio in dominios_apple:
                    print("✅ Domínio ngrok configurado no Apple Pay!")
                    return True
                else:
                    print("⚠️ Domínio ngrok NÃO configurado no Apple Pay")
                    return False
                    
            except Exception as e:
                print(f"❌ Erro ao verificar wallets: {e}")
                return False

        wallets_configuradas = verificar_config_wallets()

        # 🔥 TEXTOS TRADUZIDOS PARA O CHECKOUT (shipping_message e submit_message)
        textos_checkout_messages = {
            'portugues': {
                "shipping_message": "📦 Enviaremos o seu Cartoon personalizado para este endereço!",
                "submit_message": "✨ Obrigado! Vamos criar um Cartoon incrível para si!"
            },
            'ingles': {
                "shipping_message": "📦 We'll send your personalized Cartoon to this address!",
                "submit_message": "✨ Thank you! We'll create an amazing Cartoon for you!"
            },
            'espanhol': {
                "shipping_message": "📦 ¡Enviaremos tu Cartoon personalizado a esta dirección!",
                "submit_message": "✨ ¡Gracias! ¡Crearemos un Cartoon increíble para ti!"
            },
            'italiano': {
                "shipping_message": "📦 Spediremo il tuo Cartoon personalizzato a questo indirizzo!",
                "submit_message": "✨ Grazie! Creeremo un Cartoon incredibile per te!"
            },
            'alemao': {
                "shipping_message": "📦 Wir senden Ihren personalisierten Cartoon an diese Adresse!",
                "submit_message": "✨ Danke! Wir erstellen einen fantastischen Cartoon für Sie!"
            },
            'frances': {
                "shipping_message": "📦 Nous enverrons votre Cartoon personnalisé à cette adresse !",
                "submit_message": "✨ Merci ! Nous créerons un Cartoon incroyable pour vous !"
            }
        }
        
        textos_messages = textos_checkout_messages.get(idioma, textos_checkout_messages['portugues'])
        
        # 🔥 DESCRIÇÕES DO PRODUTO POR IDIOMA
        descricoes_produto = {
            'portugues': f"Oferta Especial Tamanho 4.5\" | 11.5cm - Pedido #{pedido_id}",
            'ingles': f"Special Size Offer 4.5\" | 11.5cm - Order #{pedido_id}",
            'espanhol': f"Oferta Especial Tamaño 4.5\" | 11.5cm - Pedido #{pedido_id}",
            'italiano': f"Offerta Speciale Dimensione 4.5\" | 11.5cm - Ordine #{pedido_id}",
            'alemao': f"Sonderangebot Größe 4.5\" | 11.5cm - Bestellung #{pedido_id}",
            'frances': f"Offre Spéciale Taille 4.5\" | 11.5cm - Commande #{pedido_id}"
        }
        
        descricao_produto = descricoes_produto.get(idioma, descricoes_produto['portugues'])
        
        # 🔥 NOMES DO PRODUTO POR IDIOMA
        nomes_produto = {
            'portugues': f"Cartoon 4.5\" | 11.5cm - {pedido['tipo_cartoon']}",
            'ingles': f"Cartoon 4.5\" | 11.5cm - {pedido['tipo_cartoon']}",
            'espanhol': f"Cartoon 4.5\" | 11.5cm - {pedido['tipo_cartoon']}",
            'italiano': f"Cartoon 4.5\" | 11.5cm - {pedido['tipo_cartoon']}",
            'alemao': f"Cartoon 4.5\" | 11.5cm - {pedido['tipo_cartoon']}",
            'frances': f"Dessin Animé 4.5\" | 11.5cm - {pedido['tipo_cartoon']}"
        }
        
        nome_produto = nomes_produto.get(idioma, nomes_produto['portugues'])
        
        # 🔥 PASSO 3: CRIAR SESSÃO STRIPE - MESMA ESTRUTURA
        print("🔗 Criando Checkout Session para oferta tamanho 4.5cm...")
        
        session_config = {
            "payment_method_types": metodos_reais,
            "mode": "payment",
            "customer_email": pedido["email"],
            
            # 🔥 CONFIGURAÇÃO PARA WALLETS
            "payment_method_options": {
                "card": {
                    "request_three_d_secure": "automatic"
                }
            },
            
            "shipping_address_collection": {
                "allowed_countries": [
                    "PT", "ES", "FR", "DE", "BE", "GB", "US", "NL", "BR", "IE", "IT", "LU", "CA"
                ]
            },
            
            # 🔥 MENSAGENS TRADUZIDAS PARA O CHECKOUT
            "custom_text": {
                "shipping_address": {
                    "message": textos_messages["shipping_message"]
                },
                "submit": {
                    "message": textos_messages["submit_message"]
                }
            },
            
            "line_items": [{
                "price_data": {
                    "currency": pedido["moeda"].lower(),
                    "product_data": {
                        "name": nome_produto,  # 🔥 NOME TRADUZIDO
                        "description": descricao_produto,  # 🔥 DESCRIÇÃO TRADUZIDA
                    },
                    "unit_amount": int(pedido["total"] * 100),
                },
                "quantity": 1
            }],
            
            # 🔥 URLs CORRETAS
            "success_url": f"https://t.me/plan3d_bot?start=payment_success_{pedido_id}",
            "cancel_url": f"https://t.me/plan3d_bot?start=payment_cancelled_{pedido_id}",
            
            "metadata": {
                "pedido_id": pedido_id,
                "chat_id": str(chat_id),
                "pais": pedido['pais'],
                "moeda": pedido["moeda"],
                "total_pago": str(pedido["total"]),
                "nome_cliente": pedido['nome'],
                "tipo_cartoon": pedido['tipo_cartoon'],
                "tamanho_cartoon": pedido['tamanho_cartoon'],
                "tipo_sessao": "oferta_tamanho_45",
                "tipo_oferta": "tamanho_45",
                "valor_original": str(pedido.get('total_original', pedido['total'])),
                "wallets_habilitadas": str(wallets_configuradas),
                "idioma": idioma  # 🔥 ADICIONAR IDIOMA AO METADATA
            },
            
            "expires_at": int((datetime.now() + timedelta(minutes=30)).timestamp()),
        }

        # 🔥 CONFIGURAÇÃO ESPECÍFICA PARA WALLETS - MESMA LÓGICA
        paises_com_wallets = ["Reino Unido", "Estados Unidos", "Brasil", "Irlanda", 
                            "França", "Alemanha", "Itália", "Espanha", "Portugal", 
                            "Países Baixos", "Bélgica", "Luxemburgo", "Canadá"]
        
        if pedido['pais'] in paises_com_wallets and "link" in metodos_reais:
            print(f"📱 Configurando Apple Pay/Google Pay para {pedido['pais']}")
            session_config["payment_method_options"]["link"] = {"persistent_token": None}

        # 🔥 CRIAR A SESSÃO
        session = stripe.checkout.Session.create(**session_config)

        print(f"✅ CHECKOUT SESSION CRIADA: {session.id}")
        print(f"🔗 URL do Checkout: {session.url}")

        # 🔥 PASSO 4: ATUALIZAR PEDIDO
        pedido["session_id_oferta"] = session.id
        pedido["payment_intent_id"] = session.payment_intent
        pedido["wallets_configuradas"] = wallets_configuradas
        pedido["data_oferta"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pedido["oferta_ativa"] = True
        
        print(f"📊 Pedido atualizado para oferta 4.5cm")

        # 🔥 PASSO 5: MENSAGEM FINAL COM INSTRUÇÕES CLARAS - COM TRADUÇÃO
        texto_metodos = formatar_metodos(metodos_reais, pedido['pais'], idioma)
        
        status_wallets = "✅ **CONFIGURADO**" if wallets_configuradas else "⚠️ **EM CONFIGURAÇÃO**"
        
        # Calcular economia para mostrar na mensagem
        economia = pedido.get('total_original', pedido['total']) - pedido['total']
        
        # 🔥 TEXTOS DA OFERTA CONFIRMADA POR IDIOMA
        textos_oferta_confirmada = {
            'portugues': {
                'titulo': "🎉 *OFERTA ESPECIAL CONFIRMADA!* 🎉",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envio:*",
                'moeda': "💰 *Moeda:*",
                'detalhes': "✨ *Detalhes da Oferta:*",
                'tipo': "• 🎨",
                'tamanho': "• 📏 Tamanho: 4.5\" | 11.5cm (Oferta Especial)",
                'economia': "• 💰 Economia:",
                'total_pagar': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido:",
                'checkout_pedido': "📋 *No checkout será pedido:*",
                'endereco': "1️⃣ **Endereço de entrega completo**",
                'metodo': "2️⃣ **Método de pagamento**",
                'metodos_disponiveis': "💳 *Métodos disponíveis:*",
                'seguro': "🔒 *Pagamento 100% seguro via Stripe*",
                'tempo': "⏰ *Tem 10 minutos para efetuar o pagamento*",
                'clique_abaixo': "Clique abaixo para pagar: 👇",
                'botao': "💳 PAGAR OFERTA →"
            },
            'ingles': {
                'titulo': "🎉 *SPECIAL OFFER CONFIRMED!* 🎉",
                'cliente': "👤 *Customer:*",
                'pais': "🌍 *Shipping Country:*",
                'moeda': "💰 *Currency:*",
                'detalhes': "✨ *Offer Details:*",
                'tipo': "• 🎨",
                'tamanho': "• 📏 Size: 4.5\" | 11.5cm (Special Offer)",
                'economia': "• 💰 Savings:",
                'total_pagar': "💳 **TOTAL TO PAY:",
                'pedido': "🆔 **Order:",
                'checkout_pedido': "📋 *In checkout you will be asked for:*",
                'endereco': "1️⃣ **Complete shipping address**",
                'metodo': "2️⃣ **Payment method**",
                'metodos_disponiveis': "💳 *Available methods:*",
                'seguro': "🔒 *100% secure payment via Stripe*",
                'tempo': "⏰ *You have 10 minutes to complete payment*",
                'clique_abaixo': "Click below to pay: 👇",
                'botao': "💳 PAY OFFER →"
            },
            'espanhol': {
                'titulo': "🎉 *¡OFERTA ESPECIAL CONFIRMADA!* 🎉",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envío:*",
                'moeda': "💰 *Moneda:*",
                'detalhes': "✨ *Detalles de la Oferta:*",
                'tipo': "• 🎨",
                'tamanho': "• 📏 Tamaño: 4.5\" | 11.5cm (Oferta Especial)",
                'economia': "• 💰 Ahorro:",
                'total_pagar': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido:",
                'checkout_pedido': "📋 *En el checkout se pedirá:*",
                'endereco': "1️⃣ **Dirección de envío completa**",
                'metodo': "2️⃣ **Método de pago**",
                'metodos_disponiveis': "💳 *Métodos disponibles:*",
                'seguro': "🔒 *Pago 100% seguro vía Stripe*",
                'tempo': "⏰ *Tienes 10 minutos para efectuar el pago*",
                'clique_abaixo': "Haz clic abajo para pagar: 👇",
                'botao': "💳 PAGAR OFERTA →"
            },
            'italiano': {
                'titulo': "🎉 *OFFERTA SPECIALE CONFERMATA!* 🎉",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *Paese di Spedizione:*",
                'moeda': "💰 *Valuta:*",
                'detalhes': "✨ *Dettagli dell'Offerta:*",
                'tipo': "• 🎨",
                'tamanho': "• 📏 Dimensione: 4.5\" | 11.5cm (Offerta Speciale)",
                'economia': "• 💰 Risparmio:",
                'total_pagar': "💳 **TOTALE DA PAGARE:",
                'pedido': "🆔 **Ordine:",
                'checkout_pedido': "📋 *Nel checkout verrà richiesto:*",
                'endereco': "1️⃣ **Indirizzo di spedizione completo**",
                'metodo': "2️⃣ **Metodo di pagamento**",
                'metodos_disponiveis': "💳 *Metodi disponibili:*",
                'seguro': "🔒 *Pagamento 100% sicuro tramite Stripe*",
                'tempo': "⏰ *Hai 10 minuti per effettuare il pagamento*",
                'clique_abaixo': "Clicca qui sotto per pagare: 👇",
                'botao': "💳 PAGA OFFERTA →"
            },
            'alemao': {
                'titulo': "🎉 *SONDERANGEBOT BESTÄTIGT!* 🎉",
                'cliente': "👤 *Kunde:*",
                'pais': "🌍 *Versandland:*",
                'moeda': "💰 *Währung:*",
                'detalhes': "✨ *Angebotsdetails:*",
                'tipo': "• 🎨",
                'tamanho': "• 📏 Größe: 4.5\" | 11.5cm (Sonderangebot)",
                'economia': "• 💰 Ersparnis:",
                'total_pagar': "💳 **GESAMTBETRAG ZU ZAHLEN:",
                'pedido': "🆔 **Bestellung:",
                'checkout_pedido': "📋 *Im Checkout wird angefordert:*",
                'endereco': "1️⃣ **Vollständige Lieferadresse**",
                'metodo': "2️⃣ **Zahlungsmethode**",
                'metodos_disponiveis': "💳 *Verfügbare Methoden:*",
                'seguro': "🔒 *100% sichere Zahlung über Stripe*",
                'tempo': "⏰ *Sie haben 10 Minuten für die Zahlung*",
                'clique_abaixo': "Klicken Sie unten zum Bezahlen: 👇",
                'botao': "💳 ANGEBOT BEZAHLEN →"
            },
            'frances': {
                'titulo': "🎉 *OFFRE SPÉCIALE CONFIRMÉE !* 🎉",
                'cliente': "👤 *Client:*",
                'pais': "🌍 *Pays de Livraison:*",
                'moeda': "💰 *Devise:*",
                'detalhes': "✨ *Détails de l'Offre:*",
                'tipo': "• 🎨",
                'tamanho': "• 📏 Taille: 4.5\" | 11.5cm (Offre Spéciale)",
                'economia': "• 💰 Économie:",
                'total_pagar': "💳 **TOTAL À PAYER:",
                'pedido': "🆔 **Commande:",
                'checkout_pedido': "📋 *Dans le checkout, il sera demandé:*",
                'endereco': "1️⃣ **Adresse de livraison complète**",
                'metodo': "2️⃣ **Méthode de paiement**",
                'metodos_disponiveis': "💳 *Méthodes disponibles:*",
                'seguro': "🔒 *Paiement 100% sécurisé via Stripe*",
                'tempo': "⏰ *Vous avez 10 minutes pour effectuer le paiement*",
                'clique_abaixo': "Cliquez ci-dessous pour payer : 👇",
                'botao': "💳 PAYER L'OFFRE →"
            }
        }
        
        textos = textos_oferta_confirmada.get(idioma, textos_oferta_confirmada['portugues'])

        await query.edit_message_text(
            text=(
                f"{textos['titulo']}\n\n"
                f"{textos['cliente']} {pedido['nome']}\n"
                f"{textos['pais']} {pedido['pais']}\n"
                f"{textos['moeda']} {pedido['moeda'].upper()} {pedido['simbolo_moeda']}\n\n"
                
                f"{textos['detalhes']}\n"
                f"{textos['tipo']} {pedido['tipo_cartoon']}\n"
                f"{textos['tamanho']}\n"
                f"{textos['economia']} {pedido['simbolo_moeda']}{economia:.2f}\n\n"
                
                f"{textos['total_pagar']} {pedido['simbolo_moeda']}{pedido['total']:.2f}**\n"
                f"{textos['pedido']} #{pedido_id}**\n\n"
                
                f"{textos['checkout_pedido']}\n"
                f"{textos['endereco']}\n"
                f"{textos['metodo']}\n\n"
                f"{textos['metodos_disponiveis']} {texto_metodos}\n"
                f"{textos['seguro']}\n\n"
                f"{textos['tempo']}\n\n"
                f"{textos['clique_abaixo']}"
            ),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos['botao'], url=session.url)]
            ])
        )
        
        print(f"✅ Usuário redirecionado para Checkout (Oferta 4.5cm) | Idioma: {idioma}")

        # 🔥 INICIAR TEMPORIZADOR PARA OFERTA
        await iniciar_temporizador_oferta(context, pedido_id, chat_id, query.message.message_id, idioma)
        
    except Exception as e:
        print(f"❌ ERRO STRIPE NA OFERTA 4.5cm: {str(e)}")
        print(f"🔍 Tipo do erro: {type(e)}")
        
        import traceback
        print(f"🔍 Traceback completo: {traceback.format_exc()}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro_pagamento = {
            'portugues': {
                'erro': "❌ Erro no processamento do pagamento.",
                'tentar': "🔄 Tentar Novamente",
                'suporte': "📞 Suporte"
            },
            'ingles': {
                'erro': "❌ Error processing payment.",
                'tentar': "🔄 Try Again",
                'suporte': "📞 Support"
            },
            'espanhol': {
                'erro': "❌ Error en el procesamiento del pago.",
                'tentar': "🔄 Intentar de Nuevo",
                'suporte': "📞 Soporte"
            },
            'italiano': {
                'erro': "❌ Errore nell'elaborazione del pagamento.",
                'tentar': "🔄 Riprova",
                'suporte': "📞 Supporto"
            },
            'alemao': {
                'erro': "❌ Fehler bei der Zahlungsverarbeitung.",
                'tentar': "🔄 Erneut versuchen",
                'suporte': "📞 Support"
            },
            'frances': {
                'erro': "❌ Erreur lors du traitement du paiement.",
                'tentar': "🔄 Réessayer",
                'suporte': "📞 Support"
            }
        }
        
        textos_erro = textos_erro_pagamento.get(idioma, textos_erro_pagamento['portugues'])
        
        await query.edit_message_text(
            f"{textos_erro['erro']}\n"
            "Por favor, tente novamente em alguns segundos.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos_erro['tentar'], callback_data=f"pagar_tamanho45_{pedido_id}")],
                [InlineKeyboardButton(textos_erro['suporte'], callback_data=f"todas_recusadas_{pedido_id}")]
            ])
        )






async def pagar_portachaves(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para pagar oferta do porta-chaves - ATUALIZADO COM MESMA ESTRUTURA E TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    print(f"🎯 PAGAR_PORTACHAVES CHAMADO - VERSÃO ATUALIZADA")
    
    # Extrair pedido_id do callback_data
    pedido_id = query.data.replace("pagar_portachaves_", "")
    print(f"🔍 Procurando pedido: {pedido_id}")
    
    if pedido_id not in PEDIDOS_REGISTO:
        print(f"❌ Pedido não encontrado no registro: {pedido_id}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro_pedido = {
            'portugues': "❌ Pedido não encontrado. Por favor, inicie um novo pedido.",
            'ingles': "❌ Order not found. Please start a new order.",
            'espanhol': "❌ Pedido no encontrado. Por favor, inicie un nuevo pedido.",
            'italiano': "❌ Ordine non trovato. Per favore, inizi un nuovo ordine.",
            'alemao': "❌ Bestellung nicht gefunden. Bitte beginnen Sie eine neue Bestellung.",
            'frances': "❌ Commande introuvable. Veuillez démarrer une nouvelle commande."
        }
        
        # 🔥 PEGAR IDIOMA DO CONTEXTO
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro_pedido.get(idioma, textos_erro_pedido['portugues']))
        return
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    chat_id = query.message.chat_id
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    # 🔥 CANCELAR QUALQUER TEMPORIZADOR ATIVO
    await cancelar_temporizadores_pedido(pedido_id)
    
    # ATUALIZAR ESTATÍSTICAS
    atualizar_estatistica("ofertas_aceites")
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: DEFINIR MOEDA CORRETA PARA STRIPE
    def determinar_moeda_stripe(pais):
        """Retorna código da moeda em minúsculas para Stripe"""
        pais_lower = pais.lower()
        if pais_lower == "estados unidos":
            return "usd", "$"
        elif pais_lower == "brasil":
            return "brl", "R$"
        elif pais_lower == "reino unido":
            return "gbp", "£"
        elif pais_lower in ["canada", "canadá"]:
            return "cad", "C$"
        else:
            return "eur", "€"  # Padrão Europa
    
    # 🔥 USAR FUNÇÃO CORRETA PARA DETERMINAR MOEDA
    codigo_moeda, simbolo_moeda = determinar_moeda_stripe(pedido["pais"])
    
    total_oferta = pedido["total"]
    valor_original = pedido.get('valor_original', pedido.get('total_original_real', pedido.get('total_original', pedido['total'])))
    economia = pedido.get('economia', valor_original - total_oferta)
    
    # 🔥 🔥 🔥 OBTER TAMANHO CORRETO DO PORTA-CHAVES
    tamanho_portachaves = pedido.get('tamanho_portachaves', '2.5" | 6.4cm')
    
    print(f"🔍 VERIFICAÇÃO MOEDA STRIPE | Idioma: {idioma}:")
    print(f"   - País: {pedido['pais']}")
    print(f"   - Código Stripe: {codigo_moeda}")
    print(f"   - Símbolo: {simbolo_moeda}")
    print(f"   - Tamanho Porta-chaves: {tamanho_portachaves}")
    print(f"   - Vai Direto: {pedido.get('vai_direto_portachaves', False)}")
    
    print(f"✅ Pedido encontrado: #{pedido_id}")
    print(f"🔍 Chat ID do cliente: {chat_id}")

    try:
        # 🔥 PASSO 1: DEFINIR MÉTODOS DE PAGAMENTO POR PAÍS
        def get_payment_methods(pais):
            """Retorna métodos de pagamento baseado no país"""
            
            def get_country_code(pais_nome):
                mapeamento_paises = {
                    "portugal": "PT",
                    "espanha": "ES", 
                    "franca": "FR",
                    "alemanha": "DE",
                    "belgica": "BE",
                    "reino unido": "GB",
                    "estados unidos": "US",
                    "paises baixos": "NL",
                    "brasil": "BR",
                    "irlanda": "IE",
                    "italia": "IT",
                    "luxemburgo": "LU",
                    "canada": "CA"
                }
                return mapeamento_paises.get(pais_nome.lower(), pais_nome.upper())
            
            country_code = get_country_code(pais)
            print(f"🔍 País recebido: '{pais}' → Código: '{country_code}'")
            
            payment_methods_by_country = {
                "PT": ["card", "paypal", "link", "klarna", "mb_way", "sepa_debit"],
                "ES": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "FR": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "DE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "BE": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "GB": ["card", "paypal", "link", "klarna"],
                "US": ["card", "paypal", "link"],
                "NL": ["card", "paypal", "link", "klarna", "ideal", "sepa_debit"],
                "BR": ["card", "link"],
                "IE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "IT": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "LU": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "CA": ["card", "paypal", "link"]
            }
            
            methods = payment_methods_by_country.get(country_code, ["card", "link"])
            print(f"💳 Métodos de pagamento para {pais} ({country_code}): {methods}")
            return methods

        # 🔥 OBTER MÉTODOS REAIS PARA ESTE PAÍS
        metodos_reais = get_payment_methods(pedido['pais'])
        
        # 🔥 CRIAR TEXTO DINÂMICO DOS MÉTODOS COM TRADUÇÃO
        def formatar_metodos(metodos, pais, idioma):
            """Formata os métodos de pagamento para exibição"""
            # 🔥 NOMES DOS MÉTODOS POR IDIOMA
            nomes_metodos_por_idioma = {
                'portugues': {
                    "card": "Cartão",
                    "paypal": "PayPal", 
                    "link": "Link (inclui Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'ingles': {
                    "card": "Card",
                    "paypal": "PayPal", 
                    "link": "Link (includes Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA Debit",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'espanhol': {
                    "card": "Tarjeta",
                    "paypal": "PayPal", 
                    "link": "Link (incluye Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'italiano': {
                    "card": "Carta",
                    "paypal": "PayPal", 
                    "link": "Link (include Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Addebito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'alemao': {
                    "card": "Karte",
                    "paypal": "PayPal", 
                    "link": "Link (inkl. Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA-Lastschrift",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'frances': {
                    "card": "Carte",
                    "paypal": "PayPal", 
                    "link": "Link (inclut Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Prélèvement SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                }
            }
            
            nomes_metodos = nomes_metodos_por_idioma.get(idioma, nomes_metodos_por_idioma['portugues'])
            textos = []
            
            for metodo in metodos:
                if metodo in nomes_metodos:
                    textos.append(nomes_metodos[metodo])
                else:
                    textos.append(metodo.capitalize())
            
            return ", ".join(textos)

        # 🔥 PASSO 2: VERIFICAR CONFIGURAÇÃO DE WALLETS
        def verificar_config_wallets():
            """Verifica se as wallets estão configuradas corretamente"""
            try:
                apple_domains = stripe.ApplePayDomain.list()
                print("🍎 Domínios Apple Pay configurados:")
                for domain in apple_domains.data:
                    print(f"   - {domain.domain}")
                
                # Verificar domínio atual
                seu_dominio = "unceased-bibliothecal-donette.ngrok-free.dev"
                dominios_apple = [d.domain for d in apple_domains.data]
                if seu_dominio in dominios_apple:
                    print("✅ Domínio ngrok configurado no Apple Pay!")
                    return True
                else:
                    print("⚠️ Domínio ngrok NÃO configurado no Apple Pay")
                    return False
                    
            except Exception as e:
                print(f"❌ Erro ao verificar wallets: {e}")
                return False

        wallets_configuradas = verificar_config_wallets()

        # 🔥 TEXTOS TRADUZIDOS PARA O CHECKOUT (shipping_message e submit_message)
        textos_checkout_messages = {
            'portugues': {
                "shipping_message": "📦 Enviaremos o seu Porta-Chaves personalizado para este endereço!",
                "submit_message": "✨ Obrigado! Vamos criar um Porta-Chaves incrível para si!"
            },
            'ingles': {
                "shipping_message": "📦 We'll send your personalized Keychain to this address!",
                "submit_message": "✨ Thank you! We'll create an amazing Keychain for you!"
            },
            'espanhol': {
                "shipping_message": "📦 ¡Enviaremos tu Llavero personalizado a esta dirección!",
                "submit_message": "✨ ¡Gracias! ¡Crearemos un Llavero increíble para ti!"
            },
            'italiano': {
                "shipping_message": "📦 Spediremo il tuo Portachiavi personalizzato a questo indirizzo!",
                "submit_message": "✨ Grazie! Creeremo un Portachiavi incredibile per te!"
            },
            'alemao': {
                "shipping_message": "📦 Wir senden Ihren personalisierten Schlüsselanhänger an diese Adresse!",
                "submit_message": "✨ Danke! Wir erstellen einen fantastischen Schlüsselanhänger für Sie!"
            },
            'frances': {
                "shipping_message": "📦 Nous enverrons votre Porte-clés personnalisé à cette adresse !",
                "submit_message": "✨ Merci ! Nous créerons un Porte-clés incroyable pour vous !"
            }
        }
        
        textos_messages = textos_checkout_messages.get(idioma, textos_checkout_messages['portugues'])
        
        # 🔥 DESCRIÇÕES DO PRODUTO POR IDIOMA
        descricoes_produto = {
            'portugues': f"Oferta Especial Porta-chaves {tamanho_portachaves} - Pedido #{pedido_id}",
            'ingles': f"Special Keychain Offer {tamanho_portachaves} - Order #{pedido_id}",
            'espanhol': f"Oferta Especial Llavero {tamanho_portachaves} - Pedido #{pedido_id}",
            'italiano': f"Offerta Speciale Portachiavi {tamanho_portachaves} - Ordine #{pedido_id}",
            'alemao': f"Sonderangebot Schlüsselanhänger {tamanho_portachaves} - Bestellung #{pedido_id}",
            'frances': f"Offre Spéciale Porte-clés {tamanho_portachaves} - Commande #{pedido_id}"
        }
        
        descricao_produto = descricoes_produto.get(idioma, descricoes_produto['portugues'])
        
        # 🔥 NOMES DO PRODUTO POR IDIOMA
        nomes_produto = {
            'portugues': f"Porta-chaves Cartoon {tamanho_portachaves}",
            'ingles': f"Keychain Cartoon {tamanho_portachaves}",
            'espanhol': f"Llavero Cartoon {tamanho_portachaves}",
            'italiano': f"Portachiavi Cartoon {tamanho_portachaves}",
            'alemao': f"Schlüsselanhänger Cartoon {tamanho_portachaves}",
            'frances': f"Porte-clés Cartoon {tamanho_portachaves}"
        }
        
        nome_produto = nomes_produto.get(idioma, nomes_produto['portugues'])
        
        # 🔥 PASSO 3: CRIAR SESSÃO STRIPE
        print("🔗 Criando Checkout Session para oferta porta-chaves...")
        
        session_config = {
            "payment_method_types": metodos_reais,
            "mode": "payment",
            "customer_email": pedido["email"],
            
            # 🔥 CONFIGURAÇÃO PARA WALLETS
            "payment_method_options": {
                "card": {
                    "request_three_d_secure": "automatic"
                }
            },
            
            "shipping_address_collection": {
                "allowed_countries": [
                    "PT", "ES", "FR", "DE", "BE", "GB", "US", "NL", "BR", "IE", "IT", "LU", "CA"
                ]
            },
            
            # 🔥 MENSAGENS TRADUZIDAS PARA O CHECKOUT
            "custom_text": {
                "shipping_address": {
                    "message": textos_messages["shipping_message"]
                },
                "submit": {
                    "message": textos_messages["submit_message"]
                }
            },
            
            "line_items": [{
                "price_data": {
                    "currency": codigo_moeda,  # 🔥 AGORA CORRETO: "eur", "usd", etc.
                    "product_data": {
                        "name": nome_produto,  # 🔥 NOME TRADUZIDO
                        "description": descricao_produto,  # 🔥 DESCRIÇÃO TRADUZIDA
                    },
                    "unit_amount": int(total_oferta * 100),
                },
                "quantity": 1
            }],
            
            # 🔥 URLs CORRETAS
            "success_url": f"https://t.me/plan3d_bot?start=payment_success_{pedido_id}",
            "cancel_url": f"https://t.me/plan3d_bot?start=payment_cancelled_{pedido_id}",
            
            "metadata": {
                "pedido_id": pedido_id,
                "chat_id": str(chat_id),
                "pais": pedido['pais'],
                "moeda": codigo_moeda.upper(),  # Para registro, pode ser maiúsculas
                "total_pago": str(total_oferta),
                "nome_cliente": pedido['nome'],
                "tipo_cartoon": pedido['tipo_cartoon'],
                "tipo_original": pedido.get('tipo_original', pedido['tipo_cartoon']),
                "tamanho_cartoon": tamanho_portachaves,
                "tipo_sessao": "oferta_portachaves",
                "tipo_oferta": "portachaves",
                "valor_original": str(valor_original),
                "economia": str(economia),
                "wallets_habilitadas": str(wallets_configuradas),
                "idioma": idioma  # 🔥 ADICIONAR IDIOMA AO METADATA
            },
            
            "expires_at": int((datetime.now() + timedelta(minutes=30)).timestamp()),
        }

        # 🔥 CONFIGURAÇÃO ESPECÍFICA PARA WALLETS
        paises_com_wallets = ["Reino Unido", "Estados Unidos", "Brasil", "Irlanda", 
                            "França", "Alemanha", "Itália", "Espanha", "Portugal", 
                            "Países Baixos", "Bélgica", "Luxemburgo", "Canadá"]
        
        if pedido['pais'] in paises_com_wallets and "link" in metodos_reais:
            print(f"📱 Configurando Apple Pay/Google Pay para {pedido['pais']}")
            session_config["payment_method_options"]["link"] = {"persistent_token": None}

        # 🔥 CRIAR A SESSÃO
        session = stripe.checkout.Session.create(**session_config)

        print(f"✅ CHECKOUT SESSION CRIADA: {session.id}")
        print(f"🔗 URL do Checkout: {session.url}")

        # 🔥 PASSO 4: ATUALIZAR PEDIDO
        pedido["session_id_oferta"] = session.id
        pedido["payment_intent_id"] = session.payment_intent
        pedido["wallets_configuradas"] = wallets_configuradas
        pedido["data_oferta"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pedido["oferta_ativa"] = True
        
        print(f"📊 Pedido atualizado para oferta porta-chaves")

        # 🔥 PASSO 5: MENSAGEM FINAL COM TRADUÇÃO
        texto_metodos = formatar_metodos(metodos_reais, pedido['pais'], idioma)
        
        # 🔥 TEXTOS DA OFERTA CONFIRMADA POR IDIOMA
        textos_oferta_confirmada = {
            'portugues': {
                'titulo': "🎉 *OFERTA ESPECIAL CONFIRMADA!* 🎉",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envio:*",
                'moeda': "💰 *Moeda:*",
                'detalhes': "✨ *Detalhes da Oferta:*",
                'tipo': "• 🔑",
                'tamanho': "• 📏 Tamanho:",
                'economia': "• 💰 Economia:",
                'total_pagar': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido:",
                'checkout_pedido': "📋 *No checkout será pedido:*",
                'endereco': "1️⃣ **Endereço de entrega completo**",
                'metodo': "2️⃣ **Método de pagamento**",
                'metodos_disponiveis': "💳 *Métodos disponíveis:*",
                'seguro': "🔒 *Pagamento 100% seguro via Stripe*",
                'tempo': "⏰ *Tem 10 minutos para efetuar o pagamento*",
                'clique_abaixo': "Clique abaixo para pagar: 👇",
                'botao': "💳 PAGAR OFERTA →"
            },
            'ingles': {
                'titulo': "🎉 *SPECIAL OFFER CONFIRMED!* 🎉",
                'cliente': "👤 *Customer:*",
                'pais': "🌍 *Shipping Country:*",
                'moeda': "💰 *Currency:*",
                'detalhes': "✨ *Offer Details:*",
                'tipo': "• 🔑",
                'tamanho': "• 📏 Size:",
                'economia': "• 💰 Savings:",
                'total_pagar': "💳 **TOTAL TO PAY:",
                'pedido': "🆔 **Order:",
                'checkout_pedido': "📋 *In checkout you will be asked for:*",
                'endereco': "1️⃣ **Complete shipping address**",
                'metodo': "2️⃣ **Payment method**",
                'metodos_disponiveis': "💳 *Available methods:*",
                'seguro': "🔒 *100% secure payment via Stripe*",
                'tempo': "⏰ *You have 10 minutes to complete payment*",
                'clique_abaixo': "Click below to pay: 👇",
                'botao': "💳 PAY OFFER →"
            },
            'espanhol': {
                'titulo': "🎉 *¡OFERTA ESPECIAL CONFIRMADA!* 🎉",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envío:*",
                'moeda': "💰 *Moneda:*",
                'detalhes': "✨ *Detalles de la Oferta:*",
                'tipo': "• 🔑",
                'tamanho': "• 📏 Tamaño:",
                'economia': "• 💰 Ahorro:",
                'total_pagar': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido:",
                'checkout_pedido': "📋 *En el checkout se pedirá:*",
                'endereco': "1️⃣ **Dirección de envío completa**",
                'metodo': "2️⃣ **Método de pago**",
                'metodos_disponiveis': "💳 *Métodos disponibles:*",
                'seguro': "🔒 *Pago 100% seguro vía Stripe*",
                'tempo': "⏰ *Tienes 10 minutos para efectuar el pago*",
                'clique_abaixo': "Haz clic abajo para pagar: 👇",
                'botao': "💳 PAGAR OFERTA →"
            },
            'italiano': {
                'titulo': "🎉 *OFFERTA SPECIALE CONFERMATA!* 🎉",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *Paese di Spedizione:*",
                'moeda': "💰 *Valuta:*",
                'detalhes': "✨ *Dettagli dell'Offerta:*",
                'tipo': "• 🔑",
                'tamanho': "• 📏 Dimensione:",
                'economia': "• 💰 Risparmio:",
                'total_pagar': "💳 **TOTALE DA PAGARE:",
                'pedido': "🆔 **Ordine:",
                'checkout_pedido': "📋 *Nel checkout verrà richiesto:*",
                'endereco': "1️⃣ **Indirizzo di spedizione completo**",
                'metodo': "2️⃣ **Metodo di pagamento**",
                'metodos_disponiveis': "💳 *Metodi disponibili:*",
                'seguro': "🔒 *Pagamento 100% sicuro tramite Stripe*",
                'tempo': "⏰ *Hai 10 minuti per effettuare il pagamento*",
                'clique_abaixo': "Clicca qui sotto per pagare: 👇",
                'botao': "💳 PAGA OFFERTA →"
            },
            'alemao': {
                'titulo': "🎉 *SONDERANGEBOT BESTÄTIGT!* 🎉",
                'cliente': "👤 *Kunde:*",
                'pais': "🌍 *Versandland:*",
                'moeda': "💰 *Währung:*",
                'detalhes': "✨ *Angebotsdetails:*",
                'tipo': "• 🔑",
                'tamanho': "• 📏 Größe:",
                'economia': "• 💰 Ersparnis:",
                'total_pagar': "💳 **GESAMTBETRAG ZU ZAHLEN:",
                'pedido': "🆔 **Bestellung:",
                'checkout_pedido': "📋 *Im Checkout wird angefordert:*",
                'endereco': "1️⃣ **Vollständige Lieferadresse**",
                'metodo': "2️⃣ **Zahlungsmethode**",
                'metodos_disponiveis': "💳 *Verfügbare Methoden:*",
                'seguro': "🔒 *100% sichere Zahlung über Stripe*",
                'tempo': "⏰ *Sie haben 10 Minuten für die Zahlung*",
                'clique_abaixo': "Klicken Sie unten zum Bezahlen: 👇",
                'botao': "💳 ANGEBOT BEZAHLEN →"
            },
            'frances': {
                'titulo': "🎉 *OFFRE SPÉCIALE CONFIRMÉE !* 🎉",
                'cliente': "👤 *Client:*",
                'pais': "🌍 *Pays de Livraison:*",
                'moeda': "💰 *Devise:*",
                'detalhes': "✨ *Détails de l'Offre:*",
                'tipo': "• 🔑",
                'tamanho': "• 📏 Taille:",
                'economia': "• 💰 Économie:",
                'total_pagar': "💳 **TOTAL À PAYER:",
                'pedido': "🆔 **Commande:",
                'checkout_pedido': "📋 *Dans le checkout, il sera demandé:*",
                'endereco': "1️⃣ **Adresse de livraison complète**",
                'metodo': "2️⃣ **Méthode de paiement**",
                'metodos_disponiveis': "💳 *Méthodes disponibles:*",
                'seguro': "🔒 *Paiement 100% sécurisé via Stripe*",
                'tempo': "⏰ *Vous avez 10 minutos para efectuar el pago*",
                'clique_abaixo': "Cliquez ci-dessous pour payer : 👇",
                'botao': "💳 PAYER L'OFFRE →"
            }
        }
        
        textos = textos_oferta_confirmada.get(idioma, textos_oferta_confirmada['portugues'])

        await query.edit_message_text(
            text=(
                f"{textos['titulo']}\n\n"
                f"{textos['cliente']} {pedido['nome']}\n"
                f"{textos['pais']} {pedido['pais']}\n"
                f"{textos['moeda']} {codigo_moeda.upper()} {simbolo_moeda}\n\n"
                
                f"{textos['detalhes']}\n"
                f"{textos['tipo']} {pedido['tipo_cartoon']}\n"
                f"{textos['tamanho']} {tamanho_portachaves}\n"
                f"{textos['economia']} {simbolo_moeda}{economia:.2f}\n\n"
                
                f"{textos['total_pagar']} {simbolo_moeda}{total_oferta:.2f}**\n"
                f"{textos['pedido']} #{pedido_id}**\n\n"
                
                f"{textos['checkout_pedido']}\n"
                f"{textos['endereco']}\n"
                f"{textos['metodo']}\n\n"
                f"{textos['metodos_disponiveis']} {texto_metodos}\n"
                f"{textos['seguro']}\n\n"
                f"{textos['tempo']}\n\n"
                f"{textos['clique_abaixo']}"
            ),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos['botao'], url=session.url)]
            ])
        )
        
        print(f"✅ Usuário redirecionado para Checkout (Oferta Porta-chaves {tamanho_portachaves}) | Idioma: {idioma}")

        # 🔥 INICIAR TEMPORIZADOR PARA OFERTA
        await iniciar_temporizador_oferta(context, pedido_id, chat_id, query.message.message_id, idioma)
        
    except Exception as e:
        print(f"❌ ERRO STRIPE NA OFERTA PORTA-CHAVES: {str(e)}")
        print(f"🔍 Tipo do erro: {type(e)}")
        
        import traceback
        print(f"🔍 Traceback completo: {traceback.format_exc()}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro_pagamento = {
            'portugues': {
                'erro': "❌ Erro no processamento do pagamento.",
                'tentar': "🔄 Tentar Novamente",
                'suporte': "📞 Suporte"
            },
            'ingles': {
                'erro': "❌ Error processing payment.",
                'tentar': "🔄 Try Again",
                'suporte': "📞 Support"
            },
            'espanhol': {
                'erro': "❌ Error en el procesamiento del pago.",
                'tentar': "🔄 Intentar de Nuevo",
                'suporte': "📞 Soporte"
            },
            'italiano': {
                'erro': "❌ Errore nell'elaborazione del pagamento.",
                'tentar': "🔄 Riprova",
                'suporte': "📞 Supporto"
            },
            'alemao': {
                'erro': "❌ Fehler bei der Zahlungsverarbeitung.",
                'tentar': "🔄 Erneut versuchen",
                'suporte': "📞 Support"
            },
            'frances': {
                'erro': "❌ Erreur lors du traitement du paiement.",
                'tentar': "🔄 Réessayer",
                'suporte': "📞 Support"
            }
        }
        
        textos_erro = textos_erro_pagamento.get(idioma, textos_erro_pagamento['portugues'])
        
        await query.edit_message_text(
            f"{textos_erro['erro']}\n"
            "Por favor, tente novamente em alguns segundos.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos_erro['tentar'], callback_data=f"pagar_portachaves_{pedido_id}")],
                [InlineKeyboardButton(textos_erro['suporte'], callback_data=f"todas_recusadas_{pedido_id}")]
            ])
        )













async def proxima_oferta(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para mostrar a próxima oferta (tamanho 4.5cm)"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("proxima_oferta_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.edit_message_text("❌ Pedido não encontrado.")
        return
    
    print(f"🔄 PRÓXIMA OFERTA: #{pedido_id}")
    
    # 🔥 CORREÇÃO: LIMPAR EMOJIS E ESPAÇOS
    tipo_cartoon = pedido['tipo_cartoon'].lower()
    estilo_cartoon = pedido.get('estilo_cartoon', '').lower()
    
    # Remover emojis e caracteres especiais
    import re
    tipo_limpo = re.sub(r'[^\w\s]', '', tipo_cartoon).strip()
    estilo_limpo = re.sub(r'[^\w\s]', '', estilo_cartoon).strip()
    
    print(f"🔍 DEBUG - Tipo: '{tipo_cartoon}' → '{tipo_limpo}'")
    print(f"🔍 DEBUG - Estilo: '{estilo_cartoon}' → '{estilo_limpo}'")
    
    # TIPOS QUE NUNCA TÊM TAMANHO 4.5cm
    tipos_proibidos = ['animal', 'personalizado']
    
    # VERIFICAR SE É BUSTO (individual + estilo bust)
    # Agora compara com as versões limpas
    eh_busto = (tipo_limpo == 'cartoon individual' and estilo_limpo == 'bust')
    
    # VERIFICAR SE PODE OFERECER TAMANHO 4.5cm
    pode_ofertar_tamanho_45 = True
    
    # Se for qualquer um dos tipos proibidos (usando versão limpa)
    for tipo_proibido in tipos_proibidos:
        if tipo_proibido in tipo_limpo:
            pode_ofertar_tamanho_45 = False
            print(f"🚫 TIPO PROIBIDO: {tipo_limpo}")
            break
    
    # Se for busto, também não oferece 4.5cm
    if eh_busto:
        pode_ofertar_tamanho_45 = False
        print(f"🚫 ESTILO BUSTO DETETADO: {tipo_limpo} + {estilo_limpo}")
    
    print(f"🎯 RESULTADO: Oferecer 4.5cm? {pode_ofertar_tamanho_45}")
    
    if pode_ofertar_tamanho_45:
        print(f"✅ OFERTANDO TAMANHO 4.5cm")
        await mostrar_oferta_tamanho_45(context, pedido, query.message.chat_id, query.message.message_id)
    else:
        print(f"🚫 PULANDO PARA PORTA-CHAVES")
        await mostrar_oferta_portachaves(context, pedido, query.message.chat_id, query.message.message_id)








async def ultima_oferta(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para mostrar a última oferta (porta-chaves)"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("ultima_oferta_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.edit_message_text("❌ Pedido não encontrado.")
        return
    
    print(f"🔄 ÚLTIMA OFERTA (PORTA-CHAVES): #{pedido_id}")
    await mostrar_oferta_portachaves(context, pedido, query.message.chat_id, query.message.message_id)








async def sair_oferta(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando usuário clica em 'Sair e Perder Esta Oferta' - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("sair_oferta_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande introuvable."
        }
        
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"😔 USUÁRIO QUER SAIR: #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 TEXTOS EMOCIONAIS POR IDIOMA
    textos_emocionais = {
        'portugues': {
            'titulo': "😔 *Quer mesmo perder esta oportunidade?*",
            'mensagem': "💭 *Pense bem...*\n"
                       "Alguns momentos merecem ser guardados para a sua história.\n"
                       "O tempo não volta, mas você pode eternizá-lo agora.\n"
                       "Vai mesmo deixar só guardado na memória ou\n"
                       "transformar em algo para ver e sorrir todos dias?\n",
            'pergunta': "*Quer viver este momento?* 👇",
            'botao_sim': "❤️ Sim, Quero Eterniza-lo!",
            'botao_nao': "❌ Não, Recusar e sair"
        },
        'ingles': {
            'titulo': "😔 *Do you really want to miss this opportunity?*",
            'mensagem': "💭 *Think carefully...*\n"
                       "Some moments deserve to be kept for your history.\n"
                       "Time doesn't go back, but you can immortalize it now.\n"
                       "Will you just leave it stored in memory or\n"
                       "turn it into something to see and smile every day?\n",
            'pergunta': "*Do you want to live this moment?* 👇",
            'botao_sim': "❤️ Yes, I Want to Immortalize It!",
            'botao_nao': "❌ No, Reject and leave"
        },
        'espanhol': {
            'titulo': "😔 *¿Realmente quieres perder esta oportunidad?*",
            'mensagem': "💭 *Piensa bien...*\n"
                       "Algunos momentos merecen ser guardados para tu historia.\n"
                       "El tiempo no vuelve, pero puedes eternizarlo ahora.\n"
                       "¿Vas a dejarlo solo guardado en la memoria o\n"
                       "transformarlo en algo para ver y sonreír todos los días?\n",
            'pergunta': "*¿Quieres vivir este momento?* 👇",
            'botao_sim': "❤️ Sí, ¡Quiero Eternizarlo!",
            'botao_nao': "❌ No, Rechazar y salir"
        },
        'italiano': {
            'titulo': "😔 *Vuoi davvero perdere questa opportunità?*",
            'mensagem': "💭 *Pensa bene...*\n"
                       "Alcuni momenti meritano di essere conservati per la tua storia.\n"
                       "Il tempo non torna indietro, ma puoi eternizzarlo ora.\n"
                       "Lo lascerai solo conservato nella memoria o\n"
                       "lo trasformerai in qualcosa da vedere e sorridere ogni giorno?\n",
            'pergunta': "*Vuoi vivere questo momento?* 👇",
            'botao_sim': "❤️ Sì, Voglio Eternizzarlo!",
            'botao_nao': "❌ No, Rifiuta ed esci"
        },
        'alemao': {
            'titulo': "😔 *Möchten Sie diese Gelegenheit wirklich verpassen?*",
            'mensagem': "💭 *Überlegen Sie gut...*\n"
                       "Einige Momente verdienen es, für Ihre Geschichte bewahrt zu werden.\n"
                       "Die Zeit kommt nicht zurück, aber Sie können sie jetzt verewigen.\n"
                       "Lassen Sie es nur im Gedächtnis gespeichert oder\n"
                       "verwandeln Sie es in etwas, das Sie jeden Tag sehen und lächeln lässt?\n",
            'pergunta': "*Möchten Sie diesen Moment erleben?* 👇",
            'botao_sim': "❤️ Ja, Ich Will Es Verewigen!",
            'botao_nao': "❌ Nein, Ablehnen und verlassen"
        },
        'frances': {
            'titulo': "😔 *Voulez-vous vraiment manquer cette opportunité ?*",
            'mensagem': "💭 *Réfléchissez bien...*\n"
                       "Certains moments méritent d'être conservés pour votre histoire.\n"
                       "Le temps ne revient pas, mais vous pouvez l'éterniser maintenant.\n"
                       "Allez-vous le laisser juste stocké dans la mémoire ou\n"
                       "le transformer en quelque chose à voir et sourire tous les jours ?\n",
            'pergunta': "*Voulez-vous vivre ce moment ?* 👇",
            'botao_sim': "❤️ Oui, Je Veux l'Éterniser !",
            'botao_nao': "❌ Non, Refuser et partir"
        }
    }
    
    textos = textos_emocionais.get(idioma, textos_emocionais['portugues'])
    
    texto_emocional = (
        f"{textos['titulo']}\n\n"
        f"{textos['mensagem']}\n"
        f"{textos['pergunta']}"
    )
    
    botoes_emocionais = [
        [InlineKeyboardButton(textos['botao_sim'], callback_data=f"recuperar_pagar_{pedido_id}")],
        [InlineKeyboardButton(textos['botao_nao'], callback_data=f"confirmar_saida_{pedido_id}")]
    ]
    
    await query.edit_message_text(
        text=texto_emocional,
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(botoes_emocionais)
    )

# 🔥 🔥 🔥 CORREÇÃO: FUNÇÃO SEPARADA - NÃO DENTRO DA OUTRA
async def confirmar_saida(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando usuário confirma que quer sair - COM CONTAGEM DE RECUSA E TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("confirmar_saida_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande introuvable."
        }
        
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"😞 USUÁRIO CONFIRMOU SAÍDA: #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 MARCAR QUE RECUSOU OFERTA ORIGINAL
    pedido["recusou_original"] = True
    print(f"📝 Pedido #{pedido_id} marcado como recusou oferta original")
    
    # 🔥 VERIFICAR SE JÁ RECUSOU OUTRAS OFERTAS
    recusou_45 = pedido.get("recusou_oferta_45", False)
    recusou_portachaves = pedido.get("recusou_portachaves", False)
    
    # Se já recusou as outras 2, contar como recusou todas
    if recusou_45 and recusou_portachaves:
        ESTATISTICAS['ofertas_recusadas'] = ESTATISTICAS.get('ofertas_recusadas', 0) + 1
        print(f"🎯 USUÁRIO RECUSOU TODAS AS 3 OFERTAS: #{pedido_id}")
        print(f"📈 Estatística atualizada: Ofertas recusadas")
    
    # 🔥 TEXTOS DAS ESPERAS POR IDIOMA
    textos_esperas = {
        'portugues': {
            'primeira': "⏳ *Espere... Deixe-me pensar um momento...*\n\nEstou a refletir na sua decisão...",
            'segunda': "💭 *Ainda estou a pensar...*\n\nHá sempre uma maneira de tornar tudo possível...",
            'solucao': "🎉 *ESPERE! Acho que encontrei uma solução!*\n\n"
                      "🌟 *Vou oferecer uma versão exclusiva somente para você,\n"
                      "com um desconto que vai surpreender!*\n\n"
                      "*O que acha disto?* 👇"
        },
        'ingles': {
            'primeira': "⏳ *Wait... Let me think for a moment...*\n\nI'm reflecting on your decision...",
            'segunda': "💭 *I'm still thinking...*\n\nThere's always a way to make everything possible...",
            'solucao': "🎉 *WAIT! I think I found a solution!*\n\n"
                      "🌟 *I'm going to offer an exclusive version just for you,\n"
                      "with a discount that will surprise you!*\n\n"
                      "*What do you think of this?* 👇"
        },
        'espanhol': {
            'primeira': "⏳ *Espere... Déjeme pensar un momento...*\n\nEstoy reflexionando sobre su decisión...",
            'segunda': "💭 *Todavía estoy pensando...*\n\nSiempre hay una manera de hacer todo posible...",
            'solucao': "🎉 *¡ESPERE! ¡Creo que encontré una solución!*\n\n"
                      "🌟 *Voy a ofrecer una versión exclusiva solo para usted,\n"
                      "¡con un descuento que le sorprenderá!*\n\n"
                      "*¿Qué le parece esto?* 👇"
        },
        'italiano': {
            'primeira': "⏳ *Aspetta... Fammi pensare un momento...*\n\nSto riflettendo sulla tua decisione...",
            'segunda': "💭 *Sto ancora pensando...*\n\nC'è sempre un modo per rendere tutto possibile...",
            'solucao': "🎉 *ASPETTA! Penso di aver trovato una soluzione!*\n\n"
                      "🌟 *Offrirò una versione esclusiva solo per te,\n"
                      "con uno sconto che ti sorprenderà!*\n\n"
                      "*Cosa ne pensi di questo?* 👇"
        },
        'alemao': {
            'primeira': "⏳ *Warten Sie... Lassen Sie mich einen Moment nachdenken...*\n\nIch überlege Ihre Entscheidung...",
            'segunda': "💭 *Ich denke noch nach...*\n\nEs gibt immer einen Weg, alles möglich zu machen...",
            'solucao': "🎉 *WARTEN SIE! Ich glaube, ich habe eine Lösung gefunden!*\n\n"
                      "🌟 *Ich werde Ihnen eine exklusive Version anbieten,\n"
                      "mit einem Rabatt, der Sie überraschen wird!*\n\n"
                      "*Was halten Sie davon?* 👇"
        },
        'frances': {
            'primeira': "⏳ *Attendez... Laissez-moi réfléchir un moment...*\n\nJe réfléchis à votre décision...",
            'segunda': "💭 *Je réfléchis encore...*\n\nIl y a toujours un moyen de tout rendre possible...",
            'solucao': "🎉 *ATTENDEZ ! Je crois que j'ai trouvé une solution !*\n\n"
                      "🌟 *Je vais vous offrir une version exclusive rien que pour vous,\n"
                      "avec une réduction qui va vous surprendre !*\n\n"
                      "*Qu'en pensez-vous ?* 👇"
        }
    }
    
    textos = textos_esperas.get(idioma, textos_esperas['portugues'])
    
    # PRIMEIRA ESPERA (1 minuto)
    await query.edit_message_text(
        text=textos['primeira'],
        parse_mode="Markdown"
    )
    
    # Espera 1 minuto
    await asyncio.sleep(60)
    
    # SEGUNDA MENSAGEM (após 1 minuto)
    await context.bot.edit_message_text(
        chat_id=query.message.chat_id,
        message_id=query.message.message_id,
        text=textos['segunda'],
        parse_mode="Markdown"
    )
    
    # Espera mais 1 minuto
    await asyncio.sleep(60)
    
    # MENSAGEM FINAL COM SOLUÇÃO
    await context.bot.edit_message_text(
        chat_id=query.message.chat_id,
        message_id=query.message.message_id,
        text=textos['solucao'],
        parse_mode="Markdown"
    )
    
    # Espera 5 segundos e mostra a segunda oferta
    await asyncio.sleep(10)
    
    # 🔥 ATUALIZAR PEDIDO ANTES DE MOSTRAR PRÓXIMA OFERTA
    pedido["recusou_original"] = True
    
    # Mostrar a próxima oferta (tamanho 4.5cm)
    await mostrar_oferta_tamanho_45(context, pedido, query.message.chat_id, query.message.message_id)






async def sair_diretoportachaves(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando usuário quer sair do fluxo direto para porta-chaves - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("sair_diretoportachaves_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande introuvable."
        }
        
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"😔 USUÁRIO QUER SAIR DO FLUXO DIRETO: #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 TEXTOS EMOCIONAIS ESPECÍFICOS PARA FLUXO DIRETO POR IDIOMA
    textos_emocionais_direto = {
        'portugues': {
            'titulo': "😔 *Quer mesmo deixar este momento especial passar?*",
            'mensagem': "💫 *Pense no significado...*\n"
                       "Ela capturou um instante único da sua vida.\n"
                       "Um momento que merece ser transformado em arte.\n"
                       "Não espere a saudade para perceber o valor do que você viveu.\n"
                       "Preencha o espaço entre a memória e a realidade hoje mesmo.\n\n",
            'pergunta': "*Quer aproveitar esta oportunidade única?* ✨",
            'botao_sim': "❤️ Sim, Pensando bem quero!",
            'botao_nao': "❌ Não, Quero Sair"
        },
        'ingles': {
            'titulo': "😔 *Do you really want to let this special moment pass?*",
            'mensagem': "💫 *Think about the meaning...*\n"
                       "It captured a unique moment of your life.\n"
                       "A moment that deserves to be transformed into art.\n"
                       "Don't wait for longing to realize the value of what you lived.\n"
                       "Fill the gap between memory and reality today.\n\n",
            'pergunta': "*Do you want to take this unique opportunity?* ✨",
            'botao_sim': "❤️ Yes, Thinking about it I want it!",
            'botao_nao': "❌ No, I Want to Leave"
        },
        'espanhol': {
            'titulo': "😔 *¿Realmente quieres dejar pasar este momento especial?*",
            'mensagem': "💫 *Piensa en el significado...*\n"
                       "Capturó un instante único de tu vida.\n"
                       "Un momento que merece ser transformado en arte.\n"
                       "No esperes a la nostalgia para darte cuenta del valor de lo que viviste.\n"
                       "Llena el espacio entre el recuerdo y la realidad hoy mismo.\n\n",
            'pergunta': "*¿Quieres aprovechar esta oportunidad única?* ✨",
            'botao_sim': "❤️ Sí, ¡Pensándolo bien quiero!",
            'botao_nao': "❌ No, Quiero Salir"
        },
        'italiano': {
            'titulo': "😔 *Vuoi davvero lasciar passare questo momento speciale?*",
            'mensagem': "💫 *Pensa al significato...*\n"
                       "Ha catturato un istante unico della tua vita.\n"
                       "Un momento che merita di essere trasformato in arte.\n"
                       "Non aspettare la nostalgia per realizzare il valore di ciò che hai vissuto.\n"
                       "Riempi lo spazio tra il ricordo e la realtà oggi stesso.\n\n",
            'pergunta': "*Vuoi cogliere questa opportunità unica?* ✨",
            'botao_sim': "❤️ Sì, Pensandoci bene lo voglio!",
            'botao_nao': "❌ No, Voglio Uscire"
        },
        'alemao': {
            'titulo': "😔 *Möchten Sie diesen besonderen Moment wirklich verpassen?*",
            'mensagem': "💫 *Denken Sie an die Bedeutung...*\n"
                       "Es hat einen einzigartigen Moment Ihres Lebens eingefangen.\n"
                       "Ein Moment, der es verdient, in Kunst verwandelt zu werden.\n"
                       "Warten Sie nicht auf Sehnsucht, um den Wert dessen zu erkennen, was Sie erlebt haben.\n"
                       "Füllen Sie die Lücke zwischen Erinnerung und Realität noch heute.\n\n",
            'pergunta': "*Möchten Sie diese einzigartige Gelegenheit nutzen?* ✨",
            'botao_sim': "❤️ Ja, Wenn ich darüber nachdenke, möchte ich es!",
            'botao_nao': "❌ Nein, Ich Möchte Verlassen"
        },
        'frances': {
            'titulo': "😔 *Voulez-vous vraiment laisser passer ce moment spécial ?*",
            'mensagem': "💫 *Pensez à la signification...*\n"
                       "Il a capturé un instant unique de votre vie.\n"
                       "Un moment qui mérite d'être transformé en art.\n"
                       "N'attendez pas la nostalgie pour réaliser la valeur de ce que vous avez vécu.\n"
                       "Comblez l'écart entre le souvenir et la réalité dès aujourd'hui.\n\n",
            'pergunta': "*Voulez-vous saisir cette opportunité unique ?* ✨",
            'botao_sim': "❤️ Oui, En y réfléchissant je le veux !",
            'botao_nao': "❌ Non, Je Veux Partir"
        }
    }
    
    textos = textos_emocionais_direto.get(idioma, textos_emocionais_direto['portugues'])
    
    texto_emocional = (
        f"{textos['titulo']}\n\n"
        f"{textos['mensagem']}\n"
        f"{textos['pergunta']}"
    )
    
    botoes_emocionais = [
        [InlineKeyboardButton(textos['botao_sim'], callback_data=f"recuperar_pagar_{pedido_id}")],
        [InlineKeyboardButton(textos['botao_nao'], callback_data=f"confirmar_saidadireta_{pedido_id}")]
    ]
    
    try:
        await query.edit_message_text(
            text=texto_emocional,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(botoes_emocionais)
        )
    except BadRequest:
        print(f"✅ Mensagem já está com o conteúdo correto - ignorando erro | Idioma: {idioma}")






async def confirmar_saidadireta(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando usuário confirma saída do fluxo direto - MARCA 4.5cm COMO RECUSADA MAS NÃO MOSTRA NO RELATÓRIO - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("confirmar_saidadireta_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande introuvable."
        }
        
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"😞 USUÁRIO CONFIRMOU SAÍDA DIRETA: #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 MARCAR COMO RECUSOU OFERTA ORIGINAL E 4.5cm (PARA CONTAR 3 RECUSAS)
    pedido["recusou_original"] = True
    pedido["recusou_oferta_45"] = True  # 🔥 MARCA COMO RECUSADA PARA CONTAGEM
    pedido["fluxo_direto"] = True  # 🔥 MARCAR COMO FLUXO DIRETO (NÃO MOSTRAR 4.5cm NO RELATÓRIO)
    print(f"📝 Pedido #{pedido_id} marcado como recusou oferta original (fluxo direto)")
    print(f"📝 Pedido #{pedido_id} marcado como recusou oferta 4.5cm (para contagem)")
    print(f"📝 Pedido #{pedido_id} marcado como fluxo_direto - NÃO MOSTRAR 4.5cm NO RELATÓRIO")
    
    # 🔥 VERIFICAR ESTADO ATUAL DAS RECUSAS
    recusou_original = pedido.get("recusou_original", False)
    recusou_45 = pedido.get("recusou_oferta_45", False)
    recusou_portachaves = pedido.get("recusou_portachaves", False)
    
    print(f"🔍 ESTADO DAS RECUSAS NO FLUXO DIRETO | Idioma: {idioma}:")
    print(f"   • Recusou original: {recusou_original}")
    print(f"   • Recusou 4.5cm: {recusou_45} (PARA CONTAGEM)") 
    print(f"   • Recusou porta-chaves: {recusou_portachaves}")
    
    # 🔥 TEXTOS DAS ESPERAS POR IDIOMA (FLUXO DIRETO)
    textos_esperas_direto = {
        'portugues': {
            'primeira': "⏳ *Um momento... Deixe-me pensar...*\n\nHá sempre uma solução para momentos especiais...",
            'segunda': "💭 *Ainda estou a refletir...*\n\nNão podemos deixar esta memória escapar assim...",
            'solucao': "🎊 *ESPERE! Encontrei algo REALMENTE ESPECIAL!*\n\n"
                      "Para o seu tipo único de cartoon, preparei uma\n"
                      "proposta que vai surpreender...\n\n"
                      "🌟 *Imagine ter esta memória sempre consigo,*\n"
                      "em formato premium, a um valor incrível!\n\n"
                      "*Quer ver esta proposta exclusiva?* 👇"
        },
        'ingles': {
            'primeira': "⏳ *One moment... Let me think...*\n\nThere's always a solution for special moments...",
            'segunda': "💭 *I'm still reflecting...*\n\nWe can't let this memory escape like this...",
            'solucao': "🎊 *WAIT! I found something TRULY SPECIAL!*\n\n"
                      "For your unique type of cartoon, I've prepared a\n"
                      "proposal that will surprise you...\n\n"
                      "🌟 *Imagine having this memory with you always,*\n"
                      "in premium format, at an incredible value!\n\n"
                      "*Do you want to see this exclusive proposal?* 👇"
        },
        'espanhol': {
            'primeira': "⏳ *Un momento... Déjeme pensar...*\n\nSiempre hay una solución para momentos especiales...",
            'segunda': "💭 *Todavía estoy reflexionando...*\n\nNo podemos dejar escapar este recuerdo así...",
            'solucao': "🎊 *¡ESPERE! ¡Encontré algo REALMENTE ESPECIAL!*\n\n"
                      "Para su tipo único de caricatura, he preparado una\n"
                      "propuesta que le sorprenderá...\n\n"
                      "🌟 *Imagine tener este recuerdo siempre con usted,*\n"
                      "en formato premium, a un valor increíble!\n\n"
                      "*¿Quiere ver esta propuesta exclusiva?* 👇"
        },
        'italiano': {
            'primeira': "⏳ *Un momento... Fammi pensare...*\n\nC'è sempre una soluzione per momenti speciali...",
            'segunda': "💭 *Sto ancora riflettendo...*\n\nNon possiamo lasciare sfuggire questo ricordo così...",
            'solucao': "🎊 *ASPETTA! Ho trovato qualcosa di VERAMENTE SPECIALE!*\n\n"
                      "Per il tuo tipo unico di cartoon, ho preparato una\n"
                      "proposta che ti sorprenderà...\n\n"
                      "🌟 *Immagina di avere questo ricordo sempre con te,*\n"
                      "in formato premium, a un valore incredibile!\n\n"
                      "*Vuoi vedere questa proposta esclusiva?* 👇"
        },
        'alemao': {
            'primeira': "⏳ *Einen Moment... Lassen Sie mich nachdenken...*\n\nEs gibt immer eine Lösung für besondere Momente...",
            'segunda': "💭 *Ich denke noch nach...*\n\nWir können diese Erinnerung nicht so entkommen lassen...",
            'solucao': "🎊 *WARTEN SIE! Ich habe etwas WIRKLICH BESONDERES gefunden!*\n\n"
                      "Für Ihre einzigartige Karikaturart habe ich einen\n"
                      "Vorschlag vorbereitet, der Sie überraschen wird...\n\n"
                      "🌟 *Stellen Sie sich vor, diese Erinnerung immer bei sich zu haben,*\n"
                      "in Premium-Format, zu einem unglaublichen Wert!\n\n"
                      "*Möchten Sie diesen exklusiven Vorschlag sehen?* 👇"
        },
        'frances': {
            'primeira': "⏳ *Un moment... Laissez-moi réfléchir...*\n\nIl y a toujours une solution pour les moments spéciaux...",
            'segunda': "💭 *Je réfléchis encore...*\n\nNous ne pouvons pas laisser ce souvenir s'échapper ainsi...",
            'solucao': "🎊 *ATTENDEZ ! J'ai trouvé quelque chose de VRAIMENT SPÉCIAL !*\n\n"
                      "Pour votre type unique de dessin animé, j'ai préparé une\n"
                      "proposition qui va vous surprendre...\n\n"
                      "🌟 *Imaginez avoir ce souvenir toujours avec vous,*\n"
                      "en format premium, à une valeur incroyable !\n\n"
                      "*Voulez-vous voir cette proposition exclusive ?* 👇"
        }
    }
    
    textos = textos_esperas_direto.get(idioma, textos_esperas_direto['portugues'])
    
    # ESPERA E REFLEXÃO
    await query.edit_message_text(
        text=textos['primeira'],
        parse_mode="Markdown"
    )
    
    # Espera 1 minuto
    await asyncio.sleep(60)
    
    # SEGUNDA MENSAGEM
    await context.bot.edit_message_text(
        chat_id=query.message.chat_id,
        message_id=query.message.message_id,
        text=textos['segunda'],
        parse_mode="Markdown"
    )
    
    # Espera mais 1 minuto
    await asyncio.sleep(60)
    
    # MENSAGEM FINAL COM PROPOSTA ESPECIAL
    await context.bot.edit_message_text(
        chat_id=query.message.chat_id,
        message_id=query.message.message_id,
        text=textos['solucao'],
        parse_mode="Markdown"
    )
    
    # Espera 3 segundos e mostra o porta-chaves
    await asyncio.sleep(10)
    await mostrar_oferta_portachaves(context, pedido, query.message.chat_id, query.message.message_id)








async def sair_poferta45(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando usuário recusa a oferta do 4.5cm - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    print(f"🎯 DEBUG: sair_poferta45 CHAMADO - data: {query.data}")  # 🔥 DEBUG
    
    pedido_id = query.data.replace("sair_poferta45_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        print(f"❌ DEBUG: Pedido {pedido_id} NÃO ENCONTRADO")  # 🔥 DEBUG
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande introuvable."
        }
        
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"😔 USUÁRIO RECUSOU OFERTA 4.5 | 11.5cm : #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 TEXTOS EMOCIONAIS POR IDIOMA
    textos_emocionais = {
        'portugues': {
            'titulo': "😔 *Tem certeza que quer recusar esta oportunidade única?*",
            'mensagem': "💭 *Pense no valor que está a deixar passar...*\n"
                       "Esta edição colecionador foi criada especialmente para você.\n"
                       "Não espere a saudade para perceber o valor do que você viveu.\n"
                       "Preencha o espaço entre a memória e a realidade hoje mesmo.\n\n",
            'pergunta': "*Vamos aproveitar esta oportunidade única?* 👇",
            'botao_sim': "❤️ Sim, Quero Aproveitar!",
            'botao_nao': "❌ Não, Recusar e Sair"
        },
        'ingles': {
            'titulo': "😔 *Are you sure you want to refuse this unique opportunity?*",
            'mensagem': "💭 *Think about the value you're letting go...*\n"
                       "This collector's edition was created especially for you.\n"
                       "Don't wait for longing to realize the value of what you lived.\n"
                       "Fill the gap between memory and reality today.\n\n",
            'pergunta': "*Shall we take this unique opportunity?* 👇",
            'botao_sim': "❤️ Yes, I Want to Take Advantage!",
            'botao_nao': "❌ No, Refuse and Leave"
        },
        'espanhol': {
            'titulo': "😔 *¿Está seguro de que quiere rechazar esta oportunidad única?*",
            'mensagem': "💭 *Piense en el valor que está dejando pasar...*\n"
                       "Esta edición de coleccionista fue creada especialmente para usted.\n"
                       "No espere a la nostalgia para darse cuenta del valor de lo que vivió.\n"
                       "Llame el espacio entre el recuerdo y la realidad hoy mismo.\n\n",
            'pergunta': "*¿Aprovechamos esta oportunidad única?* 👇",
            'botao_sim': "❤️ Sí, ¡Quiero Aprovechar!",
            'botao_nao': "❌ No, Rechazar y Salir"
        },
        'italiano': {
            'titulo': "😔 *È sicuro di voler rifiutare questa opportunità unica?*",
            'mensagem': "💭 *Pensi al valore che sta lasciando passare...*\n"
                       "Questa edizione da collezione è stata creata appositamente per lei.\n"
                       "Non aspetti la nostalgia per rendersi conto del valore di ciò che ha vissuto.\n"
                       "Riempa lo spazio tra il ricordo e la realtà oggi stesso.\n\n",
            'pergunta': "*Approfittiamo di questa opportunità unica?* 👇",
            'botao_sim': "❤️ Sì, Voglio Approfittarne!",
            'botao_nao': "❌ No, Rifiutare e Uscire"
        },
        'alemao': {
            'titulo': "😔 *Sind Sie sicher, dass Sie diese einzigartige Gelegenheit ablehnen möchten?*",
            'mensagem': "💭 *Denken Sie an den Wert, den Sie verpassen...*\n"
                       "Diese Sammlerausgabe wurde speziell für Sie erstellt.\n"
                       "Warten Sie nicht auf Sehnsucht, um den Wert dessen zu erkennen, was Sie erlebt haben.\n"
                       "Füllen Sie die Lücke zwischen Erinnerung und Realität noch heute.\n\n",
            'pergunta': "*Nutzen wir diese einzigartige Gelegenheit?* 👇",
            'botao_sim': "❤️ Ja, Ich Will Nutzen!",
            'botao_nao': "❌ Nein, Ablehnen und Verlassen"
        },
        'frances': {
            'titulo': "😔 *Êtes-vous sûr de vouloir refuser cette opportunité unique ?*",
            'mensagem': "💭 *Pensez à la valeur que vous laissez passer...*\n"
                       "Cette édition collectionneur a été créée spécialement pour vous.\n"
                       "N'attendez pas la nostalgie pour réaliser la valeur de ce que vous avez vécu.\n"
                       "Comblez l'écart entre le souvenir et la réalité dès aujourd'hui.\n\n",
            'pergunta': "*Profitons-nous de cette opportunité unique ?* 👇",
            'botao_sim': "❤️ Oui, Je Veux en Profiter !",
            'botao_nao': "❌ Non, Refuser et Partir"
        }
    }
    
    textos = textos_emocionais.get(idioma, textos_emocionais['portugues'])
    
    texto_emocional = (
        f"{textos['titulo']}\n\n"
        f"{textos['mensagem']}\n"
        f"{textos['pergunta']}"
    )
    
    botoes_emocionais = [
        [InlineKeyboardButton(textos['botao_sim'], callback_data=f"pagar_tamanho45_{pedido_id}")],
        [InlineKeyboardButton(textos['botao_nao'], callback_data=f"confirmar_saida45_{pedido_id}")]
    ]
    
    try:
        await query.edit_message_text(
            text=texto_emocional,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(botoes_emocionais)
        )
        print(f"✅ Mensagem emocional exibida | Idioma: {idioma}")
    except BadRequest:
        print(f"✅ Mensagem já está com o conteúdo correto - ignorando erro | Idioma: {idioma}")


async def confirmar_saida45(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando usuário confirma saída da oferta 4.5cm - COM CONTAGEM DE RECUSA E TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("confirmar_saida45_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande introuvable."
        }
        
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 PEGAR IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    
    print(f"😞 USUÁRIO CONFIRMOU SAÍDA 4.5cm: #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 MARCAR QUE RECUSOU OFERTA 4.5cm
    pedido["recusou_oferta_45"] = True
    print(f"📝 Pedido #{pedido_id} marcado como recusou oferta 4.5cm")
    
    # 🔥 VERIFICAR SE JÁ RECUSOU OUTRAS OFERTAS
    recusou_original = pedido.get("recusou_original", False)
    recusou_portachaves = pedido.get("recusou_portachaves", False)
    
    # Se já recusou as outras 2, contar como recusou todas
    if recusou_original and recusou_portachaves:
        ESTATISTICAS['ofertas_recusadas'] = ESTATISTICAS.get('ofertas_recusadas', 0) + 1
        print(f"🎯 USUÁRIO RECUSOU TODAS AS 3 OFERTAS: #{pedido_id}")
        print(f"📈 Estatística atualizada: Ofertas recusadas")
    
    # 🔥 TEXTOS DAS ESPERAS POR IDIOMA (PARA 4.5cm)
    textos_esperas_45 = {
        'portugues': {
            'primeira': "⏳ *Estou a processar a sua decisão...*\n\nDeixe-me ver se há outra maneira...",
            'segunda': "💭 *Ainda estou a pensar...*\n\nQuero mesmo ajudar a eternizar este momento...",
            'solucao': "🎁 *ESPERE! Tenho uma última proposta ESPECIAL!*\n\n"
                      "Percebi que não podemos deixar este momento passar...\n\n"
                      "🌟 *E se eu lhe oferecer uma maneira de levar consigo\n"
                      "esta memória para todo lado, a um valor simbólico?*\n\n"
                      "*Está pronto para ver a nossa proposta final?* 👇"
        },
        'ingles': {
            'primeira': "⏳ *I'm processing your decision...*\n\nLet me see if there's another way...",
            'segunda': "💭 *I'm still thinking...*\n\nI really want to help immortalize this moment...",
            'solucao': "🎁 *WAIT! I have a final SPECIAL proposal!*\n\n"
                      "I realized we can't let this moment pass...\n\n"
                      "🌟 *What if I offer you a way to carry this memory\n"
                      "with you everywhere, at a symbolic price?*\n\n"
                      "*Are you ready to see our final proposal?* 👇"
        },
        'espanhol': {
            'primeira': "⏳ *Estoy procesando su decisión...*\n\nDéjeme ver si hay otra manera...",
            'segunda': "💭 *Todavía estoy pensando...*\n\nRealmente quiero ayudar a eternizar este momento...",
            'solucao': "🎁 *¡ESPERE! ¡Tengo una última propuesta ESPECIAL!*\n\n"
                      "Me di cuenta de que no podemos dejar pasar este momento...\n\n"
                      "🌟 *¿Y si le ofrezco una manera de llevar consigo\n"
                      "este recuerdo a todas partes, a un precio simbólico?*\n\n"
                      "*¿Está listo para ver nuestra propuesta final?* 👇"
        },
        'italiano': {
            'primeira': "⏳ *Sto elaborando la sua decisione...*\n\nLasciami vedere se c'è un altro modo...",
            'segunda': "💭 *Sto ancora pensando...*\n\nVoglio davvero aiutare a eternizzare questo momento...",
            'solucao': "🎁 *ASPETTA! Ho una proposta FINALE SPECIALE!*\n\n"
                      "Mi sono reso conto che non possiamo lasciar passare questo momento...\n\n"
                      "🌟 *E se le offrissi un modo per portare con sé\n"
                      "questo ricordo ovunque, a un prezzo simbolico?*\n\n"
                      "*È pronto per vedere la nostra proposta finale?* 👇"
        },
        'alemao': {
            'primeira': "⏳ *Ich verarbeite Ihre Entscheidung...*\n\nLassen Sie mich sehen, ob es einen anderen Weg gibt...",
            'segunda': "💭 *Ich denke noch nach...*\n\nIch möchte wirklich helfen, diesen Moment zu verewigen...",
            'solucao': "🎁 *WARTEN SIE! Ich habe einen letzten BESONDEREN Vorschlag!*\n\n"
                      "Ich habe erkannt, dass wir diesen Moment nicht verpassen können...\n\n"
                      "🌟 *Was, wenn ich Ihnen eine Möglichkeit biete, diese Erinnerung\n"
                      "überallhin mitzunehmen, zu einem symbolischen Preis?*\n\n"
                      "*Sind Sie bereit, unseren endgültigen Vorschlag zu sehen?* 👇"
        },
        'frances': {
            'primeira': "⏳ *Je traite votre décision...*\n\nLaissez-moi voir s'il y a une autre façon...",
            'segunda': "💭 *Je réfléchis encore...*\n\nJe veux vraiment aider à éterniser ce moment...",
            'solucao': "🎁 *ATTENDEZ ! J'ai une dernière proposition SPÉCIALE !*\n\n"
                      "J'ai réalisé que nous ne pouvons pas laisser passer ce moment...\n\n"
                      "🌟 *Et si je vous offrais un moyen d'emporter ce souvenir\n"
                      "partout avec vous, à un prix symbolique ?*\n\n"
                      "*Êtes-vous prêt à voir notre proposition finale ?* 👇"
        }
    }
    
    textos = textos_esperas_45.get(idioma, textos_esperas_45['portugues'])
    
    # ESPERA E REFLEXÃO
    await query.edit_message_text(
        text=textos['primeira'],
        parse_mode="Markdown"
    )
    
    # Espera 1 minuto
    await asyncio.sleep(60)
    
    # SEGUNDA MENSAGEM
    await context.bot.edit_message_text(
        chat_id=query.message.chat_id,
        message_id=query.message.message_id,
        text=textos['segunda'],
        parse_mode="Markdown"
    )
    
    # Espera mais 1 minuto
    await asyncio.sleep(60)
    
    # MENSAGEM FINAL
    await context.bot.edit_message_text(
        chat_id=query.message.chat_id,
        message_id=query.message.message_id,
        text=textos['solucao'],
        parse_mode="Markdown"
    )
    
    # Espera 10 segundos e mostra a oferta final
    await asyncio.sleep(10)
    await mostrar_oferta_portachaves(context, pedido, query.message.chat_id, query.message.message_id)
    
    print(f"✅ Fluxo de retenção 4.5cm concluído | Usuário direcionado para porta-chaves | Idioma: {idioma}")






async def iniciar_novaencomenda(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para o botão de nova encomenda - REPETE a lógica do /start - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    chat_id = query.message.chat_id
    
    # 🔥 PEGAR IDIOMA DO USER_DATA (se existir)
    idioma = context.user_data.get('idioma', 'portugues')
    
    print(f"🔧 Botão 'nova encomenda' clicado para chat {chat_id} | Idioma: {idioma}")
    
    # 🔥 MESMA LÓGICA DO /start - LIMPAR TUDO COMPLETAMENTE
    cancelar_temporizador_30min(chat_id)
    context.user_data.clear()
    
    if 'conversation_state' in context.user_data:
        del context.user_data['conversation_state']
    
    print(f"✅ Dados limpos via botão nova encomenda para chat {chat_id}")
    
    # 🔥 TEXTOS DO MENU INICIAL POR IDIOMA
    textos_inicio = {
        'portugues': {
            'saudacao': "👋 Vamos criar o seu *novo cartoon*?",
            'botao': "CREATE MY CARTOON"
        },
        'ingles': {
            'saudacao': "👋 Let's create your *new cartoon*?",
            'botao': "CREATE MY CARTOON"
        },
        'espanhol': {
            'saudacao': "👋 ¿Vamos a crear tu *nuevo cartoon*?",
            'botao': "CREAR MI CARTOON"
        },
        'italiano': {
            'saudacao': "👋 Creiamo il tuo *nuovo cartoon*?",
            'botao': "CREA IL MIO CARTOON"
        },
        'alemao': {
            'saudacao': "👋 Lassen Sie uns Ihren *neuen Cartoon erstellen*?",
            'botao': "MEINEN CARTOON ERSTELLEN"
        },
        'frances': {
            'saudacao': "👋 Créons votre *nouveau cartoon* ?",
            'botao': "CRÉER MON CARTOON"
        }
    }
    
    textos = textos_inicio.get(idioma, textos_inicio['portugues'])
    
    # MOSTRAR MENU INICIAL (MESMO DO /start)
    texto = textos['saudacao']
    keyboard = [[InlineKeyboardButton(textos['botao'], callback_data="mycartoon")]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    try:
        # Tentar editar a mensagem atual
        await query.edit_message_text(
            text=texto,
            reply_markup=reply_markup,
            parse_mode="Markdown"
        )
        print(f"✅ Menu inicial editado via botão nova encomenda | Idioma: {idioma}")
        
    except Exception as e:
        print(f"⚠️ Erro ao editar mensagem, enviando nova: {e}")
        # Se não conseguir editar, enviar nova mensagem
        await context.bot.send_message(
            chat_id=chat_id,
            text=texto,
            reply_markup=reply_markup,
            parse_mode="Markdown"
        )
        print(f"✅ Nova mensagem enviada via botão nova encomenda | Idioma: {idioma}")


















FOTO_PROBLEMA = "foto_problema"  # ⬅️ ADICIONA AQUI com as outras strings
AGUARDANDO_REPORTE_PROBLEMA = "aguardando_reporte_problema"
AGUARDANDO_ID_PEDIDO = "aguardando_id_pedido"
AGUARDANDO_SCREENSHOT_CARTOON = "aguardando_screenshot_cartoon"
DESCRICAO = "descricao"
CORRECOES = "correcoes"  






async def help_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para o comando /help com opções de suporte - PRIMEIRO ESCOLHER IDIOMA"""
    
    chat_id = update.message.chat_id if update.message else update.callback_query.message.chat_id
    user_id = update.effective_user.id
    
    # 🔥 CANCELAR TEMPORIZADORES
    try:
        cancelar_temporizador_30min(chat_id)
    except:
        pass
    
    # 🔥 LIMPAR DADOS DE AJUDA SE NECESSÁRIO
    if 'conversation_state' in context.user_data:
        del context.user_data['conversation_state']
    
    print(f"🔧 /help chamado para chat {chat_id} por user {user_id}")
    
    # 🔥 🔥 🔥 **PRIMEIRO: PEDIR PARA ESCOLHER IDIOMA PARA AJUDA**
    texto_escolha_idioma = "🌍 *Please choose your language for help / Por favor escolha seu idioma para ajuda:*"
    
    keyboard = [
        [
            InlineKeyboardButton("🇵🇹 Português", callback_data="help_idioma_portugues"),
            InlineKeyboardButton("🇺🇸 English", callback_data="help_idioma_ingles")
        ],
        [
            InlineKeyboardButton("🇪🇸 Español", callback_data="help_idioma_espanhol"),
            InlineKeyboardButton("🇮🇹 Italiano", callback_data="help_idioma_italiano")
        ],
        [
            InlineKeyboardButton("🇩🇪 Deutsch", callback_data="help_idioma_alemao"),
            InlineKeyboardButton("🇫🇷 Français", callback_data="help_idioma_frances")
        ]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    if update.message:
        await update.message.reply_text(
            texto_escolha_idioma,
            reply_markup=reply_markup,
            parse_mode="Markdown"
        )
    else:
        query = update.callback_query
        await query.answer()
        await query.edit_message_text(
            texto_escolha_idioma,
            parse_mode="Markdown",
            reply_markup=reply_markup
        )
    
    print(f"✅ Tela de escolha de idioma para ajuda mostrada para chat {chat_id}")


async def help_selecionar_idioma(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para seleção de idioma no help"""
    query = update.callback_query
    await query.answer()
    
    # Extrair idioma do callback_data
    idioma = query.data.replace("help_idioma_", "")
    
    # Salvar idioma no user_data (sobrescreve se já existir)
    context.user_data['idioma'] = idioma
    
    print(f"✅ Idioma selecionado para ajuda: {idioma} por user {query.from_user.id}")
    
    # Mensagens de confirmação em cada idioma
    mensagens_confirmacao = {
        'portugues': "✅ *Idioma de ajuda definido para Português!*",
        'ingles': "✅ *Help language set to English!*",
        'espanhol': "✅ *¡Idioma de ayuda establecido en Español!*",
        'italiano': "✅ *Lingua di aiuto impostata su Italiano!*",
        'alemao': "✅ *Hilfesprache auf Deutsch eingestellt!*",
        'frances': "✅ *Langue d'aide définie sur Français!*"
    }
    
    # Mostrar confirmação
    await query.edit_message_text(
        text=mensagens_confirmacao.get(idioma, "✅ Idioma de ajuda selecionado!"),
        parse_mode="Markdown"
    )
    
    # Aguardar 1 segundo antes de mostrar o menu de ajuda
    await asyncio.sleep(1)
    
    # 🔥 AGORA MOSTRAR O MENU DE AJUDA NO IDIOMA ESCOLHIDO
    await mostrar_menu_ajuda(update, context, idioma)


async def mostrar_menu_ajuda(update: Update, context: ContextTypes.DEFAULT_TYPE, idioma=None):
    """Mostra o menu de ajuda no idioma selecionado"""
    
    # Se não veio com idioma, pegar do user_data
    if idioma is None:
        idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 TEXTOS DO MENU DE AJUDA POR IDIOMA
    textos_ajuda = {
        'portugues': {
            'titulo': "🆘 *CENTRO DE AJUDA GODSPLAN*",
            'instrucao': "Escolha uma das opções abaixo para obter assistência:",
            'encomenda': "📦 A minha encomenda?",
            'problema': "❌ Identificaste um problema?",
            'tempo': "⏰ Quanto tempo demora?",
            'voltar': "↩️ Voltar ao Menu"
        },
        'ingles': {
            'titulo': "🆘 *GODSPLAN HELP CENTER*",
            'instrucao': "Choose one of the options below to get assistance:",
            'encomenda': "📦 My order?",
            'problema': "❌ Found a problem?",
            'tempo': "⏰ How long does it take?",
            'voltar': "↩️ Back to Menu"
        },
        'espanhol': {
            'titulo': "🆘 *CENTRO DE AYUDA GODSPLAN*",
            'instrucao': "Elija una de las opciones siguientes para obtener asistencia:",
            'encomenda': "📦 ¿Mi pedido?",
            'problema': "❌ ¿Identificaste un problema?",
            'tempo': "⏰ ¿Cuánto tiempo tarda?",
            'voltar': "↩️ Volver al Menú"
        },
        'italiano': {
            'titulo': "🆘 *CENTRO DI AIUTO GODSPLAN*",
            'instrucao': "Scegli una delle opzioni seguenti per ottenere assistenza:",
            'encomenda': "📦 Il mio ordine?",
            'problema': "❌ Hai identificato un problema?",
            'tempo': "⏰ Quanto tempo ci vuole?",
            'voltar': "↩️ Torna al Menu"
        },
        'alemao': {
            'titulo': "🆘 *GODSPLAN HILFECENTER*",
            'instrucao': "Wählen Sie eine der folgenden Optionen, um Hilfe zu erhalten:",
            'encomenda': "📦 Meine Bestellung?",
            'problema': "❌ Ein Problem festgestellt?",
            'tempo': "⏰ Wie lange dauert es?",
            'voltar': "↩️ Zurück zum Menü"
        },
        'frances': {
            'titulo': "🆘 *CENTRE D'AIDE GODSPLAN*",
            'instrucao': "Choisissez l'une des options ci-dessous pour obtenir de l'aide:",
            'encomenda': "📦 Ma commande?",
            'problema': "❌ Vous avez identifié un problème?",
            'tempo': "⏰ Combien de temps cela prend-il?",
            'voltar': "↩️ Retour au Menu"
        }
    }
    
    textos = textos_ajuda.get(idioma, textos_ajuda['portugues'])
    
    texto = f"{textos['titulo']}\n\n{textos['instrucao']}"
    
    teclado = [
        [InlineKeyboardButton(textos['encomenda'], callback_data="help_encomenda")],
        [InlineKeyboardButton(textos['problema'], callback_data="help_problema")],
        [InlineKeyboardButton(textos['tempo'], callback_data="help_tempo")],
        [InlineKeyboardButton(textos['voltar'], callback_data="voltar_menu")]
    ]
    
    try:
        query = update.callback_query
        await query.answer()
        await query.edit_message_text(
            texto,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
    except:
        # Se não for um callback, enviar nova mensagem
        chat_id = update.message.chat_id if update.message else update.effective_chat.id
        await context.bot.send_message(
            chat_id=chat_id,
            text=texto,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(teclado)
        )
    
    print(f"✅ Menu de ajuda mostrado no idioma: {idioma}")


# --- Handler para opção "A minha encomenda?" ---
async def help_encomenda(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para consulta de encomenda - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 TEXTOS PARA CONSULTA DE ENCOMENDA POR IDIOMA
    textos_encomenda = {
        'portugues': {
            'titulo': "📦 *CONSULTA DE ENCOMENDA*",
            'instrucao': "Por favor, digite o *ID do seu pedido* que recebeu na confirmação da encomenda.",
            'exemplo': "*Exemplo:* `A1B2C3D4`",
            'final': "Iremos verificar o status e entraremos em contacto consigo!"
        },
        'ingles': {
            'titulo': "📦 *ORDER CONSULTATION*",
            'instrucao': "Please enter the *order ID* you received in the order confirmation.",
            'exemplo': "*Example:* `A1B2C3D4`",
            'final': "We will check the status and contact you!"
        },
        'espanhol': {
            'titulo': "📦 *CONSULTA DE PEDIDO*",
            'instrucao': "Por favor, introduzca el *ID de su pedido* que recibió en la confirmación del pedido.",
            'exemplo': "*Ejemplo:* `A1B2C3D4`",
            'final': "¡Verificaremos el estado y nos pondremos en contacto con usted!"
        },
        'italiano': {
            'titulo': "📦 *CONSULTA ORDINE*",
            'instrucao': "Per favore, inserisci l'*ID del tuo ordine* che hai ricevuto nella conferma dell'ordine.",
            'exemplo': "*Esempio:* `A1B2C3D4`",
            'final': "Controlleremo lo stato e ti contatteremo!"
        },
        'alemao': {
            'titulo': "📦 *BESTELLANFRAGE*",
            'instrucao': "Bitte geben Sie die *Bestell-ID* ein, die Sie in der Bestellbestätigung erhalten haben.",
            'exemplo': "*Beispiel:* `A1B2C3D4`",
            'final': "Wir werden den Status überprüfen und Sie kontaktieren!"
        },
        'frances': {
            'titulo': "📦 *CONSULTATION DE COMMANDE*",
            'instrucao': "Veuillez saisir l'*ID de votre commande* que vous avez reçu dans la confirmation de commande.",
            'exemplo': "*Exemple:* `A1B2C3D4`",
            'final': "Nous vérifierons le statut et vous contacterons!"
        }
    }
    
    textos = textos_encomenda.get(idioma, textos_encomenda['portugues'])
    
    texto = f"""{textos['titulo']}

{textos['instrucao']}

{textos['exemplo']}

{textos['final']}"""
    
    await query.edit_message_text(
        texto,
        parse_mode="Markdown"
    )
    
    context.user_data['conversation_state'] = AGUARDANDO_ID_PEDIDO
    print(f"✅ Estado definido para AGUARDANDO_ID_PEDIDO | Idioma: {idioma}")





    

# --- Handler para opção "Identificaste um problema?" ---
async def help_problema(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para reportar problema - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 TEXTOS PARA REPORTAR PROBLEMA POR IDIOMA (COM FOTO)
    textos_problema = {
        'portugues': {
            'titulo': "❌ *REPORTAR PROBLEMA*",
            'instrucao': "Por favor, descreva o problema que identificou ou envie um screenshot.",
            'final': "*A nossa equipa técnica irá resolver rapidamente!*"
        },
        'ingles': {
            'titulo': "❌ *REPORT PROBLEM*",
            'instrucao': "Please describe the problem you identified or send a screenshot.",
            'final': "*Our technical team will resolve it quickly!*"
        },
        'espanhol': {
            'titulo': "❌ *INFORMAR PROBLEMA*",
            'instrucao': "Por favor, describa el problema que identificó o envíe una captura de pantalla.",
            'final': "*¡Nuestro equipo técnico lo resolverá rápidamente!*"
        },
        'italiano': {
            'titulo': "❌ *SEGNALARE PROBLEMA*",
            'instrucao': "Per favore, descrivi il problema che hai identificato o invia uno screenshot.",
            'final': "*Il nostro team tecnico lo risolverà rapidamente!*"
        },
        'alemao': {
            'titulo': "❌ *PROBLEM MELDEN*",
            'instrucao': "Bitte beschreiben Sie das Problem, das Sie festgestellt haben, oder senden Sie einen Screenshot.",
            'final': "*Unser Technikteam wird es schnell lösen!*"
        },
        'frances': {
            'titulo': "❌ *SIGNALER UN PROBLÈME*",
            'instrucao': "Veuillez décrire le problème que vous avez identifié ou envoyer une capture d'écran.",
            'final': "*Notre équipe technique le résoudra rapidement !*"
        }
    }
    
    textos = textos_problema.get(idioma, textos_problema['portugues'])
    
    texto = f"""{textos['titulo']}

{textos['instrucao']}

{textos['final']}"""
    
    # 🔥 BOTÃO VOLTAR TRADUZIDO
    textos_botao = {
        'portugues': "↩️ Voltar",
        'ingles': "↩️ Back",
        'espanhol': "↩️ Volver",
        'italiano': "↩️ Indietro",
        'alemao': "↩️ Zurück",
        'frances': "↩️ Retour"
    }
    
    teclado = [
        [InlineKeyboardButton(textos_botao.get(idioma, "↩️ Voltar"), callback_data="help_voltar")]
    ]
    
    await query.edit_message_text(
        texto,
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(teclado)
    )
    
    # ⬅️ USA O NOVO ESTADO FOTO_PROBLEMA
    context.user_data['conversation_state'] = FOTO_PROBLEMA
    print(f"✅ Estado definido para FOTO_PROBLEMA | Idioma: {idioma}")



# --- Handler para opção "Quanto tempo demora?" ---
async def help_tempo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para informação sobre tempo de produção - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    
    # 🔥 TEXTOS SOBRE TEMPO DE PRODUÇÃO POR IDIOMA
    textos_tempo = {
        'portugues': {
            'titulo': "⏰ *TEMPO DE PRODUÇÃO*",
            'prazo': "⏳ *Por norma, a produção demora entre 2 a 4 semanas*, mas sempre estimamos um prazo de 2 semanas dependendo da demanda atual.",
            'porque': "🎨 *Porquê este tempo?*",
            'item1': "• Cada cartoon 3D é único e personalizado",
            'item2': "• Trabalho artesanal de profissionais especializados",
            'item3': "• Processo de qualidade rigoroso",
            'item4': "• Acabamento premium em cada peça",
            'tranquilo': "🚀 *Mas fique tranquilo(a)!*",
            'trabalho': "Estamos a trabalhar para que o seu *Cartoon 3D Premium* chegue até si o mais breve possível!",
            'qualidade': "✨ *A qualidade vale a espera!*"
        },
        'ingles': {
            'titulo': "⏰ *PRODUCTION TIME*",
            'prazo': "⏳ *Normally, production takes between 2 to 4 weeks*, but we always estimate a 2-week deadline depending on current demand.",
            'porque': "🎨 *Why this time?*",
            'item1': "• Each 3D cartoon is unique and personalized",
            'item2': "• Handcrafted work by specialized professionals",
            'item3': "• Rigorous quality process",
            'item4': "• Premium finishing on each piece",
            'tranquilo': "🚀 *But don't worry!*",
            'trabalho': "We are working so that your *Premium 3D Cartoon* reaches you as soon as possible!",
            'qualidade': "✨ *Quality is worth the wait!*"
        },
        'espanhol': {
            'titulo': "⏰ *TIEMPO DE PRODUCCIÓN*",
            'prazo': "⏳ *Normalmente, la producción tarda entre 2 y 4 semanas*, pero siempre estimamos un plazo de 2 semanas dependiendo de la demanda actual.",
            'porque': "🎨 *¿Por qué este tiempo?*",
            'item1': "• Cada caricatura 3D es única y personalizada",
            'item2': "• Trabajo artesanal de profesionales especializados",
            'item3': "• Proceso de calidad riguroso",
            'item4': "• Acabado premium en cada pieza",
            'tranquilo': "🚀 *¡Pero quédese tranquilo(a)!*",
            'trabalho': "¡Estamos trabajando para que su *Caricatura 3D Premium* llegue a usted lo antes posible!",
            'qualidade': "✨ *¡La calidad vale la espera!*"
        },
        'italiano': {
            'titulo': "⏰ *TEMPO DI PRODUZIONE*",
            'prazo': "⏳ *Normalmente, la produzione richiede da 2 a 4 settimane*, ma stimiamo sempre una scadenza di 2 settimane a seconda della domanda attuale.",
            'porque': "🎨 *Perché questo tempo?*",
            'item1': "• Ogni cartoon 3D è unico e personalizzato",
            'item2': "• Lavoro artigianale di professionisti specializzati",
            'item3': "• Processo di qualità rigoroso",
            'item4': "• Finitura premium su ogni pezzo",
            'tranquilo': "🚀 *Ma stia tranquillo(a)!*",
            'trabalho': "Stiamo lavorando affinché il suo *Cartoon 3D Premium* arrivi da lei il prima possibile!",
            'qualidade': "✨ *La qualità vale l'attesa!*"
        },
        'alemao': {
            'titulo': "⏰ *PRODUKTIONSZEIT*",
            'prazo': "⏳ *Normalerweise dauert die Produktion zwischen 2 und 4 Wochen*, aber wir schätzen je nach aktueller Nachfrage immer eine Frist von 2 Wochen.",
            'porque': "🎨 *Warum diese Zeit?*",
            'item1': "• Jeder 3D-Cartoon ist einzigartig und personalisiert",
            'item2': "• Handgefertigte Arbeit von spezialisierten Fachleuten",
            'item3': "• Strenger Qualitätsprozess",
            'item4': "• Premium-Finish an jedem Stück",
            'tranquilo': "🚀 *Aber seien Sie unbesorgt!*",
            'trabalho': "Wir arbeiten daran, dass Ihr *Premium 3D-Cartoon* so schnell wie möglich bei Ihnen ankommt!",
            'qualidade': "✨ *Qualität ist das Warten wert!*"
        },
        'frances': {
            'titulo': "⏰ *TEMPS DE PRODUCTION*",
            'prazo': "⏳ *Normalement, la production prend entre 2 et 4 semaines*, mais nous estimons toujours un délai de 2 semaines selon la demande actuelle.",
            'porque': "🎨 *Pourquoi ce temps ?*",
            'item1': "• Chaque dessin animé 3D est unique et personnalisé",
            'item2': "• Travail artisanal de professionnels spécialisés",
            'item3': "• Processus de qualité rigoureux",
            'item4': "• Finition premium sur chaque pièce",
            'tranquilo': "🚀 *Mais soyez tranquille !*",
            'trabalho': "Nous travaillons pour que votre *Dessin Animé 3D Premium* vous parvienne le plus rapidement possible !",
            'qualidade': "✨ *La qualité vaut l'attente !*"
        }
    }
    
    textos = textos_tempo.get(idioma, textos_tempo['portugues'])
    
    texto = f"""{textos['titulo']}

{textos['prazo']}

{textos['porque']}
{textos['item1']}
{textos['item2']}
{textos['item3']}
{textos['item4']}

{textos['tranquilo']}
{textos['trabalho']}

{textos['qualidade']}"""
    
    # 🔥 BOTÃO VOLTAR TRADUZIDO
    textos_botao = {
        'portugues': "↩️ Voltar",
        'ingles': "↩️ Back",
        'espanhol': "↩️ Volver",
        'italiano': "↩️ Indietro",
        'alemao': "↩️ Zurück",
        'frances': "↩️ Retour"
    }
    
    teclado = [
        [InlineKeyboardButton(textos_botao.get(idioma, "↩️ Voltar"), callback_data="help_voltar")]
    ]

    await query.edit_message_text(
        texto,
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(teclado)
    )
    
    print(f"✅ Informações sobre tempo de produção mostradas | Idioma: {idioma}")







# --- Handler para receber ID do pedido ---
async def receber_id_pedido(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber ID do pedido para consulta - COM TRADUÇÃO"""
    try:
        id_pedido = update.message.text.strip().upper()
        
        # 🔥 PEGAR IDIOMA DO USER_DATA
        idioma = context.user_data.get('idioma', 'portugues')
        
        # Validar formato do ID (8 caracteres alfanuméricos)
        if len(id_pedido) == 8 and id_pedido.isalnum():
            print(f"✅ ID de pedido recebido: {id_pedido} | Idioma: {idioma}")
            
            # 🔥 ENVIAR PARA O TEU CHAT ID PESSOAL COM BOTÃO DE CONTACTO
            mensagem_suporte = f"""
🆘 *NOVA CONSULTA DE ENCOMENDA*

👤 *Cliente:* {update.message.from_user.first_name} (@{update.message.from_user.username or 'N/A'})
🆔 *ID do Pedido:* `{id_pedido}`
💬 *Chat ID do Cliente:* {update.message.chat_id}
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
🌐 *Idioma:* {idioma}

*Por favor, verificar status e contactar o cliente.*"""
            
            CHAT_SUPORTE_GERAL = os.getenv("CHAT_SUPORTE_GERAL")
            # Validação obrigatória (SEM MOSTRAR O ID)
            if not CHAT_SUPORTE_GERAL:
                print("⚠️ AVISO: CHAT_SUPORTE_GERAL não configurado")
                await update.message.reply_text("❌ Erro interno. Por favor, tente mais tarde.")
                return  # ⬅️ SAI DA FUNÇÃO, MAS O BOT CONTINUA!

            try:
                CHAT_SUPORTE_GERAL = int(CHAT_SUPORTE_GERAL)
                print("✅ Canal de suporte geral: CONFIGURADO")
            except ValueError:
                print("⚠️ AVISO: CHAT_SUPORTE_GERAL inválido")
                await update.message.reply_text("❌ Erro interno. Por favor, tente mais tarde.")
                return  # ⬅️ SAI DA FUNÇÃO, MAS O BOT CONTINUA!
            
            # 🔥 BOTÃO PARA CONTACTAR CLIENTE (sempre em português para a equipa)
            keyboard_suporte = [
                [
                    InlineKeyboardButton("📞 Contactar Cliente", 
                                       url=f"tg://user?id={update.message.chat_id}")
                ]
            ]
            
            try:
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_GERAL,
                    text=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
                print(f"✅ Mensagem de consulta enviada para o suporte | Idioma cliente: {idioma}")
                
            except Exception as e:
                print(f"❌ Erro ao enviar para suporte: {e}")
                # Fallback - guardar em log
                with open("suporte_log.txt", "a", encoding="utf-8") as f:
                    f.write(f"\n{datetime.now()}: {mensagem_suporte}\n")
                    f.write(f"ERRO: {str(e)}\n")
            
            # 🔥 TEXTOS DE CONFIRMAÇÃO AO CLIENTE POR IDIOMA
            textos_confirmacao = {
                'portugues': {
                    'agradecimento': "✅ *Obrigado!* Recebemos o seu pedido de consulta.",
                    'id_pedido': f"🆔 *ID do Pedido:* `{id_pedido}`",
                    'contacto': "📞 *Iremos verificar e contactá-lo(a) brevemente!*",
                    'mais_ajuda': "*Se precisar de mais ajuda:*",
                    'start': "👉 /start - Para criar nova encomenda",
                    'help': "👉 /help - Para ver opções de ajuda",
                    'agradecimento_final': "_A equipa GodsPlan agradece a sua paciência._"
                },
                'ingles': {
                    'agradecimento': "✅ *Thank you!* We received your consultation request.",
                    'id_pedido': f"🆔 *Order ID:* `{id_pedido}`",
                    'contacto': "📞 *We will check and contact you shortly!*",
                    'mais_ajuda': "*If you need more help:*",
                    'start': "👉 /start - To create a new order",
                    'help': "👉 /help - To see help options",
                    'agradecimento_final': "_The GodsPlan team thanks you for your patience._"
                },
                'espanhol': {
                    'agradecimento': "✅ *¡Gracias!* Recibimos su solicitud de consulta.",
                    'id_pedido': f"🆔 *ID del Pedido:* `{id_pedido}`",
                    'contacto': "📞 *¡Verificaremos y lo contactaremos pronto!*",
                    'mais_ajuda': "*Si necesita más ayuda:*",
                    'start': "👉 /start - Para crear un nuevo pedido",
                    'help': "👉 /help - Para ver opciones de ayuda",
                    'agradecimento_final': "_El equipo GodsPlan agradece su paciencia._"
                },
                'italiano': {
                    'agradecimento': "✅ *Grazie!* Abbiamo ricevuto la tua richiesta di consultazione.",
                    'id_pedido': f"🆔 *ID Ordine:* `{id_pedido}`",
                    'contacto': "📞 *Controlleremo e ti contatteremo a breve!*",
                    'mais_ajuda': "*Se hai bisogno di più aiuto:*",
                    'start': "👉 /start - Per creare un nuovo ordine",
                    'help': "👉 /help - Per vedere le opzioni di aiuto",
                    'agradecimento_final': "_Il team GodsPlan ringrazia per la tua pazienza._"
                },
                'alemao': {
                    'agradecimento': "✅ *Danke!* Wir haben Ihre Beratungsanfrage erhalten.",
                    'id_pedido': f"🆔 *Bestell-ID:* `{id_pedido}`",
                    'contacto': "📞 *Wir werden prüfen und Sie bald kontaktieren!*",
                    'mais_ajuda': "*Wenn Sie mehr Hilfe benötigen:*",
                    'start': "👉 /start - Um eine neue Bestellung zu erstellen",
                    'help': "👉 /help - Um Hilfsoptionen zu sehen",
                    'agradecimento_final': "_Das GodsPlan-Team dankt Ihnen für Ihre Geduld._"
                },
                'frances': {
                    'agradecimento': "✅ *Merci !* Nous avons reçu votre demande de consultation.",
                    'id_pedido': f"🆔 *ID de Commande:* `{id_pedido}`",
                    'contacto': "📞 *Nous vérifierons et vous contacterons bientôt !*",
                    'mais_ajuda': "*Si vous avez besoin de plus d'aide :*",
                    'start': "👉 /start - Pour créer une nouvelle commande",
                    'help': "👉 /help - Pour voir les options d'aide",
                    'agradecimento_final': "_L'équipe GodsPlan vous remercie de votre patience._"
                }
            }
            
            textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
            
            # 🔥 CONFIRMAR AO CLIENTE COM TEXTO CLICÁVEL
            await update.message.reply_text(
                f"{textos['agradecimento']}\n\n"
                f"{textos['id_pedido']}\n"
                f"{textos['contacto']}\n\n"
                f"{textos['mais_ajuda']}\n"
                f"{textos['start']}\n"
                f"{textos['help']}\n\n"
                f"{textos['agradecimento_final']}",
                parse_mode="Markdown"
            )
            
            # Limpar estado
            context.user_data['conversation_state'] = None
            
        else:
            # 🔥 TEXTOS DE ERRO POR IDIOMA
            textos_erro = {
                'portugues': {
                    'titulo': "❌ *ID inválido!*",
                    'instrucao': "Por favor, digite um *ID de pedido válido* (8 caracteres alfanuméricos).",
                    'exemplo': "*Exemplo:* `A1B2C3D4`",
                    'info': "O ID foi fornecido na confirmação da sua encomenda."
                },
                'ingles': {
                    'titulo': "❌ *Invalid ID!*",
                    'instrucao': "Please enter a *valid order ID* (8 alphanumeric characters).",
                    'exemplo': "*Example:* `A1B2C3D4`",
                    'info': "The ID was provided in your order confirmation."
                },
                'espanhol': {
                    'titulo': "❌ *¡ID inválido!*",
                    'instrucao': "Por favor, introduzca un *ID de pedido válido* (8 caracteres alfanuméricos).",
                    'exemplo': "*Ejemplo:* `A1B2C3D4`",
                    'info': "El ID se proporcionó en la confirmación de su pedido."
                },
                'italiano': {
                    'titulo': "❌ *ID non valido!*",
                    'instrucao': "Per favore, inserisci un *ID ordine valido* (8 caratteri alfanumerici).",
                    'exemplo': "*Esempio:* `A1B2C3D4`",
                    'info': "L'ID è stato fornito nella conferma del tuo ordine."
                },
                'alemao': {
                    'titulo': "❌ *Ungültige ID!*",
                    'instrucao': "Bitte geben Sie eine *gültige Bestell-ID* ein (8 alphanumerische Zeichen).",
                    'exemplo': "*Beispiel:* `A1B2C3D4`",
                    'info': "Die ID wurde in Ihrer Bestellbestätigung bereitgestellt."
                },
                'frances': {
                    'titulo': "❌ *ID invalide !*",
                    'instrucao': "Veuillez saisir un *ID de commande valide* (8 caractères alphanumériques).",
                    'exemplo': "*Exemple:* `A1B2C3D4`",
                    'info': "L'ID a été fourni dans la confirmation de votre commande."
                }
            }
            
            textos = textos_erro.get(idioma, textos_erro['portugues'])
            
            await update.message.reply_text(
                f"{textos['titulo']}\n\n"
                f"{textos['instrucao']}\n"
                f"{textos['exemplo']}\n\n"
                f"{textos['info']}",
                parse_mode="Markdown"
            )
            
    except Exception as e:
        print(f"ERRO em receber_id_pedido: {e}")
        
        # 🔥 MENSAGEM DE ERRO GENÉRICA POR IDIOMA
        textos_erro_generico = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, intente de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro_generico.get(idioma, textos_erro_generico['portugues'])
        )

      







# --- Handler para receber problema ---
# --- Handler para receber problema ---
# --- Handler para receber problema ---
async def receber_problema(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber problema - com texto clicável e tradução completa"""
    try:
        print("🎯 receber_problema CHAMADO")
        
        CHAT_SUPORTE_CLIENTES = os.getenv("CHAT_SUPORTE_CLIENTES") 
        
        # ✅ ADICIONA VALIDAÇÃO:
        if not CHAT_SUPORTE_CLIENTES:
            print("⚠️ AVISO: CHAT_SUPORTE não configurado")
            await update.message.reply_text("❌ Erro interno. Por favor, tente mais tarde.")
            return
        
        try:
            CHAT_SUPORTE_CLIENTES = int(CHAT_SUPORTE_CLIENTES)
        except ValueError:
            print("⚠️ AVISO: CHAT_SUPORTE inválido")
            await update.message.reply_text("❌ Erro interno. Por favor, tente mais tarde.")
            return


        user = update.message.from_user
        chat_id = update.message.chat_id
        
        # 🔥 PEGAR IDIOMA DO USER_DATA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"📋 Idioma do cliente: {idioma}")
        
        # Verificar se é foto ou texto
        if update.message.photo:
            # É uma foto/screenshot
            print("📸 É uma FOTO")
            photo = update.message.photo[-1]
            file_id = photo.file_id
            
            legenda = update.message.caption if update.message.caption else "Sem descrição adicional"
            
            # 🔥 TEXTO PARA SUPORTE (sempre em português para a equipa)
            mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO - COM FOTO*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
🌐 *Idioma:* {idioma}

📝 *Descrição:*
{legenda}

*Cliente enviou uma foto/screenshot do problema.*"""
            
            # BOTÕES PARA SUPORTE (sempre em português)
            keyboard_suporte = [
                [
                    InlineKeyboardButton("📞 Contactar Cliente", 
                                       url=f"tg://user?id={chat_id}")
                ]
            ]
            
            try:
                await context.bot.send_photo(
                    chat_id=CHAT_SUPORTE_CLIENTES,
                    photo=file_id,
                    caption=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
                print(f"✅ Foto enviada para suporte | Idioma cliente: {idioma}")
            except Exception as e:
                print(f"❌ Erro ao enviar foto: {e}")
                # ❌ LOG DE ERRO REMOVIDO (conforme solicitado)
            
        elif update.message.text:
            # É texto
            print("📝 É TEXTO")
            problema = update.message.text
            
            # 🔥 TEXTO PARA SUPORTE (sempre em português para a equipa)
            mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
🌐 *Idioma:* {idioma}

📝 *Problema:*
{problema}"""
            
            # BOTÕES PARA SUPORTE (sempre em português)
            keyboard_suporte = [
                [
                    InlineKeyboardButton("📞 Contactar Cliente", 
                                       url=f"tg://user?id={chat_id}")
                ]
            ]
            
            try:
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_CLIENTES,
                    text=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
                print(f"✅ Problema enviado para suporte | Idioma cliente: {idioma}")
            except Exception as e:
                print(f"❌ Erro ao enviar problema: {e}")
                # ❌ LOG DE ERRO REMOVIDO (conforme solicitado)
        
        else:
            # 🔥 MENSAGEM DE ERRO POR IDIOMA
            textos_erro_formato = {
                'portugues': {
                    'titulo': "❌ *Formato não suportado!*",
                    'instrucao': "Por favor, envie uma descrição em texto ou um screenshot/foto do problema."
                },
                'ingles': {
                    'titulo': "❌ *Format not supported!*",
                    'instrucao': "Please send a text description or a screenshot/photo of the problem."
                },
                'espanhol': {
                    'titulo': "❌ *¡Formato no soportado!*",
                    'instrucao': "Por favor, envíe una descripción en texto o una captura de pantalla/foto del problema."
                },
                'italiano': {
                    'titulo': "❌ *Formato non supportato!*",
                    'instrucao': "Per favore, invia una descrizione testuale o uno screenshot/foto del problema."
                },
                'alemao': {
                    'titulo': "❌ *Format nicht unterstützt!*",
                    'instrucao': "Bitte senden Sie uma Textbeschreibung oder einen Screenshot/Foto des Problems."
                },
                'frances': {
                    'titulo': "❌ *Format non pris en charge !*",
                    'instrucao': "Veuillez envoyer une description textuelle ou une capture d'écran/photo du problème."
                }
            }
            
            textos = textos_erro_formato.get(idioma, textos_erro_formato['portugues'])
            
            await update.message.reply_text(
                f"{textos['titulo']}\n\n{textos['instrucao']}",
                parse_mode="Markdown"
            )
            return
        
        # 🔥 TEXTOS DE CONFIRMAÇÃO AO CLIENTE POR IDIOMA
        textos_confirmacao = {
            'portugues': {
                'titulo': "✅ *Problema reportado com sucesso!*",
                'resolucao': "Nossa equipa técnica vai resolver o seu problema brevemente.",
                'mais_ajuda': "*Se precisar de mais ajuda, clique em:*",
                'start': "👉 /start - Para criar uma nova encomenda",
                'help': "👉 /help - Para ver opções de ajuda",
                'agradecimento': "_Obrigado pela sua paciência._"
            },
            'ingles': {
                'titulo': "✅ *Problem reported successfully!*",
                'resolucao': "Our technical team will solve your problem shortly.",
                'mais_ajuda': "*If you need more help, click on:*",
                'start': "👉 /start - To create a new order",
                'help': "👉 /help - To see help options",
                'agradecimento': "_Thank you for your patience._"
            },
            'espanhol': {
                'titulo': "✅ *¡Problema informado con éxito!*",
                'resolucao': "Nuestro equipo técnico resolverá su problema en breve.",
                'mais_ajuda': "*Si necesita más ayuda, haga clic en:*",
                'start': "👉 /start - Para crear un nuevo pedido",
                'help': "👉 /help - Para ver opciones de ayuda",
                'agradecimento': "_Gracias por su paciencia._"
            },
            'italiano': {
                'titulo': "✅ *Problema segnalato con successo!*",
                'resolucao': "Il nostro team tecnico risolverà il tuo problema a breve.",
                'mais_ajuda': "*Se hai bisogno di più aiuto, clicca su:*",
                'start': "👉 /start - Per creare un nuovo ordine",
                'help': "👉 /help - Per vedere le opzioni di aiuto",
                'agradecimento': "_Grazie per la tua pazienza._"
            },
            'alemao': {
                'titulo': "✅ *Problem erfolgreich gemeldet!*",
                'resolucao': "Unser Technikteam wird Ihr Problem in Kürze lösen.",
                'mais_ajuda': "*Wenn Sie mehr Hilfe benötigen, klicken Sie auf:*",
                'start': "👉 /start - Um eine neue Bestellung zu erstellen",
                'help': "👉 /help - Um Hilfsoptionen zu sehen",
                'agradecimento': "_Danke für Ihre Geduld._"
            },
            'frances': {
                'titulo': "✅ *Problème signalé avec succès !*",
                'resolucao': "Notre équipe technique résoudra votre problème sous peu.",
                'mais_ajuda': "*Si vous avez besoin de plus d'aide, cliquez sur :*",
                'start': "👉 /start - Pour créer une nouvelle commande",
                'help': "👉 /help - Pour voir les options d'aide",
                'agradecimento': "_Merci pour votre patience._"
            }
        }
        
        textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
        
        # 🔥 MENSAGEM COM TEXTO CLICÁVEL - SEM BOTÕES
        mensagem_cliente = f"""{textos['titulo']}

{textos['resolucao']}

{textos['mais_ajuda']}
{textos['start']}
{textos['help']}

{textos['agradecimento']}"""

        await update.message.reply_text(
            mensagem_cliente,
            parse_mode="Markdown"
        )
        
        # Limpar estado
        context.user_data['conversation_state'] = None
        print(f"✅ Estado limpo | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_problema: {e}")
        
        # 🔥 MENSAGEM DE ERRO GENÉRICA POR IDIOMA
        textos_erro_generico = {
            'portugues': "❌ Ocorreu um erro. Por favor, tente novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, intente de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro_generico.get(idioma, textos_erro_generico['portugues'])
        )



# --- Handler para reportar problema ---
async def receber_reportar_problema(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber problema reportado - com texto clicável e tradução completa"""
    try:
        print("🎯 receber_reportar_problema CHAMADO")
        
        CHAT_SUPORTE_CLIENTES = os.getenv("CHAT_SUPORTE_CLIENTES") 
        
        # ✅ ADICIONA VALIDAÇÃO:
        if not CHAT_SUPORTE_CLIENTES:
            print("⚠️ AVISO: CHAT_SUPORTE não configurado")
            await update.message.reply_text("❌ Erro interno. Por favor, tente mais tarde.")
            return
        
        try:
            CHAT_SUPORTE_CLIENTES = int(CHAT_SUPORTE_CLIENTES)
        except ValueError:
            print("⚠️ AVISO: CHAT_SUPORTE inválido")
            await update.message.reply_text("❌ Erro interno. Por favor, tente mais tarde.")
            return


        user = update.message.from_user
        chat_id = update.message.chat_id
        
        # 🔥 PEGAR IDIOMA DO USER_DATA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"📋 Idioma do cliente: {idioma}")
        
        # 🔥 VARIÁVEL PARA ARMAZENAR TEXTO DO PROBLEMA
        texto_problema = ""
        
        if update.message.photo:
            # Foto (com ou sem legenda)
            print("📸 Foto recebida")
            photo = update.message.photo[-1]
            file_id = photo.file_id
            
            # 🔥 CAPTURAR LEGENDA DA FOTO (se houver)
            if update.message.caption:
                texto_problema = update.message.caption
                print(f"📝 Foto COM legenda: {texto_problema[:100]}...")
            else:
                print("📸 Foto SEM legenda")
            
            # 🔥 TEXTO PARA SUPORTE (sempre em português para a equipa)
            if texto_problema:
                mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO - COM FOTO E TEXTO*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
🌐 *Idioma:* {idioma}

📝 *Texto enviado com a foto:*
{texto_problema}

*Cliente enviou uma foto ilustrativa do problema.*"""
            else:
                mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO - COM FOTO*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
🌐 *Idioma:* {idioma}

*Cliente enviou uma foto do problema.*"""

            keyboard_suporte = [
                [
                    InlineKeyboardButton("📞 Contactar Cliente", 
                                       url=f"tg://user?id={chat_id}")
                ]
            ]
            
            try:
                await context.bot.send_photo(
                    chat_id=CHAT_SUPORTE_CLIENTES,
                    photo=file_id,
                    caption=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
                print(f"✅ Foto {'com texto' if texto_problema else ''} enviada para suporte | Idioma cliente: {idioma}")
            except Exception as e:
                print(f"❌ Erro ao enviar foto: {e}")
            
        elif update.message.text:
            # Texto
            print("📝 Texto recebido")
            texto_problema = update.message.text
            
            # 🔥 TEXTO PARA SUPORTE (sempre em português para a equipa)
            mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
🌐 *Idioma:* {idioma}

📝 *Problema:*
{texto_problema}

*Contactar o cliente para resolução.*"""

            keyboard_suporte = [
                [
                    InlineKeyboardButton("📞 Contactar Cliente", 
                                       url=f"tg://user?id={chat_id}")
                ]
            ]
            
            try:
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_CLIENTES,
                    text=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
                print(f"✅ Problema enviado para suporte | Idioma cliente: {idioma}")
            except Exception as e:
                print(f"❌ Erro ao enviar problema: {e}")
        
        else:
            # 🔥 MENSAGEM DE ERRO POR IDIOMA
            textos_erro_formato = {
                'portugues': {
                    'titulo': "❌ *Formato não suportado!*",
                    'instrucao': "Envie texto ou foto do problema."
                },
                'ingles': {
                    'titulo': "❌ *Format not supported!*",
                    'instrucao': "Send text or photo of the problem."
                },
                'espanhol': {
                    'titulo': "❌ *¡Formato no soportado!*",
                    'instrucao': "Envía texto o foto del problema."
                },
                'italiano': {
                    'titulo': "❌ *Formato non supportato!*",
                    'instrucao': "Invia testo o foto del problema."
                },
                'alemao': {
                    'titulo': "❌ *Format nicht unterstützt!*",
                    'instrucao': "Senden Sie Text oder Foto des Problems."
                },
                'frances': {
                    'titulo': "❌ *Format non pris en charge !*",
                    'instrucao': "Envoyez du texto ou une photo du problema."
                }
            }
            
            textos = textos_erro_formato.get(idioma, textos_erro_formato['portugues'])
            
            await update.message.reply_text(
                f"{textos['titulo']}\n\n{textos['instrucao']}",
                parse_mode="Markdown"
            )
            return
        
        # 🔥 TEXTOS DE CONFIRMAÇÃO AO CLIENTE POR IDIOMA
        textos_confirmacao = {
            'portugues': {
                'titulo': "✅ *Problema recebido!*",
                'resolucao': "Nossa equipa técnica vai resolver o seu problema brevemente.",
                'outra_coisa': "*Se quiser fazer outra coisa, clique em:*",
                'start': "👉 /start - Para nova encomenda",
                'help': "👉 /help - Para ajuda",
                'agradecimento': "_Obrigado._"
            },
            'ingles': {
                'titulo': "✅ *Problem received!*",
                'resolucao': "Our technical team will solve your problem shortly.",
                'outra_coisa': "*If you want to do something else, click on:*",
                'start': "👉 /start - For new order",
                'help': "👉 /help - For help",
                'agradecimento': "_Thank you._"
            },
            'espanhol': {
                'titulo': "✅ *¡Problema recibido!*",
                'resolucao': "Nuestro equipo técnico resolverá su problema en breve.",
                'outra_coisa': "*Si queres fazer outra cosa, haz clic en:*",
                'start': "👉 /start - Para nuevo pedido",
                'help': "👉 /help - Para ayuda",
                'agradecimento': "_Gracias._"
            },
            'italiano': {
                'titulo': "✅ *Problema ricevuto!*",
                'resolucao': "Il nostro team tecnico risolverà il tuo problema a breve.",
                'outra_coisa': "*Se vuoi fare qualcos'altro, clicca su:*",
                'start': "👉 /start - Per nuovo ordine",
                'help': "👉 /help - Per aiuto",
                'agradecimento': "_Grazie._"
            },
            'alemao': {
                'titulo': "✅ *Problem erhalten!*",
                'resolucao': "Unser Technikteam wird Ihr Problem in Kürze lösen.",
                'outra_coisa': "*Wenn Sie etwas anderes tun möchten, klicken Sie auf:*",
                'start': "👉 /start - Für neue Bestellung",
                'help': "👉 /help - Für Hilfe",
                'agradecimento': "_Danke._"
            },
            'frances': {
                'titulo': "✅ *Problème reçu !*",
                'resolucao': "Notre équipe technique résoudra votre problème sous peu.",
                'outra_coisa': "*Si vous voulez faire autre chose, cliquez sur :*",
                'start': "👉 /start - Pour nouvelle commande",
                'help': "👉 /help - Pour aide",
                'agradecimento': "_Merci._"
            }
        }
        
        textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
        
        # 🔥 MENSAGEM COM TEXTO CLICÁVEL - SEM BOTÕES
        mensagem_cliente = f"""{textos['titulo']}

{textos['resolucao']}

{textos['outra_coisa']}
{textos['start']}
{textos['help']}

{textos['agradecimento']}"""

        await update.message.reply_text(
            mensagem_cliente,
            parse_mode="Markdown"
        )
        
        context.user_data['conversation_state'] = None
        print(f"✅ Estado limpo | Idioma: {idioma}")
        
    except Exception as e:
        print(f"ERRO em receber_reportar_problema: {e}")
        
        # 🔥 MENSAGEM DE ERRO GENÉRICA POR IDIOMA
        textos_erro_generico = {
            'portugues': "❌ Erro. Por favor, tente novamente.",
            'ingles': "❌ Error. Please try again.",
            'espanhol': "❌ Error. Por favor, intente de nuevo.",
            'italiano': "❌ Errore. Per favore, riprova.",
            'alemao': "❌ Fehler. Bitte versuchen Sie es erneut.",
            'frances': "❌ Erreur. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro_generico.get(idioma, textos_erro_generico['portugues'])
        )









        

async def receber_descricao(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber descrição do cartoon - COM TRADUÇÃO"""
    try:
        print("🎯 receber_descricao CHAMADO")
        
        # 🔥 PEGAR IDIOMA DO USER_DATA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"📋 Idioma do cliente: {idioma}")
        
        if update.message.text:
            descricao = update.message.text
            context.user_data['descricao'] = descricao
            print(f"✅ Descrição recebida: {descricao[:50]}... | Idioma: {idioma}")
            
            # 🔥 TEXTOS DE CONFIRMAÇÃO POR IDIOMA
            textos_confirmacao = {
                'portugues': {
                    'titulo': "✅ *Descrição recebida!*",
                    'processando': "Agora vou processar o teu cartoon...",
                    'proxima_acao': "Por favor, aguarda enquanto crio a tua obra de arte personalizada!"
                },
                'ingles': {
                    'titulo': "✅ *Description received!*",
                    'processando': "Now I'll process your cartoon...",
                    'proxima_acao': "Please wait while I create your personalized artwork!"
                },
                'espanhol': {
                    'titulo': "✅ *¡Descripción recibida!*",
                    'processando': "Ahora procesaré tu cartoon...",
                    'proxima_acao': "¡Por favor espera mientras creo tu obra de arte personalizada!"
                },
                'italiano': {
                    'titulo': "✅ *Descrizione ricevuta!*",
                    'processando': "Ora elaborerò il tuo cartoon...",
                    'proxima_acao': "Per favore aspetta mentre creo la tua opera d'arte personalizzata!"
                },
                'alemao': {
                    'titulo': "✅ *Beschreibung erhalten!*",
                    'processando': "Jetzt verarbeite ich Ihren Cartoon...",
                    'proxima_acao': "Bitte warten Sie, während ich Ihr personalisiertes Kunstwerk erstelle!"
                },
                'frances': {
                    'titulo': "✅ *Description reçue !*",
                    'processando': "Maintenant je vais traiter votre dessin animé...",
                    'proxima_acao': "Veuillez patienter pendant que je crée votre œuvre d'art personnalisée !"
                }
            }
            
            textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
            
            await update.message.reply_text(
                f"{textos['titulo']}\n\n"
                f"{textos['processando']}\n\n"
                f"{textos['proxima_acao']}",
                parse_mode="Markdown"
            )
            
            # 🔥 AQUI PODE-SE CHAMAR A FUNÇÃO PARA CRIAR O CARTOON
            # Por exemplo: await criar_cartoon_ai(update, context, descricao)
            
            # Limpar estado
            context.user_data['conversation_state'] = None
            
    except Exception as e:
        print(f"ERRO em receber_descricao: {e}")
        
        # 🔥 MENSAGEM DE ERRO POR IDIOMA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tenta novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, intente de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )


# --- Função para processar correções ---
async def processar_correcoes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Processar correções do cartoon - COM TRADUÇÃO"""
    try:
        print("🎯 processar_correcoes CHAMADO")
        
        # 🔥 PEGAR IDIOMA DO USER_DATA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"📋 Idioma do cliente: {idioma}")
        
        if update.message.text:
            correcoes = update.message.text
            context.user_data['correcoes'] = correcoes
            print(f"✅ Correções recebidas: {correcoes[:50]}... | Idioma: {idioma}")
            
            # 🔥 TEXTOS DE CONFIRMAÇÃO POR IDIOMA
            textos_confirmacao = {
                'portugues': {
                    'titulo': "✅ *Correções recebidas!*",
                    'processando': "Vou aplicar as correções no cartoon...",
                    'agradecimento': "Obrigado pelo feedback! Vamos melhorar o cartoon para ti."
                },
                'ingles': {
                    'titulo': "✅ *Corrections received!*",
                    'processando': "I'll apply the corrections to the cartoon...",
                    'agradecimento': "Thank you for the feedback! We'll improve the cartoon for you."
                },
                'espanhol': {
                    'titulo': "✅ *¡Correcciones recibidas!*",
                    'processando': "Aplicaré las correcciones en el cartoon...",
                    'agradecimento': "¡Gracias por los comentarios! Mejoraremos el cartoon para ti."
                },
                'italiano': {
                    'titulo': "✅ *Correzioni ricevute!*",
                    'processando': "Applicherò le correzioni al cartoon...",
                    'agradecimento': "Grazie per il feedback! Miglioreremo il cartoon per te."
                },
                'alemao': {
                    'titulo': "✅ *Korrekturen erhalten!*",
                    'processando': "Ich werde die Korrekturen am Cartoon anwenden...",
                    'agradecimento': "Danke für das Feedback! Wir verbessern den Cartoon für Sie."
                },
                'frances': {
                    'titulo': "✅ *Corrections reçues !*",
                    'processando': "Je vais appliquer les corrections au dessin animé...",
                    'agradecimento': "Merci pour les commentaires ! Nous améliorerons le dessin animé pour vous."
                }
            }
            
            textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
            
            await update.message.reply_text(
                f"{textos['titulo']}\n\n"
                f"{textos['processando']}\n\n"
                f"{textos['agradecimento']}",
                parse_mode="Markdown"
            )
            
            # 🔥 AQUI PODE-SE CHAMAR A FUNÇÃO PARA APLICAR CORREÇÕES
            # Por exemplo: await aplicar_correcoes_cartoon(update, context, correcoes)
            
            # Limpar estado
            context.user_data['conversation_state'] = None
            
    except Exception as e:
        print(f"ERRO em processar_correcoes: {e}")
        
        # 🔥 MENSAGEM DE ERRO POR IDIOMA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tenta novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, intente de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )


async def receber_screenshot_cartoon(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Receber screenshot para cartoon - SEM resumo final, vai direto para descrição - COM TRADUÇÃO"""
    try:
        print("🎯 receber_screenshot_cartoon CHAMADO")
        
        # 🔥 PEGAR IDIOMA DO USER_DATA
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"📋 Idioma do cliente: {idioma}")
        
        if update.message.photo:
            # Guardar a foto no user_data
            photo = update.message.photo[-1]
            context.user_data['foto'] = photo.file_id
            context.user_data['foto_id'] = photo.file_id
            print(f"✅ Screenshot guardado para cartoon | Idioma: {idioma}")
            
            # Mudar para estado de descrição
            context.user_data['conversation_state'] = DESCRICAO
            
            # 🔥 TEXTOS PARA PEDIR DESCRIÇÃO POR IDIOMA
            textos_descricao = {
                'portugues': {
                    'titulo': "📝 *Agora descreve o que queres no cartoon:*",
                    'instrucao': "Explica o que deve aparecer na imagem, personagens, ações, etc.",
                    'exemplo': "Exemplo: \"Quero um cartoon de mim e da minha família num parque, com cachorro e sol\""
                },
                'ingles': {
                    'titulo': "📝 *Now describe what you want in the cartoon:*",
                    'instrucao': "Explain what should appear in the image, characters, actions, etc.",
                    'exemplo': "Example: \"I want a cartoon of me and my family in a park, with a dog and sun\""
                },
                'espanhol': {
                    'titulo': "📝 *Ahora describe lo que quieres en el cartoon:*",
                    'instrucao': "Explica lo que debe aparecer en la imagen, personajes, acciones, etc.",
                    'exemplo': "Ejemplo: \"Quiero un cartoon de mí y mi familia en un parque, con perro y sol\""
                },
                'italiano': {
                    'titulo': "📝 *Ora descrivi cosa vuoi nel cartoon:*",
                    'instrucao': "Spiega cosa dovrebbe apparire nell'immagine, personaggi, azioni, ecc.",
                    'exemplo': "Esempio: \"Voglio un cartoon di me e della mia famiglia in un parco, con cane e sole\""
                },
                'alemao': {
                    'titulo': "📝 *Beschreiben Sie nun, was Sie im Cartoon wollen:*",
                    'instrucao': "Erklären Sie, was im Bild erscheinen soll, Charaktere, Aktionen usw.",
                    'exemplo': "Beispiel: \"Ich möchte einen Cartoon von mir und meiner Familie in einem Park, mit Hund und Sonne\""
                },
                'frances': {
                    'titulo': "📝 *Maintenant décrivez ce que vous voulez dans le dessin animé :*",
                    'instrucao': "Expliquez ce qui doit apparaître dans l'image, personnages, actions, etc.",
                    'exemplo': "Exemple : \"Je veux un dessin animé de moi et ma famille dans un parc, avec un chien et le soleil\""
                }
            }
            
            textos = textos_descricao.get(idioma, textos_descricao['portugues'])
            
            await update.message.reply_text(
                f"{textos['titulo']}\n\n"
                f"{textos['instrucao']}\n\n"
                f"_{textos['exemplo']}_",
                parse_mode="Markdown"
            )
        else:
            # 🔥 MENSAGEM DE ERRO POR IDIOMA
            textos_erro = {
                'portugues': "❌ Por favor, envia uma screenshot para criar o cartoon.",
                'ingles': "❌ Please send a screenshot to create the cartoon.",
                'espanhol': "❌ Por favor, envía una captura de pantalla para crear el cartoon.",
                'italiano': "❌ Per favore, invia uno screenshot per creare il cartoon.",
                'alemao': "❌ Bitte senden Sie einen Screenshot, um den Cartoon zu erstellen.",
                'frances': "❌ Veuillez envoyer une capture d'écran pour créer le dessin animé."
            }
            
            await update.message.reply_text(
                textos_erro.get(idioma, textos_erro['portugues']),
                parse_mode="Markdown"
            )
            
    except Exception as e:
        print(f"ERRO em receber_screenshot_cartoon: {e}")
        
        # 🔥 MENSAGEM DE ERRO POR IDIOMA
        textos_erro = {
            'portugues': "❌ Ocorreu um erro. Por favor, tenta novamente.",
            'ingles': "❌ An error occurred. Please try again.",
            'espanhol': "❌ Ocurrió un error. Por favor, intente de nuevo.",
            'italiano': "❌ Si è verificato un errore. Per favore, riprova.",
            'alemao': "❌ Ein Fehler ist aufgetreten. Bitte versuchen Sie es erneut.",
            'frances': "❌ Une erreur s'est produite. Veuillez réessayer."
        }
        
        await update.message.reply_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )





# --- Handlers auxiliares do Help ---
async def help_mais(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Voltar para mais opções de ajuda"""
    query = update.callback_query
    await query.answer()
    await help_handler(update, context)

async def help_voltar(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Voltar ao menu principal do help"""
    query = update.callback_query
    await query.answer()
    await help_handler(update, context)

async def voltar_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Voltar ao menu inicial - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA ANTES DE LIMPAR
    idioma = context.user_data.get('idioma', 'portugues')
    
    # Limpar qualquer estado de conversação, mas manter o idioma
    context.user_data.clear()
    context.user_data['idioma'] = idioma  # 🔥 MANTER O IDIOMA SELECIONADO
    
    print(f"🔙 Voltar ao menu | Idioma mantido: {idioma}")
    
    # 🔥 TEXTOS DO MENU INICIAL POR IDIOMA
    textos_menu = {
        'portugues': {
            'saudacao': "👋 Olá! Bem-vindo à *GodsPlan*, vamos criar o seu cartoon?",
            'botao': "CREATE MY CARTOON"
        },
        'ingles': {
            'saudacao': "👋 Hello! Welcome to *GodsPlan*, shall we create your cartoon?",
            'botao': "CREATE MY CARTOON"
        },
        'espanhol': {
            'saudacao': "👋 ¡Hola! Bienvenido a *GodsPlan*, ¿vamos a crear tu cartoon?",
            'botao': "CREAR MI CARTOON"
        },
        'italiano': {
            'saudacao': "👋 Ciao! Benvenuto in *GodsPlan*, creiamo il tuo cartoon?",
            'botao': "CREA IL MIO CARTOON"
        },
        'alemao': {
            'saudacao': "👋 Hallo! Willkommen bei *GodsPlan*, sollen wir Ihren Cartoon erstellen?",
            'botao': "MEINEN CARTOON ERSTELLEN"
        },
        'frances': {
            'saudacao': "👋 Bonjour ! Bienvenue chez *GodsPlan*, allons-nous créer votre dessin animé ?",
            'botao': "CRÉER MON DESSIN ANIMÉ"
        }
    }
    
    textos = textos_menu.get(idioma, textos_menu['portugues'])
    
    texto = textos['saudacao']
    keyboard = [[InlineKeyboardButton(textos['botao'], callback_data="mycartoon")]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.edit_message_text(
        texto, 
        reply_markup=reply_markup, 
        parse_mode="Markdown"
    )
    
    print(f"✅ Menu inicial mostrado no idioma: {idioma}")












CANAL_ADMIN = os.getenv("CANAL_ADMIN")
CANAL_REQUESTS = os.getenv("CANAL_REQUESTS")



async def enviar_mensagem_automatica(context: ContextTypes.DEFAULT_TYPE):
    """Envia mensagem automaticamente para o canal (executar uma vez)"""
    try:
        keyboard = [[InlineKeyboardButton("🔐 PAINEL ADMIN", callback_data="admin_page_1")]]
        
        await context.bot.send_message(
            chat_id=CANAL_ADMIN,
            text="🛡️ *PAINEL DE CONTROLE - GODSPLAN*\n\nAcesse o painel administrativo:",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        print("✅ Mensagem automática enviada para o canal")
    except Exception as e:
        print(f"❌ Erro: {e}")




 # 🔥 ADMIN





async def admin_command(update: Update, context: ContextTypes.DEFAULT_TYPE, pagina=1):
    """Comando admin simplificado - apenas resumo financeiro e estatísticas EM EUR"""
    ADMIN_USER_IDS = os.getenv("ADMIN_USER_IDS")
    if ADMIN_USER_IDS:
        try:
           ADMIN_USER_IDS = [int(id.strip()) for id in ADMIN_USER_IDS.split(",") if id.strip()]
           print(f"✅ Administradores carregados: {len(ADMIN_USER_IDS)} usuários")
        except ValueError:
           print("❌ ERRO: ADMIN_USER_IDS contém valores não numéricos")
      
    else:
        print("⚠️ AVISO: ADMIN_USER_IDS não configurado no .env")
        
    
    
    # Verificar se é message (comando) ou callback_query (botão)
    if update.message:
        chat_id = update.message.chat_id
        message_method = update.message.reply_text
        user_id = update.effective_user.id
    elif update.callback_query:
        chat_id = update.callback_query.message.chat_id
        message_method = update.callback_query.edit_message_text
        user_id = update.callback_query.from_user.id
        await update.callback_query.answer()  # ✅ IMPORTANTE: Responder ao callback
    else:
        return
    
    print(f"🔍 ADMIN ACCESS CHECK - Chat ID: {chat_id}, User ID: {user_id}")
    print(f"🔍 ADMIN_USER_IDS: {ADMIN_USER_IDS}")
    print(f"🔍 User ID Type: {type(user_id)}, Value: {user_id}")
    
    # ✅ VERIFICAÇÃO DE ACESSO CORRIGIDA - CONVERTER PARA INT SE NECESSÁRIO
    user_id_int = int(user_id)  # 🔥 CONVERTER PARA INT PARA GARANTIR
    
    if user_id_int not in ADMIN_USER_IDS:
        print(f"❌ ACESSO NEGADO - User ID {user_id_int} não está na lista de admins")
        print(f"❌ Lista de admins: {ADMIN_USER_IDS}")
        if update.message:
            await update.message.reply_text(f"❌ Acesso negado. User ID: {user_id_int}")
        elif update.callback_query:
            await update.callback_query.answer(f"❌ Acesso negado. User ID: {user_id_int}", show_alert=True)
        return
    
    print("✅ ACESSO PERMITIDO - User é admin")
    
    # 🔥 CALCULAR ESTATÍSTICAS DE OFERTAS E PEDIDOS POR REGIÃO
    # 🔥 🔥 🔥 CORREÇÃO: USAR APENAS PEDIDOS PAGOS
    pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
    total_pedidos = len(pedidos_pagos)  # 🔥 AGORA SÓ PAGOS
    
    # 🔥 DEFINIR PAÍSES INTERNACIONAIS (Reino Unido, Estados Unidos, Brasil, Canadá) - COM CANADÁ
    paises_internacionais = ['reino unido', 'united kingdom', 'uk', 'estados unidos', 'united states', 'us', 'usa', 'brasil', 'brazil', 'canada', 'canadá']
    paises_europeus = [
        'portugal', 'espanha', 'spain', 'frança', 'france', 'franca',
        'alemanha', 'germany', 'itália', 'italia', 'bélgica', 'belgica',
        'países baixos', 'paises baixos', 'holanda', 'netherlands',
        'irlanda', 'ireland', 'luxemburgo', 'luxembourg',
        'suecia', 'sweden', 'dinamarca', 'denmark'
    ]
    
    # Calcular pedidos por região baseado no país - 🔥 AGORA SÓ PEDIDOS PAGOS
    pedidos_internacional = []
    pedidos_europeu = []
    
    for pedido in pedidos_pagos:  # 🔥 AGORA SÓ PAGOS
        pais = pedido.get('pais', '').lower()
        if any(pais_internacional in pais for pais_internacional in paises_internacionais):
            pedidos_internacional.append(pedido)
        elif any(pais_europeu in pais for pais_europeu in paises_europeus):
            pedidos_europeu.append(pedido)
        # Se não tiver país definido, considerar como europeu (default)
        elif not pais:
            pedidos_europeu.append(pedido)
    
    total_internacional = len(pedidos_internacional)
    total_europeu = len(pedidos_europeu)
    
    # 🔥 DEBUG: VERIFICAR OS TIPOS DE OFERTA NOS PEDIDOS PAGOS
    # Já temos pedidos_pagos definido acima
    
    print("🔍 DEBUG - TIPOS DE OFERTA NOS PEDIDOS PAGOS:")
    for pedido in pedidos_pagos:
        oferta_tipo = pedido.get("oferta_tipo", "N/A")
        print(f"   #{pedido.get('id', 'N/A')}: '{oferta_tipo}'")

    # 🔥 CALCULAR OFERTAS POR TIPO - CORRIGIDO PARA OS NOMES REAIS
    oferta_original_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "original")
    oferta_tamanho_45_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["tamanho_4.5", "oferta_tamanho_45"])
    oferta_portachaves_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["portachaves", "oferta_portachaves"])
    oferta_surpresa_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "oferta_surpresa")
    pagamento_direto_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "pagamento_direto")  # 🔥 NOVA ESTATÍSTICA
    oferta_recusadas_count = ESTATISTICAS['ofertas_recusadas']
    
    pedidos_pagos_count = len(pedidos_pagos)  # Esta variável agora é redundante

    print(f"🔍 DEBUG CONTAGEM:")
    print(f"   • Original: {oferta_original_count}")
    print(f"   • Tamanho 4.5: {oferta_tamanho_45_count}") 
    print(f"   • Portachaves: {oferta_portachaves_count}")
    print(f"   • Total pedidos pagos: {pedidos_pagos_count}")
    
    # 🔥 CALCULAR TOTAIS FINANCEIROS - SISTEMA SIMPLES: MANUAL + AUTOMÁTICO
    if ESTATISTICAS.get('usar_valores_manuais'):
        print("📊 Usando sistema MANUAL + AUTOMÁTICO")
        
        # 🔥 VALORES BASE MANUAIS
        subtotal_base = ESTATISTICAS.get('subtotal_manual', 0)
        impostos_base = ESTATISTICAS.get('impostos_manual', 0)
        frete_base = ESTATISTICAS.get('frete_manual', 0)
        internacional_base = ESTATISTICAS.get('internacional_manual', 0)
        europeu_base = ESTATISTICAS.get('europeu_manual', 0)
        total_final_base = ESTATISTICAS.get('total_final_manual', 0)
        
        # 🔥 CALCULAR VALORES DAS VENDAS AUTOMÁTICAS (todos os pedidos pagos)
        subtotal_vendas = 0.0
        impostos_vendas = 0.0
        frete_vendas = 0.0
        internacional_vendas = 0.0
        europeu_vendas = 0.0
        total_final_vendas = 0.0

        print(f"💰 PROCESSANDO {len(pedidos_pagos)} PEDIDOS PAGOS PARA ACRESCENTAR AOS VALORES MANUAIS")
        
        for pedido in pedidos_pagos:
            # 🔥 OBTER INFORMAÇÕES DE MOEDA DO PEDIDO
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            # 🔥 SE JÁ TEM VALOR CONVERTIDO EM EUR, USAR ESSE
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
                print(f"   ✅ Pedido #{pedido.get('id', 'N/A')}: {moeda_original} → Já convertido = €{total_eur:.2f}")
            else:
                # 🔥 CONVERTER SE NECESSÁRIO
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                    print(f"   🔄 Pedido #{pedido.get('id', 'N/A')}: {moeda_original} {total_original:.2f} → €{total_eur:.2f} EUR (taxa: 1 {moeda_original} = {taxa} EUR)")
                else:
                    total_eur = total_original
                    print(f"   ✅ Pedido #{pedido.get('id', 'N/A')}: EUR = €{total_eur:.2f}")
            
            # 🔥 DETERMINAR REGIÃO BASEADO NO PAÍS - COM CANADÁ
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                internacional_vendas += total_eur
                print(f"     🌎 Venda Internacional: {pais} = €{total_eur:.2f}")
            else:
                europeu_vendas += total_eur
                print(f"     🇪🇺 Venda Europeu: {pais if pais else 'Não especificado'} = €{total_eur:.2f}")
            
            # 🔥 CALCULAR COMPONENTES EM EUR (proporcionalmente)
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                subtotal_vendas += subtotal_eur
                impostos_vendas += imposto_eur
                frete_vendas += frete_eur
                total_final_vendas += total_eur
                
                print(f"     📊 Componentes: Subtotal €{subtotal_eur:.2f}, Imposto €{imposto_eur:.2f}, Frete €{frete_eur:.2f}")
            else:
                # 🔥 SE NÃO TEM VALORES DETALHADOS, USAR APENAS O TOTAL
                total_final_vendas += total_eur
                print(f"     ⚠️ Sem detalhes - usando apenas total: €{total_eur:.2f}")

        # 🔥 SOMAR BASE MANUAL + VENDAS AUTOMÁTICAS
        total_bruto = subtotal_base + subtotal_vendas
        total_impostos = impostos_base + impostos_vendas
        total_frete = frete_base + frete_vendas
        total_internacional_valor = internacional_base + internacional_vendas
        total_europeu_valor = europeu_base + europeu_vendas
        total_final = total_final_base + total_final_vendas
        
        print(f"💰 TOTAIS MANUAL + AUTOMÁTICO:")
        print(f"   • Base Manual: Subtotal €{subtotal_base:.2f}, Int €{internacional_base:.2f}, Eur €{europeu_base:.2f}")
        print(f"   • Vendas Automáticas: Subtotal €{subtotal_vendas:.2f}, Int €{internacional_vendas:.2f}, Eur €{europeu_vendas:.2f}")
        print(f"   • Total Final: Subtotal €{total_bruto:.2f}, Int €{total_internacional_valor:.2f}, Eur €{total_europeu_valor:.2f}")
        
    else:
        # 🔥 CÁLCULO NORMAL BASEADO APENAS EM PEDIDOS (modo automático puro)
        total_bruto = 0.0
        total_impostos = 0.0
        total_frete = 0.0
        total_final = 0.0
        total_internacional_valor = 0.0
        total_europeu_valor = 0.0

        print(f"💰 PROCESSANDO {len(pedidos_pagos)} PEDIDOS PAGOS (MODO AUTOMÁTICO PURO)")
        
        for pedido in pedidos_pagos:
            # 🔥 OBTER INFORMAÇÕES DE MOEDA DO PEDIDO
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            # 🔥 SE JÁ TEM VALOR CONVERTIDO EM EUR, USAR ESSE
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
                print(f"   ✅ Pedido #{pedido.get('id', 'N/A')}: {moeda_original} → Já convertido = €{total_eur:.2f}")
            else:
                # 🔥 CONVERTER SE NECESSÁRIO
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                    print(f"   🔄 Pedido #{pedido.get('id', 'N/A')}: {moeda_original} {total_original:.2f} → €{total_eur:.2f} EUR (taxa: 1 {moeda_original} = {taxa} EUR)")
                else:
                    total_eur = total_original
                    print(f"   ✅ Pedido #{pedido.get('id', 'N/A')}: EUR = €{total_eur:.2f}")
            
            # 🔥 DETERMINAR REGIÃO BASEADO NO PAÍS - COM CANADÁ
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                total_internacional_valor += total_eur
                print(f"     🌎 Pedido Internacional: {pais} = €{total_eur:.2f}")
            else:
                total_europeu_valor += total_eur
                print(f"     🇪🇺 Pedido Europeu: {pais if pais else 'Não especificado'} = €{total_eur:.2f}")
            
            # 🔥 CALCULAR COMPONENTES EM EUR (proporcionalmente)
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                total_bruto += subtotal_eur
                total_impostos += imposto_eur
                total_frete += frete_eur
                total_final += total_eur
                
                print(f"     📊 Componentes: Subtotal €{subtotal_eur:.2f}, Imposto €{imposto_eur:.2f}, Frete €{frete_eur:.2f}")
            else:
                # 🔥 SE NÃO TEM VALORES DETALHADOS, USAR APENAS O TOTAL
                total_final += total_eur
                print(f"     ⚠️ Sem detalhes - usando apenas total: €{total_eur:.2f}")

    print(f"💰 TOTAIS FINAIS EM EUR:")
    print(f"   • Subtotal: €{total_bruto:.2f}")
    print(f"   • Impostos: €{total_impostos:.2f}")
    print(f"   • Frete: €{total_frete:.2f}")
    print(f"   • Total Internacional (UK/US/BR/CA): €{total_internacional_valor:.2f}")  # 🔥 ATUALIZADO COM CA
    print(f"   • Total Europeu: €{total_europeu_valor:.2f}")
    print(f"   • Total Final: €{total_final:.2f}")
    
    # 🔥🔥🔥 CALCULAR USUÁRIOS ATIVOS REAIS (SUBSTITUIR O RANDOM)
    usuarios_ativos = obter_utilizadores_ativos()  # 🔥 CONTADOR REAL
    ver_estado_contador_0000()  # 🔥 Agora é 14:55

    # 🔥 CONSTRUIR MENSAGEM COM NOVA ESTRUTURA - COM CANADÁ
    mensagem = f"""
🛡️ *PAINEL ADMIN - GODSPLAN*

💼 *RESUMO FINANCEIRO (TODOS OS VALORES EM EUR):*
💰 *Total Subtotal:* €{total_bruto:.2f}
🧾 *Total Impostos:* €{total_impostos:.2f}
🚚 *Total Frete:* €{total_frete:.2f}
🌎 *Total Internacional (UK/US/BR/CA):* €{total_internacional_valor:.2f}  
🇪🇺 *Total Europeu:* €{total_europeu_valor:.2f}
💳 *Total Final (Com Tudo):* €{total_final:.2f}

📊 *ESTATÍSTICAS:*
• 📦 Total de pedidos: {total_pedidos} 
• 🌎 Total pedidos internacional (UK/US/BR/CA): {total_internacional} 
• 🇪🇺 Total pedidos europeu: {total_europeu}
• ✅ Pedidos pagos: {pedidos_pagos_count} 
• 💳 Pagamentos diretos: {pagamento_direto_count}
• 🎯 Oferta original: {oferta_original_count}
• 📏 Oferta tamanho 4.5: {oferta_tamanho_45_count}
• 🔑 Oferta portachaves: {oferta_portachaves_count}
• 🎁 Oferta surpresa: {oferta_surpresa_count} 
• ❌ Ofertas recusadas: {oferta_recusadas_count}"""

    # 🔥 CALCULAR TAXAS (agora depois das ofertas)
    if total_pedidos > 0:
        taxa_conversao = (pedidos_pagos_count / total_pedidos) * 100
        taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / total_pedidos) * 100
        mensagem += f"\n• 📈 Taxa de conversão: {taxa_conversao:.1f}%"
        mensagem += f"\n• 🔄 Taxa de recuperação: {taxa_recuperacao:.1f}%"

    # Adicionar informação do modo
    if ESTATISTICAS.get('usar_valores_manuais'):
        subtotal_base = ESTATISTICAS.get('subtotal_manual', 0)
        internacional_base = ESTATISTICAS.get('internacional_manual', 0)
        europeu_base = ESTATISTICAS.get('europeu_manual', 0)
        
        mensagem += f"\n\n🔧 *MODO: MANUAL + AUTOMÁTICO*"
        mensagem += f"\n💰 *Base Manual:* Subtotal €{subtotal_base:.2f}, Int €{internacional_base:.2f}, Eur €{europeu_base:.2f}"
        mensagem += f"\n📈 *Vendas automáticas são somadas à base*"
    else:
        mensagem += f"\n\n🤖 *MODO: AUTOMÁTICO PURO*"

    mensagem += f"\n👥 *Utilizadores Ativos Hoje:* {usuarios_ativos}"

    # 🔥 BOTÕES SIMPLIFICADOS
    keyboard = [
        [InlineKeyboardButton("🔄 Atualizar", callback_data="admin_refresh")],
        [InlineKeyboardButton("📊 Exportar", callback_data="menu_export")],
        [InlineKeyboardButton("⚙️ Opções", callback_data="btn_options")],
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)

    try:
        await message_method(
            text=mensagem,
            parse_mode="Markdown",
            reply_markup=reply_markup
        )
    except BadRequest as e:
        if "Message is not modified" in str(e):
            print("ℹ️ Mensagem não foi modificada (conteúdo igual)")
            # Não faz nada, é normal quando o conteúdo é o mesmo
        else:
            raise e


    
async def menu_export(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Menu de exportação - popup"""
    query = update.callback_query
    
    
    # Depois do popup, mostra os botões inline
    keyboard = [
        [InlineKeyboardButton("📄 CSV", callback_data="export_csv")],
        [InlineKeyboardButton("📝 TXT", callback_data="export_txt")],
        [InlineKeyboardButton("📑 PDF", callback_data="export_pdf")],
        [InlineKeyboardButton("📘 Word", callback_data="export_word")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text="*Escolha o formato:*",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )

    


async def export_csv_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exportar relatório completo para CSV"""
    query = update.callback_query
    await query.answer()
    
    try:
        # Obter os mesmos dados do admin_command
        pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
        usuarios_ativos = obter_utilizadores_ativos()
        total_pedidos = len(PEDIDOS_REGISTO)
        
        # 🔥 CALCULAR ESTATÍSTICAS (IGUAL AO ADMIN_COMMAND)
        paises_internacionais = ['reino unido', 'united kingdom', 'uk', 'estados unidos', 'united states', 'us', 'usa', 'brasil', 'brazil', 'canada', 'canadá']
        paises_europeus = [
        'portugal', 'espanha', 'spain', 'frança', 'france', 'franca',
        'alemanha', 'germany', 'itália', 'italia', 'bélgica', 'belgica',
        'países baixos', 'paises baixos', 'holanda', 'netherlands',
        'irlanda', 'ireland', 'luxemburgo', 'luxembourg',
        'suecia', 'sweden', 'dinamarca', 'denmark'
        ]
        
        # Calcular pedidos por região
        pedidos_internacional = []
        pedidos_europeu = []
        
        for pedido in PEDIDOS_REGISTO.values():
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                pedidos_internacional.append(pedido)
            elif any(pais_europeu in pais for pais_europeu in paises_europeus):
                pedidos_europeu.append(pedido)
            elif not pais:
                pedidos_europeu.append(pedido)
        
        total_internacional = len(pedidos_internacional)
        total_europeu = len(pedidos_europeu)
        
        # 🔥 CALCULAR OFERTAS POR TIPO
        oferta_original_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "original")
        oferta_tamanho_45_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["tamanho_4.5", "oferta_tamanho_45"])
        oferta_portachaves_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["portachaves", "oferta_portachaves"])
        pagamento_direto_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "pagamento_direto")
        oferta_recusadas_count = ESTATISTICAS['ofertas_recusadas']
        pedidos_pagos_count = len(pedidos_pagos)
        
        # 🔥 CALCULAR TOTAIS FINANCEIROS (IGUAL AO ADMIN_COMMAND)
        total_bruto = 0.0
        total_impostos = 0.0
        total_frete = 0.0
        total_final = 0.0
        total_internacional_valor = 0.0
        total_europeu_valor = 0.0
        
        # Para calcular subtotais por região
        subtotal_internacional = 0.0
        subtotal_europeu = 0.0
        impostos_europeus = 0.0
        fretes_internacionais = 0.0
        fretes_europeus = 0.0
        
        for pedido in pedidos_pagos:
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
            else:
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                else:
                    total_eur = total_original
            
            # Determinar região
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                total_internacional_valor += total_eur
                # Calcular componentes para internacional
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_internacional += pedido.get('subtotal', 0) * proporcao
                    fretes_internacionais += pedido.get('frete', 0) * proporcao
            else:
                total_europeu_valor += total_eur
                # Calcular componentes para europeu
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_europeu += pedido.get('subtotal', 0) * proporcao
                    impostos_europeus += pedido.get('imposto', 0) * proporcao
                    fretes_europeus += pedido.get('frete', 0) * proporcao
            
            # Calcular componentes gerais
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                total_bruto += subtotal_eur
                total_impostos += imposto_eur
                total_frete += frete_eur
                total_final += total_eur
            else:
                total_final += total_eur
        
        # Calcular taxas
        taxa_conversao = (pedidos_pagos_count / total_pedidos) * 100 if total_pedidos > 0 else 0
        taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / total_pedidos) * 100 if total_pedidos > 0 else 0
        
        # Criar CSV com TODOS os dados
        csv_content = "RELATÓRIO ADMINISTRATIVO - GODSPLAN\n"
        csv_content += "===================================\n\n"
        
        # RESUMO FINANCEIRO SEPARADO
        csv_content += "RESUMO FINANCEIRO (TODOS OS VALORES EM EUR)\n"
        csv_content += "Categoria,Valor\n"
        
        # MERCADO EUROPEU
        csv_content += "MERCADO EUROPEU,\n"
        csv_content += f"Subtotal,€{subtotal_europeu:.2f}\n"
        csv_content += f"Impostos,€{impostos_europeus:.2f}\n"
        csv_content += f"Frete,€{fretes_europeus:.2f}\n"
        csv_content += f"Total Europeu,€{total_europeu_valor:.2f}\n\n"
        
        # MERCADO INTERNACIONAL
        csv_content += "MERCADO INTERNACIONAL,\n"
        csv_content += f"Subtotal,€{subtotal_internacional:.2f}\n"
        csv_content += f"Impostos,€0.00\n"
        csv_content += f"Frete,€{fretes_internacionais:.2f}\n"
        csv_content += f"Total Internacional,€{total_internacional_valor:.2f}\n\n"
        
        # RESUMO GERAL CONJUNTO
        csv_content += "RESUMO GERAL,\n"
        csv_content += f"Subtotal Total,€{subtotal_europeu + subtotal_internacional:.2f}\n"
        csv_content += f"Impostos Total,€{impostos_europeus:.2f}\n"
        csv_content += f"Frete Total,€{fretes_europeus + fretes_internacionais:.2f}\n"
        csv_content += f"Total Final,€{total_europeu_valor + total_internacional_valor:.2f}\n\n"
        
        # ESTATÍSTICAS COMPLETAS
        csv_content += "ESTATÍSTICAS\n"
        csv_content += "Categoria,Quantidade\n"
        csv_content += f"Total de pedidos,{total_pedidos}\n"
        csv_content += f"Total pedidos internacional (UK/US/BR),{total_internacional}\n"
        csv_content += f"Total pedidos europeu,{total_europeu}\n"
        csv_content += f"Pedidos pagos,{pedidos_pagos_count}\n"
        csv_content += f"Pagamentos diretos,{pagamento_direto_count}\n"
        csv_content += f"Oferta original,{oferta_original_count}\n"
        csv_content += f"Oferta tamanho 4.5,{oferta_tamanho_45_count}\n"
        csv_content += f"Oferta portachaves,{oferta_portachaves_count}\n"
        csv_content += f"Ofertas recusadas,{oferta_recusadas_count}\n"
        csv_content += f"Utilizadores Ativos Hoje,{usuarios_ativos}\n"
        csv_content += f"Taxa de conversão,{taxa_conversao:.1f}%\n"
        csv_content += f"Taxa de recuperação,{taxa_recuperacao:.1f}%\n\n"
        
        # Data de exportação
        from datetime import datetime
        csv_content += f"Data de exportação,{datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        
        # Criar ficheiro em memória
        csv_file = io.BytesIO(csv_content.encode('utf-8'))
        csv_file.name = "relatorio_admin.csv"
        
        # Enviar ficheiro
        await context.bot.send_document(
            chat_id=query.message.chat_id,
            document=InputFile(csv_file),
            caption="📄 *Relatório Admin exportado em CSV*",
            parse_mode='Markdown'
        )
        
        await query.edit_message_text("✅ *Relatório CSV exportado com sucesso!*", parse_mode='Markdown')
        
    except Exception as e:
        await query.edit_message_text(f"❌ Erro ao exportar CSV: {e}")




async def export_txt_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exportar relatório completo para TXT"""
    query = update.callback_query
    await query.answer()
    
    try:
        # Obter os mesmos dados do admin_command
        pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
        usuarios_ativos = obter_utilizadores_ativos()
        total_pedidos = len(PEDIDOS_REGISTO)
        
        # 🔥 CALCULAR ESTATÍSTICAS (IGUAL AO ADMIN_COMMAND)
        paises_internacionais = ['reino unido', 'united kingdom', 'uk', 'estados unidos', 'united states', 'us', 'usa', 'brasil', 'brazil', 'canada', 'canadá']
        paises_europeus = [
        'portugal', 'espanha', 'spain', 'frança', 'france', 'franca',
        'alemanha', 'germany', 'itália', 'italia', 'bélgica', 'belgica',
        'países baixos', 'paises baixos', 'holanda', 'netherlands',
        'irlanda', 'ireland', 'luxemburgo', 'luxembourg',
        'suecia', 'sweden', 'dinamarca', 'denmark'
        ]

        # Calcular pedidos por região
        pedidos_internacional = []
        pedidos_europeu = []
        
        for pedido in PEDIDOS_REGISTO.values():
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                pedidos_internacional.append(pedido)
            elif any(pais_europeu in pais for pais_europeu in paises_europeus):
                pedidos_europeu.append(pedido)
            elif not pais:
                pedidos_europeu.append(pedido)
        
        total_internacional = len(pedidos_internacional)
        total_europeu = len(pedidos_europeu)
        
        # 🔥 CALCULAR OFERTAS POR TIPO
        oferta_original_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "original")
        oferta_tamanho_45_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["tamanho_4.5", "oferta_tamanho_45"])
        oferta_portachaves_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["portachaves", "oferta_portachaves"])
        pagamento_direto_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "pagamento_direto")
        oferta_recusadas_count = ESTATISTICAS['ofertas_recusadas']
        pedidos_pagos_count = len(pedidos_pagos)
        
        # 🔥 CALCULAR TOTAIS FINANCEIROS (IGUAL AO ADMIN_COMMAND)
        total_bruto = 0.0
        total_impostos = 0.0
        total_frete = 0.0
        total_final = 0.0
        total_internacional_valor = 0.0
        total_europeu_valor = 0.0
        
        # Para calcular subtotais por região
        subtotal_internacional = 0.0
        subtotal_europeu = 0.0
        impostos_europeus = 0.0
        fretes_internacionais = 0.0
        fretes_europeus = 0.0
        
        for pedido in pedidos_pagos:
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
            else:
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                else:
                    total_eur = total_original
            
            # Determinar região
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                total_internacional_valor += total_eur
                # Calcular componentes para internacional
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_internacional += pedido.get('subtotal', 0) * proporcao
                    fretes_internacionais += pedido.get('frete', 0) * proporcao
            else:
                total_europeu_valor += total_eur
                # Calcular componentes para europeu
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_europeu += pedido.get('subtotal', 0) * proporcao
                    impostos_europeus += pedido.get('imposto', 0) * proporcao
                    fretes_europeus += pedido.get('frete', 0) * proporcao
            
            # Calcular componentes gerais
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                total_bruto += subtotal_eur
                total_impostos += imposto_eur
                total_frete += frete_eur
                total_final += total_eur
            else:
                total_final += total_eur
        
        # Calcular taxas
        taxa_conversao = (pedidos_pagos_count / total_pedidos) * 100 if total_pedidos > 0 else 0
        taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / total_pedidos) * 100 if total_pedidos > 0 else 0
        
        # Criar TXT com TODOS os dados
        txt_content = "👑 RELATÓRIO ADMINISTRATIVO - GODSPLAN 👑\n\n"
        
        # RESUMO FINANCEIRO SEPARADO
        txt_content += "💰 RESUMO FINANCEIRO (TODOS OS VALORES EM EUR):\n"
        
        # MERCADO EUROPEU
        txt_content += "🇪🇺 MERCADO EUROPEU:\n"
        txt_content += f"   • Subtotal: €{subtotal_europeu:.2f}\n"
        txt_content += f"   • Impostos: €{impostos_europeus:.2f}\n"
        txt_content += f"   • Frete: €{fretes_europeus:.2f}\n"
        txt_content += f"   • Total Europeu: €{total_europeu_valor:.2f}\n\n"
        
        # MERCADO INTERNACIONAL
        txt_content += "🌎 MERCADO INTERNACIONAL:\n"
        txt_content += f"   • Subtotal: €{subtotal_internacional:.2f}\n"
        txt_content += f"   • Impostos: €0.00\n"
        txt_content += f"   • Frete: €{fretes_internacionais:.2f}\n"
        txt_content += f"   • Total Internacional: €{total_internacional_valor:.2f}\n\n"
        
        # RESUMO GERAL CONJUNTO
        txt_content += "📊 RESUMO GERAL:\n"
        txt_content += f"• Subtotal Total: €{subtotal_europeu + subtotal_internacional:.2f}\n"
        txt_content += f"• Impostos Total: €{impostos_europeus:.2f}\n"
        txt_content += f"• Frete Total: €{fretes_europeus + fretes_internacionais:.2f}\n"
        txt_content += f"• Total Final: €{total_europeu_valor + total_internacional_valor:.2f}\n\n"
        
        # ESTATÍSTICAS COMPLETAS
        txt_content += "📈 ESTATÍSTICAS:\n"
        txt_content += f"• 📦 Total de pedidos: {total_pedidos}\n"
        txt_content += f"• 🌎 Total pedidos internacional (UK/US/BR): {total_internacional}\n"
        txt_content += f"• 🇪🇺 Total pedidos europeu: {total_europeu}\n"
        txt_content += f"• ✅ Pedidos pagos: {pedidos_pagos_count}\n"
        txt_content += f"• 💳 Pagamentos diretos: {pagamento_direto_count}\n"
        txt_content += f"• 🎯 Oferta original: {oferta_original_count}\n"
        txt_content += f"• 📏 Oferta tamanho 4.5: {oferta_tamanho_45_count}\n"
        txt_content += f"• 🔑 Oferta portachaves: {oferta_portachaves_count}\n"
        txt_content += f"• ❌ Ofertas recusadas: {oferta_recusadas_count}\n"
        txt_content += f"• 👥 Utilizadores Ativos Hoje: {usuarios_ativos}\n"
        txt_content += f"• 📈 Taxa de conversão: {taxa_conversao:.1f}%\n"
        txt_content += f"• 🔄 Taxa de recuperação: {taxa_recuperacao:.1f}%\n\n"
        
        # Data de exportação
        from datetime import datetime
        txt_content += f"📅 Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        
        # Criar ficheiro em memória
        txt_file = io.BytesIO(txt_content.encode('utf-8'))
        txt_file.name = "relatorio_admin.txt"
        
        # Enviar ficheiro
        await context.bot.send_document(
            chat_id=query.message.chat_id,
            document=InputFile(txt_file),
            caption="📝 *Relatório Admin exportado em TXT*",
            parse_mode='Markdown'
        )
        
        await query.edit_message_text("✅ *Relatório TXT exportado com sucesso!*", parse_mode='Markdown')
        
    except Exception as e:
        await query.edit_message_text(f"❌ Erro ao exportar TXT: {e}")

async def export_pdf_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exportar relatório completo para PDF"""
    query = update.callback_query
    await query.answer()
    
    try:
        # Obter os mesmos dados do admin_command
        pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
        usuarios_ativos = obter_utilizadores_ativos()
        total_pedidos = len(PEDIDOS_REGISTO)
        
        # 🔥 CALCULAR ESTATÍSTICAS (IGUAL AO ADMIN_COMMAND)
        paises_internacionais = ['reino unido', 'united kingdom', 'uk', 'estados unidos', 'united states', 'us', 'usa', 'brasil', 'brazil', 'canada', 'canadá']
        paises_europeus = [
        'portugal', 'espanha', 'spain', 'frança', 'france', 'franca',
        'alemanha', 'germany', 'itália', 'italia', 'bélgica', 'belgica',
        'países baixos', 'paises baixos', 'holanda', 'netherlands',
        'irlanda', 'ireland', 'luxemburgo', 'luxembourg',
        'suecia', 'sweden', 'dinamarca', 'denmark'
        ]

        # Calcular pedidos por região
        pedidos_internacional = []
        pedidos_europeu = []
        
        for pedido in PEDIDOS_REGISTO.values():
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                pedidos_internacional.append(pedido)
            elif any(pais_europeu in pais for pais_europeu in paises_europeus):
                pedidos_europeu.append(pedido)
            elif not pais:
                pedidos_europeu.append(pedido)
        
        total_internacional = len(pedidos_internacional)
        total_europeu = len(pedidos_europeu)
        
        # 🔥 CALCULAR OFERTAS POR TIPO
        oferta_original_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "original")
        oferta_tamanho_45_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["tamanho_4.5", "oferta_tamanho_45"])
        oferta_portachaves_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["portachaves", "oferta_portachaves"])
        pagamento_direto_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "pagamento_direto")
        oferta_recusadas_count = ESTATISTICAS['ofertas_recusadas']
        pedidos_pagos_count = len(pedidos_pagos)
        
        # 🔥 CALCULAR TOTAIS FINANCEIROS (IGUAL AO ADMIN_COMMAND)
        total_bruto = 0.0
        total_impostos = 0.0
        total_frete = 0.0
        total_final = 0.0
        total_internacional_valor = 0.0
        total_europeu_valor = 0.0
        
        # Para calcular subtotais por região
        subtotal_internacional = 0.0
        subtotal_europeu = 0.0
        impostos_europeus = 0.0
        fretes_internacionais = 0.0
        fretes_europeus = 0.0
        
        for pedido in pedidos_pagos:
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
            else:
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                else:
                    total_eur = total_original
            
            # Determinar região
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                total_internacional_valor += total_eur
                # Calcular componentes para internacional
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_internacional += pedido.get('subtotal', 0) * proporcao
                    fretes_internacionais += pedido.get('frete', 0) * proporcao
            else:
                total_europeu_valor += total_eur
                # Calcular componentes para europeu
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_europeu += pedido.get('subtotal', 0) * proporcao
                    impostos_europeus += pedido.get('imposto', 0) * proporcao
                    fretes_europeus += pedido.get('frete', 0) * proporcao
            
            # Calcular componentes gerais
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                total_bruto += subtotal_eur
                total_impostos += imposto_eur
                total_frete += frete_eur
                total_final += total_eur
            else:
                total_final += total_eur
        
        # Calcular taxas
        taxa_conversao = (pedidos_pagos_count / total_pedidos) * 100 if total_pedidos > 0 else 0
        taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / total_pedidos) * 100 if total_pedidos > 0 else 0
        
        # Criar PDF com TODOS os dados
        pdf_content = "RELATÓRIO ADMINISTRATIVO - GODSPLAN\n"
        pdf_content += "===================================\n\n"
        
        # RESUMO FINANCEIRO SEPARADO
        pdf_content += "RESUMO FINANCEIRO (TODOS OS VALORES EM EUR):\n\n"
        
        # MERCADO EUROPEU
        pdf_content += "MERCADO EUROPEU:\n"
        pdf_content += f"  Subtotal: €{subtotal_europeu:.2f}\n"
        pdf_content += f"  Impostos: €{impostos_europeus:.2f}\n"
        pdf_content += f"  Frete: €{fretes_europeus:.2f}\n"
        pdf_content += f"  Total Europeu: €{total_europeu_valor:.2f}\n\n"
        
        # MERCADO INTERNACIONAL
        pdf_content += "MERCADO INTERNACIONAL:\n"
        pdf_content += f"  Subtotal: €{subtotal_internacional:.2f}\n"
        pdf_content += f"  Impostos: €0.00\n"
        pdf_content += f"  Frete: €{fretes_internacionais:.2f}\n"
        pdf_content += f"  Total Internacional: €{total_internacional_valor:.2f}\n\n"
        
        # RESUMO GERAL CONJUNTO
        pdf_content += "RESUMO GERAL:\n"
        pdf_content += f"  Subtotal Total: €{subtotal_europeu + subtotal_internacional:.2f}\n"
        pdf_content += f"  Impostos Total: €{impostos_europeus:.2f}\n"
        pdf_content += f"  Frete Total: €{fretes_europeus + fretes_internacionais:.2f}\n"
        pdf_content += f"  Total Final: €{total_europeu_valor + total_internacional_valor:.2f}\n\n"
        
        # ESTATÍSTICAS COMPLETAS
        pdf_content += "ESTATÍSTICAS:\n"
        pdf_content += f"  • Total de pedidos: {total_pedidos}\n"
        pdf_content += f"  • Total pedidos internacional (UK/US/BR): {total_internacional}\n"
        pdf_content += f"  • Total pedidos europeu: {total_europeu}\n"
        pdf_content += f"  • Pedidos pagos: {pedidos_pagos_count}\n"
        pdf_content += f"  • Pagamentos diretos: {pagamento_direto_count}\n"
        pdf_content += f"  • Oferta original: {oferta_original_count}\n"
        pdf_content += f"  • Oferta tamanho 4.5: {oferta_tamanho_45_count}\n"
        pdf_content += f"  • Oferta portachaves: {oferta_portachaves_count}\n"
        pdf_content += f"  • Ofertas recusadas: {oferta_recusadas_count}\n"
        pdf_content += f"  • Utilizadores Ativos Hoje: {usuarios_ativos}\n"
        pdf_content += f"  • Taxa de conversão: {taxa_conversao:.1f}%\n"
        pdf_content += f"  • Taxa de recuperação: {taxa_recuperacao:.1f}%\n\n"
        
        # Data de exportação
        from datetime import datetime
        pdf_content += f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        
        # Criar ficheiro em memória
        pdf_file = io.BytesIO(pdf_content.encode('utf-8'))
        pdf_file.name = "relatorio_admin.pdf"
        
        await context.bot.send_document(
            chat_id=query.message.chat_id,
            document=InputFile(pdf_file),
            caption="📑 *Relatório Admin exportado em PDF*",
            parse_mode='Markdown'
        )
        
        await query.edit_message_text("✅ *Relatório PDF exportado com sucesso!*", parse_mode='Markdown')
        
    except Exception as e:
        await query.edit_message_text(f"❌ Erro ao exportar PDF: {e}")

async def export_word_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exportar relatório completo para Word"""
    query = update.callback_query
    await query.answer()
    
    try:
        # Obter os mesmos dados do admin_command
        pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
        usuarios_ativos = obter_utilizadores_ativos()
        total_pedidos = len(PEDIDOS_REGISTO)
        
        # 🔥 CALCULAR ESTATÍSTICAS (IGUAL AO ADMIN_COMMAND)
        paises_internacionais = ['reino unido', 'united kingdom', 'uk', 'estados unidos', 'united states', 'us', 'usa', 'brasil', 'brazil', 'canada', 'canadá']
        paises_europeus = [
        'portugal', 'espanha', 'spain', 'frança', 'france', 'franca',
        'alemanha', 'germany', 'itália', 'italia', 'bélgica', 'belgica',
        'países baixos', 'paises baixos', 'holanda', 'netherlands',
        'irlanda', 'ireland', 'luxemburgo', 'luxembourg',
        'suecia', 'sweden', 'dinamarca', 'denmark'
        ]

        # Calcular pedidos por região
        pedidos_internacional = []
        pedidos_europeu = []
        
        for pedido in PEDIDOS_REGISTO.values():
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                pedidos_internacional.append(pedido)
            elif any(pais_europeu in pais for pais_europeu in paises_europeus):
                pedidos_europeu.append(pedido)
            elif not pais:
                pedidos_europeu.append(pedido)
        
        total_internacional = len(pedidos_internacional)
        total_europeu = len(pedidos_europeu)
        
        # 🔥 CALCULAR OFERTAS POR TIPO
        oferta_original_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "original")
        oferta_tamanho_45_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["tamanho_4.5", "oferta_tamanho_45"])
        oferta_portachaves_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["portachaves", "oferta_portachaves"])
        pagamento_direto_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "pagamento_direto")
        oferta_recusadas_count = ESTATISTICAS['ofertas_recusadas']
        pedidos_pagos_count = len(pedidos_pagos)
        
        # 🔥 CALCULAR TOTAIS FINANCEIROS (IGUAL AO ADMIN_COMMAND)
        total_bruto = 0.0
        total_impostos = 0.0
        total_frete = 0.0
        total_final = 0.0
        total_internacional_valor = 0.0
        total_europeu_valor = 0.0
        
        # Para calcular subtotais por região
        subtotal_internacional = 0.0
        subtotal_europeu = 0.0
        impostos_europeus = 0.0
        fretes_internacionais = 0.0
        fretes_europeus = 0.0
        
        for pedido in pedidos_pagos:
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
            else:
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                else:
                    total_eur = total_original
            
            # Determinar região
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in paises_internacionais):
                total_internacional_valor += total_eur
                # Calcular componentes para internacional
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_internacional += pedido.get('subtotal', 0) * proporcao
                    fretes_internacionais += pedido.get('frete', 0) * proporcao
            else:
                total_europeu_valor += total_eur
                # Calcular componentes para europeu
                if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                    proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                    subtotal_europeu += pedido.get('subtotal', 0) * proporcao
                    impostos_europeus += pedido.get('imposto', 0) * proporcao
                    fretes_europeus += pedido.get('frete', 0) * proporcao
            
            # Calcular componentes gerais
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                total_bruto += subtotal_eur
                total_impostos += imposto_eur
                total_frete += frete_eur
                total_final += total_eur
            else:
                total_final += total_eur
        
        # Calcular taxas
        taxa_conversao = (pedidos_pagos_count / total_pedidos) * 100 if total_pedidos > 0 else 0
        taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / total_pedidos) * 100 if total_pedidos > 0 else 0
        
        # Criar Word com TODOS os dados
        word_content = "RELATÓRIO ADMINISTRATIVO - GODSPLAN\n"
        word_content += "===================================\n\n"
        
        # RESUMO FINANCEIRO SEPARADO
        word_content += "RESUMO FINANCEIRO (TODOS OS VALORES EM EUR):\n\n"
        
        # MERCADO EUROPEU
        word_content += "MERCADO EUROPEU:\n"
        word_content += f"  • Subtotal: €{subtotal_europeu:.2f}\n"
        word_content += f"  • Impostos: €{impostos_europeus:.2f}\n"
        word_content += f"  • Frete: €{fretes_europeus:.2f}\n"
        word_content += f"  • Total Europeu: €{total_europeu_valor:.2f}\n\n"
        
        # MERCADO INTERNACIONAL
        word_content += "MERCADO INTERNACIONAL:\n"
        word_content += f"  • Subtotal: €{subtotal_internacional:.2f}\n"
        word_content += f"  • Impostos: €0.00\n"
        word_content += f"  • Frete: €{fretes_internacionais:.2f}\n"
        word_content += f"  • Total Internacional: €{total_internacional_valor:.2f}\n\n"
        
        # RESUMO GERAL CONJUNTO
        word_content += "RESUMO GERAL:\n"
        word_content += f"  • Subtotal Total: €{subtotal_europeu + subtotal_internacional:.2f}\n"
        word_content += f"  • Impostos Total: €{impostos_europeus:.2f}\n"
        word_content += f"  • Frete Total: €{fretes_europeus + fretes_internacionais:.2f}\n"
        word_content += f"  • Total Final: €{total_europeu_valor + total_internacional_valor:.2f}\n\n"
        
        # ESTATÍSTICAS COMPLETAS
        word_content += "ESTATÍSTICAS:\n"
        word_content += f"  • Total de pedidos: {total_pedidos}\n"
        word_content += f"  • Total pedidos internacional (UK/US/BR): {total_internacional}\n"
        word_content += f"  • Total pedidos europeu: {total_europeu}\n"
        word_content += f"  • Pedidos pagos: {pedidos_pagos_count}\n"
        word_content += f"  • Pagamentos diretos: {pagamento_direto_count}\n"
        word_content += f"  • Oferta original: {oferta_original_count}\n"
        word_content += f"  • Oferta tamanho 4.5: {oferta_tamanho_45_count}\n"
        word_content += f"  • Oferta portachaves: {oferta_portachaves_count}\n"
        word_content += f"  • Ofertas recusadas: {oferta_recusadas_count}\n"
        word_content += f"  • Utilizadores Ativos Hoje: {usuarios_ativos}\n"
        word_content += f"  • Taxa de conversão: {taxa_conversao:.1f}%\n"
        word_content += f"  • Taxa de recuperação: {taxa_recuperacao:.1f}%\n\n"
        
        # Data de exportação
        from datetime import datetime
        word_content += f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        
        # Criar ficheiro em memória
        word_file = io.BytesIO(word_content.encode('utf-8'))
        word_file.name = "relatorio_admin.docx"
        
        await context.bot.send_document(
            chat_id=query.message.chat_id,
            document=InputFile(word_file),
            caption="📘 *Relatório Admin exportado em Word*",
            parse_mode='Markdown'
        )
        
        await query.edit_message_text("✅ *Relatório Word exportado com sucesso!*", parse_mode='Markdown')
        
    except Exception as e:
        await query.edit_message_text(f"❌ Erro ao exportar Word: {e}")



async def admin_back_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Voltar ao menu principal do admin"""
    query = update.callback_query
    await query.answer()
    
    # Voltar para o admin_command
    await admin_command(update, context)        


#adminoptions

# Handler para o menu de opções
async def btn_options(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Menu de opções - popup"""
    query = update.callback_query
    
    # Mostra os botões inline
    keyboard = [
        [InlineKeyboardButton("Guardar Dados", callback_data="options_save")],
        [InlineKeyboardButton("Recuperar Dados", callback_data="options_restore")],
        [InlineKeyboardButton("Eliminar", callback_data="options_delete")],
        [InlineKeyboardButton("Admin", callback_data="options_edit_admin")],
        [InlineKeyboardButton("Imposto / Frete", callback_data="options_edit_taxes")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await context.bot.send_message(
        chat_id=query.message.chat_id,
        text="*OPÇÕES AVANÇADAS:*",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )




# Handlers para opções
async def options_save_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Guardar dados para persistência"""
    query = update.callback_query
    await query.answer()
    
    try:
        # 🔥 CALCULAR RESUMO FINANCEIRO ATUAL ANTES DE GUARDAR
        pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
        
        # Calcular totais financeiros
        total_bruto = 0.0
        total_impostos = 0.0
        total_frete = 0.0
        total_final = 0.0
        total_internacional_valor = 0.0
        total_europeu_valor = 0.0
        
        for pedido in pedidos_pagos:
            moeda_original = pedido.get('moeda_original', 'EUR')
            total_original = pedido.get('total_pago_original', pedido.get('total', 0))
            
            if 'total_pago_eur' in pedido:
                total_eur = pedido['total_pago_eur']
            else:
                if moeda_original != 'EUR':
                    TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
                    taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
                    taxa = float(taxa_decimal)
                    total_eur = total_original / taxa
                else:
                    total_eur = total_original
            
            # Determinar região
            pais = pedido.get('pais', '').lower()
            if any(pais_internacional in pais for pais_internacional in ['reino unido', 'united kingdom', 'uk', 'estados unidos', 'united states', 'us', 'usa', 'brasil', 'brazil']):
                total_internacional_valor += total_eur
            else:
                total_europeu_valor += total_eur
            
            # Calcular componentes
            if pedido.get('subtotal') and pedido.get('total') and pedido['total'] > 0:
                proporcao = total_eur / pedido['total'] if pedido['total'] > 0 else 1.0
                subtotal_eur = pedido.get('subtotal', 0) * proporcao
                imposto_eur = pedido.get('imposto', 0) * proporcao
                frete_eur = pedido.get('frete', 0) * proporcao
                
                total_bruto += subtotal_eur
                total_impostos += imposto_eur
                total_frete += frete_eur
                total_final += total_eur
            else:
                total_final += total_eur
        
        # 🔥 CALCULAR ESTATÍSTICAS DETALHADAS
        total_pedidos = len(PEDIDOS_REGISTO)
        pedidos_pagos_count = len(pedidos_pagos)
        
        # Ofertas por tipo
        oferta_original_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "original")
        oferta_tamanho_45_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["tamanho_4.5", "oferta_tamanho_45"])
        oferta_portachaves_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") in ["portachaves", "oferta_portachaves"])
        pagamento_direto_count = sum(1 for p in pedidos_pagos if p.get("oferta_tipo") == "pagamento_direto")
        
        # Taxas
        taxa_conversao = (pedidos_pagos_count / total_pedidos) * 100 if total_pedidos > 0 else 0
        taxa_recuperacao = (ESTATISTICAS['tentativas_recuperacao'] / total_pedidos) * 100 if total_pedidos > 0 else 0
        
        # 🔥 DADOS COMPLETOS A GUARDAR
        dados_para_guardar = {
            # Dados principais
            'PEDIDOS_REGISTO': PEDIDOS_REGISTO,
            'ESTATISTICAS': ESTATISTICAS,
            'TAXAS_PAISES': TAXAS_PAISES,
            'CONTADOR_UTILIZADORES': CONTADOR_UTILIZADORES,
            
            # 🔥 RESUMO FINANCEIRO CALCULADO
            'RESUMO_FINANCEIRO': {
                'total_bruto': total_bruto,
                'total_impostos': total_impostos,
                'total_frete': total_frete,
                'total_final': total_final,
                'total_internacional_valor': total_internacional_valor,
                'total_europeu_valor': total_europeu_valor,
                'total_pedidos': total_pedidos,
                'pedidos_pagos_count': pedidos_pagos_count
            },
            
            # 🔥 ESTATÍSTICAS DETALHADAS
            'ESTATISTICAS_DETALHADAS': {
                'oferta_original_count': oferta_original_count,
                'oferta_tamanho_45_count': oferta_tamanho_45_count,
                'oferta_portachaves_count': oferta_portachaves_count,
                'pagamento_direto_count': pagamento_direto_count,
                'oferta_recusadas_count': ESTATISTICAS['ofertas_recusadas'],
                'taxa_conversao': taxa_conversao,
                'taxa_recuperacao': taxa_recuperacao,
                'usuarios_ativos': obter_utilizadores_ativos()
            },
            
            'ultimo_backup': datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            'timestamp': datetime.now().isoformat()
        }
        
        # 🔥 GUARDAR EM JSON (legível)
        with open('backup_dados.json', 'w', encoding='utf-8') as f:
            json.dump(dados_para_guardar, f, indent=4, ensure_ascii=False, default=str)
        
        # 🔥 GUARDAR EM PICKLE (mais eficiente)
        with open('backup_dados.pkl', 'wb') as f:
            pickle.dump(dados_para_guardar, f)
        
        # 🔥 BACKUP DE SEGURANÇA COM TIMESTAMP
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        with open(f'backup_seguranca_{timestamp}.json', 'w', encoding='utf-8') as f:
            json.dump(dados_para_guardar, f, indent=4, ensure_ascii=False, default=str)
        
        print(f"✅ BACKUP COMPLETO: {len(PEDIDOS_REGISTO)} pedidos, €{total_final:.2f} total")
        
        await query.edit_message_text(
            "✅ *Dados guardados com sucesso!*\n\n"
            f"• 📦 Pedidos: {len(PEDIDOS_REGISTO)} ( {pedidos_pagos_count} pagos)\n"
            f"• 💰 Valor total: €{total_final:.2f}\n"
            f"• 📊 Estatísticas: {len(ESTATISTICAS)}\n"
            f"• 🌍 Países: {len(TAXAS_PAISES)}\n"
            f"• 👥 Utilizadores: {CONTADOR_UTILIZADORES}\n"
            f"• ⏰ Backup: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            parse_mode='Markdown'
        )
        
    except Exception as e:
        print(f"❌ ERRO NO BACKUP: {e}")
        await query.edit_message_text(f"❌ Erro ao guardar dados: {e}")

def carregar_dados_backup():
    """Carrega dados do backup quando o bot inicia"""
    global PEDIDOS_REGISTO, ESTATISTICAS, TAXAS_PAISES, CONTADOR_UTILIZADORES
    
    try:
        # Tenta carregar do pickle primeiro (mais rápido)
        with open('backup_dados.pkl', 'rb') as f:
            dados = pickle.load(f)
            
        # 🔥 CARREGAR TODOS OS DADOS
        PEDIDOS_REGISTO = dados.get('PEDIDOS_REGISTO', {})
        ESTATISTICAS = dados.get('ESTATISTICAS', {})
        TAXAS_PAISES = dados.get('TAXAS_PAISES', {})
        
        # 🔥 CARREGAR CONTADOR (pode ser dict ou int)
        contador_backup = dados.get('CONTADOR_UTILIZADORES', 0)
        if isinstance(contador_backup, dict):
            CONTADOR_UTILIZADORES = contador_backup
        else:
            # Se for int, converter para a estrutura dict
            CONTADOR_UTILIZADORES = {
                "contador": contador_backup,
                "utilizadores_unicos": set(),
                "reset_feito_hoje": False,
                "estatisticas_enviadas": False
            }
        
        # 🔥 DADOS FINANCEIROS E ESTATÍSTICAS (para referência)
        resumo_financeiro = dados.get('RESUMO_FINANCEIRO', {})
        estatisticas_detalhadas = dados.get('ESTATISTICAS_DETALHADAS', {})
        
        print(f"✅ DADOS CARREGADOS: {len(PEDIDOS_REGISTO)} pedidos, €{resumo_financeiro.get('total_final', 0):.2f} total")
        print(f"📊 Contador: {CONTADOR_UTILIZADORES}")
        print(f"🕒 Último backup: {dados.get('ultimo_backup', 'N/A')}")
        
    except FileNotFoundError:
        print("ℹ️ Nenhum backup encontrado, iniciando com dados vazios")
    except Exception as e:
        print(f"❌ ERRO AO CARREGAR BACKUP: {e}")



async def options_restore_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Restaurar dados do backup"""
    query = update.callback_query
    await query.answer()
    
    try:
        # 🔥 CARREGAR DADOS
        carregar_dados_backup()
        
        # 🔥 CALCULAR VALORES ATUAIS PARA CONFIRMAR
        pedidos_pagos = [p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"]
        pedidos_pagos_count = len(pedidos_pagos)
        
        # Calcular total financeiro atual
        total_final_atual = 0.0
        for pedido in pedidos_pagos:
            if 'total_pago_eur' in pedido:
                total_final_atual += pedido['total_pago_eur']
            else:
                total_original = pedido.get('total_pago_original', pedido.get('total', 0))
                total_final_atual += total_original
        
        # 🔥 OBTER CONTADOR CORRETAMENTE
        if isinstance(CONTADOR_UTILIZADORES, dict):
            contador_utilizadores = CONTADOR_UTILIZADORES.get("contador", 0)
        else:
            contador_utilizadores = CONTADOR_UTILIZADORES
        
        await query.edit_message_text(
            "✅ *Dados restaurados com sucesso!*\n\n"
            f"• 📦 Total de Pedidos: {len(PEDIDOS_REGISTO)}\n"
            f"• ✅ Pedidos Pagos: {pedidos_pagos_count}\n"
            f"• 💰 Valor Total Atual: €{total_final_atual:.2f}\n"
            f"• 📊 Estatísticas: {len(ESTATISTICAS)} categorias\n"
            f"• 🌍 Países Configurados: {len(TAXAS_PAISES)}\n"
            f"• 👥 Utilizadores Ativos: {contador_utilizadores}\n"
            f"• ⏰ Restaurado: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
            f"🔄 *Atualize o painel admin para ver os dados restaurados*",
            parse_mode='Markdown'
        )
        
        print(f"✅ DADOS RESTAURADOS: {len(PEDIDOS_REGISTO)} pedidos, €{total_final_atual:.2f} total")
        
    except Exception as e:
        print(f"❌ ERRO ao restaurar dados: {e}")
        await query.edit_message_text(f"❌ Erro ao restaurar dados: {e}")




async def options_delete_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Eliminar TODOS os dados - Zerar painel completo"""
    query = update.callback_query
    await query.answer()
    
    try:
        # 🔥 DECLARA GLOBAL PRIMEIRO
        global CONTADOR_UTILIZADORES
        
        # 🔥 GUARDA OS VALORES ANTES DE ELIMINAR (apenas para mostrar)
        pedidos_antes = len(PEDIDOS_REGISTO)
        estatisticas_antes = len(ESTATISTICAS)
        utilizadores_antes = CONTADOR_UTILIZADORES
        
        # 🔥 CALCULA VALORES FINANCEIROS ANTES (para mostrar no resumo)
        pedidos_pagos_antes = len([p for p in PEDIDOS_REGISTO.values() if p.get("status") == "pago"])
        total_final_antes = 0.0
        
        for pedido in PEDIDOS_REGISTO.values():
            if pedido.get("status") == "pago":
                total_original = pedido.get('total_pago_original', pedido.get('total', 0))
                if 'total_pago_eur' in pedido:
                    total_final_antes += pedido['total_pago_eur']
                else:
                    total_final_antes += total_original
        
        # 🔥 ELIMINA TUDO - ZERA COMPLETAMENTE
        # Limpa todos os pedidos (elimina resumo financeiro)
        PEDIDOS_REGISTO.clear()
        
        # Limpa e reinicia estatísticas
        ESTATISTICAS.clear()
        ESTATISTICAS.update({
            'ofertas_recusadas': 0,
            'tentativas_recuperacao': 0,
            'pedidos_abandonados': 0
        })
        
        # Zera contador de utilizadores
        CONTADOR_UTILIZADORES = 0
        
        # 🔥 ELIMINA BACKUPS ANTIGOS (opcional)
        try:
            import os
            if os.path.exists('backup_dados.json'):
                os.remove('backup_dados.json')
            if os.path.exists('backup_dados.pkl'):
                os.remove('backup_dados.pkl')
        except:
            pass  # Ignora erros na eliminação de backups
        
        # 🔥 MENSAGEM DE CONFIRMAÇÃO
        await query.edit_message_text(
            "🗑️ *TODOS OS DADOS ELIMINADOS!*\n\n"
            "✅ *Painel reiniciado com sucesso!*\n\n"
            f"• 📦 Pedidos eliminados: {pedidos_antes} → 0\n"
            f"• 💰 Valor total eliminado: €{total_final_antes:.2f} → €0.00\n"
            f"• ✅ Pedidos pagos eliminados: {pedidos_pagos_antes} → 0\n"
            f"• 📊 Estatísticas zeradas: {estatisticas_antes} → {len(ESTATISTICAS)}\n"
            f"• 👥 Utilizadores: {utilizadores_antes} → 0\n"
            f"• 🌍 Mercado Europeu: €0.00\n"
            f"• 🌎 Mercado Internacional: €0.00\n"
            f"• 🧾 Impostos: €0.00\n"
            f"• 🚚 Fretes: €0.00\n\n"
            f"⏰ Reiniciado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            parse_mode='Markdown'
        )
        
        print(f"🔥 DADOS ELIMINADOS: {pedidos_antes} pedidos, €{total_final_antes:.2f} eliminados")
        
    except Exception as e:
        await query.edit_message_text(f"❌ Erro ao eliminar dados: {e}")









async def options_edit_taxes_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Editar impostos"""
    query = update.callback_query
    await query.answer()
    
    keyboard = [
        [InlineKeyboardButton("📊 Ver Taxas Atuais", callback_data="view_taxes")],
        [InlineKeyboardButton("✏️ Editar Imposto País", callback_data="edit_tax_country")],
        [InlineKeyboardButton("🚚 Editar Frete País", callback_data="edit_frete_country")],
        [InlineKeyboardButton("📊 Ver Fretes Atuais", callback_data="view_frete")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        "💰 *GESTÃO DE IMPOSTOS E FRETES*\n\n"
        "Configurar taxas de IVA e valores de frete:",
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def view_taxes_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Mostra as taxas de imposto atuais de todos os países"""
    query = update.callback_query
    await query.answer()
    
    texto = "💰 *TAXAS DE IMPOSTO ATUAIS*\n\n"
    
    # Ordenar países alfabeticamente para melhor visualização
    paises_ordenados = sorted(TAXAS_PAISES.keys())
    
    for pais in paises_ordenados:
        imposto_percent = TAXAS_PAISES[pais]["imposto"] * 100
        texto += f"• **{pais.replace('_', ' ').title()}**: {imposto_percent:.0f}%\n"
    
    texto += f"\n📊 *Total de países:* {len(TAXAS_PAISES)}"
    
    await query.edit_message_text(texto, parse_mode='Markdown')

async def view_frete_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Mostra os valores de frete atuais de todos os países"""
    query = update.callback_query
    await query.answer()
    
    texto = "🚚 *VALORES DE FRETE ATUAIS*\n\n"
    
    # Ordenar países alfabeticamente para melhor visualização
    paises_ordenados = sorted(TAXAS_PAISES.keys())
    
    for pais in paises_ordenados:
        frete = TAXAS_PAISES[pais]["frete"]
        texto += f"• **{pais.replace('_', ' ').title()}**: €{frete:.2f}\n"
    
    texto += f"\n📊 *Total de países:* {len(TAXAS_PAISES)}"
    
    await query.edit_message_text(texto, parse_mode='Markdown')






#editadminpainel







async def options_edit_admin_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Inicia edição simplificada do painel admin - apenas totais financeiros"""
    query = update.callback_query
    await query.answer()
    
    # Buscar valores atuais das estatísticas (se existirem)
    subtotal_atual = ESTATISTICAS.get('subtotal_manual', 0)
    impostos_atual = ESTATISTICAS.get('impostos_manual', 0)
    frete_atual = ESTATISTICAS.get('frete_manual', 0)
    internacional_atual = ESTATISTICAS.get('internacional_manual', 0)
    europeu_atual = ESTATISTICAS.get('europeu_manual', 0)
    total_final_atual = subtotal_atual + impostos_atual + frete_atual
    
    user_id = query.from_user.id
    chat_id = query.message.chat.id
    
    print(f"🎯 EDITAR_PAINEL_SIMPLIFICADO CHAMADO por user {user_id} no chat {chat_id}")
    
    # 🔥 DELETAR A MENSAGEM DE OPÇÕES ANTERIOR
    try:
        await query.message.delete()
        print("✅ Mensagem de opções deletada")
    except Exception as e:
        print(f"⚠️ Não foi possível deletar a mensagem de opções: {e}")
    
    # Guardar estado de edição
    context.bot_data[f'editing_painel_{chat_id}'] = {
        'tipo': 'painel_admin_simplificado',
        'user_id': user_id,
        'chat_id': chat_id,
        'passo_atual': 'subtotal',
        'valores': {}
    }
    
    # 🔥 ENVIAR PRIMEIRA PERGUNTA E GUARDAR ID
    primeira_pergunta = await context.bot.send_message(
        chat_id=chat_id,
        text=(
            "📝 *EDIÇÃO SIMPLIFICADA DO PAINEL ADMIN*\n\n"
            f"📊 *Valores atuais:*\n"
            f"• 💰 Subtotal: €{subtotal_atual:.2f}\n"
            f"• 🧾 Impostos: €{impostos_atual:.2f}\n"
            f"• 🚚 Frete: €{frete_atual:.2f}\n"
            f"• 🌎 Internacional: €{internacional_atual:.2f}\n"
            f"• 🇪🇺 Europeu: €{europeu_atual:.2f}\n"
            f"• 💵 Total Final: €{total_final_atual:.2f}\n\n"
            "💶 *Digite o novo SUBTOTAL (em EUR):*"
        ),
        parse_mode='Markdown'
    )
    
    # 🔥 GUARDAR ID DA PRIMEIRA PERGUNTA
    context.bot_data[f'editing_painel_{chat_id}']['ultima_pergunta_id'] = primeira_pergunta.message_id

async def processar_edicao_painel_direto(update: Update, context: ContextTypes.DEFAULT_TYPE, editing_data, message):
    """Processa a edição simplificada do painel admin - COM DELETE CORRETO"""
    ADMIN_USER_IDS = os.getenv("ADMIN_USER_IDS")
    if ADMIN_USER_IDS:
        try:

           ADMIN_USER_IDS = [int(id.strip()) for id in ADMIN_USER_IDS.split(",") if id.strip()]
           print(f"✅ Administradores carregados: {len(ADMIN_USER_IDS)} usuários")
        except ValueError:
           print("❌ ERRO: ADMIN_USER_IDS contém valores não numéricos")
      
    else:
        print("⚠️ AVISO: ADMIN_USER_IDS não configurado no .env")
    
    print(f"🎯 PROCESSAR_EDIÇÃO_PAINEL_SIMPLIFICADO CHAMADO!")
    
    user_id = editing_data['user_id']
    if user_id not in ADMIN_USER_IDS:
        await message.reply_text("❌ Acesso negado.")
        return
    
    passo_atual = editing_data['passo_atual']
    valores = editing_data['valores']
    
    try:
        texto = message.text.strip().replace(',', '.')
        valor = float(texto)
        
        print(f"🔍 Passo: {passo_atual}, Valor: {valor}")
        
        # 🔥 DELETAR A MENSAGEM DE RESPOSTA DO UTILIZADOR E A PERGUNTA ANTERIOR
        try:
            # Deletar a pergunta anterior (se existir)
            if 'ultima_pergunta_id' in editing_data:
                await context.bot.delete_message(
                    chat_id=message.chat.id,
                    message_id=editing_data['ultima_pergunta_id']
                )
                print("✅ Pergunta anterior deletada")
            
            # Deletar a resposta do utilizador
            await message.delete()
            print("✅ Resposta do utilizador deletada")
            
        except Exception as e:
            print(f"⚠️ Não foi possível deletar mensagens: {e}")
        
        # Guardar o valor atual
        valores[passo_atual] = valor
        
        # Definir próximo passo (apenas totais financeiros)
        proximos_passos = {
            'subtotal': 'impostos',
            'impostos': 'frete', 
            'frete': 'internacional',
            'internacional': 'europeu',
            'europeu': 'total_final'
        }
        
        if passo_atual in proximos_passos:
            proximo_passo = proximos_passos[passo_atual]
            editing_data['passo_atual'] = proximo_passo
            
            # Atualizar no bot_data
            chat_id = message.chat.id
            context.bot_data[f'editing_painel_{chat_id}'] = editing_data
            
            perguntas = {
                'impostos': "🧾 *Digite o novo valor de IMPOSTOS (em EUR):*",
                'frete': "🚚 *Digite o novo valor de FRETE (em EUR):*", 
                'internacional': "🌎 *Digite o novo TOTAL INTERNACIONAL (em EUR):*",
                'europeu': "🇪🇺 *Digite o novo TOTAL EUROPEU (em EUR):*",
                'total_final': "💵 *Digite o novo TOTAL FINAL (em EUR):*"
            }
            
            # 🔥 ENVIAR NOVA PERGUNTA E GUARDAR ID
            nova_pergunta = await context.bot.send_message(
                chat_id=message.chat.id,
                text=perguntas[proximo_passo],
                parse_mode='Markdown'
            )
            
            # 🔥 GUARDAR ID DA NOVA PERGUNTA PARA DEPOIS DELETAR
            editing_data['ultima_pergunta_id'] = nova_pergunta.message_id
            context.bot_data[f'editing_painel_{chat_id}'] = editing_data
            
        else:
            # Último passo - aplicar todas as alterações
            # 🔥 DELETAR A ÚLTIMA PERGUNTA ANTES DE APLICAR
            try:
                if 'ultima_pergunta_id' in editing_data:
                    await context.bot.delete_message(
                        chat_id=message.chat.id,
                        message_id=editing_data['ultima_pergunta_id']
                    )
                    print("✅ Última pergunta deletada")
            except Exception as e:
                print(f"⚠️ Não foi possível deletar a última pergunta: {e}")
            
            await aplicar_alteracoes_painel_simplificado(update, context, valores, message)
            
    except ValueError:
        # 🔥 DELETAR MENSAGEM INVÁLIDA E PERGUNTA ANTERIOR
        try:
            if 'ultima_pergunta_id' in editing_data:
                await context.bot.delete_message(
                    chat_id=message.chat.id,
                    message_id=editing_data['ultima_pergunta_id']
                )
            await message.delete()
        except:
            pass
        # Enviar mensagem de erro (não vamos deletar esta)
        await context.bot.send_message(
            chat_id=message.chat.id,
            text="❌ Valor inválido. Use apenas números."
        )
    except Exception as e:
        # 🔥 DELETAR MENSAGEM COM ERRO E PERGUNTA ANTERIOR
        try:
            if 'ultima_pergunta_id' in editing_data:
                await context.bot.delete_message(
                    chat_id=message.chat.id,
                    message_id=editing_data['ultima_pergunta_id']
                )
            await message.delete()
        except:
            pass
        # Enviar mensagem de erro (não vamos deletar esta)
        await context.bot.send_message(
            chat_id=message.chat.id,
            text=f"❌ Erro: {e}"
        )



async def aplicar_alteracoes_painel_simplificado(update: Update, context: ContextTypes.DEFAULT_TYPE, valores, message):
    """Aplica alterações apenas nos totais financeiros - COM DELETE FINAL"""
    try:
        # 🔥 VALORES FINANCEIROS
        novo_subtotal = float(valores.get('subtotal', 0))
        novo_impostos = float(valores.get('impostos', 0))
        novo_frete = float(valores.get('frete', 0))
        novo_internacional = float(valores.get('internacional', 0))
        novo_europeu = float(valores.get('europeu', 0))
        novo_total_final = float(valores.get('total_final', 0))
        
        print(f"🔍 Aplicando valores manuais como BASE:")
        print(f"  - Subtotal Base: €{novo_subtotal:.2f}")
        print(f"  - Impostos Base: €{novo_impostos:.2f}")
        print(f"  - Frete Base: €{novo_frete:.2f}")
        print(f"  - Internacional Base: €{novo_internacional:.2f}")
        print(f"  - Europeu Base: €{novo_europeu:.2f}")
        print(f"  - Total Final Base: €{novo_total_final:.2f}")
        
        # 🔥 GUARDAR VALORES BASE MANUAIS
        ESTATISTICAS['subtotal_manual'] = novo_subtotal
        ESTATISTICAS['impostos_manual'] = novo_impostos
        ESTATISTICAS['frete_manual'] = novo_frete
        ESTATISTICAS['internacional_manual'] = novo_internacional
        ESTATISTICAS['europeu_manual'] = novo_europeu
        ESTATISTICAS['total_final_manual'] = novo_total_final
        
        # 🔥 MARCAR QUE USAMOS SISTEMA MANUAL + AUTOMÁTICO
        ESTATISTICAS['usar_valores_manuais'] = True
        
        print("✅ Valores manuais definidos como BASE")
        
        # 🔥 LIMPAR ESTADO
        chat_id = message.chat.id
        del context.bot_data[f'editing_painel_{chat_id}']
        
        # Salvar backup automaticamente
        salvar_dados_backup()
        
        # 🔥 DELETAR A MENSAGEM DE RESPOSTA FINAL DO UTILIZADOR
        try:
            await message.delete()
            print("✅ Resposta final do utilizador deletada")
        except Exception as e:
            print(f"⚠️ Não foi possível deletar a resposta final: {e}")
        
        mensagem = (
            f"✅ *Valores Manuais Definidos como BASE!*\n\n"
            f"📊 *Valores Base Aplicados:*\n"
            f"• 💰 Subtotal Base: €{novo_subtotal:.2f}\n"
            f"• 🧾 Impostos Base: €{novo_impostos:.2f}\n"
            f"• 🚚 Frete Base: €{novo_frete:.2f}\n"
            f"• 🌎 Internacional Base: €{novo_internacional:.2f}\n"
            f"• 🇪🇺 Europeu Base: €{novo_europeu:.2f}\n"
            f"• 💵 Total Final Base: €{novo_total_final:.2f}\n\n"
            f"📈 *Sistema Ativo: MANUAL + AUTOMÁTICO*\n"
            f"• Estes valores são a BASE\n"
            f"• As vendas automáticas serão SOMADAS a esta base\n"
            f"• 💾 Backup realizado automaticamente"
        )
        
        # 🔥 ENVIAR MENSAGEM FINAL (esta não será deletada)
        await context.bot.send_message(
            chat_id=message.chat.id,
            text=mensagem,
            parse_mode='Markdown'
        )
        print(f"✅ Sistema MANUAL + AUTOMÁTICO ativado!")
        
    except Exception as e:
        print(f"❌ Erro ao aplicar valores manuais: {e}")
        import traceback
        traceback.print_exc()
        await context.bot.send_message(
            chat_id=message.chat.id,
            text=f"❌ Erro ao aplicar valores: {str(e)}"
        )












#editfreteadmin


async def edit_frete_country_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Menu para selecionar país para editar frete"""
    query = update.callback_query
    await query.answer()
    
    keyboard = []
    # Criar botões para cada país
    for pais in sorted(TAXAS_PAISES.keys()):
        nome_pais = pais.replace('_', ' ').title()
        frete_atual = TAXAS_PAISES[pais]["frete"]
        keyboard.append([InlineKeyboardButton(f"{nome_pais} (€{frete_atual:.2f})", callback_data=f"edit_frete_{pais}")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        "✏️ *EDITAR FRETE DE PAÍS*\n\n"
        "Selecione o país que deseja editar:",
        parse_mode='Markdown',
        reply_markup=reply_markup
    )

async def edit_frete_pais_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Inicia edição do frete de um país específico"""
    query = update.callback_query
    await query.answer()
    
    # Extrair o país do callback_data (edit_frete_portugal -> portugal)
    pais = query.data.replace("edit_frete_", "")
    taxas_atual = TAXAS_PAISES[pais]
    frete_atual = taxas_atual["frete"]
    
    user_id = query.from_user.id
    chat_id = query.message.chat.id
    
    print(f"🎯 EDIT_FRETE_PAIS_HANDLER CHAMADO: {pais} por user {user_id} no chat {chat_id}")
    
    # 🔥 GUARDAR COM CHAT_ID
    context.bot_data[f'editing_frete_{chat_id}'] = {
        'pais': pais,
        'tipo': 'frete',
        'user_id': user_id,
        'chat_id': chat_id
    }
    
    print(f"🔍 bot_data guardado: editing_frete_{chat_id} = {context.bot_data.get(f'editing_frete_{chat_id}')}")
    
    await query.edit_message_text(
        f"✏️ *EDITANDO FRETE - {pais.replace('_', ' ').title()}*\n\n"
        f"🚚 *Frete atual:* €{frete_atual:.2f}\n\n"
        "💶 *Digite o novo valor de frete (em EUR):*\n"
        "🔹 *Exemplo:* `15.50` para €15.50\n"
        "🔹 *Exemplo:* `0` para frete grátis",
        parse_mode='Markdown'
    )

async def processar_edicao_frete_direto(update: Update, context: ContextTypes.DEFAULT_TYPE, editing_data, message):
    """Processa a edição de fretes quando detetada no handle_message"""
    ADMIN_USER_IDS = os.getenv("ADMIN_USER_IDS")
    if ADMIN_USER_IDS:
        try:
           ADMIN_USER_IDS = [int(id.strip()) for id in ADMIN_USER_IDS.split(",") if id.strip()]
           print(f"✅ Administradores carregados: {len(ADMIN_USER_IDS)} usuários")
        except ValueError:
           print("❌ ERRO: ADMIN_USER_IDS contém valores não numéricos")
      
    else:
        print("⚠️ AVISO: ADMIN_USER_IDS não configurado no .env")
    
    print(f"🎯 PROCESSAR_EDIÇÃO_FRETE_DIRETO CHAMADO!")
    
    user_id = editing_data['user_id']
    if user_id not in ADMIN_USER_IDS:
        await message.reply_text("❌ Acesso negado.")
        return
    
    pais = editing_data['pais']
    print(f"🔍 Editando frete para: {pais}")
    print(f"🔍 Texto recebido: '{message.text}'")
    
    try:
        texto = message.text.strip().replace(',', '.')
        novo_frete = float(texto)
        
        print(f"🔍 Novo frete convertido: {novo_frete}")
        
        # Validar o valor (não pode ser negativo)
        if novo_frete < 0:
            await message.reply_text("❌ Valor inválido. O frete não pode ser negativo.")
            return
        
        # Atualizar o frete
        TAXAS_PAISES[pais]["frete"] = novo_frete
        
        # 🔥 LIMPAR ESTADO NO bot_data
        chat_id = message.chat.id
        del context.bot_data[f'editing_frete_{chat_id}']
        
        # Salvar backup automaticamente
        salvar_dados_backup()
        
        mensagem = (
            f"✅ *Frete atualizado com sucesso!*\n\n"
            f"**País:** {pais.replace('_', ' ').title()}\n"
            f"**Novo frete:** €{novo_frete:.2f}\n\n"
            f"📊 As alterações foram guardadas automaticamente."
        )
        
        await message.reply_text(mensagem, parse_mode='Markdown')
        print(f"✅ Frete atualizado: {pais} -> €{novo_frete:.2f}")
        
    except ValueError:
        await message.reply_text(
            "❌ Valor inválido. Use apenas números.\n"
            "🔹 *Exemplo:* `15.50` para €15.50"
        )
    except Exception as e:
        await message.reply_text(f"❌ Erro: {e}")
















#editaxasadmin




async def edit_tax_country_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Menu para selecionar país para editar imposto"""
    query = update.callback_query
    await query.answer()
    
    keyboard = []
    # Criar botões para cada país
    for pais in sorted(TAXAS_PAISES.keys()):
        nome_pais = pais.replace('_', ' ').title()
        imposto_atual = TAXAS_PAISES[pais]["imposto"] * 100
        keyboard.append([InlineKeyboardButton(f"{nome_pais} ({imposto_atual:.0f}%)", callback_data=f"edit_tax_{pais}")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.edit_message_text(
        "✏️ *EDITAR IMPOSTO DE PAÍS*\n\n"
        "Selecione o país que deseja editar:",
        parse_mode='Markdown',
        reply_markup=reply_markup
    )

async def edit_tax_pais_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Inicia edição do imposto de um país específico"""
    query = update.callback_query
    await query.answer()
    
    pais = query.data.replace("edit_tax_", "")
    taxas_atual = TAXAS_PAISES[pais]
    imposto_atual = taxas_atual["imposto"] * 100
    
    user_id = query.from_user.id
    chat_id = query.message.chat.id  # 🔥 PEGAR O CHAT_ID CORRETO
    
    print(f"🎯 EDIT_TAX_PAIS_HANDLER CHAMADO: {pais} por user {user_id} no chat {chat_id}")
    
    # 🔥 GUARDAR COM CHAT_ID (IMPORTANTE!)
    context.bot_data[f'editing_tax_{chat_id}'] = {
        'pais': pais,
        'tipo': 'imposto',
        'user_id': user_id,
        'chat_id': chat_id
    }
    
    print(f"🔍 bot_data guardado: editing_tax_{chat_id} = {context.bot_data.get(f'editing_tax_{chat_id}')}")
    
    await query.edit_message_text(
        f"✏️ *EDITANDO IMPOSTO - {pais.replace('_', ' ').title()}*\n\n"
        f"💰 *Imposto atual:* {imposto_atual:.0f}%\n\n"
        "💶 *Digite o novo valor de imposto (em %):*\n"
        "🔹 *Exemplo:* `23` para 23%\n"
        "🔹 *Exemplo:* `0` para 0% (isenção)",
        parse_mode='Markdown'
    )



def salvar_dados_backup():
    """Salva os dados atuais em backup"""
    try:
        dados = {
            'PEDIDOS_REGISTO': PEDIDOS_REGISTO,
            'ESTATISTICAS': ESTATISTICAS,
            'TAXAS_PAISES': TAXAS_PAISES,
            'CONTADOR_UTILIZADORES': CONTADOR_UTILIZADORES,
            'ultimo_backup': datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        }
        
        with open('backup_dados.pkl', 'wb') as f:
            pickle.dump(dados, f)
        
        print(f"✅ Backup salvo: {len(PEDIDOS_REGISTO)} pedidos, {len(TAXAS_PAISES)} países")
        return True
        
    except Exception as e:
        print(f"❌ Erro ao salvar backup: {e}")
        return False



async def processar_edicao_imposto_direto(update: Update, context: ContextTypes.DEFAULT_TYPE, editing_data, message):
    """Processa a edição de impostos quando detetada no handle_message"""
    ADMIN_USER_IDS = os.getenv("ADMIN_USER_IDS")
    if ADMIN_USER_IDS:
        try:

           ADMIN_USER_IDS = [int(id.strip()) for id in ADMIN_USER_IDS.split(",") if id.strip()]
           print(f"✅ Administradores carregados: {len(ADMIN_USER_IDS)} usuários")
        except ValueError:
           print("❌ ERRO: ADMIN_USER_IDS contém valores não numéricos")
      
    else:
        print("⚠️ AVISO: ADMIN_USER_IDS não configurado no .env")
    
    print(f"🎯 PROCESSAR_EDIÇÃO_IMPOSTO_DIRETO CHAMADO!")
    
    user_id = editing_data['user_id']  # 🔥 PEGAR O USER_ID DO editing_data
    if user_id not in ADMIN_USER_IDS:
        await message.reply_text("❌ Acesso negado.")
        return
    
    pais = editing_data['pais']
    print(f"🔍 Editando imposto para: {pais}")
    print(f"🔍 Texto recebido: '{message.text}'")
    
    try:
        texto = message.text.strip().replace(',', '.')
        novo_imposto = float(texto)
        
        print(f"🔍 Novo imposto convertido: {novo_imposto}")
        
        # Validar o valor (entre 0% e 100%)
        if novo_imposto < 0 or novo_imposto > 100:
            await message.reply_text("❌ Valor inválido. O imposto deve estar entre 0% e 100%.")
            return
        
        # Converter porcentagem para decimal (23% -> 0.23)
        TAXAS_PAISES[pais]["imposto"] = novo_imposto / 100
        
        # 🔥 LIMPAR ESTADO NO bot_data USANDO CHAT_ID
        chat_id = message.chat.id
        del context.bot_data[f'editing_tax_{chat_id}']
        
        # Salvar backup automaticamente
        salvar_dados_backup()
        
        mensagem = (
            f"✅ *Imposto atualizado com sucesso!*\n\n"
            f"**País:** {pais.replace('_', ' ').title()}\n"
            f"**Novo imposto:** {novo_imposto:.0f}%\n\n"
            f"📊 As alterações foram guardadas automaticamente."
        )
        
        await message.reply_text(mensagem, parse_mode='Markdown')
        print(f"✅ Imposto atualizado: {pais} -> {novo_imposto}%")
        
    except ValueError:
        await message.reply_text(
            "❌ Valor inválido. Use apenas números.\n"
            "🔹 *Exemplo:* `23` para 23% ou `0` para 0%"
        )
    except Exception as e:
        await message.reply_text(f"❌ Erro: {e}")





CALLBACK_HANDLERS = {
    "view_taxes": view_taxes_handler,
    "view_frete": view_frete_handler,
    "edit_tax_country": edit_tax_country_handler,
    "edit_frete_country": edit_frete_country_handler,
    "options_edit_admin": options_edit_admin_handler,
    # ... outros handlers que já tens
}


for pais in TAXAS_PAISES.keys():
    CALLBACK_HANDLERS[f"edit_tax_{pais}"] = edit_tax_pais_handler
    CALLBACK_HANDLERS[f"edit_frete_{pais}"] = edit_frete_pais_handler
    print(f"✅ Handler registrado: edit_frete_{pais}")

print(f"✅ Total de handlers de fretes registrados: {len([k for k in CALLBACK_HANDLERS.keys() if 'edit_frete' in k])}")




print("✅ Todos os handlers de edição de impostos registrados!")

#editaradmin



async def handle_callback_query(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler principal para callback queries"""
    query = update.callback_query
    await query.answer()
    
    callback_data = query.data
    print(f"🔍 Callback recebido: {callback_data}")
    
    # Procura o handler no dicionário
    handler = CALLBACK_HANDLERS.get(callback_data)
    
    if handler:
        print(f"🎯 Handler encontrado: {handler.__name__}")
        await handler(update, context)
    else:
        print(f"❌ Handler não encontrado para: {callback_data}")
        await query.edit_message_text("❌ Comando não reconhecido.")




# --- Função para mostrar detalhes completos do pedido ---
async def show_pedido_details(update: Update, context: ContextTypes.DEFAULT_TYPE, pedido_id: str):
    """Mostra detalhes completos de um pedido específico"""
    query = update.callback_query
    await query.answer()
    
    # Encontrar o pedido
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    if not pedido:
        await query.edit_message_text("❌ Pedido não encontrado.")
        return
    
    # 🔥 CONSTRUIR MENSAGEM COMPLETA (igual à função enviar_pedido_pago_para_admin)
    mensagem_detalhes = f"""
🎉 *DETALHES COMPLETOS DO PEDIDO*

🆔 *Pedido:* #{pedido['id']}
📅 *Data do Pedido:* {pedido.get('data', 'N/A')}
💳 *Data do Pagamento:* {pedido.get('data_pagamento', 'N/A')}
💰 *Valor Pago:* {pedido.get('simbolo_moeda', '€')}{pedido.get('total', 0):.2f} {pedido.get('moeda', 'EUR')}
💬 *Chat ID do Cliente:* {pedido.get('chat_id', 'N/A')}

👤 *DADOS PESSOAIS:*
• *Nome:* {pedido.get('nome', 'N/A')}
• *Email:* {pedido.get('email', 'N/A')}
• *País:* {pedido.get('pais', 'N/A')}
• *Contacto:* {pedido.get('contacto', 'N/A')}

🎨 *DETALHES DO CARTOON:*
• *Tipo:* {pedido.get('tipo_cartoon', 'N/A')}
• *Estilo:* {pedido.get('estilo_cartoon', 'N/A')}
• *Tamanho:* {pedido.get('tamanho_cartoon', 'N/A')}"""

    # 🔥 ADICIONAR TODOS OS CAMPOS PERSONALIZADOS
    if pedido.get('nome_family'):
        mensagem_detalhes += f"\n• *Nome da Família:* {pedido['nome_family']}"
    if pedido.get('frase_family') and pedido['frase_family'] != "Não adicionou frase":
        mensagem_detalhes += f"\n• *Frase da Família:* \"{pedido['frase_family']}\""
    
    if pedido.get('tipo_personalizado'):
        mensagem_detalhes += f"\n• *Tipo de Peça:* {pedido['tipo_personalizado']}"
    if pedido.get('nome_peca_personalizado'):
        mensagem_detalhes += f"\n• *Nome da Peça:* {pedido['nome_peca_personalizado']}"
    if pedido.get('nome_personalizado'):
        mensagem_detalhes += f"\n• *Nome do Cartoon:* {pedido['nome_personalizado']}"
    if pedido.get('frase_personalizado') and pedido['frase_personalizado'] != "Não adicionou frase":
        mensagem_detalhes += f"\n• *Frase do Elemento:* \"{pedido['frase_personalizado']}\""
    
    if pedido.get('nome_cartoon'):
        mensagem_detalhes += f"\n• *Nome no Cartoon:* {pedido['nome_cartoon']}"
    if pedido.get('frase_cartoon') and pedido['frase_cartoon'] != "Não adicionou frase":
        mensagem_detalhes += f"\n• *Frase na Box:* \"{pedido['frase_cartoon']}\""
    
    if pedido.get('profissao'):
        mensagem_detalhes += f"\n• *Profissão:* {pedido['profissao']}"
    if pedido.get('objetos_office'):
        mensagem_detalhes += f"\n• *Objetos Personalizados:* {pedido['objetos_office']}"
    if pedido.get('super_heroi'):
        mensagem_detalhes += f"\n• *Super-Herói:* {pedido['super_heroi']}"
    
    if pedido.get('elementos_family'):
        mensagem_detalhes += f"\n• *Total de Elementos:* {pedido['elementos_family']}"
    if pedido.get('adultos_family'):
        mensagem_detalhes += f"\n• *Adultos:* {pedido['adultos_family']}"
    if pedido.get('criancas_family'):
        mensagem_detalhes += f"\n• *Crianças:* {pedido['criancas_family']}"
    if pedido.get('animais_family'):
        mensagem_detalhes += f"\n• *Animais:* {pedido['animais_family']}"
    if pedido.get('nome_animal'):
        mensagem_detalhes += f"\n• *Nome do Animal:* {pedido['nome_animal']}"
    if pedido.get('tipo_animal'):
        mensagem_detalhes += f"\n• *Tipo de Animal:* {pedido['tipo_animal']}"

    # 🔥 DETALHES FINANCEIROS
    mensagem_detalhes += f"\n\n💵 *DETALHES FINANCEIROS:*"
    mensagem_detalhes += f"\n• *Subtotal:* €{pedido.get('subtotal', 0):.2f}"
    mensagem_detalhes += f"\n• *Imposto ({pedido.get('taxa_imposto', 0)}%):* €{pedido.get('imposto', 0):.2f}"
    mensagem_detalhes += f"\n• *Frete:* €{pedido.get('frete', 0):.2f}"
    mensagem_detalhes += f"\n• *Total Final:* €{pedido.get('total', 0):.2f}"

    # Botão para voltar à lista
    keyboard = [
        [InlineKeyboardButton("⬅️ Voltar à Lista", callback_data="admin_back_to_list")],
        [InlineKeyboardButton("🔄 Atualizar", callback_data=f"admin_details_{pedido_id}")],
        [InlineKeyboardButton("📧 Contactar Cliente", callback_data=f"admin_contact_{pedido_id}")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Se existir foto, enviar separadamente
    if pedido.get('foto_id'):
        try:
            # Primeiro enviar a foto
            await context.bot.send_photo(
                chat_id=query.message.chat_id,
                photo=pedido['foto_id'],
                caption=f"📸 *Foto do Pedido #{pedido_id}*",
                parse_mode="Markdown"
            )
            # Depois enviar os detalhes
            await query.edit_message_text(
                text=mensagem_detalhes,
                parse_mode="Markdown",
                reply_markup=reply_markup
            )
        except Exception as e:
            # Se der erro na foto, enviar só os detalhes
            mensagem_detalhes += f"\n\n❌ *Erro ao carregar foto:* {e}"
            await query.edit_message_text(
                text=mensagem_detalhes,
                parse_mode="Markdown",
                reply_markup=reply_markup
            )
    else:
        # Se não tem foto, só enviar os detalhes
        await query.edit_message_text(
            text=mensagem_detalhes,
            parse_mode="Markdown",
            reply_markup=reply_markup
        )








async def atualizar_pedido(update: Update, context: ContextTypes.DEFAULT_TYPE, pedido_id):
    """Atualiza a mensagem do pedido"""
    query = update.callback_query
    await query.answer("🔄 Pedido atualizado!", show_alert=True)



# --- Handler para comando admin ---
async def admin_command_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para o comando /admin"""
    pagina = int(context.args[0]) if context.args and context.args[0].isdigit() else 1
    await admin_command(update, context, pagina)






async def enviar_pedido_para_canal_requests(pedido):
    """Envia o pedido COMPLETO para o canal de requests com MOEDAS CORRETAS"""
    try:
        print(f"🎯 INICIANDO enviar_pedido_para_canal_requests - Pedido #{pedido['id']}")
        
        # 🔥 OBTER INFORMAÇÕES DE MOEDA
        moeda_original = pedido.get('moeda_original', 'EUR')
        simbolo_original = get_simbolo_moeda(moeda_original.lower())
        total_original = pedido.get('total_pago_original', pedido.get('total', 0))
        total_eur = pedido.get('total_pago_eur', pedido.get('total', 0))
        
        # 🔥 OBTER TAXA DE IMPOSTO CORRETA (com fallback) - FORMATO INTEIRO
        taxa_imposto = pedido.get('taxa_imposto', 0)
        percentual_imposto = int(taxa_imposto * 100)  # 🔥 CONVERTER PARA INTEIRO
        
        print(f"💰 MOEDAS - Original: {simbolo_original}{total_original:.2f} {moeda_original} | EUR: €{total_eur:.2f}")
        print(f"💰 IMPOSTO - Taxa: {percentual_imposto}%")

        # 🔥 CONSTRUIR MENSAGEM DETALHADA
        mensagem = f"""
🎉 *NOVO PEDIDO PAGO - #{pedido['id']}*

📅 *Data do Pedido:* {pedido.get('data', 'N/A')}
💳 *Data do Pagamento:* {pedido.get('data_pagamento', 'N/A')}
💰 *Valor Pago:* {simbolo_original}{total_original:.2f} {moeda_original} (€{total_eur:.2f})

👤 *DADOS PESSOAIS:*
• *Nome:* {pedido.get('nome', 'N/A')}
• *Email:* {pedido.get('email', 'N/A')}
• *País:* {pedido.get('pais', 'N/A')}
• *Contacto:* {pedido.get('contacto', 'N/A')}
• *Chat ID:* `{pedido.get('chat_id', 'N/A')}`"""

        # 🔥 ADICIONAR DADOS DE SHIPPING SE EXISTIREM
        if pedido.get('shipping_details'):
            shipping = pedido['shipping_details']
            address = shipping.get('address', {})
            name = shipping.get('name', 'N/A')
            
            mensagem += f"\n\n🚚 *DADOS DE ENVIO:*"
            mensagem += f"\n• *Nome:* {name}"
            mensagem += f"\n• *Linha 1:* {address.get('line1', 'N/A')}"
            if address.get('line2'):
                mensagem += f"\n• *Linha 2:* {address.get('line2')}"
            mensagem += f"\n• *Cidade:* {address.get('city', 'N/A')}"
            mensagem += f"\n• *Código Postal:* {address.get('postal_code', 'N/A')}"
            mensagem += f"\n• *País:* {address.get('country', 'N/A')}"

        mensagem += f"\n\n🎨 *DETALHES DO CARTOON:*"
        mensagem += f"\n• *Tipo:* {pedido.get('tipo_cartoon', 'N/A')}"
        
        # 🔥 ESTILO SÓ APARECE PARA CARTOON INDIVIDUAL
        tipo_cartoon = pedido.get('tipo_cartoon', '').lower()
        if 'individual' in tipo_cartoon and pedido.get('estilo_cartoon'):
            mensagem += f"\n• *Estilo:* {pedido['estilo_cartoon']}"
        
        mensagem += f"\n• *Tamanho:* {pedido.get('tamanho_cartoon', 'N/A')}"

        # 🔥 ADICIONAR CAMPOS PERSONALIZADOS ESPECÍFICOS
        campos_personalizados = [
            # Campos da Family
            ('nome_family', '👨‍👩‍👧‍👦 *Nome da Família:*'),
            ('frase_family', '💬 *Frase da Família:*'),
            ('elementos_family', '👥 *Total de Elementos:*'),
            ('adultos_family', '👨‍👩 *Adultos:*'),
            ('criancas_family', '👧🧒 *Crianças:*'),
            ('animais_family', '🐱🐶 *Animais:*'),
            ('nome_animal', '🐾 *Nome do Animal:*'),
            ('tipo_animal', '🐕 *Tipo de Animal:*'),
            
            # Campos do Personalizado
            ('tipo_personalizado', '📦 *Tipo de Peça:*'),
            ('nome_peca_personalizado', '📝 *Nome da Peça:*'),
            ('nome_personalizado', '🎭 *Nome do Cartoon:*'),
            ('frase_personalizado', '💬 *Frase do Elemento:*'),
            
            # Campos de Personalização da Box
            ('nome_cartoon', '🎭 *Nome no Cartoon:*'),
            ('frase_cartoon', '💬 *Frase na Box:*'),
            
            # Campos do Office/Profissional
            ('profissao', '💼 *Profissão:*'),
            ('objetos_office', '🎯 *Objetos Personalizados:*'),
            ('super_heroi', '🦸‍♂️ *Super-Herói:*')
        ]
        
        for campo, label in campos_personalizados:
            if pedido.get(campo) and pedido[campo] != "Não adicionou frase":
                if "frase" in campo:
                    mensagem += f"\n• {label} \"{pedido[campo]}\""
                else:
                    mensagem += f"\n• {label} {pedido[campo]}"

        # 🔥 DETALHES FINANCEIROS NA MOEDA ORIGINAL + EUR
        mensagem += f"\n\n💵 *DETALHES FINANCEIROS:*"
        
        # 🔥 CORREÇÃO: USAR OS VALORES JÁ CONVERTIDOS DO PEDIDO
        if moeda_original != 'EUR':
            # 🔥 OS VALORES JÁ ESTÃO CONVERTIDOS - APENAS MOSTRAR
            subtotal_original = pedido.get('subtotal', 0)
            imposto_original = pedido.get('imposto', 0)
            frete_original = pedido.get('frete', 0)
            
            # 🔥 OBTER OS VALORES EM EUR DO PEDIDO (já calculados)
            subtotal_eur = pedido.get('subtotal_eur', pedido.get('subtotal', 0))
            imposto_eur = pedido.get('imposto_eur', pedido.get('imposto', 0))
            frete_eur = pedido.get('frete_eur', pedido.get('frete', 0))
            
            print(f"🔧 MOSTRANDO DETALHES FINANCEIROS - Valores já convertidos:")
            print(f"   • Subtotal: {simbolo_original}{subtotal_original:.2f} (€{subtotal_eur:.2f})")
            print(f"   • Imposto ({percentual_imposto}%): {simbolo_original}{imposto_original:.2f} (€{imposto_eur:.2f})")
            print(f"   • Frete: {simbolo_original}{frete_original:.2f} (€{frete_eur:.2f})")
            print(f"   • Total: {simbolo_original}{total_original:.2f} (€{total_eur:.2f})")
            
            mensagem += f"\n• *Subtotal:* {simbolo_original}{subtotal_original:.2f} (€{subtotal_eur:.2f})"
            
            # 🔥 CORREÇÃO: MOSTRAR IMPOSTO COM PERCENTUAL CORRETO (SEM DECIMAIS)
            if imposto_original > 0:
                mensagem += f"\n• *Imposto ({percentual_imposto}%):* {simbolo_original}{imposto_original:.2f} (€{imposto_eur:.2f})"
            else:
                mensagem += f"\n• *Imposto:* {simbolo_original}{imposto_original:.2f} (€{imposto_eur:.2f})"
                
            mensagem += f"\n• *Frete:* {simbolo_original}{frete_original:.2f} (€{frete_eur:.2f})"
            mensagem += f"\n• *Total Final:* {simbolo_original}{total_original:.2f} (€{total_eur:.2f})"
        else:
            # Para EUR, mostrar apenas valores em EUR
            subtotal_eur = pedido.get('subtotal', 0)
            imposto_eur = pedido.get('imposto', 0)
            frete_eur = pedido.get('frete', 0)
            
            mensagem += f"\n• *Subtotal:* €{subtotal_eur:.2f}"
            
            # 🔥 CORREÇÃO: MOSTRAR IMPOSTO COM PERCENTUAL CORRETO (SEM DECIMAIS)
            if imposto_eur > 0:
                mensagem += f"\n• *Imposto ({percentual_imposto}%):* €{imposto_eur:.2f}"
            else:
                mensagem += f"\n• *Imposto:* €{imposto_eur:.2f}"
                
            mensagem += f"\n• *Frete:* €{frete_eur:.2f}"
            mensagem += f"\n• *Total Final:* €{total_original:.2f}"

        

        # 🔥 BOTÕES DE AÇÃO
        keyboard = [
            [
                InlineKeyboardButton("📞 Contactar", 
                                   url=f"tg://user?id={pedido['chat_id']}"),
                InlineKeyboardButton("✅ Marcar como Feito", 
                                   callback_data=f"done_{pedido['id']}")
            ],
            [
                InlineKeyboardButton("📊 Exportar CSV", 
                                   callback_data=f"export_ccsv_{pedido['id']}"),
                InlineKeyboardButton("📝 Exportar TXT", 
                                   callback_data=f"export_txt_photo_{pedido['id']}")
            ],
            [               
                InlineKeyboardButton("📄 Exportar PDF", 
                                   callback_data=f"export_pdf_photo_{pedido['id']}"),
                InlineKeyboardButton("📝 Exportar Word", 
                                   callback_data=f"export_word_{pedido['id']}"),                   
            ]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)

        # 🔥 ENVIAR COM FOTO SE EXISTIR
        if pedido.get('foto_id') and pedido['foto_id'] is not None:
            try:
                print("📸 Enviando com FOTO e MOEDAS CORRETAS...")
                await bot.send_photo(
                    chat_id=CANAL_REQUESTS,
                    photo=pedido['foto_id'],
                    caption=mensagem[:1024],  # Limite do Telegram
                    parse_mode="Markdown",
                    reply_markup=reply_markup
                )
                print("✅ Foto com moedas enviada!")
                
                # Se a mensagem for muito longa
                if len(mensagem) > 1024:
                    await bot.send_message(
                        chat_id=CANAL_REQUESTS,
                        text=mensagem[1024:],
                        parse_mode="Markdown"
                    )
                    
            except Exception as e:
                print(f"❌ Erro ao enviar com foto: {e}")
                # Fallback
                await bot.send_message(
                    chat_id=CANAL_REQUESTS,
                    text=mensagem,
                    parse_mode="Markdown",
                    reply_markup=reply_markup
                )
        else:
            # Sem foto
            await bot.send_message(
                chat_id=CANAL_REQUESTS,
                text=mensagem,
                parse_mode="Markdown",
                reply_markup=reply_markup
            )
        
        print(f"🎉 ENVIO COMPLETO para pedido #{pedido['id']} com MOEDAS CORRETAS")
        
    except Exception as e:
        print(f"❌ ERRO CRÍTICO: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")












async def processar_pagamento_sucesso(pedido_id, chat_id, amount, shipping_details=None, moeda_original=None, tipo_sessao=None):
    """Função comum para processar pagamentos bem-sucedidos - COM TIPO_SESSAO E CONVERSÃO CORRETA"""
    print(f"🔍 INICIANDO processar_pagamento_sucesso")
    print(f"🔍 Pedido ID: {pedido_id}")
    print(f"🔍 Chat ID: {chat_id}")
    print(f"🔍 Amount Original: {amount} {moeda_original}")
    print(f"🔍 Tipo Sessão Recebido: {tipo_sessao}")
    print(f"🔍 Tipo: {type(tipo_sessao)}")
    print(f"🔍 Shipping Details: {shipping_details}")
    
    if not pedido_id or not chat_id:
        print("❌ Dados incompletos no metadata")
        return "OK", 200
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: SE NÃO EXISTIR NO PEDIDOS_REGISTO, CRIAR AGORA (SÓ QUANDO PAGO)
    if pedido_id not in PEDIDOS_REGISTO:
        print(f"📦 CRIANDO PEDIDO #{pedido_id} NO REGISTRO (primeiro pagamento)")
        
        # 🔥 RECRIAR PEDIDO COM DADOS BÁSICOS DO PAGAMENTO
        PEDIDOS_REGISTO[pedido_id] = {
            "id": pedido_id,
            "chat_id": chat_id,
            "status": "pendente",  # Será atualizado para pago abaixo
            "data": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
            "total_pago_original": amount,
            "moeda_original": moeda_original or 'EUR',
            "tipo_sessao": tipo_sessao or "pagamento_direto",
            "shipping_details": shipping_details
        }
        print(f"✅ Pedido #{pedido_id} criado no registro")
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    
    print(f"🔍 Pedido encontrado/criado: #{pedido_id}")
    print(f"🔍 PEDIDO ANTES DO PROCESSAMENTO:")
    print(f"   • Oferta Tipo no Pedido: {pedido.get('oferta_tipo', 'NÃO DEFINIDO')}")
    print(f"   • Tipo Cartoon: {pedido.get('tipo_cartoon', 'NÃO DEFINIDO')}")
    print(f"   • Tipo Sessão Recebido: {tipo_sessao}")
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: DETERMINAR O TIPO DE OFERTA CORRETAMENTE
    def determinar_tipo_oferta(pedido, tipo_sessao_recebido):
        """Determina o tipo de oferta baseado em múltiplas fontes"""
        
        print(f"🔍 DETERMINAR_TIPO_OFERTA - DEBUG INICIAL:")
        print(f"   • Tipo Sessão Recebido: {tipo_sessao_recebido}")
        print(f"   • Tipo: {type(tipo_sessao_recebido)}")
        print(f"   • Oferta Existente no Pedido: {pedido.get('oferta_tipo', 'N/A')}")
        print(f"   • Tipo Cartoon: {pedido.get('tipo_cartoon', 'N/A')}")
        
        # 🔥 CORREÇÃO CRÍTICA: VERIFICAR SE NÃO É NONE PRIMEIRO
        # 1. PRIMEIRO: Usar o tipo_sessao recebido (mais confiável)
        if tipo_sessao_recebido is not None:
            print(f"🎯 [PASSO 1] Usando tipo_sessao_recebido: {tipo_sessao_recebido}")
            return tipo_sessao_recebido
        
        # 2. SEGUNDO: Verificar se já existe oferta_tipo no pedido (do finalizar_compra/finalizar_gift)
        oferta_existente = pedido.get("oferta_tipo")
        if oferta_existente:
            print(f"🎯 [PASSO 2] Usando oferta_existente no pedido: {oferta_existente}")
            return oferta_existente
        
        # 3. TERCEIRO: Tentar obter da session do Stripe
        session_id = pedido.get("session_id_original") or pedido.get("session_id_recuperacao")
        if session_id:
            try:
                session = stripe.checkout.Session.retrieve(session_id)
                tipo_stripe = session.metadata.get("tipo_sessao")
                if tipo_stripe:
                    print(f"🎯 [PASSO 3] Usando tipo detectado via Stripe: {tipo_stripe}")
                    return tipo_stripe
            except Exception as e:
                print(f"⚠️ Não foi possível obter session do Stripe: {e}")
        
        # 4. QUARTO: Verificar pelo tipo de produto
        tipo_cartoon = pedido.get("tipo_cartoon", "").lower()
        print(f"🔍 Tipo Cartoon para análise: '{tipo_cartoon}'")
        
        if "porta-chaves" in tipo_cartoon or "portachaves" in tipo_cartoon:
            print(f"🎯 [PASSO 4] Tipo detectado pelo produto: portachaves")
            return "portachaves"
        
        # 5. QUINTO: Verificar se é oferta surpresa
        if "surpresa" in tipo_cartoon.lower() or "oferta_surpresa" in str(pedido).lower():
            print(f"🎯 [PASSO 5] Tipo detectado: oferta_surpresa")
            return "oferta_surpresa"
        
        # 6. DEFAULT: Pagamento direto
        print(f"🎯 [PASSO 6] Tipo de oferta padrão: pagamento_direto")
        return "pagamento_direto"

    # 🔥 USAR A FUNÇÃO CORRIGIDA
    tipo_oferta = determinar_tipo_oferta(pedido, tipo_sessao)
    
    # VERIFICAR SE JÁ ESTÁ PAGO
    if pedido.get("status") == "pago":
        print(f"❌ PEDIDO JÁ ESTÁ PAGO! Ignorando...")
        return "OK", 200
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: GUARDAR O TIPO DE OFERTA NO PEDIDO
    pedido["oferta_tipo"] = tipo_oferta
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: GUARDAR A TAXA DE IMPOSTO NO PEDIDO
    pais = pedido.get('pais', '').lower()
    if not pais and shipping_details and shipping_details.get('address'):
        # Tentar obter país do shipping details
        pais = shipping_details['address'].get('country', '').lower()
        print(f"🌍 País obtido do shipping: {pais}")
    
    if pais:
        pais_normalizado = normalizar_nome_pais(pais)
        taxas_pais = TAXAS_PAISES.get(pais_normalizado, TAXAS_PAISES["portugal"])
        taxa_imposto = taxas_pais["imposto"]  # Já está em decimal (ex: 0.23)
        pedido["taxa_imposto"] = taxa_imposto
        print(f"💰 Taxa de imposto para {pais}: {taxa_imposto*100}%")
    else:
        print("⚠️ País não encontrado para calcular imposto")
    
    # 🔥 INICIALIZAR VARIÁVEIS DE CONVERSÃO
    amount_eur = amount  # Valor padrão em EUR
    moeda_original = moeda_original or pedido.get('moeda', 'EUR')
    
    print(f"💰 Moeda original: {moeda_original}")
    
    # 🔥 CONVERTER PARA EUR PARA O SISTEMA INTERNO
    if moeda_original and moeda_original.upper() != 'EUR':
        # Usar seu sistema de conversão existente
        TAXAS_CAMBIO = obter_taxas_cambio_em_tempo_real()
        taxa_decimal = TAXAS_CAMBIO.get(moeda_original.lower(), 1.0)
        
        # 🔥 CORREÇÃO: DIVIDIR em vez de multiplicar!
        # A taxa é 1 USD = 1.1575 EUR, então para converter USD para EUR: USD ÷ taxa
        taxa = float(taxa_decimal)
        amount_eur = amount / taxa  # 🔥 MUDANÇA CRÍTICA: / em vez de *
        print(f"💰 CONVERSÃO CORRIGIDA: {amount} {moeda_original} → €{amount_eur:.2f} EUR (taxa: 1 {moeda_original} = {taxa} EUR)")
        
        # 🔥 CORREÇÃO: CONVERTER TAMBÉM OS VALORES DETALHADOS SE EXISTIREM
        if pedido.get('subtotal'):
            subtotal_eur = pedido.get('subtotal', 0) / taxa
            imposto_eur = pedido.get('imposto', 0) / taxa
            frete_eur = pedido.get('frete', 0) / taxa
            
            print(f"💰 CONVERSÃO DETALHADA:")
            print(f"   • Subtotal: {pedido.get('subtotal', 0)} {moeda_original} → €{subtotal_eur:.2f} EUR")
            print(f"   • Imposto: {pedido.get('imposto', 0)} {moeda_original} → €{imposto_eur:.2f} EUR")
            print(f"   • Frete: {pedido.get('frete', 0)} {moeda_original} → €{frete_eur:.2f} EUR")
            
            # 🔥 GUARDAR OS VALORES CONVERTIDOS EM EUR
            pedido["subtotal_eur"] = subtotal_eur
            pedido["imposto_eur"] = imposto_eur
            pedido["frete_eur"] = frete_eur
        
    else:
        amount_eur = amount
        moeda_original = 'EUR'
        print(f"💰 SEM CONVERSÃO: {amount} EUR")
        
        # Para EUR, os valores já estão corretos (se existirem)
        if pedido.get('subtotal'):
            pedido["subtotal_eur"] = pedido.get('subtotal', 0)
            pedido["imposto_eur"] = pedido.get('imposto', 0)
            pedido["frete_eur"] = pedido.get('frete', 0)
    
    # 🔥 CANCELAR TODOS OS TEMPORIZADORES
    await cancelar_temporizadores_pedido(pedido_id)
    
    # ATUALIZAR STATUS COM VALORES EM EUR
    pedido["status"] = "pago"
    pedido["data_pagamento"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    pedido["total_pago_eur"] = amount_eur  # 🔥 GUARDAR EM EUR
    pedido["total_pago_original"] = amount  # 🔥 GUARDAR VALOR ORIGINAL
    pedido["moeda_original"] = moeda_original  # 🔥 GUARDAR MOEDA ORIGINAL
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: SÓ ATUALIZAR ESTATÍSTICAS AQUI (QUANDO PAGO)
    atualizar_estatistica("total_pedidos")
    print(f"📊 ESTATÍSTICAS: Pedido #{pedido_id} AGORA contado (status: pago)")
    
    # 🔥 🔥 🔥 ATUALIZAR ESTATÍSTICAS BASEADO NO TIPO DE OFERTA - CORRIGIDO
    print(f"📊 ATUALIZANDO ESTATÍSTICAS ESPECÍFICAS PARA: {tipo_oferta}")

    if tipo_oferta == "original":
        ESTATISTICAS['ofertas_aceites'] = ESTATISTICAS.get('ofertas_aceites', 0) + 1
        print(f"📈 Estatística atualizada: Oferta original aceite")
        
    elif tipo_oferta == "tamanho_4.5" or tipo_oferta == "oferta_tamanho_45":
        ESTATISTICAS['ofertas_aceites'] = ESTATISTICAS.get('ofertas_aceites', 0) + 1
        print(f"📈 Estatística atualizada: Oferta tamanho 4.5 aceite")
        
    elif tipo_oferta == "portachaves" or tipo_oferta == "oferta_portachaves":
        ESTATISTICAS['ofertas_aceites'] = ESTATISTICAS.get('ofertas_aceites', 0) + 1
        print(f"📈 Estatística atualizada: Oferta portachaves aceite")
        
    elif tipo_oferta == "oferta_surpresa":  # 🔥 NOVA OFERTA SURPRESA
        ESTATISTICAS['ofertas_aceites'] = ESTATISTICAS.get('ofertas_aceites', 0) + 1
        print(f"📈 Estatística atualizada: Oferta surpresa aceite")
        
    elif tipo_oferta == "pagamento_direto":
        # 🔥 NÃO incrementar 'ofertas_aceites' para pagamentos diretos
        print(f"📈 Estatística atualizada: Pagamento direto (não conta como oferta aceite)")
        
    # 🔥 ATUALIZAR CONTADOR ESPECÍFICO PARA O ADMIN_COMMAND
    if 'contadores_ofertas' not in ESTATISTICAS:
        ESTATISTICAS['contadores_ofertas'] = {
            'original': 0,
            'tamanho_4.5': 0,
            'portachaves': 0,
            'oferta_surpresa': 0,  # 🔥 ADICIONADO OFERTA SURPRESA
            'pagamento_direto': 0
        }

    # Incrementar o contador específico
    if tipo_oferta in ESTATISTICAS['contadores_ofertas']:
        ESTATISTICAS['contadores_ofertas'][tipo_oferta] += 1
        print(f"📊 Contador específico atualizado: {tipo_oferta} = {ESTATISTICAS['contadores_ofertas'][tipo_oferta]}")
    
    # 🔥 GUARDAR DADOS DE SHIPPING NO PEDIDO
    if shipping_details:
        pedido["shipping_details"] = shipping_details
        print("✅ Dados de shipping guardados no pedido")
    
    print(f"✅ Pedido #{pedido_id} marcado como PAGO - Valor: {amount} {moeda_original} (€{amount_eur:.2f})")
    print(f"📊 Tipo de oferta registrado: {tipo_oferta}")
    
    # 🔥 ENVIAR PARA O CANAL DE REQUESTS (COM SHIPPING)
    print("🔄 ===== CHAMANDO enviar_pedido_para_canal_requests =====")
    try:
        await enviar_pedido_para_canal_requests(pedido)
        print("✅ ===== enviar_pedido_para_canal_requests CONCLUÍDA =====")
    except Exception as e:
        print(f"❌ ===== ERRO em enviar_pedido_para_canal_requests: {e} =====")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: ENVIAR MENSAGEM DE AGRADECIMENTO AO CLIENTE
    print("🔄 ===== CHAMANDO enviar_mensagem_agradecimento =====")
    try:
        await enviar_mensagem_agradecimento(chat_id, pedido, amount)
        print("✅ ===== enviar_mensagem_agradecimento CONCLUÍDA =====")
    except Exception as e:
        print(f"❌ ===== ERRO em enviar_mensagem_agradecimento: {e} =====")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
    
    # 🔥 NOTIFICAÇÃO PARA O TELEMÓVEL/ADMIN
    print("🔄 ===== CHAMANDO enviar_notificacoes_pagamento =====")
    try:
        await enviar_notificacoes_pagamento(chat_id, pedido, amount)
        print("✅ ===== enviar_notificacoes_pagamento CONCLUÍDA =====")
    except Exception as e:
        print(f"❌ ===== ERRO em enviar_notificacoes_pagamento: {e} =====")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
    
    print(f"🎉 ===== Processamento COMPLETO do pedido #{pedido_id} =====")
    print(f"📊 ===== Estatísticas atualizadas para {tipo_oferta} =====")
    
    return "OK", 200





async def marcar_como_feito(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Apenas troca o botão 'Marcar como Feito' para 'Concluído'"""
    query = update.callback_query
    await query.answer("✅ Pedido marcado como concluído!")
    
    # 🔥 EXTRAIR PEDIDO_ID DO CALLBACK_DATA
    pedido_id = query.data.replace("done_", "")
    
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.answer("❌ Pedido não encontrado", show_alert=True)
        return
    
    try:
        # 🔥 ATUALIZAR STATUS PARA CONCLUÍDO
        pedido['status'] = 'CONCLUÍDO ✅'
        pedido['data_conclusao'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        
        # 🔥 MANTER TODOS OS BOTÕES ORIGINAIS, APENAS TROCAR "MARCAR COMO FEITO" PARA "CONCLUÍDO"
        keyboard = [
            [
                InlineKeyboardButton("📞 Contactar Cliente", 
                                   callback_data=f"contact_{pedido['id']}"),
                
                InlineKeyboardButton("✅ Concluído", 
                                   callback_data=f"done_{pedido['id']}")                   
            ],
            [
                InlineKeyboardButton("📊 Exportar CSV", 
                                   callback_data=f"export_ccsv_{pedido['id']}"),
                InlineKeyboardButton("📃 Exportar TXT", 
                                   callback_data=f"export_txt_{pedido['id']}")                 
                
            ],
            [
                
                InlineKeyboardButton("📄 Exportar PDF", 
                                   callback_data=f"export_pdf_photo_{pedido['id']}"),

                InlineKeyboardButton("📝 Exportar Word", 
                                   callback_data=f"export_word_{pedido['id']}")                   
            ]
        ]
        
        # 🔥 APENAS ATUALIZAR OS BOTÕES, MANTENDO O TEXTO ORIGINAL
        await query.edit_message_reply_markup(
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        
        print(f"✅ Pedido #{pedido_id} marcado como CONCLUÍDO")
        
    except Exception as e:
        print(f"❌ Erro ao marcar como feito: {e}")
        await query.answer("❌ Erro ao atualizar pedido", show_alert=True)







async def exportar_csv_completo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exporta pedido para CSV VERTICAL com nomes curtos e legíveis"""
    query = update.callback_query
    await query.answer("📊 Gerando ficheiro CSV...")
    
    # 🔥 EXTRAIR PEDIDO_ID DO CALLBACK_DATA
    pedido_id = query.data.replace("export_ccsv_", "")
    
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.answer("❌ Pedido não encontrado", show_alert=True)
        return
    
    try:
        # 🔥 DICIONÁRIO COM NOMES CURTOS E LEGÍVEIS
        dados_organizados = {
            # 1. INFORMAÇÕES BÁSICAS
            'ID_Pedido': pedido['id'],
            'Data_Pedido': pedido.get('data', ''),
            'Data_Pagamento': pedido.get('data_pagamento', ''),
            'Status': pedido.get('status', ''),
            'Chat_ID': pedido.get('chat_id', ''),
            
            # 2. DADOS PESSOAIS
            'Cliente': pedido.get('nome', ''),
            'Email': pedido.get('email', ''),
            'País': pedido.get('pais', ''),
            'Contacto': pedido.get('contacto', ''),
            
            # 3. DADOS DE ENVIO
            'Nome_Envio': pedido.get('shipping_details', {}).get('name', ''),
            'Endereco_1': pedido.get('shipping_details', {}).get('address', {}).get('line1', ''),
            'Endereco_2': pedido.get('shipping_details', {}).get('address', {}).get('line2', ''),
            'Cidade': pedido.get('shipping_details', {}).get('address', {}).get('city', ''),
            'Codigo_Postal': pedido.get('shipping_details', {}).get('address', {}).get('postal_code', ''),
            'País_Envio': pedido.get('shipping_details', {}).get('address', {}).get('country', ''),
            
            # 4. PRODUTO
            'Tipo_Cartoon': pedido.get('tipo_cartoon', ''),
            'Estilo': pedido.get('estilo_cartoon', ''),
            'Tamanho': pedido.get('tamanho_cartoon', ''),
            
            # 5. PERSONALIZAÇÕES (APENAS OS PREENCHIDOS)
            'Nome_Familia': pedido.get('nome_family', ''),
            'Frase_Familia': pedido.get('frase_family', ''),
            'Total_Elementos': pedido.get('elementos_family', ''),
            'Adultos': pedido.get('adultos_family', ''),
            'Criancas': pedido.get('criancas_family', ''),
            'Animais': pedido.get('animais_family', ''),
            'Nome_Animal': pedido.get('nome_animal', ''),
            'Tipo_Animal': pedido.get('tipo_animal', ''),
            'Nome_Cartoon': pedido.get('nome_personalizado', pedido.get('nome_cartoon', '')),
            'Frase_Personalizada': pedido.get('frase_personalizado', pedido.get('frase_cartoon', '')),
            'Profissao': pedido.get('profissao', ''),
            'Objetos': pedido.get('objetos_office', ''),
            'Super_Heroi': pedido.get('super_heroi', ''),
            
            # 6. FINANCEIRO
            'Moeda': pedido.get('moeda_original', 'EUR'),
            'Subtotal': f"{get_simbolo_moeda(pedido.get('moeda_original', 'EUR').lower())}{pedido.get('subtotal', 0):.2f}",
            'Imposto': f"{get_simbolo_moeda(pedido.get('moeda_original', 'EUR').lower())}{pedido.get('imposto', 0):.2f}",
            'Frete': f"{get_simbolo_moeda(pedido.get('moeda_original', 'EUR').lower())}{pedido.get('frete', 0):.2f}",
            'Total': f"{get_simbolo_moeda(pedido.get('moeda_original', 'EUR').lower())}{pedido.get('total_pago_original', 0):.2f}",
            'Taxa_Imposto': f"{int(pedido.get('taxa_imposto', 0) * 100)}%",
            
            # 7. CONVERSÕES EUR (se aplicável)
            'Subtotal_EUR': f"€{pedido.get('subtotal_eur', pedido.get('subtotal', 0)):.2f}",
            'Imposto_EUR': f"€{pedido.get('imposto_eur', pedido.get('imposto', 0)):.2f}",
            'Frete_EUR': f"€{pedido.get('frete_eur', pedido.get('frete', 0)):.2f}",
            'Total_EUR': f"€{pedido.get('total_pago_eur', pedido.get('total_pago_original', 0)):.2f}",
            
            # 8. FOTO
            'Tem_Foto': 'SIM' if pedido.get('foto_id') else 'NÃO',
            'Nome_Ficheiro': pedido.get('nome_foto', '')
        }
        
        # 🔥 REMOVER CAMPOS VAZIOS PARA SIMPLIFICAR
        dados_finais = {chave: valor for chave, valor in dados_organizados.items() 
                       if valor not in ['', 'Não adicionou frase', 0, '0.00', '€0.00', '$0.00', '£0.00', 'R$0.00']}
        
        # 🔥 CRIAR CSV VERTICAL (Campo, Valor)
        csv_buffer = io.StringIO()
        
        # Escrever cabeçalho simples
        csv_buffer.write("CAMPO;VALOR\n")
        
        # Escrever dados no formato VERTICAL
        for campo, valor in dados_finais.items():
            # Substituir caracteres problemáticos
            valor_limpo = str(valor).replace(';', ',').replace('\n', ' | ').replace('\r', '')
            csv_buffer.write(f"{campo};{valor_limpo}\n")
        
        # 🔥 PREPARAR FICHEIRO PARA ENVIO
        csv_buffer.seek(0)
        csv_content = csv_buffer.getvalue()
        
        # 🔥 ENVIAR CSV VERTICAL
        await query.message.reply_document(
            document=io.BytesIO(csv_content.encode('utf-8-sig')),  # 🔥 utf-8-sig para Excel
            filename=f"pedido_{pedido_id}.csv",
            caption=f"📊 *CSV Organizado - Pedido #{pedido_id}*\n\n"
                   f"• Formato VERTICAL (Campo | Valor)\n"
                   f"• Nomes curtos e legíveis\n" 
                   f"• Ideal para visualização rápida\n"
                   f"• {len(dados_finais)} campos preenchidos",
            parse_mode="Markdown"
        )
        
        print(f"✅ CSV VERTICAL exportado para pedido #{pedido_id} - {len(dados_finais)} campos")
        
    except Exception as e:
        print(f"❌ Erro ao exportar CSV: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
        await query.answer("❌ Erro ao gerar ficheiro CSV", show_alert=True)

























async def exportar_word_completo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exporta pedido para documento Word com foto"""
    query = update.callback_query
    await query.answer("📝 Gerando documento Word...")
    
    # 🔥 EXTRAIR PEDIDO_ID DO CALLBACK_DATA
    pedido_id = query.data.replace("export_word_", "")
    
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.answer("❌ Pedido não encontrado", show_alert=True)
        return
    
    try:
        # Criar novo documento Word
        doc = Document()
        
        # 🔥 CONFIGURAÇÃO DA PÁGINA
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
        # 🔥 CABEÇALHO
        title = doc.add_heading('GODSPLAN - DETALHES DO PEDIDO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Cartoons Personalizados')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        
        doc.add_paragraph()  # Espaço
        
        # 🔥 INFORMAÇÕES BÁSICAS
        doc.add_heading('INFORMAÇÕES DO PEDIDO', level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        # Preencher tabela de informações
        info_cells = info_table.rows[0].cells
        info_cells[0].text = "Nº do Pedido:"
        info_cells[1].text = pedido['id']
        
        info_cells = info_table.rows[1].cells
        info_cells[0].text = "Data do Pedido:"
        info_cells[1].text = pedido.get('data', 'N/A')
        
        info_cells = info_table.rows[2].cells
        info_cells[0].text = "Data do Pagamento:"
        info_cells[1].text = pedido.get('data_pagamento', 'N/A')
        
        info_cells = info_table.rows[3].cells
        info_cells[0].text = "Status:"
        info_cells[1].text = pedido.get('status', 'N/A')
        
        doc.add_paragraph()  # Espaço
        
        # 🔥 DADOS PESSOAIS
        doc.add_heading('DADOS PESSOAIS', level=1)
        
        pessoal_table = doc.add_table(rows=5, cols=2)
        pessoal_table.style = 'Light Grid Accent 1'
        
        pessoal_cells = pessoal_table.rows[0].cells
        pessoal_cells[0].text = "Nome:"
        pessoal_cells[1].text = pedido.get('nome', 'N/A')
        
        pessoal_cells = pessoal_table.rows[1].cells
        pessoal_cells[0].text = "Email:"
        pessoal_cells[1].text = pedido.get('email', 'N/A')
        
        pessoal_cells = pessoal_table.rows[2].cells
        pessoal_cells[0].text = "País:"
        pessoal_cells[1].text = pedido.get('pais', 'N/A')
        
        pessoal_cells = pessoal_table.rows[3].cells
        pessoal_cells[0].text = "Contacto:"
        pessoal_cells[1].text = pedido.get('contacto', 'N/A')
        
        pessoal_cells = pessoal_table.rows[4].cells
        pessoal_cells[0].text = "Chat ID:"
        pessoal_cells[1].text = str(pedido.get('chat_id', 'N/A'))
        
        doc.add_paragraph()  # Espaço
        
        # 🔥 DADOS DE ENVIO
        if pedido.get('shipping_details'):
            doc.add_heading('DADOS DE ENVIO', level=1)
            
            shipping = pedido['shipping_details']
            address = shipping.get('address', {})
            
            envio_table = doc.add_table(rows=6, cols=2)
            envio_table.style = 'Light Grid Accent 1'
            
            envio_cells = envio_table.rows[0].cells
            envio_cells[0].text = "Nome:"
            envio_cells[1].text = shipping.get('name', 'N/A')
            
            envio_cells = envio_table.rows[1].cells
            envio_cells[0].text = "Linha 1:"
            envio_cells[1].text = address.get('line1', 'N/A')
            
            envio_cells = envio_table.rows[2].cells
            envio_cells[0].text = "Linha 2:"
            envio_cells[1].text = address.get('line2', 'N/A') if address.get('line2') else 'N/A'
            
            envio_cells = envio_table.rows[3].cells
            envio_cells[0].text = "Cidade:"
            envio_cells[1].text = address.get('city', 'N/A')
            
            envio_cells = envio_table.rows[4].cells
            envio_cells[0].text = "Código Postal:"
            envio_cells[1].text = address.get('postal_code', 'N/A')
            
            envio_cells = envio_table.rows[5].cells
            envio_cells[0].text = "País:"
            envio_cells[1].text = address.get('country', 'N/A')
            
            doc.add_paragraph()  # Espaço
        
        # 🔥 DETALHES DO CARTOON
        doc.add_heading('DETALHES DO CARTOON', level=1)
        
        cartoon_table = doc.add_table(rows=3, cols=2)
        cartoon_table.style = 'Light Grid Accent 1'
        
        cartoon_cells = cartoon_table.rows[0].cells
        cartoon_cells[0].text = "Tipo:"
        cartoon_cells[1].text = pedido.get('tipo_cartoon', 'N/A')
        
        cartoon_cells = cartoon_table.rows[1].cells
        cartoon_cells[0].text = "Estilo:"
        # Estilo só para Individual
        tipo_cartoon = pedido.get('tipo_cartoon', '').lower()
        if 'individual' in tipo_cartoon and pedido.get('estilo_cartoon'):
            cartoon_cells[1].text = pedido['estilo_cartoon']
        else:
            cartoon_cells[1].text = 'N/A'
        
        cartoon_cells = cartoon_table.rows[2].cells
        cartoon_cells[0].text = "Tamanho:"
        cartoon_cells[1].text = pedido.get('tamanho_cartoon', 'N/A')
        
        doc.add_paragraph()  # Espaço
        
        # 🔥 CAMPOS PERSONALIZADOS
        campos_personalizados = [
            ('nome_family', 'Nome da Família'),
            ('frase_family', 'Frase da Família'),
            ('elementos_family', 'Total de Elementos'),
            ('adultos_family', 'Adultos'),
            ('criancas_family', 'Crianças'),
            ('animais_family', 'Animais'),
            ('nome_animal', 'Nome do Animal'),
            ('tipo_animal', 'Tipo de Animal'),
            ('tipo_personalizado', 'Tipo de Peça'),
            ('nome_peca_personalizado', 'Nome da Peça'),
            ('nome_personalizado', 'Nome do Cartoon'),
            ('frase_personalizado', 'Frase do Elemento'),
            ('nome_cartoon', 'Nome no Cartoon'),
            ('frase_cartoon', 'Frase na Box'),
            ('profissao', 'Profissão'),
            ('objetos_office', 'Objetos Personalizados'),
            ('super_heroi', 'Super-Herói')
        ]
        
        campos_preenchidos = False
        for campo, label in campos_personalizados:
            if pedido.get(campo) and pedido[campo] != "Não adicionou frase":
                if not campos_preenchidos:
                    doc.add_heading('CAMPOS PERSONALIZADOS', level=1)
                    campos_preenchidos = True
                
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(pedido[campo]))
        
        if campos_preenchidos:
            doc.add_paragraph()  # Espaço
        
        # 🔥 FOTO DO CLIENTE
        doc.add_heading('FOTO DO CLIENTE', level=1)
        
        if pedido.get('foto_id'):
            try:
                # 🔥 BAIXAR A FOTO DO TELEGRAM
                photo_file = await context.bot.get_file(pedido['foto_id'])
                photo_bytes = await photo_file.download_as_bytearray()
                
                # 🔥 ADICIONAR A FOTO AO DOCUMENTO
                photo_stream = io.BytesIO(photo_bytes)
                
                # Adicionar parágrafo para a foto
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Inserir imagem (largura máxima de 12cm)
                run = p.add_run()
                run.add_picture(photo_stream, width=Cm(12))
                
                # Legenda da foto
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.add_run(f"Foto enviada pelo cliente - {pedido.get('nome_foto', 'foto.jpg')}").italic = True
                
            except Exception as e:
                doc.add_paragraph(f"❌ Erro ao carregar foto: {str(e)}")
        else:
            doc.add_paragraph("❌ Nenhuma foto recebida")
        
        doc.add_paragraph()  # Espaço
        
        # 🔥 DETALHES FINANCEIROS
        doc.add_heading('DETALHES FINANCEIROS', level=1)
        
        financeiro_table = doc.add_table(rows=4, cols=2)
        financeiro_table.style = 'Light Grid Accent 1'
        
        moeda_original = pedido.get('moeda_original', 'EUR')
        simbolo_original = get_simbolo_moeda(moeda_original.lower())
        
        finance_cells = financeiro_table.rows[0].cells
        finance_cells[0].text = "Subtotal:"
        
        finance_cells = financeiro_table.rows[1].cells
        finance_cells[0].text = f"Imposto ({pedido.get('taxa_imposto', 0)*100:.0f}%):"
        
        finance_cells = financeiro_table.rows[2].cells
        finance_cells[0].text = "Frete:"
        
        finance_cells = financeiro_table.rows[3].cells
        finance_cells[0].text = "TOTAL FINAL:"
        finance_cells[1].text = f"{simbolo_original}{pedido.get('total_pago_original', 0):.2f} {moeda_original}"
        
        if moeda_original != 'EUR':
            finance_cells = financeiro_table.rows[0].cells
            finance_cells[1].text = f"{simbolo_original}{pedido.get('subtotal', 0):.2f} (€{pedido.get('subtotal_eur', 0):.2f})"
            
            finance_cells = financeiro_table.rows[1].cells
            finance_cells[1].text = f"{simbolo_original}{pedido.get('imposto', 0):.2f} (€{pedido.get('imposto_eur', 0):.2f})"
            
            finance_cells = financeiro_table.rows[2].cells
            finance_cells[1].text = f"{simbolo_original}{pedido.get('frete', 0):.2f} (€{pedido.get('frete_eur', 0):.2f})"
            
            finance_cells = financeiro_table.rows[3].cells
            finance_cells[1].text = f"{simbolo_original}{pedido.get('total_pago_original', 0):.2f} (€{pedido.get('total_pago_eur', 0):.2f})"
        else:
            finance_cells = financeiro_table.rows[0].cells
            finance_cells[1].text = f"€{pedido.get('subtotal', 0):.2f}"
            
            finance_cells = financeiro_table.rows[1].cells
            finance_cells[1].text = f"€{pedido.get('imposto', 0):.2f}"
            
            finance_cells = financeiro_table.rows[2].cells
            finance_cells[1].text = f"€{pedido.get('frete', 0):.2f}"
            
            finance_cells = financeiro_table.rows[3].cells
            finance_cells[1].text = f"€{pedido.get('total_pago_original', 0):.2f}"
        
        # 🔥 SALVAR O DOCUMENTO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # 🔥 ENVIAR O DOCUMENTO WORD
        await query.message.reply_document(
            document=doc_buffer,
            filename=f"pedido_{pedido_id}_completo.docx",
            caption=f"📄 *Documento Word - Pedido #{pedido_id}*\n\nDocumento editável com todos os detalhes e foto incluída.",
            parse_mode="Markdown"
        )
        
        print(f"✅ Documento Word exportado para pedido #{pedido_id}")
        
    except Exception as e:
        print(f"❌ Erro ao exportar Word: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
        await query.answer("❌ Erro ao gerar documento Word", show_alert=True)






















async def exportar_pdf_com_foto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exporta pedido para PDF com estrutura COMPLETA de talão para papel 80mm"""
    query = update.callback_query
    await query.answer("🧾 Gerando talão completo para impressora...")
    
    # 🔥 EXTRAIR PEDIDO_ID DO CALLBACK_DATA
    pedido_id = query.data.replace("export_pdf_photo_", "")
    
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.answer("❌ Pedido não encontrado", show_alert=True)
        return
    
    try:
        # 🔥 TAMANHO PARA PAPEL TILL ROLL 80mm
        width = 80 * mm  # 80mm de largura
        height = 400 * mm  # Altura maior para conteúdo completo
        
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=(width, height))
        
        # 🔥 MARGENS PARA TILL ROLL
        left_margin = 5 * mm
        right_margin = width - 5 * mm
        y_position = height - 10 * mm  # Começar no topo
        
        # 🔥 CABEÇALHO DO TALÃO
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(width/2, y_position, "GODSPLAN")
        y_position -= 6 * mm
        
        c.setFont("Helvetica", 8)
        c.drawCentredString(width/2, y_position, "Cartoons Personalizados")
        y_position -= 8 * mm
        
        # Linha separadora
        c.line(left_margin, y_position, right_margin, y_position)
        y_position -= 6 * mm
        
        # 🔥 1. INFORMAÇÕES BÁSICAS DO PEDIDO
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left_margin, y_position, "INFORMAÇÕES DO PEDIDO")
        y_position -= 5 * mm
        
        c.setFont("Helvetica", 8)
        info_lines = [
            f"Pedido: {pedido['id']}",
            f"Data: {pedido.get('data_pagamento', pedido.get('data', 'N/A'))}",
            f"Status: {pedido.get('status', 'N/A')}",
            f"Chat ID: {pedido.get('chat_id', 'N/A')}"
        ]
        
        for line in info_lines:
            c.drawString(left_margin, y_position, line)
            y_position -= 4 * mm
        
        y_position -= 3 * mm
        c.line(left_margin, y_position, right_margin, y_position)
        y_position -= 4 * mm
        
        # 🔥 2. DADOS PESSOAIS COMPLETOS
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left_margin, y_position, "DADOS PESSOAIS")
        y_position -= 5 * mm
        
        c.setFont("Helvetica", 8)
        pessoal_lines = [
            f"Cliente: {pedido.get('nome', 'N/A')}",
            f"Email: {pedido.get('email', 'N/A')}",
            f"País: {pedido.get('pais', 'N/A')}",
            f"Contacto: {pedido.get('contacto', 'N/A')}"
        ]
        
        for line in pessoal_lines:
            c.drawString(left_margin, y_position, line)
            y_position -= 4 * mm
        
        y_position -= 3 * mm
        c.line(left_margin, y_position, right_margin, y_position)
        y_position -= 4 * mm
        
        # 🔥 3. DADOS DE ENVIO COMPLETOS
        if pedido.get('shipping_details'):
            c.setFont("Helvetica-Bold", 10)
            c.drawString(left_margin, y_position, "DADOS DE ENVIO")
            y_position -= 5 * mm
            
            c.setFont("Helvetica", 8)
            shipping = pedido['shipping_details']
            address = shipping.get('address', {})
            
            envio_lines = [f"Nome: {shipping.get('name', 'N/A')}"]
            
            # Endereço linha por linha
            linha1 = address.get('line1', 'N/A')
            if len(linha1) > 35:
                partes = [linha1[i:i+35] for i in range(0, len(linha1), 35)]
                for parte in partes:
                    envio_lines.append(parte)
            else:
                envio_lines.append(linha1)
            
            if address.get('line2'):
                linha2 = address.get('line2')
                if len(linha2) > 35:
                    partes = [linha2[i:i+35] for i in range(0, len(linha2), 35)]
                    for parte in partes:
                        envio_lines.append(parte)
                else:
                    envio_lines.append(linha2)
            
            cidade = f"{address.get('city', '')} {address.get('postal_code', '')}"
            if cidade.strip():
                envio_lines.append(cidade.strip())
            
            if address.get('country'):
                envio_lines.append(address['country'])
            
            for line in envio_lines:
                c.drawString(left_margin, y_position, line)
                y_position -= 4 * mm
            
            y_position -= 3 * mm
            c.line(left_margin, y_position, right_margin, y_position)
            y_position -= 4 * mm
        
        # 🔥 4. DETALHES DO CARTOON COMPLETOS
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left_margin, y_position, "DETALHES DO CARTOON")
        y_position -= 5 * mm
        
        c.setFont("Helvetica", 8)
        cartoon_lines = [f"Tipo: {pedido.get('tipo_cartoon', 'N/A')}"]
        
        # Estilo só para Individual
        tipo_cartoon = pedido.get('tipo_cartoon', '').lower()
        if 'individual' in tipo_cartoon and pedido.get('estilo_cartoon'):
            cartoon_lines.append(f"Estilo: {pedido['estilo_cartoon']}")
        
        cartoon_lines.append(f"Tamanho: {pedido.get('tamanho_cartoon', 'N/A')}")
        
        for line in cartoon_lines:
            c.drawString(left_margin, y_position, line)
            y_position -= 4 * mm
        
        y_position -= 3 * mm
        c.line(left_margin, y_position, right_margin, y_position)
        y_position -= 4 * mm
        
        # 🔥 5. CAMPOS PERSONALIZADOS COMPLETOS
        campos_personalizados = [
            ('nome_family', 'Família'),
            ('frase_family', 'Frase Família'),
            ('elementos_family', 'Elementos'),
            ('adultos_family', 'Adultos'),
            ('criancas_family', 'Crianças'),
            ('animais_family', 'Animais'),
            ('nome_animal', 'Nome Animal'),
            ('tipo_animal', 'Tipo Animal'),
            ('nome_personalizado', 'Nome Cartoon'),
            ('frase_personalizado', 'Frase Elemento'),
            ('nome_cartoon', 'Nome no Cartoon'),
            ('frase_cartoon', 'Frase na Box'),
            ('profissao', 'Profissão'),
            ('objetos_office', 'Objetos'),
            ('super_heroi', 'Super-Herói')
        ]
        
        campos_preenchidos = False
        for campo, label in campos_personalizados:
            if pedido.get(campo) and pedido[campo] != "Não adicionou frase":
                if not campos_preenchidos:
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(left_margin, y_position, "PERSONALIZAÇÕES")
                    y_position -= 5 * mm
                    campos_preenchidos = True
                
                texto = f"{label}: {pedido[campo]}"
                # Quebrar texto se for muito longo
                if len(texto) > 35:
                    partes = [texto[i:i+35] for i in range(0, len(texto), 35)]
                    for parte in partes:
                        c.setFont("Helvetica", 8)
                        c.drawString(left_margin, y_position, parte)
                        y_position -= 4 * mm
                else:
                    c.setFont("Helvetica", 8)
                    c.drawString(left_margin, y_position, texto)
                    y_position -= 4 * mm
        
        if campos_preenchidos:
            y_position -= 3 * mm
            c.line(left_margin, y_position, right_margin, y_position)
            y_position -= 4 * mm
        
        # 🔥 6. DETALHES FINANCEIROS COMPLETOS (COM PERCENTAGEM DE IMPOSTO)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left_margin, y_position, "DETALHES FINANCEIROS")
        y_position -= 5 * mm
        
        c.setFont("Helvetica", 8)
        moeda_original = pedido.get('moeda_original', 'EUR')
        simbolo_original = get_simbolo_moeda(moeda_original.lower())
        
        # 🔥 OBTER PERCENTAGEM DE IMPOSTO
        taxa_imposto = pedido.get('taxa_imposto', 0)
        percentual_imposto = int(taxa_imposto * 100)  # Converter para inteiro
        
        if moeda_original != 'EUR':
            # 🔥 FORMATO: €106.20 ($115.75) COM PERCENTAGEM DE IMPOSTO
            finance_lines = [
                f"Subtotal: €{pedido.get('subtotal_eur', 0):.2f} ({simbolo_original}{pedido.get('subtotal', 0):.2f})",
                f"Imposto ({percentual_imposto}%): €{pedido.get('imposto_eur', 0):.2f} ({simbolo_original}{pedido.get('imposto', 0):.2f})",  # 🔥 ADICIONADO PERCENTAGEM
                f"Frete: €{pedido.get('frete_eur', 0):.2f} ({simbolo_original}{pedido.get('frete', 0):.2f})",
                f"TOTAL: €{pedido.get('total_pago_eur', 0):.2f} ({simbolo_original}{pedido.get('total_pago_original', 0):.2f})"
            ]
        else:
            # 🔥 SE JÁ FOR EUR, MOSTRAR APENAS EUR COM PERCENTAGEM DE IMPOSTO
            finance_lines = [
                f"Subtotal: €{pedido.get('subtotal', 0):.2f}",
                f"Imposto ({percentual_imposto}%): €{pedido.get('imposto', 0):.2f}",  # 🔥 ADICIONADO PERCENTAGEM
                f"Frete: €{pedido.get('frete', 0):.2f}",
                f"TOTAL: €{pedido.get('total_pago_original', 0):.2f}"
            ]

        for line in finance_lines:
            if "TOTAL:" in line:
                c.setFont("Helvetica-Bold", 9)
            else:
                c.setFont("Helvetica", 8)
            c.drawString(left_margin, y_position, line)
            y_position -= 4 * mm

        # 🔥 ADICIONAR INFORMAÇÃO DE OFERTA SE APLICÁVEL
        if pedido.get('tipo_oferta') == 'portachaves':
            y_position -= 3 * mm
            c.setFont("Helvetica-Bold", 9)
            c.drawString(left_margin, y_position, "OFERTA ESPECIAL")
            y_position -= 4 * mm
            
            c.setFont("Helvetica", 8)
            oferta_lines = [
                f"Porta-chaves - 70% OFF",
                f"Economia: {simbolo_original}{pedido.get('economia', 0):.2f}"
            ]
            
            for line in oferta_lines:
                c.drawString(left_margin, y_position, line)
                y_position -= 4 * mm

        y_position -= 6 * mm
        c.line(left_margin, y_position, right_margin, y_position)
        y_position -= 4 * mm
        
        # 🔥 7. FOTO DO CLIENTE (NO FINAL - COMO PEDIDO)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left_margin, y_position, "FOTO DO CLIENTE")
        y_position -= 5 * mm
        
        if pedido.get('foto_id'):
            try:
                # 🔥 BAIXAR A FOTO DO TELEGRAM
                photo_file = await context.bot.get_file(pedido['foto_id'])
                photo_bytes = await photo_file.download_as_bytearray()
                
                # 🔥 CALCULAR TAMANHO DA FOTO PARA 80mm
                max_width = 70 * mm  # Largura máxima com margens
                max_height = 80 * mm  # Altura máxima
                
                # Criar ImageReader
                image_reader = ImageReader(io.BytesIO(photo_bytes))
                
                # Obter dimensões da imagem
                img_width, img_height = image_reader.getSize()
                
                # Calcular proporção para caber na largura
                ratio = min(max_width / img_width, max_height / img_height, 1.0)
                new_width = img_width * ratio
                new_height = img_height * ratio
                
                # Centralizar a imagem
                x_center = (width - new_width) / 2
                
                # 🔥 INSERIR A FOTO
                c.drawImage(image_reader, x_center, y_position - new_height, 
                           width=new_width, height=new_height, 
                           preserveAspectRatio=True, mask='auto')
                
                # Atualizar posição Y após a foto
                y_position -= new_height + 4 * mm
                
                # Nome do ficheiro da foto
                c.setFont("Helvetica", 7)
                c.drawCentredString(width/2, y_position, f"Ficheiro: {pedido.get('nome_foto', 'foto.jpg')}")
                y_position -= 4 * mm
                
            except Exception as e:
                c.setFont("Helvetica", 7)
                c.drawString(left_margin, y_position, f"❌ Erro ao carregar foto")
                y_position -= 4 * mm
                c.drawString(left_margin, y_position, f"Detalhe: {str(e)[:25]}...")
                y_position -= 4 * mm
        else:
            c.setFont("Helvetica", 8)
            c.drawString(left_margin, y_position, "❌ Nenhuma foto recebida")
            y_position -= 4 * mm
        
        # 🔥 RODAPÉ DO TALÃO
        y_position -= 6 * mm
        c.line(left_margin, y_position, right_margin, y_position)
        y_position -= 4 * mm
        
        c.setFont("Helvetica", 7)
        c.drawCentredString(width/2, y_position, "Obrigado pela sua encomenda! • www.godsplan.com")
        y_position -= 3 * mm
        c.drawCentredString(width/2, y_position, f"Exportado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Finalizar o PDF
        c.save()
        
        # 🔥 ENVIAR O PDF
        pdf_buffer.seek(0)
        await query.message.reply_document(
            document=pdf_buffer,
            filename=f"talao_completo_{pedido_id}.pdf",
            caption=f"🧾 *Talão COMPLETO - Pedido #{pedido_id}*\n\n📋 Inclui TODOS os dados + foto\n📏 Formato para impressora 80mm",
            parse_mode="Markdown"
        )
        
        print(f"✅ Talão COMPLETO PDF gerado para pedido #{pedido_id}")
        
    except Exception as e:
        print(f"❌ Erro ao gerar talão PDF: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
        await query.answer("❌ Erro ao gerar talão", show_alert=True)










async def exportar_txt_com_foto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Exporta pedido para TXT formatado verticalmente com informações da foto"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 EXTRAIR PEDIDO_ID DO CALLBACK_DATA
    pedido_id = query.data.replace("export_txt_photo_", "")
    
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.answer("❌ Pedido não encontrado", show_alert=True)
        return
    
    try:
        # Criar TXT formatado verticalmente
        output = io.StringIO()
        
        # 🔥 CABEÇALHO
        output.write("=" * 60 + "\n")
        output.write("PEDIDO GODSPLAN - EXPORTAÇÃO COMPLETA\n")
        output.write("=" * 60 + "\n")
        output.write(f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n\n")
        
        # 🔥 INFORMAÇÕES BÁSICAS
        output.write("INFORMAÇÕES DO PEDIDO:\n")
        output.write("-" * 30 + "\n")
        output.write(f"ID do Pedido: {pedido['id']}\n")
        output.write(f"Data do Pedido: {pedido.get('data', 'N/A')}\n")
        output.write(f"Data do Pagamento: {pedido.get('data_pagamento', 'N/A')}\n")
        output.write(f"Status: {pedido.get('status', 'N/A')}\n\n")
        
        # 🔥 DADOS PESSOAIS
        output.write("DADOS PESSOAIS:\n")
        output.write("-" * 30 + "\n")
        output.write(f"Nome: {pedido.get('nome', 'N/A')}\n")
        output.write(f"Email: {pedido.get('email', 'N/A')}\n")
        output.write(f"País: {pedido.get('pais', 'N/A')}\n")
        output.write(f"Contacto: {pedido.get('contacto', 'N/A')}\n")
        output.write(f"Chat ID: {pedido.get('chat_id', 'N/A')}\n\n")
        
        # 🔥 DADOS DE ENVIO
        if pedido.get('shipping_details'):
            shipping = pedido['shipping_details']
            address = shipping.get('address', {})
            
            output.write("DADOS DE ENVIO:\n")
            output.write("-" * 30 + "\n")
            output.write(f"Nome: {shipping.get('name', 'N/A')}\n")
            output.write(f"Linha 1: {address.get('line1', 'N/A')}\n")
            if address.get('line2'):
                output.write(f"Linha 2: {address.get('line2')}\n")
            output.write(f"Cidade: {address.get('city', 'N/A')}\n")
            output.write(f"Código Postal: {address.get('postal_code', 'N/A')}\n")
            output.write(f"País: {address.get('country', 'N/A')}\n\n")
        
        # 🔥 DETALHES DO CARTOON
        output.write("DETALHES DO CARTOON:\n")
        output.write("-" * 30 + "\n")
        output.write(f"Tipo: {pedido.get('tipo_cartoon', 'N/A')}\n")
        
        # Estilo só para Individual
        tipo_cartoon = pedido.get('tipo_cartoon', '').lower()
        if 'individual' in tipo_cartoon and pedido.get('estilo_cartoon'):
            output.write(f"Estilo: {pedido['estilo_cartoon']}\n")
        
        output.write(f"Tamanho: {pedido.get('tamanho_cartoon', 'N/A')}\n\n")
        
        # 🔥 CAMPOS PERSONALIZADOS
        campos_personalizados = [
            ('nome_family', 'Nome da Família'),
            ('frase_family', 'Frase da Família'),
            ('elementos_family', 'Total de Elementos'),
            ('adultos_family', 'Adultos'),
            ('criancas_family', 'Crianças'),
            ('animais_family', 'Animais'),
            ('nome_animal', 'Nome do Animal'),
            ('tipo_animal', 'Tipo de Animal'),
            ('tipo_personalizado', 'Tipo de Peça'),
            ('nome_peca_personalizado', 'Nome da Peça'),
            ('nome_personalizado', 'Nome do Cartoon'),
            ('frase_personalizado', 'Frase do Elemento'),
            ('nome_cartoon', 'Nome no Cartoon'),
            ('frase_cartoon', 'Frase na Box'),
            ('profissao', 'Profissão'),
            ('objetos_office', 'Objetos Personalizados'),
            ('super_heroi', 'Super-Herói')
        ]
        
        campos_preenchidos = False
        for campo, label in campos_personalizados:
            if pedido.get(campo) and pedido[campo] != "Não adicionou frase":
                if not campos_preenchidos:
                    output.write("CAMPOS PERSONALIZADOS:\n")
                    output.write("-" * 30 + "\n")
                    campos_preenchidos = True
                output.write(f"{label}: {pedido[campo]}\n")
        
        if campos_preenchidos:
            output.write("\n")
        
        # 🔥 INFORMAÇÕES DA FOTO
        output.write("INFORMAÇÕES DA FOTO:\n")
        output.write("-" * 30 + "\n")
        if pedido.get('foto_id'):
            output.write("✅ Foto recebida: SIM\n")
            output.write(f"Nome do ficheiro: {pedido.get('nome_foto', 'N/A')}\n")
            output.write("📸 A foto está disponível no sistema\n")
        else:
            output.write("❌ Foto recebida: NÃO\n")
        output.write("\n")
        
        # 🔥 DETALHES FINANCEIROS
        output.write("DETALHES FINANCEIROS:\n")
        output.write("-" * 30 + "\n")
        
        moeda_original = pedido.get('moeda_original', 'EUR')
        simbolo_original = get_simbolo_moeda(moeda_original.lower())
        
        if moeda_original != 'EUR':
            output.write(f"Subtotal: {simbolo_original}{pedido.get('subtotal', 0):.2f} (€{pedido.get('subtotal_eur', 0):.2f})\n")
            output.write(f"Imposto ({pedido.get('taxa_imposto', 0)*100:.0f}%): {simbolo_original}{pedido.get('imposto', 0):.2f} (€{pedido.get('imposto_eur', 0):.2f})\n")
            output.write(f"Frete: {simbolo_original}{pedido.get('frete', 0):.2f} (€{pedido.get('frete_eur', 0):.2f})\n")
            output.write(f"Total Final: {simbolo_original}{pedido.get('total_pago_original', 0):.2f} (€{pedido.get('total_pago_eur', 0):.2f})\n")
        else:
            output.write(f"Subtotal: €{pedido.get('subtotal', 0):.2f}\n")
            output.write(f"Imposto ({pedido.get('taxa_imposto', 0)*100:.0f}%): €{pedido.get('imposto', 0):.2f}\n")
            output.write(f"Frete: €{pedido.get('frete', 0):.2f}\n")
            output.write(f"Total Final: €{pedido.get('total_pago_original', 0):.2f}\n")
        
        output.write("\n" + "=" * 60 + "\n")
        output.write("EXPORTAÇÃO CONCLUÍDA\n")
        output.write("=" * 60 + "\n")
        
        txt_data = output.getvalue()
        output.close()
        
        # Enviar como arquivo
        await query.message.reply_document(
            document=io.BytesIO(txt_data.encode()),
            filename=f"pedido_{pedido_id}_completo.txt",
            caption=f"📝 *Exportação TXT Completa - Pedido #{pedido_id}*\n\nTodos os detalhes organizados verticalmente com informações da foto.",
            parse_mode="Markdown"
        )
        
        print(f"✅ TXT com foto exportado para pedido #{pedido_id}")
        
    except Exception as e:
        print(f"❌ Erro ao exportar TXT com foto: {e}")
        await query.answer("❌ Erro ao gerar arquivo", show_alert=True)



 # 🔥 ACABA ADMIN








async def safe_delete_message(message_or_query):
    """Deleta mensagens de forma segura, prevenindo erros"""
    try:
        if hasattr(message_or_query, 'delete_message'):
            # É um CallbackQuery
            await message_or_query.delete_message()
        elif hasattr(message_or_query, 'delete'):
            # É uma Message
            await message_or_query.delete()
        else:
            print("❌ Tipo não suportado para delete")
        return True
    except telegram.error.BadRequest as e:
        if "Message to delete not found" in str(e):
            print("⚠️ Mensagem já deletada - ignorando")
            return True  # Consideramos sucesso pois a mensagem já não existe
        else:
            print(f"❌ Erro ao deletar mensagem: {e}")
            return False
    except Exception as e:
        print(f"❌ Erro inesperado ao deletar: {e}")
        return False



# --- Handler para mensagens ---
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler principal para mensagens de TEXTO - VERSÃO SEGURA"""
    try:
        print(f"🔍 HANDLE_MESSAGE (TEXTO) CHAMADO")
        
        # 🔥 VERIFICAÇÃO 1: update existe
        if not update:
            print("⚠️ Update é None - ignorando")
            return
        
        # 🔥 DETERMINAR TIPO DE MENSAGEM
        is_channel_post = update.channel_post is not None
        message = update.message or update.channel_post
        
        if not message:
            print("⚠️ Nenhuma mensagem encontrada")
            return
        
        print(f"📨 Tipo: {'CHANNEL_POST' if is_channel_post else 'MESSAGE'} | Chat ID: {message.chat.id}")
        
        # 🔥 PRIMEIRO: VERIFICAR SE É UMA EDIÇÃO (POR CHAT_ID)
        # Isso deve ser verificado ANTES de verificar effective_user!
        chat_id = message.chat.id
        
        # 1. Verificar edição de IMPOSTOS
        editing_data = context.bot_data.get(f'editing_tax_{chat_id}')
        if editing_data:
            print(f"🎯 DETETADO MODO EDIÇÃO DE IMPOSTO")
            await processar_edicao_imposto_direto(update, context, editing_data, message)
            return
        
        # 2. Verificar edição de FRETE
        frete_editing_data = context.bot_data.get(f'editing_frete_{chat_id}')
        if frete_editing_data:
            print(f"🎯 DETETADO MODO EDIÇÃO DE FRETE")
            await processar_edicao_frete_direto(update, context, frete_editing_data, message)
            return
        
        # 3. Verificar edição do PAINEL ADMIN
        painel_editing_data = context.bot_data.get(f'editing_painel_{chat_id}')
        if painel_editing_data:
            print(f"🎯 DETETADO MODO EDIÇÃO DO PAINEL")
            await processar_edicao_painel_direto(update, context, painel_editing_data, message)
            return
        
        # 🔥 SE NÃO FOR EDIÇÃO: VERIFICAR effective_user (para mensagens normais)
        if not update.effective_user:
            print("⚠️ effective_user é None e não é edição - ignorando")
            return
        
        # ✅ É UMA MENSAGEM NORMAL DE USUÁRIO
        user_id = update.effective_user.id
        user_data = context.user_data if context.user_data is not None else {}
        
        print(f"👤 User: {update.effective_user.first_name} (ID: {user_id}): {message.text[:100]}")
        
        # 🔥 OBTER SESSÃO DO USUÁRIO
        try:
            session = get_user_session(user_id)
            session.last_activity = time.time()
        except:
            print("⚠️ Erro ao obter sessão")
            session = None
        
        # 🔥 PEGAR IDIOMA
        idioma = user_data.get('idioma', 'portugues')
        print(f"🌐 Idioma: {idioma}")
        
        # 🔥 SE FOR FOTO, RETORNAR
        if message.photo:
            print(f"📸 Foto recebida - deixando para gift_foto_handler")
            return
        
        # 🔥 VERIFICAR SE TEM TEXTO
        if not message.text:
            print("⚠️ Mensagem sem texto")
            return
        
        print(f"🔍 Processando TEXTO: '{message.text[:100]}'")
        
        # 🔥 DEPOIS: Processamento normal das mensagens
        if message.text:
            if message.text.startswith('/start'):
                await start(update, context)
                return
            elif message.text.startswith('/help'):
                await help_handler(update, context)
                return
            elif message.text.startswith('/'):
                print(f"🔧 Comando {message.text} será tratado pelo CommandHandler")
                return
        
        # 🔥 USAR DADOS DA SESSÃO
        state = None
        if session:
            state = session.get_state('conversation_state')
        
        if state is None:
            state = user_data.get('conversation_state')
        
        print(f"🔍 Estado atual: {state} | Idioma: {idioma}")

        # 🔥 ESTADOS DE PROBLEMA (TEXTO)
        if state == AGUARDANDO_REPORTE_PROBLEMA:
            if message.text:
                print(f"🎯 Processando como REPORTE DE PROBLEMA (texto)")
                await receber_problema(update, context)
                return
                
        elif state == AGUARDANDO_ID_PEDIDO:
            if message.text:
                print(f"🎯 Processando como ID PEDIDO")
                await receber_id_pedido(update, context)
                return

        elif state == 'problema_outro':
            if message.text:
                print(f"🎯 Processando como PROBLEMA_OUTRO (texto)")
                await problema_outro(update, context)
                return

        elif state == 'todas_recusadas':
            if message.text:
                print(f"🎯 Processando como todas_recusadas (texto)")
                await todas_recusadas(update, context)
                return

        elif state == FOTO_PROBLEMA:
            if message.text:
                print(f"🎯 Processando como DESCRIÇÃO após foto problema")
                await receber_problema(update, context)
                return
        
        # 🔥 ESTADOS NORMAIS DO FLUXO DE CARTOON (TEXTO)
        elif state == AGUARDANDO_SCREENSHOT_CARTOON:
            if message.text:
                print(f"🎯 Processando como TEXTO para screenshot cartoon")
                textos_erro = {
                    'portugues': "📸 *Por favor, envie uma screenshot/foto!*",
                    'ingles': "📸 *Please send a screenshot/photo!*",
                    'espanhol': "📸 *¡Por favor, envía una captura de pantalla/foto!*",
                    'italiano': "📸 *Per favore, invia uno screenshot/foto!*",
                    'alemao': "📸 *Bitte senden Sie einen Screenshot/Foto!*",
                    'frances': "📸 *Veuillez envoyer une capture d'écran/photo !*"
                }
                await message.reply_text(
                    textos_erro.get(idioma, textos_erro['portugues']),
                    parse_mode="Markdown"
                )
                return

        elif state == DESCRICAO:
            if message.text:
                print(f"🎯 Processando como DESCRICAO do cartoon")
                await receber_descricao(update, context)
                return
                
        elif state == CORRECOES:
            if message.text:
                print(f"🎯 Processando como CORRECOES do cartoon")
                await processar_correcoes(update, context)
                return

        elif state == FOTO:
            if message.text:
                print(f"🎯 Processando como TEXTO para foto do cartoon")
                textos_erro = {
                    'portugues': "📸 *Por favor, envie uma foto!*",
                    'ingles': "📸 *Please send a photo!*",
                    'espanhol': "📸 *¡Por favor, envía uma foto!*",
                    'italiano': "📸 *Per favore, invia una foto!*",
                    'alemao': "📸 *Bitte senden Sie ein Foto!*",
                    'frances': "📸 *Veuillez envoyer une photo !*"
                }
                await message.reply_text(
                    textos_erro.get(idioma, textos_erro['portugues']),
                    parse_mode="Markdown"
                )
                return

        # 🔥 ESTADOS DO FLUXO PRINCIPAL
        elif state == NOME:
            await receber_nome(update, context)
        elif state == EMAIL:
            await receber_email(update, context)
        elif state == CONTACTO:
            await receber_contacto(update, context)
        elif state == PROFISSAO:
            await receber_profissao(update, context)
        elif state == OBJETOS:
            await receber_objetos(update, context)
        elif state == SUPER_HEROI:
            await receber_super_heroi(update, context)
        elif state == ELEMENTOS_FAMILY:
            await receber_elementos_family(update, context)
        elif state == ADULTOS_FAMILY:
            await receber_adultos_family(update, context)
        elif state == CRIANCAS_FAMILY:
            await receber_criancas_family(update, context)
        elif state == ANIMAIS_FAMILY:
            await receber_animais_family(update, context)
        elif state == NOME_ANIMAL:
            await receber_nome_animal(update, context)
        elif state == NOME_PECA:
            await receber_nome_peca(update, context)

        elif state == NOME_PERSONALIZADO:
            await receber_nome_personalizado(update, context)
        elif state == FRASE_PERSONALIZADO:
            await receber_frase_personalizado(update, context)

        elif state == NOME_FAMILY:
            await receber_nome_family(update, context)
        elif state == FRASE_FAMILY:
            await receber_frase_family(update, context)

        elif state == NOME_CARTOON:
            await receber_nome_cartoon(update, context)
        elif state == FRASE_CARTOON:
            await receber_frase_cartoon(update, context)

        elif user_data.get('aguardando_pais_manual'):
            await receber_pais_manual(update, context)
        
        # 🔥 ESTADOS DE GIFT (para texto)
        elif state in [GIFT_NOME, GIFT_EMAIL, GIFT_CONTACTO, GIFT_NOME_BOX, GIFT_FRASE_BOX]:
            print(f"🎯 É estado de GIFT (texto) - deixando para gift_text_handler")
            # Deixa para o gift_text_handler específico
            return
        
        else:
            print(f"🔍 Estado desconhecido: {state} | Idioma: {idioma}")
            
            textos_resposta = {
                'portugues': "👋 *Olá!*\n\nUse /start para criar seu cartoon\nUse /help para ajuda",
                'ingles': "👋 *Hello!*\n\nUse /start to create your cartoon\nUse /help for assistance",
                'espanhol': "👋 *¡Hola!*\n\nUsa /start para crear tu cartoon\nUsa /help para ajuda",
                'italiano': "👋 *Ciao!*\n\nUsa /start per creare il tuo cartoon\nUsa /help per assistenza",
                'alemao': "👋 *Hallo!*\n\nVerwenden Sie /start, um Ihren Cartoon zu erstellen\nVerwenden Sie /help für Unterstützung",
                'frances': "👋 *Bonjour !*\n\nUtilisez /start para crear votre dessin animé\nUtilisez /help para obtener de l'aide"
            }
            
            await message.reply_text(
                textos_resposta.get(idioma, textos_resposta['portugues']),
                parse_mode="Markdown"
            )

    except Exception as e:
        print(f"❌ ERRO em handle_message: {e}")
        # NÃO re-lançar o erro - deixa o bot continuar






# --- Handler para voltar ---
async def voltar_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    context.user_data.clear()
    await menu_inicial(update, context)

# ======================= FUNÇÃO FINALIZAR COMPRA =======================
async def finalizar_compra(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    print("🔄 FINALIZAR_COMPRA INICIADO - COM TODOS OS CAMPOS DA FAMILY")
    print(f"🔍 DEBUG - Chat ID: {query.message.chat_id}")
    print(f"🔍 DEBUG - User Data keys: {list(context.user_data.keys())}")
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    print(f"🌐 Idioma detectado: {idioma}")
    
    # Verificar campos específicos
    if "nome_family" in context.user_data:
        print(f"👨‍👩‍👧‍👦 Nome family encontrado: {context.user_data['nome_family']}")
    if "frase_family" in context.user_data:
        print(f"💬 Frase family encontrada: {context.user_data['frase_family']}")
    
    if "pedido_id" in context.user_data:
        old_pedido_id = context.user_data["pedido_id"]
        print(f"🚨 ATENÇÃO: JÁ EXISTE pedido_id NO user_data: {old_pedido_id}")
        
        if old_pedido_id in PEDIDOS_REGISTO:
            status_antigo = PEDIDOS_REGISTO[old_pedido_id]["status"]
            print(f"🔍 Pedido anterior #{old_pedido_id} ainda no registo - Status: {status_antigo}")
            
            if status_antigo == "pendente":
                del PEDIDOS_REGISTO[old_pedido_id]
                print(f"🗑️ Pedido anterior #{old_pedido_id} removido do registo")
        
        del context.user_data["pedido_id"]
        print(f"✅ Pedido_id anterior #{old_pedido_id} removido do user_data")
    
    # Remover a mensagem anterior com botões
    await safe_delete_message(query)
    
    # 🔥 CALCULAR TOTAIS
    pais = context.user_data.get("pais", "portugal")
    print(f"🌍 País selecionado: {pais}")
    totais = calcular_total_por_moeda(context, pais)
    
    # 🔥 DICIONÁRIO PARA CONVERTER PAÍSES PARA INGLÊS
    PAISES_PARA_INGLES = {
        'estados_unidos': 'United States',
        'canada': 'Canada',
        'reino_unido': 'United Kingdom',
        'brasil': 'Brazil',
        'alemanha': 'Germany',
        'paises_baixos': 'Netherlands',
        'holanda': 'Netherlands',
        'franca': 'France',
        'espanha': 'Spain',
        'belgica': 'Belgium',
        'italia': 'Italy',
        'portugal': 'Portugal',
        'irlanda': 'Ireland',
        'luxemburgo': 'Luxembourg'
    }
    
    def converter_pais_para_ingles(pais_key):
        """Converte o nome/callback do país para inglês"""
        if isinstance(pais_key, str):
            # Remove "pais_" se existir
            if pais_key.startswith('pais_'):
                pais_key = pais_key[5:]
            # Remove acentos e converte para minúsculas para comparação
            pais_clean = pais_key.lower()
            # Mapeamento adicional para nomes em português
            mapeamento = {
                'bélgica': 'belgica',
                'bélgica (português)': 'belgica',
                'frança': 'franca',
                'espanha': 'espanha',
                'alemanha': 'alemanha',
                'itália': 'italia',
                'irlanda': 'irlanda',
                'luxemburgo': 'luxemburgo',
                'países baixos': 'paises_baixos',
                'holanda': 'paises_baixos',
                'reino unido': 'reino_unido',
                'estados unidos': 'estados_unidos',
                'eua': 'estados_unidos'
            }
            pais_key = mapeamento.get(pais_clean, pais_key)
        return PAISES_PARA_INGLES.get(pais_key, pais_key.title())
    
    # 🔥 CONVERTER PAÍS PARA INGLÊS
    pais_original = pais
    pais_ingles = converter_pais_para_ingles(pais_original)
    print(f"🌍 País original: {pais_original} -> Inglês: {pais_ingles}")
    
    # 🔥 OBTER TODOS OS DADOS
    nome = context.user_data.get("nome", "")
    email = context.user_data.get("email", "")
    contacto = context.user_data.get("contacto", "")
    tipo = context.user_data.get("tipo_cartoon", "")
    estilo = context.user_data.get("estilo_cartoon", "")
    tamanho = context.user_data.get("tamanho_cartoon", "")
    nome_foto = context.user_data.get("nome_foto", "foto.jpg")

    print(f"🔍 DEBUG FINALIZAR_COMPRA - Tamanho a guardar:")
    print(f"   • tamanho_cartoon: {tamanho}")
    print(f"   • tamanho_original: {tamanho} (mesmo valor)")
    
    # 🔥 VERIFICAR TIPO DE OFERTA (PAGAMENTO DIRETO OU OFERTA ESPECÍFICA)
    oferta_tipo = context.user_data.get("oferta_tipo", "pagamento_direto")
    print(f"🔍 Tipo de oferta detectado: {oferta_tipo}")
    
    # 🔥 CAMPOS PERSONALIZADOS
    profissao = context.user_data.get("profissao", "")
    objetos_office = context.user_data.get("objetos_office", "")
    super_heroi = context.user_data.get("super_heroi", "")
    elementos_family = context.user_data.get("elementos_family", "")
    adultos_family = context.user_data.get("adultos_family", "")
    criancas_family = context.user_data.get("criancas_family", "")
    animais_family = context.user_data.get("animais_family", "")
    nome_animal = context.user_data.get("nome_animal", "")
    tipo_animal = context.user_data.get("tipo_animal", "")
    
    # 🔥 CAMPOS ESPECÍFICOS DO PERSONALIZADO
    tipo_personalizado = context.user_data.get("tipo_personalizado", "")
    nome_peca_personalizado = context.user_data.get("nome_peca_personalizado", "")
    nome_personalizado = context.user_data.get("nome_personalizado", "")
    frase_personalizado = context.user_data.get("frase_personalizado", "")
    
    # 🔥 CAMPOS DE PERSONALIZAÇÃO DA BOX
    nome_cartoon = context.user_data.get("nome_cartoon", "")
    frase_cartoon = context.user_data.get("frase_cartoon", "")
    
    # 🔥 CAMPOS ESPECÍFICOS DA FAMILY
    nome_family = context.user_data.get("nome_family", "")
    frase_family = context.user_data.get("frase_family", "")
    
    foto_recebida = "✅" if "foto_id" in context.user_data else "❌"

    # GERAR ID ÚNICO DO PEDIDO
    pedido_id = str(uuid.uuid4())[:8].upper()
    data_pedido = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    
    if pedido_id in PEDIDOS_REGISTO:
        print(f"🔄 CONFLITO: ID {pedido_id} já existe, gerando novo...")
        pedido_id = str(uuid.uuid4())[:8].upper()
        print(f"🆔 NOVO ID gerado: {pedido_id}")
    
    # 🔥 🔥 🔥 CORREÇÃO: GUARDAR NO PEDIDOS_REGISTO PARA O TEMPORIZADOR FUNCIONAR
    # Mas ainda NÃO contar nas estatísticas - só quando pagar
    
    # Criar o objeto do pedido
    pedido_data = {
        "id": pedido_id,
        "data": data_pedido,
        "nome": nome,
        "email": email,
        "pais": pais_ingles,  # 🔥 GUARDAR EM INGLÊS NO PEDIDO
        "pais_original": pais_original,  # Guardar original também
        "contacto": contacto,
        "tipo_cartoon": tipo,
        "estilo_cartoon": estilo,
        "tamanho_cartoon": tamanho,
        "tamanho_original": tamanho,
        "nome_foto": nome_foto,
        "foto_id": context.user_data.get("foto_id"),
        
        # 🔥 GUARDAR IDIOMA NO PEDIDO
        "idioma": idioma,
        
        # 🔥 🔥 🔥 IDENTIFICAR TIPO DE PAGAMENTO (DIRETO OU OFERTA)
        "oferta_tipo": oferta_tipo,  # "pagamento_direto", "original", "tamanho_4.5", "portachaves"
        
        # 🔥 CAMPOS PERSONALIZADOS
        "profissao": profissao,
        "objetos_office": objetos_office,
        "super_heroi": super_heroi,
        "elementos_family": elementos_family,
        "adultos_family": adultos_family,
        "criancas_family": criancas_family,
        "animais_family": animais_family,
        "nome_animal": nome_animal,
        "tipo_animal": tipo_animal,
        
        # 🔥 CAMPOS ESPECÍFICOS DO PERSONALIZADO
        "tipo_personalizado": tipo_personalizado,
        "nome_peca_personalizado": nome_peca_personalizado,
        "nome_personalizado": nome_personalizado,
        "frase_personalizado": frase_personalizado,
        
        # 🔥 CAMPOS DE PERSONALIZAÇÃO DA BOX
        "nome_cartoon": nome_cartoon,
        "frase_cartoon": frase_cartoon,
        
        # 🔥 CAMPOS ESPECÍFICOS DA FAMILY
        "nome_family": nome_family,
        "frase_family": frase_family,
        
        # 🔥 INFORMAÇÕES DE MOEDA
        "subtotal": totais['subtotal'],
        "imposto": totais['imposto'],
        "frete": totais['frete'],
        "total": totais['total'],
        "valor_original_real": totais['total'],
        "moeda": totais['moeda'],
        "simbolo_moeda": totais['simbolo_moeda'],
        "chat_id": query.message.chat_id,
        "status": "pendente",
        "data_expiracao": datetime.now() + timedelta(minutes=10),
        "tentativas_recuperacao": 0
    }
    
    # 🔥 GUARDAR NO PEDIDOS_REGISTO (PARA TEMPORIZADOR FUNCIONAR)
    PEDIDOS_REGISTO[pedido_id] = pedido_data
    
    # 🔥 GUARDAR TAMBÉM NO USER_DATA (PARA FACILITAR ACESSO)
    context.user_data["pedido_data"] = pedido_data.copy()
    context.user_data["pedido_id"] = pedido_id

    print(f"✅ PEDIDO GUARDADO NO PEDIDOS_REGISTO: #{pedido_id}")
    print(f"📊 TIPO DE PAGAMENTO: {oferta_tipo}")
    print(f"📊 NÃO CONTADO NAS ESTATÍSTICAS (aguardando pagamento)")
    print(f"💰 Moeda do pedido: {totais['moeda']} {totais['simbolo_moeda']}")
    print(f"🌐 Idioma do pedido: {idioma}")

    # 🔥 CÁLCULO DO PREÇO ANTERIOR E KLARNA
    preco_anterior = totais['total'] / 0.70
    desconto = preco_anterior - totais['total']
    percentual_desconto = 30
    valor_klarna = totais['total'] / 3

    # 🔥 TEXTOS POR IDIOMA
    textos_resumo = {
        'portugues': {
            'titulo': "🧾 RESUMO FINAL PARA PAGAMENTO",
            'id_pedido': "🆔 ID do Pedido",
            'data': "📅 Data",
            'pais_envio': "🌍 País de Envio",
            'moeda': "💰 Moeda",
            'tempo_pagar': "⏰ Tempo para pagar",
            'dados_pessoais': "👤 DADOS PESSOAIS",
            'nome': "👤 Nome",
            'email': "📧 Email",
            'pais': "🌍 País",
            'telefone': "📱 Telefone",
            'detalhes_cartoon': "🎨 DETALHES DO CARTOON",
            'tipo': "🎨 Tipo",
            'nome_familia': "👨‍👩‍👧‍👦 Nome da Família",
            'frase_familia': "💬 Frase da Família",
            'tipo_peca': "📦 Tipo de Peça",
            'nome_peca': "📝 Nome da Peça",
            'nome_cartoon': "🎭 Nome do Cartoon",
            'frase_elemento': "💬 Frase do Elemento",
            'nome_no_cartoon': "🎭 Nome no Cartoon",
            'frase_na_box': "💬 Frase na Box",
            'estilo': "🖌 Estilo",
            'profissao': "💼 Profissão",
            'objetos_personalizados': "🎯 Objetos Personalizados",
            'super_heroi': "🦸‍♂️ Super-Herói",
            'total_elementos': "👥 Total de Elementos",
            'adultos': "👨‍👩 Adultos",
            'criancas': "👧🧒 Crianças",
            'animais': "🐱🐶 Animais",
            'nome_animal': "🐾 Nome do Animal",
            'tipo_animal': "🐕 Tipo de Animal",
            'tamanho': "📏 Tamanho",
            'foto': "📸 Foto",
            'valores': "💵 VALORES",
            'preco_anterior': "Preço anterior",
            'desconto': "Desconto",
            'total_pagar': "TOTAL A PAGAR",
            'impostos_frete': "Impostos e Frete Incluídos",
            'klarna_disponivel': "💎 KLARNA DISPONÍVEL",
            'klarna_texto': "Pague em 3x de",
            'klarna_juros': "SEM JUROS",
            'alerta_tempo': "⚠️ Tem 10 minutos para efetuar o pagamento!",
            'guardar_id': "Guarde o ID do pedido para referência futura!",
            'clique_pagar': "Clique abaixo para pagar: 👇",
            'botao_pagar': "💳 Pagar com Cartão",
            'minutos': "minutos"
        },
        'ingles': {
            'titulo': "🧾 FINAL SUMMARY FOR PAYMENT",
            'id_pedido': "🆔 Order ID",
            'data': "📅 Date",
            'pais_envio': "🌍 Shipping Country",
            'moeda': "💰 Currency",
            'tempo_pagar': "⏰ Time to pay",
            'dados_pessoais': "👤 PERSONAL DATA",
            'nome': "👤 Name",
            'email': "📧 Email",
            'pais': "🌍 Country",
            'telefone': "📱 Phone",
            'detalhes_cartoon': "🎨 CARTOON DETAILS",
            'tipo': "🎨 Type",
            'nome_familia': "👨‍👩‍👧‍👦 Family Name",
            'frase_familia': "💬 Family Phrase",
            'tipo_peca': "📦 Piece Type",
            'nome_peca': "📝 Piece Name",
            'nome_cartoon': "🎭 Cartoon Name",
            'frase_elemento': "💬 Element Phrase",
            'nome_no_cartoon': "🎭 Name in Cartoon",
            'frase_na_box': "💬 Box Phrase",
            'estilo': "🖌 Style",
            'profissao': "💼 Profession",
            'objetos_personalizados': "🎯 Custom Objects",
            'super_heroi': "🦸‍♂️ Superhero",
            'total_elementos': "👥 Total Elements",
            'adultos': "👨‍👩 Adults",
            'criancas': "👧🧒 Children",
            'animais': "🐱🐶 Animals",
            'nome_animal': "🐾 Animal Name",
            'tipo_animal': "🐕 Animal Type",
            'tamanho': "📏 Size",
            'foto': "📸 Photo",
            'valores': "💵 VALUES",
            'preco_anterior': "Previous price",
            'desconto': "Discount",
            'total_pagar': "TOTAL TO PAY",
            'impostos_frete': "Taxes and Shipping Included",
            'klarna_disponivel': "💎 KLARNA AVAILABLE",
            'klarna_texto': "Pay in 3 installments of",
            'klarna_juros': "NO INTEREST",
            'alerta_tempo': "⚠️ You have 10 minutes to make the payment!",
            'guardar_id': "Save the order ID for future reference!",
            'clique_pagar': "Click below to pay: 👇",
            'botao_pagar': "💳 Pay with Card",
            'minutos': "minutes"
        },
        'espanhol': {
            'titulo': "🧾 RESUMEN FINAL PARA PAGO",
            'id_pedido': "🆔 ID del Pedido",
            'data': "📅 Fecha",
            'pais_envio': "🌍 País de Envío",
            'moeda': "💰 Moneda",
            'tempo_pagar': "⏰ Tiempo para pagar",
            'dados_pessoais': "👤 DATOS PERSONALES",
            'nome': "👤 Nombre",
            'email': "📧 Email",
            'pais': "🌍 País",
            'telefone': "📱 Teléfono",
            'detalhes_cartoon': "🎨 DETALLES DEL CARTOON",
            'tipo': "🎨 Tipo",
            'nome_familia': "👨‍👩‍👧‍👦 Nombre de la Familia",
            'frase_familia': "💬 Frase de la Familia",
            'tipo_peca': "📦 Tipo de Pieza",
            'nome_peca': "📝 Nombre de la Pieza",
            'nome_cartoon': "🎭 Nombre del Cartoon",
            'frase_elemento': "💬 Frase del Elemento",
            'nome_no_cartoon': "🎭 Nombre en el Cartoon",
            'frase_na_box': "💬 Frase en la Caja",
            'estilo': "🖌 Estilo",
            'profissao': "💼 Profesión",
            'objetos_personalizados': "🎯 Objetos Personalizados",
            'super_heroi': "🦸‍♂️ Superhéroe",
            'total_elementos': "👥 Total de Elementos",
            'adultos': "👨‍👩 Adultos",
            'criancas': "👧🧒 Niños",
            'animais': "🐱🐶 Animales",
            'nome_animal': "🐾 Nombre del Animal",
            'tipo_animal': "🐕 Tipo de Animal",
            'tamanho': "📏 Tamaño",
            'foto': "📸 Foto",
            'valores': "💵 VALORES",
            'preco_anterior': "Precio anterior",
            'desconto': "Descuento",
            'total_pagar': "TOTAL A PAGAR",
            'impostos_frete': "Impuestos y Envío Incluidos",
            'klarna_disponivel': "💎 KLARNA DISPONIBLE",
            'klarna_texto': "Pague en 3 cuotas de",
            'klarna_juros': "SIN INTERESES",
            'alerta_tempo': "⚠️ ¡Tienes 10 minutos para efectuar el pago!",
            'guardar_id': "¡Guarde el ID del pedido para referencia futura!",
            'clique_pagar': "Haz clic abajo para pagar: 👇",
            'botao_pagar': "💳 Pagar con Tarjeta",
            'minutos': "minutos"
        },
        'italiano': {
            'titulo': "🧾 RIEPILOGO FINALE PER PAGAMENTO",
            'id_pedido': "🆔 ID Ordine",
            'data': "📅 Data",
            'pais_envio': "🌍 Paese di Spedizione",
            'moeda': "💰 Valuta",
            'tempo_pagar': "⏰ Tempo per pagare",
            'dados_pessoais': "👤 DATI PERSONALI",
            'nome': "👤 Nome",
            'email': "📧 Email",
            'pais': "🌍 Paese",
            'telefone': "📱 Telefono",
            'detalhes_cartoon': "🎨 DETTAGLI DEL CARTOON",
            'tipo': "🎨 Tipo",
            'nome_familia': "👨‍👩‍👧‍👦 Nome della Famiglia",
            'frase_familia': "💬 Frase della Famiglia",
            'tipo_peca': "📦 Tipo di Pezzo",
            'nome_peca': "📝 Nome del Pezzo",
            'nome_cartoon': "🎭 Nome del Cartoon",
            'frase_elemento': "💬 Frase dell'Elemento",
            'nome_no_cartoon': "🎭 Nome nel Cartoon",
            'frase_na_box': "💬 Frase nella Scatola",
            'estilo': "🖌 Stile",
            'profissao': "💼 Professione",
            'objetos_personalizados': "🎯 Oggetti Personalizzati",
            'super_heroi': "🦸‍♂️ Supereroe",
            'total_elementos': "👥 Totale Elementi",
            'adultos': "👨‍👩 Adulti",
            'criancas': "👧🧒 Bambini",
            'animais': "🐱🐶 Animali",
            'nome_animal': "🐾 Nome dell'Animale",
            'tipo_animal': "🐕 Tipo di Animale",
            'tamanho': "📏 Dimensione",
            'foto': "📸 Foto",
            'valores': "💵 VALORI",
            'preco_anterior': "Prezzo precedente",
            'desconto': "Sconto",
            'total_pagar': "TOTALE DA PAGARE",
            'impostos_frete': "Tasse e Spedizione Incluse",
            'klarna_disponivel': "💎 KLARNA DISPONIBILE",
            'klarna_texto': "Paga in 3 rate da",
            'klarna_juros': "SENZA INTERESSI",
            'alerta_tempo': "⚠️ Hai 10 minuti per effettuare il pagamento!",
            'guardar_id': "Conserva l'ID dell'ordine per riferimento futuro!",
            'clique_pagar': "Clicca qui sotto per pagare: 👇",
            'botao_pagar': "💳 Paga con Carta",
            'minutos': "minuti"
        },
        'alemao': {
            'titulo': "🧾 ABSCHLIESSENDE ZUSAMMENFASSUNG FÜR ZAHLUNG",
            'id_pedido': "🆔 Bestell-ID",
            'data': "📅 Datum",
            'pais_envio': "🌍 Versandland",
            'moeda': "💰 Währung",
            'tempo_pagar': "⏰ Zeit zum Bezahlen",
            'dados_pessoais': "👤 PERSÖNLICHE DATEN",
            'nome': "👤 Name",
            'email': "📧 Email",
            'pais': "🌍 Land",
            'telefone': "📱 Telefon",
            'detalhes_cartoon': "🎨 CARTOON-DETAILS",
            'tipo': "🎨 Typ",
            'nome_familia': "👨‍👩‍👧‍👦 Familienname",
            'frase_familia': "💬 Familienspruch",
            'tipo_peca': "📦 Stücktyp",
            'nome_peca': "📝 Stückname",
            'nome_cartoon': "🎭 Cartoon-Name",
            'frase_elemento': "💬 Element-Spruch",
            'nome_no_cartoon': "🎭 Name im Cartoon",
            'frase_na_box': "💬 Box-Spruch",
            'estilo': "🖌 Stil",
            'profissao': "💼 Beruf",
            'objetos_personalizados': "🎯 Benutzerdefinierte Objekte",
            'super_heroi': "🦸‍♂️ Superheld",
            'total_elementos': "👥 Gesamtelemente",
            'adultos': "👨‍👩 Erwachsene",
            'criancas': "👧🧒 Kinder",
            'animais': "🐱🐶 Tiere",
            'nome_animal': "🐾 Tiername",
            'tipo_animal': "🐕 Tierart",
            'tamanho': "📏 Größe",
            'foto': "📸 Foto",
            'valores': "💵 WERTE",
            'preco_anterior': "Vorheriger Preis",
            'desconto': "Rabatt",
            'total_pagar': "GESAMTBETRAG ZU ZAHLEN",
            'impostos_frete': "Steuern und Versand inklusive",
            'klarna_disponivel': "💎 KLARNA VERFÜGBAR",
            'klarna_texto': "Zahlen Sie in 3 Raten à",
            'klarna_juros': "OHNE ZINSEN",
            'alerta_tempo': "⚠️ Sie haben 10 Minuten, um die Zahlung vorzunehmen!",
            'guardar_id': "Bewahren Sie die Bestell-ID für zukünftige Referenz auf!",
            'clique_pagar': "Klicken Sie unten zum Bezahlen: 👇",
            'botao_pagar': "💳 Mit Karte bezahlen",
            'minutos': "Minuten"
        },
        'frances': {
            'titulo': "🧾 RÉSUMÉ FINAL POUR PAIEMENT",
            'id_pedido': "🆔 ID de Commande",
            'data': "📅 Date",
            'pais_envio': "🌍 Pays de Livraison",
            'moeda': "💰 Devise",
            'tempo_pagar': "⏰ Temps pour payer",
            'dados_pessoais': "👤 DONNÉES PERSONNELLES",
            'nome': "👤 Nom",
            'email': "📧 Email",
            'pais': "🌍 Pays",
            'telefone': "📱 Téléphone",
            'detalhes_cartoon': "🎨 DÉTAILS DU DESSIN ANIMÉ",
            'tipo': "🎨 Type",
            'nome_familia': "👨‍👩‍👧‍👦 Nom de Famille",
            'frase_familia': "💬 Phrase de Famille",
            'tipo_peca': "📦 Type de Pièce",
            'nome_peca': "📝 Nom de la Pièce",
            'nome_cartoon': "🎭 Nom du Dessin Animé",
            'frase_elemento': "💬 Phrase de l'Élément",
            'nome_no_cartoon': "🎭 Nom dans le Dessin Animé",
            'frase_na_box': "💬 Phrase sur la Boîte",
            'estilo': "🖌 Style",
            'profissao': "💼 Profession",
            'objetos_personalizados': "🎯 Objets Personnalisés",
            'super_heroi': "🦸‍♂️ Super-héros",
            'total_elementos': "👥 Total des Éléments",
            'adultos': "👨‍👩 Adultes",
            'criancas': "👧🧒 Enfants",
            'animais': "🐱🐶 Animaux",
            'nome_animal': "🐾 Nom de l'Animal",
            'tipo_animal': "🐕 Type d'Animal",
            'tamanho': "📏 Taille",
            'foto': "📸 Photo",
            'valores': "💵 VALEURS",
            'preco_anterior': "Prix précédent",
            'desconto': "Réduction",
            'total_pagar': "TOTAL À PAYER",
            'impostos_frete': "Taxes et Livraison Incluses",
            'klarna_disponivel': "💎 KLARNA DISPONIBLE",
            'klarna_texto': "Payez en 3 versements de",
            'klarna_juros': "SANS INTÉRÊTS",
            'alerta_tempo': "⚠️ Vous avez 10 minutes pour effectuer le pagamento !",
            'guardar_id': "Conservez l'ID de commande pour référence future !",
            'clique_pagar': "Cliquez ci-dessous pour payer : 👇",
            'botao_pagar': "💳 Payer avec Carte",
            'minutos': "minutes"
        }
    }
    
    textos = textos_resumo.get(idioma, textos_resumo['portugues'])
    
    # 🔥 VERSÃO COM HTML INCLUINDO TODOS OS CAMPOS
    texto = f"""<b>{textos['titulo']}</b>

<b>{textos['id_pedido']}:</b> {pedido_id}
<b>{textos['data']}:</b> {data_pedido}
<b>{textos['pais_envio']}:</b> {pais_ingles}  
<b>{textos['moeda']}:</b> {totais['moeda']} {totais['simbolo_moeda']}
<b>{textos['tempo_pagar']}:</b> 10 {textos['minutos']}

<b>{textos['dados_pessoais']}:</b>
• <b>{textos['nome']}:</b> {nome}
• <b>{textos['email']}:</b> {email}
• <b>{textos['pais']}:</b> {pais_ingles} 
• <b>{textos['telefone']}:</b> {contacto}

<b>{textos['detalhes_cartoon']}:</b>
• <b>{textos['tipo']}:</b> {tipo}"""

    # 🔥 CAMPOS ESPECÍFICOS DA FAMILY
    if nome_family:
        texto += f"\n• <b>{textos['nome_familia']}:</b> {nome_family}"
    if frase_family and frase_family != "Não adicionou frase":
        texto += f"\n• <b>{textos['frase_familia']}:</b> \"{frase_family}\""
    
    # 🔥 CAMPOS ESPECÍFICOS DO PERSONALIZADO
    if tipo_personalizado:
        texto += f"\n• <b>{textos['tipo_peca']}:</b> {tipo_personalizado}"
    if nome_peca_personalizado:
        texto += f"\n• <b>{textos['nome_peca']}:</b> {nome_peca_personalizado}"
    if nome_personalizado:
        texto += f"\n• <b>{textos['nome_cartoon']}:</b> {nome_personalizado}"
    if frase_personalizado and frase_personalizado != "Não adicionou frase":
        texto += f"\n• <b>{textos['frase_elemento']}:</b> \"{frase_personalizado}\""
    
    # 🔥 CAMPOS DE PERSONALIZAÇÃO DA BOX
    if nome_cartoon:
        texto += f"\n• <b>{textos['nome_no_cartoon']}:</b> {nome_cartoon}"
    if frase_cartoon and frase_cartoon != "Não adicionou frase":
        texto += f"\n• <b>{textos['frase_na_box']}:</b> \"{frase_cartoon}\""
    
    if estilo:
        texto += f"\n• <b>{textos['estilo']}:</b> {estilo}"
    
    # 🔥 RESTANTES CAMPOS PERSONALIZADOS
    if profissao:
        texto += f"\n• <b>{textos['profissao']}:</b> {profissao}"
    if objetos_office:
        texto += f"\n• <b>{textos['objetos_personalizados']}:</b> {objetos_office}"
    if super_heroi:
        texto += f"\n• <b>{textos['super_heroi']}:</b> {super_heroi}"
    
    # 🔥 CAMPOS DA FAMILY (ORIGINAIS)
    if elementos_family:
        texto += f"\n• <b>{textos['total_elementos']}:</b> {elementos_family}"
    if adultos_family:
        texto += f"\n• <b>{textos['adultos']}:</b> {adultos_family}"
    if criancas_family:
        texto += f"\n• <b>{textos['criancas']}:</b> {criancas_family}"
    if animais_family:
        texto += f"\n• <b>{textos['animais']}:</b> {animais_family}"
    if nome_animal:
        texto += f"\n• <b>{textos['nome_animal']}:</b> {nome_animal}"
    if tipo_animal:
        texto += f"\n• <b>{textos['tipo_animal']}:</b> {tipo_animal}"
    
    texto += f"""
• <b>{textos['tamanho']}:</b> {tamanho}
• <b>{textos['foto']}:</b> {foto_recebida} ({nome_foto})

<b>{textos['valores']}:</b>
• <b>{textos['preco_anterior']}:</b> {totais['simbolo_moeda']}{preco_anterior:.2f}❌ 
• <b>{textos['desconto']}:</b> {totais['simbolo_moeda']}{desconto:.2f} ({percentual_desconto}% OFF)
• 💰 <b>{textos['total_pagar']}: {totais['simbolo_moeda']}{totais['total']:.2f}</b>
• 📝 <b>{textos['impostos_frete']}</b>"""

    # 🔥 ADICIONAR MENSAGEM DO KLARNA APENAS PARA PAÍSES QUE SUPORTAM
    paises_sem_klarna = ["canada", "brasil", "estados unidos"]
    pais_lower = pais_ingles.lower()
    
    if pais_lower not in paises_sem_klarna:
        texto += f"""

<b>{textos['klarna_disponivel']}:</b>
{textos['klarna_texto']} <b>{totais['simbolo_moeda']}{valor_klarna:.2f} {textos['klarna_juros']}</b>"""
        print(f"✅ Klarna disponível para {pais_ingles}")
    else:
        print(f"🚫 Klarna NÃO disponível para {pais_ingles}")

    texto += f"""

<b>{textos['alerta_tempo']}</b>
<b>{textos['guardar_id']}</b>

<b>{textos['clique_pagar']}</b>"""

    # BOTÕES
    botoes = [
        [InlineKeyboardButton(textos['botao_pagar'], callback_data="pagar_stripe")]
    ]
    
    # ENVIAR MENSAGEM
    try:
        mensagem = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=texto, 
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup(botoes)
        )
        print(f"✅ Resumo de pagamento enviado com todos os campos da Family | Idioma: {idioma}")
        print(f"✅ País mostrado como: {pais_ingles} (em inglês)")
        
    except Exception as e:
        print(f"❌ Erro ao enviar com HTML: {e}")
    
    # 🔥 🔥 🔥 AGORA SIM: TEMPORIZADOR (PEDIDO JÁ ESTÁ NO REGISTRO)
    print(f"⏰ Iniciando temporizador de 10min para pedido #{pedido_id}")
    await iniciar_temporizador(context, pedido_id, query.message.chat_id, mensagem.message_id, idioma)



# ======================= SISTEMA DE TEMPORIZADOR =======================
async def iniciar_temporizador(context, pedido_id, chat_id, message_id, idioma=None):
    """Temporizador NÃO-BLOQUEANTE para pagamento - COM TRADUÇÃO"""
    try:
        # 🔥 SE NÃO VEIO COM IDIOMA, PEGAR DO PEDIDO
        if idioma is None and pedido_id in PEDIDOS_REGISTO:
            idioma = PEDIDOS_REGISTO[pedido_id].get('idioma', 'portugues')
        elif idioma is None:
            idioma = 'portugues'
            
        print(f"⏰ Temporizador INICIADO (não-bloqueante) para pedido #{pedido_id} | Idioma: {idioma}")
        
        # 🔥 TEXTOS DO TEMPORIZADOR POR IDIOMA
        textos_temporizador = {
            'portugues': {
                'expirado_titulo': "❌ *PAGAMENTO EXPIRADO*",
                'expirado_mensagem': "O pedido expirou por falta de pagamento.",
                'tentar_novamente': "💳 Tentar Novamente",
                'reportar_problema': "🤔 Reportar Problema",
                'confirmacao_expirado': "✅ Temporizador configurado - Pedido expira em 10 minutos",
                'pedido_expirado_log': "❌ PEDIDO EXPIRADO"
            },
            'ingles': {
                'expirado_titulo': "❌ *PAYMENT EXPIRED*",
                'expirado_mensagem': "The order expired due to lack of payment.",
                'tentar_novamente': "💳 Try Again",
                'reportar_problema': "🤔 Report Problem",
                'confirmacao_expirado': "✅ Timer configured - Order expires in 10 minutes",
                'pedido_expirado_log': "❌ ORDER EXPIRED"
            },
            'espanhol': {
                'expirado_titulo': "❌ *PAGO EXPIRADO*",
                'expirado_mensagem': "El pedido expiró por falta de pago.",
                'tentar_novamente': "💳 Intentar de Nuevo",
                'reportar_problema': "🤔 Informar Problema",
                'confirmacao_expirado': "✅ Temporizador configurado - Pedido expira en 10 minutos",
                'pedido_expirado_log': "❌ PEDIDO EXPIRADO"
            },
            'italiano': {
                'expirado_titulo': "❌ *PAGAMENTO SCADUTO*",
                'expirado_mensagem': "L'ordine è scaduto per mancanza di pagamento.",
                'tentar_novamente': "💳 Riprova",
                'reportar_problema': "🤔 Segnalare Problema",
                'confirmacao_expirado': "✅ Timer configurato - L'ordine scade in 10 minuti",
                'pedido_expirado_log': "❌ ORDINE SCADUTO"
            },
            'alemao': {
                'expirado_titulo': "❌ *ZAHLUNG ABGELAUFEN*",
                'expirado_mensagem': "Die Bestellung ist aufgrund fehlender Zahlung abgelaufen.",
                'tentar_novamente': "💳 Erneut Versuchen",
                'reportar_problema': "🤔 Problem Melden",
                'confirmacao_expirado': "✅ Timer konfiguriert - Bestellung läuft in 10 Minuten ab",
                'pedido_expirado_log': "❌ BESTELLUNG ABGELAUFEN"
            },
            'frances': {
                'expirado_titulo': "❌ *PAIEMENT EXPIRÉ*",
                'expirado_mensagem': "La commande a expiré en raison d'un manque de paiement.",
                'tentar_novamente': "💳 Réessayer",
                'reportar_problema': "🤔 Signaler un Problème",
                'confirmacao_expirado': "✅ Minuterie configurée - La commande expire dans 10 minutes",
                'pedido_expirado_log': "❌ COMMANDE EXPIRÉE"
            }
        }
        
        textos = textos_temporizador.get(idioma, textos_temporizador['portugues'])
        print(f"{textos['confirmacao_expirado']}")
        
        # 🔥 CRIA UMA TASK SEPARADA que não bloqueia o bot
        async def temporizador_task():
            try:
                print(f"⏰ Task temporizador iniciada para #{pedido_id} | Idioma: {idioma}")
                await asyncio.sleep(60)  # ⬅️ 10 minutos (600 segundos)
                
                # Verificar se o pedido ainda existe e está pendente
                if (pedido_id in PEDIDOS_REGISTO and 
                    PEDIDOS_REGISTO[pedido_id]["status"] == "pendente" and
                    "timer_task" in PEDIDOS_REGISTO[pedido_id]):  # ⬅️ Só expirar se timer ainda estiver ativo
                    
                    PEDIDOS_REGISTO[pedido_id]["status"] = "expirado"
                    atualizar_estatistica("pedidos_expirados")
                    
                    print(f"{textos['pedido_expirado_log']}: #{pedido_id}")
                    
                    # 🔥 ENVIAR NOTIFICAÇÃO DE EXPIRAÇÃO COM TRADUÇÃO
                    await context.bot.edit_message_text(
                        chat_id=chat_id,
                        message_id=message_id,
                        text=f"{textos['expirado_titulo']}\n\n{textos['expirado_mensagem']}",
                        parse_mode="Markdown",
                        reply_markup=InlineKeyboardMarkup([
                            [
                                InlineKeyboardButton(
                                    textos['tentar_novamente'], 
                                    callback_data=f"recuperar_pagar_{pedido_id}"
                                )
                            ],
                            [
                                InlineKeyboardButton(
                                    textos['reportar_problema'], 
                                    callback_data=f"reportar_problema_{pedido_id}"
                                )
                            ]
                        ])
                    )
                    
                    # 🔥 ADICIONAR LOG DE EXPIRAÇÃO
                    print(f"📝 Pedido #{pedido_id} marcado como expirado no sistema | Idioma: {idioma}")
                    
            except asyncio.CancelledError:
                print(f"✅ Temporizador cancelado - Pedido #{pedido_id} PAGO | Idioma: {idioma}")
            except Exception as e:
                print(f"❌ Erro na task do temporizador: {e} | Idioma: {idioma}")
        
        # 🔥 INICIA A TASK EM BACKGROUND E GUARDA REFERÊNCIA
        timer_task = asyncio.create_task(temporizador_task())
        
        # 🔥 GUARDAR INFORMAÇÕES NO PEDIDO PARA GESTÃO
        if pedido_id in PEDIDOS_REGISTO:
            PEDIDOS_REGISTO[pedido_id]["timer_task"] = timer_task  # ⬅️ GUARDAR PARA PODER CANCELAR
            PEDIDOS_REGISTO[pedido_id]["idioma_temporizador"] = idioma  # 🔥 GUARDAR IDIOMA DO TEMPORIZADOR
            PEDIDOS_REGISTO[pedido_id]["hora_inicio_temporizador"] = datetime.now().strftime("%H:%M:%S")
            PEDIDOS_REGISTO[pedido_id]["hora_expiracao_temporizador"] = (datetime.now() + timedelta(minutes=10)).strftime("%H:%M:%S")
            
            print(f"✅ Task temporizador criada em background para #{pedido_id}")
            print(f"   • Idioma: {idioma}")
            print(f"   • Início: {PEDIDOS_REGISTO[pedido_id]['hora_inicio_temporizador']}")
            print(f"   • Expira: {PEDIDOS_REGISTO[pedido_id]['hora_expiracao_temporizador']}")
        
    except Exception as e:
        print(f"❌ Erro ao iniciar temporizador: {e} | Idioma: {idioma if 'idioma' in locals() else 'não definido'}")




# ======================= SISTEMA DE RECUPERAÇÃO =======================
async def recuperar_pedido(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para o botão 'Tentar Novamente' - COM TRADUÇÃO COMPLETA"""
    query = update.callback_query
    await query.answer()
    
    print(f"🎯 RECUPERAR_PEDIDO CHAMADO - VERSÃO TRADUZIDA")
    
    # Extrair pedido_id do callback_data
    pedido_id = query.data.replace("recuperar_pagar_", "")
    print(f"🔍 Procurando pedido: {pedido_id}")
    
    if pedido_id not in PEDIDOS_REGISTO:
        print(f"❌ Pedido não encontrado no registro: {pedido_id}")
        
        # 🔥 MENSAGEM DE ERRO POR IDIOMA
        idioma = context.user_data.get('idioma', 'portugues')
        
        textos_erro = {
            'portugues': "❌ *Pedido não encontrado!*\n\nPor favor, inicie um novo pedido com /start",
            'ingles': "❌ *Order not found!*\n\nPlease start a new order with /start",
            'espanhol': "❌ *¡Pedido no encontrado!*\n\nPor favor, inicie un nuevo pedido con /start",
            'italiano': "❌ *Ordine non trovato!*\n\nPer favore, iniziare un nuovo ordine con /start",
            'alemao': "❌ *Bestellung nicht gefunden!*\n\nBitte starten Sie eine neue Bestellung mit /start",
            'frances': "❌ *Commande non trouvée !*\n\nVeuillez démarrer une nouvelle commande avec /start"
        }
        
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    chat_id = query.message.chat_id
    
    # 🔥 PEGAR IDIOMA DO PEDIDO OU DO USER_DATA
    idioma = pedido.get('idioma', context.user_data.get('idioma', 'portugues'))
    print(f"🌐 Idioma detectado: {idioma}")
    
    # 🔥 CANCELAR QUALQUER TEMPORIZADOR ATIVO
    await cancelar_temporizadores_pedido(pedido_id)
    
    # ATUALIZAR ESTATÍSTICAS
    atualizar_estatistica("tentativas_recuperacao")
    
    print(f"✅ Pedido encontrado: #{pedido_id} | Idioma: {idioma}")

    try:
        # 🔥 PASSO 1: DEFINIR MÉTODOS DE PAGAMENTO POR PAÍS - MESMA ESTRUTURA
        def get_payment_methods(pais):
            """Retorna métodos de pagamento baseado no país"""
            
            def get_country_code(pais_nome):
                mapeamento_paises = {
                    "portugal": "PT",
                    "espanha": "ES", 
                    "franca": "FR",
                    "alemanha": "DE",
                    "belgica": "BE",
                    "reino unido": "GB",
                    "estados unidos": "US",
                    "paises baixos": "NL",
                    "brasil": "BR",
                    "irlanda": "IE",
                    "italia": "IT",
                    "luxemburgo": "LU",
                    "canada": "CA"
                }
                return mapeamento_paises.get(pais_nome.lower(), pais_nome.upper())
            
            country_code = get_country_code(pais)
            print(f"🔍 País recebido: '{pais}' → Código: '{country_code}'")
            
            # 🔥 MESMO payment_methods_by_country DO pagar_stripe
            payment_methods_by_country = {
                "PT": ["card", "paypal", "link", "klarna", "mb_way", "sepa_debit"],
                "ES": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "FR": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "DE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "BE": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "GB": ["card", "paypal", "link", "klarna"],
                "US": ["card", "paypal", "link"],
                "NL": ["card", "paypal", "link", "klarna", "ideal", "sepa_debit"],
                "BR": ["card", "link"],
                "IE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "IT": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "LU": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "CA": ["card", "paypal", "link"]
            }
            
            methods = payment_methods_by_country.get(country_code, ["card", "link"])
            print(f"💳 Métodos de pagamento para {pais} ({country_code}): {methods}")
            return methods

        # 🔥 OBTER MÉTODOS REAIS PARA ESTE PAÍS
        metodos_reais = get_payment_methods(pedido['pais'])
        
        # 🔥 CRIAR TEXTO DINÂMICO DOS MÉTODOS - COM TRADUÇÃO
        def formatar_metodos(metodos, pais, idioma):
            """Formata os métodos de pagamento para exibição em diferentes idiomas"""
            
            # 🔥 DICIONÁRIO DE TRADUÇÃO DE MÉTODOS DE PAGAMENTO
            nomes_metodos = {
                'portugues': {
                    "card": "Cartão de Crédito/Débito",
                    "paypal": "PayPal", 
                    "link": "Link (inclui Apple Pay/Google Pay)",
                    "klarna": "Klarna (Pague em 3x sem juros)",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'ingles': {
                    "card": "Credit/Debit Card",
                    "paypal": "PayPal", 
                    "link": "Link (includes Apple Pay/Google Pay)",
                    "klarna": "Klarna (Pay in 3 installments)",
                    "sepa_debit": "SEPA Debit",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'espanhol': {
                    "card": "Tarjeta de Crédito/Débito",
                    "paypal": "PayPal", 
                    "link": "Link (incluye Apple Pay/Google Pay)",
                    "klarna": "Klarna (Paga en 3 cuotas)",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'italiano': {
                    "card": "Carta di Credito/Debito",
                    "paypal": "PayPal", 
                    "link": "Link (include Apple Pay/Google Pay)",
                    "klarna": "Klarna (Paga in 3 rate)",
                    "sepa_debit": "Addebito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'alemao': {
                    "card": "Kredit-/Debitkarte",
                    "paypal": "PayPal", 
                    "link": "Link (enthält Apple Pay/Google Pay)",
                    "klarna": "Klarna (In 3 Raten zahlen)",
                    "sepa_debit": "SEPA-Lastschrift",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'frances': {
                    "card": "Carte de Crédit/Débit",
                    "paypal": "PayPal", 
                    "link": "Link (comprend Apple Pay/Google Pay)",
                    "klarna": "Klarna (Payer en 3 fois)",
                    "sepa_debit": "Prélèvement SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                }
            }
            
            textos = []
            nomes = nomes_metodos.get(idioma, nomes_metodos['portugues'])
            
            for metodo in metodos:
                if metodo in nomes:
                    textos.append(nomes[metodo])
                else:
                    textos.append(metodo.capitalize())
            
            return ", ".join(textos)

        # 🔥 PASSO 2: VERIFICAR CONFIGURAÇÃO DE WALLETS
        def verificar_config_wallets():
            """Verifica se as wallets estão configuradas corretamente"""
            try:
                apple_domains = stripe.ApplePayDomain.list()
                print("🍎 Domínios Apple Pay configurados:")
                for domain in apple_domains.data:
                    print(f"   - {domain.domain}")
                
                seu_dominio = "unceased-bibliothecal-donette.ngrok-free.dev"
                dominios_apple = [d.domain for d in apple_domains.data]
                if seu_dominio in dominios_apple:
                    print("✅ Domínio ngrok configurado no Apple Pay!")
                    return True
                else:
                    print("⚠️ Domínio ngrok NÃO configurado no Apple Pay")
                    return False
                    
            except Exception as e:
                print(f"❌ Erro ao verificar wallets: {e}")
                return False

        wallets_configuradas = verificar_config_wallets()

        # 🔥 PASSO 3: CRIAR SESSÃO STRIPE
        print("🔗 Criando Checkout Session para recuperação...")
        
        # 🔥 TEXTOS DE CHECKOUT POR IDIOMA
        textos_checkout = {
            'portugues': {
                "shipping_message": "📦 Enviaremos o seu cartoon personalizado para este endereço!",
                "submit_message": "✨ Obrigado! Vamos criar um cartoon incrível para si!"
            },
            'ingles': {
                "shipping_message": "📦 We'll send your personalized cartoon to this address!",
                "submit_message": "✨ Thank you! We'll create an amazing cartoon for you!"
            },
            'espanhol': {
                "shipping_message": "📦 ¡Enviaremos tu cartoon personalizado a esta dirección!",
                "submit_message": "✨ ¡Gracias! ¡Crearemos un cartoon increíble para ti!"
            },
            'italiano': {
                "shipping_message": "📦 Spediremo il tuo cartoon personalizzato a questo indirizzo!",
                "submit_message": "✨ Grazie! Creeremo un cartoon incredibile per te!"
            },
            'alemao': {
                "shipping_message": "📦 Wir senden Ihren personalisierten Cartoon an diese Adresse!",
                "submit_message": "✨ Danke! Wir erstellen einen fantastischen Cartoon für Sie!"
            },
            'frances': {
                "shipping_message": "📦 Nous enverrons votre dessin animé personnalisé à cette adresse !",
                "submit_message": "✨ Merci ! Nous créerons un dessin animé incroyable pour vous !"
            }
        }
        
        textos = textos_checkout.get(idioma, textos_checkout['portugues'])
        
        session_config = {
            "payment_method_types": metodos_reais,
            "mode": "payment",
            "customer_email": pedido["email"],
            
            "payment_method_options": {
                "card": {
                    "request_three_d_secure": "automatic"
                }
            },
            
            "shipping_address_collection": {
                "allowed_countries": [
                    "PT", "ES", "FR", "DE", "BE", "GB", "US", "NL", "BR", "IE", "IT", "LU", "CA"
                ]
            },
            
            "custom_text": {
                "shipping_address": {
                    "message": textos["shipping_message"]
                },
                "submit": {
                    "message": textos["submit_message"]
                }
            },
            
            "line_items": [{
                "price_data": {
                    "currency": pedido["moeda"].lower(),
                    "product_data": {
                        "name": f"Cartoon Personalizado - {pedido['tipo_cartoon']}",
                        "description": f"Recuperação Pedido #{pedido_id} - Para {pedido['nome']}",
                    },
                    "unit_amount": int(pedido["total"] * 100),
                },
                "quantity": 1
            }],
            
            "success_url": f"https://t.me/plan3d_bot?start=payment_success_{pedido_id}",
            "cancel_url": f"https://t.me/plan3d_bot?start=payment_cancelled_{pedido_id}",
            
            "metadata": {
                "pedido_id": pedido_id,
                "chat_id": str(chat_id),
                "pais": pedido['pais'],
                "moeda": pedido["moeda"],
                "total_pago": str(pedido["total"]),
                "nome_cliente": pedido['nome'],
                "tipo_cartoon": pedido['tipo_cartoon'],
                "tipo_sessao": "recuperacao",
                "tentativa_numero": str(pedido.get("tentativas_recuperacao", 1)),
                "wallets_habilitadas": str(wallets_configuradas),
                "idioma": idioma,
            },
            
            "expires_at": int((datetime.now() + timedelta(minutes=30)).timestamp()),
        }

        # 🔥 CONFIGURAÇÃO ESPECÍFICA PARA WALLETS
        paises_com_wallets = ["Reino Unido", "Estados Unidos", "Brasil", "Irlanda", 
                            "França", "Alemanha", "Itália", "Espanha", "Portugal", 
                            "Países Baixos", "Bélgica", "Luxemburgo", "Canadá"]
        
        if pedido['pais'] in paises_com_wallets and "link" in metodos_reais:
            print(f"📱 Configurando Apple Pay/Google Pay para {pedido['pais']}")
            session_config["payment_method_options"]["link"] = {"persistent_token": None}

        # 🔥 CRIAR A SESSÃO
        session = stripe.checkout.Session.create(**session_config)

        print(f"✅ CHECKOUT SESSION CRIADA: {session.id} | Idioma: {idioma}")

        # 🔥 PASSO 4: ATUALIZAR PEDIDO
        pedido["session_id_recuperacao"] = session.id
        pedido["payment_intent_id"] = session.payment_intent
        pedido["wallets_configuradas"] = wallets_configuradas
        pedido["tentativas_recuperacao"] = pedido.get("tentativas_recuperacao", 0) + 1
        pedido["data_recuperacao"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        print(f"📊 Pedido atualizado para recuperação | Idioma: {idioma}")

        # 🔥 PASSO 5: MENSAGEM FINAL COM TRADUÇÃO
        texto_metodos = formatar_metodos(metodos_reais, pedido['pais'], idioma)
        
        # 🔥 TEXTOS DA MENSAGEM POR IDIOMA
        textos_mensagem = {
            'portugues': {
                "titulo": "🔄 *TENTANDO PAGAMENTO NOVAMENTE* 🔄",
                "cliente": "👤 *Cliente:*",
                "pais_envio": "🌍 *País de Envio:*",
                "moeda": "💰 *Moeda:*",
                "total": "💳 **TOTAL A PAGAR:**",
                "pedido": "🆔 **Pedido:**",
                "checkout_instrucoes": "📋 *No checkout será pedido:*",
                "endereco": "1️⃣ **Endereço de entrega completo**",
                "metodo": "2️⃣ **Método de pagamento**",
                "metodos_disponiveis": "💳 *Métodos disponíveis:*",
                "seguro": "🔒 *Pagamento 100% seguro via Stripe*",
                "tempo": "⏰ *Tem 30 minutos para efetuar o pagamento*",
                "botao_pagar": "💳 PAGAR AGORA →"
            },
            'ingles': {
                "titulo": "🔄 *TRYING PAYMENT AGAIN* 🔄",
                "cliente": "👤 *Customer:*",
                "pais_envio": "🌍 *Shipping Country:*",
                "moeda": "💰 *Currency:*",
                "total": "💳 **TOTAL TO PAY:**",
                "pedido": "🆔 **Order:**",
                "checkout_instrucoes": "📋 *In checkout you will be asked for:*",
                "endereco": "1️⃣ **Complete delivery address**",
                "metodo": "2️⃣ **Payment method**",
                "metodos_disponiveis": "💳 *Available methods:*",
                "seguro": "🔒 *100% secure payment via Stripe*",
                "tempo": "⏰ *You have 30 minutes to make the payment*",
                "botao_pagar": "💳 PAY NOW →"
            },
            'espanhol': {
                "titulo": "🔄 *INTENTANDO PAGO NUEVAMENTE* 🔄",
                "cliente": "👤 *Cliente:*",
                "pais_envio": "🌍 *País de Envío:*",
                "moeda": "💰 *Moneda:*",
                "total": "💳 **TOTAL A PAGAR:**",
                "pedido": "🆔 **Pedido:**",
                "checkout_instrucoes": "📋 *En el checkout se pedirá:*",
                "endereco": "1️⃣ **Dirección de entrega completa**",
                "metodo": "2️⃣ **Método de pago**",
                "metodos_disponiveis": "💳 *Métodos disponibles:*",
                "seguro": "🔒 *Pago 100% seguro vía Stripe*",
                "tempo": "⏰ *Tienes 30 minutos para efectuar el pago*",
                "botao_pagar": "💳 PAGAR AHORA →"
            },
            'italiano': {
                "titulo": "🔄 *TENTANDO PAGAMENTO DI NUOVO* 🔄",
                "cliente": "👤 *Cliente:*",
                "pais_envio": "🌍 *Paese di Spedizione:*",
                "moeda": "💰 *Valuta:*",
                "total": "💳 **TOTALE DA PAGARE:**",
                "pedido": "🆔 **Ordine:**",
                "checkout_instrucoes": "📋 *Nel checkout verrà richiesto:*",
                "endereco": "1️⃣ **Indirizzo di consegna completo**",
                "metodo": "2️⃣ **Metodo di pagamento**",
                "metodos_disponiveis": "💳 *Metodi disponibili:*",
                "seguro": "🔒 *Pagamento 100% sicuro via Stripe*",
                "tempo": "⏰ *Hai 30 minuti per effettuare il pagamento*",
                "botao_pagar": "💳 PAGA ORA →"
            },
            'alemao': {
                "titulo": "🔄 *ZAHLLUNG ERNEUT VERSUCHEN* 🔄",
                "cliente": "👤 *Kunde:*",
                "pais_envio": "🌍 *Versandland:*",
                "moeda": "💰 *Währung:*",
                "total": "💳 **GESAMTBETRAG ZU ZAHLEN:**",
                "pedido": "🆔 **Bestellung:**",
                "checkout_instrucoes": "📋 *Im Checkout werden Sie aufgefordert:*",
                "endereco": "1️⃣ **Vollständige Lieferadresse**",
                "metodo": "2️⃣ **Zahlungsmethode**",
                "metodos_disponiveis": "💳 *Verfügbare Methoden:*",
                "seguro": "🔒 *100% sichere Zahlung über Stripe*",
                "tempo": "⏰ *Sie haben 30 Minuten, um die Zahlung vorzunehmen*",
                "botao_pagar": "💳 JETZT BEZAHLEN →"
            },
            'frances': {
                "titulo": "🔄 *ESSAI DE PAIEMENT À NOUVEAU* 🔄",
                "cliente": "👤 *Client:*",
                "pais_envio": "🌍 *Pays de Livraison:*",
                "moeda": "💰 *Devise:*",
                "total": "💳 **TOTAL À PAYER:**",
                "pedido": "🆔 **Commande:**",
                "checkout_instrucoes": "📋 *Dans le checkout, il vous sera demandé:*",
                "endereco": "1️⃣ **Adresse de livraison complète**",
                "metodo": "2️⃣ **Méthode de paiement**",
                "metodos_disponiveis": "💳 *Méthodes disponibles:*",
                "seguro": "🔒 *Paiement 100% sécurisé via Stripe*",
                "tempo": "⏰ *Vous avez 30 minutes pour effectuer le paiement*",
                "botao_pagar": "💳 PAYER MAINTENANT →"
            }
        }
        
        textos_msg = textos_mensagem.get(idioma, textos_mensagem['portugues'])
        
        mensagem = f"""{textos_msg['titulo']}

{textos_msg['cliente']} {pedido['nome']}
{textos_msg['pais_envio']} {pedido['pais']}
{textos_msg['moeda']} {pedido['moeda'].upper()} {pedido['simbolo_moeda']}

{textos_msg['total']} {pedido['simbolo_moeda']}{pedido['total']:.2f}
{textos_msg['pedido']} #{pedido_id}

{textos_msg['checkout_instrucoes']}
{textos_msg['endereco']}
{textos_msg['metodo']}

{textos_msg['metodos_disponiveis']} {texto_metodos}

{textos_msg['seguro']}

{textos_msg['tempo']}

{textos_mensagem[idioma]['clique_abaixo'] if 'clique_abaixo' in textos_mensagem[idioma] else 'Clique abaixo para pagar: 👇'}"""

        await query.edit_message_text(
            text=mensagem,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos_msg['botao_pagar'], url=session.url)]
            ])
        )
        
        print(f"✅ Usuário redirecionado para Checkout (Recuperação) | Idioma: {idioma}")

        # 🔥 INICIAR TEMPORIZADOR PARA RECUPERAÇÃO
        await iniciar_temporizador_recuperacao(context, pedido_id, chat_id, query.message.message_id, idioma)
        
    except Exception as e:
        print(f"❌ ERRO STRIPE NA RECUPERAÇÃO: {str(e)}")
        print(f"🔍 Tipo do erro: {type(e)}")
        print(f"🌐 Idioma do erro: {idioma}")
        
        import traceback
        print(f"🔍 Traceback completo: {traceback.format_exc()}")
        
        # 🔥 TEXTOS DE ERRO POR IDIOMA
        textos_erro_final = {
            'portugues': {
                "mensagem": "❌ *Erro no processamento do pagamento!*\n\nPor favor, tente novamente em alguns segundos.",
                "tentar_novamente": "🔄 Tentar Novamente",
                "suporte": "📞 Suporte"
            },
            'ingles': {
                "mensagem": "❌ *Error processing payment!*\n\nPlease try again in a few seconds.",
                "tentar_novamente": "🔄 Try Again",
                "suporte": "📞 Support"
            },
            'espanhol': {
                "mensagem": "❌ *¡Error al procesar el pago!*\n\nPor favor, intente de nuevo en unos segundos.",
                "tentar_novamente": "🔄 Intentar de Nuevo",
                "suporte": "📞 Soporte"
            },
            'italiano': {
                "mensagem": "❌ *Errore nell'elaborazione del pagamento!*\n\nPer favore, riprova tra pochi secondi.",
                "tentar_novamente": "🔄 Riprova",
                "suporte": "📞 Supporto"
            },
            'alemao': {
                "mensagem": "❌ *Fehler bei der Zahlungsabwicklung!*\n\nBitte versuchen Sie es in einigen Sekunden erneut.",
                "tentar_novamente": "🔄 Erneut Versuchen",
                "suporte": "📞 Unterstützung"
            },
            'frances': {
                "mensagem": "❌ *Erreur de traitement du paiement !*\n\nVeuillez réessayer dans quelques secondes.",
                "tentar_novamente": "🔄 Réessayer",
                "suporte": "📞 Support"
            }
        }
        
        textos_erro = textos_erro_final.get(idioma, textos_erro_final['portugues'])
        
        await query.edit_message_text(
            textos_erro["mensagem"],
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(
                    textos_erro["tentar_novamente"], 
                    callback_data=f"recuperar_pagar_{pedido_id}"
                )],
                [InlineKeyboardButton(
                    textos_erro["suporte"], 
                    callback_data=f"todas_recusadas_{pedido_id}"
                )]
            ])
        )





async def cancelar_temporizadores_pedido(pedido_id):
    """Cancela todos os temporizadores ativos de um pedido"""
    try:
        print(f"🔍🔍🔍 CANCELAR_TEMPORIZADORES_PEDIDO chamado para: #{pedido_id}")
        
        if pedido_id in PEDIDOS_REGISTO:
            pedido = PEDIDOS_REGISTO[pedido_id]
            
            # 🔥 CANCELAR TEMPORIZADOR PRINCIPAL
            if "timer_task" in pedido:
                try:
                    print(f"⏹️ Tentando cancelar timer_task para #{pedido_id}")
                    pedido["timer_task"].cancel()
                    print(f"✅ timer_task cancelado para #{pedido_id}")
                except Exception as e:
                    print(f"⚠️ Erro ao cancelar timer principal: {e}")
                finally:
                    if "timer_task" in pedido:
                        del pedido["timer_task"]
                        print(f"🗑️ timer_task removido do pedido #{pedido_id}")
            
            # 🔥 CANCELAR TEMPORIZADOR DE RECUPERAÇÃO
            if "timer_recuperacao" in pedido:
                try:
                    print(f"⏹️ Tentando cancelar timer_recuperacao para #{pedido_id}")
                    pedido["timer_recuperacao"].cancel()
                    print(f"✅ timer_recuperacao cancelado para #{pedido_id}")
                except Exception as e:
                    print(f"⚠️ Erro ao cancelar timer recuperação: {e}")
                finally:
                    if "timer_recuperacao" in pedido:
                        del pedido["timer_recuperacao"]
                        print(f"🗑️ timer_recuperacao removido do pedido #{pedido_id}")
            
            if "timer_oferta" in pedido:
                try:
                    print(f"⏹️ Tentando cancelar timer_oferta para #{pedido_id}")
                    pedido["timer_oferta"].cancel()
                    temporizadores_cancelados += 1
                    print(f"✅ timer_oferta cancelado para #{pedido_id}")
                except Exception as e:
                    print(f"⚠️ Erro ao cancelar timer oferta: {e}")
                finally:
                    if "timer_oferta" in pedido:
                        del pedido["timer_oferta"]
                        print(f"🗑️ timer_oferta removido do pedido #{pedido_id}")
            
            print(f"✅✅✅ {temporizadores_cancelados} temporizadores cancelados para #{pedido_id}")
        else:
            print(f"❌❌❌ Pedido #{pedido_id} não encontrado no registro")
            
    except Exception as e:
        print(f"❌❌❌ Erro ao cancelar temporizadores: {e}")
        


async def iniciar_temporizador_recuperacao(context, pedido_id, chat_id, message_id, idioma=None):
    """Temporizador de 30 minutos para recuperação - ATUALIZADO"""
    try:
        print(f"⏰⏰⏰ INICIAR_TEMPORIZADOR_RECUPERACAO para #{pedido_id} (30 minutos)")
        
        # 🔥 DETECTAR IDIOMA DO PEDIDO
        idioma = 'portugues'  # padrão
        if pedido_id in PEDIDOS_REGISTO:
            idioma = PEDIDOS_REGISTO[pedido_id].get('idioma', 'portugues')
        
        print(f"🌐 Idioma detectado: {idioma}")
        
        # 🔥 TEXTOS POR IDIOMA (apenas mensagens para o cliente)
        textos = {
            'portugues': {
                'expirado_titulo': "❌ *PAGAMENTO NÃO CONCLUÍDO*",
                'expirado_mensagem': "O tempo para pagamento do pedido `{pedido_id}` expirou.",
                'interessado': "*Se ainda estiver interessado, inicie um novo pedido:* 👇",
                'novo_pedido': "🔄 Novo Pedido",
                'callback': "voltar_inicio"
            },
            'ingles': {
                'expirado_titulo': "❌ *PAYMENT NOT COMPLETED*",
                'expirado_mensagem': "The payment time for order `{pedido_id}` has expired.",
                'interessado': "*If you're still interested, start a new order:* 👇",
                'novo_pedido': "🔄 New Order",
                'callback': "start_new_order"
            },
            'espanhol': {
                'expirado_titulo': "❌ *PAGO NO COMPLETADO*",
                'expirado_mensagem': "El tiempo de pago del pedido `{pedido_id}` ha expirado.",
                'interessado': "*Si aún está interesado, inicie un nuevo pedido:* 👇",
                'novo_pedido': "🔄 Nuevo Pedido",
                'callback': "volver_inicio"
            },
            'italiano': {
                'expirado_titulo': "❌ *PAGAMENTO NON COMPLETATO*",
                'expirado_mensagem': "Il tempo di pagamento dell'ordine `{pedido_id}` è scaduto.",
                'interessado': "*Se sei ancora interessato, inizia un nuovo ordine:* 👇",
                'novo_pedido': "🔄 Nuovo Ordine",
                'callback': "torna_inizio"
            },
            'alemao': {
                'expirado_titulo': "❌ *ZAHLUNG NICHT ABGESCHLOSSEN*",
                'expirado_mensagem': "Die Zahlungsfrist für Bestellung `{pedido_id}` ist abgelaufen.",
                'interessado': "*Wenn Sie noch interessiert sind, starten Sie eine neue Bestellung:* 👇",
                'novo_pedido': "🔄 Neue Bestellung",
                'callback': "zurueck_start"
            },
            'frances': {
                'expirado_titulo': "❌ *PAIEMENT NON TERMINÉ*",
                'expirado_mensagem': "Le délai de paiement de la commande `{pedido_id}` a expiré.",
                'interessado': "*Si vous êtes toujours intéressé, commencez une nouvelle commande :* 👇",
                'novo_pedido': "🔄 Nouvelle Commande",
                'callback': "retour_debut"
            }
        }
        
        # 🔥 USAR PORTUGUÊS COMO FALLBACK
        textos_cliente = textos.get(idioma, textos['portugues'])
        
        async def temporizador_recuperacao_task():
            try:
                print(f"⏰ Task temporizador recuperação iniciada para #{pedido_id}")
                await asyncio.sleep(1800)  # 🔥 30 minutos (1800 segundos)
                
                print(f"🔍 Verificando se pedido #{pedido_id} ainda está em recuperação...")
                
                if (pedido_id in PEDIDOS_REGISTO and 
                    PEDIDOS_REGISTO[pedido_id]["status"] == "recuperando" and
                    "timer_recuperacao" in PEDIDOS_REGISTO[pedido_id]):  # ⬅️ SÓ EXPIRAR SE TIMER AINDA ESTIVER ATIVO
                    
                    pedido = PEDIDOS_REGISTO[pedido_id]
                    pedido["status"] = "expirado_definitivo"
                    
                    # REMOVER DA RECUPERAÇÃO NAS ESTATÍSTICAS
                    ESTATISTICAS["em_recuperacao"] = max(0, ESTATISTICAS["em_recuperacao"] - 1)
                    
                    print("=" * 70)
                    print(f"❌ RECUPERAÇÃO EXPIRADA: #{pedido_id}")
                    print(f"👤 {pedido['nome']} | 💰 {pedido['simbolo_moeda']}{pedido['total']:.2f} PERDIDO DEFINITIVAMENTE")
                    print("=" * 70)
                    
                    # 🔥 MENSAGEM FINAL TRADUZIDA
                    mensagem_final = (
                        f"{textos_cliente['expirado_titulo']}\n\n"
                        f"{textos_cliente['expirado_mensagem'].format(pedido_id=pedido_id)}\n\n"
                        f"{textos_cliente['interessado']}"
                    )
                    
                    await context.bot.edit_message_text(
                        chat_id=chat_id,
                        message_id=message_id,
                        text=mensagem_final,
                        parse_mode="Markdown",
                        reply_markup=InlineKeyboardMarkup([
                            [InlineKeyboardButton(textos_cliente['novo_pedido'], callback_data=textos_cliente['callback'])]
                        ])
                    )
                    
            except asyncio.CancelledError:
                print(f"✅✅✅ Temporizador recuperação CANCELADO - Pedido #{pedido_id} PAGO")
            except Exception as e:
                print(f"❌❌❌ Erro na task do temporizador de recuperação: {e}")
        
        # 🔥 INICIAR TASK E GUARDAR REFERÊNCIA
        task = asyncio.create_task(temporizador_recuperacao_task())
        PEDIDOS_REGISTO[pedido_id]["timer_recuperacao"] = task
        print(f"✅✅✅ Task temporizador recuperação criada para #{pedido_id}")
        
    except Exception as e:
        print(f"❌❌❌ Erro ao iniciar temporizador de recuperação: {e}")





        

        

# ======================= SISTEMA DE PROBLEMAS =======================
async def reportar_problema(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para quando o usuário clica em 'Não, tive um problema' - BOTÃO 'NÃO'"""
    query = update.callback_query
    await query.answer()
    
    print(f"🔴 REPORTAR_PROBLEMA chamado: {query.data}")
    
    # Extrair pedido_id - formato: "reportar_problema_ABC123"
    pedido_id = query.data.replace("reportar_problema_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        # 🔥 MENSAGEM DE ERRO POR IDIOMA
        idioma = pedido.get('idioma', context.user_data.get('idioma', 'portugues'))
        textos_erro = {
            'portugues': "❌ Pedido não encontrado.",
            'ingles': "❌ Order not found.",
            'espanhol': "❌ Pedido no encontrado.",
            'italiano': "❌ Ordine non trovato.",
            'alemao': "❌ Bestellung nicht gefunden.",
            'frances': "❌ Commande non trouvée."
        }
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    # 🔥 OBTER IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', context.user_data.get('idioma', 'portugues'))
    print(f"🔴 MOSTRANDO OPÇÕES DE PROBLEMA - #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 TEXTOS POR IDIOMA
    textos = {
        'portugues': {
            'titulo': "❌ *Identificámos um problema*\n\nPara podermos ajudar melhor, qual foi o problema?\n\n*Escolha uma opção:* 👇",
            'valor_alto': "💰 Valor muito alto",
            'outro_problema': "🔧 Outro problema"
        },
        'ingles': {
            'titulo': "❌ *We identified a problem*\n\nTo help you better, what was the problem?\n\n*Choose an option:* 👇",
            'valor_alto': "💰 Price too high",
            'outro_problema': "🔧 Other problem"
        },
        'espanhol': {
            'titulo': "❌ *Identificamos un problema*\n\nPara poder ayudarle mejor, ¿cuál fue el problema?\n\n*Elija una opción:* 👇",
            'valor_alto': "💰 Precio muy alto",
            'outro_problema': "🔧 Otro problema"
        },
        'italiano': {
            'titulo': "❌ *Abbiamo identificato un problema*\n\nPer aiutarvi meglio, qual è stato il problema?\n\n*Scegli un'opzione:* 👇",
            'valor_alto': "💰 Prezzo troppo alto",
            'outro_problema': "🔧 Altro problema"
        },
        'alemao': {
            'titulo': "❌ *Wir haben ein Problem erkannt*\n\nUm Ihnen besser helfen zu können, was war das Problem?\n\n*Wählen Sie eine Option:* 👇",
            'valor_alto': "💰 Preis zu hoch",
            'outro_problema': "🔧 Anderes Problem"
        },
        'frances': {
            'titulo': "❌ *Nous avons identifié un problème*\n\nPour mieux vous aider, quel était le problème ?\n\n*Choisissez une option :* 👇",
            'valor_alto': "💰 Prix trop élevé",
            'outro_problema': "🔧 Autre problème"
        }
    }
    
    textos_idioma = textos.get(idioma, textos['portugues'])
    
    # MOSTRAR OPÇÕES DE PROBLEMA
    botoes = [
        [InlineKeyboardButton(textos_idioma['valor_alto'], callback_data=f"problema_valor_{pedido_id}")],
        [InlineKeyboardButton(textos_idioma['outro_problema'], callback_data=f"problema_outro_{pedido_id}")]
    ]
    
    await query.edit_message_text(
        text=textos_idioma['titulo'],
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(botoes)
    )



# ======================= SISTEMA DE OFERTAS =======================





async def problema_valor(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para problema de valor - PRIMEIRA OFERTA: Justificativa + Klarna"""
    query = update.callback_query
    await query.answer()
    
    print(f"🔴 PROBLEMA_VALOR chamado: {query.data}")
    
    # Extrair pedido_id
    pedido_id = query.data.replace("problema_valor_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.edit_message_text("❌ Pedido não encontrado.")
        return
    
    # 🔥 OBTER IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    print(f"💰 PRIMEIRA OFERTA: #{pedido_id} | Idioma: {idioma}")
    
    # 🔥 VERIFICAR SE É GIFT/OFERTA_SURPRESA
    tipo_cartoon = pedido.get('tipo_cartoon', '').lower()
    is_gift = (
        pedido.get('oferta_tipo') == 'oferta_surpresa' or 
        'porta-chaves' in tipo_cartoon or 
        'portachaves' in tipo_cartoon
    )
    print(f"🎁 É GIFT? {is_gift}")
    
    # 🔥 OBTER MOEDA E SÍMBOLO CORRETOS DO PEDIDO
    moeda = pedido.get('moeda', 'EUR')
    simbolo_moeda = pedido.get('simbolo_moeda', '€')
    total_pedido = pedido.get('total_pago_original', pedido.get('total', 0))
    
    # Calcular valor em 3x sem juros
    valor_3x = total_pedido / 3
    
    # 🔥 VERIFICAR SE O PAÍS SUPORTA KLARNA
    pais_cliente = pedido.get('pais', '').lower()
    paises_sem_klarna = ["canada", "brasil", "estados unidos", "united states", "usa", "us"]
    
    # 🔥 CORREÇÃO: LIMPAR EMOJIS E ESPAÇOS PARA VERIFICAR TIPO
    tipo_cartoon_limpo = tipo_cartoon
    estilo_cartoon = pedido.get('estilo_cartoon', '').lower()
    
    # Remover emojis e caracteres especiais
    import re
    tipo_limpo = re.sub(r'[^\w\s]', '', tipo_cartoon_limpo).strip()
    estilo_limpo = re.sub(r'[^\w\s]', '', estilo_cartoon).strip()
    
    print(f"🔍 DEBUG problema_valor - Tipo original: '{tipo_cartoon_limpo}'")
    print(f"🔍 DEBUG problema_valor - Tipo limpo: '{tipo_limpo}'")
    print(f"🔍 DEBUG problema_valor - Estilo limpo: '{estilo_limpo}'")
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: MAPPING DE TIPOS EM DIFERENTES IDIOMAS
    tipos_proibidos_mapping = {
        'portugues': ['animal', 'personalizado'],
        'ingles': ['animal', 'custom', 'personalized'],
        'espanhol': ['animal', 'personalizado'],
        'italiano': ['animale', 'personalizzato'],
        'alemao': ['tier', 'personalisiert'],
        'frances': ['animal', 'personnalisé']
    }
    
    # 🔥 OBTER TIPOS PROIBIDOS PARA O IDIOMA ATUAL
    tipos_proibidos = tipos_proibidos_mapping.get(idioma, tipos_proibidos_mapping['portugues'])
    print(f"🔍 TIPOS PROIBIDOS para {idioma}: {tipos_proibidos}")
    
    # 🔥 VERIFICAR SE É BUST (é sempre "Bust" em todos os idiomas)
    eh_bust = "bust" in estilo_limpo
    
    # 🔥 VERIFICAR SE PODE OFERECER TAMANHO 4.5cm
    pode_ofertar_tamanho_45 = True
    
    # Verificar se o tipo limpo contém algum dos tipos proibidos
    for tipo_proibido in tipos_proibidos:
        if tipo_proibido in tipo_limpo.lower():
            pode_ofertar_tamanho_45 = False
            print(f"🚫 TIPO PROIBIDO DETETADO: '{tipo_proibido}' em '{tipo_limpo}'")
            break
    
    # Se for bust, também não oferece 4.5cm
    if eh_bust:
        pode_ofertar_tamanho_45 = False
        print(f"🚫 ESTILO BUST DETETADO: NÃO oferece 4.5cm")
    
    print(f"🎯 RESULTADO problema_valor: Oferecer 4.5cm? {pode_ofertar_tamanho_45}")
    
    # 🔥 🔥 🔥 DETECTAR SE É ANIMAL OU PERSONALIZADO PARA TAMANHO DO PORTA-CHAVES
    tipo_lower = tipo_limpo.lower()
    tamanho_portachaves = "1.5\" | 3.8cm"  # Default para animal
    
    # Palavras-chave para ANIMAL em todos os idiomas
    palavras_animal = [
        'animal', 'pet', 'bicho', 'animais',  # Português
        'animal', 'pet', 'creature', 'animals',  # Inglês
        'animal', 'mascota', 'animales',  # Espanhol
        'animale', 'animale domestico', 'animali',  # Italiano
        'tier', 'haustier', 'animal', 'tiere',  # Alemão
        'animal', 'animal de compagnie', 'animaux'  # Francês
    ]
    
    # Palavras-chave para PERSONALIZADO em todos os idiomas
    palavras_personalizado = [
        'personalizado', 'custom', 'especial',  # Português
        'custom', 'personalized', 'bespoke', 'special',  # Inglês
        'personalizado', 'customizado', 'especial',  # Espanhol
        'personalizzato', 'su misura', 'speciale',  # Italiano
        'personalisiert', 'individuell', 'benutzerdefiniert', 'maßgeschneidert', 'personal',  # Alemão
        'personnalisé', 'customisé', 'spécial'  # Francês
    ]
    
    # Verificar qual tipo específico foi detectado
    eh_animal = any(palavra in tipo_lower for palavra in palavras_animal)
    eh_personalizado = any(palavra in tipo_lower for palavra in palavras_personalizado)
    
    print(f"🔍 DETECÇÃO PARA TAMANHO PORTA-CHAVES:")
    print(f"   • É animal? {eh_animal}")
    print(f"   • É personalizado? {eh_personalizado}")
    
    if eh_animal:
    # 🔥 ANIMAL → SEMPRE 1.5" (não importa o tamanho pedido)
       tamanho_portachaves = "1.5\" | 3.8cm"
       print(f"🐾 ANIMAL DETECTADO → TAMANHO FIXO: {tamanho_portachaves}")
       print(f"   (Ignorando tamanho pedido: {pedido.get('tamanho_cartoon', 'N/A')})")
    
    elif eh_personalizado or eh_bust:
    # 🔥 PERSONALIZADO ou BUST → SEMPRE 2.5"
        tamanho_portachaves = "2.5\" | 6.4cm"
        print(f"🎨 PERSONALIZADO/BUST DETECTADO → TAMANHO FIXO: {tamanho_portachaves}")
    
    else:
    # Outros tipos → 2.5" como padrão
        tamanho_portachaves = "2.5\" | 6.4cm"
        print(f"📏 OUTRO TIPO → TAMANHO PADRÃO: {tamanho_portachaves}")

# 🔥 MARCAR NO PEDIDO QUAL TIPO FOI DETECTADO
    pedido['tipo_detectado_portachaves'] = 'animal' if eh_animal else ('personalizado' if eh_personalizado else ('bust' if eh_bust else 'outro'))
    pedido['tamanho_portachaves_fixo'] = tamanho_portachaves
    
    # 🔥 TEXTOS POR IDIOMA
    textos = {
        'portugues': {
            'titulo': "💎 *ENTENDEMOS SUA PREOCUPAÇÃO - E TEMOS UMA SURPRESA!*\n\n",
            'justificativa': "*Porquê o valor de {simbolo}{total:.2f}:*\n",
            'personalizado': "✨ *100% Personalizado* - Desde o estilo até a embalagem\n",
            'arte': "🎨 *Arte Exclusiva* - Pintura à mão pelos nossos artistas\n",
            'horas': "⏰ *+40 Horas de Trabalho* - Em cada peça única\n",
            'frete': "📦 *Frete Premium* - Embalagem especial e rastreio\n",
            'qualidade': "🏆 *Qualidade Premium* - Desde o início ao fim\n\n",
            'mais': "🎁 *E AINDA TEMOS MAIS PARA SI:*\n",
            'desconto': "• *Desconto de 30%* já aplicado no valor final\n",
            'klarna': "• *Klarna Disponível* - Pague em 3x de {simbolo}{valor_3x:.2f} SEM JUROS\n",
            'memoria': "🌟 *Mais que um cartoon, é uma memória!*\n\n",
            'final': "*Quer esta obra de arte de qualidade exclusiva?* 👇",
            'botao_sim': "✅ Sim, Quero!",
            'botao_nao': "❌ Não, Recusar Oferta",
            'botao_nao_direto': "❌ Não, Recusar Oferta"
        },
        'ingles': {
            'titulo': "💎 *WE UNDERSTAND YOUR CONCERN - AND WE HAVE A SURPRISE!*\n\n",
            'justificativa': "*Why the price of {simbolo}{total:.2f}:*\n",
            'personalizado': "✨ *100% Customized* - From style to packaging\n",
            'arte': "🎨 *Exclusive Art* - Hand-painted by our artists\n",
            'horas': "⏰ *+40 Hours of Work* - In each unique piece\n",
            'frete': "📦 *Premium Shipping* - Special packaging and tracking\n",
            'qualidade': "🏆 *Premium Quality* - From start to finish\n\n",
            'mais': "🎁 *AND WE HAVE MORE FOR YOU:*\n",
            'desconto': "• *30% Discount* already applied to the final price\n",
            'klarna': "• *Klarna Available* - Pay in 3 installments of {simbolo}{valor_3x:.2f} NO INTEREST\n",
            'memoria': "🌟 *More than a cartoon, it's a memory!*\n\n",
            'final': "*Do you want this exclusive quality work of art?* 👇",
            'botao_sim': "✅ Yes, I Want It!",
            'botao_nao': "❌ No, Reject Offer",
            'botao_nao_direto': "❌ No, Reject Offer"
        },
        'espanhol': {
            'titulo': "💎 *ENTENDEMOS SU PREOCUPACIÓN - ¡Y TENEMOS UNA SORPRESA!*\n\n",
            'justificativa': "*Por qué el precio de {simbolo}{total:.2f}:*\n",
            'personalizado': "✨ *100% Personalizado* - Desde el estilo hasta el embalaje\n",
            'arte': "🎨 *Arte Exclusiva* - Pintado a mano por nuestros artistas\n",
            'horas': "⏰ *+40 Horas de Trabajo* - En cada pieza única\n",
            'frete': "📦 *Envío Premium* - Embalaje especial y seguimiento\n",
            'qualidade': "🏆 *Calidad Premium* - Desde el principio hasta el final\n\n",
            'mais': "🎁 *¡Y AÚN TENEMOS MÁS PARA USTED!:*\n",
            'desconto': "• *Descuento del 30%* ya aplicado al precio final\n",
            'klarna': "• *Klarna Disponible* - Pague en 3 cuotas de {simbolo}{valor_3x:.2f} SIN INTERESES\n",
            'memoria': "🌟 *¡Más que una caricatura, es un recuerdo!*\n\n",
            'final': "*¿Quiere esta obra de arte de calidad exclusiva?* 👇",
            'botao_sim': "✅ Sí, ¡Lo Quiero!",
            'botao_nao': "❌ No, Rechazar Oferta",
            'botao_nao_direto': "❌ No, Rechazar Oferta"
        },
        'italiano': {
            'titulo': "💎 *COMPRENDIAMO LA TUA PREOCCUPAZIONE - E ABBIAMO UNA SORPRESA!*\n\n",
            'justificativa': "*Perché il prezzo di {simbolo}{total:.2f}:*\n",
            'personalizado': "✨ *100% Personalizzato* - Dallo stile all'imballaggio\n",
            'arte': "🎨 *Arte Esclusiva* - Dipinto a mano dai nostri artisti\n",
            'horas': "⏰ *+40 Ore di Lavoro* - In ogni pezzo unico\n",
            'frete': "📦 *Spedizione Premium* - Imballaggio speciale e tracciamento\n",
            'qualidade': "🏆 *Qualità Premium* - Dall'inizio alla fine\n\n",
            'mais': "🎁 *E ABBIAMO ANCORA DI PIÙ PER TE:*\n",
            'desconto': "• *Sconto del 30%* già applicato al prezzo finale\n",
            'klarna': "• *Klarna Disponibile* - Paga in 3 rate da {simbolo}{valor_3x:.2f} SENZA INTERESSI\n",
            'memoria': "🌟 *Più di un cartoon, è un ricordo!*\n\n",
            'final': "*Vuoi quest'opera d'arte de qualidade esclusiva?* 👇",
            'botao_sim': "✅ Sì, Lo Voglio!",
            'botao_nao': "❌ No, Rifiuta Offerta",
            'botao_nao_direto': "❌ No, Rifiuta Offerta"
        },
        'alemao': {
            'titulo': "💎 *WIR VERSTEHEN IHRE BEDENKEN - UND WIR HABEN EINE ÜBERRASCHUNG!*\n\n",
            'justificativa': "*Warum der Preis von {simbolo}{total:.2f}:*\n",
            'personalizado': "✨ *100% Personalisiert* - Vom Stil bis zur Verpackung\n",
            'arte': "🎨 *Exklusive Kunst* - Handgemalt von unseren Künstlern\n",
            'horas': "⏰ *+40 Arbeitsstunden* - In jedem einzigartigen Stück\n",
            'frete': "📦 *Premium-Versand* - Spezielle Verpackung und Sendungsverfolgung\n",
            'qualidade': "🏆 *Premium-Qualität* - Vom Anfang bis zum Ende\n\n",
            'mais': "🎁 *UND WIR HABEN NOCH MEHR FÜR SIE:*\n",
            'desconto': "• *30% Rabatt* bereits auf den Endpreis angewendet\n",
            'klarna': "• *Klarna Verfügbar* - Bezahlen Sie in 3 Raten von {simbolo}{valor_3x:.2f} OHNE ZINSEN\n",
            'memoria': "🌟 *Mehr als ein Cartoon, es ist eine Erinnerung!*\n\n",
            'final': "*Möchten Sie dieses Kunstwerk exklusiver Qualität?* 👇",
            'botao_sim': "✅ Ja, Ich Will Es!",
            'botao_nao': "❌ Nein, Angebot Ablehnen",
            'botao_nao_direto': "❌ Nein, Angebot Ablehnen"
        },
        'frances': {
            'titulo': "💎 *NOUS COMPRENONS VOTRE INQUIÉTUDE - ET NOUS AVONS UNE SURPRISE !*\n\n",
            'justificativa': "*Pourquoi le prix de {simbolo}{total:.2f} :*\n",
            'personalizado': "✨ *100% Personnalisé* - Du style à l'emballage\n",
            'arte': "🎨 *Art Exclusif* - Peint à la main par nos artistas\n",
            'horas': "⏰ *+40 Heures de Travail* - Dans chaque pièce unique\n",
            'frete': "📦 *Livraison Premium* - Emballage spécial et suivi\n",
            'qualidade': "🏆 *Qualité Premium* - Du début à la fin\n\n",
            'mais': "🎁 *ET NOUS AVONS ENCORE PLUS POUR VOUS :*\n",
            'desconto': "• *Réduction de 30%* já aplicée au prix final\n",
            'klarna': "• *Klarna Disponible* - Payez en 3 fois de {simbolo}{valor_3x:.2f} SANS INTÉRÊTS\n",
            'memoria': "🌟 *Plus qu'un cartoon, c'est un souvenir !*\n\n",
            'final': "*Voulez-vous cette œuvre d'art de qualité exclusive ?* 👇",
            'botao_sim': "✅ Oui, Je Le Veux !",
            'botao_nao': "❌ Non, Refuser l'Offre",
            'botao_nao_direto': "❌ Non, Refuser l'Offre"
        }
    }
    
    textos_idioma = textos.get(idioma, textos['portugues'])
    
    # 🔥 CONSTRUIR TEXTO
    texto = textos_idioma['titulo']
    texto += textos_idioma['justificativa'].format(simbolo=simbolo_moeda, total=total_pedido)
    texto += textos_idioma['personalizado']
    texto += textos_idioma['arte']
    texto += textos_idioma['horas']
    texto += textos_idioma['frete']
    texto += textos_idioma['qualidade']
    texto += textos_idioma['mais']
    texto += textos_idioma['desconto']
    
    # 🔥 ADICIONAR KLARNA APENAS PARA PAÍSES QUE SUPORTAM
    tem_klarna = True
    for pais_sem_klarna in paises_sem_klarna:
        if pais_sem_klarna in pais_cliente:
            tem_klarna = False
            print(f"🚫 Klarna NÃO disponível para {pais_cliente}")
            break
    
    if tem_klarna:
        texto += textos_idioma['klarna'].format(simbolo=simbolo_moeda, valor_3x=valor_3x)
        print(f"✅ Klarna disponível para {pais_cliente}")
    
    texto += textos_idioma['memoria']
    texto += textos_idioma['final']
    
    print(f"💰 DEBUG problema_valor - Moeda: {moeda} {simbolo_moeda}")
    print(f"🌍 DEBUG problema_valor - País: {pais_cliente}")
    print(f"🌐 DEBUG problema_valor - Idioma: {idioma}")
    print(f"🎯 DEBUG problema_valor - Tipo limpo: {tipo_limpo}")
    print(f"🎯 DEBUG problema_valor - Estilo limpo: {estilo_limpo}")
    print(f"🎯 DEBUG problema_valor - É bust: {eh_bust}")
    print(f"🎯 DEBUG problema_valor - Pode oferecer 4.5cm: {pode_ofertar_tamanho_45}")
    print(f"📏 TAMANHO PORTA-CHAVES DEFINIDO: {tamanho_portachaves}")
    
    # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: BOTÕES DINÂMICOS BASEADOS NO TIPO (GIFT ou NORMAL)
    if is_gift:
        # 🔥 SE É GIFT → BOTÕES ESPECIAIS PARA GIFT COM callbacks CORRETOS
        botoes = [
            [InlineKeyboardButton(textos_idioma['botao_sim'], callback_data=f"pagar_gift_{pedido_id}")],
            [InlineKeyboardButton(textos_idioma['botao_nao'], callback_data=f"recusar_gift_{pedido_id}")]
        ]
        
        # Marcar como gift para referência futura
        pedido['eh_gift'] = True
        pedido['vai_direto_portachaves'] = True
        pedido['tamanho_portachaves'] = "2.5\" | 6.4cm" 
        
        print(f"🎁 BOTÕES ESPECIAIS PARA GIFT ATIVADOS (2 botões)")
        print(f"🎁 CALLBACKS: pagar_gift_{pedido_id} e recusar_gift_{pedido_id}")
        
    elif pode_ofertar_tamanho_45:
        # FLUXO NORMAL: Primeira oferta → Tamanho 4.5cm → Porta-chaves
        botoes = [
            [InlineKeyboardButton(textos_idioma['botao_sim'], callback_data=f"pagar_original_{pedido_id}")],
            [InlineKeyboardButton(textos_idioma['botao_nao'], callback_data=f"sair_oferta_{pedido_id}")]
        ]
        print(f"✅ OFERTANDO TAMANHO 4.5cm")
        pedido['vai_direto_portachaves'] = False
        pedido['tamanho_portachaves'] = "2.5\" | 6.4cm" 

    else:
        # FLUXO DIRETO: Primeira oferta → Porta-chaves (pula tamanho 4.5cm)
        botoes = [
            [InlineKeyboardButton(textos_idioma['botao_sim'], callback_data=f"pagar_original_{pedido_id}")],
            [InlineKeyboardButton(textos_idioma['botao_nao_direto'], callback_data=f"sair_diretoportachaves_{pedido_id}")]
        ]
        print(f"🚫 PULANDO DIRETO PARA PORTA-CHAVES")
        pedido['vai_direto_portachaves'] = True
        pedido['tamanho_portachaves'] = tamanho_portachaves 
        pedido['eh_bust_animal_personalizado'] = True  
        print(f"📏 TAMANHO PORTA-CHAVES ATRIBUÍDO AO PEDIDO: {tamanho_portachaves}")
    
    await query.edit_message_text(
        text=texto,
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(botoes)
    )





async def pagar_gift_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando clica em SIM para gift - COM MESMA ESTRUTURA DO pagar_original"""
    query = update.callback_query
    await query.answer()
    
    print(f"🎯 PAGAR_GIFT CHAMADO - MESMA ESTRUTURA DO PAGAR_ORIGINAL")
    
    # Extrair pedido_id do callback_data
    pedido_id = query.data.replace("pagar_gift_", "")
    print(f"🔍 Procurando pedido GIFT: {pedido_id}")
    
    if pedido_id not in PEDIDOS_REGISTO:
        print(f"❌ Pedido não encontrado no registro: {pedido_id}")
        await query.edit_message_text("❌ Pedido não encontrado. Por favor, inicie um novo pedido.")
        return
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    chat_id = query.message.chat_id
    
    # 🔥 OBTER IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    print(f"🌐 Idioma do pedido GIFT: {idioma}")
    
    # 🔥 VERIFICAR SE É REALMENTE UM GIFT
    if not (pedido.get('oferta_tipo') == 'oferta_surpresa' or 'porta-chaves' in pedido.get('tipo_cartoon', '').lower()):
        print(f"⚠️ ATENÇÃO: Pedido #{pedido_id} não é um GIFT, mas foi chamado como pagar_gift")
    
    # 🔥 CANCELAR QUALQUER TEMPORIZADOR ATIVO
    await cancelar_temporizadores_pedido(pedido_id)
    
    print(f"✅ Pedido GIFT encontrado: #{pedido_id}")
    print(f"🎁 Tipo de oferta: {pedido.get('oferta_tipo', 'oferta_surpresa')}")
    print(f"🔍 Chat ID do cliente: {chat_id}")

    try:
        # 🔥 PASSO 1: DEFINIR MÉTODOS DE PAGAMENTO POR PAÍS - MESMA ESTRUTURA DO pagar_stripe
        def get_payment_methods(pais):
            """Retorna métodos de pagamento baseado no país"""
            
            def get_country_code(pais_nome):
                mapeamento_paises = {
                    "portugal": "PT",
                    "espanha": "ES", 
                    "franca": "FR",
                    "alemanha": "DE",
                    "belgica": "BE",
                    "reino unido": "GB",
                    "estados unidos": "US",
                    "paises baixos": "NL",
                    "brasil": "BR",
                    "irlanda": "IE",
                    "italia": "IT",
                    "luxemburgo": "LU",
                    "canada": "CA"
                }
                return mapeamento_paises.get(pais_nome.lower(), pais_nome.upper())
            
            country_code = get_country_code(pais)
            print(f"🔍 País recebido: '{pais}' → Código: '{country_code}'")
            
            # 🔥 MESMO payment_methods_by_country DO pagar_stripe
            payment_methods_by_country = {
                "PT": ["card", "paypal", "link", "klarna", "mb_way", "sepa_debit"],
                "ES": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "FR": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "DE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "BE": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "GB": ["card", "paypal", "link", "klarna"],
                "US": ["card", "paypal", "link"],
                "NL": ["card", "paypal", "link", "klarna", "ideal", "sepa_debit"],
                "BR": ["card", "link"],
                "IE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "IT": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "LU": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "CA": ["card", "paypal", "link"]
            }
            
            methods = payment_methods_by_country.get(country_code, ["card", "link"])
            print(f"💳 Métodos de pagamento para {pais} ({country_code}): {methods}")
            return methods

        # 🔥 OBTER MÉTODOS REAIS PARA ESTE PAÍS
        metodos_reais = get_payment_methods(pedido['pais'])
        
        # 🔥 CRIAR TEXTO DINÂMICO DOS MÉTODOS - COM TRADUÇÃO
        def formatar_metodos(metodos, pais, idioma='portugues'):
            """Formata os métodos de pagamento para exibição com tradução"""
            
            # 🔥 NOMES DOS MÉTODOS POR IDIOMA
            nomes_metodos = {
                'portugues': {
                    "card": "Cartão",
                    "paypal": "PayPal", 
                    "link": "Link (inclui Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'ingles': {
                    "card": "Card",
                    "paypal": "PayPal", 
                    "link": "Link (includes Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA Debit",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'espanhol': {
                    "card": "Tarjeta",
                    "paypal": "PayPal", 
                    "link": "Enlace (incluye Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'italiano': {
                    "card": "Carta",
                    "paypal": "PayPal", 
                    "link": "Collegamento (include Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Addebito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'alemao': {
                    "card": "Karte",
                    "paypal": "PayPal", 
                    "link": "Link (enthält Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA-Lastschrift",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'frances': {
                    "card": "Carte",
                    "paypal": "PayPal", 
                    "link": "Lien (comprend Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Prélèvement SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                }
            }
            
            textos = []
            nomes_idioma = nomes_metodos.get(idioma, nomes_metodos['portugues'])
            
            for metodo in metodos:
                if metodo in nomes_idioma:
                    textos.append(nomes_idioma[metodo])
                else:
                    textos.append(metodo.capitalize())
            
            return ", ".join(textos)

        # 🔥 PASSO 2: VERIFICAR CONFIGURAÇÃO DE WALLETS - MESMA LÓGICA
        def verificar_config_wallets():
            """Verifica se as wallets estão configuradas corretamente"""
            try:
                apple_domains = stripe.ApplePayDomain.list()
                print("🍎 Domínios Apple Pay configurados:")
                for domain in apple_domains.data:
                    print(f"   - {domain.domain}")
                
                # Verificar domínio atual
                seu_dominio = "unceased-bibliothecal-donette.ngrok-free.dev"
                dominios_apple = [d.domain for d in apple_domains.data]
                if seu_dominio in dominios_apple:
                    print("✅ Domínio ngrok configurado no Apple Pay!")
                    return True
                else:
                    print("⚠️ Domínio ngrok NÃO configurado no Apple Pay")
                    return False
                    
            except Exception as e:
                print(f"❌ Erro ao verificar wallets: {e}")
                return False

        wallets_configuradas = verificar_config_wallets()

        # 🔥 PASSO 3: CRIAR SESSÃO STRIPE - COM METADATA IGUAL AO PAGAR_STRIPE
        print("🔗 Criando Checkout Session para GIFT com metadata correto...")
        
        # 🔥 🔥 🔥 CRÍTICO: USAR VARIÁVEL ÚNICA PARA OFERTA_TIPO (IGUAL AO PAGAR_STRIPE)
        oferta_tipo_stripe = "oferta_surpresa"  
        
        # 🔥 TEXTOS TRADUZIDOS PARA O CHECKOUT
        textos_checkout_messages = {
            'portugues': {
                "shipping_message": "📦 Enviaremos o seu Porta-Chaves personalizado para este endereço!",
                "submit_message": "✨ Obrigado! Vamos criar um Porta-Chaves incrível para si!"
            },
            'ingles': {
                "shipping_message": "📦 We'll send your personalized Keychain to this address!",
                "submit_message": "✨ Thank you! We'll create an amazing Keychain for you!"
            },
            'espanhol': {
                "shipping_message": "📦 ¡Enviaremos tu Llavero personalizado a esta dirección!",
                "submit_message": "✨ ¡Gracias! ¡Crearemos un Llavero increíble para ti!"
            },
            'italiano': {
                "shipping_message": "📦 Spediremo il tuo Portachiavi personalizzato a questo indirizzo!",
                "submit_message": "✨ Grazie! Creeremo un Portachiavi incredibile per te!"
            },
            'alemao': {
                "shipping_message": "📦 Wir senden Ihren personalisierten Schlüsselanhänger an diese Adresse!",
                "submit_message": "✨ Danke! Wir erstellen einen fantastischen Schlüsselanhänger für Sie!"
            },
            'frances': {
                "shipping_message": "📦 Nous enverrons votre Porte-clés personnalisé à cette adresse !",
                "submit_message": "✨ Merci ! Nous créerons un Porte-clés incroyable pour vous !"
            }
        }
        
        textos_messages = textos_checkout_messages.get(idioma, textos_checkout_messages['portugues'])
        
        session_config = {
            "payment_method_types": metodos_reais,
            "mode": "payment",
            "customer_email": pedido["email"],
            
            # 🔥 CONFIGURAÇÃO PARA WALLETS
            "payment_method_options": {
                "card": {
                    "request_three_d_secure": "automatic"
                }
            },
            
            "shipping_address_collection": {
                "allowed_countries": [
                    "PT", "ES", "FR", "DE", "BE", "GB", "US", "NL", "BR", "IE", "IT", "LU", "CA"
                ]
            },
            
            # 🔥 MENSAGENS TRADUZIDAS PARA O CHECKOUT
            "custom_text": {
                "shipping_address": {
                    "message": textos_messages["shipping_message"]
                },
                "submit": {
                    "message": textos_messages["submit_message"]
                }
            },
            
            "line_items": [{
                "price_data": {
                    "currency": pedido["moeda"].lower(),
                    "product_data": {
                        "name": f"Porta-Chaves Personalizado",
                        "description": f"Pedido #{pedido_id} - Presente para {pedido['nome']}",
                    },
                    "unit_amount": int(pedido["total"] * 100),
                },
                "quantity": 1
            }],
            
            # 🔥 URLs CORRETAS
            "success_url": f"https://t.me/plan3d_bot?start=payment_success_{pedido_id}",
            "cancel_url": f"https://t.me/plan3d_bot?start=payment_cancelled_{pedido_id}",
            
            "metadata": {
                # 🔥 🔥 🔥 METADATA IDÊNTICO AO PAGAR_STRIPE
                "chat_id": str(chat_id),
                "moeda": pedido.get('moeda', 'eur'),
                "nome_cliente": pedido['nome'],
                "oferta_tipo": oferta_tipo_stripe, 
                "pais": pedido['pais'],
                "pedido_id": pedido_id,
                "tipo_cartoon": pedido.get('tipo_cartoon', 'Porta-Chaves 🎁'),
                "tipo_sessao": oferta_tipo_stripe, 
                "total_pago": str(pedido.get('total', 0)),
                "produto_tipo": "portachaves", 
                
                "wallets_habilitadas": str(wallets_configuradas),
            },
            
            "expires_at": int((datetime.now() + timedelta(minutes=30)).timestamp()),
        }

        # 🔥 CONFIGURAÇÃO ESPECÍFICA PARA WALLETS - MESMA LÓGICA
        paises_com_wallets = ["Reino Unido", "Estados Unidos", "Brasil", "Irlanda", 
                            "França", "Alemanha", "Itália", "Espanha", "Portugal", 
                            "Países Baixos", "Bélgica", "Luxemburgo", "Canadá"]
        
        if pedido['pais'] in paises_com_wallets and "link" in metodos_reais:
            print(f"📱 Configurando Apple Pay/Google Pay para {pedido['pais']}")
            session_config["payment_method_options"]["link"] = {"persistent_token": None}

        # 🔥 CRIAR A SESSÃO
        session = stripe.checkout.Session.create(**session_config)

        print(f"✅ CHECKOUT SESSION CRIADA: {session.id}")
        print(f"🔗 URL do Checkout: {session.url}")

        # 🔥 PASSO 4: ATUALIZAR PEDIDO COM INFO DE GIFT E METADATA
        pedido["session_id_gift"] = session.id
        pedido["payment_intent_id"] = session.payment_intent
        pedido["wallets_configuradas"] = wallets_configuradas
        pedido["data_pagamento_gift"] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        # 🔥 🔥 🔥 ATUALIZAR METADATA NO PEDIDO TAMBÉM
        pedido["tipo_sessao"] = oferta_tipo_stripe   
        pedido["oferta_tipo"] = oferta_tipo_stripe  
        pedido["tipo_sessao_webhook"] = oferta_tipo_stripe  
        
        print(f"📊 Pedido GIFT atualizado com METADATA:")
        print(f"   • tipo_sessao: {pedido.get('tipo_sessao')}")
        print(f"   • oferta_tipo: {pedido.get('oferta_tipo')}")
        print(f"   • produto_tipo: {pedido.get('produto_tipo', 'portachaves')}")

        # 🔥 PASSO 5: MENSAGEM FINAL COM INSTRUÇÕES CLARAS - TRADUZIDA
        texto_metodos = formatar_metodos(metodos_reais, pedido['pais'], idioma)
        
        # 🔥 TEXTOS POR IDIOMA
        textos_checkout = {
            'portugues': {
                'titulo': "🎁 *CHECKOUT PARA PORTA-CHAVES PERSONALIZADO* 🎁\n\n",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envio:*",
                'moeda': "💰 *Moeda:*",
                'total': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido: #",
                'info': "📋 *No checkout será pedido:*\n",
                'item1': "1️⃣ **Endereço de entrega completo**\n",
                'item2': "2️⃣ **Método de pagamento**\n\n",
                'metodos': "💳 *Métodos disponíveis:*",
                'seguro': "🔒 *Pagamento 100% seguro via Stripe*\n\n",
                'tempo': "⏰ *Tem 30 minutos para efetuar o pagamento*\n\n",
                'final': "Clique abaixo para pagar: 👇",
                'botao': "💳 PAGAR AGORA →"
            },
            'ingles': {
                'titulo': "🎁 *CHECKOUT FOR PERSONALIZED KEYCHAIN* 🎁\n\n",
                'cliente': "👤 *Client:*",
                'pais': "🌍 *Shipping Country:*",
                'moeda': "💰 *Currency:*",
                'total': "💳 **TOTAL TO PAY:",
                'pedido': "🆔 **Order: #",
                'info': "📋 *In checkout you will be asked for:*\n",
                'item1': "1️⃣ **Complete delivery address**\n",
                'item2': "2️⃣ **Payment method**\n\n",
                'metodos': "💳 *Available methods:*",
                'seguro': "🔒 *100% secure payment via Stripe*\n\n",
                'tempo': "⏰ *You have 30 minutes to complete payment*\n\n",
                'final': "Click below to pay: 👇",
                'botao': "💳 PAY NOW →"
            },
            'espanhol': {
                'titulo': "🎁 *CHECKOUT PARA PORTA-LLAVES PERSONALIZADO* 🎁\n\n",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envío:*",
                'moeda': "💰 *Moneda:*",
                'total': "💳 **TOTAL A PAGAR:",
                'pedido': "🆔 **Pedido: #",
                'info': "📋 *En el checkout se le pedirá:*\n",
                'item1': "1️⃣ **Dirección de entrega completa**\n",
                'item2': "2️⃣ **Método de pago**\n\n",
                'metodos': "💳 *Métodos disponibles:*",
                'seguro': "🔒 *Pago 100% seguro a través de Stripe*\n\n",
                'tempo': "⏰ *Tiene 30 minutos para efectuar el pago*\n\n",
                'final': "Haga clic abajo para pagar: 👇",
                'botao': "💳 PAGAR AHORA →"
            },
            'italiano': {
                'titulo': "🎁 *CHECKOUT PER PORTA-CHIAVI PERSONALIZZATO* 🎁\n\n",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *Paese di Spedizione:*",
                'moeda': "💰 *Valuta:*",
                'total': "💳 **TOTALE DA PAGARE:",
                'pedido': "🆔 **Ordine: #",
                'info': "📋 *Nel checkout verrà chiesto:*\n",
                'item1': "1️⃣ **Indirizzo di consegna completo**\n",
                'item2': "2️⃣ **Metodo di pagamento**\n\n",
                'metodos': "💳 *Metodi disponibili:*",
                'seguro': "🔒 *Pagamento 100% sicuro tramite Stripe*\n\n",
                'tempo': "⏰ *Hai 30 minuti per effettuare il pagamento*\n\n",
                'final': "Clicca sotto per pagare: 👇",
                'botao': "💳 PAGA ORA →"
            },
            'alemao': {
                'titulo': "🎁 *CHECKOUT FÜR PERSONALISIERTEN SCHLÜSSELANHÄNGER* 🎁\n\n",
                'cliente': "👤 *Kunde:*",
                'pais': "🌍 *Versandland:*",
                'moeda': "💰 *Währung:*",
                'total': "💳 **ZU ZAHLENDER BETRAG:",
                'pedido': "🆔 **Bestellung: #",
                'info': "📋 *Im Checkout werden Sie gefragt:*\n",
                'item1': "1️⃣ **Vollständige Lieferadresse**\n",
                'item2': "2️⃣ **Zahlungsmethode**\n\n",
                'metodos': "💳 *Verfügbare Methoden:*",
                'seguro': "🔒 *100% sichere Zahlung über Stripe*\n\n",
                'tempo': "⏰ *Sie haben 30 Minuten, um die Zahlung durchzuführen*\n\n",
                'final': "Klicken Sie unten, um zu bezahlen: 👇",
                'botao': "💳 JETZT BEZAHLEN →"
            },
            'frances': {
                'titulo': "🎁 *CHECKOUT POUR PORTA-CLÉS PERSONNALISÉ* 🎁\n\n",
                'cliente': "👤 *Client:*",
                'pais': "🌍 *Pays de Livraison:*",
                'moeda': "💰 *Devise:*",
                'total': "💳 **MONTANT TOTAL À PAYER:",
                'pedido': "🆔 **Commande: #",
                'info': "📋 *Dans le checkout, il vous sera demandé:*\n",
                'item1': "1️⃣ **Adresse de livraison complète**\n",
                'item2': "2️⃣ **Méthode de paiement**\n\n",
                'metodos': "💳 *Méthodes disponibles:*",
                'seguro': "🔒 *Paiement 100% sécurisé via Stripe*\n\n",
                'tempo': "⏰ *Vous avez 30 minutes pour effectuer le paiement*\n\n",
                'final': "Cliquez ci-dessous pour payer: 👇",
                'botao': "💳 PAYER MAINTENANT →"
            }
        }
        
        textos = textos_checkout.get(idioma, textos_checkout['portugues'])
        
        # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
        mensagem = (
            f"{textos['titulo']}"
            f"{textos['cliente']} {pedido['nome']}\n"
            f"{textos['pais']} {pedido['pais']}\n"
            f"{textos['moeda']} {pedido['moeda'].upper()} {pedido['simbolo_moeda']}\n\n"
            f"{textos['total']} {pedido['simbolo_moeda']}{pedido['total']:.2f}**\n"
            f"{textos['pedido']}{pedido_id}**\n\n"
            f"{textos['info']}"
            f"{textos['item1']}"
            f"{textos['item2']}"
            f"{textos['metodos']} {texto_metodos}\n"
            f"{textos['seguro']}"
            f"{textos['tempo']}"
            f"{textos['final']}"
        )
        
        # 🔥 BOTÃO TRADUZIDO
        botoes_traduzidos = {
            'portugues': "💳 PAGAR AGORA →",
            'ingles': "💳 PAY NOW →",
            'espanhol': "💳 PAGAR AHORA →",
            'italiano': "💳 PAGA ORA →",
            'alemao': "💳 JETZT BEZAHLEN →",
            'frances': "💳 PAYER MAINTENANT →"
        }
        
        texto_botao = botoes_traduzidos.get(idioma, "💳 PAGAR AGORA →")

        await query.edit_message_text(
            text=mensagem,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(texto_botao, url=session.url)]
            ])
        )
        
        print(f"✅ Usuário redirecionado para Checkout GIFT em {idioma} - Metadata: {oferta_tipo_stripe}")

        # 🔥 INICIAR TEMPORIZADOR (será cancelado pelo webhook quando pagamento for feito)
        await iniciar_temporizador_pagamento_original(context, pedido_id, chat_id, query.message.message_id)
        
    except Exception as e:
        print(f"❌ ERRO STRIPE NO PAGAMENTO GIFT: {str(e)}")
        print(f"🔍 Tipo do erro: {type(e)}")
        
        import traceback
        print(f"🔍 Traceback completo: {traceback.format_exc()}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Erro no processamento do pagamento. Por favor, tente novamente em alguns segundos.",
            'ingles': "❌ Payment processing error. Please try again in a few seconds.",
            'espanhol': "❌ Error en el procesamiento del pago. Por favor, intente de nuevo en unos segundos.",
            'italiano': "❌ Errore nell'elaborazione del pagamento. Per favore, riprova tra qualche secondo.",
            'alemao': "❌ Fehler bei der Zahlungsverarbeitung. Bitte versuchen Sie es in einigen Sekunden erneut.",
            'frances': "❌ Erreur de traitement du paiement. Veuillez réessayer dans quelques secondes."
        }
        
        textos_tentar = {
            'portugues': "🔄 Tentar Novamente",
            'ingles': "🔄 Try Again",
            'espanhol': "🔄 Intentar de Nuevo",
            'italiano': "🔄 Riprova",
            'alemao': "🔄 Erneut Versuchen",
            'frances': "🔄 Réessayer"
        }
        
        texto_erro = textos_erro.get(idioma, textos_erro['portugues'])
        texto_tentar = textos_tentar.get(idioma, textos_tentar['portugues'])
        
        await query.edit_message_text(
            texto_erro,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(texto_tentar, callback_data=f"pagar_gift_{pedido_id}")],
                [InlineKeyboardButton("📞 Suporte", callback_data=f"recusar_gift_{pedido_id}")]
            ])
        )


    

async def recusar_gift_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler quando recusa o GIFT - CONTA COMO 3 OFERTAS E ENVIA RELATÓRIO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("recusar_gift_", "")
    chat_id = query.message.chat_id
    
    print(f"🎁 USUÁRIO RECUSOU GIFT - Pedido #{pedido_id}")
    
    if pedido_id not in PEDIDOS_REGISTO:
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Pedido não encontrado no sistema.",
            'ingles': "❌ Order not found in the system.",
            'espanhol': "❌ Pedido no encontrado en el sistema.",
            'italiano': "❌ Ordine non trovato nel sistema.",
            'alemao': "❌ Bestellung nicht im System gefunden.",
            'frances': "❌ Commande non trouvée dans le système."
        }
        
        # Tentar obter idioma do contexto
        idioma = context.user_data.get('idioma', 'portugues')
        await query.edit_message_text(textos_erro.get(idioma, textos_erro['portugues']))
        return
    
    pedido = PEDIDOS_REGISTO[pedido_id]
    
    # 🔥 OBTER IDIOMA DO PEDIDO
    idioma = pedido.get('idioma', 'portugues')
    print(f"🌐 Idioma do pedido GIFT recusado: {idioma}")
    
    # 🔥 MARCAR QUE RECUSOU TODAS AS 3 OFERTAS (GIFT CONTA COMO 3)
    pedido["recusou_gift"] = True
    pedido["recusou_portachaves"] = True
    pedido["recusou_original"] = True  
    pedido["recusou_oferta_45"] = True  
    
    # 🔥 GARANTIR QUE É MARCADO COMO GIFT NO PEDIDO
    pedido["oferta_tipo"] = "oferta_surpresa"
    pedido["tipo_sessao"] = "oferta_surpresa"
    pedido["tipo_recusa"] = "gift" 
    
    print(f"📊 Gift recusado - Marcado como 3 ofertas recusadas: #{pedido_id}")
    
    # 🔥 ATUALIZAR ESTATÍSTICAS
    ESTATISTICAS['ofertas_recusadas'] = ESTATISTICAS.get('ofertas_recusadas', 0) + 1
    print(f"📈 Estatística atualizada: Ofertas recusadas = {ESTATISTICAS['ofertas_recusadas']}")
    
    # 🔥 GARANTIR QUE O CHAT_ID ESTÁ NO PEDIDO
    if 'chat_id' not in pedido:
        pedido['chat_id'] = chat_id
        print(f"💾 Chat ID guardado no pedido: {chat_id}")
    
    # 🔥 ENVIAR RELATÓRIO PARA SUPORTE (A FUNÇÃO JÁ DETECTA QUE É GIFT)
    print(f"📨 Enviando relatório de recusa de GIFT: #{pedido_id}")
    await enviar_relatorio_gift_suporte(pedido_id, pedido, context)
    
    # 🔥 MOVER PARA PEDIDOS_RECUSADOS
    PEDIDOS_RECUSADOS[pedido_id] = {
        **pedido,
        "data_recusa": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "tipo_recusa": "gift"
    }
    
    # 🔥 REMOVER DO REGISTRO ATIVO
    del PEDIDOS_REGISTO[pedido_id]
    print(f"📦 Gift #{pedido_id} movido para PEDIDOS_RECUSADOS")
    
    # 🔥 TEXTOS POR IDIOMA PARA MENSAGEM FINAL
    textos_mensagem = {
        'portugues': {
            'titulo': "🎁 *Lamentamos que não queira o nosso porta-chaves personalizado!*\n\n",
            'outras_opcoes': "🌟 *Mas temos outras opções incríveis para si:*\n",
            'opcao1': "🎨 `/start` - Crie um cartoon personalizado único\n",
            'opcao2': "🎭 `/gift` - Outro porta-chaves personalizado\n",
            'final': "🎯 *Ou volte quando quiser!*\n\n",
            'tempo': "⏰ *Voltaremos ao início em 30 minutos*\n",
            'agradecimento': "*Obrigado pela sua consideração!* 👋",
            'botao_suporte': "💬 Falar com Suporte"
        },
        'ingles': {
            'titulo': "🎁 *We're sorry you don't want our personalized keychain!*\n\n",
            'outras_opcoes': "🌟 *But we have other amazing options for you:*\n",
            'opcao1': "🎨 `/start` - Create a unique personalized cartoon\n",
            'opcao2': "🎭 `/gift` - Another personalized keychain\n",
            'final': "🎯 *Or come back whenever you want!*\n\n",
            'tempo': "⏰ *We'll return to the beginning in 30 minutes*\n",
            'agradecimento': "*Thank you for your consideration!* 👋",
            'botao_suporte': "💬 Talk to Support"
        },
        'espanhol': {
            'titulo': "🎁 *¡Lamentamos que no quiera nuestro llavero personalizado!*\n\n",
            'outras_opcoes': "🌟 *¡Pero tenemos otras opciones increíbles para usted!*\n",
            'opcao1': "🎨 `/start` - Crea una caricatura personalizada única\n",
            'opcao2': "🎭 `/gift` - Otro llavero personalizado\n",
            'final': "🎯 *¡O regrese cuando quiera!*\n\n",
            'tempo': "⏰ *Volveremos al inicio en 30 minutos*\n",
            'agradecimento': "*¡Gracias por su consideración!* 👋",
            'botao_suporte': "💬 Hablar con Soporte"
        },
        'italiano': {
            'titulo': "🎁 *Ci dispiace che non desideri il nostro portachiavi personalizzato!*\n\n",
            'outras_opcoes': "🌟 *Ma abbiamo altre fantastiche opzioni per te:*\n",
            'opcao1': "🎨 `/start` - Crea un cartone animato personalizzato unico\n",
            'opcao2': "🎭 `/gift` - Un altro portachiavi personalizzato\n",
            'final': "🎯 *O torna quando vuoi!*\n\n",
            'tempo': "⏰ *Torniamo all'inizio tra 30 minuti*\n",
            'agradecimento': "*Grazie per la tua considerazione!* 👋",
            'botao_suporte': "💬 Parlare con il Supporto"
        },
        'alemao': {
            'titulo': "🎁 *Es tut uns leid, dass Sie unseren personalisierten Schlüsselanhänger nicht möchten!*\n\n",
            'outras_opcoes': "🌟 *Aber wir haben andere tolle Optionen für Sie:*\n",
            'opcao1': "🎨 `/start` - Erstellen Sie ein einzigartiges personalisiertes Cartoon\n",
            'opcao2': "🎭 `/gift` - Ein weiterer personalisierter Schlüsselanhänger\n",
            'final': "🎯 *Oder kommen Sie zurück, wann immer Sie wollen!*\n\n",
            'tempo': "⏰ *Wir kehren in 30 Minuten zum Anfang zurück*\n",
            'agradecimento': "*Vielen Dank für Ihre Rücksichtnahme!* 👋",
            'botao_suporte': "💬 Mit Support sprechen"
        },
        'frances': {
            'titulo': "🎁 *Nous sommes désolés que vous ne vouliez pas notre porte-clés personnalisé !*\n\n",
            'outras_opcoes': "🌟 *Mais nous avons d'autres options incroyables pour vous :*\n",
            'opcao1': "🎨 `/start` - Créez un dessin animé personnalisé unique\n",
            'opcao2': "🎭 `/gift` - Un autre porte-clés personnalisé\n",
            'final': "🎯 *Ou revenez quand vous voulez !*\n\n",
            'tempo': "⏰ *Nous reviendrons au début dans 30 minutes*\n",
            'agradecimento': "*Merci pour votre considération !* 👋",
            'botao_suporte': "💬 Parler au Support"
        }
    }
    
    textos = textos_mensagem.get(idioma, textos_mensagem['portugues'])
    
    # 🔥 CONSTRUIR MENSAGEM FINAL TRADUZIDA
    mensagem_final = (
        f"{textos['titulo']}"
        f"{textos['outras_opcoes']}"
        f"{textos['opcao1']}"
        f"{textos['opcao2']}"
        f"{textos['final']}"
        f"{textos['tempo']}"
        f"{textos['agradecimento']}"
    )
    
    # 🔥 ENVIAR MENSAGEM FINAL PARA O CLIENTE
    try:
        await query.edit_message_text(
            text=mensagem_final,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos['botao_suporte'], callback_data=f"todas_recusadas_{pedido_id}")]
            ])
        )
    except BadRequest:
        print("✅ Mensagem já está com o conteúdo correto - ignorando erro")
    
    # 🔥 INICIAR TEMPORIZADOR DE 30 MINUTOS
    asyncio.create_task(iniciar_temporizador_limpeza_30min(context, chat_id, query.message.message_id))















async def aceitar_oferta_especifica(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para aceitar oferta do porta-chaves"""
    query = update.callback_query
    await query.answer()
    
    print(f"🎯 OFERTA ACEITA: {query.data}")
    
    # Extrair pedido_id
    pedido_id = query.data.replace("oferta_portachaves_", "")
    pedido = PEDIDOS_REGISTO.get(pedido_id)
    
    if not pedido:
        await query.edit_message_text("❌ Pedido não encontrado.")
        return
    
    # CALCULAR NOVO PREÇO DO PORTA-CHAVES
    oferta = calcular_oferta_portachaves(pedido)
    
    # ATUALIZAR PEDIDO COM PORTA-CHAVES
    pedido["total_original"] = pedido["total"]
    pedido["total"] = oferta["total"]
    pedido["tipo_original"] = pedido["tipo_cartoon"]
    pedido["tipo_cartoon"] = "Porta-chaves"
    pedido["tamanho_original"] = pedido["tamanho_cartoon"]
    pedido["tamanho_cartoon"] = "portachaves"
    pedido["tipo_oferta"] = "portachaves"
    pedido["nome_oferta"] = oferta["nome"]
    pedido["economia"] = oferta["economia"]
    
    # ATUALIZAR ESTATÍSTICAS
    atualizar_estatistica("ofertas_aceites")
    atualizar_estatistica("em_recuperacao")
    
    print(f"🎉 PORTA-CHAVES ACEITO: #{pedido_id} | €{oferta['total']:.2f}")
    
    # CRIAR SESSÃO STRIPE PARA PORTA-CHAVES
    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="payment",
            customer_email=pedido["email"],
            line_items=[{
                "price_data": {
                    "currency": "eur",
                    "product_data": {
                        "name": f"Porta-chaves Cartoon - {pedido['estilo_cartoon']}",
                        "description": f"Oferta Especial | {pedido['nome']}"
                    },
                    "unit_amount": int(oferta["total"] * 100)
                },
                "quantity": 1
            }],
            success_url="https://teusite.com/sucesso",
            cancel_url="https://teusite.com/cancelado"
        )
        
        await query.edit_message_text(
            text=f"🎉 *EXCELENTE ESCOLHA!* 🎊\n\n"
                 f"*{oferta['nome']} Selecionado:*\n"
                 f"• Preço base: €{oferta['preco_base']:.2f}\n"
                 f"• Frete: €{oferta['frete']:.2f}\n"
                 f"• Imposto: €{oferta['imposto']:.2f}\n"
                 f"• *Total Final: €{oferta['total']:.2f}*\n"
                 f"• Economia: €{oferta['economia']:.2f} 💰\n\n"
                 f"*Pedido #{pedido_id}*\n"
                 f"Estilo: {pedido['estilo_cartoon']}\n\n"
                 "*Clique abaixo para pagar:* 👇",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("💳 Pagar Agora", url=session.url)],
                [InlineKeyboardButton("📞 Suporte", callback_data=f"todas_recusadas_{pedido_id}")]
            ])
        )
        
        # INICIAR TEMPORIZADOR PARA OFERTA
        await iniciar_temporizador_oferta(context, pedido_id, query.message.chat_id, query.message.message_id)
        
    except Exception as e:
        print(f"❌ Erro Stripe: {e}")
        await query.edit_message_text(
            text="❌ *Erro no processamento*\n\n"
                 "Desculpe, houve um erro. Contacte o suporte.",
            parse_mode="Markdown"
        )


async def iniciar_temporizador_oferta(context, pedido_id, chat_id, message_id, idioma=None):
    """Temporizador de 30 minutos para oferta especial - VERSÃO COMPLETA E CORRIGIDA"""
    try:
        print(f"⏰⏰⏰ INICIAR_TEMPORIZADOR_OFERTA para #{pedido_id} (30 minutos) | Idioma: {idioma}")
        
        # Se idioma não foi fornecido, pegar do pedido ou do user_data
        if idioma is None:
            if pedido_id in PEDIDOS_REGISTO and 'idioma' in PEDIDOS_REGISTO[pedido_id]:
                idioma = PEDIDOS_REGISTO[pedido_id]['idioma']
            else:
                idioma = context.user_data.get('idioma', 'portugues')
        
        print(f"🌐 Idioma definido para temporizador: {idioma}")
        
        async def temporizador_oferta_task():
            try:
                print(f"⏰ Task temporizador oferta iniciada para #{pedido_id}")
                await asyncio.sleep(1800)  # 30 minutos
                
                print(f"🔍 Verificando se oferta #{pedido_id} ainda está ativa...")
                
                if (pedido_id in PEDIDOS_REGISTO and 
                    PEDIDOS_REGISTO[pedido_id].get("oferta_tipo") and  
                    PEDIDOS_REGISTO[pedido_id]["status"] != "pago"):
                    
                    pedido = PEDIDOS_REGISTO[pedido_id]
                    
                    # 🔥 OBTER IDIOMA DO PEDIDO
                    idioma_pedido = pedido.get('idioma', idioma)
                    print(f"🌐 Idioma do pedido oferta expirada: {idioma_pedido}")
                    
                    # REMOVER DA RECUPERAÇÃO NAS ESTATÍSTICAS
                    ESTATISTICAS["em_recuperacao"] = max(0, ESTATISTICAS["em_recuperacao"] - 1)
                    
                    # 🔥 DETERMINAR TIPO DE OFERTA PARA MENSAGEM
                    tipo_oferta = pedido.get('oferta_tipo', '')
                    print(f"🔍 Tipo de oferta expirada: {tipo_oferta}")
                    
                    # 🔥 NOMES DAS OFERTAS POR IDIOMA
                    nomes_ofertas = {
                        'tamanho_4.5': {
                            'portugues': "Personalizado 4.5cm",
                            'ingles': "4.5cm Custom",
                            'espanhol': "Personalizado 4.5cm", 
                            'italiano': "Personalizzato 4.5cm",
                            'alemao': "Individuell 4.5cm",
                            'frances': "Personnalisé 4.5cm"
                        },
                        'oferta_surpresa': {
                            'portugues': "Porta-Chaves Surpresa",
                            'ingles': "Surprise Keychain",
                            'espanhol': "Llavero Sorpresa",
                            'italiano': "Portachiavi Sorpresa",
                            'alemao': "Überraschungs-Schlüsselanhänger",
                            'frances': "Porte-clés Surprise"
                        },
                        'original': {
                            'portugues': "Original 10cm",
                            'ingles': "Original 10cm", 
                            'espanhol': "Original 10cm",
                            'italiano': "Originale 10cm",
                            'alemao': "Original 10cm",
                            'frances': "Original 10cm"
                        },
                        'portachaves': {
                            'portugues': "Porta-Chaves",
                            'ingles': "Keychain",
                            'espanhol': "Llavero",
                            'italiano': "Portachiavi",
                            'alemao': "Schlüsselanhänger",
                            'frances': "Porte-clés"
                        }
                    }
                    
                    nome_oferta = "Oferta Especial"
                    if tipo_oferta in nomes_ofertas:
                        nome_oferta = nomes_ofertas[tipo_oferta].get(idioma_pedido, nomes_ofertas[tipo_oferta]['portugues'])
                    
                    print("=" * 70)
                    print(f"❌ OFERTA EXPIRADA: #{pedido_id}")
                    print(f"👤 {pedido.get('nome', 'N/A')} | 🎁 {nome_oferta} | 💰 {pedido.get('simbolo_moeda', '€')}{pedido.get('total', 0):.2f} PERDIDA")
                    print("=" * 70)
                    
                    # 🔥 TEXTOS POR IDIOMA PARA MENSAGEM DE EXPIRAÇÃO
                    textos_expiracao = {
                        'portugues': {
                            'titulo': "⏰ *OFERTA EXPIRADA*\n\n",
                            'texto1': f"A oferta especial do pedido `{pedido_id}` expirou.\n\n",
                            'oferta': "*Oferta selecionada:*",
                            'valor': "*Valor da oferta:*",
                            'final': "\n*Se ainda estiver interessado, inicie um novo pedido.*",
                            'botao': "🔄 Novo Pedido"
                        },
                        'ingles': {
                            'titulo': "⏰ *OFFER EXPIRED*\n\n",
                            'texto1': f"The special offer for order `{pedido_id}` has expired.\n\n",
                            'oferta': "*Selected offer:*",
                            'valor': "*Offer value:*",
                            'final': "\n*If you're still interested, start a new order.*",
                            'botao': "🔄 New Order"
                        },
                        'espanhol': {
                            'titulo': "⏰ *OFERTA EXPIRADA*\n\n",
                            'texto1': f"La oferta especial del pedido `{pedido_id}` ha expirado.\n\n",
                            'oferta': "*Oferta seleccionada:*",
                            'valor': "*Valor de la oferta:*",
                            'final': "\n*Si aún está interesado, inicie un nuevo pedido.*",
                            'botao': "🔄 Nuevo Pedido"
                        },
                        'italiano': {
                            'titulo': "⏰ *OFFERTA SCADUTA*\n\n",
                            'texto1': f"L'offerta speciale per l'ordine `{pedido_id}` è scaduta.\n\n",
                            'oferta': "*Offerta selezionata:*",
                            'valor': "*Valore dell'offerta:*",
                            'final': "\n*Se sei ancora interessato, inizia un nuovo ordine.*",
                            'botao': "🔄 Nuovo Ordine"
                        },
                        'alemao': {
                            'titulo': "⏰ *ANGEBOT ABGELAUFEN*\n\n",
                            'texto1': f"Das Sonderangebot für Bestellung `{pedido_id}` ist abgelaufen.\n\n",
                            'oferta': "*Ausgewähltes Angebot:*",
                            'valor': "*Angebotswert:*",
                            'final': "\n*Wenn Sie noch interessiert sind, starten Sie eine neue Bestellung.*",
                            'botao': "🔄 Neue Bestellung"
                        },
                        'frances': {
                            'titulo': "⏰ *OFFRE EXPIRÉE*\n\n",
                            'texto1': f"L'offre spéciale pour la commande `{pedido_id}` a expiré.\n\n",
                            'oferta': "*Offre sélectionnée:*",
                            'valor': "*Valeur de l'offre:*",
                            'final': "\n*Si vous êtes toujours intéressé, commencez une nouvelle commande.*",
                            'botao': "🔄 Nouvelle Commande"
                        }
                    }
                    
                    textos = textos_expiracao.get(idioma_pedido, textos_expiracao['portugues'])
                    
                    # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
                    simbolo_moeda = pedido.get('simbolo_moeda', '€')
                    total = pedido.get('total', 0)
                    
                    mensagem_expiracao = (
                        f"{textos['titulo']}"
                        f"{textos['texto1']}"
                        f"{textos['oferta']} {nome_oferta}\n"
                        f"{textos['valor']} {simbolo_moeda}{total:.2f}\n"
                        f"{textos['final']}"
                    )
                    
                    # 🔥 ENVIAR MENSAGEM DE EXPIRAÇÃO TRADUZIDA
                    try:
                        await context.bot.edit_message_text(
                            chat_id=chat_id,
                            message_id=message_id,
                            text=mensagem_expiracao,
                            parse_mode="Markdown",
                            reply_markup=InlineKeyboardMarkup([
                                [InlineKeyboardButton(textos['botao'], callback_data="voltar_inicio")]
                            ])
                        )
                        print(f"✅✅✅ Mensagem de expiração enviada para oferta #{pedido_id} | Idioma: {idioma_pedido}")
                    except Exception as e:
                        print(f"⚠️ Erro ao enviar mensagem de expiração: {e}")
                        # Tentar enviar mensagem simples
                        try:
                            await context.bot.send_message(
                                chat_id=chat_id,
                                text=f"⏰ A oferta #{pedido_id} expirou. Se ainda estiver interessado, inicie um novo pedido."
                            )
                        except:
                            pass
                    
                    # 🔥 ATUALIZAR STATUS DO PEDIDO
                    pedido["status"] = "expirado"
                    PEDIDOS_REGISTO[pedido_id] = pedido
                    
                    # 🔥 REMOVER TIMER
                    if "timer_oferta" in pedido:
                        del pedido["timer_oferta"]
                        
                else:
                    print(f"✅ Oferta #{pedido_id} já foi paga ou não existe mais")
                    
            except asyncio.CancelledError:
                print(f"✅✅✅ Temporizador oferta CANCELADO - Pedido #{pedido_id} PAGO")
            except Exception as e:
                print(f"❌❌❌ Erro na task do temporizador de oferta: {e}")
                import traceback
                traceback.print_exc()
        
        # 🔥 GARANTIR QUE O PEDIDO TEM OS DADOS NECESSÁRIOS
        if pedido_id in PEDIDOS_REGISTO:
            pedido = PEDIDOS_REGISTO[pedido_id]
            # Garantir que tem idioma
            if 'idioma' not in pedido:
                pedido['idioma'] = idioma
            # Garantir que tem oferta_tipo (para compatibilidade)
            if 'oferta_tipo' not in pedido and 'tipo_oferta' in pedido:
                pedido['oferta_tipo'] = pedido['tipo_oferta']
            PEDIDOS_REGISTO[pedido_id] = pedido
        
        # 🔥 INICIAR TASK E GUARDAR REFERÊNCIA
        task = asyncio.create_task(temporizador_oferta_task())
        
        # 🔥 GARANTIR QUE O PEDIDO EXISTE ANTES DE ADICIONAR TIMER
        if pedido_id in PEDIDOS_REGISTO:
            PEDIDOS_REGISTO[pedido_id]["timer_oferta"] = task
            print(f"✅✅✅ Task temporizador oferta criada para #{pedido_id}")
        else:
            print(f"❌❌❌ ERRO: Pedido #{pedido_id} não existe para adicionar timer!")
            task.cancel()  # Cancelar task se pedido não existe
        
    except Exception as e:
        print(f"❌❌❌ Erro ao iniciar temporizador de oferta: {e}")
        import traceback
        traceback.print_exc()





async def recusar_oferta(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para recusar oferta - COM ENVIO PARA SUPORTE E TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    pedido_id = query.data.replace("recusar_oferta_", "")
    chat_id = query.message.chat_id
    
    print(f"🔴 OFERTA RECUSADA: #{pedido_id} no chat {chat_id}")
    
    # 🔥 VERIFICAR SE RECUSOU TODAS AS 3 OFERTAS
    if pedido_id in PEDIDOS_REGISTO:
        pedido = PEDIDOS_REGISTO[pedido_id]
        pedido["recusou_portachaves"] = True
        
        # 🔥 OBTER IDIOMA DO PEDIDO
        idioma = pedido.get('idioma', 'portugues')
        print(f"🌐 Idioma do pedido recusado: {idioma}")
        
        recusou_original = pedido.get("recusou_original", False)
        recusou_45 = pedido.get("recusou_oferta_45", False)
        
        if recusou_original and recusou_45:
            # 🔥 DEBUG PARA VERIFICAR SE CHEGA ATÉ AQUI
            print(f"🔍 DEBUG - Chegou ao ponto de enviar relatório: #{pedido_id}")
            print(f"🔍 DEBUG - Recusou original: {recusou_original}")
            print(f"🔍 DEBUG - Recusou 4.5cm: {recusou_45}")
            print(f"🔍 DEBUG - Vai chamar enviar_relatorio_recusa_suporte")
            
            # 🔥 GARANTIR QUE O CHAT_ID ESTÁ NO PEDIDO
            if 'chat_id' not in pedido:
                pedido['chat_id'] = chat_id
                print(f"💾 Chat ID guardado no pedido: {chat_id}")
            
            # 🔥 MOVER PARA PEDIDOS_RECUSADOS E REMOVER DO PEDIDOS_REGISTO
            PEDIDOS_RECUSADOS[pedido_id] = {
                **pedido,
                "data_recusa": datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            }
            
            # 🔥 ENVIAR RELATÓRIO PARA SUPORTE ANTES DE REMOVER
            await enviar_relatorio_recusa_suporte(pedido_id, pedido, context)
            
            # 🔥 REMOVER DO REGISTRO APÓS ENVIAR RELATÓRIO
            del PEDIDOS_REGISTO[pedido_id]
            
            # 🔥 CONTAR NAS ESTATÍSTICAS
            ESTATISTICAS['ofertas_recusadas'] = ESTATISTICAS.get('ofertas_recusadas', 0) + 1
            print(f"🎯 USUÁRIO RECUSOU TODAS AS 3 OFERTAS: #{pedido_id}")
            print(f"📈 Estatística atualizada: Ofertas recusadas = {ESTATISTICAS['ofertas_recusadas']}")
        else:
            print(f"🔍 Usuário recusou porta-chaves mas ainda não recusou todas:")
            print(f"   • Recusou original: {recusou_original}")
            print(f"   • Recusou 4.5cm: {recusou_45}")
    
    # 🔥 TEXTOS POR IDIOMA PARA MENSAGEM DE RECUSA
    textos_recusa = {
        'portugues': {
            'titulo': "😔 *Entendo sua decisão*\n\n",
            'texto': "A oferta do porta-chaves não foi do seu interesse.\n\n",
            'suporte': "*Nosso suporte pode ajudar com alternativas personalizadas.*\n\n",
            'tempo': "⏰ *Voltaremos ao início em 30 minutos*\n",
            'final': "*Ou use /start a qualquer momento*",
            'botao_suporte': "💬 Falar com Suporte"
        },
        'ingles': {
            'titulo': "😔 *I understand your decision*\n\n",
            'texto': "The keychain offer was not to your liking.\n\n",
            'suporte': "*Our support can help with personalized alternatives.*\n\n",
            'tempo': "⏰ *We'll return to the beginning in 30 minutes*\n",
            'final': "*Or use /start at any time*",
            'botao_suporte': "💬 Talk to Support"
        },
        'espanhol': {
            'titulo': "😔 *Entiendo su decisión*\n\n",
            'texto': "La oferta del llavero no fue de su interés.\n\n",
            'suporte': "*Nuestro soporte puede ayudar con alternativas personalizadas.*\n\n",
            'tempo': "⏰ *Volveremos al inicio en 30 minutos*\n",
            'final': "*O use /start en cualquier momento*",
            'botao_suporte': "💬 Hablar con Soporte"
        },
        'italiano': {
            'titulo': "😔 *Capisco la tua decisione*\n\n",
            'texto': "L'offerta del portachiavi non ti ha interessato.\n\n",
            'suporte': "*Il nostro supporto può aiutarti con alternative personalizzate.*\n\n",
            'tempo': "⏰ *Torneremo all'inizio in 30 minuti*\n",
            'final': "*Oppure usa /start in qualsiasi momento*",
            'botao_suporte': "💬 Parlare con il Supporto"
        },
        'alemao': {
            'titulo': "😔 *Ich verstehe Ihre Entscheidung*\n\n",
            'texto': "Das Schlüsselanhänger-Angebot hat Sie nicht angesprochen.\n\n",
            'suporte': "*Unser Support kann mit personalisierten Alternativen helfen.*\n\n",
            'tempo': "⏰ *Wir kehren in 30 Minuten zum Anfang zurück*\n",
            'final': "*Oder verwenden Sie /start jederzeit*",
            'botao_suporte': "💬 Mit Support sprechen"
        },
        'frances': {
            'titulo': "😔 *Je comprends votre décision*\n\n",
            'texto': "L'offre de porte-clés ne vous a pas intéressé.\n\n",
            'suporte': "*Notre support peut vous aider avec des alternatives personnalisées.*\n\n",
            'tempo': "⏰ *Nous reviendrons au début dans 30 minutes*\n",
            'final': "*Ou utilisez /start à tout moment*",
            'botao_suporte': "💬 Parler au Support"
        }
    }
    
    textos = textos_recusa.get(idioma, textos_recusa['portugues'])
    
    # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
    mensagem_final = (
        f"{textos['titulo']}"
        f"{textos['texto']}"
        f"{textos['suporte']}"
        f"{textos['tempo']}"
        f"{textos['final']}"
    )
    
    try:
        await query.edit_message_text(
            text=mensagem_final,
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton(textos['botao_suporte'], callback_data=f"todas_recusadas_{pedido_id}")]
            ])
        )
    except BadRequest:
        print("✅ Mensagem já está com o conteúdo correto - ignorando erro")
    
    # INICIAR TEMPORIZADOR DE 30 MINUTOS
    asyncio.create_task(iniciar_temporizador_limpeza_30min(context, chat_id, query.message.message_id))



async def enviar_relatorio_recusa_suporte(pedido_id, pedido, context):
    """Envia relatório completo das ofertas recusadas para o suporte - MOSTRA APENAS OFERTAS APRESENTADAS"""
    try:
        CHAT_SUPORTE_PAGOS = os.getenv("CHAT_SUPORTE_PAGOS")  # Ou o canal apropriado
        
        if not CHAT_SUPORTE_PAGOS:
            print("⚠️ AVISO: CHAT_SUPORTE_PAGOS não configurado")
            return  # Sai da função sem enviar
        
        try:
            CHAT_SUPORTE_PAGOS = int(CHAT_SUPORTE_PAGOS)
        except ValueError:
            print("⚠️ AVISO: CHAT_SUPORTE_PAGOS inválido")
            return  # Sai da função sem enviar
        
        print(f"📨 ENVIANDO RELATÓRIO DE RECUSA PARA SUPORTE: #{pedido_id}")
        
        # 🔥 VERIFICAR SE FOI FLUXO DIRETO (PULOU 4.5cm)
        fluxo_direto = pedido.get("fluxo_direto", False)
        print(f"🔍 TIPO DE FLUXO: {'DIRETO (pulou 4.5cm)' if fluxo_direto else 'NORMAL (3 ofertas)'}")
        
        # 🔥 OBTER INFORMAÇÕES DE MOEDA
        moeda_original = pedido.get('moeda', 'EUR')
        simbolo_original = pedido.get('simbolo_moeda', '€')
        total_original = pedido.get('valor_original_real', pedido.get('total', 0))
        
        # 🔥 CONVERTER PARA EUR USANDO API FRANKFURTER - COM CANADÁ
        def converter_para_eur(valor, codigo_moeda_origem):
            if codigo_moeda_origem.upper() == 'EUR':
                return valor
            try:
                response = requests.get(f"https://api.frankfurter.app/latest?from={codigo_moeda_origem.upper()}&to=EUR", timeout=10)
                response.raise_for_status()
                data = response.json()
                taxa = data['rates']['EUR']
                return valor * taxa
            except:
                # 🔥 TAXAS FALLBACK COM CANADÁ ADICIONADO
                taxas_fallback = {
                    'USD': 0.85, 'GBP': 1.15, 'BRL': 0.17, 
                    'CAD': 0.68,  # 🔥 NOVO: Dólar Canadiano
                    'AUD': 0.60, 'CHF': 0.95
                }
                taxa = taxas_fallback.get(codigo_moeda_origem.upper(), 1.0)
                return valor * taxa

        def obter_codigo_moeda(simbolo_ou_codigo):
            mapeamento = {
                '$': 'USD', 'US$': 'USD', 'USD': 'USD',
                '€': 'EUR', 'EUR': 'EUR', 
                '£': 'GBP', 'GBP': 'GBP',
                'R$': 'BRL', 'BRL': 'BRL',
                'C$': 'CAD', 'CAD': 'CAD'  # 🔥 NOVO: Dólar Canadiano
            }
            if len(simbolo_ou_codigo) == 3 and simbolo_ou_codigo.isalpha():
                return simbolo_ou_codigo.upper()
            return mapeamento.get(simbolo_ou_codigo, 'EUR')

        codigo_moeda_original = obter_codigo_moeda(moeda_original)
        total_eur = converter_para_eur(total_original, codigo_moeda_original)
        
        # 🔥 OBTER VALORES DAS OFERTAS
        valor_oferta_45_real = pedido.get('valor_oferta_45_real', total_original * 0.80)
        valor_oferta_portachaves_real = pedido.get('valor_oferta_portachaves_real', total_original * 0.30)
        
        # Converter ofertas para EUR
        valor_45_eur = converter_para_eur(valor_oferta_45_real, codigo_moeda_original)
        valor_portachaves_eur = converter_para_eur(valor_oferta_portachaves_real, codigo_moeda_original)

        chat_id_cliente = pedido.get('chat_id')
        nome_cliente = pedido.get('nome', 'Cliente')
        
        print(f"💰 MOEDAS - Original: {simbolo_original}{total_original:.2f} {codigo_moeda_original} | EUR: €{total_eur:.2f}")

        # 🔥 FORMATAR VALORES COM EUR ENTRE PARÊNTESIS SE NÃO FOR EUR
        def formatar_valor(valor_original, valor_eur, simbolo_original, codigo_moeda):
            if codigo_moeda.upper() == 'EUR':
                return f"{simbolo_original}{valor_original:.2f}"
            else:
                return f"{simbolo_original}{valor_original:.2f} {codigo_moeda} (€{valor_eur:.2f})"

        valor_original_formatado = formatar_valor(total_original, total_eur, simbolo_original, codigo_moeda_original)
        valor_45_formatado = formatar_valor(valor_oferta_45_real, valor_45_eur, simbolo_original, codigo_moeda_original)
        valor_portachaves_formatado = formatar_valor(valor_oferta_portachaves_real, valor_portachaves_eur, simbolo_original, codigo_moeda_original)

        # 🔥 CONSTRUIR MENSAGEM DAS OFERTAS RECUSADAS CONFORME O FLUXO
        ofertas_recusadas = ""
        
        if fluxo_direto:
            # 🔥 FLUXO DIRETO: MOSTRA APENAS 2 OFERTAS (ORIGINAL + PORTA-CHAVES)
            ofertas_recusadas = f"""1️⃣ *ORIGINAL*: {valor_original_formatado} ❌
2️⃣ *PORTA-CHAVES (70% OFF)*: {valor_portachaves_formatado} ❌"""
            
            print("🔍 FLUXO DIRETO: Oferta 4.5cm NÃO mostrada no relatório (foi pulada)")
        else:
            # 🔥 FLUXO NORMAL: MOSTRA AS 3 OFERTAS
            ofertas_recusadas = f"""1️⃣ *ORIGINAL*: {valor_original_formatado} ❌
2️⃣ *TAMANHO 4.5 (20% OFF)*: {valor_45_formatado} ❌
3️⃣ *PORTA-CHAVES (70% OFF)*: {valor_portachaves_formatado} ❌"""
            
            print("🔍 FLUXO NORMAL: Mostrando todas as 3 ofertas")

        # 🔥 MENSAGEM COMPACTA COM VALORES EM EUR
        mensagem = f"""🚨 *RELATÓRIO DE RECUSA - #{pedido_id}*

*👤 CLIENTE*
Nome: {nome_cliente}
Email: {pedido.get('email', 'N/A')}
País: {pedido.get('pais', 'N/A')}
Contacto: {pedido.get('contacto', 'N/A')}

*🎨 PRODUTO*
{_obter_detalhes_adicionais_pedido(pedido)}

*💰 OFERTAS RECUSADAS*
{ofertas_recusadas}

🚨 *RECUSOU TODAS AS OFERTAS APRESENTADAS*"""

        # Verificar tamanho da mensagem
        print(f"📏 Tamanho da mensagem: {len(mensagem)} caracteres")

        # 🔥 BOTÃO PARA CONTACTAR DIRETAMENTE O CLIENTE
        keyboard = []
        if chat_id_cliente:
            keyboard.append([
                InlineKeyboardButton("📞 Contactar Cliente", url=f"tg://user?id={chat_id_cliente}")
            ])
        
        reply_markup = InlineKeyboardMarkup(keyboard) if keyboard else None

        # 🔥 ENVIAR APENAS 1 MENSAGEM
        if pedido.get('foto_id'):
            try:
                print("📸 Enviando 1 MENSAGEM com foto...")
                await context.bot.send_photo(
                    chat_id=CHAT_SUPORTE_PAGOS,
                    photo=pedido['foto_id'],
                    caption=mensagem,
                    parse_mode="Markdown",
                    reply_markup=reply_markup
                )
                print("✅ 1 mensagem com foto enviada com sucesso!")
                
            except Exception as e:
                print(f"❌ Erro ao enviar com foto: {e}")
                # Fallback: enviar sem foto
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_PAGOS,
                    text=mensagem,
                    parse_mode="Markdown",
                    reply_markup=reply_markup
                )
        else:
            print("🖼️ Nenhuma imagem, enviando 1 mensagem sem foto")
            await context.bot.send_message(
                chat_id=CHAT_SUPORTE_PAGOS,
                text=mensagem,
                parse_mode="Markdown",
                reply_markup=reply_markup
            )
        
        print(f"✅ Relatório de recusa enviado para suporte: #{pedido_id}")
        
    except Exception as e:
        print(f"❌ ERRO ao enviar relatório para suporte: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")




def _obter_detalhes_adicionais_pedido(pedido):
    """Retorna detalhes adicionais específicos do tipo de pedido"""
    detalhes = ""
    
    # 🔥 DEBUG: VERIFICAR CAMPOS IMPORTANTES
    print(f"🔍 DEBUG _obter_detalhes_adicionais_pedido:")
    print(f"   • tipo_animal: {pedido.get('tipo_animal', 'NÃO ENCONTRADO')}")
    print(f"   • nome_animal: {pedido.get('nome_animal', 'NÃO ENCONTRADO')}")
    print(f"   • tipo_cartoon: {pedido.get('tipo_cartoon', 'NÃO ENCONTRADO')}")
    print(f"   • estilo_cartoon: {pedido.get('estilo_cartoon', 'NÃO ENCONTRADO')}")
    print(f"   • tamanho_original: {pedido.get('tamanho_original', 'NÃO ENCONTRADO')}")
    
    # 🔥 TIPO E ESTILO DO CARTOON
    tipo_cartoon = pedido.get('tipo_cartoon', '')
    estilo_cartoon = pedido.get('estilo_cartoon', '')
    
    if tipo_cartoon:
        detalhes += f"• Tipo: {tipo_cartoon}\n"
    if estilo_cartoon:
        detalhes += f"• Estilo: {estilo_cartoon}\n"
    
    # 🔥 TAMANHO ORIGINAL
    tamanho_original = pedido.get('tamanho_original', '')
    if tamanho_original:
        detalhes += f"• Tamanho: {tamanho_original}\n"
    else:
        # Fallback para tamanho_cartoon
        tamanho_fallback = pedido.get('tamanho_cartoon', '')
        if tamanho_fallback:
            # Limpar texto de oferta
            if "(Oferta Especial)" in tamanho_fallback:
                tamanho_fallback = tamanho_fallback.replace("(Oferta Especial)", "").strip()
            if "Oferta" in tamanho_fallback:
                tamanho_fallback = tamanho_fallback.replace("Oferta", "").strip()
            detalhes += f"• Tamanho: {tamanho_fallback}\n"
    
    # 🔥 NOME E FRASE DO CARTOON
    if pedido.get('nome_cartoon'):
        detalhes += f"• Nome: {pedido.get('nome_cartoon')}\n"
    if pedido.get('frase_cartoon'):
        detalhes += f"• Frase: {pedido.get('frase_cartoon')}\n"
    
    # 🔥 DETALHES DA FAMÍLIA
    if pedido.get('nome_family'):
        detalhes += f"• Nome da Família: {pedido.get('nome_family')}\n"
    if pedido.get('frase_family') and pedido.get('frase_family') != "Não adicionou frase":
        detalhes += f"• Frase da Família: \"{pedido.get('frase_family')}\"\n"
    if pedido.get('elementos_family'):
        detalhes += f"• Elementos: {pedido.get('elementos_family')}\n"
    if pedido.get('adultos_family'):
        detalhes += f"• Adultos: {pedido.get('adultos_family')}\n"
    if pedido.get('criancas_family'):
        detalhes += f"• Crianças: {pedido.get('criancas_family')}\n"
    if pedido.get('animais_family'):
        detalhes += f"• Animais da Família: {pedido.get('animais_family')}\n"
    
    # 🔥 DETALHES PROFISSIONAIS
    if pedido.get('profissao'):
        detalhes += f"• Profissão: {pedido.get('profissao')}\n"
    if pedido.get('objetos_office'):
        detalhes += f"• Objetos: {pedido.get('objetos_office')}\n"
    if pedido.get('super_heroi'):
        detalhes += f"• Super-Herói: {pedido.get('super_heroi')}\n"
    
    # 🔥 DETALHES PERSONALIZADOS
    if pedido.get('tipo_personalizado'):
        detalhes += f"• Tipo de Peça: {pedido.get('tipo_personalizado')}\n"
    if pedido.get('nome_peca_personalizado'):
        detalhes += f"• Nome da Peça: {pedido.get('nome_peca_personalizado')}\n"
    
    # 🔥 DETALHES DE ANIMAIS (APENAS nome_animal E tipo_animal COMO SOLICITADO)
    
    # 1. TIPO DE ANIMAL
    if pedido.get('tipo_animal'):
        tipo_animal = pedido.get('tipo_animal')
        # Verificar se é string ou lista
        if isinstance(tipo_animal, list):
            if tipo_animal:  # Se não estiver vazio
                tipo_animal_str = ", ".join(tipo_animal)
                detalhes += f"• Tipo de Animal: {tipo_animal_str}\n"
        elif isinstance(tipo_animal, str) and tipo_animal.strip():
            detalhes += f"• Tipo de Animal: {tipo_animal}\n"
    
    # 2. NOME DO ANIMAL
    if pedido.get('nome_animal'):
        nome_animal = pedido.get('nome_animal')
        # Verificar se é string ou lista
        if isinstance(nome_animal, list):
            if nome_animal:  # Se não estiver vazio
                nome_animal_str = ", ".join(nome_animal)
                detalhes += f"• Nome do Animal: {nome_animal_str}\n"
        elif isinstance(nome_animal, str) and nome_animal.strip():
            detalhes += f"• Nome do Animal: {nome_animal}\n"
    
    # Se não houver detalhes
    if not detalhes:
        detalhes = "• Sem detalhes adicionais"
    
    print(f"📊 DETALHES FINAIS: {detalhes}")
    return detalhes






async def enviar_relatorio_gift_suporte(pedido_id, pedido, context):
    """Envia relatório ESPECÍFICO para recusa de GIFT (1 oferta)"""
    try:
        CHAT_SUPORTE_PAGOS = os.getenv("CHAT_SUPORTE_PAGOS")  # Ou o canal apropriado
        
        if not CHAT_SUPORTE_PAGOS:
            print("⚠️ AVISO: CHAT_SUPORTE_PAGOS não configurado")
            return  # Sai da função sem enviar
        
        try:
            CHAT_SUPORTE_PAGOS = int(CHAT_SUPORTE_PAGOS)
        except ValueError:
            print("⚠️ AVISO: CHAT_SUPORTE_PAGOS inválido")
            return  # Sai da função sem enviar
        
        print(f"🎁 ENVIANDO RELATÓRIO DE RECUSA DE GIFT: #{pedido_id}")
        
        # 🔥 OBTER DADOS DO GIFT
        nome_cliente = pedido.get('nome', 'N/A')
        email_cliente = pedido.get('email', 'N/A')
        pais_cliente = pedido.get('pais', 'N/A')
        contacto_cliente = pedido.get('contacto', 'N/A')
        total_gift = pedido.get('total', 0)
        simbolo_moeda = pedido.get('simbolo_moeda', '€')
        
        # 🔥 DETALHES ESPECÍFICOS DO GIFT
        nome_gift = pedido.get('nome_gift', 'Sem nome')
        frase_gift = pedido.get('frase_gift', 'Sem frase')
        
        # 🔥 FUNÇÃO PARA ESCAPAR CARACTERES PROBLEMÁTICOS
        def limpar_texto(texto):
            if not texto:
                return "N/A"
            # Remover caracteres que causam problemas no Markdown
            caracteres_problematicos = ['*', '_', '`', '[', ']', '(', ')', '~', '>', '#', '+', '-', '=', '|', '{', '}']
            for char in caracteres_problematicos:
                texto = texto.replace(char, '')
            return texto.strip()
        
        # 🔥 LIMPAR TEXTOS
        nome_cliente_limpo = limpar_texto(nome_cliente)
        nome_gift_limpo = limpar_texto(nome_gift)
        frase_gift_limpo = limpar_texto(frase_gift)
        
        # 🔥 MENSAGEM PARA GIFT (APENAS 1 OFERTA - 30% OFF)
        mensagem = f"""🚨 *RELATÓRIO DE RECUSA DE GIFT* - #{pedido_id}

*👤 CLIENTE*
Nome: {nome_cliente_limpo}
Email: {email_cliente}
País: {pais_cliente}
Contacto: {contacto_cliente}

*🎁 DETALHES DO PORTA-CHAVES*
• Tamanho: 2.5" | 6.4cm
• Nome na Box: {nome_gift_limpo}
• Frase na Box: "{frase_gift_limpo}"
• Tipo: Porta-Chaves Personalizado

*💰 OFERTA RECUSADA*
🎁 *Porta-Chaves Personalizado (30% OFF)*: {simbolo_moeda}{total_gift:.2f} ❌

🚨 *CLIENTE RECUSOU OFERTA SURPRESA DE PORTA-CHAVES PERSONALIZADO*"""

        print(f"📏 Tamanho da mensagem Gift: {len(mensagem)} caracteres")
        
        # 🔥 BOTÃO PARA CONTACTAR CLIENTE
        keyboard = []
        chat_id_cliente = pedido.get('chat_id')
        if chat_id_cliente:
            keyboard.append([
                InlineKeyboardButton("📞 Contactar Cliente", url=f"tg://user?id={chat_id_cliente}")
            ])
        
        reply_markup = InlineKeyboardMarkup(keyboard) if keyboard else None
        
        # 🔥 ENVIAR GIFT COM FOTO
        if pedido.get('foto_id'):
            try:
                print("📸 Enviando Gift com foto...")
                await context.bot.send_photo(
                    chat_id=CHAT_SUPORTE_PAGOS,
                    photo=pedido['foto_id'],
                    caption=mensagem[:1024],  # Limitar caption a 1024 caracteres
                    parse_mode="Markdown",
                    reply_markup=reply_markup
                )
                print("✅ Gift com foto enviado!")
                
                # Se a mensagem for muito longa, enviar o resto
                if len(mensagem) > 1024:
                    resto = mensagem[1024:]
                    await context.bot.send_message(
                        chat_id=CHAT_SUPORTE_PAGOS,
                        text=resto[:4096],
                        parse_mode="Markdown"
                    )
                    print(f"✅ Texto adicional enviado ({len(resto)} chars)")
                    
            except Exception as e:
                print(f"❌ Erro ao enviar Gift com foto: {e}")
                # Fallback: enviar sem foto
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_PAGOS,
                    text=mensagem[:4096],
                    parse_mode="Markdown",
                    reply_markup=reply_markup
                )
        else:
            print("🖼️ Gift sem foto, enviando apenas texto")
            await context.bot.send_message(
                chat_id=CHAT_SUPORTE_PAGOS,
                text=mensagem[:4096],
                parse_mode="Markdown",
                reply_markup=reply_markup
            )
        
        print(f"✅ Relatório Gift enviado: #{pedido_id}")
        
    except Exception as e:
        print(f"❌ ERRO ao enviar relatório Gift: {e}")
        import traceback
        print(f"🔍 Traceback Gift: {traceback.format_exc()}")








async def todas_recusadas(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para outros problemas - CORRIGIDO COM TRADUÇÃO"""
    
    CHAT_SUPORTE_PAGOS = os.getenv("CHAT_SUPORTE_PAGOS")  # Ou o canal apropriado
        
    if not CHAT_SUPORTE_PAGOS:
         print("⚠️ AVISO: CHAT_SUPORTE_PAGOS não configurado")
         return  # Sai da função sem enviar
        
    try:
         CHAT_SUPORTE_PAGOS = int(CHAT_SUPORTE_PAGOS)
    except ValueError:
            print("⚠️ AVISO: CHAT_SUPORTE_PAGOS inválido")
            return  # Sai da função sem enviar
    
    try:
        # 🔥 1. SE É CLIQUE NO BOTÃO - INICIAR ESTADO
        if update.callback_query:
            query = update.callback_query
            await query.answer()
            
            pedido_id = query.data.replace("todas_recusadas_", "")
            
            print(f"🔴 todas_recusadas INICIADO: #{pedido_id}")
            
            # 🔥 OBTER IDIOMA DO PEDIDO OU DO CONTEXT
            pedido = PEDIDOS_RECUSADOS.get(pedido_id) or PEDIDOS_REGISTO.get(pedido_id, {})
            idioma = pedido.get('idioma', context.user_data.get('idioma', 'portugues'))
            print(f"🌐 Idioma detectado para todas_recusadas: {idioma}")
            
            # 🔥 DEFINIR ESTADO DE FORMA EXPLÍCITA
            context.user_data['conversation_state'] = 'todas_recusadas'
            context.user_data['ultimo_pedido_problema'] = pedido_id
            context.user_data['idioma'] = idioma  # Guardar idioma para a resposta
            
            # 🔥 DEBUG PARA CONFIRMAR
            print(f"🟢 ESTADO DEFINIDO: {context.user_data.get('conversation_state')}")
            print(f"🟢 user_data: {context.user_data}")
            
            # 🔥 TEXTOS POR IDIOMA PARA INÍCIO DO RELATO
            textos_inicio = {
                'portugues': {
                    'titulo': "📝 *Descreva o seu problema*\n\n",
                    'texto1': "Vimos que recusou o seu pedido!\n",
                    'texto2': "Por favor, descreva o que aconteceu!\n\n",
                    'pode_enviar': "*Pode enviar:*\n",
                    'opcao1': "• Texto com a descrição\n",
                    'opcao2': "• Foto/screenshot do problema\n\n",
                    'final': "*A nossa equipa irá ajudar!*"
                },
                'ingles': {
                    'titulo': "📝 *Describe your problem*\n\n",
                    'texto1': "We saw that you declined your order!\n",
                    'texto2': "Please describe what happened!\n\n",
                    'pode_enviar': "*You can send:*\n",
                    'opcao1': "• Text description\n",
                    'opcao2': "• Photo/screenshot of the problem\n\n",
                    'final': "*Our team will help!*"
                },
                'espanhol': {
                    'titulo': "📝 *Describa su problema*\n\n",
                    'texto1': "¡Vimos que rechazó su pedido!\n",
                    'texto2': "¡Por favor, describa lo que sucedió!\n\n",
                    'pode_enviar': "*Puede enviar:*\n",
                    'opcao1': "• Texto con la descripción\n",
                    'opcao2': "• Foto/captura de pantalla del problema\n\n",
                    'final': "¡*Nuestro equipo le ayudará!*"
                },
                'italiano': {
                    'titulo': "📝 *Descrivi il tuo problema*\n\n",
                    'texto1': "Abbiamo visto che hai rifiutato il tuo ordine!\n",
                    'texto2': "Per favore, descrivi cosa è successo!\n\n",
                    'pode_enviar': "*Puoi inviare:*\n",
                    'opcao1': "• Testo con la descrizione\n",
                    'opcao2': "• Foto/screenshot del problema\n\n",
                    'final': "*Il nostro team ti aiuterà!*"
                },
                'alemao': {
                    'titulo': "📝 *Beschreiben Sie Ihr Problem*\n\n",
                    'texto1': "Wir haben gesehen, dass Sie Ihre Bestellung abgelehnt haben!\n",
                    'texto2': "Bitte beschreiben Sie, was passiert ist!\n\n",
                    'pode_enviar': "*Sie können senden:*\n",
                    'opcao1': "• Text mit Beschreibung\n",
                    'opcao2': "• Foto/Screenshot des Problems\n\n",
                    'final': "*Unser Team wird helfen!*"
                },
                'frances': {
                    'titulo': "📝 *Décrivez votre problème*\n\n",
                    'texto1': "Nous avons vu que vous avez refusé votre commande!\n",
                    'texto2': "Veuillez décrire ce qui s'est passé!\n\n",
                    'pode_enviar': "*Vous pouvez envoyer:*\n",
                    'opcao1': "• Texte avec description\n",
                    'opcao2': "• Photo/capture d'écran du problème\n\n",
                    'final': "*Notre équipe vous aidera!*"
                }
            }
            
            textos = textos_inicio.get(idioma, textos_inicio['portugues'])
            
            # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
            mensagem_inicio = (
                f"{textos['titulo']}"
                f"{textos['texto1']}"
                f"{textos['texto2']}"
                f"{textos['pode_enviar']}"
                f"{textos['opcao1']}"
                f"{textos['opcao2']}"
                f"{textos['final']}"
            )
            
            await query.edit_message_text(
                text=mensagem_inicio,
                parse_mode="Markdown"
            )
            return
        
        # 🔥 2. SE É MENSAGEM - VERIFICAR SE ESTÁ NO ESTADO CORRETO
        elif update.message:
            current_state = context.user_data.get('conversation_state')
            print(f"🔍 todas_recusadas recebeu mensagem - Estado: {current_state}")
            
            # 🔥 SÓ PROCESSAR SE ESTIVER NO ESTADO todas_recusadas
            if current_state != 'todas_recusadas':
                print(f"⚠️ todas_recusadas IGNORADO - estado incorreto: {current_state}")
                return
                
            user = update.message.from_user
            chat_id = update.message.chat_id
            pedido_id = context.user_data.get('ultimo_pedido_problema', 'N/A')
            
            # 🔥 OBTER IDIOMA SALVO NO CONTEXT
            idioma = context.user_data.get('idioma', 'portugues')
            
            # Processar texto OU foto
            if update.message.text:
                problema = update.message.text
                print(f"📝 todas_recusadas CAPTUROU TEXTO: {problema}")
            elif update.message.photo:
                problema = {
                    'portugues': "📸 Foto enviada como problema",
                    'ingles': "📸 Photo sent as problem",
                    'espanhol': "📸 Foto enviada como problema",
                    'italiano': "📸 Foto inviata come problema",
                    'alemao': "📸 Foto als Problem gesendet",
                    'frances': "📸 Photo envoyée comme problème"
                }.get(idioma, "📸 Foto enviada como problema")
                print(f"📸 todas_recusadas CAPTUROU FOTO")
            else:
                return
            
            # 🔥 ENVIAR PARA SUPORTE
            pedido = PEDIDOS_RECUSADOS.get(pedido_id) or PEDIDOS_REGISTO.get(pedido_id, {})
            
            mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO - PEDIDO #{pedido_id}*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}

📦 *Detalhes do Pedido:*
• ID: {pedido_id}
• Cliente: {pedido.get('nome', 'N/A')}
• Email: {pedido.get('email', 'N/A')}
• Produto: {pedido.get('tipo_cartoon', 'N/A')}
• Valor: {pedido.get('simbolo_moeda', '')}{pedido.get('total', 0):.2f}

📝 *Problema:*
{problema}"""
            
            keyboard_suporte = [
                [InlineKeyboardButton("📞 Contactar", url=f"tg://user?id={chat_id}")]
            ]
            
            # Se for foto, enviar a foto também
            if update.message.photo:
                photo_file = await update.message.photo[-1].get_file()
                await context.bot.send_photo(
                    chat_id=CHAT_SUPORTE_PAGOS,
                    photo=photo_file.file_id,
                    caption=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
            else:
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_PAGOS,
                    text=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
            
            # 🔥 TEXTOS POR IDIOMA PARA CONFIRMAÇÃO AO USUÁRIO
            textos_confirmacao = {
                'portugues': {
                    'titulo': "✅ *Problema reportado com sucesso!*\n\n",
                    'texto1': "Nossa equipa vai resolver o seu problema brevemente.\n\n",
                    'ajuda': "*Se precisar de mais ajuda, clique em:*\n",
                    'opcao1': "👉 /start - Para criar uma nova encomenda\n",
                    'opcao2': "👉 /help - Para ver opções de ajuda\n\n",
                    'final': "*Obrigado pela sua paciência.*"
                },
                'ingles': {
                    'titulo': "✅ *Problem reported successfully!*\n\n",
                    'texto1': "Our team will resolve your problem shortly.\n\n",
                    'ajuda': "*If you need more help, click on:*\n",
                    'opcao1': "👉 /start - To create a new order\n",
                    'opcao2': "👉 /help - To see help options\n\n",
                    'final': "*Thank you for your patience.*"
                },
                'espanhol': {
                    'titulo': "✅ *¡Problema reportado con éxito!*\n\n",
                    'texto1': "Nuestro equipo resolverá su problema en breve.\n\n",
                    'ajuda': "*Si necesita más ayuda, haga clic en:*\n",
                    'opcao1': "👉 /start - Para crear un nuevo pedido\n",
                    'opcao2': "👉 /help - Para ver opciones de ayuda\n\n",
                    'final': "*Gracias por su paciencia.*"
                },
                'italiano': {
                    'titulo': "✅ *Problema segnalato con successo!*\n\n",
                    'texto1': "Il nostro team risolverà il tuo problema a breve.\n\n",
                    'ajuda': "*Se hai bisogno di ulteriore aiuto, clicca su:*\n",
                    'opcao1': "👉 /start - Per creare un nuovo ordine\n",
                    'opcao2': "👉 /help - Per vedere le opzioni di aiuto\n\n",
                    'final': "*Grazie per la tua pazienza.*"
                },
                'alemao': {
                    'titulo': "✅ *Problem erfolgreich gemeldet!*\n\n",
                    'texto1': "Unser Team wird Ihr Problem bald lösen.\n\n",
                    'ajuda': "*Wenn Sie weitere Hilfe benötigen, klicken Sie auf:*\n",
                    'opcao1': "👉 /start - Um eine neue Bestellung zu erstellen\n",
                    'opcao2': "👉 /help - Um Hilfeoptionen zu sehen\n\n",
                    'final': "*Danke für Ihre Geduld.*"
                },
                'frances': {
                    'titulo': "✅ *Problème signalé avec succès!*\n\n",
                    'texto1': "Notre équipe résoudra votre problème sous peu.\n\n",
                    'ajuda': "*Si vous avez besoin de plus d'aide, cliquez sur:*\n",
                    'opcao1': "👉 /start - Pour créer une nouvelle commande\n",
                    'opcao2': "👉 /help - Pour voir les options d'aide\n\n",
                    'final': "*Merci pour votre patience.*"
                }
            }
            
            textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
            
            # 🔥 CONSTRUIR MENSAGEM DE CONFIRMAÇÃO TRADUZIDA
            mensagem_confirmacao = (
                f"{textos['titulo']}"
                f"{textos['texto1']}"
                f"{textos['ajuda']}"
                f"{textos['opcao1']}"
                f"{textos['opcao2']}"
                f"{textos['final']}"
            )
            
            # CONFIRMAR AO USUÁRIO
            await update.message.reply_text(
                mensagem_confirmacao,
                parse_mode="Markdown"
            )
            
            # 🔥 LIMPAR ESTADO COMPLETAMENTE
            context.user_data.pop('conversation_state', None)
            context.user_data.pop('ultimo_pedido_problema', None)
            context.user_data.pop('idioma', None)
            print("✅ ESTADO LIMPO: todas_recusadas concluído")
            
    except Exception as e:
        print(f"❌ ERRO em todas_recusadas: {e}")
        import traceback
        traceback.print_exc()









async def problema_outro(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler para outros problemas - CORRIGIDO COM TRADUÇÃO"""
    
    CHAT_SUPORTE_CLIENTES = os.getenv("CHAT_SUPORTE_CLIENTES")  # Canal apropriado
    
    if not CHAT_SUPORTE_CLIENTES:
        print("⚠️ AVISO: CHAT_SUPORTE_CLIENTES não configurado")
        await update.callback_query.answer("❌ Erro interno. Tente mais tarde.")
        return
    
    try:
        CHAT_SUPORTE_CLIENTES = int(CHAT_SUPORTE_CLIENTES)
    except ValueError:
        print("⚠️ AVISO: CHAT_SUPORTE_CLIENTES inválido")
        await update.callback_query.answer("❌ Erro interno. Tente mais tarde.")
        return
    
    try:
        # 🔥 1. SE É CLIQUE NO BOTÃO - INICIAR ESTADO
        if update.callback_query:
            query = update.callback_query
            await query.answer()
            
            pedido_id = query.data.replace("problema_outro_", "")
            
            print(f"🔴 PROBLEMA_OUTRO INICIADO: #{pedido_id}")
            
            # 🔥 OBTER IDIOMA DO PEDIDO OU DO CONTEXT
            pedido = PEDIDOS_REGISTO.get(pedido_id, {})
            idioma = pedido.get('idioma', context.user_data.get('idioma', 'portugues'))
            print(f"🌐 Idioma detectado para problema_outro: {idioma}")
            
            # 🔥 DEFINIR ESTADO DE FORMA EXPLÍCITA
            context.user_data['conversation_state'] = 'problema_outro'
            context.user_data['ultimo_pedido_problema'] = pedido_id
            context.user_data['idioma'] = idioma  # Guardar idioma para a resposta
            
            # 🔥 DEBUG PARA CONFIRMAR
            print(f"🟢 ESTADO DEFINIDO: {context.user_data.get('conversation_state')}")
            print(f"🟢 user_data: {context.user_data}")
            
            # 🔥 TEXTOS POR IDIOMA PARA INÍCIO DO RELATO
            textos_inicio = {
                'portugues': {
                    'titulo': "📝 *Descreva o seu problema*\n\n",
                    'texto': "Por favor, descreva o problema que está a enfrentar.\n\n",
                    'pode_enviar': "*Pode enviar:*\n",
                    'opcao1': "• Texto com a descrição\n",
                    'opcao2': "• Foto/screenshot do problema\n\n",
                    'final': "*A nossa equipa irá ajudar!*"
                },
                'ingles': {
                    'titulo': "📝 *Describe your problem*\n\n",
                    'texto': "Please describe the problem you are facing.\n\n",
                    'pode_enviar': "*You can send:*\n",
                    'opcao1': "• Text description\n",
                    'opcao2': "• Photo/screenshot of the problem\n\n",
                    'final': "*Our team will help!*"
                },
                'espanhol': {
                    'titulo': "📝 *Describa su problema*\n\n",
                    'texto': "Por favor, describa el problema que está enfrentando.\n\n",
                    'pode_enviar': "*Puede enviar:*\n",
                    'opcao1': "• Texto con la descripción\n",
                    'opcao2': "• Foto/captura de pantalla del problema\n\n",
                    'final': "*¡Nuestro equipo le ayudará!*"
                },
                'italiano': {
                    'titulo': "📝 *Descrivi il tuo problema*\n\n",
                    'texto': "Per favore, descrivi il problema che stai affrontando.\n\n",
                    'pode_enviar': "*Puoi inviare:*\n",
                    'opcao1': "• Testo con la descrizione\n",
                    'opcao2': "• Foto/screenshot del problema\n\n",
                    'final': "*Il nostro team ti aiuterà!*"
                },
                'alemao': {
                    'titulo': "📝 *Beschreiben Sie Ihr Problem*\n\n",
                    'texto': "Bitte beschreiben Sie das Problem, mit dem Sie konfrontiert sind.\n\n",
                    'pode_enviar': "*Sie können senden:*\n",
                    'opcao1': "• Text mit Beschreibung\n",
                    'opcao2': "• Foto/Screenshot des Problems\n\n",
                    'final': "*Unser Team wird helfen!*"
                },
                'frances': {
                    'titulo': "📝 *Décrivez votre problème*\n\n",
                    'texto': "Veuillez décrire le problème auquel vous êtes confronté.\n\n",
                    'pode_enviar': "*Vous pouvez envoyer:*\n",
                    'opcao1': "• Texte avec description\n",
                    'opcao2': "• Photo/capture d'écran du problème\n\n",
                    'final': "*Notre équipe vous aidera!*"
                }
            }
            
            textos = textos_inicio.get(idioma, textos_inicio['portugues'])
            
            # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
            mensagem_inicio = (
                f"{textos['titulo']}"
                f"{textos['texto']}"
                f"{textos['pode_enviar']}"
                f"{textos['opcao1']}"
                f"{textos['opcao2']}"
                f"{textos['final']}"
            )
            
            await query.edit_message_text(
                text=mensagem_inicio,
                parse_mode="Markdown"
            )
            return
        
        # 🔥 2. SE É MENSAGEM - VERIFICAR SE ESTÁ NO ESTADO CORRETO
        elif update.message:
            current_state = context.user_data.get('conversation_state')
            print(f"🔍 problema_outro recebeu mensagem - Estado: {current_state}")
            
            # 🔥 SÓ PROCESSAR SE ESTIVER NO ESTADO problema_outro
            if current_state != 'problema_outro':
                print(f"⚠️ problema_outro IGNORADO - estado incorreto: {current_state}")
                return
                
            user = update.message.from_user
            chat_id = update.message.chat_id
            pedido_id = context.user_data.get('ultimo_pedido_problema', 'N/A')
            
            # 🔥 OBTER IDIOMA SALVO NO CONTEXT
            idioma = context.user_data.get('idioma', 'portugues')
            
            # Processar texto OU foto
            if update.message.text:
                problema = update.message.text
                print(f"📝 problema_outro CAPTUROU TEXTO: {problema}")
            elif update.message.photo:
                problema = {
                    'portugues': "📸 Foto enviada como problema",
                    'ingles': "📸 Photo sent as problem",
                    'espanhol': "📸 Foto enviada como problema",
                    'italiano': "📸 Foto inviata come problema",
                    'alemao': "📸 Foto als Problem gesendet",
                    'frances': "📸 Photo envoyée comme problème"
                }.get(idioma, "📸 Foto enviada como problema")
                print(f"📸 problema_outro CAPTUROU FOTO")
            else:
                return
            
            # 🔥 ENVIAR PARA SUPORTE
            pedido = PEDIDOS_REGISTO.get(pedido_id, {})
            
            mensagem_suporte = f"""
🚨 *PROBLEMA REPORTADO - PEDIDO #{pedido_id}*

👤 *Cliente:* {user.first_name} (@{user.username or 'N/A'})
💬 *Chat ID:* `{chat_id}`
⏰ *Data:* {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}

📦 *Detalhes do Pedido:*
• ID: {pedido_id}
• Cliente: {pedido.get('nome', 'N/A')}
• Email: {pedido.get('email', 'N/A')}
• Produto: {pedido.get('tipo_cartoon', 'N/A')}
• Valor: {pedido.get('simbolo_moeda', '')}{pedido.get('total', 0):.2f}

📝 *Problema:*
{problema}"""
            
            keyboard_suporte = [
                [InlineKeyboardButton("📞 Contactar", url=f"tg://user?id={chat_id}")]
            ]
            
            # Se for foto, enviar a foto também
            if update.message.photo:
                photo_file = await update.message.photo[-1].get_file()
                await context.bot.send_photo(
                    chat_id=CHAT_SUPORTE_CLIENTES,
                    photo=photo_file.file_id,
                    caption=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
            else:
                await context.bot.send_message(
                    chat_id=CHAT_SUPORTE_CLIENTES,
                    text=mensagem_suporte,
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup(keyboard_suporte)
                )
            
            # 🔥 TEXTOS POR IDIOMA PARA CONFIRMAÇÃO AO USUÁRIO
            textos_confirmacao = {
                'portugues': {
                    'titulo': "✅ *Problema reportado com sucesso!*\n\n",
                    'texto': "Nossa equipa técnica vai resolver o seu problema brevemente.\n\n",
                    'ajuda': "*Se precisar de mais ajuda, clique em:*\n",
                    'opcao1': "👉 /start - Para criar uma nova encomenda\n",
                    'opcao2': "👉 /help - Para ver opções de ajuda\n\n",
                    'final': "*Obrigado pela sua paciência.*"
                },
                'ingles': {
                    'titulo': "✅ *Problem reported successfully!*\n\n",
                    'texto': "Our technical team will resolve your problem shortly.\n\n",
                    'ajuda': "*If you need more help, click on:*\n",
                    'opcao1': "👉 /start - To create a new order\n",
                    'opcao2': "👉 /help - To see help options\n\n",
                    'final': "*Thank you for your patience.*"
                },
                'espanhol': {
                    'titulo': "✅ *¡Problema reportado con éxito!*\n\n",
                    'texto': "Nuestro equipo técnico resolverá su problema en breve.\n\n",
                    'ajuda': "*Si necesita más ayuda, haga clic en:*\n",
                    'opcao1': "👉 /start - Para crear un nuevo pedido\n",
                    'opcao2': "👉 /help - Para ver opciones de ayuda\n\n",
                    'final': "*Gracias por su paciencia.*"
                },
                'italiano': {
                    'titulo': "✅ *Problema segnalato con successo!*\n\n",
                    'texto': "Il nostro team tecnico risolverà il tuo problema a breve.\n\n",
                    'ajuda': "*Se hai bisogno di ulteriore aiuto, clicca su:*\n",
                    'opcao1': "👉 /start - Per creare un nuovo ordine\n",
                    'opcao2': "👉 /help - Per vedere le opzioni di aiuto\n\n",
                    'final': "*Grazie per la tua pazienza.*"
                },
                'alemao': {
                    'titulo': "✅ *Problem erfolgreich gemeldet!*\n\n",
                    'texto': "Unser technisches Team wird Ihr Problem bald lösen.\n\n",
                    'ajuda': "*Wenn Sie weitere Hilfe benötigen, klicken Sie auf:*\n",
                    'opcao1': "👉 /start - Um eine neue Bestellung zu erstellen\n",
                    'opcao2': "👉 /help - Um Hilfeoptionen zu sehen\n\n",
                    'final': "*Danke für Ihre Geduld.*"
                },
                'frances': {
                    'titulo': "✅ *Problème signalé avec succès!*\n\n",
                    'texto': "Notre équipe technique résoudra votre problème sous peu.\n\n",
                    'ajuda': "*Si vous avez besoin de plus d'aide, cliquez sur:*\n",
                    'opcao1': "👉 /start - Pour créer une nouvelle commande\n",
                    'opcao2': "👉 /help - Pour voir les options d'aide\n\n",
                    'final': "*Merci pour votre patience.*"
                }
            }
            
            textos = textos_confirmacao.get(idioma, textos_confirmacao['portugues'])
            
            # 🔥 CONSTRUIR MENSAGEM DE CONFIRMAÇÃO TRADUZIDA
            mensagem_confirmacao = (
                f"{textos['titulo']}"
                f"{textos['texto']}"
                f"{textos['ajuda']}"
                f"{textos['opcao1']}"
                f"{textos['opcao2']}"
                f"{textos['final']}"
            )
            
            # CONFIRMAR AO USUÁRIO
            await update.message.reply_text(
                mensagem_confirmacao,
                parse_mode="Markdown"
            )
            
            # 🔥 LIMPAR ESTADO COMPLETAMENTE
            context.user_data.pop('conversation_state', None)
            context.user_data.pop('ultimo_pedido_problema', None)
            context.user_data.pop('idioma', None)
            print("✅ ESTADO LIMPO: problema_outro concluído")
            
    except Exception as e:
        print(f"❌ ERRO em problema_outro: {e}")
        import traceback
        traceback.print_exc()








async def iniciar_temporizador_limpeza_30min(context, chat_id, message_id):
    """Temporizador de 30 minutos - VOLTA AO INÍCIO AUTOMATICAMENTE"""
    try:
        print(f"⏰ Temporizador 30min INICIADO para chat {chat_id}")
        
        # Criar task assíncrona
        task = asyncio.create_task(temporizador_30min_task(context, chat_id, message_id))
        TEMPORIZADORES_ATIVOS[chat_id] = task
        
        # Esperar a task completar (ou ser cancelada)
        await task
        
    except asyncio.CancelledError:
        print(f"✅ Temporizador 30min CANCELADO para chat {chat_id}")
    except Exception as e:
        print(f"❌ Erro no temporizador 30min: {e}")




async def temporizador_30min_task(context, chat_id, message_id):
    """Task separada para o temporizador de 30min"""
    try:
        # Esperar 30 minutos (1800 segundos)
        await asyncio.sleep(1800)
        
        # Verificar se ainda está ativo (pode ter sido cancelado)
        if chat_id in TEMPORIZADORES_ATIVOS:
            print(f"🕒 TEMPORIZADOR 30min EXPIRADO - Voltando ao início para chat {chat_id}")
            
            # 🔥 VOLTAR AO INÍCIO AUTOMATICAMENTE
            await voltar_ao_inicio_automatico(context, chat_id)
                
    except asyncio.CancelledError:
        print(f"✅ Task temporizador 30min CANCELADA para chat {chat_id}")
        raise  # Re-lançar a exceção





async def voltar_ao_inicio_automatico(context, chat_id):
    """Volta ao início automaticamente após 30min COM TRADUÇÃO"""
    try:
        print(f"🔄 Voltando ao início automaticamente para chat {chat_id}")
        
        # 🔥 OBTER IDIOMA DO USER_DATA OU DO CONTEXT
        idioma = context.user_data.get('idioma', 'portugues')
        print(f"🌐 Idioma detectado para voltar_ao_inicio_automatico: {idioma}")
        
        # 🔥 REMOVER TEMPORIZADOR DA LISTA
        if chat_id in TEMPORIZADORES_ATIVOS:
            del TEMPORIZADORES_ATIVOS[chat_id]
        
        # 🔥 LIMPAR DADOS (usar approach diferente)
        try:
            # Tentar limpar user_data de forma mais agressiva
            context.user_data.clear()
            print(f"✅ Dados limpos automaticamente para chat {chat_id}")
        except:
            pass
        
        # 🔥 TEXTOS POR IDIOMA PARA VOLTA AO INÍCIO
        textos_inicio = {
            'portugues': {
                'mensagem': "👋 Olá! Bem-vindo à *GodsPlan*, vamos criar o seu cartoon?",
                'botao': "CREATE MY CARTOON"
            },
            'ingles': {
                'mensagem': "👋 Hello! Welcome to *GodsPlan*, let's create your cartoon?",
                'botao': "CREATE MY CARTOON"
            },
            'espanhol': {
                'mensagem': "👋 ¡Hola! Bienvenido a *GodsPlan*, ¿vamos a crear tu caricatura?",
                'botao': "CREATE MY CARTOON"
            },
            'italiano': {
                'mensagem': "👋 Ciao! Benvenuto a *GodsPlan*, creiamo il tuo cartoon?",
                'botao': "CREATE MY CARTOON"
            },
            'alemao': {
                'mensagem': "👋 Hallo! Willkommen bei *GodsPlan*, erstellen wir Ihren Cartoon?",
                'botao': "CREATE MY CARTOON"
            },
            'frances': {
                'mensagem': "👋 Bonjour! Bienvenue à *GodsPlan*, créons votre dessin animé?",
                'botao': "CREATE MY CARTOON"
            }
        }
        
        textos = textos_inicio.get(idioma, textos_inicio['portugues'])
        
        # 🔥 CONSTRUIR MENSAGEM TRADUZIDA
        mensagem = textos['mensagem']
        texto_botao = textos['botao']
        
        keyboard = [[InlineKeyboardButton(texto_botao, callback_data="mycartoon")]]
        reply_markup = InlineKeyboardMarkup(keyboard)

        await context.bot.send_message(
            chat_id=chat_id,
            text=mensagem,
            reply_markup=reply_markup,
            parse_mode="Markdown"
        )
        
        print(f"✅ Menu inicial automático mostrado para chat {chat_id} em {idioma}")
        
    except Exception as e:
        print(f"❌ Erro ao voltar ao início automático: {e}")






async def enviar_mensagem_agradecimento(chat_id, pedido_pago, amount):
    """Função async para enviar mensagem de agradecimento com todos os detalhes COM TRADUÇÃO"""
    try:
        # 🔥 OBTER IDIOMA DO PEDIDO
        idioma = pedido_pago.get('idioma', 'portugues')
        print(f"🌐 Idioma detectado para agradecimento: {idioma}")
        
        # 🔥 TEXTOS POR IDIOMA PARA OS EMOJIS INICIAIS
        textos_confetes = {
            'portugues': "🎊 *Parabéns pela sua encomenda!* 🎊",
            'ingles': "🎊 *Congratulations on your order!* 🎊",
            'espanhol': "🎊 *¡Felicitaciones por tu pedido!* 🎊",
            'italiano': "🎊 *Congratulazioni per il tuo ordine!* 🎊",
            'alemao': "🎊 *Herzlichen Glückwunsch zu Ihrer Bestellung!* 🎊",
            'frances': "🎊 *Félicitations pour votre commande !* 🎊"
        }
        
        # 🔥 PRIMEIRO: ENVIAR EMOJI DE CONFETES
        try:
            await bot.send_animation(
                chat_id=chat_id,
                animation="https://media.giphy.com/media/xT0xeuOy2Fcl9vDGiA/giphy.gif",
                caption=textos_confetes.get(idioma, textos_confetes['portugues'])
            )
        except:
            await bot.send_message(chat_id=chat_id, text="🎊🎉🎊🎉🎊🎉🎊🎉")
        
        # 🔥 SEGUNDO: OBTER INFORMAÇÕES DE MOEDA DO PEDIDO
        moeda_original = pedido_pago.get('moeda_original', 'EUR')
        simbolo_original = get_simbolo_moeda(moeda_original.lower())
        total_original = pedido_pago.get('total_pago_original', pedido_pago.get('total', 0))
        texto_valor_cliente = f"{simbolo_original}{total_original:.2f}"
        
        # 🔥 CONSTRUIR DETALHES ESPECÍFICOS BASEADOS NO TIPO
        tipo_cartoon = pedido_pago['tipo_cartoon']
        detalhes_especificos = ""
        
        # 🔥 TEXTOS POR IDIOMA PARA OS CAMPOS DE DETALHES
        textos_detalhes = {
            'portugues': {
                'estilo': "🖌 *Estilo:*",
                'profissao': "💼 *Profissão:*",
                'super_heroi': "🦸 *Super-Herói:*",
                'nome_cartoon': "📛 *Nome no Cartoon:*",
                'frase_cartoon': "💬 *Frase na Box:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nome da Família:*",
                'frase_family': "💬 *Frase da Família:*",
                'elementos_family': "👥 *Total de Elementos:*",
                'adultos_family': "👨‍👩 *Adultos:*",
                'criancas_family': "👧🧒 *Crianças:*",
                'animais_family': "🐱🐶 *Animais:*",
                'nome_animal': "🐾 *Nome do Animal:*",
                'tipo_animal': "🐕 *Tipo de Animal:*",
                'tipo_personalizado': "🎨 *Tipo de Peça:*",
                'nome_peca_personalizado': "📝 *Nome da Peça:*",
                'nome_personalizado': "🎭 *Nome do Cartoon:*",
                'frase_personalizado': "💭 *Frase do Elemento:*",
                'tamanho': "📏 *Tamanho:*",
                'data': "📅 *Data:*",
                'valor': "💰 *Valor Pago:*",
                'padrao': "Padrão"
            },
            'ingles': {
                'estilo': "🖌 *Style:*",
                'profissao': "💼 *Profession:*",
                'super_heroi': "🦸 *Superhero:*",
                'nome_cartoon': "📛 *Name in Cartoon:*",
                'frase_cartoon': "💬 *Phrase on Box:*",
                'nome_family': "👨‍👩‍👧‍👦 *Family Name:*",
                'frase_family': "💬 *Family Phrase:*",
                'elementos_family': "👥 *Total Elements:*",
                'adultos_family': "👨‍👩 *Adults:*",
                'criancas_family': "👧🧒 *Children:*",
                'animais_family': "🐱🐶 *Animals:*",
                'nome_animal': "🐾 *Animal Name:*",
                'tipo_animal': "🐕 *Animal Type:*",
                'tipo_personalizado': "🎨 *Piece Type:*",
                'nome_peca_personalizado': "📝 *Piece Name:*",
                'nome_personalizado': "🎭 *Cartoon Name:*",
                'frase_personalizado': "💭 *Element Phrase:*",
                'tamanho': "📏 *Size:*",
                'data': "📅 *Date:*",
                'valor': "💰 *Amount Paid:*",
                'padrao': "Standard"
            },
            'espanhol': {
                'estilo': "🖌 *Estilo:*",
                'profissao': "💼 *Profesión:*",
                'super_heroi': "🦸 *Superhéroe:*",
                'nome_cartoon': "📛 *Nombre en Caricatura:*",
                'frase_cartoon': "💬 *Frase en la Caja:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nombre de Familia:*",
                'frase_family': "💬 *Frase Familiar:*",
                'elementos_family': "👥 *Elementos Totales:*",
                'adultos_family': "👨‍👩 *Adultos:*",
                'criancas_family': "👧🧒 *Niños:*",
                'animais_family': "🐱🐶 *Animales:*",
                'nome_animal': "🐾 *Nombre del Animal:*",
                'tipo_animal': "🐕 *Tipo de Animal:*",
                'tipo_personalizado': "🎨 *Tipo de Pieza:*",
                'nome_peca_personalizado': "📝 *Nombre de la Pieza:*",
                'nome_personalizado': "🎭 *Nombre del Cartoon:*",
                'frase_personalizado': "💭 *Frase del Elemento:*",
                'tamanho': "📏 *Tamaño:*",
                'data': "📅 *Fecha:*",
                'valor': "💰 *Valor Pagado:*",
                'padrao': "Estándar"
            },
            'italiano': {
                'estilo': "🖌 *Stile:*",
                'profissao': "💼 *Professione:*",
                'super_heroi': "🦸 *Supereroe:*",
                'nome_cartoon': "📛 *Nome nel Cartoon:*",
                'frase_cartoon': "💬 *Frase sulla Scatola:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nome della Famiglia:*",
                'frase_family': "💬 *Frase Familiare:*",
                'elementos_family': "👥 *Elementi Totali:*",
                'adultos_family': "👨‍👩 *Adulti:*",
                'criancas_family': "👧🧒 *Bambini:*",
                'animais_family': "🐱🐶 *Animali:*",
                'nome_animal': "🐾 *Nome dell'Animale:*",
                'tipo_animal': "🐕 *Tipo di Animale:*",
                'tipo_personalizado': "🎨 *Tipo di Pezzo:*",
                'nome_peca_personalizado': "📝 *Nome del Pezzo:*",
                'nome_personalizado': "🎭 *Nome del Cartoon:*",
                'frase_personalizado': "💭 *Frase dell'Elemento:*",
                'tamanho': "📏 *Dimensione:*",
                'data': "📅 *Data:*",
                'valor': "💰 *Importo Pagato:*",
                'padrao': "Standard"
            },
            'alemao': {
                'estilo': "🖌 *Stil:*",
                'profissao': "💼 *Beruf:*",
                'super_heroi': "🦸 *Superheld:*",
                'nome_cartoon': "📛 *Name im Cartoon:*",
                'frase_cartoon': "💬 *Satz auf der Box:*",
                'nome_family': "👨‍👩‍👧‍👦 *Familienname:*",
                'frase_family': "💬 *Familiensatz:*",
                'elementos_family': "👥 *Gesamtelemente:*",
                'adultos_family': "👨‍👩 *Erwachsene:*",
                'criancas_family': "👧🧒 *Kinder:*",
                'animais_family': "🐱🐶 *Tiere:*",
                'nome_animal': "🐾 *Tiername:*",
                'tipo_animal': "🐕 *Tierart:*",
                'tipo_personalizado': "🎨 *Stücktyp:*",
                'nome_peca_personalizado': "📝 *Stückname:*",
                'nome_personalizado': "🎭 *Cartoon-Name:*",
                'frase_personalizado': "💭 *Elementsatz:*",
                'tamanho': "📏 *Größe:*",
                'data': "📅 *Datum:*",
                'valor': "💰 *Bezahlter Betrag:*",
                'padrao': "Standard"
            },
            'frances': {
                'estilo': "🖌 *Style:*",
                'profissao': "💼 *Profession:*",
                'super_heroi': "🦸 *Super-héros:*",
                'nome_cartoon': "📛 *Nom dans le Dessin Animé:*",
                'frase_cartoon': "💬 *Phrase sur la Boîte:*",
                'nome_family': "👨‍👩‍👧‍👦 *Nom de Famille:*",
                'frase_family': "💬 *Phrase Familiale:*",
                'elementos_family': "👥 *Éléments Totaux:*",
                'adultos_family': "👨‍👩 *Adultes:*",
                'criancas_family': "👧🧒 *Enfants:*",
                'animais_family': "🐱🐶 *Animaux:*",
                'nome_animal': "🐾 *Nom de l'Animal:*",
                'tipo_animal': "🐕 *Type d'Animal:*",
                'tipo_personalizado': "🎨 *Type de Pièce:*",
                'nome_peca_personalizado': "📝 *Nom de la Pièce:*",
                'nome_personalizado': "🎭 *Nom du Dessin Animé:*",
                'frase_personalizado': "💭 *Phrase de l'Élément:*",
                'tamanho': "📏 *Taille:*",
                'data': "📅 *Date:*",
                'valor': "💰 *Montant Payé:*",
                'padrao': "Standard"
            }
        }
        
        textos = textos_detalhes.get(idioma, textos_detalhes['portugues'])
        
        if "individual" in tipo_cartoon.lower():
            # DETALHES PARA CARTOON INDIVIDUAL
            detalhes_especificos = f"{textos['estilo']} {pedido_pago.get('estilo_cartoon', textos['padrao'])}\n"
            
            # CAMPOS PERSONALIZADOS PARA INDIVIDUAL
            if pedido_pago.get('profissao'):
                detalhes_especificos += f"{textos['profissao']} {pedido_pago['profissao']}\n"
            if pedido_pago.get('super_heroi'):
                detalhes_especificos += f"{textos['super_heroi']} {pedido_pago['super_heroi']}\n"
            if pedido_pago.get('nome_cartoon'):
                detalhes_especificos += f"{textos['nome_cartoon']} {pedido_pago['nome_cartoon']}\n"
            if pedido_pago.get('frase_cartoon'):
                detalhes_especificos += f"{textos['frase_cartoon']} {pedido_pago['frase_cartoon']}\n"
                
        elif "family" in tipo_cartoon.lower() or "família" in tipo_cartoon.lower() or "familia" in tipo_cartoon.lower():
            # DETALHES PARA CARTOON FAMILIAR
            if pedido_pago.get('nome_family'):
                detalhes_especificos += f"{textos['nome_family']} {pedido_pago['nome_family']}\n"
            if pedido_pago.get('frase_family'):
                detalhes_especificos += f"{textos['frase_family']} {pedido_pago['frase_family']}\n"
            if pedido_pago.get('elementos_family'):
                detalhes_especificos += f"{textos['elementos_family']} {pedido_pago['elementos_family']}\n"
            if pedido_pago.get('adultos_family'):
                detalhes_especificos += f"{textos['adultos_family']} {pedido_pago['adultos_family']}\n"
            if pedido_pago.get('criancas_family'):
                detalhes_especificos += f"{textos['criancas_family']} {pedido_pago['criancas_family']}\n"
            if pedido_pago.get('animais_family'):
                detalhes_especificos += f"{textos['animais_family']} {pedido_pago['animais_family']}\n"
            if pedido_pago.get('nome_animal'):
                detalhes_especificos += f"{textos['nome_animal']} {pedido_pago['nome_animal']}\n"
            if pedido_pago.get('tipo_animal'):
                detalhes_especificos += f"{textos['tipo_animal']} {pedido_pago['tipo_animal']}\n"
                
        elif "personalizado" in tipo_cartoon.lower() or "custom" in tipo_cartoon.lower():
            # DETALHES PARA PEÇA PERSONALIZADA
            if pedido_pago.get('tipo_personalizado'):
                detalhes_especificos += f"{textos['tipo_personalizado']} {pedido_pago['tipo_personalizado']}\n"
            if pedido_pago.get('nome_peca_personalizado'):
                detalhes_especificos += f"{textos['nome_peca_personalizado']} {pedido_pago['nome_peca_personalizado']}\n"
            if pedido_pago.get('nome_personalizado'):
                detalhes_especificos += f"{textos['nome_personalizado']} {pedido_pago['nome_personalizado']}\n"
            if pedido_pago.get('frase_personalizado'):
                detalhes_especificos += f"{textos['frase_personalizado']} {pedido_pago['frase_personalizado']}\n"
        
        # 🔥 TEXTOS POR IDIOMA PARA A MENSAGEM PRINCIPAL
        textos_principal = {
            'portugues': {
                'titulo': "🎉 *PAGAMENTO CONFIRMADO!* 🎉\n\n",
                'parabens': "✨ *Parabéns pela sua encomenda!* ✨\n\n",
                'detalhes': "📋 *DETALHES DA SUA ENCOMENDA:*\n",
                'pedido': "🆔 *Pedido:*",
                'tipo': "🎨 *Tipo:*",
                'tamanho': "📏 *Tamanho:*",
                'data': "📅 *Data:*",
                'valor': "💰 *Valor Pago:*",
                'trabalho': "🛠️ *A nossa equipa já começou a trabalhar no seu cartoon!*\n\n",
                'outra_encomenda': "*Deseja fazer outra encomenda?* 👇",
                'botao': "🔄 FAZER NOVA ENCOMENDA"
            },
            'ingles': {
                'titulo': "🎉 *PAYMENT CONFIRMED!* 🎉\n\n",
                'parabens': "✨ *Congratulations on your order!* ✨\n\n",
                'detalhes': "📋 *YOUR ORDER DETAILS:*\n",
                'pedido': "🆔 *Order:*",
                'tipo': "🎨 *Type:*",
                'tamanho': "📏 *Size:*",
                'data': "📅 *Date:*",
                'valor': "💰 *Amount Paid:*",
                'trabalho': "🛠️ *Our team has already started working on your cartoon!*\n\n",
                'outra_encomenda': "*Would you like to place another order?* 👇",
                'botao': "🔄 PLACE NEW ORDER"
            },
            'espanhol': {
                'titulo': "🎉 *¡PAGO CONFIRMADO!* 🎉\n\n",
                'parabens': "✨ *¡Felicitaciones por tu pedido!* ✨\n\n",
                'detalhes': "📋 *DETALLES DE TU PEDIDO:*\n",
                'pedido': "🆔 *Pedido:*",
                'tipo': "🎨 *Tipo:*",
                'tamanho': "📏 *Tamaño:*",
                'data': "📅 *Fecha:*",
                'valor': "💰 *Valor Pagado:*",
                'trabalho': "🛠️ *¡Nuestro equipo ya empezó a trabajar en tu caricatura!*\n\n",
                'outra_encomenda': "*¿Deseas hacer otro pedido?* 👇",
                'botao': "🔄 HACER NUEVO PEDIDO"
            },
            'italiano': {
                'titulo': "🎉 *PAGAMENTO CONFERMATO!* 🎉\n\n",
                'parabens': "✨ *Congratulazioni per il tuo ordine!* ✨\n\n",
                'detalhes': "📋 *DETTAGLI DEL TUO ORDINE:*\n",
                'pedido': "🆔 *Ordine:*",
                'tipo': "🎨 *Tipo:*",
                'tamanho': "📏 *Dimensione:*",
                'data': "📅 *Data:*",
                'valor': "💰 *Importo Pagato:*",
                'trabalho': "🛠️ *Il nostro team ha già iniziato a lavorare sul tuo cartoon!*\n\n",
                'outra_encomenda': "*Vuoi fare un altro ordine?* 👇",
                'botao': "🔄 FARE NUOVO ORDINE"
            },
            'alemao': {
                'titulo': "🎉 *ZAHLUNG BESTÄTIGT!* 🎉\n\n",
                'parabens': "✨ *Herzlichen Glückwunsch zu Ihrer Bestellung!* ✨\n\n",
                'detalhes': "📋 *IHRE BESTELLDETAILS:*\n",
                'pedido': "🆔 *Bestellung:*",
                'tipo': "🎨 *Typ:*",
                'tamanho': "📏 *Größe:*",
                'data': "📅 *Datum:*",
                'valor': "💰 *Bezahlter Betrag:*",
                'trabalho': "🛠️ *Unser Team hat bereits mit der Arbeit an Ihrem Cartoon begonnen!*\n\n",
                'outra_encomenda': "*Möchten Sie eine weitere Bestellung aufgeben?* 👇",
                'botao': "🔄 NEUE BESTELLUNG AUFGEBEN"
            },
            'frances': {
                'titulo': "🎉 *PAIEMENT CONFIRMÉ !* 🎉\n\n",
                'parabens': "✨ *Félicitations pour votre commande !* ✨\n\n",
                'detalhes': "📋 *DÉTAILS DE VOTRE COMMANDE:*\n",
                'pedido': "🆔 *Commande:*",
                'tipo': "🎨 *Type:*",
                'tamanho': "📏 *Taille:*",
                'data': "📅 *Date:*",
                'valor': "💰 *Montant Payé:*",
                'trabalho': "🛠️ *Notre équipe a déjà commencé à travailler sur votre dessin animé !*\n\n",
                'outra_encomenda': "*Souhaitez-vous passer une autre commande ?* 👇",
                'botao': "🔄 PASSER NOUVELLE COMMANDE"
            }
        }
        
        textos_msg = textos_principal.get(idioma, textos_principal['portugues'])
        
        # 🔥 MENSAGEM COMPLETA COM TODOS OS DETALHES (AGORA SEGUNDA)
        mensagem_agradecimento = (
            f"{textos_msg['titulo']}"
            f"{textos_msg['parabens']}"
            
            f"{textos_msg['detalhes']}"
            f"{textos_msg['pedido']} #{pedido_pago['id']}\n"
            f"{textos_msg['tipo']} {pedido_pago['tipo_cartoon']}\n"
            f"{detalhes_especificos}"
            f"{textos_msg['tamanho']} {pedido_pago.get('tamanho_cartoon', textos['padrao'])}\n"
            f"{textos_msg['data']} {pedido_pago.get('data_pagamento', 'Hoje')}\n"
            f"{textos_msg['valor']} {texto_valor_cliente}\n\n"
            
            f"{textos_msg['trabalho']}"
            
            f"{textos_msg['outra_encomenda']}"
        )
        
        keyboard = [[InlineKeyboardButton(textos_msg['botao'], callback_data="iniciar_novaencomenda")]]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        # ENVIAR MENSAGEM COMPLETA (DEPOIS DOS CONFETES)
        await bot.send_message(
            chat_id=chat_id,
            text=mensagem_agradecimento,
            parse_mode="Markdown",
            reply_markup=reply_markup
        )
            
        print("✅ Mensagem de agradecimento com detalhes enviada")
        
    except Exception as e:
        print(f"❌ Erro ao enviar mensagem de agradecimento: {e}")





async def enviar_notificacoes_pagamento(chat_id, pedido_pago, amount):
    """Envia APENAS notificação para o telemóvel/admin"""
    try:
        print(f"🎯 Enviando notificação para telemóvel - {pedido_pago['nome']}")
        
        # 🔥 OBTER INFORMAÇÕES DE MOEDA DO PEDIDO
        moeda_original = pedido_pago.get('moeda_original', 'EUR')
        simbolo_original = get_simbolo_moeda(moeda_original.lower())
        total_original = pedido_pago.get('total_pago_original', pedido_pago.get('total', 0))
        total_eur = pedido_pago.get('total_pago_eur', pedido_pago.get('total', 0))
        
        print(f"💰 NOTIFICAÇÃO TELEMÓVEL - Moeda: {moeda_original}, Original: {simbolo_original}{total_original:.2f}, EUR: €{total_eur:.2f}")
        
        # 🔥 USAR BOT ASSÍNCRONO CORRETAMENTE
        bot = application.bot

        CANAL_NOTIFICACOES = os.getenv("CANAL_NOTIFICACOES")
        
        if not CANAL_NOTIFICACOES:
            print("⚠️ AVISO: CANAL_NOTIFICACOES não configurado")
            return
        
        try:
            CANAL_NOTIFICACOES = int(CANAL_NOTIFICACOES)
        except ValueError:
            print("⚠️ AVISO: CANAL_NOTIFICACOES inválido")
            return
        # NOTIFICAÇÃO PARA O ADMIN (TELEMÓVEL)
        try:
            texto_valor_admin = f"€{total_eur:.2f}"
            
            await bot.send_message(
                chat_id=CANAL_NOTIFICACOES,
                text=(
                    f"*NOVA ENCOMENDA DE {texto_valor_admin}!*\n\n"
                    f"🆔 *Pedido:* #{pedido_pago['id']}\n"
                    f"👤 *Cliente:* {pedido_pago['nome']}\n"
                    f"📧 *Email:* {pedido_pago['email']}\n"
                    f"📱 *Contacto:* {pedido_pago.get('contacto', 'N/A')}\n"
                    f"🌍 *País:* {pedido_pago.get('pais', 'N/A')}\n"
                    f"💵 *Valor:* €{total_eur:.2f}" + (f" ({simbolo_original}{total_original:.2f})" if moeda_original != 'EUR' else "") + "\n"
                    f"💰 *Moeda:* {moeda_original}"
                ),
                parse_mode="Markdown"
            )
            print("📱 Notificação admin enviada para telemóvel")
        except Exception as admin_error:
            print(f"⚠️ Erro ao enviar notificação admin: {admin_error}")
        
        print(f"✅ Notificação telemóvel enviada para {pedido_pago['nome']}")
        
    except Exception as e:
        print(f"❌ Erro na notificação telemóvel: {e}")









# ======================= HANDLERS ADICIONAIS =======================



async def pagar_stripe(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """VERSÃO COM INSTRUÇÕES CLARAS PARA BROWSERS - COM TRADUÇÃO"""
    query = update.callback_query
    await query.answer()
    
    # 🔥 PEGAR IDIOMA DO USER_DATA
    idioma = context.user_data.get('idioma', 'portugues')
    
    print(f"🎯 PAGAR_STRIPE - INSTRUÇÕES PARA BROWSERS | Idioma: {idioma}")
    print("🎯 PAGAR_STRIPE - INICIANDO")
    print(f"🔍 Context user_data NO PAGAR_STRIPE: {context.user_data}")
    print(f"🎯 OFERTA_TIPO NO PAGAR_STRIPE: {context.user_data.get('oferta_tipo', 'NÃO DEFINIDO')}")
    
    try:
        # 🔥 PASSO 1: BUSCAR PEDIDO DO REGISTRO
        pedido_id = context.user_data.get("pedido_id")
        print(f"🔍 Procurando pedido: {pedido_id}")
        
        if not pedido_id or pedido_id not in PEDIDOS_REGISTO:
            print("❌ Pedido não encontrado no registro")
            
            # 🔥 MENSAGEM DE ERRO TRADUZIDA
            textos_erro_pedido = {
                'portugues': "❌ Pedido não encontrado. Por favor, inicie um novo pedido. /start",
                'ingles': "❌ Order not found. Please start a new order. /start",
                'espanhol': "❌ Pedido no encontrado. Por favor, inicie un nuevo pedido. /start",
                'italiano': "❌ Ordine non trovato. Per favore, inizi un nuovo ordine. /start",
                'alemao': "❌ Bestellung nicht gefunden. Bitte beginnen Sie eine neue Bestellung. /start",
                'frances': "❌ Commande introuvable. Veuillez démarrer una nueva orden. /start"
            }
            
            await query.edit_message_text(textos_erro_pedido.get(idioma, textos_erro_pedido['portugues']))
            return
        
        pedido = PEDIDOS_REGISTO[pedido_id]
        pais_cliente = pedido["pais"]
        email = pedido["email"]
        tipo = pedido["tipo_cartoon"]
        nome_cliente = pedido["nome"]

        chat_id = query.message.chat_id

        print(f"✅ Pedido encontrado: #{pedido_id}")
        print(f"🔍 Chat ID do cliente: {chat_id}")
        
        # 🔥 CORREÇÃO CRÍTICA: DETERMINAR OFERTA_TIPO CORRETAMENTE
        # 1. PRIMEIRO: Tentar do context.user_data (mais recente)
        oferta_tipo_context = context.user_data.get('oferta_tipo')
        # 2. SEGUNDO: Tentar do pedido no registro
        oferta_tipo_pedido = pedido.get('oferta_tipo')
        # 3. TERCEIRO: Fallback inteligente baseado no tipo de produto
        if 'porta-chaves' in tipo.lower() or 'portachaves' in tipo.lower():
            oferta_tipo_fallback = 'oferta_surpresa'
        else:
            oferta_tipo_fallback = 'pagamento_direto'
        
        # 🔥 DECISÃO FINAL: Prioridade context > pedido > fallback
        oferta_tipo_final = oferta_tipo_context or oferta_tipo_pedido or oferta_tipo_fallback
        
        print(f"🎯 OFERTA_TIPO DETECTADO:")
        print(f"   • Context: {oferta_tipo_context}")
        print(f"   • Pedido: {oferta_tipo_pedido}")
        print(f"   • Fallback: {oferta_tipo_fallback}")
        print(f"   • FINAL: {oferta_tipo_final}")
        
        # 🔥 PASSO 2: CANCELAR O TIMER
        if "timer_task" in pedido:
            try:
                pedido["timer_task"].cancel()
                del pedido["timer_task"]
                print(f"⏹️ TIMER CANCELADO - Usuário clicou em pagar para pedido #{pedido_id}")
            except Exception as e:
                print(f"⚠️ Erro ao cancelar timer: {e}")

        # 🔥 PASSO 3: CALCULAR TOTAL NA MOEDA DO PAÍS
        totais = calcular_total_por_moeda(context, pais_cliente)
        total_na_moeda = totais['total']
        currency = totais['moeda'].lower()
        simbolo = totais['simbolo_moeda']
        
        print(f"💰 Total na moeda local: {simbolo}{total_na_moeda:.2f} {currency.upper()}")

        # 🔥 PASSO 4: DEFINIR MÉTODOS DE PAGAMENTO POR PAÍS
        def get_payment_methods(pais):
            """Retorna métodos de pagamento baseado no país"""
            
            def get_country_code(pais_nome):
                mapeamento_paises = {
                    "portugal": "PT",
                    "espanha": "ES", 
                    "franca": "FR",
                    "alemanha": "DE",
                    "belgica": "BE",
                    "reino unido": "GB",
                    "estados unidos": "US",
                    "paises baixos": "NL",
                    "brasil": "BR",
                    "irlanda": "IE",
                    "italia": "IT",
                    "luxemburgo": "LU",
                    "canada": "CA"
                }
                return mapeamento_paises.get(pais_nome.lower(), pais_nome.upper())
            
            country_code = get_country_code(pais)
            print(f"🔍 País recebido: '{pais}' → Código: '{country_code}'")
            
            payment_methods_by_country = {
                "PT": ["card", "paypal", "link", "klarna", "mb_way", "sepa_debit"],
                "ES": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "FR": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "DE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "BE": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "GB": ["card", "paypal", "link", "klarna"],
                "US": ["card", "paypal", "link"],
                "NL": ["card", "paypal", "link", "klarna", "ideal", "sepa_debit"],
                "BR": ["card", "link"],
                "IE": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "IT": ["card", "paypal", "link", "klarna", "sepa_debit"],
                "LU": ["card", "paypal", "link", "klarna", "bancontact", "sepa_debit"],
                "CA": ["card", "paypal", "link"]
            }
            
            methods = payment_methods_by_country.get(country_code, ["card", "link"])
            print(f"💳 Métodos de pagamento para {pais} ({country_code}): {methods}")
            return methods

        # 🔥 OBTER MÉTODOS REAIS PARA ESTE PAÍS
        metodos_reais = get_payment_methods(pais_cliente)
        
        # 🔥 CRIAR TEXTO DINÂMICO DOS MÉTODOS COM TRADUÇÃO
        def formatar_metodos(metodos, pais, idioma):
            """Formata os métodos de pagamento para exibição"""
            # 🔥 NOMES DOS MÉTODOS POR IDIOMA
            nomes_metodos_por_idioma = {
                'portugues': {
                    "card": "Cartão",
                    "paypal": "PayPal", 
                    "link": "Link (inclui Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'ingles': {
                    "card": "Card",
                    "paypal": "PayPal", 
                    "link": "Link (includes Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA Debit",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'espanhol': {
                    "card": "Tarjeta",
                    "paypal": "PayPal", 
                    "link": "Link (incluye Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Débito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'italiano': {
                    "card": "Carta",
                    "paypal": "PayPal", 
                    "link": "Link (include Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Addebito SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'alemao': {
                    "card": "Karte",
                    "paypal": "PayPal", 
                    "link": "Link (inkl. Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "SEPA-Lastschrift",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                },
                'frances': {
                    "card": "Carte",
                    "paypal": "PayPal", 
                    "link": "Link (inclut Apple Pay/Google Pay)",
                    "klarna": "Klarna",
                    "sepa_debit": "Prélèvement SEPA",
                    "mb_way": "MB Way",
                    "bancontact": "Bancontact",
                    "ideal": "iDEAL"
                }
            }
            
            nomes_metodos = nomes_metodos_por_idioma.get(idioma, nomes_metodos_por_idioma['portugues'])
            textos = []
            
            for metodo in metodos:
                if metodo in nomes_metodos:
                    textos.append(nomes_metodos[metodo])
                else:
                    textos.append(metodo.capitalize())
            
            return ", ".join(textos)

        # 🔥 TEXTOS TRADUZIDOS PARA O CHECKOUT (shipping_message e submit_message)
        textos_checkout_messages = {
            'portugues': {
                "shipping_message": "📦 Enviaremos o seu cartoon personalizado para este endereço!",
                "submit_message": "✨ Obrigado! Vamos criar um cartoon incrível para si!"
            },
            'ingles': {
                "shipping_message": "📦 We'll send your personalized cartoon to this address!",
                "submit_message": "✨ Thank you! We'll create an amazing cartoon for you!"
            },
            'espanhol': {
                "shipping_message": "📦 ¡Enviaremos tu cartoon personalizado a esta dirección!",
                "submit_message": "✨ ¡Gracias! ¡Crearemos un cartoon increíble para ti!"
            },
            'italiano': {
                "shipping_message": "📦 Spediremo il tuo cartoon personalizzato a questo indirizzo!",
                "submit_message": "✨ Grazie! Creeremo un cartoon incredibile per te!"
            },
            'alemao': {
                "shipping_message": "📦 Wir senden Ihren personalisierten Cartoon an diese Adresse!",
                "submit_message": "✨ Danke! Wir erstellen einen fantastischen Cartoon für Sie!"
            },
            'frances': {
                "shipping_message": "📦 Nous enverrons votre dessin animé personnalisé à cette adresse !",
                "submit_message": "✨ Merci ! Nous créerons un dessin animé incroyable pour vous !"
            }
        }
        
        textos_messages = textos_checkout_messages.get(idioma, textos_checkout_messages['portugues'])
        
        # 🔥 DESCRIÇÕES DO PRODUTO POR IDIOMA
        descricoes_produto = {
            'portugues': f"Pedido #{pedido_id} - Para {nome_cliente}",
            'ingles': f"Order #{pedido_id} - For {nome_cliente}",
            'espanhol': f"Pedido #{pedido_id} - Para {nome_cliente}",
            'italiano': f"Ordine #{pedido_id} - Per {nome_cliente}",
            'alemao': f"Bestellung #{pedido_id} - Für {nome_cliente}",
            'frances': f"Commande #{pedido_id} - Pour {nome_cliente}"
        }
        
        descricao_produto = descricoes_produto.get(idioma, descricoes_produto['portugues'])
        
        # 🔥 NOMES DO PRODUTO POR IDIOMA
        nomes_produto = {
            'portugues': f"Cartoon Personalizado - {tipo}",
            'ingles': f"Personalized Cartoon - {tipo}",
            'espanhol': f"Cartoon Personalizado - {tipo}",
            'italiano': f"Cartoon Personalizzato - {tipo}",
            'alemao': f"Personalisierter Cartoon - {tipo}",
            'frances': f"Dessin Animé Personnalisé - {tipo}"
        }
        
        nome_produto = nomes_produto.get(idioma, nomes_produto['portugues'])
        
        # 🔥 CONFIGURAÇÃO PRINCIPAL - CHECKOUT SESSION PADRÃO (FUNCIONAL)
        print("🔗 Criando Checkout Session padrão...")
        
        # 🔥 CORREÇÃO CRÍTICA: USAR VARIÁVEL ÚNICA PARA EVITAR DUPLICAÇÃO
        oferta_tipo_stripe = oferta_tipo_final
        
        session_config = {
            "payment_method_types": metodos_reais,
            "mode": "payment",
            "customer_email": email,
            
            "payment_method_options": {
                "card": {
                    "request_three_d_secure": "automatic"
                }
            },
            
            "shipping_address_collection": {
                "allowed_countries": [
                    "PT", "ES", "FR", "DE", "BE", "GB", "US", "NL", "BR", "IE", "IT", "LU", "CA"
                ]
            },
            
            # 🔥 MENSAGENS TRADUZIDAS PARA O CHECKOUT
            "custom_text": {
                "shipping_address": {
                    "message": textos_messages["shipping_message"]
                },
                "submit": {
                    "message": textos_messages["submit_message"]
                }
            },
            
            "line_items": [{
                "price_data": {
                    "currency": currency,
                    "product_data": {
                        "name": nome_produto,  
                        "description": descricao_produto,  
                    },
                    "unit_amount": int(total_na_moeda * 100),
                },
                "quantity": 1
            }],
            
            "success_url": f"https://t.me/plan3d_bot?start=payment_success_{pedido_id}",
            "cancel_url": f"https://t.me/plan3d_bot?start=payment_cancelled_{pedido_id}",
            
            "metadata": {
                "pedido_id": pedido_id,
                "chat_id": str(chat_id),
                "pais": pais_cliente,
                "moeda": currency,
                "total_pago": str(total_na_moeda),
                "nome_cliente": nome_cliente,
                "tipo_cartoon": tipo,
                "tipo_sessao": oferta_tipo_stripe,  # 🔥 USAR VARIÁVEL ÚNICA
                "oferta_tipo": oferta_tipo_stripe,   # 🔥 USAR VARIÁVEL ÚNICA
                "idioma": idioma  # 🔥 ADICIONAR IDIOMA AO METADATA
            },
            
            "expires_at": int((datetime.now() + timedelta(minutes=30)).timestamp()),
        }

        # 🔥 CONFIGURAÇÃO ESPECÍFICA PARA WALLETS
        paises_com_wallets = ["Reino Unido", "Estados Unidos", "Brasil", "Irlanda", 
                            "França", "Alemanha", "Itália", "Espanha", "Portugal", 
                            "Países Baixos", "Bélgica", "Luxemburgo", "Canadá"]
        
        if pais_cliente in paises_com_wallets and "link" in metodos_reais:
            print(f"📱 Configurando Apple Pay/Google Pay para {pais_cliente}")
            session_config["payment_method_options"]["link"] = {"persistent_token": None}

        # 🔥 CRIAR A SESSÃO
        session = stripe.checkout.Session.create(**session_config)

        print(f"✅ CHECKOUT SESSION CRIADA: {session.id}")
        print(f"🔗 URL do Checkout: {session.url}")

        # 🔥 PASSO 7: ATUALIZAR PEDIDO
        pedido["moeda"] = currency
        pedido["total_na_moeda"] = total_na_moeda
        pedido["simbolo_moeda"] = simbolo
        pedido["session_id"] = session.id
        pedido["payment_intent_id"] = session.payment_intent
        pedido["idioma"] = idioma  # 🔥 GUARDAR IDIOMA NO PEDIDO
        
        print(f"📊 Pedido atualizado")

        # 🔥 PASSO 8: MENSAGEM COM INSTRUÇÕES CLARAS - COM TRADUÇÃO
        texto_metodos = formatar_metodos(metodos_reais, pais_cliente, idioma)
        
        # 🔥 TEXTOS DO PAGAMENTO POR IDIOMA
        textos_pagamento = {
            'portugues': {
                'titulo': "💳 *FINALIZAR PAGAMENTO* 💳",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envio:*",
                'total': "💰 *Total:*",
                'pedido': "🆔 *Pedido:*",
                'metodos': "💳 *Métodos disponíveis:*",
                'seguro': "🔒 *Pagamento 100% seguro via Stripe*",
                'valido': "⏰ *Válido por 30 minutos*",
                'instrucao': "Clique abaixo e siga os passos: 👇",
                'botao': "💳 PAGAR AGORA"
            },
            'ingles': {
                'titulo': "💳 *COMPLETE PAYMENT* 💳",
                'cliente': "👤 *Customer:*",
                'pais': "🌍 *Shipping Country:*",
                'total': "💰 *Total:*",
                'pedido': "🆔 *Order:*",
                'metodos': "💳 *Available methods:*",
                'seguro': "🔒 *100% secure payment via Stripe*",
                'valido': "⏰ *Valid for 30 minutes*",
                'instrucao': "Click below and follow the steps: 👇",
                'botao': "💳 PAY NOW"
            },
            'espanhol': {
                'titulo': "💳 *FINALIZAR PAGO* 💳",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *País de Envío:*",
                'total': "💰 *Total:*",
                'pedido': "🆔 *Pedido:*",
                'metodos': "💳 *Métodos disponibles:*",
                'seguro': "🔒 *Pago 100% seguro vía Stripe*",
                'valido': "⏰ *Válido por 30 minutos*",
                'instrucao': "Haz clic abajo y sigue los pasos: 👇",
                'botao': "💳 PAGAR AHORA"
            },
            'italiano': {
                'titulo': "💳 *COMPLETA PAGAMENTO* 💳",
                'cliente': "👤 *Cliente:*",
                'pais': "🌍 *Paese di Spedizione:*",
                'total': "💰 *Totale:*",
                'pedido': "🆔 *Ordine:*",
                'metodos': "💳 *Metodi disponibili:*",
                'seguro': "🔒 *Pagamento 100% sicuro tramite Stripe*",
                'valido': "⏰ *Valido per 30 minuti*",
                'instrucao': "Clicca qui sotto e segui i passaggi: 👇",
                'botao': "💳 PAGA ORA"
            },
            'alemao': {
                'titulo': "💳 *ZAHLUNG ABSCHLIESSEN* 💳",
                'cliente': "👤 *Kunde:*",
                'pais': "🌍 *Versandland:*",
                'total': "💰 *Gesamt:*",
                'pedido': "🆔 *Bestellung:*",
                'metodos': "💳 *Verfügbare Methoden:*",
                'seguro': "🔒 *100% sichere Zahlung über Stripe*",
                'valido': "⏰ *30 Minuten gültig*",
                'instrucao': "Klicken Sie unten und folgen Sie den Schritten: 👇",
                'botao': "💳 JETZT BEZAHLEN"
            },
            'frances': {
                'titulo': "💳 *FINALISER LE PAIEMENT* 💳",
                'cliente': "👤 *Client:*",
                'pais': "🌍 *Pays de Livraison:*",
                'total': "💰 *Total:*",
                'pedido': "🆔 *Commande:*",
                'metodos': "💳 *Méthodes disponibles:*",
                'seguro': "🔒 *Paiement 100% sécurisé via Stripe*",
                'valido': "⏰ *Valable 30 minutes*",
                'instrucao': "Cliquez ci-dessous et suivez les étapes : 👇",
                'botao': "💳 PAYER MAINTENANT"
            }
        }
        
        textos = textos_pagamento.get(idioma, textos_pagamento['portugues'])

        await query.edit_message_text(
            text=(
                f"{textos['titulo']}\n\n"
                f"{textos['cliente']} {nome_cliente}\n"
                f"{textos['pais']} {pais_cliente}\n"
                f"{textos['total']} {simbolo}{total_na_moeda:.2f} {currency.upper()}\n"
                f"{textos['pedido']} #{pedido_id}\n\n"
                f"{textos['metodos']} {texto_metodos}\n\n"
                f"{textos['seguro']}\n"
                f"{textos['valido']}\n\n"
                f"{textos['instrucao']}"
            ),
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                # 🔥 APENAS 1 BOTÃO - COMO ORIGINAL
                [InlineKeyboardButton(textos['botao'], url=session.url)]
            ])
        )
        
        print(f"✅ Instruções claras para browsers adicionadas | Idioma: {idioma}")

    except Exception as e:
        print(f"❌ ERRO STRIPE: {str(e)}")
        print(f"🔍 Tipo do erro: {type(e)}")
        
        import traceback
        print(f"🔍 Traceback completo: {traceback.format_exc()}")
        
        # 🔥 MENSAGEM DE ERRO TRADUZIDA
        textos_erro = {
            'portugues': "❌ Erro no processamento do pagamento. Por favor, tente novamente em alguns segundos.",
            'ingles': "❌ Error processing payment. Please try again in a few seconds.",
            'espanhol': "❌ Error en el procesamiento del pago. Por favor, intente de nuevo en unos segundos.",
            'italiano': "❌ Errore nell'elaborazione del pagamento. Per favore, riprova tra qualche secondo.",
            'alemao': "❌ Fehler bei der Zahlungsverarbeitung. Bitte versuchen Sie es in einigen Sekunden erneut.",
            'frances': "❌ Erreur lors du traitement du paiement. Veuillez réessayer dans quelques secondes."
        }
        
        await query.edit_message_text(
            textos_erro.get(idioma, textos_erro['portugues'])
        )



        

app = Flask(__name__)

# 🔥 ADICIONE ESTAS ROTAS:
@app.route("/")
def home():
    print("✅ ROTA / ACESSADA!")
    return "✅ GodsPlan Bot Online 24/7! 🚀", 200

@app.route("/test")
def test():
    print("✅ ROTA /test ACESSADA!")
    return "✅ Test route working! 🎉", 200


# ======================= WEBHOOK STRIPE =======================
@app.route("/stripe_webhook", methods=["POST", "GET"])
def stripe_webhook():
    print(f"🎯 WEBHOOK ACESSADO! Método: {request.method}")
    
    if request.method == "GET":
        return "✅ Webhook route working! 🎯", 200
    
    # SEU CÓDIGO ORIGINAL DO WEBHOOK AQUI
    print("📦 Payload recebido do Stripe!")
    print(f"🔍 Headers: {dict(request.headers)}")

     # 🔥 🔥 🔥 ADICIONE ESTES 2 PRINTS AQUI 🔥 🔥 🔥
    print(f"🔍 Secret sendo usado: {STRIPE_WEBHOOK_SECRET}")
    print(f"🔍 Secret length: {len(STRIPE_WEBHOOK_SECRET)}")
    # 🔥 🔥 🔥 FIM DO DEBUG 🔥 🔥 🔥

    print(f"🔍 Webhook chamado - Secret configurado: {bool(STRIPE_WEBHOOK_SECRET)}")




    payload = request.data
    sig = request.headers.get("Stripe-Signature")
    
    try:
        event = stripe.Webhook.construct_event(payload, sig, STRIPE_WEBHOOK_SECRET)
    except Exception as e:
        print(f"❌ Erro verificação webhook: {e}")
        return "Erro", 400

    print(f"✅ Evento Stripe recebido: {event['type']}")
    print(f"🔍 Evento completo: {event}")
    # 🔥 🔥 🔥 TRATAR AMBOS OS TIPOS DE EVENTO COM ASYNC 🔥 🔥 🔥
    
    if event["type"] == "checkout.session.completed":
        print("🎯 Processando checkout.session.completed")
        return processar_checkout_completed_async(event)
        
    elif event["type"] == "payment_intent.succeeded":
        print("🎯 Processando payment_intent.succeeded")
        return processar_payment_intent_succeeded_async(event)
        
    else:
        print(f"⚠️ Evento não tratado: {event['type']}")
        return "OK", 200


def processar_checkout_completed_async(event):
    """Processa checkout.session.completed de forma assíncrona - COM TIPO_SESSAO CORRIGIDO"""
    try:
        session = event['data']['object']
        metadata = session.get('metadata', {})
        pedido_id = metadata.get('pedido_id')
        chat_id = metadata.get('chat_id')
        amount = session['amount_total'] / 100
        currency = session.get('currency', 'eur').upper()
        
        # 🔥 🔥 🔥 CORREÇÃO CRÍTICA: DEFAULT DEVE SER "pagamento_direto"
        tipo_sessao = metadata.get('tipo_sessao', 'pagamento_direto')  # ✅ CORRIGIDO
        
        print(f"🎉 Pagamento confirmado via Checkout: {amount} {currency} | Tipo: {tipo_sessao}")
        print(f"📋 Metadata: {metadata}")

        # 🔥 DEBUG: VERIFICAR SE O TIPO_SESSAO ESTÁ CORRETO
        print(f"🔍 DEBUG TIPO_SESSAO: '{tipo_sessao}'")
        print(f"🔍 METADATA COMPLETO DO STRIPE: {metadata}")
        print(f"🎯 OFERTA_TIPO NO WEBHOOK: {metadata.get('oferta_tipo', 'NÃO ENCONTRADO')}")
        print(f"🎯 TIPO_SESSAO NO WEBHOOK: {metadata.get('tipo_sessao', 'NÃO ENCONTRADO')}")
        # 🔥 CAPTURAR DADOS DE SHIPPING
        shipping_details = None
        if session.get('collected_information') and session['collected_information'].get('shipping_details'):
            shipping_details = session['collected_information']['shipping_details']
            print(f"🚚 Dados de Shipping (collected_information): {shipping_details}")
        elif session.get('shipping_details'):
            shipping_details = session['shipping_details']
            print(f"🚚 Dados de Shipping (shipping_details): {shipping_details}")
        
        if pedido_id and chat_id:
            # 🔥 ATUALIZAR ESTATÍSTICAS AQUI - SÓ QUANDO PAGO
            atualizar_estatistica("total_pedidos")
            print(f"📊 ESTATÍSTICAS: Pedido #{pedido_id} AGORA contado (status: pago)")
            
            # 🔥 EXECUTAR FUNÇÃO ASSÍNCRONA PASSANDO O TIPO_SESSAO
            import asyncio
            
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
            
            # 🔥 AGORA PASSAMOS O TIPO_SESSAO CORRETAMENTE
            if loop.is_running():
                asyncio.create_task(processar_pagamento_sucesso(
                    pedido_id=pedido_id,
                    chat_id=chat_id,
                    amount=amount,
                    shipping_details=shipping_details,
                    moeda_original=currency,
                    tipo_sessao=tipo_sessao  # 🔥 AGORA CORRETO
                ))
                print(f"🔄 Task criada para: {pedido_id} - Tipo: {tipo_sessao}")
            else:
                loop.run_until_complete(processar_pagamento_sucesso(
                    pedido_id=pedido_id,
                    chat_id=chat_id,
                    amount=amount,
                    shipping_details=shipping_details,
                    moeda_original=currency,
                    tipo_sessao=tipo_sessao  
                ))
                print(f"✅ Processado sincronamente: {pedido_id} - Tipo: {tipo_sessao}")
        else:
            print("❌ Metadata incompleto no webhook")
            
    except Exception as e:
        print(f"❌ Erro ao processar checkout: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
    
    return "OK", 200





def processar_payment_intent_succeeded_async(event):
    """Processa payment_intent.succeeded de forma assíncrona - COM SHIPPING E MOEDA"""
    try:
        payment_intent = event['data']['object']
        amount = payment_intent['amount'] / 100
        currency = payment_intent.get('currency', 'eur').upper()  # 🔥 CAPTURAR A MOEDA
        charge_id = payment_intent.get('latest_charge')
        
        print(f"🎉 Pagamento confirmado via Payment Intent: {amount} {currency}")
        print(f"⚡ Charge ID: {charge_id}")
        
        # 🔥 TENTAR OBTER METADATA E SHIPPING DO PAYMENT_INTENT
        shipping_details = None
        
        if payment_intent.get('shipping'):
            shipping_details = payment_intent['shipping']
            print(f"🚚 Dados de Shipping (payment_intent): {shipping_details}")
        
        # 🔥 PRECISAMOS DO PEDIDO_ID E CHAT_ID - tentar da session
        if payment_intent.get('metadata') and payment_intent['metadata'].get('pedido_id'):
            pedido_id = payment_intent['metadata']['pedido_id']
            chat_id = payment_intent['metadata']['chat_id']
            
            import asyncio
            
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
            
            if loop.is_running():
                asyncio.create_task(processar_pagamento_sucesso(pedido_id, chat_id, amount, shipping_details, currency))
                print(f"🔄 Task assíncrona criada a partir do payment_intent COM MOEDA {currency}")
            else:
                loop.run_until_complete(processar_pagamento_sucesso(pedido_id, chat_id, amount, shipping_details, currency))
                print(f"✅ Pagamento processado a partir do payment_intent COM MOEDA {currency}")
        else:
            print("⚠️ Payment Intent sem metadata - dependendo do checkout.session.completed")
            
    except Exception as e:
        print(f"❌ Erro ao processar payment intent: {e}")
    
    return "OK", 200








# ======================= COMANDO ESTATÍSTICAS =======================



# --- Função principal ---
def main():
    """Função principal corrigida para suportar múltiplos clientes"""
    TOKEN = TELEGRAM_TOKEN
    
    # 🔥 CONFIGURAR LOOP DE EVENTOS CORRETAMENTE
    try:
        # Para Windows, usar política correta (com tratamento para deprecation)
      #   try:
      #      if sys.platform == 'win32':
      #           asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
      #   except AttributeError:
            # Para versões mais recentes do Python
      #       pass
        
        # 🔥 PRIMEIRO CRIAR A APPLICATION
        application = (
            ApplicationBuilder()
            .token(TOKEN)
            .concurrent_updates(True)  # 🔥 CRÍTICO: Permitir múltiplos clientes
            .pool_timeout(60)
            .connect_timeout(60)
            .read_timeout(60)
            .write_timeout(60)
            .get_updates_read_timeout(60)
            .build()
        )
        
        print("🚀 INICIANDO BOT COM SUPORTE A MÚLTIPLOS CLIENTES...")
        print("🧪 Iniciando testes dos canais...")
        
        # 🔥 AGORA CRIAR O LOOP E EXECUTAR A MENSAGEM AUTOMÁTICA
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        # 🔥 EXECUTAR MENSAGEM AUTOMÁTICA DENTRO DO LOOP
        try:
            loop.run_until_complete(enviar_mensagem_automatica(application))
            print("✅ Mensagem automática enviada com sucesso!")
        except Exception as e:
            print(f"⚠️ Erro ao enviar mensagem automática: {e}")
        
        # 🔥 AGORA ADICIONAR SEUS HANDLERS (MANTENDO SUA ESTRUTURA)
        
        # 1. HANDLERS DE COMANDOS
        application.add_handler(CommandHandler("start", start))
        application.add_handler(CommandHandler("skip", pular_frase_handler))   
        application.add_handler(CommandHandler("help", help_handler))
        application.add_handler(CommandHandler("admin", admin_command))
        application.add_handler(CommandHandler("gift", gift_command))
       
        
        # 2. HANDLERS DE CALLBACK (MAIS ESPECÍFICOS PRIMEIRO)
        application.add_handler(CallbackQueryHandler(selecionar_idioma, pattern="^idioma_"))
        application.add_handler(CallbackQueryHandler(gift_selecionar_idioma, pattern="^gift_idioma_"))

        # 🔥 HANDLERS DE GIFT (PRIORIDADE)
        application.add_handler(CallbackQueryHandler(recusar_gift_handler, pattern="^recusar_gift_"))
        application.add_handler(CallbackQueryHandler(pagar_gift_handler, pattern="^pagar_gift_"))
        
        # 🔥 HANDLERS DE OFERTAS
        application.add_handler(CallbackQueryHandler(recusar_oferta, pattern="^recusar_oferta_"))
        application.add_handler(CallbackQueryHandler(pagar_original, pattern="^pagar_original_"))
        application.add_handler(CallbackQueryHandler(pagar_tamanho45, pattern="^pagar_tamanho45_"))
        application.add_handler(CallbackQueryHandler(pagar_portachaves, pattern="^pagar_portachaves_"))
        
        # 🔥 HANDLERS DE SAÍDA/CONFIRMAÇÃO
        application.add_handler(CallbackQueryHandler(sair_oferta, pattern="^sair_oferta_"))
        application.add_handler(CallbackQueryHandler(confirmar_saida, pattern="^confirmar_saida_"))
        application.add_handler(CallbackQueryHandler(sair_poferta45, pattern="^sair_poferta45_"))
        application.add_handler(CallbackQueryHandler(confirmar_saida45, pattern="^confirmar_saida45_"))
        application.add_handler(CallbackQueryHandler(sair_diretoportachaves, pattern="^sair_diretoportachaves_"))
        application.add_handler(CallbackQueryHandler(confirmar_saidadireta, pattern="^confirmar_saidadireta_"))
        
        # 🔥 HANDLERS DE RECUPERAÇÃO
        application.add_handler(CallbackQueryHandler(recuperar_pedido, pattern="^recuperar_pagar_"))
        application.add_handler(CallbackQueryHandler(reportar_problema, pattern="^reportar_problema_"))
        application.add_handler(CallbackQueryHandler(problema_valor, pattern="^problema_valor_"))
        application.add_handler(CallbackQueryHandler(problema_outro, pattern="^problema_outro_"))
        application.add_handler(CallbackQueryHandler(todas_recusadas, pattern="^todas_recusadas_"))
        
        # 🔥 HANDLERS DE OFERTAS ESPECÍFICAS
        application.add_handler(CallbackQueryHandler(aceitar_oferta_especifica, pattern="^oferta_portachaves_"))
        application.add_handler(CallbackQueryHandler(proxima_oferta, pattern="^proxima_oferta_"))
        application.add_handler(CallbackQueryHandler(ultima_oferta, pattern="^ultima_oferta_"))
        
        # 🔥 HANDLERS DE CARTOON
        application.add_handler(CallbackQueryHandler(iniciar_cartoon, pattern="^(mycartoon|iniciar_cartoon)$"))
        application.add_handler(CallbackQueryHandler(cartoon_handler, pattern="^cartoon_"))
        application.add_handler(CallbackQueryHandler(estilo_handler, pattern="^estilo_"))
        application.add_handler(CallbackQueryHandler(tamanho_handler, pattern="^tamanho_"))
        
        # 🔥 HANDLERS DE PAÍS
        application.add_handler(CallbackQueryHandler(selecionar_pais, pattern="^pais_"))
        application.add_handler(CallbackQueryHandler(selecionar_gift_pais, pattern="^gift_pais_"))
        
        # 🔥 HANDLERS DE VOLTAR/NAVEGAÇÃO
        application.add_handler(CallbackQueryHandler(voltar_handler, pattern="^voltar$"))
        application.add_handler(CallbackQueryHandler(voltar_inicio, pattern="^voltar_inicio$"))
        application.add_handler(CallbackQueryHandler(voltar_menu, pattern="^voltar_menu$"))
        application.add_handler(CallbackQueryHandler(help_voltar, pattern="^help_voltar$"))
        
        # 🔥 HANDLERS DE FINALIZAÇÃO
        application.add_handler(CallbackQueryHandler(finalizar_compra, pattern="^finalizar_compra$"))
        application.add_handler(CallbackQueryHandler(finalizar_gift, pattern="^finalizar_gift$"))
        
        # 🔥 HANDLERS DE FOTO
        application.add_handler(CallbackQueryHandler(mudar_foto, pattern="^mudar_foto$"))
        application.add_handler(CallbackQueryHandler(mudar_gift_foto, pattern="^mudar_gift_foto$"))
        
        # 🔥 HANDLERS DE ADMIN
        application.add_handler(CallbackQueryHandler(admin_command, pattern="^admin_page_"))
        application.add_handler(CallbackQueryHandler(admin_command, pattern="^admin_refresh$"))
        application.add_handler(CallbackQueryHandler(admin_back_handler, pattern="^admin_back$"))
        
        # 🔥 HANDLERS DE MENU ADMIN
        application.add_handler(CallbackQueryHandler(menu_export, pattern="^menu_export$"))
        application.add_handler(CallbackQueryHandler(btn_options, pattern="^btn_options$"))
        
        # 🔥 HANDLERS DE EXPORTAÇÃO
        application.add_handler(CallbackQueryHandler(export_csv_handler, pattern="^export_csv$"))
        application.add_handler(CallbackQueryHandler(export_txt_handler, pattern="^export_txt$"))
        application.add_handler(CallbackQueryHandler(export_pdf_handler, pattern="^export_pdf$"))
        application.add_handler(CallbackQueryHandler(export_word_handler, pattern="^export_word$"))
        application.add_handler(CallbackQueryHandler(exportar_csv_completo, pattern="^export_ccsv_"))
        application.add_handler(CallbackQueryHandler(exportar_word_completo, pattern="^export_word_"))
        application.add_handler(CallbackQueryHandler(exportar_pdf_com_foto, pattern="^export_pdf_photo_"))
        application.add_handler(CallbackQueryHandler(exportar_txt_com_foto, pattern="^export_txt_photo_"))
        
        # 🔥 HANDLERS DE REQUESTS
        application.add_handler(CallbackQueryHandler(marcar_como_feito, pattern="^done_"))
        
        # 🔥 HANDLERS DE PAGAMENTO
        application.add_handler(CallbackQueryHandler(pagar_stripe, pattern="^pagar_stripe$"))
        
        # 🔥 HANDLERS DE NOVA ENCOMENDA
        application.add_handler(CallbackQueryHandler(iniciar_novaencomenda, pattern="^iniciar_novaencomenda$"))
        
        # 🔥 HANDLERS DE ANIMAL/PERSONALIZADO
        application.add_handler(CallbackQueryHandler(tipo_animal_handler, pattern="^tipo_"))
        application.add_handler(CallbackQueryHandler(tipo_personalizado_handler, pattern="^personalizado_"))
        
        # 🔥 HANDLERS DE OPÇÕES
        application.add_handler(CallbackQueryHandler(options_save_handler, pattern="^options_save$"))
        application.add_handler(CallbackQueryHandler(options_restore_handler, pattern="^options_restore$"))
        application.add_handler(CallbackQueryHandler(options_delete_handler, pattern="^options_delete$"))
        application.add_handler(CallbackQueryHandler(options_edit_taxes_handler, pattern="^options_edit_taxes$"))
        application.add_handler(CallbackQueryHandler(view_taxes_handler, pattern="^view_taxes$"))
        application.add_handler(CallbackQueryHandler(view_frete_handler, pattern="^view_frete$"))
        application.add_handler(CallbackQueryHandler(edit_tax_country_handler, pattern="^edit_tax_country$"))
        application.add_handler(CallbackQueryHandler(edit_frete_country_handler, pattern="^edit_frete_country$"))
        application.add_handler(CallbackQueryHandler(options_edit_admin_handler, pattern="^options_edit_admin$"))
        
        # 🔥 HANDLERS DE HELP
        application.add_handler(CallbackQueryHandler(help_selecionar_idioma, pattern="^help_idioma_"))
        application.add_handler(CallbackQueryHandler(help_encomenda, pattern="^help_encomenda$"))
        application.add_handler(CallbackQueryHandler(help_problema, pattern="^help_problema$"))
        application.add_handler(CallbackQueryHandler(help_tempo, pattern="^help_tempo$"))
        
        # 3. HANDLERS DE MENSAGENS (ÚLTIMOS - MAIS GENÉRICOS)

        # 🔥 HANDLER DE FOTOS NORMAS (ÚLTIMO - MAIS GENÉRICO)
        application.add_handler(MessageHandler(filters.PHOTO & filters.ChatType.PRIVATE, gift_foto_handler))

# 🔥 3. HANDLER DE TEXTO DO GIFT (ESPECÍFICO)
        application.add_handler(MessageHandler(filters.TEXT & filters.ChatType.PRIVATE, gift_text_handler))

        
        # 🔥 HANDLERS GIFT (ESPECÍFICOS)        
        # 🔥 HANDLERS PRINCIPAIS DE MENSAGEM
        application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
        application.add_handler(MessageHandler(filters.CONTACT, receber_contacto))
        
        
        
        # 🔥 HANDLER DE CALLBACK GENÉRICO (ÚLTIMO)
        application.add_handler(CallbackQueryHandler(handle_callback_query))
        
        # 🔥 VERIFICAR HANDLERS
        print("\n" + "="*50)
        print("✅ HANDLERS REGISTRADOS COM SUCESSO")
        print("="*50)
        
        total_handlers = 0
        for group in sorted(application.handlers.keys()):
            count = len(application.handlers[group])
            total_handlers += count
            print(f"📋 Grupo {group}: {count} handlers")
        
        print(f"\n🎯 TOTAL: {total_handlers} handlers registrados")
        print("🤖 Bot está a funcionar...")
        
        # 🔥 INICIAR POLLING
        loop.run_until_complete(
            application.run_polling(
                poll_interval=0.1,  # 🔥 Resposta rápida para múltiplos clientes
                timeout=30,
                drop_pending_updates=True,
                allowed_updates=Update.ALL_TYPES,
                close_loop=False  # 🔥 IMPORTANTE: Não fechar o loop
            )
        )
        
    except KeyboardInterrupt:
        print("\n🛑 Bot interrompido pelo usuário")
    except Exception as e:
        print(f"\n❌ ERRO CRÍTICO: {e}")
        import traceback
        traceback.print_exc()
    finally:
        print("\n👋 Bot terminado")
if __name__ == "__main__":
    # Iniciar Flask em thread separada (se necessário)
    import threading
    
    # Verificar se a função run_flask existe antes de criar thread
    if 'run_flask' in globals():
        threading.Thread(target=run_flask, daemon=True).start()
    
    # Executar bot principal

    main()        
